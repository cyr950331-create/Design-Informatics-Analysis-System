import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from tkinter import messagebox, PhotoImage, filedialog
import subprocess
import os
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from Source.Toolkit import ToolkitApp
import tkinter as tk

# 获取当前脚本所在的目录
current_dir = os.path.dirname(os.path.abspath(__file__))
project_root = current_dir

# 定义语言字典
LANGUAGES = {
    'zh': {
        'title': "设计信息学分析系统",
        'copy_success': "内容已复制到剪贴板",
        'copyright': "cyr",
        'data_cleaning': "数据清洗",
        'select_file': "选择文件",
        'clean_data': "清洗数据",
        'select_file_first': "请先选择文件",
        'file_selected': "已选择文件: ",
        'clean_success': "数据清洗成功，已保存为: ",
        'invalid_format': "不支持的文件格式",
        'clean_error': "数据清洗失败: ",
        'col_clean_success': "列名清洗完成",
        'word_report': "清洗报告已保存为: ",
        'original_rows': "原始数据行数: ",
        'duplicate_rows': "重复行行数: ",
        'numeric_cols': "数值型列数: ",
        'filled_numeric': "填充缺失值的数值型列数: ",
        'categorical_cols': "分类型列数: ",
        'filled_categorical': "填充缺失值的分类型列数: ",
        'col_rename': "列名修改情况: ",
        'data_issues': "数据存在的问题: ",
        'switch_language': "切换语言"
    },
    'en': {
        'title': "Design Informatics Analysis System",
        'copy_success': "Content has been copied to clipboard",
        'copyright': "cyr",
        'data_cleaning': "Data Cleaning",
        'select_file': "Select File",
        'clean_data': "Clean Data",
        'select_file_first': "Please select a file first",
        'file_selected': "Selected file: ",
        'clean_success': "Data cleaning successful, saved as: ",
        'invalid_format': "Unsupported file format",
        'clean_error': "Data cleaning failed: ",
        'col_clean_success': "Column name cleaning completed",
        'word_report': "Cleaning report saved as: ",
        'original_rows': "Original number of rows: ",
        'duplicate_rows': "Number of duplicate rows: ",
        'numeric_cols': "Number of numeric columns: ",
        'filled_numeric': "Number of numeric columns with missing values filled: ",
        'categorical_cols': "Number of categorical columns: ",
        'filled_categorical': "Number of categorical columns with missing values filled: ",
        'col_rename': "Column name changes: ",
        'data_issues': "Data issues identified: ",
        'switch_language': "Switch Language"
    }
}

# 定义按钮文本的语言字典
BUTTON_TEXTS = {
    'zh': {
        "工具包": "工具包",
        "用户手册": "用户手册",
        "数据清洗": "数据清洗"
    },
    'en': {
        "工具包": "Toolkit",
        "用户手册": "User Manual",
        "数据清洗": "Data Cleaning"
    }
}

# 当前语言
current_language = 'en'

# 图片切换相关变量
current_image_index = 0  # 当前显示的图片索引
image_paths = [
    os.path.join(project_root, "icon", "icon1.gif"),
    os.path.join(project_root, "icon", "icon2.gif"),
    os.path.join(project_root, "icon", "icon3.gif"),
    os.path.join(project_root, "icon", "icon4.gif")
]  # 四张图片的路径列表

# 用于存储每个按钮的最大宽度
button_max_widths = []

# 记录是否已经显示过详情
has_shown_details = False

def run_script(file_path):
    try:
        if file_path == "Toolkit":
            ToolkitApp(ttk.Toplevel(root))
        elif file_path == "UserManual":
            pdf_path = os.path.join(current_dir, "user_manual.pdf")
            if not os.path.exists(pdf_path):
                messagebox.showerror("错误", f"文件不存在: {pdf_path}")
                return
            if os.name == 'nt':
                os.startfile(pdf_path)
            else:
                subprocess.run(['open' if os.name == 'posix' else 'xdg-open', pdf_path])
        elif file_path == "DataCleaning":
            open_data_cleaning_dialog()
        else:
            import sys
            subprocess.Popen([sys.executable, file_path])
    except Exception as e:
        messagebox.showerror("错误", f"运行脚本时出错: {e}")

def clean_column_names(df):
    """清洗列名：去除空格、特殊字符，并记录修改情况"""
    original_columns = df.columns.tolist()
    changes = []

    # 处理每一列
    new_columns = []
    for col in original_columns:
        new_col = col.strip()  # 去除前后空格
        new_col = new_col.replace(' ', '_')  # 替换空格为下划线
        new_col = ''.join([c for c in new_col if c.isalnum() or c == '_'])  # 去除特殊字符
        if new_col != col:
            changes.append((col, new_col))
        new_columns.append(new_col)
    df.columns = new_columns
    return df, changes


def create_cleaning_report(original_path, cleaned_path, stats, issues):
    """创建数据清洗报告Word文档"""
    doc = Document()

    # 设置标题
    title = doc.add_heading(LANGUAGES[current_language]['data_cleaning'] + "报告", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加原始文件信息
    doc.add_heading("1. " + "文件信息", level=1)
    doc.add_paragraph(f"原始文件: {original_path}")
    doc.add_paragraph(f"清洗后文件: {cleaned_path}")

    # 添加数据统计信息
    doc.add_heading("2. " + "数据统计", level=1)
    doc.add_paragraph(f"{LANGUAGES[current_language]['original_rows']}{stats['original_rows']}")
    doc.add_paragraph(f"{LANGUAGES[current_language]['duplicate_rows']}{stats['duplicate_rows']}")
    doc.add_paragraph(f"{LANGUAGES[current_language]['numeric_cols']}{stats['numeric_cols']}")
    doc.add_paragraph(f"{LANGUAGES[current_language]['filled_numeric']}{stats['filled_numeric']}")
    doc.add_paragraph(f"{LANGUAGES[current_language]['categorical_cols']}{stats['categorical_cols']}")
    doc.add_paragraph(f"{LANGUAGES[current_language]['filled_categorical']}{stats['filled_categorical']}")

    # 添加列名修改情况
    doc.add_heading("3. " + "列名修改", level=1)
    if stats['col_changes']:
        for original, new in stats['col_changes']:
            doc.add_paragraph(f"• {original} → {new}")
    else:
        doc.add_paragraph("无列名修改")

    # 添加数据清洗操作说明
    doc.add_heading("4. " + "清洗操作说明", level=1)
    doc.add_paragraph("1. 列名清洗：去除空格、特殊字符，转换为小写，确保列名规范")
    doc.add_paragraph("2. 重复值处理：删除完全重复的行")
    doc.add_paragraph("3. 缺失值处理：")
    doc.add_paragraph("   - 数值型列：使用列均值填充缺失值")
    doc.add_paragraph("   - 分类型列：使用列众数填充缺失值")

    # 添加数据问题说明
    doc.add_heading("5. " + LANGUAGES[current_language]['data_issues'], level=1)
    if issues:
        for issue in issues:
            doc.add_paragraph(f"• {issue}")
    else:
        doc.add_paragraph("未发现明显数据问题")

    # 保存报告
    base, ext = os.path.splitext(original_path)
    report_path = f"{base}_cleaning_report.docx"
    doc.save(report_path)
    return report_path

def open_data_cleaning_dialog():
    """打开数据清洗对话框"""
    dialog = ttk.Toplevel(root)
    dialog.title(LANGUAGES[current_language]['data_cleaning'])
    dialog_width = int(screen_width * 0.4)
    dialog_height = int(screen_height * 0.3)
    dialog.geometry(f"{dialog_width}x{dialog_height}")

    # 先隐藏窗口
    dialog.withdraw()

    # 子窗口同步任务栏图标
    try:
        icon_path = os.path.join(project_root, "icon", "icon.ico")
        if os.path.exists(icon_path):
            dialog.iconbitmap(icon_path)  # 子窗口任务栏图标
            icon_img = PhotoImage(file=icon_path)
            dialog.iconphoto(True, icon_img)
    except Exception as e:
        print(f"子窗口任务栏图标设置失败: {str(e)}")

    # 设置窗口居中
    dialog.update_idletasks()  # 确保窗口尺寸已计算
    width = dialog.winfo_width()
    height = dialog.winfo_height()
    x = (root.winfo_screenwidth() - width) // 2
    y = (root.winfo_screenheight() - height) // 2
    dialog.geometry(f"+{x}+{y}")

    dialog.transient(root)  # 设置为主窗口的子窗口
    dialog.grab_set()  # 模态窗口

    # 存储选中的文件路径
    selected_file = [None]

    # 创建按钮框架
    button_frame = ttk.Frame(dialog)
    button_frame.pack(pady=50)

    # 创建显示选中文件的标签
    file_label = ttk.Label(dialog, text="", wraplength=550)
    file_label.pack(pady=10)

    # 最后显示窗口
    dialog.deiconify()

    def select_file():
        """选择要清洗的数据文件"""
        file_path = filedialog.askopenfilename(
            filetypes=[
                ("All files", "*.*"),
                ("Excel files", "*.xlsx;*.xls"),
                ("CSV files", "*.csv")
            ]
        )
        if file_path:
            selected_file[0] = file_path
            file_label.config(text=f"{LANGUAGES[current_language]['file_selected']}{file_path}")

    def clean_data():
        """执行数据清洗操作"""
        if not selected_file[0]:
            messagebox.showinfo("", LANGUAGES[current_language]['select_file_first'])
            return

        try:
            # 初始化统计信息和问题列表
            stats = {}
            issues = []

            # 根据文件扩展名读取数据
            if selected_file[0].endswith('.csv'):
                df = pd.read_csv(selected_file[0])
            elif selected_file[0].endswith(('.xlsx', '.xls')):
                df = pd.read_excel(selected_file[0])
            else:
                messagebox.showerror("错误", LANGUAGES[current_language]['invalid_format'])
                return

            # 记录原始行数
            stats['original_rows'] = len(df)

            # 清洗列名
            df, col_changes = clean_column_names(df)
            stats['col_changes'] = col_changes

            # 处理重复行
            duplicates_before = df.duplicated().sum()
            stats['duplicate_rows'] = duplicates_before
            if duplicates_before > 0:
                issues.append(f"存在{duplicates_before}行重复数据")
                df = df.drop_duplicates()

            # 处理缺失值
            # 数值型列
            numeric_cols = df.select_dtypes(include=['number']).columns
            stats['numeric_cols'] = len(numeric_cols)

            # 记录有缺失值的数值列
            numeric_with_missing = [col for col in numeric_cols if df[col].isnull().sum() > 0]
            stats['filled_numeric'] = len(numeric_with_missing)

            if numeric_with_missing:
                issues.append(f"数值型列中有{len(numeric_with_missing)}列存在缺失值")
                df[numeric_cols] = df[numeric_cols].fillna(df[numeric_cols].mean())

            # 分类型列
            categorical_cols = df.select_dtypes(include=['object']).columns
            stats['categorical_cols'] = len(categorical_cols)

            # 记录有缺失值的分类型列
            categorical_with_missing = []
            for col in categorical_cols:
                if df[col].isnull().sum() > 0 and not df[col].mode().empty:
                    categorical_with_missing.append(col)

            stats['filled_categorical'] = len(categorical_with_missing)

            if categorical_with_missing:
                issues.append(f"分类型列中有{len(categorical_with_missing)}列存在缺失值")
                for col in categorical_cols:
                    if not df[col].mode().empty:
                        df[col] = df[col].fillna(df[col].mode().iloc[0])

            # 保存清洗后的数据
            base, ext = os.path.splitext(selected_file[0])
            cleaned_path = f"{base}_cleaned{ext}"

            if selected_file[0].endswith('.csv'):
                df.to_csv(cleaned_path, index=False)
            else:
                df.to_excel(cleaned_path, index=False)

            # 生成清洗报告
            report_path = create_cleaning_report(selected_file[0], cleaned_path, stats, issues)

            messagebox.showinfo(
                "成功",
                f"{LANGUAGES[current_language]['clean_success']}{cleaned_path}\n"
                f"{LANGUAGES[current_language]['word_report']}{report_path}"
            )
            dialog.destroy()

        except Exception as e:
            messagebox.showerror("错误", f"{LANGUAGES[current_language]['clean_error']}{str(e)}")

    # 创建"选择文件"按钮
    select_btn = ttk.Button(
        button_frame,
        text=LANGUAGES[current_language]['select_file'],
        bootstyle=PRIMARY,
        command=select_file,
        width=15
    )
    select_btn.pack(side=LEFT, padx=20)

    # 创建"清洗数据"按钮
    clean_btn = ttk.Button(
        button_frame,
        text=LANGUAGES[current_language]['clean_data'],
        bootstyle=SUCCESS,
        command=clean_data,
        width=15
    )
    clean_btn.pack(side=LEFT, padx=20)

    # 添加状态标签
    status_label = ttk.Label(
        dialog,
        text="Remove the spaces before and after the table header (usually the first row), and replace the remaining spaces with underscores;\nRemove all special characters that are not letters, numbers, or underscores from the table header;"
             "\nFill missing values using the mean or mode of the column;\nGenerate a cleaned data file.",
        foreground="gray",
        wraplength=500  # 可选：限制每行最大宽度（像素），超过则自动换行
    )
    status_label.pack(pady=20)

def switch_language():
    global current_language, has_shown_details
    current_language = 'zh' if current_language == 'en' else 'en'
    root.title(LANGUAGES[current_language]['title'])
    copyright_label.config(text=LANGUAGES[current_language]['copyright'])
    #language_label.config(text=LANGUAGES[current_language]['switch_language'])

    # 更新按钮文本
    for index, button in enumerate(button_list):
        original_text = button_texts[index]
        display_text = BUTTON_TEXTS[current_language][original_text]
        button.config(text=display_text, width=button_max_widths[index])

    # 更新所有打开的示例数据窗口的关闭按钮文本
    for window in root.winfo_children():
        if isinstance(window, ttk.Toplevel) and hasattr(window, 'close_btn'):
            window.close_btn.config(text=LANGUAGES[current_language]['close'])

    has_shown_details = False

def show_details(event, text):
    # 移除了详情显示功能，因为详情框已删除
    pass

def hide_details(event):
    # 移除了详情隐藏功能，因为详情框已删除
    pass

# 创建主窗口
root = ttk.Window(themename="flatly")
root.title(LANGUAGES[current_language]['title'])
root.withdraw()  # 隐藏窗口

# 设置窗口图标（同时关联任务栏图标）
try:
    # 定义ICO图标路径（使用icon.ico作为任务栏图标）
    icon_path = os.path.join(project_root, "icon", "icon.ico")
    if os.path.exists(icon_path):
        root.iconbitmap(icon_path)
        icon_img = PhotoImage(file=icon_path)
        root.iconphoto(True, icon_img)
    else:
        print(f"图标文件不存在: {icon_path}，请检查路径是否正确")
except Exception as e:
    print(f"设置任务栏图标失败: {str(e)}")

# 获取屏幕的宽度和高度
screen_width = root.winfo_screenwidth()
screen_height = root.winfo_screenheight()

# 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
window_width = int(screen_width * 0.6)
window_height = int(screen_height * 0.6)

# 限制最小窗口尺寸，避免过小
min_width = 600
min_height = 420
window_width = max(window_width, min_width)
window_height = max(window_height, min_height)

# 计算窗口应该放置的位置
x = (screen_width - window_width) // 2
y = (screen_height - window_height) // 2

# 设置窗口的位置和大小
root.geometry(f"{window_width}x{window_height}+{x}+{y}")
root.minsize(min_width, min_height)  # 添加最小尺寸限制

# 创建一个主框架，用于居中内容
main_frame = ttk.Frame(root)
main_frame.pack(expand=True, fill=BOTH, anchor='n')

# 添加图片（在按钮上方）
image_frame = ttk.Frame(main_frame)
image_frame.pack(pady=(35, 5))  # 增加上下边距

# 初始化图片对象列表（存储四张图片）
image_objects = []
try:
    for path in image_paths:
        img = PhotoImage(file=path)
        scaled_img = img.subsample(5, 5)  # 缩小为原来的1/5（与原代码保持一致）
        image_objects.append(scaled_img)

    # 初始显示第一张图片
    image_label = ttk.Label(image_frame, image=image_objects[0])
    image_label.image = image_objects[0]
    image_label.pack()

    # 定义图片切换函数
    def switch_image(event):
        global current_image_index
        current_image_index = (current_image_index + 1) % 4
        new_image = image_objects[current_image_index]
        image_label.config(image=new_image)
        image_label.image = new_image

    # 绑定点击事件（点击图片时切换）
    image_label.bind("<Button-1>", switch_image)

except Exception as e:
    print(f"图片加载失败: {str(e)}")
    ttk.Label(image_frame, text=f"无法加载图片\n请检查icon文件夹下是否有icon1-4.gif").pack()

# 创建一个框架来放置按钮（原分组框架合并）
buttons_frame = ttk.Frame(main_frame)
buttons_frame.pack(expand=True, anchor='center')

# 存储所有按钮的列表
button_list = []
button_texts = []
file_paths = []

# 最后显示窗口
root.deiconify()

def create_buttons(frame, texts, paths, bootstyle=PRIMARY):
    # 只创建一个行框架，所有按钮都放在这里
    current_row_frame = ttk.Frame(frame)
    current_row_frame.pack(anchor='center')
    # 统一按钮宽度为20（可根据需要调整）
    fixed_width = 12
    for text, path in zip(texts, paths):
        zh_text = BUTTON_TEXTS['zh'][text]
        en_text = BUTTON_TEXTS['en'][text]

        display_text = BUTTON_TEXTS[current_language][text]
        button = ttk.Button(current_row_frame, text=display_text, bootstyle=bootstyle, width=fixed_width, cursor="hand2")
        button.pack(side=ttk.LEFT, padx=5, pady=5)
        button.bind("<Button-1>", lambda event, p=path: run_script(p))
        button.bind("<Enter>", lambda event, t=text: show_details(event, t))
        button.bind("<Leave>", hide_details)
        button_list.append(button)
        button_texts.append(text)
        file_paths.append(path)
        button.update_idletasks()
        button_max_widths.append(fixed_width)

create_buttons(buttons_frame, ["用户手册", "数据清洗", "工具包"],
               ['UserManual', 'DataCleaning', 'Toolkit'],
               bootstyle=SUCCESS)

# 创建语言切换标签
#language_label = ttk.Label(root, text=LANGUAGES[current_language]['switch_language'], foreground="gray", cursor="hand2")
#language_label.bind("<Button-1>", lambda e: switch_language())
#language_label.pack(pady=5)

# 创建版权标签
copyright_label = ttk.Label(root, text=LANGUAGES[current_language]['copyright'], foreground='#f4f4f4', font=('TkDefaultFont', 8))
copyright_label.pack(pady=5)

# 运行主循环
root.mainloop()
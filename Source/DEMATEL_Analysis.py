import tkinter as tk
from tkinter import filedialog, messagebox
import os
import numpy as np
import pandas as pd
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt


# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典，所有键使用英文
languages = {
    "zh": {
        "title": "DEMATEL 分析",
        "select_button": "选择文件",
        "analyze_button": "分析文件",
        "file_not_found": "请选择文件。",
        "analysis_success": "分析完成，结果已保存到 {}\n",
        "no_save_path": "未选择保存路径，结果未保存。",
        "analysis_error": "分析文件时出错: {}",
        "switch_language": "中/英",
        'open_excel_button_text': "示例数据",
        "file_entry_placeholder": "请输入待分析 Excel 文件的完整路径",
        "explanation": {
            "comprehensive_impact_matrix": "反映因素之间综合影响关系的矩阵",
            "causal_degree": "衡量因素对其他因素影响程度的指标",
            "centrality": "衡量因素在系统中重要程度的指标"
        },
        "interpretation": {
            "comprehensive_impact_matrix": "矩阵元素值越大，对应因素之间的影响越强",
            "causal_degree": "原因度为正，该因素为原因因素；原因度为负，该因素为结果因素",
            "centrality": "中心度越大，该因素在系统中越重要"
        },
        "non_square": "输入数据必须是方阵（行数=列数）",
        "non_numeric": "数据中包含非数值内容，请检查",
        "matrix_singular": "矩阵计算异常（可能不可逆），尝试调整输入数据",
        "normalization_method": "归一化方式：除以行和与列和的最大值",
        "mismatch_row_col": "行标题数量与列标题数量不匹配，请检查数据",
        "factor": "因素",
        "causal_degree_chart": "原因度柱状图",
        "centrality_chart": "中心度柱状图",
        "network_diagram": "因素影响关系网络图",
        "network_note": "注：仅显示超过平均影响值的关系",
        "network_import_error": "提示：未安装networkx库，无法生成因素关系网络图。可通过pip install networkx安装。",
        "network_error": "生成网络图时出错：{}",
        "analysis_results": "DEMATEL 分析结果",
        "causal_centrality_title": "原因度与中心度（按中心度排序）",
        "causal_centrality_charts": "原因度和中心度柱状图"
    },
    "en": {
        "title": "DEMATEL Analysis",
        "select_button": "Select File",
        "analyze_button": "Analyze File",
        "file_not_found": "The file does not exist. Please select again.",
        "analysis_success": "Analysis completed. The results have been saved to {}\n",
        "no_save_path": "No save path selected. The results were not saved.",
        "analysis_error": "An error occurred while analyzing the file: {}",
        "switch_language": "Chinese/English",
        'open_excel_button_text': "Example data",
        "file_entry_placeholder": "Please enter the full path of the Excel file to be analyzed",
        "explanation": {
            "comprehensive_impact_matrix": "A matrix reflecting the comprehensive influence relationship between factors",
            "causal_degree": "An indicator to measure the influence degree of a factor on other factors",
            "centrality": "An indicator to measure the importance of a factor in the system"
        },
        "interpretation": {
            "comprehensive_impact_matrix": "The larger the matrix element value, the stronger the influence between corresponding factors",
            "causal_degree": "If the causal degree is positive, the factor is a causal factor; if negative, it is a result factor",
            "centrality": "The larger the centrality, the more important the factor in the system"
        },
        "non_square": "Input data must be a square matrix (rows = columns)",
        "non_numeric": "Data contains non-numeric content, please check",
        "matrix_singular": "Matrix calculation error (may be non-invertible), try adjusting input data",
        "normalization_method": "Normalization method: Divide by max of row sums and column sums",
        "mismatch_row_col": "Number of row headers does not match column headers, please check data",
        "factor": "Factor",
        "causal_degree_chart": "Bar Chart of Causal Degree",
        "centrality_chart": "Bar Chart of Centrality",
        "network_diagram": "Factor Influence Network Diagram",
        "network_note": "Note: Only relationships exceeding the average influence value are shown",
        "network_import_error": "Note: networkx library is not installed, cannot generate factor relationship network graph. Install with pip install networkx.",
        "network_error": "Error generating network graph: {}",
        "analysis_results": "DEMATEL Analysis Results",
        "causal_centrality_title": "Causal Degree and Centrality (sorted by Centrality)",
        "causal_centrality_charts": "Causal Degree and Centrality Bar Charts"
    }
}


class DEMATELAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为中文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data55.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def dematel_analysis(self, data):
        """
        进行 DEMATEL 分析
        :param data: 直接影响矩阵数据
        :return: 综合影响矩阵、原因度、中心度
        """
        # 归一化直接影响矩阵
        n = data.shape[0]
        max_sum_row = np.max(np.sum(data, axis=1))
        max_sum_col = np.max(np.sum(data, axis=0))
        max_value = max(max_sum_row, max_sum_col)

        # 避免除以零
        if max_value == 0:
            max_value = 1e-10
        D = data / max_value

        # 计算综合影响矩阵（增加异常处理）
        I = np.eye(n)
        try:
            # 增加微小扰动避免奇异矩阵
            inv_matrix = np.linalg.inv(I - D + np.eye(n) * 1e-10)
            T = np.dot(D, inv_matrix)
        except np.linalg.LinAlgError:
            raise ValueError(languages[self.current_language]["matrix_singular"])

        # 计算原因度和中心度
        sum_row = np.sum(T, axis=1)
        sum_col = np.sum(T, axis=0)
        causal_degree = sum_row - sum_col
        centrality = sum_row + sum_col

        return T, causal_degree, centrality

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]['file_not_found'])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]['file_not_found'])
            return
        try:
            # 读取Excel文件，第一行为列标题，第一列为行标题
            df = pd.read_excel(file_path, header=0, index_col=0)
            # 获取列标题（因素名）
            col_names = df.columns.tolist()
            # 获取行标题（因素名）
            row_names = df.index.tolist()

            # 数据校验：检查行列标题数量是否一致
            if len(row_names) != len(col_names):
                messagebox.showerror("Error", languages[self.current_language]["mismatch_row_col"])
                return

            # 提取数据矩阵（不包含表头）
            data = df.values

            # 数据校验：检查是否为方阵
            if data.shape[0] != data.shape[1]:
                messagebox.showerror("Error", languages[self.current_language]["non_square"])
                return

            # 数据校验：检查是否包含非数值
            if not np.issubdtype(data.dtype, np.number):
                messagebox.showerror("Error", languages[self.current_language]["non_numeric"])
                return

            # 处理可能的空值
            if np.isnan(data).any():
                data = np.nan_to_num(data)  # 将空值替换为0

            # 进行 DEMATEL 分析
            T, causal_degree, centrality = self.dematel_analysis(data)

            # 整理数据（使用实际因素名，若无则用默认名）
            if len(col_names) != data.shape[0]:
                factor_names = [f"{languages[self.current_language]['factor']}{i + 1}" for i in range(data.shape[0])]
            else:
                # 使用行标题作为因素名（与列标题保持一致）
                factor_names = row_names

            # 保留4位小数
            T_df = pd.DataFrame(T.round(4), index=row_names, columns=col_names)

            # 原因度和中心度排序（按中心度降序）
            result_df = pd.DataFrame({
                languages[self.current_language]['factor']: factor_names,
                "causal_degree": causal_degree.round(4),
                "centrality": centrality.round(4)
            }).sort_values(by="centrality", ascending=False)

            # 添加解释说明
            explanations = languages[self.current_language]['explanation']
            interpretations = languages[self.current_language]['interpretation']

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建一个新的 Word 文档
                document = Document()

                # 添加标题
                title = document.add_heading(languages[self.current_language]['analysis_results'], level=1)
                title.alignment = 1  # 居中对齐

                # 添加归一化说明
                document.add_paragraph(languages[self.current_language]["normalization_method"])

                # 添加综合影响矩阵
                document.add_heading(explanations["comprehensive_impact_matrix"], level=2)
                document.add_paragraph(interpretations["comprehensive_impact_matrix"])
                table = document.add_table(rows=len(T_df) + 1, cols=len(T_df.columns) + 1)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = ''  # 第一列的表头为空，后续行将填充行标题
                for col_idx, col_name in enumerate(T_df.columns):
                    hdr_cells[col_idx + 1].text = col_name
                for row_idx, row in enumerate(T_df.values):
                    row_cells = table.rows[row_idx + 1].cells
                    row_cells[0].text = T_df.index[row_idx]  # 填充行标题
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx + 1].text = str(value)

                # 添加原因度和中心度（合并展示并排序）
                document.add_heading(languages[self.current_language]['causal_centrality_title'], level=2)
                document.add_paragraph(explanations["causal_degree"])
                document.add_paragraph(interpretations["causal_degree"])
                document.add_paragraph(explanations["centrality"])
                document.add_paragraph(interpretations["centrality"])

                table = document.add_table(rows=len(result_df) + 1, cols=3)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = languages[self.current_language]['factor']
                hdr_cells[1].text = "causal_degree"  # 保持英文，作为专业术语
                hdr_cells[2].text = "centrality"      # 保持英文，作为专业术语
                for row_idx, row in result_df.iterrows():
                    row_cells = table.rows[row_idx + 1].cells
                    row_cells[0].text = row[languages[self.current_language]['factor']]
                    row_cells[1].text = str(row["causal_degree"])
                    row_cells[2].text = str(row["centrality"])

                # 生成原因度和中心度柱状图（使用实际因素名）
                fig, axes = plt.subplots(2, 1, figsize=(8, 10))
                axes[0].bar(result_df[languages[self.current_language]['factor']], result_df["causal_degree"])
                axes[0].set_title(languages[self.current_language]['causal_degree_chart'])
                axes[0].set_xlabel(languages[self.current_language]['factor'])
                axes[0].set_ylabel("causal_degree")  # 保持英文，作为专业术语
                plt.setp(axes[0].get_xticklabels(), rotation=0, ha='center')  # 标签旋转避免重叠

                axes[1].bar(result_df[languages[self.current_language]['factor']], result_df["centrality"])
                axes[1].set_title(languages[self.current_language]['centrality_chart'])
                axes[1].set_xlabel(languages[self.current_language]['factor'])
                axes[1].set_ylabel("centrality")      # 保持英文，作为专业术语
                plt.setp(axes[1].get_xticklabels(), rotation=0, ha='center')

                # 生成因素关系网络图
                try:
                    import networkx as nx

                    # 选择显著影响关系（例如，取综合影响矩阵中值大于平均值的关系）
                    threshold = np.mean(T)
                    G = nx.DiGraph()

                    # 添加节点
                    for name in factor_names:
                        G.add_node(name)

                    # 添加边（只添加超过阈值的影响关系）
                    for i, from_node in enumerate(factor_names):
                        for j, to_node in enumerate(factor_names):
                            if i != j and T[i, j] > threshold:
                                G.add_edge(from_node, to_node, weight=round(T[i, j], 3))

                    # 绘制网络图
                    plt.figure(figsize=(10, 8))
                    pos = nx.spring_layout(G, k=0.5, iterations=50)
                    nx.draw_networkx_nodes(G, pos, node_size=1000, node_color='lightblue')
                    nx.draw_networkx_labels(G, pos, font_size=10)

                    # 绘制边和权重
                    edges = G.edges(data=True)
                    nx.draw_networkx_edges(G, pos, edgelist=edges, arrowstyle='->', arrowsize=20)
                    edge_labels = {(u, v): d['weight'] for u, v, d in edges}
                    nx.draw_networkx_edge_labels(G, pos, edge_labels=edge_labels, font_size=8)

                    plt.title(languages[self.current_language]['network_diagram'])
                    plt.axis('off')

                    # 保存网络图
                    network_img_path = os.path.splitext(save_path)[0] + '_network.png'
                    plt.tight_layout()
                    plt.savefig(network_img_path, dpi=300)
                    plt.close()

                    # 将网络图插入到Word文档
                    document.add_heading(languages[self.current_language]['network_diagram'], level=2)
                    document.add_paragraph(languages[self.current_language]['network_note'])
                    document.add_picture(network_img_path)
                except ImportError:
                    document.add_paragraph(languages[self.current_language]['network_import_error'])
                except Exception as e:
                    document.add_paragraph(languages[self.current_language]['network_error'].format(str(e)))

                # 保存图片
                img_path = os.path.splitext(save_path)[0] + '_charts.png'
                plt.tight_layout()
                plt.savefig(img_path, dpi=300)
                plt.close()

                # 将图片插入到 Word 文档中
                document.add_heading(languages[self.current_language]['causal_centrality_charts'], level=2)
                document.add_picture(img_path)

                # 保存 Word 文档
                document.save(save_path)

                result_msg = languages[self.current_language]['analysis_success'].format(save_path)
                self.result_label.config(text=result_msg, wraplength=400)

            else:
                self.result_label.config(text=languages[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]['analysis_error'].format(str(e)))
            messagebox.showerror("Error", str(e))

    def switch_language(self, event=None):
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'
        self.root.title(languages[self.current_language]['title'])
        self.select_button.config(text=languages[self.current_language]['select_button'])
        self.analyze_button.config(text=languages[self.current_language]['analyze_button'])
        self.switch_language_label.config(text=languages[self.current_language]['switch_language'])
        self.open_excel_label.config(text=languages[self.current_language]["open_excel_button_text"])

        # 切换语言时更新提示信息
        current_text = self.file_entry.get()
        if current_text == languages['zh' if self.current_language == 'en' else 'en']["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=languages[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=languages[self.current_language]["switch_language"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = DEMATELAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()
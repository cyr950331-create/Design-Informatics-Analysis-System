import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
from scipy import stats
import matplotlib.pyplot as plt
import pathlib
import pingouin as pg
from docx import Document
from docx.shared import Inches

# 设置支持中文的字体
plt.rcParams['font.family'] = 'SimHei'  # 使用黑体字体，可根据系统情况修改
plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

# 定义语言字典，包含所有需要翻译的文本
languages = {
    'zh': {
        'title': "协方差分析",
        'select_button_text': "选择文件",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'analyze_button_text': "分析文件",
        'no_file_selected': "请选择文件。",
        'file_not_exists': "文件不存在，请重新选择。",
        'analysis_error': "分析文件时出错: {}",
        'analysis_complete': "分析完成，结果已保存到 {}，相关图片已保存到 {}",
        'no_save_path_selected': "未选择保存路径，结果未保存。",
        'switch_language_button_text': "中/英",
        'open_excel_button_text': "示例数据",
        'explanation': {
            "covariance_analysis": "在控制一个或多个协变量的影响下，分析不同组之间因变量的均值是否存在显著差异。",
            "group_effect": "分组效应：不同组别之间的差异对因变量的影响",
            "covariate_effect": "协变量效应：协变量对因变量的影响",
            "interaction_effect": "交互效应：分组与协变量的交互作用对因变量的影响",
            "residual": "残差：模型无法解释的变异部分，反映了数据中未被模型捕捉的随机波动"
        },
        'interpretation': {
            "source": "来源：指变异的来源，如分组、协变量、交互效应或残差",
            "df": "自由度：反映了可以自由变化的观测值数量，用于统计检验",
            "F": "F 统计量，用于检验相应来源变异的显著性。F值越大，表明该来源的效应越可能显著。",
            "p_unc": "未经校正的 p 值，小于显著性水平（通常为 0.05）时，表明相应来源的效应显著。",
            "np2": "偏 eta 平方，反映了该来源的变异在总变异中所占的比例，值越大表明效应越强。",
            "ci_2.5": "效应量的95%置信区间下限，区间不包含0表明效应显著",
            "ci_97.5": "效应量的95%置信区间上限，区间宽度反映了估计的精确度"
        },
        'word_heading': '协方差分析结果',
        'statistical_results': '统计结果',
        'explanation_heading': '变异来源说明',
        'interpretation_heading': '统计量解读',
        'plot_title': '协方差分析结果',
        'source': '来源',
        'df': '自由度',
        'F': 'F值',
        'p_unc': '未校正p值',
        'np2': '偏eta平方',
        'residual': '残差',
        'ci_2.5': '95%置信区间下限',
        'ci_97.5': '95%置信区间上限'
    },
    'en': {
        'title': "Analysis of Covariance",
        'select_button_text': "Select File",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'analyze_button_text': "Analyze File",
        'no_file_selected': "Please select a valid file path.",
        'file_not_exists': "The file does not exist. Please select again.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'analysis_complete': "Analysis completed. The results have been saved to {}, and the image has been saved to {}",
        'no_save_path_selected': "No save path selected. The results were not saved.",
        'switch_language_button_text': "Chinese/English",
        'open_excel_button_text': "Example data",
        'explanation': {
            "covariance_analysis": "Analyze whether there are significant differences in the means of the dependent variable between different groups while controlling for the effects of one or more covariates.",
            "group_effect": "Group effect: The impact of differences between groups on the dependent variable",
            "covariate_effect": "Covariate effect: The impact of covariates on the dependent variable",
            "interaction_effect": "Interaction effect: The combined impact of group and covariate interactions on the dependent variable",
            "residual": "Residual: The portion of variance not explained by the model, reflecting random fluctuations not captured by the model"
        },
        'interpretation': {
            "source": "Source: The origin of variation, such as group, covariate, interaction effect, or residual",
            "df": "Degrees of freedom: Reflects the number of observations that can vary freely, used in statistical tests",
            "F": "F statistic, used to test the significance of variation from the corresponding source. A larger F value indicates a more likely significant effect.",
            "p_unc": "Uncorrected p-value. When less than the significance level (usually 0.05), it indicates a significant effect from the corresponding source.",
            "np2": "Partial eta squared, reflecting the proportion of variance from this source in the total variance. A larger value indicates a stronger effect.",
            "ci_2.5": "95% confidence interval lower bound of effect size. An interval not containing 0 indicates a significant effect",
            "ci_97.5": "95% confidence interval upper bound of effect size. The width of the interval reflects the precision of the estimate"
        },
        'word_heading': 'Analysis of Covariance (ANCOVA) Results',
        'statistical_results': 'Statistical Results',
        'explanation_heading': 'Explanation of Variance Sources',
        'interpretation_heading': 'Statistic Interpretation',
        'plot_title': 'Analysis of Covariance (ANCOVA) Results',
        'source': 'Source',
        'df': 'df',
        'F': 'F',
        'p_unc': 'p-unc',
        'np2': 'np2',
        'residual': 'Residual',
        'ci_2.5': 'CI[2.5%]',
        'ci_97.5': 'CI[97.5%]'
    }
}


class ANCOVAAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为中文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data27.xls")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(
                    text=f"{languages[self.current_language]['file_not_exists'].split('，')[0]}：{excel_path}")
        except Exception as e:
            self.result_label.config(
                text=f"{languages[self.current_language]['analysis_error'].split('：')[0]}：{str(e)}")

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 假设第一列是分组变量，最后一列是因变量，其余列是协变量
            group_var = df.columns[0]
            dep_var = df.columns[-1]
            covar_vars = df.columns[1:-1]

            # 进行协方差分析
            # 添加效应量置信区间和交互效应检验
            ancova = pg.ancova(
                data=df,
                dv=dep_var,
                between=group_var,
                covar=covar_vars.tolist(),
                effsize="np2"
            )

            # 重命名列名以支持多语言
            ancova = ancova.rename(columns={
                'Source': languages[self.current_language]['source'],
                'df': languages[self.current_language]['df'],
                'F': languages[self.current_language]['F'],
                'p-unc': languages[self.current_language]['p_unc'],
                'np2': languages[self.current_language]['np2'],
                'Residual': languages[self.current_language]['residual'],
                'CI[2.5%]': languages[self.current_language]['ci_2.5'],
                'CI[97.5%]': languages[self.current_language]['ci_97.5']
            })

            # 添加解释说明和解读
            explanations = languages[self.current_language]['explanation']
            interpretations = languages[self.current_language]['interpretation']

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建一个新的Word文档
                doc = Document()

                # 添加标题
                doc.add_heading(languages[self.current_language]['word_heading'], 0)

                # 添加统计结果表格
                doc.add_heading(languages[self.current_language]['statistical_results'], level=1)
                columns = ancova.columns.tolist()

                # 创建表格，行数为数据行数+1（表头）
                table = doc.add_table(rows=1, cols=len(columns), style='Table Grid')

                # 设置表头
                hdr_cells = table.rows[0].cells
                for i, col_name in enumerate(columns):
                    hdr_cells[i].text = col_name

                # 填充表格数据
                for idx, row in ancova.iterrows():
                    row_cells = table.add_row().cells
                    for i, col_name in enumerate(columns):
                        # 格式化数值，保留4位小数
                        value = row[col_name]
                        if isinstance(value, float):
                            row_cells[i].text = f"{value:.4f}"
                        else:
                            row_cells[i].text = str(value)

                # 调整表格列宽使其更紧凑
                for col in table.columns:
                    for cell in col.cells:
                        cell.width = Inches(1.2)  # 根据需要调整列宽

                # 添加变异来源说明（与表格中的来源一一对应）
                doc.add_heading(languages[self.current_language]['explanation_heading'], level=1)
                expl_para = doc.add_paragraph()
                expl_para.add_run(f"• {explanations['covariance_analysis']}\n\n")

                # 添加分组效应说明
                expl_para.add_run(
                    f"• {languages[self.current_language]['source']}: {group_var} - {explanations['group_effect']}\n")

                # 添加协变量效应说明
                for covar in covar_vars:
                    expl_para.add_run(
                        f"• {languages[self.current_language]['source']}: {covar} - {explanations['covariate_effect']}\n")

                # 添加交互效应说明（如果存在）
                interaction_sources = [source for source in ancova[languages[self.current_language]['source']]
                                       if '*' in str(source) or '交互' in str(source)]
                for source in interaction_sources:
                    expl_para.add_run(
                        f"• {languages[self.current_language]['source']}: {source} - {explanations['interaction_effect']}\n")

                # 添加残差说明
                expl_para.add_run(
                    f"• {languages[self.current_language]['source']}: {languages[self.current_language]['residual']} - {explanations['residual']}\n")

                # 添加统计量解读（与表格中的统计量一一对应）
                doc.add_heading(languages[self.current_language]['interpretation_heading'], level=1)
                interp_para = doc.add_paragraph()
                for col_name in columns:
                    # 找到对应的解读键
                    if col_name == languages[self.current_language]['source']:
                        key = 'source'
                    elif col_name == languages[self.current_language]['df']:
                        key = 'df'
                    elif col_name == languages[self.current_language]['F']:
                        key = 'F'
                    elif col_name == languages[self.current_language]['p_unc']:
                        key = 'p_unc'
                    elif col_name == languages[self.current_language]['np2']:
                        key = 'np2'
                    elif col_name == languages[self.current_language]['ci_2.5']:
                        key = 'ci_2.5'
                    elif col_name == languages[self.current_language]['ci_97.5']:
                        key = 'ci_97.5'
                    else:
                        continue

                    interp_para.add_run(f"• {col_name}: {interpretations[key]}\n")

                # 生成结果图片
                plot_path = os.path.splitext(save_path)[0] + '_ancova_plot.png'
                plt.figure()
                for group in df[group_var].unique():
                    group_data = df[df[group_var] == group]
                    plt.scatter(group_data[covar_vars[0]], group_data[dep_var], label=str(group))

                # 使用当前语言设置图表标签
                plt.xlabel(str(covar_vars[0]))  # 协变量名保持原样（来自数据）
                plt.ylabel(str(dep_var))  # 因变量名保持原样（来自数据）
                plt.title(languages[self.current_language]['plot_title'])
                plt.legend(title=group_var)  # 分组变量名保持原样（来自数据）
                plt.savefig(plot_path)
                plt.close()

                # 将图片插入到Word文档中
                doc.add_picture(plot_path, width=Inches(6))

                # 保存Word文档
                doc.save(save_path)

                # 更新结果标签
                self.result_label.config(
                    text=languages[self.current_language]["analysis_complete"].format(save_path, plot_path),
                    wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]["no_save_path_selected"])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]["analysis_error"].format(str(e)))

    def switch_language(self, event):
        # 切换当前语言
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        current_entry_text = self.file_entry.get()
        # 只有当输入框显示的是占位符时才更新
        if current_entry_text in [languages['zh']["file_entry_placeholder"], languages['en']["file_entry_placeholder"]]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])
        self.open_excel_label.config(text=languages[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY, cursor="hand2")
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS, cursor="hand2")
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=languages[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame,
                                               text=languages[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = ANCOVAAnalysisApp()
    app.run()

if __name__ == "__main__":
    run_app()
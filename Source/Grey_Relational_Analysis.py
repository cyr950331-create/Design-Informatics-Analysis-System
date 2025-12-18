import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog, ttk as ttk_tk
import tkinter as tk
import matplotlib.pyplot as plt
import seaborn as sns
import pathlib
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典，所有显示文本均通过语言字典管理
LANGUAGES = {
    'zh': {
        'title': "灰色关联分析法",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "请选择文件。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "中/英",
        'open_excel_button_text': "示例数据",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'min_columns': "至少需要1个参考序列和1个比较序列（至少2列数据）",
        'no_numeric_data': "文件中没有有效的数值数据",
        'initialization': "初值化",
        'mean_normalization': "均值化",
        'interval_normalization': "区间化",
        'reference_sequence': "参考序列",
        'comparison_sequences': "比较序列",
        'relational_coefficient_matrix': "关联系数矩阵",
        'relational_degree': "关联度",
        'ranking_result': "关联度排序结果",
        'statistics': "统计量",
        'statistic_value': "统计量值",
        'explanation': "解释说明",
        'interpretation': "结果解读",
        'analysis_results_heading': "灰色关联分析结果",
        'relational_degree_chart': "关联度柱状图",
        'relational_coefficient_heatmap': "关联系数矩阵热力图",
        'time_point': "时间点",
        'xlabel_comparison': "比较序列",
        'ylabel_degree': "关联度",
        'xlabel_time': "时间点",
        'ylabel_comparison': "比较序列"
    },
    'en': {
        'title': "Grey Relational Analysis",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Chinese/English",
        'open_excel_button_text': "Example data",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'min_columns': "At least 1 reference sequence and 1 comparison sequence (minimum 2 columns) are required",
        'no_numeric_data': "There is no valid numerical data in the file",
        'initialization': "Initialization",
        'mean_normalization': "Mean Normalization",
        'interval_normalization': "Interval Normalization",
        'reference_sequence': "Reference Sequence",
        'comparison_sequences': "Comparison Sequences",
        'relational_coefficient_matrix': "Relational Coefficient Matrix",
        'relational_degree': "Relational Degree",
        'ranking_result': "Ranking Result",
        'statistics': "Statistic",
        'statistic_value': "Statistic Value",
        'explanation': "Explanation",
        'interpretation': "Interpretation",
        'analysis_results_heading': "Grey Relational Analysis Results",
        'relational_degree_chart': "Bar Chart of Relational Degree",
        'relational_coefficient_heatmap': "Heatmap of Relational Coefficient Matrix",
        'time_point': "Time Point",
        'xlabel_comparison': "Comparison Sequence",
        'ylabel_degree': "Relational Degree",
        'xlabel_time': "Time Point",
        'ylabel_comparison': "Comparison Sequence"
    }
}


class GreyRelationalAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为中文
        self.current_language = "en"
        # 固定参数设置
        self.has_header = True  # 默认包含表头
        self.preprocessing_method = 'initialization'  # 固定为初值化
        self.rho = 0.5  # 固定分辨系数为0.5

        # 存储预处理方法的内部标识与显示文本映射
        self.preprocessing_methods = self._get_preprocessing_methods()

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def _get_preprocessing_methods(self):
        """获取预处理方法的内部标识与显示文本映射"""
        return {
            'initialization': LANGUAGES[self.current_language]['initialization'],
            'mean_normalization': LANGUAGES[self.current_language]['mean_normalization'],
            'interval_normalization': LANGUAGES[self.current_language]['interval_normalization']
        }

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data47.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        """当用户点击输入框时，清除提示信息"""
        if self.file_entry.get() == LANGUAGES[self.current_language]['file_entry_placeholder']:
            self.file_entry.delete(0, tk.END)
            self.file_entry.configure(style="TEntry")

    def on_focusout(self, event):
        """当用户离开输入框时，如果没有输入内容，恢复提示信息"""
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
            self.file_entry.configure(style="Gray.TEntry")

    def preprocess_data(self, data, method):
        """数据预处理方法"""
        if method == "initialization":  # 初值化
            return data / data[0, :] if data.ndim == 2 else data / data[0]
        elif method == "mean_normalization":  # 均值化
            mean = np.mean(data, axis=0)
            return data / mean if data.ndim == 2 else data / mean
        elif method == "interval_normalization":  # 区间化(0-1标准化)
            min_val = np.min(data, axis=0)
            max_val = np.max(data, axis=0)
            return (data - min_val) / (max_val - min_val + 1e-10) if data.ndim == 2 else (data - min_val) / (
                    max_val - min_val + 1e-10)
        else:
            # 添加默认处理，抛出异常避免返回None
            raise ValueError(f"不支持的预处理方法: {method}")

    def grey_relational_analysis(self, reference_sequence, comparison_sequences, rho):
        # 检查输入序列有效性
        if reference_sequence.size == 0:
            raise ValueError(LANGUAGES[self.current_language]['reference_sequence'] + "不能为空")
        if comparison_sequences.size == 0:
            raise ValueError(LANGUAGES[self.current_language]['comparison_sequences'] + "不能为空")
        if reference_sequence.ndim != 1:
            raise ValueError(LANGUAGES[self.current_language]['reference_sequence'] + "必须是一维数组")

        # 使用固定的预处理方法
        preprocessing_method = self.preprocessing_method

        # 数据预处理
        reference_sequence = self.preprocess_data(reference_sequence, preprocessing_method)
        comparison_sequences = self.preprocess_data(comparison_sequences, preprocessing_method)

        # 计算差序列
        diff_matrix = np.abs(comparison_sequences - reference_sequence)

        # 计算两级最小差和两级最大差
        min_min_diff = np.min(np.min(diff_matrix))
        max_max_diff = np.max(np.max(diff_matrix))

        # 计算关联系数矩阵
        relational_coefficient_matrix = (min_min_diff + rho * max_max_diff) / (diff_matrix + rho * max_max_diff)

        # 计算关联度
        relational_degree = np.mean(relational_coefficient_matrix, axis=1)

        # 对关联度进行排序
        ranking = np.argsort(-relational_degree) + 1

        return relational_coefficient_matrix, relational_degree, ranking

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]['file_entry_placeholder']:
            file_path = ""
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 使用固定参数
            has_header = self.has_header
            rho = self.rho

            # 打开 Excel 文件，支持表头设置
            df = pd.read_excel(file_path, header=0 if has_header else None)

            # 数据清洗：仅保留数值列
            df = df.select_dtypes(include=[np.number])
            if df.empty:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_numeric_data'])
                return

            # 检查列数是否足够
            if len(df.columns) < 2:
                self.result_label.config(text=LANGUAGES[self.current_language]['min_columns'])
                return

            # 获取列名（表头信息）
            column_names = df.columns.tolist()
            reference_name = column_names[0]  # 参考序列名称
            comparison_names = column_names[1:]  # 比较序列名称

            data = df.values

            # 转换为浮点类型
            data = data.astype(float)

            # 第一列为参考序列，其余列为比较序列（转置后行变为列）
            reference_sequence = data[:, 0]  # 第一列
            comparison_sequences = data[:, 1:].T  # 其余列转置为行

            # 检查比较序列是否有效
            if comparison_sequences.size == 0:
                self.result_label.config(text=LANGUAGES[self.current_language]['min_columns'])
                return

            # 进行灰色关联分析
            relational_coefficient_matrix, relational_degree, ranking = self.grey_relational_analysis(
                reference_sequence, comparison_sequences, rho)

            # 整理数据（确保列数与表头一致）
            data = [
                [f"{LANGUAGES[self.current_language]['reference_sequence']} ({reference_name})",
                 reference_sequence.tolist()],
                [LANGUAGES[self.current_language]['comparison_sequences'],
                 [f"{name}: {seq.tolist()}" for name, seq in zip(comparison_names, comparison_sequences)]],
                [LANGUAGES[self.current_language]['relational_coefficient_matrix'],
                 relational_coefficient_matrix.tolist()],
                [LANGUAGES[self.current_language]['relational_degree'],
                 [f"{name}: {degree:.4f}" for name, degree in zip(comparison_names, relational_degree)]],
                [LANGUAGES[self.current_language]['ranking_result'],
                 [
                     f"{ranking[i]}: {comparison_names[i]} ({LANGUAGES[self.current_language]['relational_degree']}: {relational_degree[i]:.4f})"
                     for i in range(len(ranking))]]
            ]
            headers = [LANGUAGES[self.current_language]['statistics'],
                       LANGUAGES[self.current_language]['statistic_value']]

            df_result = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            explanations = {
                LANGUAGES[self.current_language]['reference_sequence']:
                    f"{LANGUAGES[self.current_language]['reference_sequence']}作为衡量其他序列关联程度的标准"
                    if self.current_language == 'zh'
                    else f"{LANGUAGES[self.current_language]['reference_sequence']} as a standard for measuring the degree of association of other sequences",

                LANGUAGES[self.current_language]['comparison_sequences']:
                    f"需要分析与{LANGUAGES[self.current_language]['reference_sequence']}关联程度的序列"
                    if self.current_language == 'zh'
                    else f"Sequences whose degree of association with the {LANGUAGES[self.current_language]['reference_sequence']} needs to be analyzed",

                LANGUAGES[self.current_language]['relational_coefficient_matrix']:
                    f"数值越大，该时刻{LANGUAGES[self.current_language]['comparison_sequences']}与{LANGUAGES[self.current_language]['reference_sequence']}的关联程度越高"
                    if self.current_language == 'zh'
                    else f"The larger the value, the higher the degree of association between the {LANGUAGES[self.current_language]['comparison_sequences']} and {LANGUAGES[self.current_language]['reference_sequence']} at that time point",

                LANGUAGES[self.current_language]['relational_degree']:
                    f"值越大，说明{LANGUAGES[self.current_language]['comparison_sequences']}与{LANGUAGES[self.current_language]['reference_sequence']}的整体关联程度越高"
                    if self.current_language == 'zh'
                    else f"The larger the value, the higher the overall degree of association between the {LANGUAGES[self.current_language]['comparison_sequences']} and {LANGUAGES[self.current_language]['reference_sequence']}",

                LANGUAGES[self.current_language]['ranking_result']:
                    f"排名越靠前，与{LANGUAGES[self.current_language]['reference_sequence']}的关联程度越高"
                    if self.current_language == 'zh'
                    else f"The higher the ranking, the higher the degree of association with the {LANGUAGES[self.current_language]['reference_sequence']}"
            }

            interpretation_df = pd.DataFrame([explanations])
            interpretation_df = interpretation_df.reindex(
                columns=[
                    LANGUAGES[self.current_language]['reference_sequence'],
                    LANGUAGES[self.current_language]['comparison_sequences'],
                    LANGUAGES[self.current_language]['relational_coefficient_matrix'],
                    LANGUAGES[self.current_language]['relational_degree'],
                    LANGUAGES[self.current_language]['ranking_result']
                ])
            interpretation_df.insert(0,
                                     f"{LANGUAGES[self.current_language]['statistics']}_{LANGUAGES[self.current_language]['explanation']}",
                                     LANGUAGES[self.current_language]['explanation'])

            # 添加分析结果解读（与解释说明内容相同，保持原结构）
            interpretation_df2 = pd.DataFrame([explanations])
            interpretation_df2 = interpretation_df2.reindex(
                columns=[
                    LANGUAGES[self.current_language]['reference_sequence'],
                    LANGUAGES[self.current_language]['comparison_sequences'],
                    LANGUAGES[self.current_language]['relational_coefficient_matrix'],
                    LANGUAGES[self.current_language]['relational_degree'],
                    LANGUAGES[self.current_language]['ranking_result']
                ])
            interpretation_df2.insert(0,
                                      f"{LANGUAGES[self.current_language]['statistics']}_{LANGUAGES[self.current_language]['interpretation']}",
                                      LANGUAGES[self.current_language]['interpretation'])

            # 合并数据、解释说明和结果解读
            combined_df = pd.concat([df_result, interpretation_df, interpretation_df2], ignore_index=True)

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建一个新的 Word 文档
                document = Document()

                # 添加标题
                document.add_heading(LANGUAGES[self.current_language]['analysis_results_heading'], 0)

                # 添加序列信息说明
                document.add_paragraph(f"{LANGUAGES[self.current_language]['reference_sequence']}: {reference_name}")
                document.add_paragraph(
                    f"{LANGUAGES[self.current_language]['comparison_sequences']}: {', '.join(comparison_names)}")
                document.add_paragraph("")  # 空行

                # 添加分析结果表格
                table = document.add_table(rows=1, cols=len(headers))
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(headers):
                    hdr_cells[col_idx].text = header

                for row in combined_df.values.tolist():
                    row_cells = table.add_row().cells
                    for col_idx, value in enumerate(row[:len(headers)]):  # 只取与表头匹配的列数
                        row_cells[col_idx].text = str(value)

                # 生成关联度柱状图
                fig, ax = plt.subplots(figsize=(12, 8))
                bars = ax.bar(range(len(relational_degree)), relational_degree)
                ax.set_title(LANGUAGES[self.current_language]['relational_degree_chart'])
                ax.set_xlabel(LANGUAGES[self.current_language]['xlabel_comparison'])
                ax.set_ylabel(LANGUAGES[self.current_language]['ylabel_degree'])
                # 设置x轴标签为比较序列名称
                ax.set_xticks(range(len(comparison_names)))
                ax.set_xticklabels(comparison_names, rotation=0, ha='right')
                # 在柱状图上显示关联度值
                for bar, degree in zip(bars, relational_degree):
                    height = bar.get_height()
                    ax.text(bar.get_x() + bar.get_width() / 2., height,
                            f'{degree:.4f}',
                            ha='center', va='bottom')

                # 保存图片
                bar_img_path = os.path.splitext(save_path)[0] + '_relational_degree.png'
                plt.tight_layout()  # 调整布局，防止标签被截断
                plt.savefig(bar_img_path,dpi=300)
                plt.close()

                # 生成关联系数矩阵热力图
                fig, ax = plt.subplots(figsize=(12, 8))
                sns.heatmap(relational_coefficient_matrix, annot=True, cmap="YlGnBu", ax=ax,
                            xticklabels=[f'{LANGUAGES[self.current_language]["time_point"]} {i + 1}'
                                         for i in range(relational_coefficient_matrix.shape[1])],
                            yticklabels=comparison_names)
                ax.set_title(LANGUAGES[self.current_language]['relational_coefficient_heatmap'])
                ax.set_xlabel(LANGUAGES[self.current_language]['xlabel_time'])
                ax.set_ylabel(LANGUAGES[self.current_language]['ylabel_comparison'])
                plt.tight_layout()  # 调整布局
                heatmap_img_path = os.path.splitext(save_path)[0] + '_coefficient_matrix.png'
                plt.savefig(heatmap_img_path,dpi=300)
                plt.close()

                # 在 Word 文档中插入图片
                document.add_heading(LANGUAGES[self.current_language]['relational_degree_chart'], level=1)
                document.add_picture(bar_img_path, width=Inches(6))

                document.add_heading(LANGUAGES[self.current_language]['relational_coefficient_heatmap'], level=1)
                document.add_picture(heatmap_img_path, width=Inches(6))

                # 保存 Word 文档
                document.save(save_path)

                # 删除临时图片文件
                os.remove(bar_img_path)
                os.remove(heatmap_img_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(save_path)
                self.result_label.config(text=result_msg, wraplength=400)

            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'
        self.root.title(LANGUAGES[self.current_language]['title'])
        self.select_button.config(text=LANGUAGES[self.current_language]['select_button'])
        self.analyze_button.config(text=LANGUAGES[self.current_language]['analyze_button'])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]['switch_language'])
        self.open_excel_label.config(text=LANGUAGES[self.current_language]["open_excel_button_text"])

        # 切换语言时更新提示信息
        current_text = self.file_entry.get()
        if current_text == LANGUAGES['zh' if self.current_language == 'en' else 'en']['file_entry_placeholder']:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建自定义样式
        style = ttk.Style()
        style.configure("Gray.TEntry", foreground="gray")

        # 创建主容器，用于居中所有元素
        main_frame = ttk.Frame(self.root)
        main_frame.pack(expand=True)  # 让主容器扩展以填充窗口

        # 文件选择区域 - 垂直布局
        file_frame = ttk.Frame(main_frame)
        file_frame.pack(pady=10)  # 上下边距

        # 创建文件选择按钮（放在输入框上方）
        self.select_button = ttk.Button(
            file_frame,
            text=LANGUAGES[self.current_language]['select_button'],
            command=self.select_file,
            bootstyle=PRIMARY
        )
        self.select_button.pack(pady=(0, 5))  # 按钮下方留白

        # 创建文件路径输入框（在按钮下方）
        self.file_entry = ttk.Entry(file_frame, width=50, style="Gray.TEntry")
        self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
        self.file_entry.pack()
        self.file_entry.bind("<FocusIn>", self.on_entry_click)
        self.file_entry.bind("<FocusOut>", self.on_focusout)

        # 创建分析按钮（在输入框下方）
        self.analyze_button = ttk.Button(
            main_frame,
            text=LANGUAGES[self.current_language]['analyze_button'],
            command=self.analyze_file,
            bootstyle=SUCCESS
        )
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            main_frame,
            text=LANGUAGES[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建语言切换标签（在分析按钮下方）
        self.switch_language_label = ttk.Label(
            main_frame,
            text=LANGUAGES[self.current_language]['switch_language'],
            foreground="gray",
            cursor="hand2"
        )
        self.switch_language_label.pack(pady=5)
        self.switch_language_label.bind("<Button-1>", self.switch_language)

        # 创建结果显示标签（在语言切换标签下方）
        self.result_label = ttk.Label(main_frame, text="", justify=tk.LEFT, wraplength=500)
        self.result_label.pack(pady=10, padx=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = GreyRelationalAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()
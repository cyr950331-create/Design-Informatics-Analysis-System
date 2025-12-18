import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
from scipy import stats
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from scipy.stats import bootstrap

# 设置支持中文的字体
plt.rcParams['font.family'] = 'SimHei'
# 解决负号显示问题
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典，包含所有需要翻译的内容
LANGUAGES = {
    'zh': {
        'title': "配对样本 Wilcoxon 检验",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "请选择文件。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "中/英",
        'open_excel_button_text': "示例数据",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'stats': {
            "w_stat": "W统计量",
            "z_stat": "Z统计量",
            "p_value": "p值",
            "significance": "显著性",
            "effect_size": "效应量(r)",
            "sample_size": "样本量",
            "median": "中位数",
            "confidence_interval": "均值差异的置信区间"
        },
        'headers': ["统计量", "结果"],
        'explanation_title': "解释说明",
        'interpretation_title': "结果解读",
        'charts_title': "图表",
        'explanation': {
            "w_stat": "用于检验配对样本之间是否存在显著差异。",
            "sample_size": "配对样本中的观测值对数。",
            "median": "配对样本差值数据的中间值，将数据分为上下两部分。",
            "p_value": "p值小于显著性水平（通常为0.05）时，拒绝原假设，认为配对样本之间存在显著差异；否则，接受原假设，认为无显著差异。",
            "confidence_interval": "包含真实均值差异的一个区间，反映了估计的不确定性。"
        },
        'interpretation': {
            "w_stat": "w统计量的绝对值越大，说明配对样本之间的差异越显著。",
            "p_value": "用于判断配对样本之间是否存在显著差异的依据。",
            "sample_size": "样本量的大小会影响统计检验的功效，较大的样本量通常能提供更准确的结果。",
            "median": "中位数反映了配对样本差值数据的中心位置，可用于比较配对样本之间的差异。",
            "confidence_interval": "如果置信区间不包含0，说明配对样本之间存在显著差异。"
        },
        'chart_titles': {
            'bar_chart': '均值比较柱状图',
            'error_bar_chart': '误差线图',
            'box_plot': '箱线图',
            'line_chart': '折线图'
        },
        'axis_labels': {
            'mean': '均值',
            'value': '数值',
            'sample1': '样本1',
            'sample2': '样本2'
        }
    },
    'en': {
        'title': "Paired Sample Wilcoxon Test",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Chinese/English",
        'open_excel_button_text': "Example data",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'stats': {
            "w_stat": "W Statistic",
            "z_stat": "Z Statistic",
            "p_value": "p-value",
            "significance": "Significance",
            "effect_size": "Effect Size (r)",
            "sample_size": "Sample Size",
            "median": "Median",
            "confidence_interval": "Confidence Interval of Mean Difference"
        },
        'headers': ["Statistic", "Result"],
        'explanation_title': "Explanation",
        'interpretation_title': "Interpretation",
        'charts_title': "Charts",
        'explanation': {
            "w_stat": "Used to test whether there is a significant difference between paired samples.",
            "sample_size": "The number of pairs of observations in the paired samples.",
            "median": "The middle value of the difference data of the paired samples, dividing the data into two parts.",
            "p_value": "When the p-value is less than the significance level (usually 0.05), the null hypothesis is rejected, indicating a significant difference between paired samples; otherwise, the null hypothesis is accepted, indicating no significant difference.",
            "confidence_interval": "An interval that contains the true mean difference, reflecting the uncertainty of the estimate."
        },
        'interpretation': {
            "w_stat": "The larger the absolute value of the W-statistic, the more significant the difference between paired samples.",
            "p_value": "The basis for determining whether there is a significant difference between paired samples.",
            "sample_size": "The sample size affects the power of the statistical test. A larger sample size usually provides more accurate results.",
            "median": "The median reflects the central position of the difference data of the paired samples and can be used to compare the difference between paired samples.",
            "confidence_interval": "If the confidence interval does not contain 0, it indicates a significant difference between paired samples."
        },
        'chart_titles': {
            'bar_chart': 'Bar Chart of Mean Comparison',
            'error_bar_chart': 'Error Bar Chart',
            'box_plot': 'Box Plot',
            'line_chart': 'Line Chart'
        },
        'axis_labels': {
            'mean': 'Mean',
            'value': 'Value',
            'sample1': 'Sample 1',
            'sample2': 'Sample 2'
        }
    }
}

class PairedSampleWilcoxonTestAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为中文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data26.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"{LANGUAGES[self.current_language]['file_not_found']}：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"{LANGUAGES[self.current_language]['analysis_error'].format(str(e))}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 检查数据是否为数值类型
            numerical_df = df.select_dtypes(include=[np.number])
            if numerical_df.empty:
                raise ValueError(f"{LANGUAGES['zh']['analysis_error'].format('数据中没有数值列，无法进行配对样本Wilcoxon检验。')}"
                                 if self.current_language == 'zh'
                                 else f"{LANGUAGES['en']['analysis_error'].format('There are no numerical columns in the data for paired sample Wilcoxon test.')}")
            if numerical_df.shape[1] != 2:
                raise ValueError(f"{LANGUAGES['zh']['analysis_error'].format('数据必须包含两列数值数据，用于配对样本Wilcoxon检验。')}"
                                 if self.current_language == 'zh'
                                 else f"{LANGUAGES['en']['analysis_error'].format('The data must contain two columns of numerical data for paired sample Wilcoxon test.')}")

            # 执行配对样本Wilcoxon检验
            wilcoxon_result = stats.wilcoxon(numerical_df.iloc[:, 0], numerical_df.iloc[:, 1])

            # 提取W统计量和p值（兼容所有scipy版本）
            w_stat = wilcoxon_result.statistic  # W统计量
            p_value = wilcoxon_result.pvalue  # p值

            # 计算样本量、中位数
            sample_size = numerical_df.count().values[0]
            differences = numerical_df.iloc[:, 0] - numerical_df.iloc[:, 1]
            median = differences.median()

            # 计算自由度
            degrees_of_freedom = sample_size - 1

            # 定义差异函数
            def diff_func(a, b):
                return np.mean(a - b)

            # 准备数据
            data = (numerical_df.iloc[:, 0].values, numerical_df.iloc[:, 1].values)

            # 计算bootstrap置信区间
            boot_ci = bootstrap(data, diff_func, confidence_level=0.95, method='percentile')
            if hasattr(boot_ci.confidence_interval, 'low'):
                confidence_interval = (boot_ci.confidence_interval.low, boot_ci.confidence_interval.high)
            else:
                # 旧版本 scipy 使用 .low 和 .high 可能存放在 .confidence_interval 元组中
                confidence_interval = (boot_ci.confidence_interval[0], boot_ci.confidence_interval[1])

            # 计算效应量r = Z / sqrt(n)
            z_stat = (w_stat - (sample_size * (sample_size + 1) / 4)) / np.sqrt(
                sample_size * (sample_size + 1) * (2 * sample_size + 1) / 24)
            effect_size_r = abs(z_stat) / np.sqrt(sample_size)

            # 显著性标记
            if self.current_language == 'zh':
                p_significance = "p < 0.001" if p_value < 0.001 else "p < 0.01" if p_value < 0.01 else "p < 0.05" if p_value < 0.05 else "p ≥ 0.05"
            else:
                p_significance = "p < 0.001" if p_value < 0.001 else "p < 0.01" if p_value < 0.01 else "p < 0.05" if p_value < 0.05 else "p ≥ 0.05"

            # 更新数据列表，使用语言相关的统计量名称
            ci_str = f"({confidence_interval[0]:.4f}, {confidence_interval[1]:.4f})"
            data = [
                [LANGUAGES[self.current_language]['stats']["w_stat"], w_stat],
                [LANGUAGES[self.current_language]['stats']["z_stat"], z_stat],
                [LANGUAGES[self.current_language]['stats']["p_value"], p_value],
                [LANGUAGES[self.current_language]['stats']["significance"], p_significance],
                [LANGUAGES[self.current_language]['stats']["effect_size"], effect_size_r],
                [LANGUAGES[self.current_language]['stats']["sample_size"], sample_size],
                [LANGUAGES[self.current_language]['stats']["median"], median],
                [LANGUAGES[self.current_language]['stats']["confidence_interval"], ci_str]
            ]
            headers = LANGUAGES[self.current_language]['headers']
            df_results = pd.DataFrame(data, columns=headers)

            # 添加解释说明和解读
            explanations = LANGUAGES[self.current_language]['explanation']
            interpretations = LANGUAGES[self.current_language]['interpretation']

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建一个新的Word文档
                doc = Document()

                # 添加标题
                doc.add_heading(LANGUAGES[self.current_language]["title"], 0)

                # 添加分析结果表格
                table = doc.add_table(rows=1, cols=len(headers))
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                for index, row in df_results.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)

                # 添加解释说明
                doc.add_heading(LANGUAGES[self.current_language]['explanation_title'], level=1)
                explanation_paragraph = doc.add_paragraph()
                for key, value in explanations.items():
                    stat_name = LANGUAGES[self.current_language]['stats'][key]
                    explanation_paragraph.add_run(f"• {stat_name}：{value}\n")

                # 添加结果解读
                doc.add_heading(LANGUAGES[self.current_language]['interpretation_title'], level=1)
                interpretation_paragraph = doc.add_paragraph()
                for key, value in interpretations.items():
                    stat_name = LANGUAGES[self.current_language]['stats'][key]
                    interpretation_paragraph.add_run(f"• {stat_name}：{value}\n")

                # 绘制柱状图
                plt.figure(figsize=(8, 6))
                bars = plt.bar(
                    [LANGUAGES[self.current_language]['axis_labels']['sample1'],
                     LANGUAGES[self.current_language]['axis_labels']['sample2']],
                    [numerical_df.iloc[:, 0].mean(), numerical_df.iloc[:, 1].mean()]
                )
                plt.title(LANGUAGES[self.current_language]['chart_titles']['bar_chart'])
                plt.ylabel(LANGUAGES[self.current_language]['axis_labels']['mean'])
                # 添加显著性标记
                if p_value < 0.05:
                    y_max = max(numerical_df.iloc[:, 0].mean(), numerical_df.iloc[:, 1].mean()) * 1.1
                    plt.plot(
                        [0, 1],  # x轴坐标
                        [y_max * 0.95, y_max * 0.95],  # y轴坐标
                        color='black',  # 黑色线条
                        linestyle='-'  # 实线
                    )
                plt.tight_layout()
                bar_chart_path = save_path.replace('.docx', '_bar_chart.png')
                plt.savefig(bar_chart_path, dpi=300)
                plt.close()

                # 绘制误差线图
                plt.figure(figsize=(8, 6))
                plt.errorbar(
                    x=[LANGUAGES[self.current_language]['axis_labels']['sample1'],
                       LANGUAGES[self.current_language]['axis_labels']['sample2']],
                    y=[numerical_df.iloc[:, 0].mean(), numerical_df.iloc[:, 1].mean()],
                    yerr=[
                        numerical_df.iloc[:, 0].std() / np.sqrt(sample_size),
                        numerical_df.iloc[:, 1].std() / np.sqrt(sample_size)
                    ],
                    fmt='o'  # 确保 fmt 参数正确（标记样式）
                )
                plt.title(LANGUAGES[self.current_language]['chart_titles']['error_bar_chart'])
                plt.ylabel(LANGUAGES[self.current_language]['axis_labels']['mean'])
                error_bar_chart_path = save_path.replace('.docx', '_error_bar_chart.png')
                plt.savefig(error_bar_chart_path)
                plt.close()

                # 绘制箱线图
                plt.figure(figsize=(8, 6))
                numerical_df.plot(kind='box')
                plt.title(LANGUAGES[self.current_language]['chart_titles']['box_plot'])
                plt.ylabel(LANGUAGES[self.current_language]['axis_labels']['value'])
                box_plot_path = save_path.replace('.docx', '_box_plot.png')
                plt.savefig(box_plot_path)
                plt.close()

                # 绘制折线图
                plt.figure(figsize=(8, 6))
                numerical_df.plot(kind='line')
                plt.title(LANGUAGES[self.current_language]['chart_titles']['line_chart'])
                plt.ylabel(LANGUAGES[self.current_language]['axis_labels']['value'])
                line_chart_path = save_path.replace('.docx', '_line_chart.png')
                plt.savefig(line_chart_path)
                plt.close()

                # 在Word文档中插入图表
                doc.add_heading(LANGUAGES[self.current_language]['charts_title'], 1)
                doc.add_picture(bar_chart_path, width=Inches(6))
                doc.add_picture(error_bar_chart_path, width=Inches(6))
                doc.add_picture(box_plot_path, width=Inches(6))
                doc.add_picture(line_chart_path, width=Inches(6))

                # 保存Word文档
                doc.save(save_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(
                    save_path)
                self.result_label.config(text=result_msg, wraplength=400)

            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(LANGUAGES[self.current_language]["title"])
        self.select_button.config(text=LANGUAGES[self.current_language]["select_button"])
        current_entry_text = self.file_entry.get()
        # 保留实际输入的文件路径，仅替换占位符
        if current_entry_text == LANGUAGES['zh']["file_entry_placeholder"] or current_entry_text == LANGUAGES['en']["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=LANGUAGES[self.current_language]["analyze_button"])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]["switch_language"])
        self.open_excel_label.config(text=LANGUAGES[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["select_button"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["analyze_button"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=LANGUAGES[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=LANGUAGES[self.current_language]["switch_language"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = PairedSampleWilcoxonTestAnalysisApp()
    app.run()

if __name__ == "__main__":
    run_app()
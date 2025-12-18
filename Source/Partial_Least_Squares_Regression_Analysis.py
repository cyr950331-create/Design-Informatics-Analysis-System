import tkinter as tk
from tkinter import filedialog
import os
import pandas as pd
import numpy as np
from sklearn.cross_decomposition import PLSRegression
from sklearn.metrics import mean_squared_error, r2_score, mean_absolute_error, median_absolute_error, \
    explained_variance_score
from sklearn.preprocessing import StandardScaler
from sklearn.model_selection import KFold
import matplotlib.pyplot as plt
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from docx import Document
from docx.shared import Inches
from scipy import stats


# 设置支持中文的字体
plt.rcParams["font.family"] = ["Microsoft YaHei", "SimSun", "WenQuanYi Micro Hei", "Heiti TC", "Arial Unicode MS"]
plt.rcParams["axes.unicode_minus"] = False  # 解决负号显示问题

# 定义语言字典，所有键使用英文
LANGUAGES = {
    'zh': {
        'title': "偏最小二乘回归分析",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "文件不存在，请重新选择。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'images_saved': "图片已保存到 {}",
        'switch_language': "中/英",
        'open_excel_button_text': "示例数据",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'regression_equation': "回归方程",
        'regression_coefficients_heading': "回归系数",
        'model_metrics_heading': "模型指标",
        'explanations_heading': "指标说明",
        'components_vs_q2_heading': "Q²与主成分数量关系",
        'actual_vs_predicted_heading': "实际值与预测值对比",
        'residual_plot_heading': "残差图",
        'standardized_coefficients_heading': "标准化回归系数",
        'number_of_components': "主成分数量",
        'q2_score': "交叉验证决定系数",
        'q2_vs_components_title': "Q²与主成分数量关系",
        'optimal_components': "最佳主成分: {}",
        'actual_value': "实际{}值",
        'predicted_value': "预测{}值",
        'actual_vs_predicted_title': "实际值与预测值对比",
        'predicted_values': "预测值",
        'residuals': "残差",
        'residual_plot_title': "残差图",
        'standardized_coefficient_value': "标准化系数值",
        'standardized_coefficients_title': "标准化回归系数",
        'outlier_message': "检测到{}个潜在异常值，已在分析中保留。建议检查数据质量。",
        'explanation': {
            "系数（Coefficients）": "回归系数，表示每个自变量对因变量的影响程度。",
            "标准化系数（Standardized Coefficients）": "标准化回归系数，消除量纲影响后的系数。",
            "截距（Intercept）": "截距，是当所有自变量为 0 时因变量的预测值。",
            "均方误差（MSE）": "均方误差，衡量预测值与真实值之间的平均误差。",
            "平均绝对误差（MAE）": "平均绝对误差，衡量预测值与真实值之间的平均绝对差值。",
            "中位数绝对误差": "中位数绝对误差，对异常值更稳健的误差衡量指标。",
            "解释方差分数": "解释方差分数，衡量模型解释的因变量方差比例。",
            "决定系数（R²）": "决定系数，取值范围在 0 到 1 之间，越接近 1 表示模型拟合效果越好。",
            "交叉验证决定系数（Q²）": "交叉验证决定系数，衡量模型的预测能力，越接近 1 越好。",
            "最佳主成分数": "通过交叉验证确定的最优主成分数量。",
            "随机种子": "用于保证结果可复现的随机数种子。"
        }
    },
    'en': {
        'title': "Partial Least Squares Regression",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'images_saved': "Images have been saved to {}",
        'switch_language': "Chinese/English",
        'open_excel_button_text': "Example data",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'regression_equation': "Regression Equation",
        'regression_coefficients_heading': "Regression Coefficients",
        'model_metrics_heading': "Model Metrics",
        'explanations_heading': "Explanations",
        'components_vs_q2_heading': "Q² vs Number of Components",
        'actual_vs_predicted_heading': "Actual vs Predicted Values",
        'residual_plot_heading': "Residual Plot",
        'standardized_coefficients_heading': "Standardized Regression Coefficients",
        'number_of_components': "Number of Components",
        'q2_score': "Q²",
        'q2_vs_components_title': "Q² vs Number of Components",
        'optimal_components': "Optimal: {}",
        'actual_value': "Actual {}",
        'predicted_value': "Predicted {}",
        'actual_vs_predicted_title': "Actual vs Predicted Values",
        'predicted_values': "Predicted Values",
        'residuals': "Residuals",
        'residual_plot_title': "Residual Plot",
        'standardized_coefficient_value': "Standardized Coefficient Value",
        'standardized_coefficients_title': "Standardized Regression Coefficients",
        'outlier_message': "Detected {} potential outliers, retained in the analysis.建议检查数据质量。It is recommended to check data quality.",
        'explanation': {
            "Coefficients": "Regression coefficients, indicating the influence of each independent variable on the dependent variable.",
            "Standardized Coefficients": "Standardized regression coefficients, eliminating the influence of dimension.",
            "Intercept": "Intercept, which is the predicted value of the dependent variable when all independent variables are 0.",
            "Mean Squared Error (MSE)": "Mean squared error, measuring the average error between the predicted and actual values.",
            "Mean Absolute Error (MAE)": "Mean absolute error, measuring the average absolute difference between predicted and actual values.",
            "Median Absolute Error": "Median absolute error, a more robust error measure for outliers.",
            "Explained Variance Score": "Explained variance score, measuring the proportion of variance explained by the model.",
            "R-squared (R²)": "Coefficient of determination, ranging from 0 to 1. A value closer to 1 indicates a better fit of the model.",
            "Q²": "Cross-validation coefficient of determination, measuring the predictive ability of the model, closer to 1 is better.",
        }
    }
}


class PartialLeastSquaresRegressionAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为中文
        self.current_language = "en"
        # 设置随机种子，确保结果可复现
        self.random_seed = 42
        np.random.seed(self.random_seed)

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data32.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件，第一行为表头
            df = pd.read_excel(file_path)

            # 第一列为因素列（不作为变量），其余列中最后一列为因变量，其他为自变量
            X = df.iloc[:, 1:-1].values  # 自变量（排除第一列和最后一列）
            y = df.iloc[:, -1].values  # 因变量（最后一列）
            var_names = df.columns[1:-1].tolist()  # 自变量名称
            target_name = df.columns[-1]  # 因变量名称

            # 检查并处理异常值
            z_scores = np.abs(stats.zscore(X))
            threshold = 3.0  # 3σ法则
            outlier_mask = (z_scores > threshold).any(axis=1)

            if np.sum(outlier_mask) > 0:
                outlier_count = np.sum(outlier_mask)
                self.result_label.config(
                    text=LANGUAGES[self.current_language]['outlier_message'].format(outlier_count)
                )

            # 对数据进行标准化处理
            scaler = StandardScaler()
            X_scaled = scaler.fit_transform(X)

            # 确定最佳主成分数量（通过改进的交叉验证）
            max_components = min(X.shape[0] - 1, X.shape[1], 20)  # 适当增加最大主成分限制
            best_components = 1
            best_q2 = -np.inf
            q2_scores = []  # 存储各主成分数量对应的Q²值

            # 使用更稳定的交叉验证方式
            kf = KFold(n_splits=5, shuffle=True, random_state=self.random_seed)

            for n in range(1, max_components + 1):
                pls = PLSRegression(n_components=n)
                scores = []
                for train_index, test_index in kf.split(X_scaled):
                    X_train, X_test = X_scaled[train_index], X_scaled[test_index]
                    y_train, y_test = y[train_index], y[test_index]
                    pls.fit(X_train, y_train)
                    y_pred = pls.predict(X_test)
                    scores.append(r2_score(y_test, y_pred))
                q2 = np.mean(scores)
                q2_scores.append(q2)

                if q2 > best_q2:
                    best_q2 = q2
                    best_components = n

            # 使用最佳主成分数量进行PLS回归
            pls = PLSRegression(n_components=best_components)
            pls.fit(X_scaled, y)
            y_pred = pls.predict(X_scaled)

            # 计算模型指标
            coefficients = pls.coef_.flatten()
            intercept = pls.intercept_[0] if hasattr(pls, 'intercept_') else 0
            mse = mean_squared_error(y, y_pred)
            mae = mean_absolute_error(y, y_pred)
            medae = median_absolute_error(y, y_pred)
            evs = explained_variance_score(y, y_pred)
            r2 = r2_score(y, y_pred)
            q2 = best_q2  # 使用交叉验证得到的Q²

            # 标准化系数（已使用标准化数据拟合，直接使用即可）
            standardized_coefficients = coefficients

            # 生成回归方程（使用原始数据的系数）
            pls_raw = PLSRegression(n_components=best_components)
            pls_raw.fit(X, y)
            raw_coefficients = pls_raw.coef_.flatten()
            raw_intercept = pls_raw.intercept_[0] if hasattr(pls_raw, 'intercept_') else 0

            equation_terms = []
            # 添加截距项
            if raw_intercept >= 0:
                equation_terms.append(f"{raw_intercept:.4f}")
            else:
                equation_terms.append(f"- {abs(raw_intercept):.4f}")

            # 添加各个自变量项
            for i, var in enumerate(var_names):
                coef = raw_coefficients[i]
                if i == 0 and raw_intercept == 0:
                    equation_terms.append(f"{coef:.4f} × {var}")
                else:
                    if coef >= 0:
                        equation_terms.append(f"+ {coef:.4f} × {var}")
                    else:
                        equation_terms.append(f"- {abs(coef):.4f} × {var}")

            # 组合成完整方程
            regression_equation = f"{target_name} = " + " ".join(equation_terms)

            # 准备结果数据
            # 1. 变量系数表
            coef_data = []
            for i, var in enumerate(var_names):
                coef_data.append([
                    var,
                    round(raw_coefficients[i], 4),  # 原始系数
                    round(standardized_coefficients[i], 4)  # 标准化系数
                ])

            if self.current_language == 'zh':
                coef_columns = ["变量", "系数", "标准化系数"]
            else:
                coef_columns = ["Variable", "Coefficients", "Standardized Coefficients"]
            coef_df = pd.DataFrame(
                coef_data,
                columns=coef_columns
            )

            # 2. 模型指标表
            if self.current_language == 'zh':
                metrics_columns = ["指标", "数值"]
                metrics_data = [
                    ["截距", round(raw_intercept, 4)],
                    ["均方误差（MSE）", round(mse, 4)],
                    ["平均绝对误差（MAE）", round(mae, 4)],
                    ["中位数绝对误差", round(medae, 4)],
                    ["解释方差分数", round(evs, 4)],
                    ["决定系数（R²）", round(r2, 4)],
                    ["交叉验证决定系数（Q²）", round(q2, 4)],
                    ["最佳主成分数", best_components],
                    ["随机种子", self.random_seed]
                ]
            else:
                metrics_columns = ["Metric", "Value"]
                metrics_data = [
                    ["Intercept", round(raw_intercept, 4)],
                    ["Mean Squared Error (MSE)", round(mse, 4)],
                    ["Mean Absolute Error (MAE)", round(mae, 4)],
                    ["Median Absolute Error", round(medae, 4)],
                    ["Explained Variance Score", round(evs, 4)],
                    ["R-squared (R²)", round(r2, 4)],
                    ["Q² (Cross-validation)", round(q2, 4)],
                    ["Optimal Components", best_components],
                    ["Random Seed", self.random_seed]
                ]
            metrics_df = pd.DataFrame(metrics_data, columns=metrics_columns)

            # 3. 解释说明表
            explanations = LANGUAGES[self.current_language]['explanation']
            expl_data = [[k, v] for k, v in explanations.items()]
            if self.current_language == 'zh':
                expl_df = pd.DataFrame(expl_data, columns=["术语", "说明"])
            else:
                expl_df = pd.DataFrame(expl_data, columns=["Term", "Explanation"])

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建Word文档
                doc = Document()
                doc.add_heading(LANGUAGES[self.current_language]["title"], level=1)

                # 添加回归方程
                doc.add_heading(LANGUAGES[self.current_language]["regression_equation"], level=2)
                doc.add_paragraph(regression_equation)

                # 添加变量系数表
                doc.add_heading(LANGUAGES[self.current_language]["regression_coefficients_heading"], level=2)
                table = doc.add_table(rows=coef_df.shape[0] + 1, cols=coef_df.shape[1])
                for col_idx, col_name in enumerate(coef_df.columns):
                    table.rows[0].cells[col_idx].text = col_name
                for row_idx, row in coef_df.iterrows():
                    for col_idx, value in enumerate(row):
                        table.rows[row_idx + 1].cells[col_idx].text = str(value)

                # 添加模型指标表
                doc.add_heading(LANGUAGES[self.current_language]["model_metrics_heading"], level=2)
                table = doc.add_table(rows=metrics_df.shape[0] + 1, cols=metrics_df.shape[1])
                for col_idx, col_name in enumerate(metrics_df.columns):
                    table.rows[0].cells[col_idx].text = col_name
                for row_idx, row in metrics_df.iterrows():
                    for col_idx, value in enumerate(row):
                        table.rows[row_idx + 1].cells[col_idx].text = str(value)

                # 添加项目符号列表
                doc.add_heading(LANGUAGES[self.current_language]["explanations_heading"], level=2)
                # 创建一个段落作为列表容器
                p = doc.add_paragraph()
                for term, explanation in expl_data:
                    # 每个术语和解释作为一个列表项
                    p.add_run(f"• {term}: {explanation}\n")

                # 获取保存路径的目录
                save_dir = os.path.dirname(save_path)

                # 生成主成分数量与Q²的关系图
                plt.figure(figsize=(10, 6))
                plt.plot(range(1, max_components + 1), q2_scores, 'bo-')
                plt.xlabel(LANGUAGES[self.current_language]["number_of_components"])
                plt.ylabel(LANGUAGES[self.current_language]["q2_score"])
                plt.title(LANGUAGES[self.current_language]["q2_vs_components_title"])
                plt.axvline(
                    x=best_components,
                    color='r',
                    linestyle='--',
                    label=LANGUAGES[self.current_language]["optimal_components"].format(best_components)
                )
                plt.legend()
                components_img_path = os.path.join(save_dir, "components_vs_q2.png")
                plt.savefig(components_img_path, bbox_inches='tight')
                plt.close()

                # 在Word文档中添加该图
                doc.add_heading(LANGUAGES[self.current_language]["components_vs_q2_heading"], level=2)
                doc.add_picture(components_img_path, width=Inches(6))

                # 生成散点图（使用当前语言的标签）
                plt.figure(figsize=(10, 6))
                plt.scatter(y, y_pred, alpha=0.6)
                plt.plot([y.min(), y.max()], [y.min(), y.max()], 'r--', lw=2)
                plt.xlabel(LANGUAGES[self.current_language]["actual_value"].format(target_name))
                plt.ylabel(LANGUAGES[self.current_language]["predicted_value"].format(target_name))
                plt.title(LANGUAGES[self.current_language]["actual_vs_predicted_title"])
                img_name = "pls_regression_scatter.png"
                img_path = os.path.join(save_dir, img_name)
                plt.savefig(img_path, bbox_inches='tight')
                plt.close()

                # 将图片添加到Word文档中
                doc.add_heading(LANGUAGES[self.current_language]["actual_vs_predicted_heading"], level=2)
                doc.add_picture(img_path, width=Inches(6))

                # 残差图
                residuals = y - y_pred.flatten()
                plt.figure(figsize=(10, 6))
                plt.scatter(y_pred, residuals, alpha=0.6)
                plt.axhline(y=0, color='r', linestyle='--')
                plt.xlabel(LANGUAGES[self.current_language]["predicted_values"])
                plt.ylabel(LANGUAGES[self.current_language]["residuals"])
                plt.title(LANGUAGES[self.current_language]["residual_plot_title"])
                residual_img_path = os.path.join(save_dir, "residual_plot.png")
                plt.savefig(residual_img_path, bbox_inches='tight')
                plt.close()

                # 在Word文档中添加残差图
                doc.add_heading(LANGUAGES[self.current_language]["residual_plot_heading"], level=2)
                doc.add_picture(residual_img_path, width=Inches(6))

                # 系数条形图
                plt.figure(figsize=(10, 6))
                indices = np.arange(len(standardized_coefficients))
                plt.bar(indices, standardized_coefficients)
                plt.xticks(indices, var_names, rotation=0, ha='center')
                plt.ylabel(LANGUAGES[self.current_language]["standardized_coefficient_value"])
                plt.title(LANGUAGES[self.current_language]["standardized_coefficients_title"])
                plt.tight_layout()
                coef_img_path = os.path.join(save_dir, "coefficients_plot.png")
                plt.savefig(coef_img_path, bbox_inches='tight')
                plt.close()

                # 在Word文档中添加系数图
                doc.add_heading(LANGUAGES[self.current_language]["standardized_coefficients_heading"], level=2)
                doc.add_picture(coef_img_path, width=Inches(6))

                # 保存Word文档
                doc.save(save_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(save_path)
                result_msg += LANGUAGES[self.current_language]['images_saved'].format(save_dir)
                self.result_label.config(text=result_msg, wraplength=400)
            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))
            # 添加错误日志记录
            import traceback
            with open("error_log.txt", "a") as f:
                f.write(f"Error at {pd.Timestamp.now()}: {traceback.format_exc()}\n")

    def switch_language(self, event=None):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(LANGUAGES[self.current_language]["title"])
        self.select_button.config(text=LANGUAGES[self.current_language]["select_button"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=LANGUAGES[self.current_language]["analyze_button"])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]["switch_language"])
        self.open_excel_label.config(text=LANGUAGES[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["select_button"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["analyze_button"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=LANGUAGES[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=LANGUAGES[self.current_language]["switch_language"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = PartialLeastSquaresRegressionAnalysisApp()
    app.run()

if __name__ == "__main__":
    run_app()
import tkinter as tk
from tkinter import filedialog
import openpyxl
import os
import pandas as pd
from scipy.stats import bartlett
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches

# 定义语言字典
languages = {
    "zh": {
        "title": "Bartlett 检验",
        "select_button_text": "选择文件",
        "file_entry_placeholder": "请输入待分析 Excel 文件的完整路径",
        "analyze_button_text": "分析文件",
        "no_file_selected": "请选择文件。",
        "file_not_exists": "文件不存在，请重新选择。",
        "analysis_error": "分析文件时出错: {}",
        "analysis_complete": "分析完成，结果已保存到 {}，箱线图已保存。",
        "no_save_path_selected": "未选择保存路径，结果未保存。",
        "columns_stats": ["Bartlett 统计量", "Bartlett P 值", "结果解读"],
        "interpretation_accept": "在 0.05 的显著性水平下，不能拒绝原假设，各样本方差具有齐性。",
        "interpretation_reject": "在 0.05 的显著性水平下，拒绝原假设，各样本方差不具有齐性。",
        'open_excel_button_text': "示例数据",
        "switch_language_button_text": "中/英",
        "boxplot_title": "Bartlett 检验箱线图",
        "boxplot_xlabel": "变量",
        "boxplot_ylabel": "数值",
        "word_heading": "Bartlett 检验结果",
        "overall": "总体",
        "table_header_column": "列名"
    },
    "en": {
        "title": "Bartlett Test",
        "select_button_text": "Select File",
        "file_entry_placeholder": "Please enter the full path of the Excel file to be analyzed",
        "analyze_button_text": "Analyze File",
        "no_file_selected": "Please select a valid file path.",
        "file_not_exists": "The file does not exist. Please select again.",
        "analysis_error": "An error occurred while analyzing the file: {}",
        "analysis_complete": "Analysis completed. The results have been saved to {}, and the box plots have been saved.",
        "no_save_path_selected": "No save path selected. The results were not saved.",
        "columns_stats": ["Bartlett Statistic", "Bartlett P-value", "Result Interpretation"],
        "interpretation_accept": "At the 0.05 significance level, the null hypothesis cannot be rejected. The variances of the samples are homogeneous.",
        "interpretation_reject": "At the 0.05 significance level, the null hypothesis is rejected. The variances of the samples are not homogeneous.",
        'open_excel_button_text': "Example data",
        "switch_language_button_text": "Chinese/English",
        "boxplot_title": "Bartlett Test Boxplot",
        "boxplot_xlabel": "Variables",
        "boxplot_ylabel": "Values",
        "overall": "Overall",
        "word_heading": "Bartlett Test Results",
        "table_header_column": "Column Name"

    }
}

class BartlettTestApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"
        
        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])
            
        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data2.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)
            df = df.select_dtypes(include=[float, int]).dropna()

            if df.shape[1] < 2:
                raise ValueError("文件中至少需要两列数值型数据才能进行 Bartlett 检验。")

            data = []
            columns_stats = languages[self.current_language]["columns_stats"]

            column_names = df.columns.tolist()
            column_data = [df[col].dropna().tolist() for col in column_names]

            if len(column_data) > 1:
                # 进行 Bartlett 检验
                bartlett_statistic, bartlett_p_value = bartlett(*column_data)

                # 根据 P 值进行结果解读
                if bartlett_p_value > 0.05:
                    interpretation = languages[self.current_language]["interpretation_accept"]
                else:
                    interpretation = languages[self.current_language]["interpretation_reject"]

                values = [bartlett_statistic, bartlett_p_value, interpretation]
                data.append([languages[self.current_language]["overall"]] + values)

                # 在绘图前设置中文字体，防止乱码
                plt.rcParams['font.sans-serif'] = ['SimHei']  # 支持中文
                plt.rcParams['axes.unicode_minus'] = False  # 正常显示负号

                # 绘制箱线图（优化版）
                plt.figure(figsize=(8, 6))  # 设置图形尺寸，防止标签挤在一起
                plt.boxplot(column_data, tick_labels=column_names, patch_artist=True)  # 用 labels 替代 tick_labels
                plt.title(languages[self.current_language]["boxplot_title"])
                plt.xlabel(languages[self.current_language]["boxplot_xlabel"])
                plt.ylabel(languages[self.current_language]["boxplot_ylabel"])

                # 旋转横轴标签，防止重叠
                plt.xticks(rotation=45, ha='right')

                # 调整边距，防止标题或标签被裁剪
                plt.tight_layout()

                # 保存图片
                boxplot_path = os.path.splitext(file_path)[0] + '_boxplot.png'
                plt.savefig(boxplot_path, dpi=300)
                plt.close()

            headers = [languages[self.current_language]["table_header_column"]] + columns_stats
            df = pd.DataFrame(data, columns=headers)

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加标题
                doc.add_heading(languages[self.current_language]["word_heading"], 0)

                # 添加表格
                table = doc.add_table(rows=df.shape[0]+1, cols=df.shape[1])
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(headers):
                    hdr_cells[col_idx].text = header

                for row_idx in range(df.shape[0]):
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx in range(df.shape[1]):
                        row_cells[col_idx].text = str(df.iloc[row_idx, col_idx])

                # 添加箱线图
                doc.add_picture(boxplot_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                # 设置 wraplength 属性让文本自动换行
                self.result_label.config(text=languages[self.current_language]["analysis_complete"].format(save_path), wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]["no_save_path_selected"])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]["analysis_error"].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])
        self.open_excel_label.config(text=languages[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"], 
                                       command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"], 
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=languages[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=languages[self.current_language]["switch_language_button_text"],
                                              foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)
        
    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = BartlettTestApp()
    app.run()

if __name__ == "__main__":
    run_app()
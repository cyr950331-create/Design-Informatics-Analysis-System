import tkinter as tk
from tkinter import filedialog
import os
import pandas as pd
import numpy as np
from scipy.stats import chi2
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from docx import Document

# 定义语言字典
languages = {
    'zh': {
        'title': "Kendall 协调系数",
        'select_button_text': "选择文件",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'analyze_button_text': "分析文件",
        'no_file_selected': "请选择文件。",
        'file_not_exists': "文件不存在，请重新选择。",
        'analysis_error': "分析文件时出错: {}",
        'analysis_success': "分析完成，结果已保存到 {}",
        'no_save_path': "未选择保存路径，结果未保存。",
        'open_excel_button_text': "示例数据",
        'switch_language_button_text': "中/英",
        'explanation': {
            "Kendall协和系数": "用于衡量多个评价者对多个项目的排序一致性。",
            "评价者数量": "参与评价的人数。",
            "项目数量": "被评价的项目总数。",
            "总样本量": "评价数据的总观测数（评价者数量 × 项目数量）。",
            "中位数（按项目）": "每个项目评分的中位数。"
        },
        'interpretation': {
            "统计量": "Kendall协和系数的值，范围从 0 到 1，越接近 1 表示一致性越高。",
            "p值": "p值小于显著性水平（通常为0.05）时，拒绝原假设，认为评价者之间存在显著一致性；否则，接受原假设，认为评价者之间无显著一致性。",
            "评价者数量": "参与评价的人数，影响一致性分析的可靠性。",
            "项目数量": "被评价的项目总数，数量越多结果越可靠。",
            "总样本量": "总观测数 = 评价者数量 × 项目数量，样本量越大统计检验功效越高。",
            "中位数（按项目）": "每个项目评分的中间值，反映该项目的整体评价水平。"
        }
    },
    'en': {
        'title': "Kendall's Coordination Coefficient",
        'select_button_text': "Select File",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'analyze_button_text': "Analyze File",
        'no_file_selected': "Please select a valid file path.",
        'file_not_exists': "The file does not exist. Please select again.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'analysis_success': "Analysis completed. The results have been saved to {}",
        'no_save_path': "No save path selected. The results were not saved.",
        'open_excel_button_text': "Example data",
        'switch_language_button_text': "Chinese/English",
        'explanation': {
            "Kendall's W Coefficient": "Used to measure the consistency of rankings of multiple items by multiple raters.",
            "Number of Raters": "The number of raters participating in the evaluation.",
            "Number of Items": "The total number of items being evaluated.",
            "Total Sample Size": "Total number of observations (number of raters × number of items).",
            "Median (by Item)": "Median score for each item."
        },
        'interpretation': {
            "Statistic": "The value of Kendall's Coordination Coefficient, ranging from 0 to 1. A value closer to 1 indicates higher consistency.",
            "p-value": "When the p-value is less than the significance level (usually 0.05), the null hypothesis is rejected, indicating significant consistency among raters; otherwise, the null hypothesis is accepted.",
            "Number of Raters": "The number of raters affects the reliability of the consistency analysis.",
            "Number of Items": "The total number of items being evaluated; more items lead to more reliable results.",
            "Total Sample Size": "Total observations = number of raters × number of items. Larger sample size increases statistical test power.",
            "Median (by Item)": "The median score for each item, reflecting the overall evaluation level of the item."
        }
    }
}


class KendallsCoordinationCoefficientApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data39.xls")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 检查数据是否为数值类型（仅项目相关数据）
            # 假设第一列是评委信息，从第二列开始是项目数据
            if df.shape[1] < 2:
                raise ValueError("数据格式不正确，至少需要包含1列评委信息和1列项目数据")

            # 提取项目相关的数值列（排除第一列评委信息）
            numerical_df = df.iloc[:, 1:].select_dtypes(include=[np.number])
            if numerical_df.empty:
                raise ValueError("项目数据中没有数值列，无法进行Kendall协和系数分析。")

            # 进行Kendall协和系数分析
            num_raters = numerical_df.shape[0]  # 评价者数量（行）
            num_items = numerical_df.shape[1]  # 项目数量（排除评委列后的数值列）

            # 计算排序
            ranks = numerical_df.rank(axis=1)

            # 计算S值（各项目秩和与平均秩和的偏差平方和）
            mean_rank_sum = num_raters * (num_items + 1) / 2
            s = ((ranks.sum(axis=0) - mean_rank_sum) ** 2).sum()

            # 计算Kendall协和系数W
            w = (12 * s) / (num_raters ** 2 * (num_items ** 3 - num_items))

            # 计算显著性p值（基于卡方分布）
            chi2_stat = w * num_raters * (num_items - 1)
            df_chi2 = num_items - 1  # 自由度
            p_value = 1 - chi2.cdf(chi2_stat, df_chi2)  # 计算p值

            # 计算总样本量和各项目中位数
            total_sample_size = num_raters * num_items  # 总样本量
            item_medians = numerical_df.median()  # 各项目中位数

            # 整理数据
            if self.current_language == "zh":
                stat_names = {
                    "kendall_w": "Kendall协和系数 (W)",
                    "num_raters": "评价者数量",
                    "num_items": "项目数量",
                    "total_sample": "总样本量",
                    "median_by_item": "中位数（按项目）"
                }
                headers = ["统计量", "统计量值", "p值"]
            else:
                stat_names = {
                    "kendall_w": "Kendall's W Coefficient",
                    "num_raters": "Number of Raters",
                    "num_items": "Number of Items",
                    "total_sample": "Total Sample Size",
                    "median_by_item": "Median (by Item)"
                }
                headers = ["Statistic", "Statistic Value", "p-value"]

            # 跟随语言切换的数据列表
            data = [
                [stat_names["kendall_w"], round(w, 4), round(p_value, 6)],
                [stat_names["num_raters"], num_raters, ""],
                [stat_names["num_items"], num_items, ""],
                [stat_names["total_sample"], total_sample_size, ""],
                [stat_names["median_by_item"], "", ""]  # 占位行，后续单独处理
            ]

            df_result = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            explanations = languages[self.current_language]['explanation']
            # 根据当前语言获取对应的列名列表
            if self.current_language == "zh":
                explanation_columns = ["Kendall协和系数", "评价者数量", "项目数量", "总样本量", "中位数（按项目）"]
            else:
                explanation_columns = ["Kendall's W Coefficient", "Number of Raters", "Number of Items",
                                       "Total Sample Size", "Median (by Item)"]
            # 使用动态列名重新索引
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(columns=explanation_columns)
            explanation_df.insert(0, "统计量_解释说明", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([languages[self.current_language]['interpretation']])

            # 根据当前语言设置解读表格的列名
            if self.current_language == "zh":
                interpretation_columns = ["统计量", "p值", "评价者数量", "项目数量", "总样本量", "中位数（按项目）"]
            else:
                interpretation_columns = ["Statistic", "p-value", "Number of Raters", "Number of Items",
                                          "Total Sample Size", "Median (by Item)"]

            # 重新索引并插入表头
            interpretation_df = interpretation_df.reindex(columns=interpretation_columns)
            interpretation_df.insert(
                0,
                "统计量_结果解读" if self.current_language == 'zh' else "Statistic_Interpretation",
                "结果解读" if self.current_language == 'zh' else "Interpretation"
            )

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建一个新的 Word 文档
                doc = Document()

                # 添加标题
                # 根据当前语言设置标题
                if self.current_language == "zh":
                    heading_text = "Kendall协和系数分析结果"
                else:
                    heading_text = "Kendall's W Coefficient Analysis Results"

                doc.add_heading(heading_text, 0)

                # 添加分析结果表格
                table = doc.add_table(rows=1, cols=len(headers))
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                for idx, row in df_result.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)
                    # 单独处理中位数行，添加具体项目的中位数
                    if idx == 4:  # 中位数行
                        for item, median in item_medians.items():
                            med_row = table.add_row().cells
                            med_row[0].text = f"  - {item}"
                            med_row[1].text = f"{median:.4f}"

                # 添加解释说明（替换原解释说明表格部分）
                # 根据当前语言设置统计量解释说明的标题
                if self.current_language == "zh":
                    explanation_heading = "统计量解释说明"
                else:
                    explanation_heading = "Statistic Explanation"

                doc.add_heading(explanation_heading, 1)
                explanation_items = languages[self.current_language]['explanation']
                for key, value in explanation_items.items():
                    doc.add_paragraph(f"- {key}: {value}", style='ListBullet')

                # 添加结果解读（替换原结果解读表格部分）
                # 根据当前语言设置统计量结果解读的标题
                if self.current_language == "zh":
                    interpretation_heading = "统计量结果解读"
                else:
                    interpretation_heading = "Statistic Interpretation"

                doc.add_heading(interpretation_heading, 1)
                interpretation_items = languages[self.current_language]['interpretation']
                for key, value in interpretation_items.items():
                    doc.add_paragraph(f"- {key}: {value}", style='ListBullet')

                # 保存 Word 文档
                doc.save(save_path)

                result_msg = languages[self.current_language]['analysis_success'].format(save_path)
                self.result_label.config(text=result_msg, wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])
        self.open_excel_label.config(text=languages[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=languages[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame,
                                               text=languages[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = KendallsCoordinationCoefficientApp()
    app.run()


if __name__ == "__main__":
    run_app()
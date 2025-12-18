import tkinter as tk
from tkinter import filedialog
import os
import pandas as pd
import numpy as np
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
import tkinter.simpledialog
import pingouin as pg
from docx import Document
from docx.shared import Inches
from scipy import stats

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'  # 设置字体为黑体，可根据系统情况修改为其他支持中文的字体
plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

# 定义语言字典
languages = {
    "zh": {
        "title": "组内评分者信度",
        "select_button_text": "选择文件",
        "file_entry_placeholder": "请输入待分析 Excel 文件的完整路径",
        "analyze_button_text": "分析文件",
        "no_file_selected": "请选择文件。",
        "file_not_exists": "文件不存在，请重新选择。",
        "analysis_error": "分析文件时出错: {}",
        "analysis_complete": "分析完成，结果已保存到 {}，相关图片已保存。",
        "no_save_path_selected": "未选择保存路径，结果未保存。",
        "switch_language_button_text": "中/英",
        'open_excel_button_text': "示例数据",
        "group_column_prompt": "请输入分组列的列名",
        "rating_column_prompt": "请输入评分列的列名（点击取消结束输入）",
        "no_group_column": "未输入有效的分组列名，分析取消。",
        "no_rating_columns": "未输入有效的评分列名，分析取消。",
        "empty_rating_column": "输入的列名不能为空，请重新输入。",
        "group_less_than_2": "组 {} 的样本数少于2，跳过...",
        "data_transformation_check": "数据转换验证: 原始数据{}行{}列，转换后数据{}行{}列，验证{}",
        "anova_assumption_normality": "ANOVA正态性检验（Shapiro-Wilk）: p值={:.4f}，{}",
        "anova_assumption_homoscedasticity": "ANOVA方差齐性检验（Levene）: p值={:.4f}，{}",
        "assumption_met": "满足前提条件",
        "assumption_not_met": "不满足前提条件",
        "using_welch_anova": "组间方差不齐，使用Welch ANOVA",
        "using_kruskal_wallis": "数据不满足正态性，使用Kruskal-Wallis检验",
        "explanation": {
            "rwg值": "组内评分者信度rwg用于评估组内成员评分的一致性，值越接近1表示一致性越高。",
            "Rwg值标准差SD": "Rwg值的标准差，反映了Rwg值的离散程度。",
            "P25": "Rwg值的第25百分位数。",
            "中位数": "Rwg值的中位数。",
            "P75": "Rwg值的第75百分位数。",
            "ICC1": "组内相关系数1，用于衡量组内评分者之间的一致性。",
            "ICC2": "组内相关系数2，考虑了评分者和项目的交互作用。",
            "MSB": "组间均方，反映了组间差异。",
            "MSW": "组内均方，反映了组内差异。",
            "F值": "F检验统计量，用于检验组间差异是否显著。",
            "p值": "F检验的p值，用于判断组间差异是否显著。",
            "Welch_F值": "Welch F检验统计量，用于方差不齐时的组间差异检验。",
            "Welch_p值": "Welch F检验的p值。",
            "Kruskal-Wallis_H值": "Kruskal-Wallis检验统计量，用于非参数组间差异检验。",
            "Kruskal-Wallis_p值": "Kruskal-Wallis检验的p值。"
        },
        "interpretation": {
            "rwg值": "rwg值越接近1，说明组内成员的评分越一致；值越低，说明组内成员的评分差异越大。",
            "Rwg值标准差SD": "标准差越大，说明Rwg值的离散程度越大。",
            "P25": "第25百分位数较低表示有25%的Rwg值低于该值。",
            "中位数": "中位数反映了Rwg值的中间水平。",
            "P75": "第75百分位数较高表示有75%的Rwg值低于该值。",
            "ICC1": "ICC1值越接近1，组内评分者之间的一致性越高。",
            "ICC2": "ICC2值越接近1，考虑交互作用后组内评分者之间的一致性越高。",
            "MSB": "MSB值越大，组间差异越明显。",
            "MSW": "MSW值越大，组内差异越明显。",
            "F值": "F值越大，说明组间差异越可能显著。",
            "p值": "p值小于0.05时，说明组间差异显著。",
            "Welch_F值": "Welch F值越大，说明组间差异越可能显著（方差不齐时）。",
            "Welch_p值": "Welch p值小于0.05时，说明组间差异显著（方差不齐时）。",
            "Kruskal-Wallis_H值": "Kruskal-Wallis H值越大，说明组间差异越可能显著（非参数检验）。",
            "Kruskal-Wallis_p值": "Kruskal-Wallis p值小于0.05时，说明组间差异显著（非参数检验）。"
        }
    },
    "en": {
        "title": "Within-Group Inter-Rater Reliability",
        "select_button_text": "Select Files",
        "file_entry_placeholder": "Please enter the full path of the Excel file to be analyzed",
        "analyze_button_text": "Analyze Files",
        "no_file_selected": "Please select a valid file path.",
        "file_not_exists": "The file does not exist. Please select again.",
        "analysis_error": "An error occurred while analyzing the file: {}",
        "analysis_complete": "Analysis completed. The results have been saved to {}, and the relevant images have been saved.",
        "no_save_path_selected": "No save path selected. The results were not saved.",
        "switch_language_button_text": "Chinese/English",
        'open_excel_button_text': "Example data",
        "group_column_prompt": "Please enter the name of the grouping column",
        "rating_column_prompt": "Please enter the name of the rating column (click Cancel to finish input)",
        "no_group_column": "No valid grouping column name entered. Analysis canceled.",
        "no_rating_columns": "No valid rating column names entered. Analysis canceled.",
        "empty_rating_column": "The column name entered cannot be empty. Please re-enter.",
        "group_less_than_2": "Group {} has less than 2 samples. Skipping...",
        "data_transformation_check": "Data transformation check: Original data {} rows {} cols, Transformed data {} rows {} cols, Validation {}",
        "anova_assumption_normality": "ANOVA normality test (Shapiro-Wilk): p-value={:.4f}, {}",
        "anova_assumption_homoscedasticity": "ANOVA homoscedasticity test (Levene): p-value={:.4f}, {}",
        "assumption_met": "assumption met",
        "assumption_not_met": "assumption not met",
        "using_welch_anova": "Unequal variances detected, using Welch ANOVA",
        "using_kruskal_wallis": "Normality assumption not met, using Kruskal-Wallis test",
        "explanation": {
            "rwg Value": "The within-group inter-rater reliability rwg is used to evaluate the consistency of ratings within a group. A value closer to 1 indicates higher consistency.",
            "Rwg SD": "The standard deviation of the rwg values, reflecting the dispersion of the rwg values.",
            "P25": "The 25th percentile of the rwg values.",
            "Median": "The median of the rwg values.",
            "P75": "The 75th percentile of the rwg values.",
            "ICC1": "Intraclass correlation coefficient 1, used to measure the consistency between raters within a group.",
            "ICC2": "Intraclass correlation coefficient 2, considering the interaction between raters and items.",
            "MSB": "Mean square between groups, reflecting the differences between groups.",
            "MSW": "Mean square within groups, reflecting the differences within groups.",
            "F Value": "F-test statistic, used to test whether the differences between groups are significant.",
            "p Value": "The p-value of the F-test, used to determine whether the differences between groups are significant.",
            "Welch F Value": "Welch F-test statistic, used for testing group differences when variances are unequal.",
            "Welch p Value": "The p-value of the Welch F-test.",
            "Kruskal-Wallis H Value": "Kruskal-Wallis test statistic, used for non-parametric testing of group differences.",
            "Kruskal-Wallis p Value": "The p-value of the Kruskal-Wallis test."
        },
        "interpretation": {
            "rwg Value": "The closer the rwg value is to 1, the more consistent the ratings within the group; the lower the value, the greater the difference in ratings within the group.",
            "Rwg SD": "A larger standard deviation indicates a greater dispersion of the rwg values.",
            "P25": "A lower 25th percentile means that 25% of the rwg values are below this value.",
            "Median": "The median reflects the middle level of the rwg values.",
            "P75": "A higher 75th percentile means that 75% of the rwg values are below this value.",
            "ICC1": "The closer the ICC1 value is to 1, the higher the consistency between raters within the group.",
            "ICC2": "The closer the ICC2 value is to 1, the higher the consistency between raters within the group considering the interaction.",
            "MSB": "A larger MSB value indicates more obvious differences between groups.",
            "MSW": "A larger MSW value indicates more obvious differences within groups.",
            "F Value": "A larger F value indicates that the differences between groups are more likely to be significant.",
            "p Value": "When the p-value is less than 0.05, the differences between groups are significant.",
            "Welch F Value": "A larger Welch F value indicates that the differences between groups are more likely to be significant (when variances are unequal).",
            "Welch p Value": "When the Welch p-value is less than 0.05, the differences between groups are significant (when variances are unequal).",
            "Kruskal-Wallis H Value": "A larger Kruskal-Wallis H value indicates that the differences between groups are more likely to be significant (non-parametric test).",
            "Kruskal-Wallis p Value": "When the Kruskal-Wallis p-value is less than 0.05, the differences between groups are significant (non-parametric test)."
        }
    }
}


class WithinGroupInterRaterReliabilityRwgAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data17.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def rwg_analysis(self, data, group_column, rating_columns):
        all_results = []
        rwg_values = []
        groups = data[group_column].unique()

        for group in groups:
            group_data = data[data[group_column] == group][rating_columns]
            # 检查数据是否为空或者只有一个样本
            if group_data.shape[0] < 2:
                print(languages[self.current_language]["group_less_than_2"].format(group))
                continue
            k = group_data.shape[1]  # 评分者数量
            n = group_data.shape[0]  # 项目数量
            var_within = group_data.var(axis=1).mean()
            expected_var = (k ** 2 - 1) / 12
            rwg = 1 - (var_within / expected_var)
            rwg_values.append(rwg)
            if self.current_language == "zh":
                rwg_suffix = "_rwg值"
            else:
                rwg_suffix = "_rwg Value"

            # 跟随语言切换定义结果字典键名
            result = {
                f"{group}{rwg_suffix}": rwg
            }
            all_results.append(result)

        # 计算 Rwg 值的统计量
        if rwg_values:
            rwg_sd = np.std(rwg_values)
            rwg_p25 = np.percentile(rwg_values, 25)
            rwg_median = np.median(rwg_values)
            rwg_p75 = np.percentile(rwg_values, 75)

            # 计算 ICC1 和 ICC2
            icc_data = pd.melt(data, id_vars=[group_column], value_vars=rating_columns)
            icc_data.columns = ['Group', 'Rater', 'Score']

            # 验证数据转换是否正确（宽→长格式）
            original_rows = data.shape[0]
            original_cols = len(rating_columns)
            transformed_rows = icc_data.shape[0]
            transformed_cols = icc_data.shape[1]
            validation = "成功" if (original_rows * original_cols == transformed_rows) else "失败"
            print(languages[self.current_language]["data_transformation_check"].format(
                original_rows, original_cols, transformed_rows, transformed_cols, validation))

            icc = pg.intraclass_corr(data=icc_data, targets='Group', raters='Rater', ratings='Score')
            icc1 = icc[icc['Type'] == 'ICC1']['ICC'].values[0]
            icc2 = icc[icc['Type'] == 'ICC2']['ICC'].values[0]

            try:
                icc = pg.intraclass_corr(data=icc_data, targets='Group', raters='Rater', ratings='Score')
                # 验证ICC结果是否有效
                if len(icc[icc['Type'] == 'ICC1']) == 0 or len(icc[icc['Type'] == 'ICC2']) == 0:
                    raise ValueError("ICC计算结果不完整")

                icc1 = icc[icc['Type'] == 'ICC1']['ICC'].values[0]
                icc2 = icc[icc['Type'] == 'ICC2']['ICC'].values[0]

                # 添加ICC的置信区间
                icc1_ci = icc[icc['Type'] == 'ICC1'][['CI95%_L', 'CI95%_U']].values[0]
                icc2_ci = icc[icc['Type'] == 'ICC2'][['CI95%_L', 'CI95%_U']].values[0]

            except Exception as e:
                print(f"ICC计算失败: {str(e)}")
                icc1 = np.nan
                icc2 = np.nan
                icc1_ci = [np.nan, np.nan]
                icc2_ci = [np.nan, np.nan]

            # 检验ANOVA前提条件
            # 正态性检验（Shapiro-Wilk）
            stat, p_norm = stats.shapiro(icc_data['Score'])
            normality_met = p_norm > 0.05
            print(languages[self.current_language]["anova_assumption_normality"].format(
                p_norm, languages[self.current_language]["assumption_met"] if normality_met else
                languages[self.current_language]["assumption_not_met"]))

            # 方差齐性检验（Levene）
            groups = [icc_data[icc_data['Group'] == g]['Score'] for g in icc_data['Group'].unique()]
            stat, p_hom = stats.levene(*groups)
            homoscedasticity_met = p_hom > 0.05
            print(languages[self.current_language]["anova_assumption_homoscedasticity"].format(
                p_hom, languages[self.current_language]["assumption_met"] if homoscedasticity_met else
                languages[self.current_language]["assumption_not_met"]))

            # 根据前提条件选择合适的检验方法
            additional_stats = {}
            if not normality_met:
                # 非参数检验：Kruskal-Wallis
                print(languages[self.current_language]["using_kruskal_wallis"])
                stat, p = stats.kruskal(*groups)
                if self.current_language == "zh":
                    h_key = "Kruskal-Wallis_H值"
                    p_key = "Kruskal-Wallis_p值"
                else:
                    h_key = "Kruskal-Wallis H Value"
                    p_key = "Kruskal-Wallis p Value"
                additional_stats[h_key] = stat
                additional_stats[p_key] = p
                # 无法计算MSB和MSW，设置为NaN
                additional_stats["MSB"] = np.nan
                additional_stats["MSW"] = np.nan
            else:
                if homoscedasticity_met:
                    # 标准ANOVA
                    anova = pg.anova(data=icc_data, dv='Score', between='Group')
                    msb = anova.loc[anova['Source'] == 'Group', 'MS'].values[0]  # 组间均方
                    msw = anova.loc[anova['Source'] == 'Residual', 'MS'].values[0]  # 残差（组内）均方
                    f_value = anova['F'][0]
                    p_value = anova['p-unc'][0]
                    additional_stats["MSB"] = msb
                    additional_stats["MSW"] = msw
                    if self.current_language == "zh":
                        f_key = "F值"
                        p_key = "p值"
                    else:
                        f_key = "F Value"
                        p_key = "p Value"

                    # 跟随语言切换更新统计量字典
                    additional_stats[f_key] = f_value
                    additional_stats[p_key] = p_value
                else:
                    # 方差不齐，使用Welch ANOVA
                    print(languages[self.current_language]["using_welch_anova"])
                    welch_anova = pg.welch_anova(data=icc_data, dv='Score', between='Group')
                    additional_stats["MSB"] = np.nan  # Welch ANOVA不提供传统MSB
                    additional_stats["MSW"] = np.nan  # Welch ANOVA不提供传统MSW
                    # 根据当前语言设置统计量键名
                    if self.current_language == "zh":
                        welch_f_key = "Welch_F值"
                        welch_p_key = "Welch_p值"
                    else:
                        welch_f_key = "Welch F Value"
                        welch_p_key = "Welch p Value"

                    # 跟随语言切换更新统计量字典
                    additional_stats[welch_f_key] = welch_anova['F'][0]
                    additional_stats[welch_p_key] = welch_anova['p-unc'][0]

            # 添加通用统计量
            # 根据当前语言定义统计量键名
            if self.current_language == "zh":
                stats_keys = {
                    "rwg_sd": "Rwg值标准差SD",
                    "rwg_p25": "P25",
                    "rwg_median": "中位数",
                    "rwg_p75": "P75",
                    "icc1": "ICC1",
                    "icc2": "ICC2"
                }
            else:
                stats_keys = {
                    "rwg_sd": "Rwg SD",
                    "rwg_p25": "P25",
                    "rwg_median": "Median",
                    "rwg_p75": "P75",
                    "icc1": "ICC1",
                    "icc2": "ICC2"
                }

            # 更新统计量字典（键名随语言切换）
            additional_stats.update({
                stats_keys["rwg_sd"]: rwg_sd,
                stats_keys["rwg_p25"]: rwg_p25,
                stats_keys["rwg_median"]: rwg_median,
                stats_keys["rwg_p75"]: rwg_p75,
                stats_keys["icc1"]: icc1,
                stats_keys["icc2"]: icc2
            })

            all_results.append(additional_stats)

        return all_results, rwg_values

    def analyze_file(self):
        file_paths = self.file_entry.get().split(", ")
        if file_paths[0] == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return

        for file_path in file_paths:
            if not os.path.exists(file_path):
                self.result_label.config(text=languages[self.current_language]["file_not_exists"])
                return

        try:
            # 获取分组列名
            group_column = tkinter.simpledialog.askstring("输入信息",
                                                          languages[self.current_language]["group_column_prompt"],
            parent=self.root)
            if not group_column:
                self.result_label.config(text=languages[self.current_language]["no_group_column"])
                return

            # 获取评分列名
            rating_columns = []
            while True:
                rating_column = tkinter.simpledialog.askstring("输入信息",
                                                               languages[self.current_language]["rating_column_prompt"],
            parent=self.root)
                if rating_column is None:
                    break
                if rating_column.strip():
                    rating_columns.append(rating_column.strip())
                else:
                    print(languages[self.current_language]["empty_rating_column"])

            if not rating_columns:
                self.result_label.config(text=languages[self.current_language]["no_rating_columns"])
                return

            all_results = []
            file_names = []
            rwg_values_all = []
            group_names_all = []

            for file_path in file_paths:
                # 打开 Excel 文件
                df = pd.read_excel(file_path)

                # 进行rwg分析
                rwg_results, rwg_values = self.rwg_analysis(df, group_column, rating_columns)

                # 检查是否存在ICC计算失败的情况（通过查找结果中是否有NaN的ICC值）
                has_icc_error = any(
                    "ICC1" in res and np.isnan(res["ICC1"])
                    for res in rwg_results
                    if isinstance(res, dict)
                )

                all_results.extend(rwg_results)
                file_names.extend([os.path.basename(file_path)] * len(rwg_results))

                # 收集rwg值和组名用于绘图
                if rwg_values:
                    for i, result in enumerate(rwg_results):
                        if i < len(rwg_results) - 1:  # 最后一个是统计结果
                            key = list(result.keys())[0]
                            group = key.split('_')[0]
                            rwg_values_all.append(result[key])
                            group_names_all.append(group)

            # 整理数据
            all_data = []
            for i, results in enumerate(all_results):
                if results:
                    data = []
                    for key, value in results.items():
                        data.append([f"{file_names[i]}_{key}", value])
                    all_data.extend(data)

            if self.current_language == "zh":
                headers = ["指标", "数值"]
            else:
                headers = ["Indicator", "Value"]
            df_result = pd.DataFrame(all_data, columns=headers)

            # 根据当前语言定义列名和表头文本
            if self.current_language == "zh":
                # 中文列名
                columns = [
                    "rwg值", "Rwg值标准差SD", "P25", "中位数", "P75",
                    "ICC1", "ICC2", "MSB", "MSW", "F值", "p值",
                    "Welch_F值", "Welch_p值", "Kruskal-Wallis_H值", "Kruskal-Wallis_p值"
                ]
                explanation_header = "指标_解释说明"
                explanation_text = "解释说明"
                interpretation_header = "指标_结果解读"
                interpretation_text = "结果解读"
            else:
                # 英文列名（国际统计学术语标准翻译）
                columns = [
                    "rwg Value", "Rwg SD", "P25", "Median", "P75",
                    "ICC1", "ICC2", "MSB", "MSW", "F Value", "p Value",
                    "Welch F Value", "Welch p Value", "Kruskal-Wallis H Value", "Kruskal-Wallis p Value"
                ]
                explanation_header = "Indicator_Explanation"
                explanation_text = "Explanation"
                interpretation_header = "Indicator_Interpretation"
                interpretation_text = "Interpretation"

            # 添加解释说明
            explanations = languages[self.current_language]['explanation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(columns=columns)
            explanation_df.insert(0, explanation_header, explanation_text)

            # 添加分析结果解读
            interpretations = languages[self.current_language]['interpretation']
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(columns=columns)
            interpretation_df.insert(0, interpretation_header, interpretation_text)

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(
                defaultextension=".docx", filetypes=[("Word files", "*.docx")])

            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加分析结果表格
                doc.add_heading("分析结果" if self.current_language == 'zh' else "Analysis Results", level=1)
                table = doc.add_table(rows=1, cols=2)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "统计量" if self.current_language == 'zh' else "Statistic"
                hdr_cells[1].text = "统计量值" if self.current_language == 'zh' else "Value"

                for index, row in df_result.iterrows():
                    row_cells = table.add_row().cells
                    # 提取统计量名称（去除文件名前缀）
                    stat_name = row[headers[0]].split('_', 1)[1] if '_' in row[headers[0]] else row[headers[0]]
                    row_cells[0].text = stat_name
                    # 处理数值格式化
                    value = row[headers[1]]
                    if isinstance(value, float):
                        if np.isnan(value):
                            row_cells[1].text = "NaN"
                        else:
                            row_cells[1].text = f"{value:.4f}"
                    else:
                        row_cells[1].text = str(value)

                # 添加解释说明（项目符号列表）
                doc.add_heading("解释说明" if self.current_language == 'zh' else "Explanation", level=1)
                explanations = languages[self.current_language]['explanation']
                for key, value in explanations.items():
                    doc.add_paragraph(f"{key}: {value}", style='ListBullet')

                # 添加分析结果解读（项目符号列表）
                doc.add_heading("结果解读" if self.current_language == 'zh' else "Interpretation", level=1)
                interpretations = languages[self.current_language]['interpretation']
                for key, value in interpretations.items():
                    doc.add_paragraph(f"{key}: {value}", style='ListBullet')

                # 生成rwg值柱状图
                if rwg_values_all:
                    fig, ax = plt.subplots(figsize=(10, 6))
                    ax.bar(group_names_all, rwg_values_all)
                    ax.set_title(
                        '组内评分者信度rwg值' if self.current_language == 'zh' else 'Within-Group Inter-Rater Reliability rwg Values')
                    ax.set_ylabel('rwg值' if self.current_language == 'zh' else 'rwg Value')
                    ax.set_xlabel('分组' if self.current_language == 'zh' else 'Group')
                    plt.xticks(rotation=0, ha='right')
                    plt.tight_layout()

                    # 保存图片
                    img_path = os.path.splitext(save_path)[0] + '_rwg.png'
                    plt.savefig(img_path)
                    plt.close()

                    # 将图片插入 Word 文档
                    doc.add_heading(
                        "组内评分者信度rwg值柱状图" if self.current_language == 'zh' else "Within-Group Inter-Rater Reliability rwg Values Bar Chart",
                        level=2)
                    doc.add_picture(img_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                self.result_label.config(text=languages[self.current_language]["analysis_complete"].format(save_path),
                                         wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]["no_save_path_selected"])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]["analysis_error"].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])
        self.open_excel_label.config(text=languages[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=languages[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame,
                                               text=languages[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = WithinGroupInterRaterReliabilityRwgAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()
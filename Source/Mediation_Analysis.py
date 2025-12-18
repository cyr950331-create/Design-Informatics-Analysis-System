import tkinter as tk
from tkinter import filedialog
import os
import numpy as np
import pandas as pd
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
import tkinter.simpledialog
import statsmodels.api as sm
from docx import Document
from docx.shared import Inches
from scipy.stats import norm

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
languages = {
    'zh': {
        'title': "中介作用",
        'select_button_text': "选择文件",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'analyze_button_text': "分析文件",
        'no_file_selected': "请选择文件。",
        'file_not_exists': "文件不存在，请重新选择。",
        'analysis_error': "分析文件时出错: {}",
        'analysis_complete': "分析完成，结果已保存到 {}，相关图片已保存。",
        'no_save_path_selected': "未选择保存路径，结果未保存。",
        'open_excel_button_text': "示例数据",
        'switch_language_button_text': "中/英",
        'input_info': "输入信息",
        'input_ind_var': "请输入自变量的列名（多个用逗号分隔）",
        'input_med_var': "请输入中介变量的列名（多个用逗号分隔）",
        'input_dep_var': "请输入因变量的列名",
        'input_incomplete': "未输入完整的变量名，分析取消。",
        'labels': {
            'plot_title': '中介作用分析结果',
            'plot_ylabel': '效应值',
            'table_headers': ['统计类型', '相关变量', '效应值', 'p值', '显著性'],
            'explanation_heading': '说明：',
            'explanation1': '路径a表示自变量对中介变量的影响；路径b表示中介变量对因变量的影响（控制自变量）；',
            'explanation2': '路径c为自变量对因变量的总效应，路径c\'为控制中介后的直接效应。',
            'explanation3': '若间接效应（a×b）显著，说明存在显著的中介作用。'
        }
    },
    'en': {
        'title': "Mediation",
        'select_button_text': "Select File",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'analyze_button_text': "Analyze File",
        'no_file_selected': "Please select a valid file path.",
        'file_not_exists': "The file does not exist. Please select again.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'analysis_complete': "Analysis completed. The results have been saved to {}, and the relevant images have been saved.",
        'no_save_path_selected': "No save path selected. The results were not saved.",
        'open_excel_button_text': "Example data",
        'switch_language_button_text': "Chinese/English",
        'input_info': "Input Information",
        'input_ind_var': "Please enter the column names of independent variables (separated by commas)",
        'input_med_var': "Please enter the column names of mediator variables (separated by commas)",
        'input_dep_var': "Please enter the column name of the dependent variable",
        'input_incomplete': "Incomplete variable names entered, analysis canceled.",
        'labels': {
            'plot_title': 'Mediation Analysis Results',
            'plot_ylabel': 'Effect Value',
            'table_headers': ['Statistical Type', 'Related Variables', 'Effect Value', 'p-value', 'Significance'],
            'explanation_heading': 'Notes:',
            'explanation1': 'Path a: effect of X on M; Path b: effect of M on Y controlling for X;',
            'explanation2': 'Path c: total effect of X on Y; Path c\': direct effect controlling for M.',
            'explanation3': 'A significant indirect effect (a×b) indicates a mediation effect.'
        }
    }
}


class MediationAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data8.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def mediation_analysis(self, data, ind_vars, med_vars, dep_var):

        results = []

        # 拆分输入变量（允许多重中介）
        ind_vars = [v.strip() for v in ind_vars.split(',')]
        med_vars = [v.strip() for v in med_vars.split(',')]

        # === 数据清洗 ===
        data = data[[*ind_vars, *med_vars, dep_var]].apply(pd.to_numeric, errors='coerce').dropna()

        # 遍历每对 X、M 的组合
        for ind_var in ind_vars:
            for med_var in med_vars:
                # 路径 a：X → M
                X_a = sm.add_constant(data[ind_var])
                model_a = sm.OLS(data[med_var], X_a).fit()

                # 路径 b 和 c'：Y ~ X + M
                X_b = sm.add_constant(data[[ind_var, med_var]])
                model_b = sm.OLS(data[dep_var], X_b).fit()

                # 路径 c：总效应（Y ~ X）
                X_c = sm.add_constant(data[ind_var])
                model_c = sm.OLS(data[dep_var], X_c).fit()

                # 提取系数
                a = model_a.params[ind_var]
                b = model_b.params[med_var]
                c = model_c.params[ind_var]
                c_prime = model_b.params[ind_var]

                # Sobel 检验：间接效应显著性
                sa = model_a.bse[ind_var]
                sb = model_b.bse[med_var]
                sobel_se = np.sqrt(b ** 2 * sa ** 2 + a ** 2 * sb ** 2 + sa ** 2 * sb ** 2)
                sobel_z = (a * b) / sobel_se
                sobel_p = 2 * (1 - norm.cdf(abs(sobel_z)))

                # 根据当前语言选择对应的路径类型描述
                if self.current_language == 'zh':
                    path_a = "路径a (X→M)"
                    path_b = "路径b (M→Y)"
                    path_c = "路径c (总效应)"
                    path_c_prime = "路径c' (直接效应)"
                    indirect_effect = "间接效应 (a×b)"
                else:
                    path_a = "Path a (X→M)"
                    path_b = "Path b (M→Y)"
                    path_c = "Path c (Total Effect)"
                    path_c_prime = "Path c' (Direct Effect)"
                    indirect_effect = "Indirect Effect (a×b)"

                # 保存结果
                results.extend([
                    {"类型": path_a, "变量": f"{ind_var}→{med_var}", "效应值": a,
                     "p值": model_a.pvalues[ind_var]},
                    {"类型": path_b, "变量": f"{med_var}→{dep_var}", "效应值": b,
                     "p值": model_b.pvalues[med_var]},
                    {"类型": path_c, "变量": f"{ind_var}→{dep_var}", "效应值": c,
                     "p值": model_c.pvalues[ind_var]},
                    {"类型": path_c_prime, "变量": f"{ind_var}→{dep_var}", "效应值": c_prime,
                     "p值": model_b.pvalues[ind_var]},
                    {"类型": indirect_effect, "变量": f"{ind_var}→{med_var}→{dep_var}", "效应值": a * b,
                     "p值": sobel_p},
                ])

        # 根据当前语言选择样本量的描述
        sample_size_label = "样本量" if self.current_language == 'zh' else "Sample Size"
        results.append({"类型": sample_size_label, "变量": "", "效应值": len(data), "p值": ""})
        return results

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 让用户输入自变量、中介变量和因变量的列名
            ind_vars = tkinter.simpledialog.askstring(languages[self.current_language]['input_info'],
                                                      languages[self.current_language]['input_ind_var'])
            med_vars = tkinter.simpledialog.askstring(languages[self.current_language]['input_info'],
                                                      languages[self.current_language]['input_med_var'])
            dep_var = tkinter.simpledialog.askstring(languages[self.current_language]['input_info'],
                                                     languages[self.current_language]['input_dep_var'])

            if not ind_vars or not med_vars or not dep_var:
                self.result_label.config(text=languages[self.current_language]['input_incomplete'])
                return

            # 进行中介作用分析
            analysis_results = self.mediation_analysis(df, ind_vars, med_vars, dep_var)

            # 整理数据用于输出
            data = []
            effects_for_plot = []
            labels_for_plot = []

            for item in analysis_results:
                # 跳过样本量（中英文）
                if item["类型"] not in ["样本量", "Sample Size"]:
                    effects_for_plot.append(item["效应值"])
                    # 标签包括路径类型和变量，方便区分
                    labels_for_plot.append(f"{item['类型']} {item['变量']}".strip())

            table_headers = languages[self.current_language]['labels']['table_headers']

            def sig_label(p):
                """根据p值返回显著性标记"""
                if p == "" or pd.isna(p):
                    return ""
                if p < 0.001:
                    return "***"
                elif p < 0.01:
                    return "**"
                elif p < 0.05:
                    return "*"
                else:
                    return "ns"

            # 遍历分析结果列表，为每个结果字典添加“显著性”键
            for item in analysis_results:
                if item["p值"] != "":
                    item["显著性"] = sig_label(item["p值"])
                else:
                    item["显著性"] = ""

            df_result = pd.DataFrame(analysis_results)

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(
                defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()
                doc.add_heading(languages[self.current_language]["title"], level=1)
                doc.add_paragraph(languages[self.current_language]['labels']['explanation_heading'])
                doc.add_paragraph(languages[self.current_language]['labels']['explanation1'])
                doc.add_paragraph(languages[self.current_language]['labels']['explanation2'])
                doc.add_paragraph(languages[self.current_language]['labels']['explanation3'])

                # 添加表格
                table = doc.add_table(rows=df_result.shape[0] + 1, cols=df_result.shape[1])
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(table_headers):
                    hdr_cells[col_idx].text = header

                for row_idx in range(df_result.shape[0]):
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx in range(df_result.shape[1]):
                        row_cells[col_idx].text = str(df_result.iloc[row_idx, col_idx])

                if self.current_language == 'zh':
                    plt.rcParams['font.family'] = ['SimHei']  # 中文
                else:
                    plt.rcParams['font.family'] = ['Arial']  # 英文
                plt.rcParams['axes.unicode_minus'] = False

                # 图像宽度根据标签数量自动调整，避免拥挤
                fig, ax = plt.subplots(figsize=(max(12, len(labels_for_plot) * 0.6), 6))

                # 简化标签内容，避免过长
                short_labels = []
                for label in labels_for_plot:
                    if self.current_language == 'zh':
                        short_label = label.replace("路径a (X→M)", "a") \
                            .replace("路径b (M→Y)", "b") \
                            .replace("路径c (总效应)", "c") \
                            .replace("路径c' (直接效应)", "c′") \
                            .replace("间接效应 (a×b)", "a×b")
                    else:
                        short_label = label.replace("Path a (X→M)", "a") \
                            .replace("Path b (M→Y)", "b") \
                            .replace("Path c (Total Effect)", "c") \
                            .replace("Path c' (Direct Effect)", "c′") \
                            .replace("Indirect Effect (a×b)", "a×b")
                    short_labels.append(short_label)

                # 绘制柱状图
                ax.bar(range(len(short_labels)), effects_for_plot, color='#4e79a7')
                ax.axhline(0, color='gray', linewidth=1)

                # 设置横轴刻度与标签
                ax.set_xticks(range(len(short_labels)))
                ax.set_xticklabels(short_labels, fontsize=10, rotation=45, ha='right')

                # 添加效应值标签
                for i, v in enumerate(effects_for_plot):
                    ax.text(i, v + np.sign(v) * 0.05, f"{v:.2f}", ha='center', fontsize=9)

                # 设置标题与标签
                ax.set_title(languages[self.current_language]['labels']['plot_title'], fontsize=12)
                ax.set_ylabel(languages[self.current_language]['labels']['plot_ylabel'], fontsize=11)

                # 调整布局避免遮挡
                plt.tight_layout()
                plt.subplots_adjust(bottom=0.25)

                # 添加效应值标签
                for i, v in enumerate(effects_for_plot):
                    ax.text(i, v + np.sign(v) * 0.05, f"{v:.2f}", ha='center', fontsize=9)

                # 设置标题与标签
                ax.set_title(languages[self.current_language]['labels']['plot_title'])
                ax.set_ylabel(languages[self.current_language]['labels']['plot_ylabel'])

                # 优化横轴标签
                plt.xticks(rotation=60, ha='right')
                plt.tight_layout()
                plt.subplots_adjust(bottom=0.25)

                # 保存图片
                img_path = os.path.splitext(save_path)[0] + '.png'
                plt.savefig(img_path, dpi=300)  #
                plt.close()

                # 将图片插入 Word 文档
                doc.add_picture(img_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                self.result_label.config(text=languages[self.current_language]['analysis_complete'].format(save_path),
                                         wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]['no_save_path_selected'])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])
        self.open_excel_label.config(text=languages[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=languages[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame,
                                               text=languages[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = MediationAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()
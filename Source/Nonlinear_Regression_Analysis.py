import tkinter as tk
from tkinter import filedialog, messagebox
import os
import pandas as pd
import numpy as np
from sklearn.metrics import mean_squared_error, r2_score
import matplotlib.pyplot as plt
import matplotlib
import statsmodels.api as sm
from scipy.optimize import curve_fit, OptimizeWarning
from scipy import stats
import warnings
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from docx import Document
from docx.shared import Inches

matplotlib.rcParams["font.family"] = ["WenQuanYi Micro Hei", "Heiti TC", "Microsoft YaHei", "Arial Unicode MS", "SimHei"]
matplotlib.rcParams["axes.unicode_minus"] = False  # 正确显示负号

# ---------- 配置 ----------
BOOTSTRAP_ITERS = 200  # 可调整。越高置信区间越稳，但耗时越多。
MODEL_PARAM_CONFIG = {
    "exponential": {"param_count": 2, "default_params": [1.0, 0.1]},
    "logarithmic": {"param_count": 2, "default_params": [1.0, 0.1]},
    "power": {"param_count": 2, "default_params": [1.0, 1.0]},
    "quadratic": {"param_count": 3, "default_params": [0.1, 0.1, 1.0]},
    "sigmoid": {"param_count": 3, "default_params": [1.0, 0.1, 0.0]},
    "multivariate": {"base_param_count": 3, "combine_param_count": 1}
}

LANGUAGES = {
    'zh': {
        'title': "非线性回归",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "请选择文件。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'images_saved': "图片已保存到 {}",
        'switch_language': "中/英",
        'open_excel_button_text': "示例数据",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'explanation': {
            "Parameters": "模型参数，反映自变量与因变量的关系强度和方向。",
            "Mean Squared Error (MSE)": "均方误差，衡量预测值与真实值之间的平均误差。",
            "R-squared (R²)": "决定系数，取值范围在 0 到 1 之间，越接近 1 表示模型拟合效果越好。",
            "P-value": "p 值（非线性回归中通常不可靠，改为报告置信区间）。",
            "Best Model": "最佳模型，基于综合指标自动选择的最优拟合模型",
            "Regression Equation": "回归方程，反映自变量与因变量之间的函数关系",
            "AIC": "赤池信息准则（基于 RSS），值越小越好",
            "Residual Normality": "残差正态性检验，p值>0.05表示残差近似正态分布"
        },
        'table_headers': {
            "Model": "模型",
            "Parameter": "参数",
            "Value": "值",
            "Std Error": "标准误差",
            "t-value": "t值",
            "P-value": "p值"
        },
        'plot_labels': {
            'actual_data': '实际数据',
            'predicted_data': '预测数据'
        },
        'model_warning': "模型 {} 因数据限制无法使用: {}\n建议处理数据后重试",
        'fit_failed': "模型 {} 拟合失败: {}\n",
        'missing_values': "数据中存在缺失值，已自动删除含缺失值的行",
        'outlier_removed': "已移除 {} 行异常值（基于联合IQR+3σ方法）",
        'independent_var': "自变量列名",
        'dependent_var': "因变量列名",
        'exp_data_transform': "指数模型数据预处理: 对全部因变量取负值以满足正值要求（若原始全为负）",
        'multi_var_warning': "检测到多自变量，将分别为每个自变量构建单变量模型并尝试构建多变量组合模型",
        'select_image_dir': "选择图片保存目录",
        'too_many_vars': "自变量数量过多（{}个），建议不超过5个以保证分析效果",
        'multi_model_title': "多变量组合模型",
        'feature_selection': "基于单变量模型性能筛选出的关键变量: {}"
    },
    'en': {
        'title': "Nonlinear Regression",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'images_saved': "Images have been saved to {}",
        'switch_language': "Chinese/English",
        'open_excel_button_text': "Example data",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'explanation': {
            "Parameters": "Model parameters, reflecting the strength and direction of the relationship between independent and dependent variables.",
            "Mean Squared Error (MSE)": "Mean squared error, measuring the average error between the predicted and actual values.",
            "R-squared (R²)": "Coefficient of determination, ranging from 0 to 1. A value closer to 1 indicates a better fit of the model.",
            "P-value": "p value (often unreliable for nonlinear models; we report confidence intervals instead).",
            "Best Model": "Best model, automatically selected based on comprehensive indicators",
            "Regression Equation": "Regression equation, reflecting the functional relationship between independent and dependent variables",
            "AIC": "Akaike Information Criterion (computed from RSS), smaller is better",
            "Residual Normality": "Residual normality test, p-value>0.05 indicates normal distribution"
        },
        'table_headers': {
            "Model": "Model",
            "Parameter": "Parameter",
            "Value": "Value",
            "Std Error": "Std Error",
            "t-value": "t-value",
            "P-value": "P-value"
        },
        'plot_labels': {
            'actual_data': 'Actual Data',
            'predicted_data': 'Predicted Data'
        },
        'model_warning': "Model {} cannot be used due to data restrictions: {}\nRecommend processing data and trying again",
        'fit_failed': "Model {} fitting failed: {}\n",
        'missing_values': "Missing values found in data, rows with missing values have been automatically removed",
        'outlier_removed': "{} rows removed as outliers (joint IQR + 3σ method)",
        'independent_var': "Independent variable column name",
        'dependent_var': "Dependent variable column name",
        'exp_data_transform': "Exponential model preprocessing: invert sign if all y are negative",
        'multi_var_warning': "Multiple independent variables detected, will build models for each and attempt a multivariate combined model",
        'select_image_dir': "Select image save directory",
        'too_many_vars': "Too many independent variables ({}), recommend no more than 5 for best results",
        'multi_model_title': "Multivariate Combined Model",
        'feature_selection': "Key variables selected based on univariate performance: {}"
    },
}

# ---------- 模型函数 ----------
def exponential_model(x, a, b):
    return a * np.exp(b * x)

def logarithmic_model(x, a, b):
    # x 必须 > 0
    return a + b * np.log(x)

def power_model(x, a, b):
    return a * (x ** b)

def quadratic_model(x, a, b, c):
    return a * (x ** 2) + b * x + c

def sigmoid_model(x, a, b, c):
    exponent = -b * (x - c)
    exponent = np.clip(exponent, -709, 709)
    return a / (1 + np.exp(exponent))

def multivariate_combined_model(X, *params):
    # X: (n_samples, n_vars)
    n_vars = X.shape[1]
    params_per_var = 3
    expected = n_vars * params_per_var + 1
    # 保证参数数量
    if len(params) < expected:
        params = tuple(list(params) + [0.1] * (expected - len(params)))
    var_params = params[:-1]
    k = params[-1]
    contributions = []
    for i in range(n_vars):
        a = var_params[i * params_per_var + 0]
        b = var_params[i * params_per_var + 1]
        c = var_params[i * params_per_var + 2]
        contributions.append(quadratic_model(X[:, i], a, b, c))
    combined = np.sum(contributions, axis=0)
    # 保留原设计，但给k较小的限制
    return combined * (1 + np.tanh(k * combined))

# ---------- 工具函数 ----------
def safe_clip_p0(p0, bounds):
    # bounds: (low_list, high_list)
    low, high = bounds
    p0_arr = np.array(p0, dtype=float)
    low_arr = np.array(low, dtype=float)
    high_arr = np.array(high, dtype=float)
    # 如果长度不同，尝试广播或截断
    L = min(len(p0_arr), len(low_arr), len(high_arr))
    p0_adj = p0_arr.copy()
    p0_adj[:L] = np.minimum(np.maximum(p0_arr[:L], low_arr[:L]), high_arr[:L])
    return p0_adj.tolist()

def compute_aic_from_rss(n, rss, k):
    # AIC = n*ln(RSS/n) + 2k
    # 如果 rss==0，使用机器最小正值避免-log(0)
    rss_safe = max(rss, 1e-12)
    return n * np.log(rss_safe / n) + 2 * k

def param_ci_from_cov(params, cov, dof=1):
    # 若 cov 可用：95% CI ~ param ± t_{0.975, dof} * se
    if cov is None:
        return None
    se = np.sqrt(np.abs(np.diag(cov)))
    # 当自由度很大时，用正态近似
    t = stats.t.ppf(0.975, dof) if dof > 0 else 1.96
    lower = params - t * se
    upper = params + t * se
    return np.vstack([lower, upper]).T  # shape (len(params), 2)

def calculate_normality_pvalue(residuals):
    # 对残差进行正态性检验：小样本用Shapiro，大样本用D'Agostino
    try:
        if len(residuals) <= 500:
            _, p = stats.shapiro(residuals)
        else:
            _, p = stats.normaltest(residuals)
        return p
    except Exception:
        return np.nan

def bootstrap_param_cis(model_func, X, y, p0, bounds, method, n_boot=BOOTSTRAP_ITERS):
    """对单变量/多变量模型进行bootstrap以估计参数分布（返回 95% CI)"""
    n = len(y)
    params_boot = []
    for i in range(n_boot):
        # 重采样索引（有放回）
        idx = np.random.choice(np.arange(n), size=n, replace=True)
        Xb = X[idx]
        yb = y[idx]
        try:
            params_b, _ = curve_fit(
                model_func, Xb, yb, p0=p0, bounds=bounds, maxfev=50000, method=method
            )
            params_boot.append(params_b)
        except Exception:
            continue
    if not params_boot:
        return None
    arr = np.vstack(params_boot)
    lower = np.percentile(arr, 2.5, axis=0)
    upper = np.percentile(arr, 97.5, axis=0)
    return np.vstack([lower, upper]).T

# ---------- 主应用类 ----------
class NonlinearRegressionAnalysisApp:
    def __init__(self, root=None):
        self.current_language = "en"
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])
        self.create_ui()

    def open_template_excel(self):
        try:
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)
            excel_path = os.path.join(parent_dir, "Sample_data", "Data33.xlsx")
            if os.path.exists(excel_path):
                os.startfile(excel_path)
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        self.root.transient()
        file_path = filedialog.askopenfilename(parent=self.root, filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def remove_outliers(self, X, y):
        """统一生成掩码（按行），避免逐列剔除导致行错位"""
        n = X.shape[0]
        mask = np.ones(n, dtype=bool)

        # X 每列分别计算掩码，然后取交集
        for i in range(X.shape[1]):
            col = X[:, i]
            skewness = stats.skew(col) if len(col) > 2 else 0
            if abs(skewness) < 1:
                mean = np.mean(col)
                std = np.std(col)
                m = np.abs((col - mean) / (std if std > 0 else 1)) < 3
            else:
                q1, q3 = np.percentile(col, [25, 75])
                iqr = q3 - q1
                lower = q1 - 1.5 * iqr
                upper = q3 + 1.5 * iqr
                m = (col >= lower) & (col <= upper)
            mask &= m

        # y 单独处理
        skew_y = stats.skew(y) if len(y) > 2 else 0
        if abs(skew_y) < 1:
            mean_y = np.mean(y)
            std_y = np.std(y)
            m_y = np.abs((y - mean_y) / (std_y if std_y > 0 else 1)) < 3
        else:
            q1y, q3y = np.percentile(y, [25, 75])
            iqr_y = q3y - q1y
            lower_y = q1y - 1.5 * iqr_y
            upper_y = q3y + 1.5 * iqr_y
            m_y = (y >= lower_y) & (y <= upper_y)

        mask &= m_y
        removed = np.sum(~mask)
        if removed > 0:
            self.result_label.config(text=LANGUAGES[self.current_language]['outlier_removed'].format(removed))
        return X[mask], y[mask]

    def get_regression_equation(self, model_name, params, independent_names, dependent_name, is_multivariate=False):
        # 格式化
        pv = [f"{v:.6f}".rstrip('0').rstrip('.') for v in params]
        if is_multivariate and independent_names:
            terms = []
            n_vars = len(independent_names)
            for i, var in enumerate(independent_names):
                a, b, c = pv[i*3:(i+1)*3]
                terms.append(f"({a})*{var}^2 + ({b})*{var} + ({c})")
            k = pv[-1]
            combined = " + ".join(terms)
            return f"{dependent_name} = ({combined}) * (1 + tanh({k} * ({combined})))"
        if model_name == "exponential":
            return f"{dependent_name} = {pv[0]} * exp({pv[1]} * x)"
        if model_name == "logarithmic":
            return f"{dependent_name} = {pv[0]} + {pv[1]} * ln(x)"
        if model_name == "power":
            return f"{dependent_name} = {pv[0]} * x^{ {pv[1]} }"
        if model_name == "quadratic":
            return f"{dependent_name} = {pv[0]}*x^2 + {pv[1]}*x + {pv[2]}"
        if model_name == "sigmoid":
            return f"{dependent_name} = {pv[0]} / (1 + exp(-{pv[1]}*(x - {pv[2]})))"
        return "Unknown"

    def analyze_single_independent_variable(self, X_single, y, independent_name, dependent_name):
        """为单个自变量拟合所有候选模型并返回结果字典"""
        X_single = np.asarray(X_single).ravel()
        y = np.asarray(y).ravel()
        model_results = {}
        warnings_text = ""

        # 初始值估计函数
        def exp_guess(x, y):
            # 如果 y 包含非正数会导致 log 问题
            if np.all(y > 0):
                try:
                    logy = np.log(y)
                    X_lin = sm.add_constant(x)
                    res = sm.OLS(logy, X_lin).fit()
                    a0 = np.exp(res.params[0])
                    b0 = res.params[1]
                    return [np.clip(a0, 1e-6, 1e6), np.clip(b0, -10, 10)]
                except Exception:
                    pass
            return MODEL_PARAM_CONFIG["exponential"]["default_params"]

        def power_guess(x, y):
            if np.all(x > 0) and np.all(y > 0):
                try:
                    lx, ly = np.log(x), np.log(y)
                    res = sm.OLS(ly, sm.add_constant(lx)).fit()
                    a0 = np.exp(res.params[0])
                    b0 = res.params[1]
                    return [np.clip(a0, 1e-6, 1e6), np.clip(b0, -10, 10)]
                except Exception:
                    pass
            return MODEL_PARAM_CONFIG["power"]["default_params"]

        def linear_guess(x, y):
            try:
                res = sm.OLS(y, sm.add_constant(x)).fit()
                c = res.params
                # map to quadratic initial: a~0, b~slope, c~intercept
                return [0.0, c[1], c[0]]
            except Exception:
                return MODEL_PARAM_CONFIG["quadratic"]["default_params"]

        # 基于数据范围设置边界与 p0
        y_min, y_max = np.min(y), np.max(y)
        y_mean, y_std = np.mean(y), np.std(y)
        x_min, x_max = np.min(X_single), np.max(X_single)
        x_mean, x_std = np.mean(X_single), np.std(X_single)
        scale = 2.0

        models = {
            "exponential": {
                "func": exponential_model,
                "valid": True,  # 允许负y，但会做预处理
                "display_name": "指数模型",
                "param_names": ["a", "b"],
                "bounds": (
                    [ -1e6, -10 ],
                    [ 1e6, 10 ]
                ),
                "p0": exp_guess(X_single, y)
            },
            "logarithmic": {
                "func": logarithmic_model,
                "valid": np.all(X_single > 0),
                "display_name": "对数模型",
                "param_names": ["a", "b"],
                "bounds": (
                    [y_mean - scale*y_std, -10],
                    [y_mean + scale*y_std, 10]
                ),
                "p0": MODEL_PARAM_CONFIG["logarithmic"]["default_params"]
            },
            "power": {
                "func": power_model,
                "valid": np.all(X_single > 0) and np.all(y > 0),
                "display_name": "幂函数模型",
                "param_names": ["a", "b"],
                "bounds": (
                    [1e-9, -10],
                    [1e9, 10]
                ),
                "p0": power_guess(X_single, y)
            },
            "quadratic": {
                "func": quadratic_model,
                "valid": True,
                "display_name": "二次模型",
                "param_names": ["a", "b", "c"],
                "bounds": (
                    [-abs(y_max) * 10, -abs(y_max / (x_max + 1e-10)) * 10, y_mean - scale*y_std],
                    [abs(y_max) * 10, abs(y_max / (x_max + 1e-10)) * 10, y_mean + scale*y_std]
                ),
                "p0": linear_guess(X_single, y)
            },
            "sigmoid": {
                "func": sigmoid_model,
                "valid": True,
                "display_name": "Sigmoid 模型",
                "param_names": ["a", "b", "c"],
                "bounds": (
                    [1e-9, 1e-6, x_min - scale * x_std],
                    [abs(y_max) * 10 + 1e-3, 50, x_max + scale * x_std]
                ),
                "p0": [y_max if y_max>0 else 1.0, 1.0, x_mean]
            }
        }

        # 如果 y 全为负数，尝试对指数模型取负再拟合（并记录说明）
        exp_data_adjusted = False
        y_for_exp = y.copy()
        if np.all(y < 0):
            y_for_exp = -y
            exp_data_adjusted = True
            warnings_text += LANGUAGES[self.current_language]['exp_data_transform'] + "\n"

        # 拟合每个模型
        for name, info in models.items():
            if not info["valid"]:
                warnings_text += LANGUAGES[self.current_language]['model_warning'].format(info["display_name"], "数据限制")
                continue
            try:
                # 选择优化器：如果有 bounds 则优先用 trf/dogbox
                method = 'trf'
                # 准备 p0 与 bounds
                p0 = safe_clip_p0(info["p0"], info["bounds"])
                bounds = info["bounds"]
                y_use = y_for_exp if name == "exponential" else y

                with warnings.catch_warnings():
                    warnings.filterwarnings("error", category=OptimizeWarning)
                    params, cov = curve_fit(
                        info["func"], X_single, y_use,
                        p0=p0, bounds=bounds, maxfev=50000, method=method
                    )
                    # 若指数模型之前做了取反，则还原 a 的符号
                    if name == "exponential" and exp_data_adjusted:
                        params[0] = -params[0]

                # 预测与指标
                y_pred = info["func"](X_single, *params)
                rss = np.sum((y - y_pred) ** 2)
                mse = mean_squared_error(y, y_pred)
                r2 = r2_score(y, y_pred)
                n = len(y)
                k = len(params)
                aic = compute_aic_from_rss(n, rss, k)
                normality_p = calculate_normality_pvalue(y - y_pred)

                # 参数标准误和 CI（若 cov 有效）
                param_se = np.sqrt(np.abs(np.diag(cov))) if cov is not None else np.array([np.nan]*len(params))
                ci_cov = param_ci_from_cov(params, cov, dof=max(n - k, 1))

                # 计算 bootstrap CI（如果cov不可用或用户需要稳健CI）
                try:
                    ci_boot = bootstrap_param_cis(info["func"], X_single, y_use, p0, bounds, method, n_boot=BOOTSTRAP_ITERS)
                except Exception:
                    ci_boot = None

                model_results[name] = {
                    "params": params,
                    "params_cov": cov,
                    "param_se": param_se,
                    "ci_cov": ci_cov,
                    "ci_boot": ci_boot,
                    "mse": mse,
                    "r2": r2,
                    "aic": aic,
                    "rss": rss,
                    "normality_pvalue": normality_p,
                    "func": info["func"],
                    "display_name": info["display_name"],
                    "param_names": info["param_names"],
                    "model_name": name,
                    "independent_name": independent_name
                }
            except Exception as e:
                warnings_text += LANGUAGES[self.current_language]['fit_failed'].format(info["display_name"], str(e)) + "\n"
                continue

        return model_results, warnings_text

    def _generate_multivariate_param_names(self, var_names):
        param_names = []
        for i, var in enumerate(var_names):
            param_names.extend([f"a_{i + 1} ({var})", f"b_{i + 1} ({var})", f"c_{i + 1} ({var})"])
        param_names.append("k (组合系数)")
        return param_names

    def build_multivariate_model(self, X, y, independent_names, var_performances):
        # 最多使用前5个变量
        n_vars = len(independent_names)
        if n_vars <= 1:
            return None, "自变量不足以构建多变量模型"

        # 计算每个变量的最佳性能（用 aic 和 r2 的简单评分）
        var_scores = []
        for name in independent_names:
            perf = var_performances.get(name, {})
            if not perf:
                var_scores.append((name, -np.inf))
                continue
            # 选 model 的 score：用 aic 和 r2 简单合成（越大越好）
            scores = []
            for m, res in perf.items():
                # 负 aic 越高越优
                scores.append(-res["aic"] + res["r2"] * 10.0)
            var_scores.append((name, max(scores) if scores else -np.inf))

        var_scores.sort(key=lambda x: x[1], reverse=True)
        n_selected = max(2, min(len(var_scores), int(len(var_scores) * 0.5)))
        selected_vars = [v for v, _ in var_scores[:n_selected]]
        selected_indices = [independent_names.index(v) for v in selected_vars]
        X_sel = X[:, selected_indices]

        params_per_var = 3
        total_params = params_per_var * len(selected_vars) + 1

        # initial params: 尝试用各变量的 quadratic 参数，否则默认
        initial = []
        for v in selected_vars:
            q = var_performances.get(v, {}).get("quadratic")
            if q is not None:
                p = q["params"]
                if len(p) >= 3:
                    initial.extend(p[:3])
                else:
                    initial.extend(MODEL_PARAM_CONFIG["quadratic"]["default_params"])
            else:
                initial.extend(MODEL_PARAM_CONFIG["quadratic"]["default_params"])
        # append k
        initial.append(0.1)

        # bounds
        y_range = np.max(y) - np.min(y) if np.max(y) != np.min(y) else 1.0
        low = []
        high = []
        for _ in range(len(selected_vars)):
            low.extend([-y_range * 10, -y_range * 10, np.min(y) - abs(y_range) * 10])
            high.extend([y_range * 10, y_range * 10, np.max(y) + abs(y_range) * 10])
        low.append(-5.0)
        high.append(5.0)
        bounds = (low, high)

        # fit
        try:
            method = 'trf'
            p0 = safe_clip_p0(initial, bounds)
            with warnings.catch_warnings():
                warnings.filterwarnings("error", category=OptimizeWarning)
                params, cov = curve_fit(multivariate_combined_model, X_sel, y, p0=p0, bounds=bounds, maxfev=100000, method=method)

            y_pred = multivariate_combined_model(X_sel, *params)
            rss = np.sum((y - y_pred) ** 2)
            mse = mean_squared_error(y, y_pred)
            r2 = r2_score(y, y_pred)
            aic = compute_aic_from_rss(len(y), rss, len(params))
            normality_p = calculate_normality_pvalue(y - y_pred)
            se = np.sqrt(np.abs(np.diag(cov))) if cov is not None else [np.nan] * len(params)

            return {
                "params": params,
                "param_errors": se,
                "mse": mse,
                "r2": r2,
                "aic": aic,
                "normality_pvalue": normality_p,
                "selected_vars": selected_vars,
                "func": multivariate_combined_model,
                "display_name": LANGUAGES[self.current_language]['multi_model_title'],
                "param_names": self._generate_multivariate_param_names(selected_vars),
                "model_name": "multivariate"
            }, LANGUAGES[self.current_language]['feature_selection'].format(", ".join(selected_vars))
        except Exception as e:
            return None, f"多变量模型拟合失败: {str(e)}"

    def auto_detect_best_model(self, X, y, independent_names, dependent_name):
        all_results = {}
        all_warnings = ""
        var_performances = {name: {} for name in independent_names}
        num_vars = X.shape[1]

        if num_vars > 5:
            all_warnings += LANGUAGES[self.current_language]['too_many_vars'].format(num_vars) + "\n"

        if num_vars > 1:
            messagebox.showinfo(LANGUAGES[self.current_language]["title"], LANGUAGES[self.current_language]['multi_var_warning'])

        for i in range(min(num_vars, 5)):
            X_single = X[:, i]
            ind_name = independent_names[i]
            results, warn = self.analyze_single_independent_variable(X_single, y, ind_name, dependent_name)
            var_performances[ind_name] = results
            all_warnings += warn
            for mname, res in results.items():
                all_results[f"{ind_name}__{mname}"] = res

        if not all_results:
            raise ValueError(f"所有模型均无法有效拟合数据\n{all_warnings}")

        multi_model = None
        multi_model_msg = ""
        if num_vars > 1:
            multi_model, multi_model_msg = self.build_multivariate_model(X, y, independent_names, var_performances)
            if multi_model:
                all_results["multivariate_model"] = multi_model
                all_warnings += multi_model_msg + "\n"
            else:
                all_warnings += f"多变量模型构建失败: {multi_model_msg}\n"

        if all_warnings:
            messagebox.showwarning(LANGUAGES[self.current_language]["title"], all_warnings)

        # 模型选择：基于 aic、r2、normality 归一化后加权
        keys = list(all_results.keys())
        aic_vals = np.array([all_results[k]["aic"] for k in keys])
        r2_vals = np.array([all_results[k]["r2"] for k in keys])
        norm_vals = np.array([all_results[k]["normality_pvalue"] for k in keys])

        # 归一化（使用 ptp 以保持数值稳定）
        def safe_norm(arr):
            rng = np.ptp(arr)
            if rng == 0:
                return np.ones_like(arr) * 0.5
            return (arr - np.min(arr)) / rng

        aic_norm = (np.max(aic_vals) - aic_vals) / (np.ptp(aic_vals) if np.ptp(aic_vals) != 0 else 1.0)
        r2_norm = safe_norm(r2_vals)
        normality_norm = safe_norm(norm_vals)

        # 权重（可调整）
        scores = 0.6 * aic_norm + 0.3 * r2_norm + 0.1 * normality_norm
        best_idx = int(np.argmax(scores))
        best_key = keys[best_idx]
        return all_results, all_results[best_key], multi_model

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]["file_entry_placeholder"] or not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return

        self.result_label.config(text="正在分析...")

        try:
            df = pd.read_excel(file_path)
            if df.shape[1] < 2:
                raise ValueError("数据至少需要包含一列自变量和一列因变量")
            column_names = df.columns.tolist()
            independent_names = column_names[:-1]
            dependent_name = column_names[-1]

            initial_rows = df.shape[0]
            df = df.dropna()
            if df.shape[0] < initial_rows:
                messagebox.showinfo(LANGUAGES[self.current_language]["title"], LANGUAGES[self.current_language]['missing_values'])

            X = df.iloc[:, :-1].values
            y = df.iloc[:, -1].values

            X_clean, y_clean = self.remove_outliers(X, y)

            all_models, best_model, multi_model = self.auto_detect_best_model(X_clean, y_clean, independent_names, dependent_name)

            model_func = best_model["func"]
            model_display_name = best_model["display_name"]
            model_name = best_model["model_name"]
            params = best_model["params"]
            mse = best_model["mse"]
            r2 = best_model["r2"]
            aic = best_model["aic"]
            param_names = best_model["param_names"]
            independent_name = best_model.get("independent_name", ", ".join(best_model.get("selected_vars", [])))
            normality_pvalue = best_model["normality_pvalue"]

            regression_eq = self.get_regression_equation(model_name, params, best_model.get("selected_vars", [independent_name]), dependent_name, model_name=="multivariate")

            # 准备输出表
            headers = [
                LANGUAGES[self.current_language]['table_headers']["Model"],
                LANGUAGES[self.current_language]['table_headers']["Parameter"],
                LANGUAGES[self.current_language]['table_headers']["Value"],
            ]
            data = []
            data.append([LANGUAGES[self.current_language]['independent_var'], independent_name, ""])
            data.append([LANGUAGES[self.current_language]['dependent_var'], dependent_name, ""])
            data.append([LANGUAGES[self.current_language]['explanation']["Regression Equation"], regression_eq, ""])
            data.append(["", "", ""])  # 空行

            # 参数行
            for i, p in enumerate(params):
                pname = param_names[i] if i < len(param_names) else f"param_{i}"
                data.append([model_display_name, pname, f"{p:.6f}"])

            data.append([model_display_name, "MSE", f"{mse:.6f}"])
            data.append([model_display_name, "R-squared (R²)", f"{r2:.6f}"])
            data.append([model_display_name, "AIC", f"{aic:.6f}"])
            data.append([model_display_name, LANGUAGES[self.current_language]['explanation']["Residual Normality"],
                         f"p={normality_pvalue:.6f}"])
            data.append(["", LANGUAGES[self.current_language]['explanation']["Best Model"], model_display_name])
            data.append(["", "", ""])

            df_result = pd.DataFrame(data, columns=headers)

            # 多变量模型结果（若存在且非最佳模型）
            if multi_model and model_name != "multivariate":
                multi_params = multi_model["params"]
                multi_mse = multi_model["mse"]
                multi_r2 = multi_model["r2"]
                multi_aic = multi_model["aic"]
                multi_normality = multi_model["normality_pvalue"]
                multi_param_names = multi_model["param_names"]
                multi_display_name = multi_model["display_name"]
                multi_regression_eq = self.get_regression_equation("multivariate", multi_params,
                                                                   multi_model["selected_vars"], dependent_name, True)

                # 只保留前三列数据，移除后三列空值
                data.append(
                    ["", LANGUAGES[self.current_language]['explanation']["Regression Equation"], multi_regression_eq])
                data.append(["", "", ""])  # 空行
                for i, p in enumerate(multi_params):  # 不再处理标准误差
                    pname = multi_param_names[i] if i < len(multi_param_names) else f"mp_{i}"
                    data.append([multi_display_name, pname, f"{p:.6f}"])  # 仅保留三列
                data.append([multi_display_name, "MSE", f"{multi_mse:.6f}"])
                data.append([multi_display_name, "R-squared (R²)", f"{multi_r2:.6f}"])
                data.append([multi_display_name, "AIC", f"{multi_aic:.6f}"])
                data.append([multi_display_name, LANGUAGES[self.current_language]['explanation']["Residual Normality"],
                             f"p={multi_normality:.6f}"])
                data.append(["", "", ""])  # 空行

            explanation_data = []
            explanations = LANGUAGES[self.current_language]['explanation']
            for key, val in explanations.items():
                explanation_data.append(["解释说明" if self.current_language=='zh' else "Explanation", key, val])
            explanation_df = pd.DataFrame(explanation_data, columns=headers)

            save_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                                     filetypes=[("Word files", "*.docx"), ("Excel files", "*.xlsx"), ("CSV files", "*.csv")])
            # 保存结果

            if save_path:
                desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
                image_dir = desktop_path if os.path.isdir(desktop_path) else os.path.dirname(save_path)
                file_ext = os.path.splitext(save_path)[1].lower()

                if file_ext == ".docx":
                    doc = Document()
                    # 添加标题
                    doc.add_heading(LANGUAGES[self.current_language]['title'], level=1)

                    # 添加模型结果表格
                    doc.add_heading("模型结果", level=2)
                    table_data = data
                    table = doc.add_table(rows=len(table_data) + 1, cols=len(headers))

                    # 设置表头
                    hdr = table.rows[0].cells
                    for i, h in enumerate(headers):
                        hdr[i].text = str(h)

                    # 填充表格数据
                    for r, row in enumerate(table_data):
                        cells = table.rows[r + 1].cells
                        for c, v in enumerate(row):
                            cells[c].text = str(v)

                    # 添加解释说明（项目列表形式）
                    doc.add_heading("解释说明" if self.current_language == 'zh' else "Explanation", level=2)
                    explanations = LANGUAGES[self.current_language]['explanation']
                    explanation_para = doc.add_paragraph()
                    for key, val in explanations.items():
                        # 添加项目符号列表
                        explanation_para.add_run(f"• {key}: {val}\n")

                    doc.save(save_path)
                elif file_ext == ".xlsx":
                    # Excel格式保持不变，仍使用表格
                    combined_df = pd.concat([pd.DataFrame(data, columns=headers), explanation_df], ignore_index=True)
                    combined_df.to_excel(save_path, index=False)
                elif file_ext == ".csv":
                    # CSV格式保持不变
                    combined_df = pd.concat([pd.DataFrame(data, columns=headers), explanation_df], ignore_index=True)
                    combined_df.to_csv(save_path, index=False, encoding='utf-8-sig')

                # 绘图
                plt.figure(figsize=(10,6))
                if X_clean.shape[1] > 1:
                    x_plot = np.mean(X_clean, axis=1)
                    sorted_idx = np.argsort(x_plot)
                    plt.scatter(x_plot, y_clean, label=LANGUAGES[self.current_language]['plot_labels']['actual_data'])
                    y_smooth = model_func(X_clean[sorted_idx] if model_name!="multivariate" else X_clean[sorted_idx], *params)
                    plt.plot(x_plot[sorted_idx], y_smooth, 'r-',
                             label=LANGUAGES[self.current_language]['plot_labels']['predicted_data'])
                    plt.xlabel("自变量组合" if self.current_language=='zh' else "Combined independent variable")
                else:
                    Xs = X_clean[:, 0]
                    idx = np.argsort(Xs)
                    plt.scatter(Xs, y_clean, label=LANGUAGES[self.current_language]['plot_labels']['actual_data'])
                    xs_sorted = Xs[idx]
                    y_smooth = model_func(xs_sorted, *params)
                    plt.plot(xs_sorted, y_smooth, 'r-',
                             label=LANGUAGES[self.current_language]['plot_labels']['predicted_data'])
                    plt.xlabel(independent_name)
                plt.ylabel(dependent_name)
                plt.title(model_display_name)
                plt.legend()
                best_img_path = os.path.join(image_dir, "best_model.png")
                plt.savefig(best_img_path)
                plt.close()

                if file_ext == ".docx":
                    doc = Document(save_path)
                    doc.add_picture(best_img_path, width=Inches(6))
                    if multi_model and model_name != "multivariate":
                        plt.figure(figsize=(10,6))
                        plt.scatter(np.mean(X_clean, axis=1), y_clean,
                                    label=LANGUAGES[self.current_language]['plot_labels']['actual_data'])
                        x_plot = np.mean(X_clean, axis=1)
                        idx = np.argsort(x_plot)
                        multi_params = multi_model["params"]
                        y_multi = multi_model["func"](X_clean[idx], *multi_params)
                        plt.plot(x_plot[idx], y_multi, 'g-',
                                 label=LANGUAGES[self.current_language]['plot_labels']['predicted_data'])
                        plt.xlabel("自变量组合" if self.current_language == 'zh' else "Combined independent variable")
                        plt.ylabel(dependent_name)
                        plt.legend()
                        multi_img_path = os.path.join(image_dir, "multivariate_model.png")
                        plt.savefig(multi_img_path)
                        plt.close()
                        doc.add_picture(multi_img_path, width=Inches(6))
                    doc.save(save_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(save_path)
                result_msg += LANGUAGES[self.current_language]['images_saved'].format(image_dir)
                self.result_label.config(text=result_msg, wraplength=400)
            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        self.current_language = "en" if self.current_language == "zh" else "zh"
        self.root.title(LANGUAGES[self.current_language]["title"])
        self.select_button.config(text=LANGUAGES[self.current_language]["select_button"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=LANGUAGES[self.current_language]["analyze_button"])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]["switch_language"])
        self.open_excel_label.config(text=LANGUAGES[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)
        min_width, min_height = 500, 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        frame = ttk.Frame(self.root)
        frame.pack(expand=True, padx=20)

        self.select_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["select_button"], command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        self.file_entry = ttk.Entry(frame, width=60)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        self.analyze_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["analyze_button"], command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        self.open_excel_label = ttk.Label(frame, text=LANGUAGES[self.current_language]["open_excel_button_text"], foreground="gray", cursor="hand2")
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        self.switch_language_label = ttk.Label(frame, text=LANGUAGES[self.current_language]["switch_language"], foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=5)

        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10, padx=20)

    def run(self):
        self.root.mainloop()

def run_app():
    app = NonlinearRegressionAnalysisApp()
    app.run()

if __name__ == "__main__":
    run_app()

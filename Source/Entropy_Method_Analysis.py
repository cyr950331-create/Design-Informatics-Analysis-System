import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
import matplotlib.pyplot as plt
import pathlib
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典，包含所有需要翻译的文本
LANGUAGES = {
    'zh': {
        'title': "熵值法",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "请选择文件。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "中/英",
        'open_excel_button_text': "示例数据",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'explanation': {
            "indicator_weight": "通过熵值法计算得到的各指标的权重",
            "indicator_entropy": "各指标的熵值，反映指标的信息无序程度"
        },
        'interpretation': {
            "indicator_weight": "权重越大，说明该指标在综合评价中越重要",
            "indicator_entropy": "熵值越大，说明该指标的信息无序程度越高，提供的有效信息越少"
        },
        'non_numeric_error': "数据中包含非数值内容，请检查文件",
        'empty_value_error': "数据中包含空值，请检查文件",
        'stats_type': {
            "indicator_weight": "指标权重",
            "indicator_entropy": "指标熵值"
        },
        'doc_title': "熵值法分析结果",
        'stats_results': "统计量结果",
        'explanation_title': "解释说明",
        'interpretation_title': "结果解读",
        'chart_title': "指标权重与熵值对比图",
        'weight_chart_title': "指标权重柱状图",
        'entropy_chart_title': "指标熵值柱状图",
        'x_axis_label': "指标",
        'weight_y_label': "权重",
        'entropy_y_label': "熵值",
        'table_headers': {
            "stats_type": "统计量类型",
            "serial_number": "序号",
            "indicator_name": "指标名称",
            "value": "数值"
        }
    },
    'en': {
        'title': "Entropy Method",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Chinese/English",
        'open_excel_button_text': "Example data",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'explanation': {
            "indicator_weight": "The weights of each indicator calculated by the entropy method",
            "indicator_entropy": "The entropy value of each indicator, reflecting the degree of information disorder of the indicator"
        },
        'interpretation': {
            "indicator_weight": "The larger the weight, the more important the indicator is in the comprehensive evaluation",
            "indicator_entropy": "The larger the entropy value, the higher the degree of information disorder of the indicator and the less effective information it provides"
        },
        'non_numeric_error': "Data contains non-numeric content, please check the file",
        'empty_value_error': "Data contains empty values, please check the file",
        'stats_type': {
            "indicator_weight": "Indicator Weight",
            "indicator_entropy": "Indicator Entropy"
        },
        'doc_title': "Entropy Method Analysis Results",
        'stats_results': "Statistical Results",
        'explanation_title': "Explanations",
        'interpretation_title': "Interpretations",
        'chart_title': "Comparison of Indicator Weights and Entropy Values",
        'weight_chart_title': "Bar Chart of Indicator Weights",
        'entropy_chart_title': "Bar Chart of Indicator Entropy",
        'x_axis_label': "Indicators",
        'weight_y_label': "Weights",
        'entropy_y_label': "Entropy",
        'table_headers': {
            "stats_type": "Statistic Type",
            "serial_number": "Serial Number",
            "indicator_name": "Indicator Name",
            "value": "Value"
        }
    }
}


class EntropyMethodAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为中文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data28.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"{LANGUAGES[self.current_language]['file_not_found']}：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"{LANGUAGES[self.current_language]['analysis_error'].format(str(e))}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def entropy_method(self, data):
        """
        进行熵值法分析
        :param data: 输入数据，每行代表一个样本，每列代表一个指标
        :return: 指标权重，指标熵值
        """
        # 改进：使用min-max标准化处理负数问题
        min_vals = data.min(axis=0)
        max_vals = data.max(axis=0)
        ranges = max_vals - min_vals

        # 处理极差为0的情况
        ranges[ranges == 0] = 1e-8

        # 标准化到[0.0001, 1.0001]区间，避免0值
        standardized_data = (data - min_vals) / ranges
        standardized_data = standardized_data * 1.0000 + 0.0001  # 映射到[0.0001, 1.0001]

        # 计算每个指标的熵值
        # 改进：使用更合理的极小值避免对数计算问题
        eps = 1e-10
        entropy = -np.sum(standardized_data * np.log(standardized_data + eps), axis=0) / np.log(data.shape[0])

        # 计算指标的权重
        weight = (1 - entropy) / np.sum(1 - entropy)

        return weight, entropy

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            file_path = ""
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件，自动检测并使用表头
            df = pd.read_excel(file_path)

            # 保存表头信息
            headers = df.columns.tolist()

            # 数据验证：检查是否有非数值数据
            non_numeric_cols = df.select_dtypes(exclude=['number']).columns
            if not non_numeric_cols.empty:
                self.result_label.config(text=LANGUAGES[self.current_language]['non_numeric_error'])
                return

            # 检查是否存在空值
            if df.isnull().any().any():
                self.result_label.config(text=LANGUAGES[self.current_language]['empty_value_error'])
                return

            data = df.values

            # 进行熵值法分析
            weight, entropy = self.entropy_method(data)

            # 整理数据，使用表头信息
            weight_data = list(zip(headers, weight.tolist()))
            entropy_data = list(zip(headers, entropy.tolist()))

            # 构建指标权重和熵值的列表数据
            weight_rows = [[LANGUAGES[self.current_language]['stats_type']['indicator_weight'], idx + 1, name, value]
                          for idx, (name, value) in enumerate(weight_data)]
            entropy_rows = [[LANGUAGES[self.current_language]['stats_type']['indicator_entropy'], idx + 1, name, value]
                           for idx, (name, value) in enumerate(entropy_data)]

            # 合并并创建DataFrame
            df_result = pd.DataFrame(
                weight_rows + entropy_rows,
                columns=[
                    LANGUAGES[self.current_language]['table_headers']['stats_type'],
                    LANGUAGES[self.current_language]['table_headers']['serial_number'],
                    LANGUAGES[self.current_language]['table_headers']['indicator_name'],
                    LANGUAGES[self.current_language]['table_headers']['value']
                ]
            )

            # 获取解释说明和结果解读
            explanations = LANGUAGES[self.current_language]['explanation']
            interpretations = LANGUAGES[self.current_language]['interpretation']

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加标题
                doc.add_heading(LANGUAGES[self.current_language]['doc_title'], 0)

                # 添加表格
                doc.add_heading(LANGUAGES[self.current_language]['stats_results'], level=1)
                stats_table = doc.add_table(rows=1, cols=4)
                stats_hdr = stats_table.rows[0].cells
                stats_hdr[0].text = LANGUAGES[self.current_language]['table_headers']['stats_type']
                stats_hdr[1].text = LANGUAGES[self.current_language]['table_headers']['serial_number']
                stats_hdr[2].text = LANGUAGES[self.current_language]['table_headers']['indicator_name']
                stats_hdr[3].text = LANGUAGES[self.current_language]['table_headers']['value']

                for index, row in df_result.iterrows():
                    cells = stats_table.add_row().cells
                    cells[0].text = str(row[LANGUAGES[self.current_language]['table_headers']['stats_type']])
                    cells[1].text = str(row[LANGUAGES[self.current_language]['table_headers']['serial_number']])
                    cells[2].text = str(row[LANGUAGES[self.current_language]['table_headers']['indicator_name']])
                    cells[3].text = f"{row[LANGUAGES[self.current_language]['table_headers']['value']]:.6f}"

                # 添加解释说明（项目符号列表）
                doc.add_heading(LANGUAGES[self.current_language]['explanation_title'], level=1)
                explanation_para = doc.add_paragraph()
                for key, value in explanations.items():
                    explanation_para.add_run(f"• {LANGUAGES[self.current_language]['stats_type'][key]}: {value}\n")

                # 添加结果解读（项目符号列表）
                doc.add_heading(LANGUAGES[self.current_language]['interpretation_title'], level=1)
                interpretation_para = doc.add_paragraph()
                for key, value in interpretations.items():
                    interpretation_para.add_run(f"• {LANGUAGES[self.current_language]['stats_type'][key]}: {value}\n")

                # 生成指标权重柱状图，使用表头作为x轴标签
                fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 6))

                # 权重图
                ax1.bar(headers, weight)
                ax1.set_title(LANGUAGES[self.current_language]['weight_chart_title'])
                ax1.set_xlabel(LANGUAGES[self.current_language]['x_axis_label'])
                ax1.set_ylabel(LANGUAGES[self.current_language]['weight_y_label'])
                plt.setp(ax1.get_xticklabels(), rotation=0, ha='right')  # 旋转标签避免重叠

                # 熵值图
                ax2.bar(headers, entropy)
                ax2.set_title(LANGUAGES[self.current_language]['entropy_chart_title'])
                ax2.set_xlabel(LANGUAGES[self.current_language]['x_axis_label'])
                ax2.set_ylabel(LANGUAGES[self.current_language]['entropy_y_label'])
                plt.setp(ax2.get_xticklabels(), rotation=0, ha='right')  # 旋转标签避免重叠

                plt.tight_layout()
                # 保存图片
                img_path = os.path.splitext(save_path)[0] + '_analysis.png'
                plt.savefig(img_path)
                plt.close()

                # 将图片插入到 Word 文档中
                doc.add_heading(LANGUAGES[self.current_language]['chart_title'], level=1)
                doc.add_picture(img_path, width=Inches(6))

                # 保存 Word 文件
                doc.save(save_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(save_path)
                self.result_label.config(text=result_msg, wraplength=400)

            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        # 保存当前输入的文件路径
        current_text = self.file_entry.get()
        placeholder = LANGUAGES[self.current_language]["file_entry_placeholder"]

        # 切换语言
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'
        self.root.title(LANGUAGES[self.current_language]['title'])
        self.select_button.config(text=LANGUAGES[self.current_language]['select_button'])
        self.analyze_button.config(text=LANGUAGES[self.current_language]['analyze_button'])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]['switch_language'])
        self.open_excel_label.config(text=LANGUAGES[self.current_language]["open_excel_button_text"])

        # 恢复文件路径
        self.file_entry.delete(0, tk.END)
        if current_text != placeholder:  # 如果不是占位文本则保留
            self.file_entry.insert(0, current_text)
            self.file_entry.config(foreground='black')
        else:  # 否则显示新语言的占位文本
            self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
            self.file_entry.config(foreground='gray')

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["select_button"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["analyze_button"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=LANGUAGES[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=LANGUAGES[self.current_language]["switch_language"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = EntropyMethodAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()
import tkinter as tk
from tkinter import filedialog, messagebox
import os
import threading
import tempfile
import numpy as np
import pandas as pd
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib
matplotlib.use('Agg')  # 禁用GUI后端以避免线程警告
import matplotlib.pyplot as plt
from docx import Document
from sklearn.cluster import DBSCAN
from sklearn.metrics import silhouette_score
from sklearn.preprocessing import StandardScaler
from sklearn.decomposition import PCA
from sklearn.neighbors import NearestNeighbors

# Matplotlib 中文支持（如需）
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 语言字典（使用英文键名）
languages = {
    "zh": {
        "title": "密度聚类",
        "select_button": "选择文件",
        "analyze_button": "分析文件",
        "file_not_found": "请选择文件。",
        "analysis_success": "分析完成，结果已保存到：{}",
        "no_save_path": "未选择保存路径，结果未保存。",
        "analysis_error": "分析文件时出错: {}",
        "switch_language": "中/英",
        'open_excel_button_text': "示例数据",
        "file_entry_placeholder": "请输入待分析 Excel 文件的完整路径",
        "estimating_eps": "正在估计 DBSCAN eps 参数...",
        "processing": "正在处理数据与聚类（后台）...",
        "pc1_label": "主成分1",
        "pc2_label": "主成分2",
        "explanation": {
            "cluster_results": "每个样本所属的聚类标签（-1 表示噪声）",
            "cluster_scatter": "样本在二维投影空间（PCA）中的分布及聚类结果",
            "silhouette_score": "聚类质量指标（靠近1较好）。"
        },
        "interpretation": {
            "cluster_results": "可用于区分不同样本所属的类别，-1 表示噪声点",
            "cluster_scatter": "二维投影用于辅助直观理解聚类分布",
            "silhouette_score": "若不可用（N/A），表示簇数或簇样本量不足，无法计算"
        },
        "report_title": "密度聚类分析报告",
        "summary": "摘要",
        "file": "文件",
        "rows": "行数",
        "columns": "列数",
        "dbscan_eps": "DBSCAN eps（估计值）",
        "min_samples": "最小样本数",
        "pca_explained": "PCA解释方差（2个成分）",
        "silhouette": "轮廓系数",
        "cluster_distribution": "聚类分布",
        "samples": "个样本",
        "cluster_scatter_plot": "聚类散点图（二维投影）",
        "cluster_feature_means": "聚类特征均值（显示前10个特征）",
        "cluster": "聚类",
        "feature_means_truncated": "特征均值（截断显示）",
        "visualization_failed": "可视化失败"
    },
    "en": {
        "title": "Density Based Clustering",
        "select_button": "Select File",
        "analyze_button": "Analyze File",
        "file_not_found": "The file does not exist. Please select again.",
        "analysis_success": "Analysis completed. Results saved to: {}",
        "no_save_path": "No save path selected. Results not saved.",
        "analysis_error": "An error occurred while analyzing the file: {}",
        "switch_language": "Chinese/English",
        'open_excel_button_text': "Example data",
        "file_entry_placeholder": "Please enter the full path of the Excel file to be analyzed",
        "estimating_eps": "Estimating DBSCAN eps parameter...",
        "processing": "Processing data and clustering (in background)...",
        "pc1_label": "Principal Component 1",
        "pc2_label": "Principal Component 2",
        "explanation": {
            "cluster_results": "The cluster label of each sample (-1 means noise)",
            "cluster_scatter": "Samples in 2D PCA projection and cluster assignment",
            "silhouette_score": "Clustering quality metric (closer to 1 is better)."
        },
        "interpretation": {
            "cluster_results": "Can be used to distinguish categories; -1 stands for noise",
            "cluster_scatter": "2D projection helps visually inspect cluster distribution",
            "silhouette_score": "If N/A, indicates insufficient clusters/sizes for computation"
        },
        "report_title": "Density-Based Clustering Analysis",
        "summary": "Summary",
        "file": "File",
        "rows": "Rows",
        "columns": "Columns",
        "dbscan_eps": "DBSCAN eps (estimated)",
        "min_samples": "min_samples",
        "pca_explained": "PCA explained variance (2 comp)",
        "silhouette": "Silhouette",
        "cluster_distribution": "Cluster Distribution",
        "samples": "samples",
        "cluster_scatter_plot": "Cluster Scatter Plot (2D projection)",
        "cluster_feature_means": "Cluster Feature Means (first 10 features shown)",
        "cluster": "Cluster",
        "feature_means_truncated": "Feature means (truncated)",
        "visualization_failed": "Visualization failed"
    }
}


class DensityBasedClusteringAnalysisApp:
    def __init__(self, root=None):
        self.current_language = "en"
        if root is None:
            self.root = ttk.Window(themename="flatly")
        else:
            self.root = root
        self.root.title(languages[self.current_language]["title"])
        self._create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data36.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"{languages[self.current_language]['file_not_found']}：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    # UI
    def _create_ui(self):
        screen_w = self.root.winfo_screenwidth()
        screen_h = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        w = int(screen_w * 0.4)
        h = int(screen_h * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        screen_w = max(screen_w, min_width)
        screen_h = max(screen_h, min_height)

        x = (screen_w - w) // 2
        y = (screen_h - h) // 2
        self.root.geometry(f"{w}x{h}+{x}+{y}")

        main = ttk.Frame(self.root, padding=12)
        main.pack(expand=True, fill=tk.BOTH)

        # 居中容器（关键）
        center_frame = ttk.Frame(main)
        center_frame.place(relx=0.5, rely=0.5, anchor="center")

        # 选择文件按钮
        self.select_button = ttk.Button(center_frame, text=languages[self.current_language]["select_button"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=(4, 8))

        # 文件路径输入框
        self.file_entry = ttk.Entry(center_frame, width=60)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self._on_entry_click)
        self.file_entry.bind('<FocusOut>', self._on_focusout)
        self.file_entry.pack(pady=4, fill=tk.X)

        # 分析按钮
        self.analyze_button = ttk.Button(center_frame, text=languages[self.current_language]["analyze_button"],
                                         command=self._on_analyze_clicked, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=(8, 6))

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            center_frame,
            text=languages[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 切换语言链接
        self.switch_language_label = ttk.Label(center_frame, text=languages[self.current_language]["switch_language"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self._switch_language)
        self.switch_language_label.pack(pady=(4, 6))

        # 结果信息标签
        self.result_label = ttk.Label(center_frame, text="", justify=tk.LEFT, wraplength=500)
        self.result_label.pack(pady=8, fill=tk.X)

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def _on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def _on_focusout(self, event):
        if self.file_entry.get().strip() == "":
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def _switch_language(self, event=None):
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button"])
        self.analyze_button.config(text=languages[self.current_language]["analyze_button"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language"])
        # 更新 placeholder
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.open_excel_label.config(text=languages[self.current_language]["open_excel_button_text"])

    # ---------- DBSCAN 辅助函数 ----------
    def _estimate_eps(self, data, k=5, percentile=50):
        """
        使用 k-最近邻距离的分位数估计 eps：
        - k = min(5, n_samples-1)
        - 返回 percentile 分位数（默认 90%）
        """
        try:
            n = data.shape[0]
            k_nn = min(max(1, k), max(1, n - 1))
            neigh = NearestNeighbors(n_neighbors=k_nn)
            neigh.fit(data)
            distances, _ = neigh.kneighbors(data)
            # 取每个点到第 k_nn 个近邻的距离
            d_k = np.sort(distances[:, k_nn - 1])
            eps = float(np.percentile(d_k, percentile))
            # 若 eps 非正，则用中位数回退
            if eps <= 0 or np.isnan(eps):
                eps = float(np.median(d_k))
            # 若仍然不合理，设一个小的默认值
            if not np.isfinite(eps) or eps <= 0:
                eps = 0.5
            return eps
        except Exception:
            return 0.5

    def _estimate_min_samples(self, data):
        """
        估计 min_samples 的简单策略：
        - 推荐为 2 * data.shape[1]（维度相关的常用启发式）
        - 与样本数成比例限制，至少为2，至多为样本数的一小部分
        """
        n, dim = data.shape[0], data.shape[1]
        base = max(2, 2 * dim)
        # 限制为不超过样本数 // 2, 且至少 2
        ms = int(min(max(2, base), max(2, n // 2)))
        return ms

    # ---------- 聚类与可视化 ----------
    def _run_dbscan(self, data, eps=None, min_samples=None):
        """
        运行 DBSCAN 并返回 labels 与模型
        """
        if eps is None:
            eps = self._estimate_eps(data)
        if min_samples is None:
            min_samples = self._estimate_min_samples(data)

        model = DBSCAN(eps=eps, min_samples=min_samples)
        labels = model.fit_predict(data)
        return labels, model, eps, min_samples

    def _safe_silhouette(self, data, labels):
        """
        在安全条件下计算轮廓系数，否则返回 np.nan
        """
        try:
            unique_labels, counts = np.unique(labels, return_counts=True)
            # silhouette 要求至少 2 个簇且每簇至少含有 2 个样本
            valid = (len(unique_labels) >= 2) and np.all(counts >= 2)
            if valid:
                return float(silhouette_score(data, labels))
            else:
                return np.nan
        except Exception:
            return np.nan

    def _plot_and_save(self, data_2d, labels, save_basepath):
        """
        绘制 2D 散点图并保存，返回图片路径
        """
        try:
            plt.figure(figsize=(8, 5))
            unique_labels = np.unique(labels)
            colors = [plt.cm.Spectral(each) for each in np.linspace(0, 1, len(unique_labels))]
            for k, col in zip(unique_labels, colors):
                if k == -1:
                    col = [0, 0, 0, 1]  # 噪声用黑色
                mask = (labels == k)
                xy = data_2d[mask]
                plt.scatter(xy[:, 0], xy[:, 1], c=[tuple(col)], edgecolors='k', s=40,
                           label=f"{languages[self.current_language]['cluster']} {k}")
            plt.title(languages[self.current_language]['cluster_scatter_plot'])
            plt.xlabel(languages[self.current_language]["pc1_label"])
            plt.ylabel(languages[self.current_language]["pc2_label"])
            plt.legend(markerscale=0.7, fontsize='small', loc='best')
            plt.grid(alpha=0.3)

            img_path = os.path.splitext(save_basepath)[0] + '_density_clustering_scatter.png'
            plt.tight_layout()
            plt.savefig(img_path, dpi=200)
            plt.close()
            return img_path
        except Exception:
            try:
                # 回退：简单保存空白图
                img_path = os.path.splitext(save_basepath)[0] + '_density_clustering_scatter.png'
                plt.figure(figsize=(6, 3))
                plt.text(0.5, 0.5, languages[self.current_language]['visualization_failed'],
                        ha='center', va='center')
                plt.axis('off')
                plt.savefig(img_path)
                plt.close()
                return img_path
            except Exception:
                return None

    # ---------- 分析工作线程 ----------
    def _analysis_worker(self, file_path, save_path):
        """
        在后台线程中执行完整分析：
        - 读取文件，清洗并数值化
        - 标准化
        - 估计 eps/min_samples 并运行 DBSCAN
        - 计算安全的轮廓系数
        - PCA->绘图（若维度>2）
        - 写入 Word 报告
        最后通过 root.after 回调更新 UI。
        """
        update_text = lambda txt: self.root.after(0, lambda: self.result_label.config(text=txt))

        try:
            update_text(languages[self.current_language]["processing"])
            # 读取数据（尝试第一行作为 header；若失败则 header=None）
            try:
                df = pd.read_excel(file_path, header=0)
            except Exception:
                df = pd.read_excel(file_path, header=None)

            # 强制转换为数值（非数值变为 NaN），删除全空列
            df_num = df.apply(pd.to_numeric, errors='coerce').dropna(axis=1, how='all')

            if df_num.shape[1] == 0 or df_num.shape[0] == 0:
                raise ValueError("输入数据不包含任何数值列或行。")

            # 用列均值填补缺失值
            df_num = df_num.fillna(df_num.mean())

            # 标准化（Z-score）
            scaler = StandardScaler()
            data_scaled = scaler.fit_transform(df_num.values)

            # 估计 eps 与 min_samples
            update_text(languages[self.current_language]["estimating_eps"])
            eps_est = self._estimate_eps(data_scaled)
            ms_est = self._estimate_min_samples(data_scaled)

            # 运行 DBSCAN
            labels, model, eps_used, min_samples_used = self._run_dbscan(data_scaled, eps=eps_est, min_samples=ms_est)

            # 计算安全的轮廓系数
            silhouette_avg = self._safe_silhouette(data_scaled, labels)

            # 计算每簇样本数
            unique, counts = np.unique(labels, return_counts=True)
            cluster_counts = dict(zip(unique.tolist(), counts.tolist()))

            # 若维度>2，PCA降至2维以绘图；否则直接使用原始2维数据
            if data_scaled.shape[1] > 2:
                try:
                    pca = PCA(n_components=2)
                    data_2d = pca.fit_transform(data_scaled)
                    pca_info = f"{np.sum(pca.explained_variance_ratio_):.2%}"
                except Exception:
                    # 降维失败则用前两列（已标准化）
                    data_2d = data_scaled[:, :2]
                    pca_info = "N/A"
            else:
                data_2d = data_scaled[:, :2]
                pca_info = "N/A"

            # 生成并保存图像（以 save_path 为基名）
            img_path = None
            if save_path:
                img_path = self._plot_and_save(data_2d, labels, save_path)

            # 生成 Word 报告
            if save_path:
                try:
                    doc = Document()
                    doc.add_heading(languages[self.current_language]['report_title'],
                                   level=1 if self.current_language == 'en' else 0)

                    # 基本信息
                    doc.add_heading(languages[self.current_language]['summary'], level=1)
                    doc.add_paragraph(f"{languages[self.current_language]['file']}: {os.path.basename(file_path)}")
                    doc.add_paragraph(f"{languages[self.current_language]['rows']}: {df_num.shape[0]}, "
                                     f"{languages[self.current_language]['columns']}: {df_num.shape[1]}")
                    doc.add_paragraph(f"{languages[self.current_language]['dbscan_eps']}: {eps_used:.6g}, "
                                     f"{languages[self.current_language]['min_samples']}: {min_samples_used}")
                    doc.add_paragraph(f"{languages[self.current_language]['pca_explained']}: {pca_info}")
                    doc.add_paragraph(f"{languages[self.current_language]['silhouette']}: "
                                     f"{silhouette_avg if not np.isnan(silhouette_avg) else 'N/A'}")

                    # 聚类分布
                    doc.add_heading(languages[self.current_language]['cluster_distribution'], level=1)
                    for cl, cnt in sorted(cluster_counts.items(), key=lambda x: (x[0] != -1, x[0])):
                        doc.add_paragraph(f"{languages[self.current_language]['cluster']} {cl}: {cnt} "
                                         f"{languages[self.current_language]['samples']}")

                    # 插入散点图
                    if img_path and os.path.exists(img_path):
                        doc.add_heading(languages[self.current_language]['cluster_scatter_plot'], level=1)
                        try:
                            doc.add_picture(img_path)
                        except Exception:
                            # 插入失败忽略
                            pass

                    # 添加每簇的特征均值（基于原始数值列）
                    doc.add_heading(languages[self.current_language]['cluster_feature_means'], level=1)
                    means_tbl = doc.add_table(rows=1, cols=2)
                    hdr = means_tbl.rows[0].cells
                    hdr[0].text = languages[self.current_language]['cluster']
                    hdr[1].text = languages[self.current_language]['feature_means_truncated']
                    # 仅显示前10列值以避免过长
                    max_show_cols = min(10, df_num.shape[1])
                    for cl in sorted(cluster_counts.keys()):
                        row_cells = means_tbl.add_row().cells
                        row_cells[0].text = str(cl)
                        mask = (labels == cl)
                        if np.sum(mask) > 0:
                            vals = df_num.values[mask][:, :max_show_cols].mean(axis=0)
                            row_cells[1].text = ", ".join([f"{v:.4f}" for v in vals])
                        else:
                            row_cells[1].text = "N/A"

                    # 保存文档
                    doc.save(save_path)
                except Exception as e:
                    # 写文档失败不影响程序继续，但记录错误
                    update_text(f"报告生成失败: {e}")

            # 清理临时图片
            msg_lines = [
                f"{languages[self.current_language]['analysis_success'].format(save_path if save_path else 'N/A')}",
                f"DBSCAN eps: {eps_used:.6g}, min_samples: {min_samples_used}",
                f"{languages[self.current_language]['silhouette']}: {silhouette_avg if not np.isnan(silhouette_avg) else 'N/A'}",
                f"{languages[self.current_language]['cluster_distribution']}: {cluster_counts}"
            ]
            final_msg = "\n".join(msg_lines)
            update_text(final_msg)
        except Exception as e:
            update_text(languages[self.current_language]['analysis_error'].format(str(e)))

    # 启动后台线程
    def _on_analyze_clicked(self):
        file_path = self.file_entry.get().strip()
        if file_path == "" or file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["file_not_found"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_found"])
            return

        # 让用户选择保存文件（Word）
        save_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                                 filetypes=[("Word files", "*.docx")])
        if not save_path:
            self.result_label.config(text=languages[self.current_language]["no_save_path"])
            return

        # 禁用按钮避免重复点击
        self.analyze_button.config(state=tk.DISABLED)
        self.select_button.config(state=tk.DISABLED)
        self.result_label.config(text=languages[self.current_language]["processing"])

        # 启动后台线程
        worker = threading.Thread(target=self._analysis_thread_entry, args=(file_path, save_path), daemon=True)
        worker.start()

    def _analysis_thread_entry(self, file_path, save_path):
        try:
            self._analysis_worker(file_path, save_path)
        finally:
            # 在主线程恢复按钮状态
            self.root.after(0, lambda: self.analyze_button.config(state=tk.NORMAL))
            self.root.after(0, lambda: self.select_button.config(state=tk.NORMAL))

    def run(self):
        self.root.mainloop()


def run_app():
    app = DensityBasedClusteringAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()
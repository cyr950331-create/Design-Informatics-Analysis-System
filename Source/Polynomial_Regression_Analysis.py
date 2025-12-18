import tkinter as tk
from tkinter import filedialog, simpledialog
import os
import pandas as pd
import numpy as np
from sklearn.linear_model import LinearRegression
from sklearn.preprocessing import PolynomialFeatures
from sklearn.metrics import mean_squared_error, r2_score
from sklearn.model_selection import cross_val_score
import matplotlib.pyplot as plt
import statsmodels.api as sm
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from docx import Document
from docx.shared import Inches

# 定义语言字典（所有键使用英文）
LANGUAGES = {
    'zh': {
        'title': "多项式回归",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "请选择文件。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'images_saved': "图片已保存到 {}",
        'switch_language': "中/英",
        'open_excel_button_text': "示例数据",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'explanation': {
            "coefficients": "表示每个自变量对因变量的影响程度。",
            "intercept": "所有自变量为 0 时因变量的预测值。",
            "mse": "衡量预测值与真实值之间的平均误差。",
            "r_squared": "取值范围在 0 到 1 之间，越接近 1 表示模型拟合效果越好。",
            "adjusted_r_squared": "考虑了模型中自变量的数量，对模型的拟合优度进行了调整。",
            "f_value": "用于检验整个回归模型的显著性。",
            "t_value": "用于检验每个自变量的显著性。",
            "p_value": "用于判断自变量的显著性，p 值越小，自变量越显著。",
            "regression_equation": "用于根据自变量预测因变量的数学表达式。"
        },
        'metrics': {
            "intercept": "截距",
            "mse": "均方误差",
            "r_squared": "决定系数",
            "adjusted_r_squared": "调整决定系数",
            "f_value": "F统计量",
            "regression_equation": "回归方程",
            "coefficients": "回归系数 ({})",  # 带占位符用于格式化
            "t_value": "t统计量 ({})",
            "p_value": "p值 ({})"
        },
        'dv_prompt': "请输入因变量所在列的索引（从0开始）:",
        'dv_error': "因变量索引无效",
        'data_error': "数据校验失败: {}",
        'insufficient_cols': "数据至少需要2列（1个自变量，1个因变量）",
        'save_results': "是否保存分析结果？",
        'save_now': "立即保存",
        'discard': "放弃",
        'cross_val_progress': "正在进行交叉验证选择最佳阶数...",
        'regression_equation_label': "回归方程",
        'cv_results': "交叉验证结果 (MSE)",
        'degree': "阶数",
        'average_mse': "平均MSE",
        'statistical_results': "统计结果",
        'actual_vs_predicted': "实际值 vs 预测值",
        'actual_values': "实际值",
        'predicted_values': "预测值",
        'residual_plot': "残差图",
        'residuals': "残差",
        'cv_mse_comparison': "交叉验证MSE对比",
        'polynomial_degree': "多项式阶数",
        'cross_validation_mse': "交叉验证MSE",
        'report_title': "分析报告",
        'best_degree': "最佳阶数",
        'different_degrees_comparison': "不同阶数的交叉验证MSE对比",
        'variable_relationship': "自变量与因变量关系",
        'independent_vs_dependent': "自变量与因变量关系",
        'original_data': "原始数据",
        'fit_curve': "阶拟合曲线",
        'metric': "指标",
        'explanation_label': "解释说明"
    },
    'en': {
        'title': "Polynomial Regression",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'images_saved': "Images have been saved to {}",
        'switch_language': "Chinese/English",
        'open_excel_button_text': "Example data",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'explanation': {
            "coefficients": "Indicating the influence of each independent variable on the dependent variable.",
            "intercept": "Which is the predicted value of the dependent variable when all independent variables are 0.",
            "mse": "Measuring the average error between the predicted and actual values.",
            "r_squared": "Ranging from 0 to 1. A value closer to 1 indicates a better fit of the model.",
            "adjusted_r_squared": "Which takes into account the number of independent variables in the model and adjusts the goodness of fit of the model.",
            "f_value": "Used to test the significance of the entire regression model.",
            "t_value": "Used to test the significance of each independent variable.",
            "p_value": "Used to determine the significance of the independent variable. The smaller the p value, the more significant the independent variable.",
            "regression_equation": "A mathematical expression used to predict the dependent variable based on independent variables."
        },
        'metrics': {
            "intercept": "Intercept",
            "mse": "Mean Squared Error",
            "r_squared": "R-squared",
            "adjusted_r_squared": "Adjusted R-squared",
            "f_value": "F-value",
            "regression_equation": "Regression Equation",
            "coefficients": "Coefficient ({})",  # 带占位符用于格式化
            "t_value": "t-value ({})",
            "p_value": "p-value ({})"
        },
        'dv_prompt': "Please enter the index of dependent variable column (starting from 0):",
        'dv_error': "Invalid dependent variable index",
        'data_error': "Data validation failed: {}",
        'insufficient_cols': "Data requires at least 2 columns (1 independent, 1 dependent variable)",
        'save_results': "Would you like to save the analysis results?",
        'save_now': "Save Now",
        'discard': "Discard",
        'cross_val_progress': "Performing cross-validation to select optimal degree...",
        'regression_equation_label': "Regression Equation",
        'cv_results': "Cross-Validation Results (MSE)",
        'degree': "Degree",
        'average_mse': "Average MSE",
        'statistical_results': "Statistical Results",
        'actual_vs_predicted': "Actual vs Predicted Values",
        'actual_values': "Actual Values",
        'predicted_values': "Predicted Values",
        'residual_plot': "Residual Plot",
        'residuals': "Residuals",
        'report_title': "Analysis Report",
        'best_degree': "Optimal Degree",
        'cv_mse_comparison': "Cross-Validation MSE Comparison",
        'polynomial_degree': "Polynomial Degree",
        'cross_validation_mse': "Cross-Validation MSE",
        'different_degrees_comparison': "Cross-Validation MSE for Different Degrees",
        'variable_relationship': "Independent vs Dependent Variable",
        'independent_vs_dependent': "Independent vs Dependent Variable",
        'original_data': "Original Data",
        'fit_curve': "th Degree Fit",
        'metric': "Metric",
        'explanation_label': "Explanation"
    }
}


class PolynomialRegressionAnalysisApp:
    def __init__(self, root=None):
        self.current_language = "en"
        self.analysis_results = None  # 保存分析结果用于后续保存

        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data33.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def validate_data(self, df):
        # 检查列数是否足够
        if len(df.columns) < 2:
            return False, LANGUAGES[self.current_language]['insufficient_cols']

        # 检查是否有非数值数据
        if not all(df.dtypes.apply(pd.api.types.is_numeric_dtype)):
            non_numeric = [col for col in df.columns if not pd.api.types.is_numeric_dtype(df[col])]
            return False, f"非数值列: {', '.join(non_numeric)}"

        # 检查是否有缺失值
        if df.isnull().any().any():
            missing = [col for col in df.columns if df[col].isnull().any()]
            return False, f"包含缺失值的列: {', '.join(missing)}"

        return True, "数据有效"

    def save_results(self, results):
        save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        if save_path:
            self._generate_report(save_path, results)
            return True, save_path
        return False, None

    def _generate_report(self, save_path, results):
        # 解包结果
        (df, X, y, X_poly, poly, model, y_pred, coefficients, intercept, mse, r2, adjusted_r2,
         t_values, p_values, f_value, transposed_df, feature_names, cv_mse_scores, regression_eq) = results

        # 设置中文字体支持
        plt.rcParams["font.family"] = ["SimHei", "WenQuanYi Micro Hei", "Heiti TC", "DejaVu Sans"]
        plt.rcParams["axes.unicode_minus"] = False  # 解决负号显示问题

        doc = Document()
        doc.add_heading(
            f"{LANGUAGES[self.current_language]['title']} {LANGUAGES[self.current_language]['report_title']} ({LANGUAGES[self.current_language]['best_degree']}: {poly.degree})",
            level=1
        )

        # 添加回归方程
        doc.add_heading(LANGUAGES[self.current_language]['regression_equation_label'], level=2)
        doc.add_paragraph(regression_eq)

        # 添加交叉验证结果
        doc.add_heading(LANGUAGES[self.current_language]['cv_results'], level=2)
        if not cv_mse_scores:  # 检查是否有数据
            doc.add_paragraph("无交叉验证结果数据")
        else:
            cv_table = doc.add_table(rows=len(cv_mse_scores) + 1, cols=2)
            cv_table.rows[0].cells[0].text = LANGUAGES[self.current_language]['degree']
            cv_table.rows[0].cells[1].text = LANGUAGES[self.current_language]['average_mse']

            # 确保按阶数排序，避免索引混乱
            for i, (degree, score) in enumerate(sorted(cv_mse_scores.items()), 1):
                if i < len(cv_table.rows):  # 确保不超出表格行数
                    cv_table.rows[i].cells[0].text = str(degree)
                    cv_table.rows[i].cells[1].text = f"{score:.4f}"

        # 添加统计结果
        doc.add_heading(LANGUAGES[self.current_language]['statistical_results'], level=2)

        # 筛选出统计结果行
        stats_only_df = transposed_df[transposed_df[LANGUAGES[self.current_language]['metric']].apply(
            lambda x: x in ['intercept', 'mse', 'r_squared', 'adjusted_r_squared', 'f_value'] or
                      x.startswith(('coefficients', 't_value', 'p_value'))
        )]

        if len(stats_only_df) > 3:
            stats_only_df = stats_only_df.iloc[:-3, :]

        # 只保留前两列
        stats_only_df = stats_only_df.iloc[:, :2]

        # 创建统计结果表格
        stats_table = doc.add_table(rows=stats_only_df.shape[0] + 1, cols=stats_only_df.shape[1])
        hdr_cells = stats_table.rows[0].cells

        # 设置表头为当前语言
        hdr_cells[0].text = LANGUAGES[self.current_language]['metric']  # 指标列
        hdr_cells[1].text = LANGUAGES[self.current_language]['best_degree']  # 最佳阶数列

        # 填充表格内容
        for row_idx, row in enumerate(stats_only_df.itertuples(index=False)):
            table_row_idx = row_idx + 1
            if table_row_idx < len(stats_table.rows):
                row_cells = stats_table.rows[table_row_idx].cells

                # 获取原始指标名
                original_metric = row[0]

                # 解析指标名和特征名（处理带括号的情况，如"coefficients (x)"）
                if '(' in original_metric and ')' in original_metric:
                    metric_base = original_metric.split('(')[0].strip()
                    feature_name = original_metric.split('(')[1].rstrip(')')
                    # 使用带占位符的格式字符串
                    display_name = LANGUAGES[self.current_language]['metrics'].get(metric_base, original_metric).format(
                        feature_name)
                else:
                    # 直接使用语言字典中的对应名称
                    display_name = LANGUAGES[self.current_language]['metrics'].get(original_metric, original_metric)

                # 设置指标名称
                row_cells[0].text = display_name

                # 设置指标值
                value = row[1]
                if isinstance(value, (int, float)):
                    row_cells[1].text = f"{value:.4f}"
                else:
                    row_cells[1].text = str(value)

        # 在_generate_report方法中，找到添加解释说明的部分，替换为以下代码
        doc.add_heading(LANGUAGES[self.current_language]['explanation_label'], level=2)
        explanations = LANGUAGES[self.current_language]['explanation']

        # 创建项目符号列表
        explanation_para = doc.add_paragraph()
        for key, value in explanations.items():
            # 对于中文，使用对应的中文键名；对于英文，保持原键名
            if self.current_language == 'zh':
                # 定义英文键到中文键的映射
                key_mapping = {
                    "coefficients": "回归系数",
                    "intercept": "截距",
                    "mse": "均方误差",
                    "r_squared": "决定系数",
                    "adjusted_r_squared": "调整决定系数",
                    "f_value": "F统计量",
                    "t_value": "t统计量",
                    "p_value": "p值",
                    "regression_equation": "回归方程"
                }
                display_key = key_mapping.get(key, key)  # 使用映射的中文键名
            else:
                display_key = key  # 英文保持原键名
            run = explanation_para.add_run(f"• {display_key}: {value}\n")
            run.bold = False

        # 获取保存目录
        save_dir = os.path.dirname(save_path)

        # 实际值 vs 预测值散点图
        plt.figure(figsize=(10, 6))
        plt.scatter(y, y_pred, alpha=0.6)
        plt.plot([y.min(), y.max()], [y.min(), y.max()], 'r--', lw=2)
        plt.xlabel(LANGUAGES[self.current_language]['actual_values'])
        plt.ylabel(LANGUAGES[self.current_language]['predicted_values'])
        plt.title(LANGUAGES[self.current_language]['actual_vs_predicted'])
        img1 = os.path.join(save_dir, "actual_vs_predicted.png")
        plt.savefig(img1)
        plt.close()
        doc.add_heading(LANGUAGES[self.current_language]['actual_vs_predicted'], level=2)
        doc.add_picture(img1, width=Inches(6))

        # 残差图
        plt.figure(figsize=(10, 6))
        residuals = y - y_pred
        plt.scatter(y_pred, residuals, alpha=0.6)
        plt.axhline(y=0, color='r', linestyle='--')
        plt.xlabel(LANGUAGES[self.current_language]['predicted_values'])
        plt.ylabel(LANGUAGES[self.current_language]['residuals'])
        plt.title(LANGUAGES[self.current_language]['residual_plot'])
        img2 = os.path.join(save_dir, "residuals.png")
        plt.savefig(img2)
        plt.close()
        doc.add_heading(LANGUAGES[self.current_language]['residual_plot'], level=2)
        doc.add_picture(img2, width=Inches(6))

        # 交叉验证MSE对比图
        plt.figure(figsize=(10, 6))
        degrees = list(cv_mse_scores.keys())
        scores = list(cv_mse_scores.values())
        plt.plot(degrees, scores, 'bo-')
        plt.xlabel(LANGUAGES[self.current_language]['polynomial_degree'])
        plt.ylabel(LANGUAGES[self.current_language]['cross_validation_mse'])
        plt.title(LANGUAGES[self.current_language]['different_degrees_comparison'])
        plt.xticks(degrees)
        img4 = os.path.join(save_dir, "cv_mse_comparison.png")
        plt.savefig(img4)
        plt.close()
        doc.add_heading(LANGUAGES[self.current_language]['cv_mse_comparison'], level=2)
        doc.add_picture(img4, width=Inches(6))

        # 自变量与因变量的拟合曲线（仅对单个自变量绘制）
        if X.shape[1] == 1:
            plt.figure(figsize=(10, 6))
            x_range = np.linspace(X.min(), X.max(), 100).reshape(-1, 1)
            x_poly_range = poly.transform(x_range)
            y_pred_range = model.predict(x_poly_range)

            plt.scatter(X, y, alpha=0.6, label=LANGUAGES[self.current_language]['original_data'])
            plt.plot(x_range, y_pred_range, 'r-', lw=2,
                     label=f'{poly.degree}{LANGUAGES[self.current_language]["fit_curve"]}')
            plt.xlabel(df.columns[:-1][0])
            plt.ylabel(df.columns[-1])
            plt.title(LANGUAGES[self.current_language]['independent_vs_dependent'])
            plt.legend()
            img3 = os.path.join(save_dir, "variable_relationship.png")
            plt.savefig(img3)
            plt.close()
            doc.add_heading(LANGUAGES[self.current_language]['variable_relationship'], level=2)
            doc.add_picture(img3, width=Inches(6))

        doc.save(save_path)

    def find_best_degree(self, X, y, degrees=range(1, 6)):
        """使用交叉验证找到最佳多项式阶数"""
        cv_mse_scores = {}

        for degree in degrees:
            poly = PolynomialFeatures(degree=degree)
            X_poly = poly.fit_transform(X)

            # 使用5折交叉验证计算MSE
            model = LinearRegression()
            scores = cross_val_score(
                model, X_poly, y,
                cv=5,
                scoring='neg_mean_squared_error'
            )

            # 转换为MSE（取平均值并取相反数）
            mse = -np.mean(scores)
            cv_mse_scores[degree] = mse

        # 选择MSE最小的阶数作为最佳阶数
        best_degree = min(cv_mse_scores, key=cv_mse_scores.get)
        return best_degree, cv_mse_scores

    def generate_regression_equation(self, intercept, coefficients, feature_names, poly):
        """生成多项式回归方程的字符串表示（修复运算符显示问题）"""
        # 获取多项式特征名称
        poly_feature_names = poly.get_feature_names_out(feature_names)

        # 构建方程字符串
        eq_parts = []
        # 添加截距项（保留4位小数）
        eq_parts.append(f"{intercept:.4f}")

        # 添加各个特征项
        for i, name in enumerate(poly_feature_names):
            if i < len(coefficients):  # 确保索引有效
                coef = coefficients[i]
                # 处理系数符号和格式
                if coef > 0:
                    # 正数前加+号（除了第一个特征项，因为截距项已经是开头）
                    eq_parts.append(f"+ {coef:.4f} × {name}")
                elif coef < 0:
                    # 负数前加-号，并显示绝对值
                    eq_parts.append(f"- {abs(coef):.4f} × {name}")
                else:
                    # 系数为0时省略该项
                    continue

        # 组合成完整方程（使用因变量名称）
        dv_name = "y"  # 可根据实际需求替换为因变量列名
        equation = f"{dv_name} = " + " ".join(eq_parts)

        return equation

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return

        try:
            # 读取数据
            df = pd.read_excel(file_path)

            # 数据校验
            is_valid, msg = self.validate_data(df)
            if not is_valid:
                self.result_label.config(text=f"{LANGUAGES[self.current_language]['data_error']} {msg}")
                return

            # 显示交叉验证进度
            self.result_label.config(text=LANGUAGES[self.current_language]['cross_val_progress'])
            self.root.update_idletasks()  # 更新UI显示

            # 自动将最后一列设为因变量
            dv_idx = len(df.columns) - 1  # 最后一列的索引

            # 划分自变量和因变量
            X = df.drop(df.columns[dv_idx], axis=1).values
            y = df.iloc[:, dv_idx].values
            feature_names = df.drop(df.columns[dv_idx], axis=1).columns.tolist()
            dv_name = df.columns[dv_idx]

            # 使用交叉验证确定最佳多项式阶数（1-5阶）
            best_degree, cv_mse_scores = self.find_best_degree(X, y, degrees=range(1, 6))

            # 使用最佳阶数进行多项式回归分析
            poly = PolynomialFeatures(degree=best_degree)
            X_poly = poly.fit_transform(X)
            model = LinearRegression()
            model.fit(X_poly, y)
            y_pred = model.predict(X_poly)

            # 计算统计指标
            coefficients = model.coef_
            intercept = model.intercept_
            mse = mean_squared_error(y, y_pred)
            r2 = r2_score(y, y_pred)
            n = len(y)
            p = X_poly.shape[1]
            adjusted_r2 = 1 - (1 - r2) * (n - 1) / (n - p - 1)

            # 计算t值、p值和F值
            X_with_const = sm.add_constant(X_poly)
            sm_model = sm.OLS(y, X_with_const).fit()
            t_values = sm_model.tvalues
            p_values = sm_model.pvalues
            f_value = sm_model.fvalue

            # 生成多项式回归方程
            regression_eq = self.generate_regression_equation(intercept, coefficients, feature_names, poly)

            # 生成特征名称（包含多项式项）
            poly_feature_names = poly.get_feature_names_out(feature_names)

            # 准备结果数据
            model_data = {
                "Model": "Polynomial Regression",
                "intercept": intercept,
                "mse": mse,
                "r_squared": r2,
                "adjusted_r_squared": adjusted_r2,
                "f_value": f_value,
                "regression_equation": regression_eq
            }

            # 添加系数、t值和p值
            for i, name in enumerate(poly_feature_names):
                # 检查索引是否在有效范围内
                if i < len(coefficients) and (i + 1) < len(t_values) and (i + 1) < len(p_values):
                    model_data[f"coefficients ({name})"] = coefficients[i]
                    model_data[f"t_value ({name})"] = t_values[i + 1]  # +1因为t_values包含常数项
                    model_data[f"p_value ({name})"] = p_values[i + 1]

            # 准备解释说明
            explanations = LANGUAGES[self.current_language]['explanation']
            explanation_data = {"Model": LANGUAGES[self.current_language]['explanation_label']}
            for key, value in explanations.items():
                explanation_data[key] = value

            # 创建结果数据框
            df_result = pd.DataFrame([model_data])
            explanation_df = pd.DataFrame([explanation_data])

            # 合并并转置
            combined_df = pd.concat([df_result, explanation_df], ignore_index=True)
            transposed_df = combined_df.set_index('Model').T.reset_index().rename(
                columns={'index': LANGUAGES[self.current_language]['metric']})

            # 保存分析结果（包含回归方程）
            self.analysis_results = (df, X, y, X_poly, poly, model, y_pred, coefficients, intercept, mse, r2,
                                     adjusted_r2, t_values, p_values, f_value, transposed_df, feature_names,
                                     cv_mse_scores, regression_eq)

            # 询问用户是否保存
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                self._generate_report(save_path, self.analysis_results)
                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(save_path)
                result_msg += LANGUAGES[self.current_language]['images_saved'].format(os.path.dirname(save_path))
                self.result_label.config(text=result_msg, wraplength=400)
            else:
                # 提供二次保存机会
                answer = simpledialog.askstring(
                    LANGUAGES[self.current_language]['save_results'],
                    f"{LANGUAGES[self.current_language]['no_save_path']} {LANGUAGES[self.current_language]['save_results']}",
                    parent=self.root,
                    initialvalue=LANGUAGES[self.current_language]['save_now']
                )
                if answer and LANGUAGES[self.current_language]['save_now'].lower() in answer.lower():
                    success, path = self.save_results(self.analysis_results)
                    if success:
                        self.result_label.config(text=LANGUAGES[self.current_language]['analysis_success'].format(path),
                                                 wraplength=400)
                    else:
                        self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])
                else:
                    self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(LANGUAGES[self.current_language]["title"])
        self.select_button.config(text=LANGUAGES[self.current_language]["select_button"])
        current_text = self.file_entry.get()
        if current_text == LANGUAGES['zh']["file_entry_placeholder"] or current_text == LANGUAGES['en'][
            "file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=LANGUAGES[self.current_language]["analyze_button"])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]["switch_language"])
        self.open_excel_label.config(text=LANGUAGES[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        self.select_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["select_button"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        self.analyze_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["analyze_button"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=LANGUAGES[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        self.switch_language_label = ttk.Label(frame, text=LANGUAGES[self.current_language]["switch_language"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        self.root.mainloop()


def run_app():
    app = PolynomialRegressionAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()
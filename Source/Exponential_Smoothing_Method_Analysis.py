import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
import matplotlib.pyplot as plt
import pathlib
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典，所有键使用英文
LANGUAGES = {
    'zh': {
        'title': "指数平滑法",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "请选择文件。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "中/英",
        'open_excel_button_text': "示例数据",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'alpha_label': "平滑系数 (0-1):",
        'alpha_error': "平滑系数必须是0到1之间的数字",
        'data_error': "数据校验失败: {}",
        'forecast_periods_label': "预测未来时间点数量:",
        'forecast_periods_error': "预测数量必须是正整数",
        'forecast_values_heading': "预测值",
        'explanation': {
            "original_data": "输入的待分析数据",
            "first_order_smooth": "通过指数平滑法计算得到的一次平滑值序列",
            "second_order_smooth": "通过指数平滑法计算得到的二次平滑值序列",
            "forecast_values": "基于二次指数平滑值得到的预测值（未来多个时间点）",
            "forecast_chart": "展示原始数据和预测值的折线图"
        },
        'interpretation': {
            "original_data": "作为分析的基础数据",
            "first_order_smooth": "反映数据的平滑趋势",
            "second_order_smooth": "进一步平滑并捕捉趋势分量",
            "forecast_values": "反映未来趋势的预测结果（包含趋势变化）",
            "forecast_chart": "直观展示原始数据和预测值的变化趋势"
        },
        # 新增图表和文档相关文本
        'time': "时间",
        'value': "数值",
        'smoothed_value': "平滑值",
        'forecast_period': "预测时间点",
        'forecast_value': "预测值",
        'statistic': "统计量",
        'explanation_title': "解释说明",
        'interpretation_title': "结果解读",
        'analysis_results': "二次指数平滑法分析结果",
        'original_data_heading': "原始数据",
        'first_order_smooth_heading': "一次指数平滑值",
        'second_order_smooth_heading': "二次指数平滑值",
        'forecast_results_heading': "预测结果",
        'results_explanation_heading': "结果解释",
        'forecast_chart_heading': "预测结果折线图",
        'original_data_label': "原始数据",
        'first_order_smooth_label': "一次指数平滑值",
        'second_order_smooth_label': "二次指数平滑值",
        'forecast_values_label': "预测值",
        'chart_title': "预测结果折线图",
        'x_axis_label': "时间",
        'y_axis_label': "值",
        'future_period': "未来第{}期"
    },
    'en': {
        'title': "Exponential Smoothing Method",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Chinese/English",
        'open_excel_button_text': "Sample Data",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'alpha_label': "Smoothing factor (0-1):",
        'alpha_error': "Smoothing factor must be a number between 0 and 1",
        'data_error': "Data validation failed: {}",
        'forecast_periods_label': "Number of future forecast periods:",
        'forecast_periods_error': "Forecast periods must be a positive integer",
        'forecast_values_heading': "Forecast Values",
        'explanation': {
            "original_data": "The input data to be analyzed",
            "first_order_smooth": "The first-order exponentially smoothed value sequence",
            "second_order_smooth": "The second-order exponentially smoothed value sequence",
            "forecast_values": "Predicted values (multiple future periods) based on second-order exponential smoothing",
            "forecast_chart": "A line chart showing the original data and predicted values"
        },
        'interpretation': {
            "original_data": "As the basic data for analysis",
            "first_order_smooth": "Reflects the smoothing trend of the data",
            "second_order_smooth": "Further smooths and captures trend components",
            "forecast_values": "Predicted results reflecting future trends (including trend changes)",
            "forecast_chart": "Visually display the changing trends of the original data and predicted values"
        },
        # 新增图表和文档相关文本
        'time': "Time",
        'value': "Value",
        'smoothed_value': "Smoothed Value",
        'forecast_period': "Forecast Period",
        'forecast_value': "Forecast Value",
        'statistic': "Statistic",
        'explanation_title': "Explanation",
        'interpretation_title': "Interpretation",
        'analysis_results': "Second-order Exponential Smoothing Analysis Results",
        'original_data_heading': "Original Data",
        'first_order_smooth_heading': "First-order Smoothed Values",
        'second_order_smooth_heading': "Second-order Smoothed Values",
        'forecast_results_heading': "Forecast Results",
        'results_explanation_heading': "Results Explanation",
        'forecast_chart_heading': "Forecast Results Chart",
        'original_data_label': "Original Data",
        'first_order_smooth_label': "First-order Smoothed Values",
        'second_order_smooth_label': "Second-order Smoothed Values",
        'forecast_values_label': "Forecast Value",
        'chart_title': "Line Chart of Prediction Results",
        'x_axis_label': "Time",
        'y_axis_label': "Value",
        'future_period': "Future Period {}"
    }
}


class ExponentialSmoothingMethodAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为中文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data54.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def exponential_smoothing(self, x, alpha):
        """
        二次指数平滑法
        :param x: 原始数据序列
        :param alpha: 平滑系数
        :return: 一次平滑值序列、二次平滑值序列
        """
        # 一次指数平滑
        s1 = [x[0]]
        for i in range(1, len(x)):
            s1.append(alpha * x[i] + (1 - alpha) * s1[i - 1])

        # 二次指数平滑（对一次平滑结果再做一次平滑）
        s2 = [s1[0]]
        for i in range(1, len(s1)):
            s2.append(alpha * s1[i] + (1 - alpha) * s2[i - 1])

        return s1, s2

    def validate_alpha(self):
        """验证平滑系数是否有效"""
        try:
            alpha = float(self.alpha_entry.get())
            if 0 <= alpha <= 1:
                return alpha
            return None
        except ValueError:
            return None

    def validate_forecast_periods(self):
        """验证预测时间点数量是否有效"""
        try:
            periods = int(self.forecast_periods_entry.get())
            if periods > 0:
                return periods
            return None
        except ValueError:
            return None

    def validate_data(self, data):
        """数据校验：检查是否有缺失值和非数值"""
        if pd.isnull(data).any():
            return False, "存在缺失值，请处理后再分析" if self.current_language == 'zh' else "There are missing values, please handle them before analysis"

        if not np.issubdtype(data.dtype, np.number):
            return False, "数据包含非数值类型，请检查" if self.current_language == 'zh' else "Data contains non-numeric types, please check"

        if len(data) < 2:
            return False, "数据量过少，无法进行分析" if self.current_language == 'zh' else "Insufficient data for analysis"

        return True, ""

    def analyze_file(self):
        # 验证平滑系数
        alpha = self.validate_alpha()
        if alpha is None:
            self.result_label.config(text=LANGUAGES[self.current_language]['alpha_error'])
            return

        # 验证预测时间点数量
        forecast_periods = self.validate_forecast_periods()
        if forecast_periods is None:
            self.result_label.config(text=LANGUAGES[self.current_language]['forecast_periods_error'])
            return

        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            file_path = ""
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 读取Excel文件，不将第一行作为表头（header=None）
            df = pd.read_excel(file_path, header=None)

            # 检查是否有至少一列数据
            if len(df.columns) >= 2:
                time_data = df.iloc[:, 0].values
                value_data = df.iloc[:, 1].values
            else:
                # 如果只有一列，使用默认索引作为时间
                value_data = df.iloc[:, 0].values
                time_data = np.arange(1, len(value_data) + 1)

            # 数据校验
            is_valid, msg = self.validate_data(value_data)
            if not is_valid:
                self.result_label.config(text=f"{LANGUAGES[self.current_language]['data_error']}\n{msg}")
                return

            # 进行二次指数平滑分析
            s1, s2 = self.exponential_smoothing(value_data, alpha)

            # 计算预测值（二次指数平滑预测公式）
            forecast_values = []
            n = len(s1)
            # 计算预测公式中的两个参数
            a_t = 2 * s1[-1] - s2[-1]  # 截距项
            b_t = (alpha / (1 - alpha)) * (s1[-1] - s2[-1])  # 趋势项

            # 预测未来k期的值：y(t+k) = a_t + b_t * k
            for k in range(1, forecast_periods + 1):
                forecast = a_t + b_t * k
                forecast_values.append(forecast)

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加标题
                doc.add_heading(LANGUAGES[self.current_language]['analysis_results'], 0)
                doc.add_paragraph(
                    f"平滑系数 α = {alpha}" if self.current_language == 'zh' else f"Smoothing factor α = {alpha}")
                doc.add_paragraph(
                    f"预测未来时间点数量 = {forecast_periods}" if self.current_language == 'zh' else f"Number of future forecast periods = {forecast_periods}")

                # 添加原始数据表格
                doc.add_heading(LANGUAGES[self.current_language]['original_data_heading'], level=1)
                data_table = doc.add_table(rows=1, cols=2)
                data_hdr = data_table.rows[0].cells
                data_hdr[0].text = LANGUAGES[self.current_language]['time']
                data_hdr[1].text = LANGUAGES[self.current_language]['value']
                for t, v in zip(time_data, value_data):
                    row_cells = data_table.add_row().cells
                    row_cells[0].text = str(t)
                    row_cells[1].text = str(v)

                # 添加一次平滑值表格
                doc.add_heading(LANGUAGES[self.current_language]['first_order_smooth_heading'], level=1)
                smooth_table1 = doc.add_table(rows=1, cols=2)
                smooth_hdr1 = smooth_table1.rows[0].cells
                smooth_hdr1[0].text = LANGUAGES[self.current_language]['time']
                smooth_hdr1[1].text = LANGUAGES[self.current_language]['smoothed_value']
                for t, s in zip(time_data, s1):
                    row_cells = smooth_table1.add_row().cells
                    row_cells[0].text = str(t)
                    row_cells[1].text = f"{s:.4f}"

                # 添加二次平滑值表格
                doc.add_heading(LANGUAGES[self.current_language]['second_order_smooth_heading'], level=1)
                smooth_table2 = doc.add_table(rows=1, cols=2)
                smooth_hdr2 = smooth_table2.rows[0].cells
                smooth_hdr2[0].text = LANGUAGES[self.current_language]['time']
                smooth_hdr2[1].text = LANGUAGES[self.current_language]['smoothed_value']
                for t, s in zip(time_data, s2):
                    row_cells = smooth_table2.add_row().cells
                    row_cells[0].text = str(t)
                    row_cells[1].text = f"{s:.4f}"

                # 添加预测值
                doc.add_heading(LANGUAGES[self.current_language]['forecast_results_heading'], level=1)
                forecast_table = doc.add_table(rows=1, cols=2)
                forecast_hdr = forecast_table.rows[0].cells
                forecast_hdr[0].text = LANGUAGES[self.current_language]['forecast_period']
                forecast_hdr[1].text = LANGUAGES[self.current_language]['forecast_value']

                # 生成预测时间点
                forecast_times = []
                if isinstance(time_data[-1], (int, float)):
                    for i in range(1, forecast_periods + 1):
                        forecast_times.append(time_data[-1] + i)
                else:
                    for i in range(1, forecast_periods + 1):
                        forecast_times.append(LANGUAGES[self.current_language]['future_period'].format(i))

                for t, v in zip(forecast_times, forecast_values):
                    row_cells = forecast_table.add_row().cells
                    row_cells[0].text = str(t)
                    row_cells[1].text = f"{v:.4f}"

                # 添加解释说明
                doc.add_heading(LANGUAGES[self.current_language]['results_explanation_heading'], level=1)
                expl_table = doc.add_table(rows=1, cols=2)
                expl_hdr = expl_table.rows[0].cells
                expl_hdr[0].text = LANGUAGES[self.current_language]['statistic']
                expl_hdr[1].text = LANGUAGES[self.current_language]['explanation_title']

                # 使用英文键名访问解释说明
                for key in ["original_data", "first_order_smooth", "second_order_smooth", "forecast_values"]:
                    row_cells = expl_table.add_row().cells
                    # 根据当前语言获取对应统计量的显示名称
                    row_cells[0].text = LANGUAGES[self.current_language][key + "_heading"]
                    row_cells[1].text = LANGUAGES[self.current_language]['explanation'][key]

                # 生成预测结果折线图
                plt.figure(figsize=(10, 6))
                plt.plot(time_data, value_data, label=LANGUAGES[self.current_language]['original_data_label'],
                         marker='o')
                plt.plot(time_data, s1,
                         label=LANGUAGES[self.current_language]['first_order_smooth_label'],
                         linestyle='-', marker='s')
                plt.plot(time_data, s2,
                         label=LANGUAGES[self.current_language]['second_order_smooth_label'],
                         linestyle='-.', marker='^')

                # 绘制预测值点
                plt.scatter(forecast_times, forecast_values, color='red', s=100,
                            label=LANGUAGES[self.current_language]['forecast_values_label'])
                # 连接最后一个实际值与第一个预测值，使图表更连贯
                if isinstance(time_data[-1], (int, float)) and isinstance(forecast_times[0], (int, float)):
                    plt.plot([time_data[-1], forecast_times[0]], [value_data[-1], forecast_values[0]],
                             color='gray', linestyle='--')

                plt.title(LANGUAGES[self.current_language]['chart_title'])
                plt.xlabel(LANGUAGES[self.current_language]['x_axis_label'])
                plt.ylabel(LANGUAGES[self.current_language]['y_axis_label'])
                plt.xticks(rotation=30)
                plt.tight_layout()
                plt.legend()

                # 保存图片
                img_path = os.path.splitext(save_path)[0] + '_prediction_chart.png'
                plt.savefig(img_path, dpi=300)
                plt.close()

                # 在 Word 文档中插入图片
                doc.add_heading(LANGUAGES[self.current_language]['forecast_chart_heading'], level=1)
                doc.add_picture(img_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(save_path)
                self.result_label.config(text=result_msg)

            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'
        self.root.title(LANGUAGES[self.current_language]['title'])
        self.select_button.config(text=LANGUAGES[self.current_language]['select_button'])
        self.analyze_button.config(text=LANGUAGES[self.current_language]['analyze_button'])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]['switch_language'])
        self.alpha_label.config(text=LANGUAGES[self.current_language]['alpha_label'])
        self.forecast_periods_label.config(text=LANGUAGES[self.current_language]['forecast_periods_label'])
        self.open_excel_label.config(text=LANGUAGES[self.current_language]["open_excel_button_text"])

        # 切换语言时更新提示信息
        current_text = self.file_entry.get()
        if current_text == LANGUAGES['zh' if self.current_language == 'en' else 'en']["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
            self.file_entry.config(foreground='gray')

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建主框架，使用grid布局让内容居中
        main_frame = ttk.Frame(self.root)
        main_frame.pack(expand=True, fill=BOTH)

        # 创建一个居中的框架
        center_frame = ttk.Frame(main_frame)
        center_frame.place(relx=0.5, rely=0.5, anchor=CENTER)

        # 创建文件选择按钮
        self.select_button = ttk.Button(center_frame, text=LANGUAGES[self.current_language]["select_button"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(center_frame, width=50)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 平滑系数输入
        alpha_frame = ttk.Frame(center_frame)
        alpha_frame.pack(pady=5, fill=X)

        self.alpha_label = ttk.Label(alpha_frame, text=LANGUAGES[self.current_language]["alpha_label"])
        self.alpha_label.pack(side=LEFT, padx=5)

        self.alpha_entry = ttk.Entry(alpha_frame, width=10)
        self.alpha_entry.insert(0, "0.3")  # 默认值
        self.alpha_entry.pack(side=LEFT)

        # 预测时间点数量输入
        forecast_periods_frame = ttk.Frame(center_frame)
        forecast_periods_frame.pack(pady=5, fill=X)

        self.forecast_periods_label = ttk.Label(forecast_periods_frame,
                                                text=LANGUAGES[self.current_language]["forecast_periods_label"])
        self.forecast_periods_label.pack(side=LEFT, padx=5)

        self.forecast_periods_entry = ttk.Entry(forecast_periods_frame, width=10)
        self.forecast_periods_entry.insert(0, "1")  # 默认预测1个时间点
        self.forecast_periods_entry.pack(side=LEFT)

        # 创建分析按钮
        self.analyze_button = ttk.Button(center_frame, text=LANGUAGES[self.current_language]["analyze_button"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            center_frame,
            text=LANGUAGES[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(center_frame, text=LANGUAGES[self.current_language]["switch_language"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=5)

        # 创建结果显示标签
        self.result_label = ttk.Label(center_frame, text="", wraplength=400, justify=CENTER)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = ExponentialSmoothingMethodAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()
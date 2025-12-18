import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import openpyxl
import os
import numpy as np
import statsmodels.formula.api as smf
import pandas as pd
from tkinter import filedialog
import tkinter as tk
from sklearn.metrics import mean_squared_error, r2_score
import matplotlib.pyplot as plt
import statsmodels.api as sm
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# 设置支持中文的字体
plt.rcParams["font.sans-serif"] = ["Arial Unicode MS", "SimHei", "Microsoft YaHei"]
plt.rcParams["axes.unicode_minus"] = False  # 解决负号显示问题

# 定义语言字典，所有键使用英文
LANGUAGES = {
    'zh': {
        'title': "普通最小二乘线性回归分析",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "请选择文件。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'images_saved': "图片已保存到 {}",
        'switch_language': "中/英",
        'open_excel_button_text': "示例数据",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'regression_equation': "回归方程:",
        'significance_test': "显著性检验结论:",
        'significant_model': "回归模型整体显著 (p值: {:.4f})",
        'not_significant_model': "回归模型整体不显著 (p值: {:.4f})",
        'significant_variable': "显著影响因素 (p < 0.05):",
        'no_significant_variable': "无显著影响因素 (p < 0.05)",
        'var_analysis_results': "变量分析结果",
        'model_statistics': "模型统计量",
        'index_explanation': "指标解释说明",
        'anova_table': "方差分析表 (ANOVA)",
        'anova_explanation': """方差分析表解释：
                - df: 自由度
                - sum_sq: 平方和
                - mean_sq: 均方
                - F: F统计量
                - PR(>F): F检验的p值，用于判断自变量对因变量的整体影响是否显著""",
        'residual_analysis': "残差分析",
        'residual_hist_desc': "残差分布图显示了模型残差的分布情况，理想情况下应接近正态分布：",
        'residual_scatter_desc': "残差与拟合值散点图用于检测残差是否具有同方差性：",
        'residual_homoscedasticity': "若点随机分布在y=0水平线附近且无明显模式，则说明满足同方差性假设。",
        'actual_vs_predicted': "实际值与预测值对比",
        'var_name': "变量名称",
        'coefficient': "系数 (Coefficient)",
        't_value': "t值 (t-value)",
        'p_value': "p值 (p-value)",
        'ci_lower': "95%置信区间下限",
        'ci_upper': "95%置信区间上限",
        'statistic': "统计量",
        'value': "值",
        'source': "来源",
        'intercept': "截距 (Intercept)",
        'mse': "均方误差 (MSE)",
        'r_squared': "决定系数 (R²)",
        'adjusted_r_squared': "调整决定系数",
        'f_statistic': "F统计量",
        'f_pvalue': "F统计量p值",
        'residuals': "残差 (Residuals)",
        'frequency': "频率 (Frequency)",
        'residual_hist_title': "残差分布图",
        'fitted_values': "拟合值 (Fitted Values)",
        'residual_scatter_title': "残差与拟合值散点图",
        'actual_value': "实际值 ({})",
        'predicted_value': "预测值 ({})",
        'actual_vs_predicted_title': "实际值 vs 预测值",
        'explanation': {
            "Coefficients": "回归系数，表示每个自变量对因变量的影响程度。",
            "Intercept": "截距，是当所有自变量为 0 时因变量的预测值。",
            "Mean Squared Error (MSE)": "均方误差，衡量预测值与真实值之间的平均误差。",
            "R-squared (R²)": "决定系数，取值范围在 0 到 1 之间，越接近 1 表示模型拟合效果越好。",
            "Adjusted R-squared": "调整决定系数，考虑了模型中自变量的数量，对模型的拟合优度进行了调整。",
            "F-value": "F 统计量，用于检验整个回归模型的显著性。",
            "t-value": "t 统计量，用于检验每个自变量的显著性。",
            "p-value": "p 值，用于判断自变量的显著性，p 值越小，自变量越显著。通常 p < 0.05 认为显著。"
        }
    },
    'en': {
        'title': "Ordinary Least Squares Linear Regression Analysis",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "Please select a file.",
        'analysis_success': "Analysis completed. Results saved to {}\n",
        'no_save_path': "No save path selected. Results not saved.",
        'analysis_error': "Error analyzing file: {}",
        'images_saved': "Images saved to {}",
        'switch_language': "Chinese/English",
        'open_excel_button_text': "Example Data",
        'file_entry_placeholder': "Please enter the full path of the Excel file to analyze",
        'regression_equation': "Regression Equation:",
        'significance_test': "Significance Test Conclusion:",
        'significant_model': "The regression model is overall significant (p-value: {:.4f})",
        'not_significant_model': "The regression model is not overall significant (p-value: {:.4f})",
        'significant_variable': "Significant factors (p < 0.05):",
        'no_significant_variable': "No significant factors (p < 0.05)",
        'var_analysis_results': "Variable Analysis Results",
        'model_statistics': "Model Statistics",
        'index_explanation': "Index Explanation",
        'anova_table': "Analysis of Variance (ANOVA)",
        'anova_explanation': """ANOVA Table Explanation:
                - df: Degrees of Freedom
                - sum_sq: Sum of Squares
                - mean_sq: Mean Squares
                - F: F-statistic
                - PR(>F): p-value for F-test, used to determine if independent variables significantly affect the dependent variable""",
        'residual_analysis': "Residual Analysis",
        'residual_hist_desc': "Residual histogram shows the distribution of model residuals, which should ideally be approximately normal:",
        'residual_scatter_desc': "Residual vs fitted values scatter plot is used to check for homoscedasticity:",
        'residual_homoscedasticity': "If points are randomly distributed around y=0 with no obvious pattern, the homoscedasticity assumption is satisfied.",
        'actual_vs_predicted': "Actual vs Predicted Values",
        'var_name': "Variable Name",
        'coefficient': "Coefficient",
        't_value': "t-value",
        'p_value': "p-value",
        'ci_lower': "95% CI Lower",
        'ci_upper': "95% CI Upper",
        'statistic': "Statistic",
        'value': "Value",
        'source': "Source",
        'intercept': "Intercept",
        'mse': "Mean Squared Error (MSE)",
        'r_squared': "R-squared (R²)",
        'adjusted_r_squared': "Adjusted R-squared",
        'f_statistic': "F-statistic",
        'f_pvalue': "F-statistic p-value",
        'residuals': "Residuals",
        'frequency': "Frequency",
        'residual_hist_title': "Residual Histogram",
        'fitted_values': "Fitted Values",
        'residual_scatter_title': "Residuals vs Fitted Values",
        'actual_value': "Actual Value ({})",
        'predicted_value': "Predicted Value ({})",
        'actual_vs_predicted_title': "Actual vs Predicted Values",
        'explanation': {
            "Coefficients": "Regression coefficients, indicating the influence of each independent variable on the dependent variable.",
            "Intercept": "Intercept, which is the predicted value of the dependent variable when all independent variables are 0.",
            "Mean Squared Error (MSE)": "Mean squared error, measuring the average error between the predicted and actual values.",
            "R-squared (R²)": "Coefficient of determination, ranging from 0 to 1. A value closer to 1 indicates a better fit of the model.",
            "Adjusted R-squared": "Adjusted coefficient of determination, which takes into account the number of independent variables in the model and adjusts the goodness of fit of the model.",
            "F-value": "F statistic, used to test the significance of the entire regression model.",
            "t-value": "t statistic, used to test the significance of each independent variable.",
            "p-value": "p value, used to determine the significance of the independent variable. The smaller the p value, the more significant the independent variable. Usually p < 0.05 is considered significant."
        }
    }
}


class OrdinaryLeastSquaresLinearRegressionAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为中文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data32.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件，第一行为表头
            df = pd.read_excel(file_path)

            # 获取因素名称（第一列）和因变量名称（最后一列的表头）
            independent_vars = df.columns[1:-1].tolist()  # 自变量列名（排除第一列和最后一列）
            dependent_var = df.columns[-1]  # 因变量列名（最后一列）

            # 准备自变量和因变量数据（排除第一列因素列）
            X = df.iloc[:, 1:-1].values
            y = df.iloc[:, -1].values
            var_names = df.columns[1:-1].tolist()

            # 添加常数项
            X_with_const = sm.add_constant(X)

            # 进行普通最小二乘回归分析
            model = sm.OLS(y, X_with_const).fit()
            # 模型拟合检查
            if model.mse_resid < 0:  # 均方误差不可能为负
                raise ValueError("模型拟合错误：均方误差为负")
            if len(model.params) != X_with_const.shape[1]:  # 系数数量应与自变量+常数项数量一致
                raise ValueError("模型系数数量与自变量不匹配")
            y_pred = model.predict(X_with_const)

            # ANOVA代码
            formula = f"{dependent_var} ~ {' + '.join(independent_vars)}"
            model = smf.ols(formula=formula, data=df).fit()
            anova_table = sm.stats.anova_lm(model)

            # 置信区间
            conf_int = model.conf_int()  # 获取95%置信区间

            # 残差分析
            residuals = model.resid  # 获取残差

            # 计算指标
            coefficients = model.params[1:]  # 排除常数项
            intercept = model.params.iloc[0]
            mse = mean_squared_error(y, y_pred)
            r2 = model.rsquared
            adjusted_r2 = model.rsquared_adj
            f_value = model.fvalue
            f_pvalue = model.f_pvalue  # 模型整体显著性p值
            t_values = model.tvalues[1:]  # 排除常数项
            p_values = model.pvalues[1:]  # 排除常数项

            # 构建回归方程
            equation_terms = [f"{intercept:.4f}"]
            for i, var_name in enumerate(independent_vars):
                coef = coefficients.iloc[i]
                if coef >= 0:
                    equation_terms.append(f"+ {coef:.4f}*{var_name}")
                else:
                    equation_terms.append(f"- {abs(coef):.4f}*{var_name}")
            regression_equation = f"{dependent_var} = {' '.join(equation_terms)}"

            # 确定显著变量
            if f_pvalue < 0.05:
                significant_vars = [independent_vars[i] for i, p in enumerate(p_values) if p < 0.05]
            else:
                significant_vars = []

            # 准备结果表格数据
            var_results = []
            for i, var_name in enumerate(independent_vars):
                var_results.append({
                    "var_name": var_name,
                    "coefficient": coefficients.iloc[i],
                    "t_value": t_values.iloc[i],
                    "p_value": p_values.iloc[i]
                })

            # 模型统计量
            model_stats = {
                LANGUAGES[self.current_language]["intercept"]: intercept,
                LANGUAGES[self.current_language]["mse"]: mse,
                LANGUAGES[self.current_language]["r_squared"]: r2,
                LANGUAGES[self.current_language]["adjusted_r_squared"]: adjusted_r2,
                LANGUAGES[self.current_language]["f_statistic"]: f_value,
                LANGUAGES[self.current_language]["f_pvalue"]: f_pvalue
            }

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加标题
                doc.add_heading(LANGUAGES[self.current_language]["title"], level=1)

                # 添加回归方程
                doc.add_heading(LANGUAGES[self.current_language]["regression_equation"], level=2)
                p = doc.add_paragraph(regression_equation)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # 添加变量分析结果表格
                doc.add_heading(LANGUAGES[self.current_language]["var_analysis_results"], level=2)
                var_table = doc.add_table(rows=1, cols=6)
                var_table.style = 'Table Grid'

                # 变量表格表头
                var_hdr = var_table.rows[0].cells
                var_hdr[0].text = LANGUAGES[self.current_language]["var_name"]
                var_hdr[1].text = LANGUAGES[self.current_language]["coefficient"]
                var_hdr[2].text = LANGUAGES[self.current_language]["t_value"]
                var_hdr[3].text = LANGUAGES[self.current_language]["p_value"]
                var_hdr[4].text = LANGUAGES[self.current_language]["ci_lower"]
                var_hdr[5].text = LANGUAGES[self.current_language]["ci_upper"]

                param_names = model.params.index.tolist()  # 包含 const 和所有自变量
                conf_values = conf_int.values  # 转成 NumPy 数组方便按位置索引

                for i, var in enumerate(var_results):
                    row_cells = var_table.add_row().cells
                    row_cells[0].text = var["var_name"]
                    row_cells[1].text = f"{var['coefficient']:.4f}"
                    row_cells[2].text = f"{var['t_value']:.4f}"
                    row_cells[3].text = f"{var['p_value']:.4f}"

                    # 注意：conf_int[0] 对应常数项，conf_int[1:] 对应自变量
                    ci_low, ci_high = conf_values[i + 1]  # 跳过常数项
                    row_cells[4].text = f"{ci_low:.4f}"
                    row_cells[5].text = f"{ci_high:.4f}"

                # 添加模型统计量表格
                doc.add_heading(LANGUAGES[self.current_language]["model_statistics"], level=2)
                stats_table = doc.add_table(rows=1, cols=2)
                stats_table.style = 'Table Grid'

                # 统计量表格表头
                stats_hdr = stats_table.rows[0].cells
                stats_hdr[0].text = LANGUAGES[self.current_language]["statistic"]
                stats_hdr[1].text = LANGUAGES[self.current_language]["value"]

                # 添加统计量数据
                for name, value in model_stats.items():
                    row_cells = stats_table.add_row().cells
                    row_cells[0].text = name
                    row_cells[1].text = f"{value:.4f}"

                # 添加显著性检验结论
                doc.add_heading(LANGUAGES[self.current_language]["significance_test"], level=2)
                if f_pvalue < 0.05:
                    doc.add_paragraph(LANGUAGES[self.current_language]["significant_model"].format(f_pvalue))
                else:
                    doc.add_paragraph(LANGUAGES[self.current_language]["not_significant_model"].format(f_pvalue))

                # 添加显著变量
                if significant_vars:
                    doc.add_paragraph(
                        f"{LANGUAGES[self.current_language]['significant_variable']} {', '.join(significant_vars)}")
                else:
                    doc.add_paragraph(LANGUAGES[self.current_language]["no_significant_variable"])

                # 添加解释说明
                doc.add_heading(LANGUAGES[self.current_language]["index_explanation"], level=2)
                expl_paragraph = doc.add_paragraph()
                explanations = LANGUAGES[self.current_language]['explanation']
                for name, desc in explanations.items():
                    # 添加项目符号
                    run = expl_paragraph.add_run(f"• {name}: {desc}\n")

                # 获取保存路径的目录
                save_dir = os.path.dirname(save_path)

                # 生成散点图
                plt.figure(figsize=(10, 6))
                plt.close()

                # 残差分布图
                plt.figure(figsize=(10, 6))
                plt.hist(residuals, bins=20, alpha=0.7)
                plt.xlabel(LANGUAGES[self.current_language]["residuals"])
                plt.ylabel(LANGUAGES[self.current_language]["frequency"])
                plt.title(LANGUAGES[self.current_language]["residual_hist_title"])
                residual_hist_path = os.path.join(save_dir, "residual_histogram.png")
                plt.savefig(residual_hist_path, bbox_inches='tight')
                plt.close()

                # 残差与拟合值散点图
                plt.figure(figsize=(10, 6))
                plt.scatter(y_pred, residuals, alpha=0.6)
                plt.axhline(y=0, color='r', linestyle='--')
                plt.xlabel(LANGUAGES[self.current_language]["fitted_values"])
                plt.ylabel(LANGUAGES[self.current_language]["residuals"])
                plt.title(LANGUAGES[self.current_language]["residual_scatter_title"])
                residual_scatter_path = os.path.join(save_dir, "residual_scatter.png")
                plt.savefig(residual_scatter_path, bbox_inches='tight')
                plt.close()


                plt.scatter(y, y_pred, alpha=0.6)
                plt.plot([y.min(), y.max()], [y.min(), y.max()], 'r--', lw=2)
                plt.xlabel(LANGUAGES[self.current_language]["actual_value"].format(dependent_var))
                plt.ylabel(LANGUAGES[self.current_language]["predicted_value"].format(dependent_var))
                plt.title(LANGUAGES[self.current_language]["actual_vs_predicted_title"])
                img_name = "ols_regression_scatter.png"
                img_path = os.path.join(save_dir, img_name)
                plt.savefig(img_path, bbox_inches='tight')
                plt.close()

                doc.add_heading(LANGUAGES[self.current_language]["anova_table"], level=2)
                # 创建ANOVA表格
                anova_table_doc = doc.add_table(rows=1, cols=len(anova_table.columns) + 1)
                anova_table_doc.style = 'Table Grid'

                # 添加表头
                hdr_cells = anova_table_doc.rows[0].cells
                hdr_cells[0].text = LANGUAGES[self.current_language]["source"]  # 行索引列标题
                for col_idx, col_name in enumerate(anova_table.columns):
                    hdr_cells[col_idx + 1].text = str(col_name)

                # 添加数据行
                for idx, row in anova_table.iterrows():
                    row_cells = anova_table_doc.add_row().cells
                    row_cells[0].text = str(idx)
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx + 1].text = f"{value:.4f}"
                doc.add_paragraph(LANGUAGES[self.current_language]["anova_explanation"])

                # 添加残差分析内容
                doc.add_heading(LANGUAGES[self.current_language]["residual_analysis"], level=2)
                doc.add_paragraph(LANGUAGES[self.current_language]["residual_hist_desc"])
                doc.add_picture(residual_hist_path, width=Inches(6))

                doc.add_paragraph(LANGUAGES[self.current_language]["residual_scatter_desc"])
                doc.add_paragraph(LANGUAGES[self.current_language]["residual_homoscedasticity"])
                doc.add_picture(residual_scatter_path, width=Inches(6))

                # 在 Word 文档中插入实际值与预测值对比图片
                doc.add_heading(LANGUAGES[self.current_language]["actual_vs_predicted"], level=2)
                doc.add_picture(img_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                result_msg = (
                        LANGUAGES[self.current_language]['analysis_success'].format(save_path) + "\n" +
                        LANGUAGES[self.current_language]['images_saved'].format(save_dir)
                )
                self.result_label.config(text=result_msg, wraplength=400)
            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(LANGUAGES[self.current_language]["title"])
        self.select_button.config(text=LANGUAGES[self.current_language]["select_button"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=LANGUAGES[self.current_language]["analyze_button"])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]["switch_language"])
        self.open_excel_label.config(text=LANGUAGES[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口的 x 和 y 坐标，使其居中
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)  # 使用 expand 选项使框架在上下方向上居中

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["select_button"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["analyze_button"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=LANGUAGES[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=LANGUAGES[self.current_language]["switch_language"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = OrdinaryLeastSquaresLinearRegressionAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()
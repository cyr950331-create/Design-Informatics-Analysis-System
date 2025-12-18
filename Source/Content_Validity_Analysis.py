import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from statsmodels.stats.proportion import proportions_ztest

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'  # 设置字体为黑体，可根据系统情况修改为其他支持中文的字体
plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

# 定义语言字典
LANGUAGES = {
    'zh': {
        'title': "内容效度",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "请选择文件。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "中/英",
        'open_excel_button_text': "示例数据",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'explanation': {
            "平均内容效度比（CVR）": "平均内容效度比用于衡量测量工具中各个题项与测量内容的相关性，取值范围在 -1 到 1 之间，越接近 1 表示相关性越强。",
            "样本量": "每个样本中的观测值数量。",
            "均值": "样本数据的平均值。",
            "标准差": "样本数据的离散程度。",
            "中位数": "样本数据的中间值。",
            "偏度": "样本数据分布的偏斜程度。",
            "峰度": "样本数据分布的峰态程度。",
            "项目内容效度指数（I-CVI）": "每个题项获得80%以上专家认可的比例，用于衡量单个题项的内容效度。",
            "量表内容效度指数（S-CVI/UA）": "所有题项均达到I-CVI临界值的比例，反映量表整体内容效度。",
            "量表内容效度指数（S-CVI/Ave）": "所有题项I-CVI的平均值，更稳健地反映量表整体内容效度。"
        },
        'interpretation': {
            "平均内容效度比（CVR）": "平均内容效度比越接近 1，说明测量工具的内容与所测量的概念或领域相关性越强，内容效度越高。",
            "样本量": "样本量的大小会影响统计检验的稳定性，较大的样本量通常能提供更可靠的结果。",
            "均值": "均值反映了数据的平均水平，可用于比较不同变量的集中趋势。",
            "标准差": "标准差越大，说明数据的离散程度越大。",
            "中位数": "中位数不受极端值的影响，能更好地反映数据的中间水平。",
            "偏度": "偏度为正表示数据右偏，偏度为负表示数据左偏。",
            "峰度": "峰度大于 3 表示数据分布比正态分布更尖峭，峰度小于 3 表示数据分布比正态分布更平坦。",
            "项目内容效度指数（I-CVI）": "I-CVI≥0.8表示题项内容效度可接受，数值越高表示题项与测量内容的相关性越强。",
            "量表内容效度指数（S-CVI/UA）": "S-CVI/UA越接近1，说明通过I-CVI检验的题项比例越高，量表整体内容效度越好。",
            "量表内容效度指数（S-CVI/Ave）": "S-CVI/Ave≥0.9表示量表内容效度优秀，0.8-0.9之间表示良好。"
        },
        'expert_count': "专家数量",
        'cvr_critical': "CVR临界值",
        'i_cvi_threshold': "I-CVI临界值",
        'item_analysis': "题项分析结果",
        'item': "题项",
        'status': "状态",
        'passed': "通过",
        'failed': "未通过",
        'cvr_chart_title': "各题项CVR值",
        'icvi_chart_title': "各题项I-CVI值",
        'items': "题项",
        'empty_data': "数据为空，无法进行分析。",
        'insufficient_experts': "专家数量不足，至少需要3名专家才能进行内容效度分析。",
        'file_permission_error': "没有权限访问该文件，请检查文件权限。",
        'empty_excel_file': "Excel文件为空，无法进行分析。",
        's_cvi_ua': "量表内容效度指数（S-CVI/UA）",
        's_cvi_ave': "量表内容效度指数（S-CVI/Ave）",
    },
    'en': {
        'title': "Content Validity",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Chinese/English",
        'open_excel_button_text': "Example data",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'explanation': {
            "Average Content Validity Ratio (CVR)": "The average content validity ratio (CVR) is used to measure the correlation between each item in the measurement tool and the measured content. The value ranges from -1 to 1, and the closer it is to 1, the stronger the correlation.",
            "Sample Size": "The number of observations in each sample.",
            "Mean": "The average value of the sample data.",
            "Standard Deviation": "The degree of dispersion of the sample data.",
            "Median": "The median value of the sample data.",
            "Skewness": "The degree of skewness of the sample data distribution.",
            "Kurtosis": "The degree of kurtosis of the sample data distribution.",
            "Item Content Validity Index (I-CVI)": "The proportion of each item recognized by more than 80% of experts, used to measure the content validity of individual items.",
            "Scale Content Validity Index (S-CVI/UA)": "The proportion of all items that meet the I-CVI threshold, reflecting the overall content validity of the scale.",
            "Scale Content Validity Index (S-CVI/Ave)": "The average of I-CVI for all items, more robustly reflecting the overall content validity of the scale."
        },
        'interpretation': {
            "Average Content Validity Ratio (CVR)": "The closer the average content validity ratio (CVR) is to 1, the stronger the correlation between the content of the measurement tool and the measured concept or domain, and the higher the content validity.",
            "Sample Size": "The sample size affects the stability of the statistical test. A larger sample size usually provides more reliable results.",
            "Mean": "The mean reflects the average level of the data and can be used to compare the central tendencies of different variables.",
            "Standard Deviation": "A larger standard deviation indicates a greater degree of dispersion of the data.",
            "Median": "The median is not affected by extreme values and can better reflect the middle level of the data.",
            "Skewness": "A positive skewness indicates a right-skewed distribution, while a negative skewness indicates a left-skewed distribution.",
            "Kurtosis": "A kurtosis greater than 3 indicates a more peaked distribution than the normal distribution, while a kurtosis less than 3 indicates a flatter distribution than the normal distribution.",
            "Item Content Validity Index (I-CVI)": "I-CVI≥0.8 indicates acceptable content validity for the item, with higher values indicating stronger relevance to the measured content.",
            "Scale Content Validity Index (S-CVI/UA)": "The closer S-CVI/UA is to 1, the higher the proportion of items passing the I-CVI test, indicating better overall content validity.",
            "Scale Content Validity Index (S-CVI/Ave)": "S-CVI/Ave≥0.9 indicates excellent scale content validity, and 0.8-0.9 indicates good validity."
        },
        'expert_count': "Number of Experts",
        'cvr_critical': "CVR Critical Value",
        'i_cvi_threshold': "I-CVI Threshold",
        'item_analysis': "Item Analysis Results",
        'item': "Item",
        'status': "Status",
        'passed': "Passed",
        'failed': "Failed",
        'cvr_chart_title': "CVR Values for Each Item",
        'icvi_chart_title': "I-CVI Values for Each Item",
        'items': "Items",
        'empty_data': "The data is empty and cannot be analyzed.",
        'insufficient_experts': "Insufficient number of experts. At least 3 experts are needed to conduct content validity analysis.",  # 移除中文句号
        'file_permission_error': "No permission to access the file, please check the file permissions.",  # 移除中文句号
        'empty_excel_file': "The Excel file is empty and cannot be analyzed.",
        's_cvi_ua': "Scale Content Validity Index (S-CVI/UA)",
        's_cvi_ave': "Scale Content Validity Index (S-CVI/Ave)",
    }
}

class ContentValidityAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"
        
        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])
            
        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data13.xls")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def content_validity_analysis(self, data):
        cvr_values = []
        critical_values = {  # 基于Lawshe表的临界值
            3: 0.99, 4: 0.75, 5: 0.62, 6: 0.52, 7: 0.45, 8: 0.40,
            9: 0.37, 10: 0.34, 11: 0.32, 12: 0.30, 13: 0.28, 14: 0.27,
            15: 0.25, 16: 0.24, 17: 0.23, 18: 0.22, 19: 0.21, 20: 0.20
        }

        n = data.shape[0]  # 专家数量
        # 确定临界值，如果专家数量超过20，使用近似值
        if n >= 20:
            cvr_critical = 1.44 / (1 + n / 2)  # 近似公式
        else:
            cvr_critical = critical_values.get(n, 0.5)  # 默认为0.5

        item_results = []
        for column in data.columns:
            # 根据数据类型计算相关专家数量
            if set(data.values.flatten()) <= {0, 1}:
                # 0/1数据：1表示相关
                ne = data[column].sum()
            else:
                # Likert量表（1-4）：3-4表示相关
                ne = sum(data[column] >= 3)

            # 计算CVR
            if n == 0:
                cvr = 0.0
            else:
                cvr = (ne - n / 2) / (n / 2)
            cvr_values.append(cvr)
            # 判断是否通过
            passed = cvr >= cvr_critical
            item_results.append({
                "item": column,
                "cvr": cvr,
                "passed": passed
            })

        # 计算平均CVR（仅包含通过的题项）
        passed_cvrs = [item["cvr"] for item in item_results if item["passed"]]
        average_cvr = np.mean(passed_cvrs) if passed_cvrs else 0.0

        # 计算CVR的统计显著性
        p_values = {}
        for column in data.columns:
            # 统一判断数据类型，避免重复计算set
            if set(data.values.flatten()) <= {0, 1}:
                ne = data[column].sum()
            else:
                ne = sum(data[column] >= 3)
            stat, p_value = proportions_ztest(ne, n, 0.5, alternative='larger')
            p_values[column] = p_value

        return {
            "average_cvr": average_cvr,
            "item_results": item_results,
            "cvr_critical": cvr_critical,
            "expert_count": n,
            "p_values": p_values  # 包含p_values
        }

    def calculate_cvi(self, data):
        """计算项目内容效度指数(I-CVI)和量表内容效度指数(S-CVI)"""
        n_experts = data.shape[0]  # 专家数量
        threshold = 0.8  # 80%专家认可阈值

        # 计算每个题项的I-CVI（假设3-4分为相关，1-2分为不相关）
        i_cvi = {}
        for column in data.columns:
            # 统计认为相关的专家数量（评分≥3）
            relevant_count = sum(data[column] >= 3)
            i_cvi[column] = relevant_count / n_experts if n_experts > 0 else 0

        # 计算S-CVI/UA（所有题项I-CVI均≥0.8的比例）
        passing_items = sum(1 for cvi in i_cvi.values() if cvi >= threshold)
        s_cvi_ua = passing_items / len(i_cvi) if i_cvi else 0

        # 计算S-CVI/Ave（所有题项I-CVI的平均值）
        s_cvi_ave = np.mean(list(i_cvi.values())) if i_cvi else 0

        return {
            "i_cvi": i_cvi,
            "s_cvi_ua": s_cvi_ua,
            "s_cvi_ave": s_cvi_ave,
            "threshold": threshold
        }

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        """当用户点击输入框时，清除提示信息"""
        if self.file_entry.get() == LANGUAGES[self.current_language]['file_entry_placeholder']:
            self.file_entry.delete(0, tk.END)
            self.file_entry.configure(style="TEntry")  # 恢复默认样式

    def on_focusout(self, event):
        """当用户离开输入框时，如果没有输入内容，恢复提示信息"""
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
            self.file_entry.configure(style="Gray.TEntry")

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]['file_entry_placeholder']:
            file_path = ""
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 缺失值处理 - 列均值填充
            if df.isnull().any().any():
                # 对数值列使用均值填充
                num_cols = df.select_dtypes(include=[np.number]).columns
                df[num_cols] = df[num_cols].fillna(df[num_cols].mean())

            # 检查数据是否为数值类型
            numerical_df = df.select_dtypes(include=[np.number])
            if numerical_df.empty:
                raise ValueError("数据中没有数值列，无法进行内容效度分析。")

            # 检查数据是否为0/1或Likert量表数据（1-4级）
            all_values = set(numerical_df.stack().unique())  # 获取所有数值的集合
            is_binary = all_values.issubset({0, 1})  # 是否为0/1数据
            is_likert = all_values.issubset({1, 2, 3, 4})  # 是否为1-4的Likert数据

            if not (is_binary or is_likert):
                invalid = all_values - {0, 1, 2, 3, 4}
                raise ValueError(f"数据包含无效值 {invalid}。请确保数据为0/1（二分类）或1-4（Likert量表）。")

            # 在内容效度分析后添加变量赋值
            analysis_result = self.content_validity_analysis(numerical_df)
            cvi_result = self.calculate_cvi(numerical_df)

            # 计算更多的统计指标
            sample_sizes = numerical_df.count()
            means = numerical_df.mean()
            stds = numerical_df.std()
            medians = numerical_df.median()
            skewnesses = numerical_df.skew()
            kurtoses = numerical_df.kurt()
            mins = numerical_df.min()
            maxs = numerical_df.max()
            ranges = maxs - mins
            q1 = numerical_df.quantile(0.25)
            q3 = numerical_df.quantile(0.75)
            iqr = q3 - q1  # 四分位距
            cv = stds / means  # 变异系数

            # 整理数据
            # 根据当前语言定义统计量名称
            if self.current_language == "zh":
                stats_names = {
                    "average_cvr": "平均内容效度比（CVR）",
                    "s_cvi_ua": "量表内容效度指数（S-CVI/UA）",
                    "s_cvi_ave": "量表内容效度指数（S-CVI/Ave）",
                    "expert_count": "专家数量",
                    "cvr_critical": "CVR临界值",
                    "i_cvi_threshold": "I-CVI临界值",
                    "sample_sizes": "样本量",
                    "means": "均值",
                    "stds": "标准差",
                    "medians": "中位数",
                    "skewnesses": "偏度",
                    "kurtoses": "峰度",
                    "cv": "变异系数",
                    "p_values": "p值（p<0.05表示统计显著）"
                }
            else:
                stats_names = {
                    "average_cvr": "Average Content Validity Ratio (CVR)",
                    "s_cvi_ua": "Scale Content Validity Index (S-CVI/UA)",
                    "s_cvi_ave": "Scale Content Validity Index (S-CVI/Ave)",
                    "expert_count": "Number of Experts",
                    "cvr_critical": "CVR Critical Value",
                    "i_cvi_threshold": "I-CVI Threshold",
                    "sample_sizes": "Sample Size",
                    "means": "Mean",
                    "stds": "Standard Deviation",
                    "medians": "Median",
                    "skewnesses": "Skewness",
                    "kurtoses": "Kurtosis",
                    "cv": "Coefficient of Variation",
                    "p_values": "p-value (p<0.05 indicates statistical significance)"
                }

            # 构建多语言数据列表
            data = [
                [stats_names["average_cvr"], f"{analysis_result['average_cvr']:.3f}"],
                [stats_names["s_cvi_ua"], f"{cvi_result['s_cvi_ua']:.3f}"],
                [stats_names["s_cvi_ave"], f"{cvi_result['s_cvi_ave']:.3f}"],
                [stats_names["expert_count"], f"{analysis_result['expert_count']}"],
                [stats_names["cvr_critical"], f"{analysis_result['cvr_critical']:.3f}"],
                [stats_names["i_cvi_threshold"], f"{cvi_result['threshold']:.1f}"],
                [stats_names["sample_sizes"], ", ".join([f"{k}: {v}" for k, v in sample_sizes.items()])],
                [stats_names["means"], ", ".join([f"{k}: {v:.3f}" for k, v in means.items()])],
                [stats_names["stds"], ", ".join([f"{k}: {v:.3f}" for k, v in stds.items()])],
                [stats_names["medians"], ", ".join([f"{k}: {v:.3f}" for k, v in medians.items()])],
                [stats_names["skewnesses"], ", ".join([f"{k}: {v:.3f}" for k, v in skewnesses.items()])],
                [stats_names["kurtoses"], ", ".join([f"{k}: {v:.3f}" for k, v in kurtoses.items()])],
                [stats_names["cv"], ", ".join([f"{k}: {v:.3f}" for k, v in cv.items()])],
                [stats_names["p_values"], ", ".join([f"{k}: {v:.3f}" for k, v in analysis_result['p_values'].items()])]
            ]
            if self.current_language == "zh":
                headers = ["统计量", "统计量值"]
            else:
                headers = ["Statistic", "Value"]
            df = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            explanations = LANGUAGES[self.current_language]['explanation']
            interpretations = LANGUAGES[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(columns=["平均内容效度比（CVR）", "样本量", "均值", "标准差", "中位数", "偏度", "峰度"])
            explanation_df.insert(0, "统计量_解释说明", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(columns=["平均内容效度比（CVR）", "样本量", "均值", "标准差", "中位数", "偏度", "峰度"])
            interpretation_df.insert(0, "统计量_结果解读", "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 合并数据、解释说明和结果解读
            combined_df = pd.concat([df, explanation_df, interpretation_df], ignore_index=True)

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:

                # 创建 Word 文档
                doc = Document()
                report_suffix = "分析报告" if self.current_language == "zh" else "Analysis Report"
                # 拼接标题并添加
                doc.add_heading(f"{LANGUAGES[self.current_language]['title']} {report_suffix}", level=1)

                # 添加方法学说明
                if self.current_language == "zh":
                    doc.add_heading("分析方法说明", level=2)
                    methodology = "本分析采用Lawshe内容效度比(CVR)和内容效度指数(CVI)方法，" \
                                  "CVR临界值基于Lawshe表确定，I-CVI临界值设为0.8。" \
                                  "缺失值采用列均值填充，统计显著性水平设为p<0.05。"
                else:
                    doc.add_heading("Explanation of Analysis Methods", level=2)
                    methodology = "This analysis uses the Lawshe Content Validity Ratio (CVR) and Content Validity Index (CVI) methods. " \
                                  "The CVR critical value is determined based on the Lawshe table, and the I-CVI critical value is set to 0.8. " \
                                  "Missing values are imputed with column means, and the statistical significance level is set to p<0.05."
                doc.add_paragraph(methodology)

                # 添加基本信息
                if self.current_language == "zh":
                    doc.add_heading("基本信息", level=2)
                else:
                    doc.add_heading("Basic Information", level=2)
                doc.add_paragraph(
                    f"{LANGUAGES[self.current_language]['expert_count']}: {analysis_result['expert_count']}")
                doc.add_paragraph(
                    f"{LANGUAGES[self.current_language]['cvr_critical']}: {analysis_result['cvr_critical']:.3f}")
                doc.add_paragraph(f"{LANGUAGES[self.current_language]['i_cvi_threshold']}: {cvi_result['threshold']}")
                doc.add_paragraph(f"{LANGUAGES[self.current_language]['s_cvi_ua']}: {cvi_result['s_cvi_ua']:.3f}")
                doc.add_paragraph(f"{LANGUAGES[self.current_language]['s_cvi_ave']}: {cvi_result['s_cvi_ave']:.3f}")

                # 添加题项分析结果表格
                doc.add_heading(LANGUAGES[self.current_language]['item_analysis'], level=2)
                item_table = doc.add_table(rows=len(analysis_result['item_results']) + 1, cols=5)
                item_hdr = item_table.rows[0].cells
                item_hdr[0].text = LANGUAGES[self.current_language]['item']
                item_hdr[1].text = "CVR"
                item_hdr[2].text = "I-CVI"
                if self.current_language == "zh":
                    item_hdr[3].text = "p值"
                else:
                    item_hdr[3].text = "p-value"
                item_hdr[4].text = LANGUAGES[self.current_language]['status']

                for idx, item in enumerate(analysis_result['item_results']):
                    row = item_table.rows[idx + 1].cells
                    row[0].text = item['item']
                    row[1].text = f"{item['cvr']:.3f}"
                    row[2].text = f"{cvi_result['i_cvi'][item['item']]:.3f}"
                    row[3].text = f"{analysis_result['p_values'][item['item']]:.3f}"
                    row[4].text = LANGUAGES[self.current_language]['passed'] if item['passed'] else \
                    LANGUAGES[self.current_language]['failed']

                # 单独添加统计结果表格部分
                if self.current_language == "zh":
                    heading_text = "统计结果"
                else:
                    heading_text = "Statistical Results"
                doc.add_heading(heading_text, level=2)
                stats_table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
                stats_hdr = stats_table.rows[0].cells
                for col_idx, header in enumerate(df.columns):
                    stats_hdr[col_idx].text = header

                for row_idx, row in df.iterrows():
                    row_cells = stats_table.rows[row_idx + 1].cells
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx].text = str(value)

                # 定义中英文统计量名称映射（顺序保持一致）
                if self.current_language == "zh":
                    stats_list = ["平均内容效度比（CVR）", "样本量", "均值", "标准差", "中位数", "偏度", "峰度"]
                else:
                    stats_list = ["Average Content Validity Ratio (CVR)", "Sample Size", "Mean", "Standard Deviation",
                                  "Median", "Skewness", "Kurtosis"]

                # 单独添加解释说明部分
                if self.current_language == "zh":
                    doc.add_heading("统计量解释说明", level=2)
                else:
                    doc.add_heading("Explanation of Statistics", level=2)
                expl_list = doc.add_paragraph()
                for stat in stats_list:
                    # 根据当前语言获取对应的解释文本（需确保explanations字典包含中英文键）
                    run = expl_list.add_run(f"• {stat}: {explanations[stat]}\n")

                # 单独添加结果解读部分
                if self.current_language == "zh":
                    doc.add_heading("统计结果解读", level=2)
                else:
                    doc.add_heading("Interpretation of Statistical Results", level=2)
                interp_list = doc.add_paragraph()
                for stat in stats_list:
                    # 根据当前语言获取对应的解读文本（需确保interpretations字典包含中英文键）
                    run = interp_list.add_run(f"• {stat}: {interpretations[stat]}\n")

                # 生成CVR条形图
                fig, ax = plt.subplots(figsize=(10, 6))
                items = [item['item'] for item in analysis_result['item_results']]
                cvrs = [item['cvr'] for item in analysis_result['item_results']]
                ax.bar(items, cvrs)
                ax.axhline(y=analysis_result['cvr_critical'], color='r', linestyle='--',
                           label=f"{LANGUAGES[self.current_language]['cvr_critical']}: {analysis_result['cvr_critical']:.3f}")
                ax.set_title(LANGUAGES[self.current_language]['cvr_chart_title'])
                ax.set_xlabel(LANGUAGES[self.current_language]['items'])
                ax.set_ylabel("CVR")
                ax.legend()
                # 使用分析结果中的p_values
                for i, item in enumerate(analysis_result['item_results']):
                    p_val = analysis_result['p_values'][item['item']]
                    significance = "*" if p_val < 0.05 else "ns"
                    ax.text(i, cvrs[i] + 0.05, significance, ha='center')
                plt.xticks(rotation=0)
                plt.tight_layout()
                cvr_img_path = os.path.splitext(save_path)[0] + '_cvr.png'
                plt.savefig(cvr_img_path)
                plt.close()
                doc.add_picture(cvr_img_path, width=Inches(6))

                # 生成I-CVI条形图
                fig, ax = plt.subplots(figsize=(10, 6))
                i_cvis = list(cvi_result['i_cvi'].values())
                ax.bar(items, i_cvis)
                ax.axhline(y=cvi_result['threshold'], color='r', linestyle='--',
                           label=f"{LANGUAGES[self.current_language]['i_cvi_threshold']}: {cvi_result['threshold']}")
                ax.set_title(LANGUAGES[self.current_language]['icvi_chart_title'])
                ax.set_xlabel(LANGUAGES[self.current_language]['items'])
                ax.set_ylabel("I-CVI")
                ax.legend()
                plt.xticks(rotation=0)
                plt.tight_layout()
                icvi_img_path = os.path.splitext(save_path)[0] + '_icvi.png'
                plt.savefig(icvi_img_path)
                plt.close()
                doc.add_picture(icvi_img_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(
                    save_path)
                self.result_label.config(text=result_msg, wraplength=400)

            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'
        self.root.title(LANGUAGES[self.current_language]['title'])
        self.select_button.config(text=LANGUAGES[self.current_language]['select_button'])
        self.analyze_button.config(text=LANGUAGES[self.current_language]['analyze_button'])
        self.language_label.config(text=LANGUAGES[self.current_language]['switch_language'])
        self.open_excel_label.config(text=LANGUAGES[self.current_language]["open_excel_button_text"])

        # 切换语言时更新提示信息
        self.file_entry.delete(0, tk.END)
        if self.current_language == 'zh':
            self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
            self.file_entry.configure(style="Gray.TEntry")
        else:
            self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
            self.file_entry.configure(style="Gray.TEntry")

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建自定义样式
        style = ttk.Style()
        style.configure("Gray.TEntry", foreground="gray")

        # 添加语言切换标签的样式
        style.configure(
            "Language.TLabel",
            foreground="gray"
        )

        # 创建一个主框架用于居中所有元素
        main_frame = ttk.Frame(self.root)
        main_frame.pack(expand=True)  # 允许框架扩展以填充空间

        # 创建文件选择按钮（放在主框架中）
        self.select_button = ttk.Button(
            main_frame,
            text=LANGUAGES[self.current_language]['select_button'],
            command=self.select_file,
            bootstyle=PRIMARY
        )
        self.select_button.pack(pady=10)

        # 创建文件路径输入框（放在主框架中）
        self.file_entry = ttk.Entry(main_frame, width=50, style="Gray.TEntry")
        self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
        self.file_entry.pack(pady=5)
        self.file_entry.bind("<FocusIn>", self.on_entry_click)
        self.file_entry.bind("<FocusOut>", self.on_focusout)

        # 创建分析按钮（放在主框架中）
        self.analyze_button = ttk.Button(
            main_frame,
            text=LANGUAGES[self.current_language]['analyze_button'],
            command=self.analyze_file,
            bootstyle=SUCCESS
        )
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            main_frame,
            text=LANGUAGES[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建语言切换标签（放在主框架中）
        self.language_label = ttk.Label(
            main_frame,
            text=LANGUAGES[self.current_language]['switch_language'],
            cursor = "hand2",
            style = "Language.TLabel"
        )
        self.language_label.pack(pady=10)
        self.language_label.bind("<Button-1>", self.switch_language)

        # 创建结果显示标签（放在主窗口中，保持在底部）
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10, fill=X, padx=10)
        
    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = ContentValidityAnalysisApp()
    app.run()

if __name__ == "__main__":
    run_app()
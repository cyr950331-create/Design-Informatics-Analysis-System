import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
from sklearn.linear_model import Lasso, LassoCV, lasso_path
from sklearn.metrics import mean_squared_error, r2_score
from sklearn.preprocessing import StandardScaler
from sklearn.model_selection import train_test_split, cross_val_score
import matplotlib.pyplot as plt
import statsmodels.api as sm
from docx import Document
from docx.shared import Inches
from scipy import stats


# 设置支持中文的字体
plt.rcParams["font.family"] = ["SimHei", "WenQuanYi Micro Hei", "Heiti TC", "Arial Unicode MS"]
plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

# 定义语言字典，所有键使用英文
LANGUAGES = {
    'zh': {
        'title': "Lasso回归分析",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "请选择文件。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'images_saved': "图片已保存到 {}",
        'switch_language': "中/英",
        'open_excel_button_text': "示例数据",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'explanation': {
            "coefficients": "回归系数，表示每个自变量对因变量的影响程度。",
            "intercept": "截距，是当所有自变量为 0 时因变量的预测值。",
            "mse": "均方误差，衡量预测值与真实值之间的平均误差。",
            "r_squared": "决定系数，取值范围在 0 到 1 之间，越接近 1 表示模型拟合效果越好。",
            "adjusted_r_squared": "调整决定系数，考虑了模型中自变量的数量，对模型的拟合优度进行了调整。",
            "f_value": "F 统计量，用于检验整个回归模型的显著性。",
            "t_value": "t 统计量，用于检验每个自变量的显著性。",
            "p_value": "p 值，用于判断自变量的显著性，p 值越小，自变量越显著。",
            "cv_r_squared": "交叉验证决定系数，评估模型的泛化能力。"
        },
        'explanation_terms': {
            "coefficients": "回归系数",
            "intercept": "截距",
            "mse": "均方误差",
            "r_squared": "决定系数",
            "adjusted_r_squared": "调整决定系数",
            "f_value": "F统计量",
            "t_value": "t统计量",
            "p_value": "p值",
            "cv_r_squared": "交叉验证决定系数"
        },
        'actual_values': '实际值',
        'predicted_values': '预测值',
        'actual_vs_predicted': '实际值 vs 预测值',
        'alpha': '正则化参数',
        'lasso_coefficients': '套索回归系数随正则化参数变化',
        'regression_formula': '回归方程',
        'dependent_variable': '因变量',
        'independent_variable': '自变量',
        'regression_coefficient': '回归系数',
        'training_set': '训练集',
        'test_set': '测试集',
        'model_info': '模型信息',
        'optimal_alpha': '最优Alpha',
        'intercept_label': '截距',
        'mse_train_label': '训练集 均方误差 (MSE)',
        'mse_test_label': '测试集 均方误差 (MSE)',
        'r2_train_label': '训练集 决定系数 (R²)',
        'r2_test_label': '测试集 决定系数 (R²)',
        'adjusted_r2_label': '调整后决定系数',
        'f_statistic_label': 'F统计量',
        'cv_r2_mean_label': '交叉验证决定系数 均值',
        'cv_r2_std_label': '交叉验证决定系数 标准差',
        'residual_normality_test': '残差正态性检验',
        'shapiro_wilk_stat': 'Shapiro-Wilk统计量',
        'feature_label': '特征',
        'coefficient_label': '系数',
        'standardized_coefficient_label': '标准化系数',
        'statistical_tests': '统计检验（仅非零系数特征）',
        'constant_t_p': '常数项 t值/p值',
        'explanation_title': '解释说明',
        'visualization_results': '可视化结果',
        'analysis_results': '分析结果',
        'data_warning': '数据警告',
        'missing_values_detected': '检测到缺失值：\n{}\n将使用均值填充数值型变量'
    },
    'en': {
        'title': "Lasso Regression Analysis",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "Please select a file.",
        'analysis_success': "Analysis completed. Results saved to {}\n",
        'no_save_path': "No save path selected. Results not saved.",
        'analysis_error': "Error analyzing file: {}",
        'images_saved': "Images saved to {}",
        'switch_language': "Chinese/English",
        'open_excel_button_text': "Example Data",
        'file_entry_placeholder': "Please enter the full path of the Excel file to analyze",
        'explanation': {
            "coefficients": "Regression coefficients, indicating the influence of each independent variable on the dependent variable.",
            "intercept": "Intercept, which is the predicted value of the dependent variable when all independent variables are 0.",
            "mse": "Mean squared error, measuring the average error between predicted and actual values.",
            "r_squared": "Coefficient of determination, ranging from 0 to 1. Closer to 1 indicates better model fit.",
            "adjusted_r_squared": "Adjusted coefficient of determination, accounting for the number of independent variables.",
            "f_value": "F statistic, used to test the significance of the entire regression model.",
            "t_value": "t statistic, used to test the significance of each independent variable.",
            "p_value": "p value, used to determine the significance of the independent variable. Smaller values indicate greater significance.",
            "cv_r_squared": "Cross-validation R-squared, evaluating the generalization ability of the model."
        },
        'explanation_terms': {
            "coefficients": "Coefficients",
            "intercept": "Intercept",
            "mse": "MSE",
            "r_squared": "R-squared",
            "adjusted_r_squared": "Adjusted R-squared",
            "f_value": "F-value",
            "t_value": "t-value",
            "p_value": "p-value",
            "cv_r_squared": "CV R-squared"
        },
        'actual_values': 'Actual Values',
        'predicted_values': 'Predicted Values',
        'actual_vs_predicted': 'Actual vs Predicted Values',
        'alpha': 'Alpha',
        'lasso_coefficients': 'Lasso coefficients as a function of regularization',
        'regression_formula': 'Regression Formula',
        'dependent_variable': 'Dependent Variable',
        'independent_variable': 'Independent Variable',
        'regression_coefficient': 'Regression Coefficient',
        'training_set': 'Training Set',
        'test_set': 'Test Set',
        'model_info': 'Model Information',
        'optimal_alpha': 'Optimal Alpha',
        'intercept_label': 'Intercept',
        'mse_train_label': 'Training Set Mean Squared Error (MSE)',
        'mse_test_label': 'Test Set Mean Squared Error (MSE)',
        'r2_train_label': 'Training Set R-squared (R²)',
        'r2_test_label': 'Test Set R-squared (R²)',
        'adjusted_r2_label': 'Adjusted R-squared',
        'f_statistic_label': 'F Statistic',
        'cv_r2_mean_label': 'Cross-validation R² Mean',
        'cv_r2_std_label': 'Cross-validation R² Std',
        'residual_normality_test': 'Residual Normality Test',
        'shapiro_wilk_stat': 'Shapiro-Wilk Statistic',
        'feature_label': 'Feature',
        'coefficient_label': 'Coefficient',
        'standardized_coefficient_label': 'Standardized Coefficient',
        'statistical_tests': 'Statistical Tests (Non-zero coefficients only)',
        'constant_t_p': 'Constant Term t-value/p-value',
        'explanation_title': 'Explanations',
        'visualization_results': 'Visualization Results',
        'analysis_results': 'Analysis Results',
        'data_warning': 'Data Warning',
        'missing_values_detected': 'Missing values detected:\n{}\nWill fill numeric variables with mean values'
    }
}


class LassoRegressionAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为中文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]['title'])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]['title'])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data32.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        current_placeholder = LANGUAGES[self.current_language]['file_entry_placeholder']
        if self.file_entry.get() == current_placeholder:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')
            self.file_entry.configure(style="TEntry")  # 恢复默认样式

    def on_focusout(self, event):
        current_placeholder = LANGUAGES[self.current_language]['file_entry_placeholder']
        if self.file_entry.get().strip() == "":
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, current_placeholder)
            self.file_entry.config(foreground='gray')
            self.file_entry.configure(style="Gray.TEntry")

    def format_p_value(self, p):
        """格式化p值并添加显著性标记"""
        if p < 0.001:
            return f"{p:.4f}***"
        elif p < 0.01:
            return f"{p:.4f}**"
        elif p < 0.05:
            return f"{p:.4f}*"
        else:
            return f"{p:.4f}"

    def analyze_file(self):
        file_path = self.file_entry.get()
        current_placeholder = LANGUAGES[self.current_language]['file_entry_placeholder']
        if file_path == current_placeholder:
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件，第一行为表头
            df = pd.read_excel(file_path)

            # 检查并处理缺失值
            if df.isnull().any().any():
                missing_report = df.isnull().sum().to_string()
                Messagebox.showwarning(
                    LANGUAGES[self.current_language]['data_warning'],
                    LANGUAGES[self.current_language]['missing_values_detected'].format(missing_report)
                )
                # 数值型变量均值填充
                df = df.fillna(df.select_dtypes(include=[np.number]).mean())

            # 数据类型校验
            # 因变量为最后一列，需为数值型
            if not pd.api.types.is_numeric_dtype(df.iloc[:, -1]):
                raise ValueError(f"{LANGUAGES[self.current_language]['dependent_variable']}（最后一列）必须为数值型数据")
            # 自变量为中间列（第2列到倒数第2列），需为数值型
            non_numeric = df.columns[1:-1][~df.iloc[:, 1:-1].apply(pd.api.types.is_numeric_dtype)]
            if len(non_numeric) > 0:
                raise ValueError(f"{LANGUAGES[self.current_language]['independent_variable']}（中间列）不是数值型：{', '.join(non_numeric)}")

            # 获取因变量名称（最后一列列名）
            dependent_var_name = df.columns[-1]
            # 提取自变量名称（第2列到倒数第2列的列名）
            feature_names = df.columns[1:-1].tolist()
            # 自变量为中间列，因变量为最后一列
            X = df.iloc[:, 1:-1].values  # 自变量：第2列到倒数第2列
            y = df.iloc[:, -1].values  # 因变量：最后一列

            # 数据集划分
            X_train, X_test, y_train, y_test = train_test_split(
                X, y, test_size=0.2, random_state=42
            )

            # 使用交叉验证自动选择最优alpha
            lasso_cv = LassoCV(alphas=np.logspace(-4, 4, 100), cv=5, random_state=42)
            lasso_cv.fit(X_train, y_train)
            optimal_alpha = lasso_cv.alpha_

            # 使用最优alpha进行套索回归
            lasso = Lasso(alpha=optimal_alpha)
            lasso.fit(X_train, y_train)

            # 预测（训练集和测试集）
            y_pred_train = lasso.predict(X_train)
            y_pred_test = lasso.predict(X_test)

            # 计算指标
            coefficients = lasso.coef_
            intercept = lasso.intercept_

            # 训练集指标
            mse_train = mean_squared_error(y_train, y_pred_train)
            r2_train = r2_score(y_train, y_pred_train)

            # 测试集指标
            mse_test = mean_squared_error(y_test, y_pred_test)
            r2_test = r2_score(y_test, y_pred_test)

            n = len(y_train)
            p = X_train.shape[1]
            adjusted_r2 = 1 - (1 - r2_train) * (n - 1) / (n - p - 1)

            # 交叉验证评估
            cv_scores = cross_val_score(lasso, X_train, y_train, cv=5, scoring='r2')
            cv_r2_mean = cv_scores.mean()
            cv_r2_std = cv_scores.std()

            # 生成回归公式
            formula_parts = [f"{intercept:.4f}"]
            for i, name in enumerate(feature_names):
                coef = coefficients[i]
                if coef != 0:  # 只包含非零系数的特征
                    if coef > 0:
                        formula_parts.append(f"+ {coef:.4f} × {name}")
                    else:
                        formula_parts.append(f"- {abs(coef):.4f} × {name}")

            # 处理特殊情况：如果没有自变量被选中
            if len(formula_parts) == 1:
                regression_formula = f"{dependent_var_name} = {formula_parts[0]}"
            else:
                regression_formula = f"{dependent_var_name} = " + " ".join(formula_parts)

            # 使用套索回归后的特征重新计算统计量（只保留非零系数特征）
            non_zero_mask = coefficients != 0
            X_selected = X_train[:, non_zero_mask]
            feature_names_selected = [feature_names[i] for i in range(len(feature_names)) if non_zero_mask[i]]

            # 对选择后的特征做OLS计算统计量
            X_selected_with_const = sm.add_constant(X_selected)
            model_selected = sm.OLS(y_train, X_selected_with_const).fit()
            t_values = model_selected.tvalues
            p_values = model_selected.pvalues
            f_value = model_selected.fvalue

            # 计算标准化系数
            scaler = StandardScaler()
            X_scaled = scaler.fit_transform(X_train)
            lasso_scaled = Lasso(alpha=optimal_alpha)
            lasso_scaled.fit(X_scaled, y_train)
            standardized_coefficients = lasso_scaled.coef_

            # 残差正态性检验
            residuals = y_train - y_pred_train
            stat, p_shapiro = stats.shapiro(residuals)

            # 准备结果数据
            result_data = {
                LANGUAGES[self.current_language]['model_info']: ["Lasso Regression", f"{LANGUAGES[self.current_language]['optimal_alpha']}: {optimal_alpha:.6f}"],
                LANGUAGES[self.current_language]['regression_formula']: [regression_formula, ""],
                LANGUAGES[self.current_language]['dependent_variable']: [dependent_var_name, ""],
                LANGUAGES[self.current_language]['intercept_label']: [intercept, ""],
                LANGUAGES[self.current_language]['mse_train_label']: [mse_train, ""],
                LANGUAGES[self.current_language]['mse_test_label']: [mse_test, ""],
                LANGUAGES[self.current_language]['r2_train_label']: [r2_train, ""],
                LANGUAGES[self.current_language]['r2_test_label']: [r2_test, ""],
                LANGUAGES[self.current_language]['adjusted_r2_label']: [adjusted_r2, ""],
                LANGUAGES[self.current_language]['f_statistic_label']: [f_value, ""],
                LANGUAGES[self.current_language]['cv_r2_mean_label']: [cv_r2_mean, ""],
                LANGUAGES[self.current_language]['cv_r2_std_label']: [cv_r2_std, ""],
                LANGUAGES[self.current_language]['residual_normality_test']: [
                    f"{LANGUAGES[self.current_language]['shapiro_wilk_stat']}: {stat:.4f}",
                    f"{LANGUAGES[self.current_language]['explanation']['p_value']}: {self.format_p_value(p_shapiro)}"
                ]
            }

            # 添加特征相关指标
            for i, name in enumerate(feature_names):
                result_data[f"{LANGUAGES[self.current_language]['feature_label']}: {name}"] = [
                    f"{LANGUAGES[self.current_language]['coefficient_label']}: {coefficients[i]:.4f}",
                    f"{LANGUAGES[self.current_language]['standardized_coefficient_label']}: {standardized_coefficients[i]:.4f}"
                ]

            # 添加选择后特征的t值和p值（带显著性标记）
            result_data[LANGUAGES[self.current_language]['statistical_tests']] = ["", ""]
            result_data[LANGUAGES[self.current_language]['constant_t_p']] = [f"{t_values[0]:.4f}", self.format_p_value(p_values[0])]
            for i, name in enumerate(feature_names_selected):
                result_data[f"{LANGUAGES[self.current_language]['feature_label']}: {name} t/p"] = [
                    f"{t_values[i + 1]:.4f}",
                    self.format_p_value(p_values[i + 1])
                ]

            # 转换为DataFrame
            df_result = pd.DataFrame.from_dict(result_data, orient='index', columns=['值', '补充信息'])

            # 添加解释说明
            explanations = LANGUAGES[self.current_language]['explanation']
            explanation_terms = LANGUAGES[self.current_language]['explanation_terms']  # 获取术语翻译
            explanation_data = {
                LANGUAGES[self.current_language]['explanation_title']: ["", ""]
            }
            for key, value in explanations.items():
                # 使用当前语言的术语作为键
                explanation_data[explanation_terms[key]] = [value, ""]
            explanation_df = pd.DataFrame.from_dict(explanation_data, orient='index', columns=['说明', ''])

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建一个新的 Word 文档
                doc = Document()

                # 添加标题
                doc.add_heading(LANGUAGES[self.current_language]['title'], 0)
                doc.add_paragraph(f"{LANGUAGES[self.current_language]['optimal_alpha']}: {optimal_alpha:.6f}")
                # 单独添加回归公式，确保其显示
                doc.add_paragraph(f"{LANGUAGES[self.current_language]['regression_formula']}: {regression_formula}")

                # 添加结果表格
                doc.add_heading(LANGUAGES[self.current_language]['analysis_results'], level=1)
                table = doc.add_table(rows=1, cols=2)
                hdr_cells = table.rows[0].cells
                # 使用多语言变量作为表头
                hdr_cells[0].text = LANGUAGES[self.current_language]['feature_label']  # 指标 -> 特征（使用现有语言项）
                hdr_cells[1].text = LANGUAGES[self.current_language]['explanation_title']  # 详情 -> 解释说明（使用现有语言项）

                for index, row in df_result.iterrows():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(index)
                    row_cells[1].text = f"{row['值']}  {row['补充信息']}".strip()

                # 添加解释说明表格（去掉说明标题）
                doc.add_heading(LANGUAGES[self.current_language]['explanation_title'], level=1)
                expl_table = doc.add_table(rows=0, cols=1)  # 初始行数设为0，不创建标题行

                for index, row in explanation_df.iterrows():
                    # 跳过标题行（如果有）
                    if index == LANGUAGES[self.current_language]['explanation_title']:
                        continue
                    row_cells = expl_table.add_row().cells
                    row_cells[0].text = f"{index}: {row['说明']}".strip()

                # 获取保存路径的目录
                save_dir = os.path.dirname(save_path)

                # 生成散点图（区分训练集和测试集）
                plt.figure(figsize=(10, 6))
                plt.scatter(y_train, y_pred_train, alpha=0.6, label=LANGUAGES[self.current_language]['training_set'])
                plt.scatter(y_test, y_pred_test, alpha=0.6, c='red', label=LANGUAGES[self.current_language]['test_set'])
                plt.plot([y.min(), y.max()], [y.min(), y.max()], 'r--', lw=2)
                plt.xlabel(LANGUAGES[self.current_language]['actual_values'])
                plt.ylabel(LANGUAGES[self.current_language]['predicted_values'])
                plt.title(LANGUAGES[self.current_language]['actual_vs_predicted'])
                plt.legend()
                img_name = "lasso_regression_scatter.png"
                img_path = os.path.join(save_dir, img_name)
                plt.savefig(img_path)
                plt.close()

                # 生成套索路径图（高效计算并标记最优alpha）
                alphas, coefs, _ = lasso_path(X_train, y_train, alphas=np.logspace(-4, 4, 50))
                coefs = coefs.T  # 转置以匹配alpha维度

                plt.figure(figsize=(10, 6))
                ax = plt.gca()
                ax.plot(alphas, coefs)
                ax.set_xscale('log')
                plt.xlabel(LANGUAGES[self.current_language]['alpha'])
                plt.ylabel(LANGUAGES[self.current_language]['regression_coefficient'])
                plt.title(LANGUAGES[self.current_language]['lasso_coefficients'])
                plt.axis('tight')
                # 标记最优alpha位置
                plt.axvline(x=optimal_alpha, color='black', linestyle='--',
                            label=f'{LANGUAGES[self.current_language]["optimal_alpha"]}: {optimal_alpha:.2e}')
                # 添加特征名称图例
                plt.legend(feature_names, loc='upper right', bbox_to_anchor=(1.2, 1))
                img_name_lasso_trace = "lasso_regression_trace.png"
                img_path_lasso_trace = os.path.join(save_dir, img_name_lasso_trace)
                plt.savefig(img_path_lasso_trace, bbox_inches='tight')
                plt.close()

                # 添加图片到文档
                doc.add_heading(LANGUAGES[self.current_language]['visualization_results'], level=1)
                doc.add_picture(img_path, width=Inches(6))
                doc.add_picture(img_path_lasso_trace, width=Inches(6))

                # 保存更新后的文档
                doc.save(save_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(save_path)
                result_msg += LANGUAGES[self.current_language]['images_saved'].format(save_dir)
                self.result_label.config(text=result_msg, wraplength=400)
            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except ValueError as ve:
            self.result_label.config(text=f"数据格式错误: {str(ve)}")
        except Exception as e:
            # 记录详细错误日志
            import traceback
            with open("error_log.txt", "a") as f:
                f.write(traceback.format_exc())
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(LANGUAGES[self.current_language]['title'])
        self.select_button.config(text=LANGUAGES[self.current_language]['select_button'])
        current_text = self.file_entry.get()
        current_placeholder = LANGUAGES['zh']['file_entry_placeholder'] if self.current_language == 'en' else \
            LANGUAGES['en']['file_entry_placeholder']
        new_placeholder = LANGUAGES[self.current_language]['file_entry_placeholder']
        if current_text == current_placeholder:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, new_placeholder)
        self.file_entry.config(foreground='gray' if self.file_entry.get() == new_placeholder else 'black')
        self.file_entry.configure(style="Gray.TEntry" if self.file_entry.get() == new_placeholder else "TEntry")
        self.analyze_button.config(text=LANGUAGES[self.current_language]['analyze_button'])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]['switch_language'])
        self.open_excel_label.config(text=LANGUAGES[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)  # 使用 expand 选项使框架在上下方向上居中

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=LANGUAGES[self.current_language]['select_button'],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建自定义样式
        style = ttk.Style()
        style.configure("Gray.TEntry", foreground="gray")
        style.configure("Gray.TLabel", foreground="gray")

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50, style="Gray.TEntry")
        self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
        self.file_entry.pack(pady=5)
        self.file_entry.bind("<FocusIn>", self.on_entry_click)
        self.file_entry.bind("<FocusOut>", self.on_focusout)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=LANGUAGES[self.current_language]['analyze_button'],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=LANGUAGES[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建语言切换标签
        self.switch_language_label = ttk.Label(frame, text=LANGUAGES[self.current_language]['switch_language'],
                                               style="Gray.TLabel", cursor="hand2")
        self.switch_language_label.pack(pady=10)
        self.switch_language_label.bind("<Button-1>", self.switch_language)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = LassoRegressionAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()
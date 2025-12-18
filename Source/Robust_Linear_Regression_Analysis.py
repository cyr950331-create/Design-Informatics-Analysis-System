import tkinter as tk
from tkinter import filedialog
import os
import pandas as pd
import numpy as np
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
from sklearn.metrics import mean_squared_error, r2_score
import statsmodels.api as sm
from docx import Document
from docx.shared import Inches

# 定义语言字典
LANGUAGES = {
    'zh': {
        'title': "Robust",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "请选择文件。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'images_saved': "图片已保存到 {}",
        'switch_language': "中/英",
        'open_excel_button_text': "示例数据",
        'regression_results': "回归分析结果：",
        'analysis_charts': "分析图表：",
        'file_entry_placeholder': "请输入待分析文件的完整路径",
        'regression_equation': "回归方程",
        'explanation': {
            "Coefficients": "回归系数，表示每个自变量对因变量的影响程度。",
            "Intercept": "截距，是当所有自变量为 0 时因变量的预测值。",
            "Mean Squared Error (MSE)": "均方误差，衡量预测值与真实值之间的平均误差。",
            "R-squared (R²)": "决定系数，取值范围在 0 到 1 之间，越接近 1 表示模型拟合效果越好。",
            "Adjusted R-squared": "调整决定系数，考虑了模型中自变量的数量和稳健回归特性。",
            "t-value": "t 统计量，用于检验每个自变量的显著性。",
            "p-value": "p 值，用于判断自变量的显著性，p 值越小，自变量越显著。",
            "Regression Equation": "表示因变量与自变量之间关系的数学表达式。"
        }
    },
    'en': {
        'title': "Robust",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'images_saved': "Images have been saved to {}",
        'switch_language': "Chinese/English",
        'open_excel_button_text': "Example data",
        'regression_results': "Regression Analysis Results:",
        'analysis_charts': "Analysis Charts:",
        'file_entry_placeholder': "Please enter the full path of the file to be analyzed",
        'regression_equation': "Regression Equation",
        'explanation': {
            "Coefficients": "Regression coefficients, indicating the influence of each independent variable on the dependent variable.",
            "Intercept": "Intercept, which is the predicted value of the dependent variable when all independent variables are 0.",
            "Mean Squared Error (MSE)": "Mean squared error, measuring the average error between the predicted and actual values.",
            "R-squared (R²)": "Coefficient of determination, ranging from 0 to 1. A value closer to 1 indicates a better fit of the model.",
            "Adjusted R-squared": "Adjusted coefficient of determination, considering the number of independent variables and robust regression characteristics.",
            "t-value": "t statistic, used to test the significance of each independent variable.",
            "p-value": "p value, used to determine the significance of the independent variable. The smaller the p value, the more significant the independent variable.",
            "Regression Equation": "Mathematical expression representing the relationship between dependent and independent variables."
        }
    }
}


plt.rcParams["font.family"] = ["SimHei", "WenQuanYi Micro Hei", "Heiti TC", "Arial Unicode MS"]
plt.rcParams["axes.unicode_minus"] = False  # 解决负号显示问题

class RobustLinearRegressionAnalysisApp:
    def __init__(self, root=None):
        self.current_language = "en"
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])
        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data32.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 支持Excel和CSV文件
            if file_path.endswith(('.xlsx', '.xls')):
                df = pd.read_excel(file_path)
            elif file_path.endswith('.csv'):
                df = pd.read_csv(file_path)
            else:
                raise ValueError("不支持的文件格式，仅支持Excel和CSV")

            # 使用第一行作为表头，第一列作为索引（因素列）
            df = df.set_index(df.columns[0])
            X = df.iloc[:, :-1].values  # 自变量（除最后一列）
            y = df.iloc[:, -1].values  # 因变量（最后一列）
            feature_names = df.columns[:-1].tolist()  # 获取自变量名称
            dependent_var_name = df.columns[-1]  # 获取因变量名称

            # 添加常数项
            X_with_const = sm.add_constant(X)

            # 稳健线性回归分析
            rlm_model = sm.RLM(y, X_with_const).fit()
            y_pred = rlm_model.predict(X_with_const)
            residuals = y - y_pred  # 计算残差

            # 计算统计指标
            coefficients = rlm_model.params[1:]
            intercept = rlm_model.params[0]
            mse = mean_squared_error(y, y_pred)
            r2 = r2_score(y, y_pred)
            n = len(y)
            p = X.shape[1]

            # 稳健回归调整R²计算
            robust_resid_var = np.median(residuals **2)  # 稳健残差方差估计
            total_var = np.var(y, ddof=1)  # 总方差
            adjusted_r2 = 1 - (robust_resid_var / total_var) * (n - 1) / (n - p - 1)

            # 获取t值和p值
            t_values = rlm_model.tvalues[1:]
            p_values = rlm_model.pvalues[1:]

            # 构建回归方程
            equation_parts = [f"{dependent_var_name} = {intercept:.4f}"]
            for i, name in enumerate(feature_names):
                coef = coefficients[i]
                # 根据系数正负决定符号
                if coef >= 0:
                    equation_parts.append(f"+ {coef:.4f}*{name}")
                else:
                    equation_parts.append(f"- {abs(coef):.4f}*{name}")
            regression_equation = " ".join(equation_parts)

            # 准备结果数据
            model_data = {
                "Model": "Robust Linear Regression",
                LANGUAGES[self.current_language]['regression_equation']: regression_equation,
                "Intercept": intercept,
                "Mean Squared Error (MSE)": mse,
                "R-squared (R²)": r2,
                "Adjusted R-squared": adjusted_r2
            }
            # 为每个自变量添加系数、t值、p值
            for i, name in enumerate(feature_names):
                model_data[f"Coefficient ({name})"] = coefficients[i]
                model_data[f"t-value ({name})"] = t_values[i]
                model_data[f"p-value ({name})"] = p_values[i]

            # 转换为DataFrame并转置
            df_result = pd.DataFrame([model_data]).transpose()
            # 重命名列名
            df_result.columns = ["Value"]

            # 准备解释说明
            explanations = LANGUAGES[self.current_language]['explanation']
            explanation_data = []
            # 为每个指标生成解释
            for index in df_result.index:
                if index == "Model":
                    explanation = "解释说明" if self.current_language == 'zh' else "Explanation"
                elif index == LANGUAGES[self.current_language]['regression_equation']:
                    explanation = explanations["Regression Equation"]
                elif "Coefficient" in index:
                    explanation = explanations["Coefficients"]
                elif "t-value" in index:
                    explanation = explanations["t-value"]
                elif "p-value" in index:
                    explanation = explanations["p-value"]
                else:
                    explanation = explanations.get(index, "")
                explanation_data.append(explanation)

            # 添加解释列为新的列
            df_result["Explanation"] = explanation_data

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                doc = Document()
                doc.add_heading(LANGUAGES[self.current_language]["title"], level=1)

                # 添加结果表格（转置后的表格）
                doc.add_paragraph(LANGUAGES[self.current_language]['regression_results'], style='Heading 2')
                # 行：指标数 + 表头，列：3列（指标名、值、解释）
                table = doc.add_table(rows=df_result.shape[0] + 1, cols=3)
                # 表头
                table.rows[0].cells[0].text = "指标" if self.current_language == 'zh' else "Indicator"
                table.rows[0].cells[1].text = "数值" if self.current_language == 'zh' else "Value"
                table.rows[0].cells[2].text = "解释说明" if self.current_language == 'zh' else "Explanation"

                # 数据行
                for row_idx, (index, row) in enumerate(df_result.iterrows()):
                    table.rows[row_idx + 1].cells[0].text = str(index)
                    # 格式化数值（回归方程不格式化）
                    if index == LANGUAGES[self.current_language]['regression_equation']:
                        table.rows[row_idx + 1].cells[1].text = str(row['Value'])
                    else:
                        table.rows[row_idx + 1].cells[1].text = f"{row['Value']:.4f}" if isinstance(row['Value'],
                                                                                                    float) else str(
                            row['Value'])
                    table.rows[row_idx + 1].cells[2].text = row['Explanation']

                # 获取保存目录
                save_dir = os.path.dirname(save_path)

                # 生成图表
                # 1. 实际值vs预测值
                plt.figure(figsize=(10, 6))
                plt.scatter(y, y_pred, alpha=0.6)
                plt.plot([y.min(), y.max()], [y.min(), y.max()], 'r--', lw=2)
                plt.xlabel('实际值' if self.current_language == 'zh' else 'Actual Values')
                plt.ylabel('预测值' if self.current_language == 'zh' else 'Predicted Values')
                plt.title('实际值 vs 预测值' if self.current_language == 'zh' else 'Actual vs Predicted Values')
                scatter_path = os.path.join(save_dir, "actual_vs_predicted.png")
                plt.savefig(scatter_path)
                plt.close()

                # 2. 残差图
                plt.figure(figsize=(10, 6))
                plt.scatter(y_pred, residuals, alpha=0.6)
                plt.axhline(y=0, color='r', linestyle='--', lw=2)
                plt.xlabel('预测值' if self.current_language == 'zh' else 'Predicted Values')
                plt.ylabel('残差' if self.current_language == 'zh' else 'Residuals')
                plt.title('残差图' if self.current_language == 'zh' else 'Residual Plot')
                residual_path = os.path.join(save_dir, "residual_plot.png")
                plt.savefig(residual_path)
                plt.close()

                # 3. 残差直方图
                plt.figure(figsize=(10, 6))
                plt.hist(residuals, bins=15, alpha=0.6)
                plt.xlabel('残差' if self.current_language == 'zh' else 'Residuals')
                plt.ylabel('频率' if self.current_language == 'zh' else 'Frequency')
                plt.title('残差分布' if self.current_language == 'zh' else 'Residual Distribution')
                hist_path = os.path.join(save_dir, "residual_histogram.png")
                plt.savefig(hist_path)
                plt.close()

                # 插入图表到Word
                doc.add_paragraph(LANGUAGES[self.current_language]['analysis_charts'], style='Heading 2')
                doc.add_picture(scatter_path, width=Inches(6))
                doc.add_picture(residual_path, width=Inches(6))
                doc.add_picture(hist_path, width=Inches(6))

                # 保存文档
                doc.save(save_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(save_path)
                result_msg += LANGUAGES[self.current_language]['images_saved'].format(save_dir)
                self.result_label.config(text=result_msg, wraplength=400)
            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event):
        self.current_language = "zh" if self.current_language == "en" else "en"
        self.root.title(LANGUAGES[self.current_language]["title"])
        self.select_button.config(text=LANGUAGES[self.current_language]["select_button"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=LANGUAGES[self.current_language]["analyze_button"])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]["switch_language"])
        self.open_excel_label.config(text=LANGUAGES[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        self.select_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["select_button"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        self.analyze_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["analyze_button"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=LANGUAGES[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        self.switch_language_label = ttk.Label(frame, text=LANGUAGES[self.current_language]["switch_language"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        self.root.mainloop()


def run_app():
    app = RobustLinearRegressionAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()
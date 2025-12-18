import warnings

warnings.filterwarnings("ignore", category=UserWarning, module="docx.styles.styles")
import tkinter as tk
from tkinter import filedialog, messagebox
import os
import numpy as np
import pandas as pd
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = ['SimHei', 'WenQuanYi Micro Hei', 'Heiti TC', 'Arial Unicode MS']
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典 - 包含所有需要翻译的文本
languages = {
    "zh": {
        "title": "德尔菲专家法",
        "select_button": "选择文件",
        "analyze_button": "分析文件",
        "file_not_found": "请选择文件。",
        "analysis_success": "分析完成，结果已保存到 {}\n",
        "no_save_path": "未选择保存路径，结果未保存。",
        "analysis_error": "分析文件时出错: {}",
        "switch_language": "中/英",
        'open_excel_button_text': "示例数据",
        "file_entry_placeholder": "请输入待分析 Excel 文件的完整路径",
        "statistic_labels": {
            "mean_per_round": "各轮评分均值",
            "std_per_round": "各轮评分标准差",
            "final_consensus": "最终共识评分",
            "convergence_status": "评分收敛情况",
            "data_range": "数据范围",
            "conv_threshold": "收敛阈值",
            "expert_deviation": "专家意见偏离度"
        },
        "explanation": {
            "mean_per_round": "每一轮专家评分的平均值",
            "std_per_round": "每一轮专家评分的离散程度",
            "final_consensus": "经过多轮反馈后，专家达成共识的评分",
            "convergence_status": "判断专家评分是否达到收敛标准",
            "expert_deviation": "各专家与群体均值的平均偏离程度"
        },
        "interpretation": {
            "mean_per_round": "反映每一轮专家对问题的整体评价",
            "std_per_round": "标准差越小，说明专家意见越集中",
            "final_consensus": "作为最终的决策参考",
            "convergence_status": "若收敛，则表示专家意见达成一致；否则，需要进一步讨论",
            "expert_deviation": "值越大表示该专家意见与群体差异越显著"
        },
        "chart_titles": {
            "last_round_hist": "最后一轮（{}）评分分布柱状图",
            "mean_trend": "各轮评分均值变化趋势",
            "std_trend": "各轮评分标准差变化趋势",
            "boxplot": "各轮评分分布对比",
            "deviation": "专家意见偏离度"
        },
        "axis_labels": {
            "scores": "评分",
            "frequency": "频数",
            "round": "轮次",
            "mean": "均值",
            "std_dev": "标准差",
            "experts": "专家",
            "avg_deviation": "平均偏离度"
        },
        "other_texts": {
            "explanation_heading": "解释说明",
            "interpretation_heading": "结果解读",
            "convergence_threshold_label": "收敛阈值: {:.2f}",
            "headers": ["统计量", "统计量值", "p值"],
            "stat_explanation": "统计量_解释说明",
            "stat_interpretation": "统计量_结果解读"
        }
    },
    "en": {
        "title": "Delphi Method",
        "select_button": "Select File",
        "analyze_button": "Analyze File",
        "file_not_found": "Please select a file.",
        "analysis_success": "Analysis completed. Results saved to {}\n",
        "no_save_path": "No save path selected. Results not saved.",
        "analysis_error": "Error analyzing file: {}",
        "switch_language": "Chinese/English",
        'open_excel_button_text': "Example data",
        "file_entry_placeholder": "Please enter the full path of the Excel file to analyze",
        "statistic_labels": {
            "mean_per_round": "Average Score per Round",
            "std_per_round": "Standard Deviation per Round",
            "final_consensus": "Final Consensus Score",
            "convergence_status": "Score Convergence Status",
            "data_range": "Data Range",
            "conv_threshold": "Convergence Threshold",
            "expert_deviation": "Expert Opinion Deviation"
        },
        "explanation": {
            "mean_per_round": "The average score of experts in each round",
            "std_per_round": "The dispersion degree of experts' scores in each round",
            "final_consensus": "The consensus score reached by experts after multiple rounds of feedback",
            "convergence_status": "Determine whether the experts' scores have reached the convergence criterion",
            "expert_deviation": "Average deviation of each expert from the group mean"
        },
        "interpretation": {
            "mean_per_round": "Reflects the overall evaluation of experts on the problem in each round",
            "std_per_round": "The smaller the standard deviation, the more concentrated the experts' opinions",
            "final_consensus": "Serves as the final decision-making reference",
            "convergence_status": "If converged, it indicates that the experts have reached an agreement; otherwise, further discussion is required",
            "expert_deviation": "Larger values indicate greater divergence from the group opinion"
        },
        "chart_titles": {
            "last_round_hist": "Histogram of Scores in Last Round ({})",
            "mean_trend": "Trend of Mean Scores by Round",
            "std_trend": "Trend of Standard Deviation by Round",
            "boxplot": "Score Distribution by Round",
            "deviation": "Expert Opinion Deviation"
        },
        "axis_labels": {
            "scores": "Scores",
            "frequency": "Frequency",
            "round": "Round",
            "mean": "Mean",
            "std_dev": "Std Deviation",
            "experts": "Experts",
            "avg_deviation": "Avg Deviation"
        },
        "other_texts": {
            "explanation_heading": "Explanation",
            "interpretation_heading": "Interpretation",
            "convergence_threshold_label": "Convergence Threshold: {:.2f}",
            "headers": ["Statistic", "Value", "p-value"],
            "stat_explanation": "Statistic_Explanation",
            "stat_interpretation": "Statistic_Interpretation"
        }
    }
}


class DelphiMethodAnalysisApp:
    def __init__(self, root=None):
        # 初始化根窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages["zh"]["title"])  # 初始使用中文标题
        else:
            self.root = root
            self.root.title(languages["zh"]["title"])

        # 设置默认参数
        self.current_language = "en"  # 当前语言，默认为中文
        self.data_format = "row"  # 默认行代表轮次
        self.convergence_threshold_pct = 20.0  # 默认收敛阈值20%
        self.weights = [0.2, 0.3, 0.5]  # 默认轮次权重（最后3轮）
        self.create_ui()

    def center_dialog(self, dialog):
        """将对话框居中显示在屏幕上"""
        dialog.update_idletasks()

        # 获取屏幕和对话框的尺寸
        screen_width = dialog.winfo_screenwidth()
        screen_height = dialog.winfo_screenheight()
        dialog_width = dialog.winfo_width()
        dialog_height = dialog.winfo_height()

        # 计算居中位置
        x = (screen_width - dialog_width) // 2
        y = (screen_height - dialog_height) // 2

        # 设置对话框位置
        dialog.geometry(f"+{x}+{y}")

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data45.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 保存主窗口当前的topmost状态
        original_topmost = self.root.attributes("-topmost")

        # 确保主窗口暂时不置顶，避免遮挡对话框
        self.root.attributes("-topmost", False)

        # 打开文件对话框
        file_path = filedialog.askopenfilename(
            parent=self.root,
            filetypes=[("Excel files", "*.xlsx;*.xls")],
            title=languages[self.current_language]["select_button"]
        )

        # 恢复主窗口原来的topmost状态
        self.root.attributes("-topmost", original_topmost)
        self.root.lift()

        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def calculate_expert_deviation(self, data, means):
        """计算专家意见偏离度"""
        deviations = []
        for expert_idx in range(data.shape[1]):
            # 计算该专家每轮评分与对应轮次均值的绝对差
            expert_scores = [data[round_idx][expert_idx] for round_idx in range(len(means))]
            abs_diff = [abs(score - means[round_idx]) for round_idx, score in enumerate(expert_scores)]
            deviations.append(np.mean(abs_diff))  # 平均偏离度
        return deviations

    def delphi_analysis(self, data):
        """
        进行德尔菲专家法分析
        :param data: 专家多轮评分数据，每一行代表一轮评分
        :return: 各轮评分均值、各轮评分标准差、最终共识评分、评分收敛情况、原始数据范围、专家偏离度等
        """
        # 使用原始数据，不进行缺失值填充和异常值处理
        data_with_nan = np.array(data, dtype=np.float64)  # 转换为float以支持NaN
        cleaned_data = data_with_nan

        # 计算每轮均值和标准差
        means = np.array([np.nanmean(round_data) for round_data in cleaned_data])
        stds = np.array([np.nanstd(round_data) for round_data in cleaned_data])

        # 计算专家意见偏离度
        expert_deviations = self.calculate_expert_deviation(np.array(cleaned_data), means)

        # 计算整体数据范围（用于动态阈值）
        all_scores = np.concatenate(cleaned_data)
        data_range = np.nanmax(all_scores) - np.nanmin(all_scores) if len(all_scores) > 0 else 1

        # 动态计算收敛阈值
        convergence_threshold = (self.convergence_threshold_pct / 100) * data_range

        # 计算最终共识评分（动态权重）
        n_rounds = len(means)
        if n_rounds == 1:
            # 仅一轮时直接使用该轮均值
            final_consensus_score = means[0]
        elif n_rounds == 2:
            # 两轮时使用简单加权
            final_consensus_score = 0.3 * means[0] + 0.7 * means[1]
        else:
            # 三轮及以上使用动态权重
            selected_weights = self.weights[-3:]
            weight_sum = sum(selected_weights)
            normalized_weights = [w / weight_sum for w in selected_weights]
            final_consensus_score = np.sum(means[-3:] * normalized_weights)

        # 收敛判断
        convergence = False
        if len(stds) >= 2:
            # 条件1：最后一轮标准差低于阈值
            last_one_below = (stds[-1] < convergence_threshold)

            # 条件2：最后两轮标准差变化率小于10%
            std_change_rate = abs(stds[-1] - stds[-2]) / stds[-2] if stds[-2] != 0 else 0
            small_change = std_change_rate < 0.15

            # 条件3：整体呈非递增趋势
            trend_stable = True
            if len(stds) >= 3:
                trend_stable = (stds[-3] >= stds[-2] and stds[-2] >= stds[-1])
            else:
                trend_stable = (stds[-2] >= stds[-1])

            # 同时满足三个放宽后的条件即判定为收敛
            convergence = (last_one_below and small_change) or \
                          (last_one_below and trend_stable) or \
                          (small_change and trend_stable)
        return means, stds, final_consensus_score, convergence, data_range, convergence_threshold, expert_deviations

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]['file_not_found'])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件，第一行作为表头（专家标识），保留空值为NaN
            df = pd.read_excel(file_path, header=0, na_values=['', ' '])

            # 提取第一列作为轮次标识（不参与计算），从第二列开始为评分数据
            if df.shape[1] < 2:
                raise ValueError("数据格式不正确，需要至少包含1列轮次标识和1列专家评分" if self.current_language == "zh"
                                 else "Incorrect data format, need at least 1 round identifier column and 1 expert score column")

            # 分离轮次标识和评分数据
            round_labels = df.iloc[:, 0].tolist()  # 第一列作为轮次标识
            expert_labels = df.columns[1:].tolist()  # 第一行（表头）作为专家标识
            data = df.iloc[:, 1:].values  # 从第二列开始为评分数据

            # 数据格式验证与转换
            if self.data_format == "col":
                data = data.T

            # 验证数据有效性
            if data.size == 0:
                raise ValueError("文件中没有有效评分数据" if self.current_language == "zh"
                                 else "No valid score data in the file")
            if len(data.shape) != 2:
                raise ValueError("数据格式应为二维表格" if self.current_language == "zh"
                                 else "Data format should be a 2D table")
            if data.shape[0] < 1 or data.shape[1] < 1:
                raise ValueError("数据维度不正确，需要至少1轮和1位专家" if self.current_language == "zh"
                                 else "Incorrect data dimensions, need at least 1 round and 1 expert")

            # 进行德尔菲分析
            (means, stds, final_consensus_score, convergence,
             data_range, conv_threshold, expert_deviations) = self.delphi_analysis(data)

            # 获取当前语言的统计标签
            stats = languages[self.current_language]["statistic_labels"]

            # 整理数据（包含轮次标识和专家标识）
            result_data = [
                [stats["mean_per_round"], [f"{round_labels[i]}: {round(m, 2)}" for i, m in enumerate(means.tolist())],
                 ""],
                [stats["std_per_round"], [f"{round_labels[i]}: {round(s, 2)}" for i, s in enumerate(stds.tolist())],
                 ""],
                [stats["final_consensus"], [round(final_consensus_score, 2)], ""],
                [stats["convergence_status"],
                 ["是" if convergence else "否" if self.current_language == "zh" else "Yes" if convergence else "No"],
                 ""],
                [stats["data_range"], [f"{np.nanmin(np.concatenate(data)):.1f}-{np.nanmax(np.concatenate(data)):.1f}"],
                 ""],
                [stats["conv_threshold"], [f"{conv_threshold:.2f}"], ""],
                [stats["expert_deviation"],
                 [f"{expert_labels[i]}: {round(d, 2)}" for i, d in enumerate(expert_deviations)], ""]
            ]
            headers = languages[self.current_language]["other_texts"]["headers"]
            df_result = pd.DataFrame(result_data, columns=headers)

            # 添加解释说明
            explanations = languages[self.current_language]['explanation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(
                columns=[stats["mean_per_round"], stats["std_per_round"], stats["final_consensus"],
                         stats["convergence_status"], stats["expert_deviation"]])
            explanation_df.insert(0, languages[self.current_language]["other_texts"]["stat_explanation"],
                                  languages[self.current_language]["other_texts"]["explanation_heading"])

            # 添加分析结果解读
            interpretations = languages[self.current_language]['interpretation']
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(
                columns=[stats["mean_per_round"], stats["std_per_round"], stats["final_consensus"],
                         stats["convergence_status"], stats["expert_deviation"]])
            interpretation_df.insert(0, languages[self.current_language]["other_texts"]["stat_interpretation"],
                                     languages[self.current_language]["other_texts"]["interpretation_heading"])

            # 创建临时窗口用于居中显示保存对话框
            temp_root = tk.Toplevel(self.root)
            temp_root.withdraw()  # 隐藏临时窗口
            self.center_dialog(temp_root)  # 将临时窗口居中

            # 让用户选择保存路径，使用临时窗口作为父窗口
            save_path = filedialog.asksaveasfilename(
                parent=temp_root,
                defaultextension=".docx",
                filetypes=[("Word files", "*.docx")]
            )

            # 销毁临时窗口
            temp_root.destroy()

            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加标题
                doc.add_heading(languages[self.current_language]['title'], 0)

                # 添加分析结果表格
                table = doc.add_table(rows=1, cols=len(headers))
                hdr_cells = table.rows[0].cells
                for col, header in enumerate(headers):
                    hdr_cells[col].text = header
                for row in df_result.values:
                    row_cells = table.add_row().cells
                    for col, value in enumerate(row):
                        row_cells[col].text = str(value)

                # 添加解释说明（项目符号列表）
                doc.add_heading(languages[self.current_language]["other_texts"]["explanation_heading"], 1)
                stats = languages[self.current_language]["statistic_labels"]
                for key, value in explanations.items():
                    para = doc.add_paragraph(style='ListBullet')
                    # 使用统计标签的中文名称替代键名
                    para.add_run(f"{stats[key]}: ").bold = True
                    para.add_run(value)

                # 添加结果解读（项目符号列表）
                doc.add_heading(languages[self.current_language]["other_texts"]["interpretation_heading"], 1)
                for key, value in interpretations.items():
                    para = doc.add_paragraph(style='ListBullet')
                    # 使用统计标签的中文名称替代键名
                    para.add_run(f"{stats[key]}: ").bold = True
                    para.add_run(value)

                # 生成图表
                img_paths = []
                chart_titles = languages[self.current_language]["chart_titles"]
                axis_labels = languages[self.current_language]["axis_labels"]

                # 最后一轮评分分布柱状图
                fig, ax = plt.subplots()
                ax.hist(data[-1][~np.isnan(data[-1])], bins=10)
                ax.set_title(chart_titles["last_round_hist"].format(round_labels[-1]))
                ax.set_xlabel(axis_labels["scores"])
                ax.set_ylabel(axis_labels["frequency"])
                img_path1 = os.path.splitext(save_path)[0] + '_last_round_hist.png'
                plt.savefig(img_path1)
                plt.close()
                img_paths.append(img_path1)

                # 多轮均值变化趋势图
                fig, ax = plt.subplots()
                ax.plot(range(1, len(means) + 1), means, marker='o', linestyle='-')
                ax.set_xticks(range(1, len(means) + 1))
                ax.set_xticklabels(round_labels, rotation=0)
                ax.set_title(chart_titles["mean_trend"])
                ax.set_xlabel(axis_labels["round"])
                ax.set_ylabel(axis_labels["mean"])
                ax.grid(True)
                plt.tight_layout()
                img_path2 = os.path.splitext(save_path)[0] + '_mean_trend.png'
                plt.savefig(img_path2)
                plt.close()
                img_paths.append(img_path2)

                # 多轮标准差变化趋势图
                fig, ax = plt.subplots()
                ax.plot(range(1, len(stds) + 1), stds, marker='s', linestyle='-', color='orange')
                ax.axhline(y=conv_threshold, color='r', linestyle='--',
                           label=languages[self.current_language]["other_texts"]["convergence_threshold_label"].format(
                               conv_threshold))
                ax.set_xticks(range(1, len(stds) + 1))
                ax.set_xticklabels(round_labels, rotation=0)
                ax.set_title(chart_titles["std_trend"])
                ax.set_xlabel(axis_labels["round"])
                ax.set_ylabel(axis_labels["std_dev"])
                ax.legend()
                ax.grid(True)
                plt.tight_layout()
                img_path3 = os.path.splitext(save_path)[0] + '_std_trend.png'
                plt.savefig(img_path3)
                plt.close()
                img_paths.append(img_path3)

                # 每轮评分箱线图对比
                fig, ax = plt.subplots()
                ax.boxplot([round_data[~np.isnan(round_data)] for round_data in data])
                ax.set_xticklabels(round_labels, rotation=0)
                ax.set_title(chart_titles["boxplot"])
                ax.set_xlabel(axis_labels["round"])
                ax.set_ylabel(axis_labels["scores"])
                ax.grid(True)
                plt.tight_layout()
                img_path4 = os.path.splitext(save_path)[0] + '_boxplot.png'
                plt.savefig(img_path4)
                plt.close()
                img_paths.append(img_path4)

                # 专家意见偏离度柱状图
                fig, ax = plt.subplots()
                ax.bar(expert_labels, expert_deviations)
                ax.set_title(chart_titles["deviation"])
                ax.set_xlabel(axis_labels["experts"])
                ax.set_ylabel(axis_labels["avg_deviation"])
                plt.xticks(rotation=0)
                ax.grid(True, axis='y')
                plt.tight_layout()
                img_path5 = os.path.splitext(save_path)[0] + '_deviation.png'
                plt.savefig(img_path5)
                plt.close()
                img_paths.append(img_path5)

                # 将图片插入到 Word 文档中
                for i, img_path in enumerate(img_paths):
                    if i == 0:
                        doc.add_heading(chart_titles["last_round_hist"].format(round_labels[-1]), 1)
                    elif i == 1:
                        doc.add_heading(chart_titles["mean_trend"], 1)
                    elif i == 2:
                        doc.add_heading(chart_titles["std_trend"], 1)
                    elif i == 3:
                        doc.add_heading(chart_titles["boxplot"], 1)
                    else:
                        doc.add_heading(chart_titles["deviation"], 1)
                    doc.add_picture(img_path, width=Inches(6))

                # 保存 Word 文件
                doc.save(save_path)

                # 清理临时图片文件
                for img_path in img_paths:
                    if os.path.exists(img_path):
                        os.remove(img_path)

                result_msg = languages[self.current_language]['analysis_success'].format(save_path)
                self.result_label.config(text=result_msg, wraplength=400)

            else:
                self.result_label.config(text=languages[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]['analysis_error'].format(str(e)))
            messagebox.showerror("Error", str(e))

    def switch_language(self, event=None):
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'
        self.root.title(languages[self.current_language]['title'])
        self.select_button.config(text=languages[self.current_language]['select_button'])
        self.analyze_button.config(text=languages[self.current_language]['analyze_button'])
        self.switch_language_label.config(text=languages[self.current_language]['switch_language'])
        self.open_excel_label.config(text=languages[self.current_language]["open_excel_button_text"])

        # 切换语言时更新提示信息
        current_text = self.file_entry.get()
        placeholder = languages[self.current_language]['file_entry_placeholder']
        other_placeholder = languages['zh' if self.current_language == 'en' else 'en']['file_entry_placeholder']
        if current_text == other_placeholder:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, placeholder)
            self.file_entry.config(foreground='gray')

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建框架
        main_frame = ttk.Frame(self.root, padding=10)
        main_frame.pack(expand=True, fill=tk.BOTH, anchor="center")

        inner_frame = ttk.Frame(main_frame)
        inner_frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(inner_frame, text=languages[self.current_language]["select_button"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(inner_frame, width=70)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(inner_frame, text=languages[self.current_language]["analyze_button"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=20)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            inner_frame,
            text=languages[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(inner_frame, text=languages[self.current_language]["switch_language"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=5)

        # 创建结果显示标签
        self.result_label = ttk.Label(inner_frame, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10, fill=tk.X)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = DelphiMethodAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()
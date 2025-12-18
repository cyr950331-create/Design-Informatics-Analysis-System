import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from scipy import stats

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = ['SimHei', 'WenQuanYi Micro Hei', 'Heiti TC']
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典，包含所有需要翻译的文本
languages = {
    'zh': {
        'title': "CRITIC 权重法",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "请选择文件。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "中/英",
        'open_excel_button_text': "示例数据",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'missing_value': "数据中存在缺失值，请处理后重试。",
        'data_dim_error': "数据维度错误，至少需要2个指标（列）和2个样本（行）。",
        'zero_info_error': "所有指标信息量均为0，无法计算权重（可能所有指标值相同）。",
        'statistic': "统计量",
        'statistic_value': "统计量值",
        'analysis_results': "分析结果",
        'explanations': "解释说明",
        'interpretations': "结果解读",
        'row_title': "行标题",
        'col_title': "列标题",
        'data': "数据",
        'indicator_title': "指标标题",
        'weight_sum': "权重和",
        'weight_chart': "指标权重柱状图",
        'std_chart': "指标标准差柱状图",
        'corr_chart': "相关系数矩阵热力图",
        'indicator': "指标",
        'indicator_weight': "指标权重",
        'standard_deviation': "标准差",
        'original_data_matrix': "原始数据矩阵",
        'std_matrix': "标准差矩阵",
        'corr_matrix': "相关系数矩阵",
        'corr_significance': "相关系数显著性",
        'info_matrix': "信息量矩阵",
        'indicator_weights': "指标权重",
        'explanation': {
            "original_data_matrix": "从 Excel 文件中读取的原始数据矩阵",
            "std_matrix": "各指标的标准差矩阵，反映指标的对比强度",
            "corr_matrix": "各指标之间的相关系数矩阵，反映指标之间的冲突性",
            "corr_significance": "各指标间相关系数的统计显著性（p值）",
            "info_matrix": "结合标准差和相关系数计算得到的各指标信息量矩阵",
            "indicator_weights": "根据信息量矩阵计算得到的各指标权重"
        },
        'interpretation': {
            "original_data_matrix": "用于后续分析的基础数据",
            "std_matrix": "标准差越大，该指标的对比强度越大，在综合评价中越重要",
            "corr_matrix": "相关系数越小，指标之间的冲突性越大，该指标在综合评价中越重要",
            "corr_significance": "p值 < 0.05 表示相关关系显著，否则不显著",
            "info_matrix": "反映各指标包含的信息量，信息量越大，该指标越重要",
            "indicator_weights": "各指标在综合评价中的相对重要程度，权重越大越重要"
        }
    },
    'en': {
        'title': "CRITIC Weight Method",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Chinese/English",
        'open_excel_button_text': "Example data",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'missing_value': "Missing values found in data. Please handle them and try again.",
        'data_dim_error': "Data dimension error: at least 2 indicators (columns) and 2 samples (rows) are required.",
        'zero_info_error': "All indicators have zero information. Cannot calculate weights (possibly all indicator values are the same).",
        'statistic': "Statistic",
        'statistic_value': "Statistic Value",
        'analysis_results': "Analysis Results",
        'explanations': "Explanations",
        'interpretations': "Interpretations",
        'row_title': "Row Titles",
        'col_title': "Column Titles",
        'data': "Data",
        'indicator_title': "Indicator Titles",
        'weight_sum': "Sum of Weights",
        'weight_chart': "Bar Chart of Indicator Weights",
        'std_chart': "Bar Chart of Indicator Std Dev",
        'corr_chart': "Correlation Coefficient Heatmap",
        'indicator': "Indicator",
        'indicator_weight': "Indicator Weight",
        'standard_deviation': "Standard Deviation",
        'original_data_matrix': "Original Data Matrix",
        'std_matrix': "Standard Deviation Matrix",
        'corr_matrix': "Correlation Coefficient Matrix",
        'corr_significance': "Correlation Significance",
        'info_matrix': "Information Matrix",
        'indicator_weights': "Indicator Weights",
        'explanation': {
            "original_data_matrix": "The original data matrix read from the Excel file",
            "std_matrix": "The standard deviation matrix of each indicator, reflecting the contrast intensity of the indicators",
            "corr_matrix": "The correlation coefficient matrix between each indicator, reflecting the conflict between the indicators",
            "corr_significance": "Statistical significance (p-value) of correlation coefficients between indicators",
            "info_matrix": "The information matrix of each indicator calculated by combining the standard deviation and correlation coefficient",
            "indicator_weights": "The weight of each indicator calculated based on the information matrix"
        },
        'interpretation': {
            "original_data_matrix": "The basic data for subsequent analysis",
            "std_matrix": "The larger the standard deviation, the greater the contrast intensity of the indicator, and the more important it is in the comprehensive evaluation",
            "corr_matrix": "The smaller the correlation coefficient, the greater the conflict between the indicators, and the more important the indicator is in the comprehensive evaluation",
            "corr_significance": "p-value < 0.05 indicates a significant correlation, otherwise not significant",
            "info_matrix": "Reflects the information contained in each indicator. The greater the information, the more important the indicator",
            "indicator_weights": "The relative importance of each indicator in the comprehensive evaluation. The larger the weight, the more important it is"
        }
    }
}


class CRITICWeightingMethodAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为中文
        self.current_language = "en"
        self.file_path = ""  # 保存当前文件路径

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data28.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"{languages[self.current_language]['file_not_found']}：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"{languages[self.current_language]['analysis_error'].format(str(e))}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
            self.file_path = file_path  # 选择文件后立即更新文件路径
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        current_text = self.file_entry.get()
        if current_text == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        current_text = self.file_entry.get()
        if current_text == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')
        else:
            self.file_path = current_text  # 更新文件路径

    def critic_weight_method(self, data):
        # 新增数据标准化处理
        from sklearn.preprocessing import StandardScaler
        scaler = StandardScaler()
        data_standardized = scaler.fit_transform(data)  # 标准化数据
        data = data_standardized

        # 计算标准差矩阵
        std_matrix = np.std(data, axis=0, ddof=1)

        # 计算相关系数矩阵
        corr_matrix = np.corrcoef(data, rowvar=False)

        # 计算p值矩阵（使用更高效的向量化方法）
        n = data.shape[0]
        pvalue_matrix = np.zeros_like(corr_matrix)
        for i in range(corr_matrix.shape[0]):
            for j in range(i, corr_matrix.shape[1]):
                if i == j:
                    pvalue_matrix[i, j] = 0.0
                else:
                    corr, p_val = stats.pearsonr(data[:, i], data[:, j])
                    pvalue_matrix[i, j] = p_val
                    pvalue_matrix[j, i] = p_val  # 对称矩阵

        # 计算冲突性（仅考虑显著相关的指标）
        # 对于p值>0.05的相关系数，视为不显著，冲突性设为1（最大冲突）
        conflict = np.where(pvalue_matrix > 0.05, 1.0, 1 - corr_matrix)

        # 计算信息量矩阵
        info_matrix = std_matrix * np.sum(conflict, axis=0)

        # 权重计算及校验
        total_info = np.sum(info_matrix)
        # 添加数值稳定性处理
        epsilon = 1e-10
        if np.isclose(total_info, 0, atol=epsilon):
            raise ValueError("zero_info")

        # 防止数值溢出
        info_matrix = np.clip(info_matrix, epsilon, None)
        total_info = np.sum(info_matrix)
        weights = info_matrix / total_info

        # 确保权重和为1（处理浮点误差）
        weights = weights / np.sum(weights)

        return std_matrix, corr_matrix, pvalue_matrix, info_matrix, weights

    def analyze_file(self):
        file_path = self.file_path
        if not file_path or file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["file_not_found"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]['file_not_found'])
            return
        try:
            # 固定设置：只读取第一行作为表头，不读取第一列作为表头
            header_row = True
            header_col = False

            # 读取Excel文件（根据固定表头设置处理）
            if header_row and header_col:
                df = pd.read_excel(file_path, header=0, index_col=0)
                data = df.values
                # 获取行列标题
                col_names = df.columns.tolist()
                row_names = df.index.tolist()
            elif header_row:
                df = pd.read_excel(file_path, header=0, index_col=None)
                data = df.values
                col_names = df.columns.tolist()
                row_names = [f"{languages[self.current_language]['indicator']} {i+1}" for i in range(data.shape[0])]
            elif header_col:
                df = pd.read_excel(file_path, header=None, index_col=0)
                data = df.values
                col_names = [f"{languages[self.current_language]['indicator']} {i+1}" for i in range(data.shape[1])]
                row_names = df.index.tolist()
            else:
                df = pd.read_excel(file_path, header=None, index_col=None)
                data = df.values
                col_names = [f"{languages[self.current_language]['indicator']} {i+1}" for i in range(data.shape[1])]
                row_names = [f"{languages[self.current_language]['indicator']} {i+1}" for i in range(data.shape[0])]

            # 缺失值检测
            if np.isnan(data).any():
                self.result_label.config(text=languages[self.current_language]['missing_value'])
                return

            # 数据类型转换
            data = data.astype(float)

            # 数据维度校验
            if data.shape[0] < 2 or data.shape[1] < 2:
                self.result_label.config(text=languages[self.current_language]['data_dim_error'])
                return

            # 进行 CRITIC 权重法分析
            std_matrix, corr_matrix, pvalue_matrix, info_matrix, weights = self.critic_weight_method(data)

            # 整理数据，包含表头信息
            data = [
                [languages[self.current_language]['original_data_matrix'],
                 f"{languages[self.current_language]['row_title']}: {row_names}\n{languages[self.current_language]['col_title']}: {col_names}\n{languages[self.current_language]['data']}: {np.round(data, 6).tolist()}"],
                [languages[self.current_language]['std_matrix'],
                 f"{languages[self.current_language]['indicator_title']}: {col_names}\n{languages[self.current_language]['data']}: {np.round(std_matrix, 6).tolist()}"],
                [languages[self.current_language]['corr_matrix'],
                 f"{languages[self.current_language]['indicator_title']}: {col_names}\n{languages[self.current_language]['data']}: {np.round(corr_matrix, 6).tolist()}"],
                [languages[self.current_language]['corr_significance'],
                 f"{languages[self.current_language]['indicator_title']}: {col_names}\n{languages[self.current_language]['data']}: {np.round(pvalue_matrix, 6).tolist()}"],
                [languages[self.current_language]['info_matrix'],
                 f"{languages[self.current_language]['indicator_title']}: {col_names}\n{languages[self.current_language]['data']}: {np.round(info_matrix, 6).tolist()}"],
                [languages[self.current_language]['indicator_weights'],
                 f"{languages[self.current_language]['indicator_title']}: {col_names}\n{languages[self.current_language]['data']}: {np.round(weights, 6).tolist()}\n{languages[self.current_language]['weight_sum']}: {np.round(np.sum(weights), 6)}"]
            ]
            headers = [languages[self.current_language]['statistic'], languages[self.current_language]['statistic_value']]
            df_result = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            explanations = languages[self.current_language]['explanation']
            interpretations = languages[self.current_language]['interpretation']

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加分析结果表格
                doc.add_heading(languages[self.current_language]['analysis_results'], level=1)
                table = doc.add_table(rows=1, cols=len(headers))
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                for _, row in df_result.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)

                # 添加解释说明
                doc.add_heading(languages[self.current_language]['explanations'], level=1)
                for key in ["original_data_matrix", "std_matrix", "corr_matrix", "corr_significance", "info_matrix", "indicator_weights"]:
                    para = doc.add_paragraph(style='List Bullet')
                    para.add_run(f"{languages[self.current_language][key]}: ").bold = True
                    para.add_run(explanations[key])

                # 添加结果解读
                doc.add_heading(languages[self.current_language]['interpretations'], level=1)
                for key in ["original_data_matrix", "std_matrix", "corr_matrix", "corr_significance", "info_matrix", "indicator_weights"]:
                    para = doc.add_paragraph(style='List Bullet')
                    para.add_run(f"{languages[self.current_language][key]}: ").bold = True
                    para.add_run(interpretations[key])

                # 生成可视化图表
                base_img_path = os.path.splitext(save_path)[0]

                # 1. 权重柱状图 - 添加指标标题
                fig, ax = plt.subplots(figsize=(12, 8))  # 更大的图表尺寸
                ax.bar(range(len(weights)), weights, color='skyblue', edgecolor='black')
                ax.set_title(
                    languages[self.current_language]['weight_chart'],
                    fontsize=14, pad=20)
                ax.set_xlabel(languages[self.current_language]['indicator'], fontsize=12)
                ax.set_ylabel(languages[self.current_language]['indicator_weight'], fontsize=12)
                ax.set_xticks(range(len(col_names)))
                ax.set_xticklabels(col_names, rotation=0, ha='center')
                ax.grid(axis='y', linestyle='--', alpha=0.7)  # 添加网格线
                # 在柱状图上显示具体数值
                for i, v in enumerate(weights):
                    ax.text(i, v + 0.01, f'{v:.4f}', ha='center', fontsize=10)
                plt.tight_layout()
                img_path_weight = f"{base_img_path}_weight.png"  # 定义权重图路径
                plt.savefig(img_path_weight, dpi=300)  # 保存图片
                plt.close()  # 关闭图表

                # 2. 标准差柱状图 - 添加指标标题
                fig, ax = plt.subplots(figsize=(12, 8))
                ax.bar(range(len(std_matrix)), std_matrix)
                ax.set_title(
                    languages[self.current_language]['std_chart'])
                ax.set_xlabel(languages[self.current_language]['indicator'])
                ax.set_ylabel(languages[self.current_language]['standard_deviation'])
                ax.set_xticks(range(len(col_names)))
                ax.set_xticklabels(col_names, rotation=0, ha='center')  # 显示指标标题
                img_path_std = f"{base_img_path}_std.png"
                plt.tight_layout()
                plt.savefig(img_path_std, dpi=300)
                plt.close()

                # 3. 相关系数热力图 - 添加指标标题
                fig, ax = plt.subplots(figsize=(12, 8))
                im = ax.imshow(corr_matrix, cmap='coolwarm', vmin=-1, vmax=1)
                plt.colorbar(im)
                ax.set_title(
                    languages[self.current_language]['corr_chart'])
                ax.set_xticks(range(len(corr_matrix)))
                ax.set_yticks(range(len(corr_matrix)))
                ax.set_xticklabels(col_names, rotation=0, ha='center')  # 显示指标标题
                ax.set_yticklabels(col_names)  # 显示指标标题
                img_path_corr = f"{base_img_path}_corr.png"
                plt.tight_layout()
                plt.savefig(img_path_corr, dpi=300)
                plt.close()

                # 在 Word 文档中插入图片
                doc.add_heading(languages[self.current_language]['weight_chart'], level=1)
                doc.add_picture(img_path_weight, width=Inches(6))  # 使用保存的权重图路径

                doc.add_heading(languages[self.current_language]['std_chart'], level=1)
                doc.add_picture(img_path_std, width=Inches(6))

                doc.add_heading(languages[self.current_language]['corr_chart'], level=1)
                doc.add_picture(img_path_corr, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                result_msg = languages[self.current_language]['analysis_success'].format(save_path)
                self.result_label.config(text=result_msg, wraplength=400)

            else:
                self.result_label.config(text=languages[self.current_language]['no_save_path'])

        except ValueError as e:
            if str(e) == "zero_info":
                self.result_label.config(text=languages[self.current_language]['zero_info_error'])
            else:
                self.result_label.config(text=languages[self.current_language]['analysis_error'].format(str(e)))
        except Exception as e:
            self.result_label.config(text=languages[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event):
        # 保存当前输入的文件路径
        current_text = self.file_entry.get()
        if current_text != languages[self.current_language]["file_entry_placeholder"]:
            self.file_path = current_text

        # 切换语言
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button"])
        self.analyze_button.config(text=languages[self.current_language]["analyze_button"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language"])
        self.open_excel_label.config(text=languages[self.current_language]["open_excel_button_text"])

        # 恢复文件路径（保留用户输入）
        self.file_entry.delete(0, tk.END)
        if self.file_path:
            self.file_entry.insert(0, self.file_path)
            self.file_entry.config(foreground='black')
        else:
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True, padx=20)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=60)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=languages[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=languages[self.current_language]["switch_language"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=5)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10, padx=20)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = CRITICWeightingMethodAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()
import tkinter as tk
from tkinter import filedialog, messagebox
import os
import pandas as pd
import numpy as np
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典，使用英文键
languages = {
    'zh': {
        'title': "TOPSIS",
        'select_button_text': "选择文件",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'analyze_button_text': "分析文件",
        'no_file_selected': "请选择文件。",
        'file_not_exists': "文件不存在，请重新选择。",
        'analysis_error': "分析文件时出错: {}",
        'analysis_complete': "分析完成，结果已保存到 {}，相关图片已保存。",
        'no_save_path_selected': "未选择保存路径，结果未保存。",
        'switch_language_button_text': "中/英",
        'open_excel_button_text': "示例数据",
        'explanation': {
            'standardized_matrix': "对原始决策矩阵进行标准化处理后的矩阵",
            'weighted_matrix': "考虑各属性权重后的标准化决策矩阵",
            'positive_ideal': "各属性的最优值构成的向量",
            'negative_ideal': "各属性的最劣值构成的向量",
            'distance_positive': "各方案与正理想解的欧几里得距离",
            'distance_negative': "各方案与负理想解的欧几里得距离",
            'relative_closeness': "反映各方案与正理想解的相对接近程度",
            'ranking_result': "根据相对贴近度对各方案进行排序的结果"
        },
        'interpretation': {
            'standardized_matrix': "消除不同属性量纲的影响",
            'weighted_matrix': "体现各属性在决策中的重要性",
            'positive_ideal': "作为衡量各方案优劣的最优参考点",
            'negative_ideal': "作为衡量各方案优劣的最劣参考点",
            'distance_positive': "距离越小，方案越优",
            'distance_negative': "距离越大，方案越优",
            'relative_closeness': "值越接近 1，方案越优",
            'ranking_result': "排名越靠前，方案越优"
        },
        'error_messages': {
            'negative_weights': "权重包含负值，请检查数据。权重应为非负数。",
            'empty_data': "Excel文件数据为空，请检查文件。",
            'insufficient_rows': "数据行数不足，至少需要2行（1行权重+1行方案）。",
            'non_numeric': "数据中包含非数值内容，请检查单元格：{}",
            'weight_normalization': "权重之和不为1，已自动归一化处理。",
            'index_type_mismatch': "指标类型数量与权重数量不匹配，请检查输入。",
            'empty_factors': "第一列（因素列）包含空值，请检查数据。"
        },
        'result_headers': {
            'statistic': "统计量",
            'value': "统计量值"
        },
        'document_headings': {
            'explanation': "解释说明",
            'interpretation': "结果解读",
            'relative_closeness_chart': "各方案相对贴近度柱状图",
            'ranking_chart': "方案排序条形图"
        },
        'chart_labels': {
            'x_alternative': "方案名称",
            'y_closeness': "相对贴近度",
            'relative_closeness_title': "各方案相对贴近度柱状图",
            'ranking_title': "方案排序结果"
        }
    },
    'en': {
        'title': "TOPSIS",
        'select_button_text': "Select File",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'analyze_button_text': "Analyze File",
        'no_file_selected': "Please select a valid file path.",
        'file_not_exists': "The file does not exist. Please select again.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'analysis_complete': "Analysis completed. The results have been saved to {}, and the relevant images have been saved.",
        'no_save_path_selected': "No save path selected. The results were not saved.",
        'switch_language_button_text': "Chinese/English",
        'open_excel_button_text': "Example data",
        'explanation': {
            'standardized_matrix': "The matrix after standardizing the original decision matrix",
            'weighted_matrix': "The standardized decision matrix considering the weights of each attribute",
            'positive_ideal': "The vector composed of the optimal values of each attribute",
            'negative_ideal': "The vector composed of the worst values of each attribute",
            'distance_positive': "The Euclidean distance between each alternative and the positive ideal solution",
            'distance_negative': "The Euclidean distance between each alternative and the negative ideal solution",
            'relative_closeness': "Reflects the relative closeness of each alternative to the positive ideal solution",
            'ranking_result': "The result of ranking each alternative according to the relative closeness"
        },
        'interpretation': {
            'standardized_matrix': "Eliminate the influence of different attribute dimensions",
            'weighted_matrix': "Reflect the importance of each attribute in the decision-making",
            'positive_ideal': "As the optimal reference point for measuring the advantages and disadvantages of each alternative",
            'negative_ideal': "As the worst reference point for measuring the advantages and disadvantages of each alternative",
            'distance_positive': "The smaller the distance, the better the alternative",
            'distance_negative': "The larger the distance, the better the alternative",
            'relative_closeness': "The closer the value is to 1, the better the alternative",
            'ranking_result': "The higher the ranking, the better the alternative"
        },
        'error_messages': {
            'negative_weights': "Weights contain negative values. Please check data. Weights should be non-negative.",
            'empty_data': "Excel file data is empty. Please check the file.",
            'insufficient_rows': "Insufficient data rows. At least 2 rows are required (1 weight row + 1 alternative row).",
            'non_numeric': "Data contains non-numeric content. Please check cell: {}",
            'weight_normalization': "The sum of weights is not 1, automatic normalization has been performed.",
            'index_type_mismatch': "The number of index types does not match the number of weights. Please check input.",
            'empty_factors': "The first column (factor column) contains empty values. Please check the data."
        },
        'result_headers': {
            'statistic': "Statistic",
            'value': "Value"
        },
        'document_headings': {
            'explanation': "Explanation",
            'interpretation': "Interpretation",
            'relative_closeness_chart': "Bar Chart of Relative Closeness of Each Alternative",
            'ranking_chart': "Alternative Ranking Bar Chart"
        },
        'chart_labels': {
            'x_alternative': "Alternative Name",
            'y_closeness': "Relative Closeness",
            'relative_closeness_title': "Bar Chart of Relative Closeness of Each Alternative",
            'ranking_title': "Alternative Ranking Results"
        }
    }
}


class TOPSISMethodAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为中文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data50.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def topsis_method(self, decision_matrix, weight_vector, index_types):
        # 标准化决策矩阵
        standardized_matrix = decision_matrix / np.sqrt(np.sum(decision_matrix **2, axis=0))

        # 加权标准化决策矩阵
        weighted_matrix = standardized_matrix * weight_vector

        # 正理想解和负理想解（根据指标类型区分）
        positive_ideal_solution = np.zeros(weighted_matrix.shape[1])
        negative_ideal_solution = np.zeros(weighted_matrix.shape[1])

        for i in range(weighted_matrix.shape[1]):
            if index_types[i] == 'benefit':
                positive_ideal_solution[i] = np.max(weighted_matrix[:, i])
                negative_ideal_solution[i] = np.min(weighted_matrix[:, i])
            else:  # cost type
                positive_ideal_solution[i] = np.min(weighted_matrix[:, i])
                negative_ideal_solution[i] = np.max(weighted_matrix[:, i])

        # 各方案到正理想解和负理想解的距离
        distances_to_positive = np.sqrt(np.sum((weighted_matrix - positive_ideal_solution)** 2, axis=1))
        distances_to_negative = np.sqrt(np.sum((weighted_matrix - negative_ideal_solution) **2, axis=1))

        # 各方案的相对贴近度
        relative_closeness = distances_to_negative / (distances_to_positive + distances_to_negative)

        # 方案排序结果
        ranking = np.argsort(-relative_closeness) + 1

        return standardized_matrix, weighted_matrix, positive_ideal_solution, negative_ideal_solution, \
            distances_to_positive, distances_to_negative, relative_closeness, ranking

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 打开 Excel 文件（第一列作为因素列保留）
            df = pd.read_excel(file_path, header=None)
            data = df.values

            # 数据校验
            if data.size == 0:
                messagebox.showerror("Error", languages[self.current_language]['error_messages']['empty_data'])
                return

            if data.shape[0] < 2:
                messagebox.showerror("Error", languages[self.current_language]['error_messages']['insufficient_rows'])
                return

            # 提取第一列作为因素名称（方案名称）
            factors = data[:, 0]  # 第一列所有行
            # 检查因素列是否有空值
            if any(pd.isna(factors)) or any(str(f).strip() == '' for f in factors):
                messagebox.showerror("Error", languages[self.current_language]['error_messages']['empty_factors'])
                return

            # 提取数值数据（从第二列开始）
            numeric_data = data[:, 1:]  # 所有行，第二列及以后

            # 检查非数值数据
            for i in range(numeric_data.shape[0]):
                for j in range(numeric_data.shape[1]):
                    if not isinstance(numeric_data[i, j], (int, float, np.number)):
                        try:
                            float(numeric_data[i, j])
                        except:
                            # 转换为Excel单元格格式（考虑第一列是非数据列，列索引+1）
                            cell_ref = f"{chr(65 + j + 1)}{i + 1}"
                            messagebox.showerror("Error", languages[self.current_language]['error_messages'][
                                'non_numeric'].format(cell_ref))
                            return

            # 将数值数据转换为浮点类型
            numeric_data = numeric_data.astype(float)

            # 获取指标类型（第二行为指标类型：1表示效益型，0表示成本型）
            # 第一行为权重，第三行及以后为决策矩阵
            index_types = ['benefit' if x == 1 else 'cost' for x in numeric_data[1]]
            weight_vector = numeric_data[0]
            decision_matrix = numeric_data[2:]  # 从第三行开始是方案数据

            # 权重校验与处理
            if np.any(weight_vector < 0):
                messagebox.showerror("Error", languages[self.current_language]['error_messages']['negative_weights'])
                return

            weight_sum = np.sum(weight_vector)
            if not np.isclose(weight_sum, 1.0):
                messagebox.showwarning("Warning",
                                       languages[self.current_language]['error_messages']['weight_normalization'])
                weight_vector = weight_vector / weight_sum  # 归一化处理

            # 检查指标类型与权重数量是否匹配
            if len(index_types) != len(weight_vector):
                messagebox.showerror("Error", languages[self.current_language]['error_messages']['index_type_mismatch'])
                return

            # 进行 TOPSIS 分析
            standardized_matrix, weighted_matrix, positive_ideal_solution, negative_ideal_solution, \
                distances_to_positive, distances_to_negative, relative_closeness, ranking = self.topsis_method(
                decision_matrix,
                weight_vector,
                index_types)

            # 整理数据（格式化矩阵显示）
            def format_matrix(matrix):
                if isinstance(matrix, list):
                    return [[round(num, 4) for num in row] for row in matrix]
                elif isinstance(matrix, np.ndarray):
                    if matrix.ndim == 1:
                        return [round(num, 4) for num in matrix.tolist()]
                    else:
                        return [[round(num, 4) for num in row] for row in matrix.tolist()]
                return matrix

            # 准备带因素名称的结果数据
            alternative_names = factors[2:]  # 第三行及以后的第一列是方案名称
            named_closeness = [f"{name}: {round(val, 4)}" for name, val in zip(alternative_names, relative_closeness)]
            named_ranking = [f"{name} (排名: {rank})" for name, rank in zip(alternative_names, ranking)]

            data = [
                [languages[self.current_language]['explanation']['standardized_matrix'], format_matrix(standardized_matrix)],
                [languages[self.current_language]['explanation']['weighted_matrix'], format_matrix(weighted_matrix)],
                [languages[self.current_language]['explanation']['positive_ideal'], format_matrix(positive_ideal_solution)],
                [languages[self.current_language]['explanation']['negative_ideal'], format_matrix(negative_ideal_solution)],
                [languages[self.current_language]['explanation']['distance_positive'],
                 [f"{name}: {round(val, 4)}" for name, val in zip(alternative_names, distances_to_positive)]],
                [languages[self.current_language]['explanation']['distance_negative'],
                 [f"{name}: {round(val, 4)}" for name, val in zip(alternative_names, distances_to_negative)]],
                [languages[self.current_language]['explanation']['relative_closeness'], named_closeness],
                [languages[self.current_language]['explanation']['ranking_result'], named_ranking]
            ]
            headers = [
                languages[self.current_language]['result_headers']['statistic'],
                languages[self.current_language]['result_headers']['value']
            ]
            df_result = pd.DataFrame(data, columns=headers)

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加结果表格
                table = doc.add_table(rows=df_result.shape[0] + 1, cols=df_result.shape[1])
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(headers):
                    hdr_cells[col_idx].text = header
                for row_idx, row in df_result.iterrows():
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx, value in enumerate(row):
                        # 格式化矩阵为更易读的形式
                        if isinstance(value, list):
                            if isinstance(value[0], list):  # 二维矩阵
                                cell_text = "\n".join([", ".join(map(str, sublist)) for sublist in value])
                            else:  # 一维向量
                                cell_text = "\n".join(map(str, value))  # 每个元素占一行
                        else:
                            cell_text = str(value)
                        row_cells[col_idx].text = cell_text

                # 添加解释说明（原代码是添加表格，替换为项目符号列表）
                doc.add_paragraph()
                doc.add_heading(languages[self.current_language]['document_headings']['explanation'], level=2)
                explanation_para = doc.add_paragraph()
                for key, value in languages[self.current_language]['explanation'].items():
                    # 添加项目符号
                    run = explanation_para.add_run(f"• {languages[self.current_language]['explanation'][key]}: {value}\n")
                    run.font.name = 'SimHei' if self.current_language == 'zh' else 'Arial'

                # 添加分析结果解读（同样替换为项目符号列表）
                doc.add_paragraph()
                doc.add_heading(languages[self.current_language]['document_headings']['interpretation'], level=2)
                interpretation_para = doc.add_paragraph()
                for key, value in languages[self.current_language]['interpretation'].items():
                    # 添加项目符号
                    run = interpretation_para.add_run(f"• {languages[self.current_language]['explanation'][key]}: {value}\n")
                    run.font.name = 'SimHei' if self.current_language == 'zh' else 'Arial'

                # 生成各方案相对贴近度柱状图（使用因素名称）
                fig, ax = plt.subplots(figsize=(10, 6))
                ax.bar(alternative_names, relative_closeness)
                ax.set_title(languages[self.current_language]['chart_labels']['relative_closeness_title'])
                ax.set_xlabel(languages[self.current_language]['chart_labels']['x_alternative'])
                ax.set_ylabel(languages[self.current_language]['chart_labels']['y_closeness'])
                plt.xticks(rotation=0, ha='right')  # 旋转标签避免重叠
                plt.tight_layout()  # 调整布局
                # 保存图片
                img_path1 = os.path.splitext(save_path)[0] + '_relative_closeness.png'
                plt.savefig(img_path1)
                plt.close()

                # 生成方案排序条形图（使用因素名称）
                fig, ax = plt.subplots(figsize=(10, 6))
                sorted_indices = np.argsort(-relative_closeness)
                ax.barh([alternative_names[i] for i in sorted_indices], relative_closeness[sorted_indices])
                ax.set_title(languages[self.current_language]['chart_labels']['ranking_title'])
                ax.set_xlabel(languages[self.current_language]['chart_labels']['y_closeness'])
                plt.tight_layout()
                img_path2 = os.path.splitext(save_path)[0] + '_ranking.png'
                plt.savefig(img_path2)
                plt.close()

                # 将图片插入到 Word 文档中
                doc.add_paragraph()
                doc.add_heading(languages[self.current_language]['document_headings']['relative_closeness_chart'], level=2)
                doc.add_picture(img_path1, width=Inches(6))

                doc.add_paragraph()
                doc.add_heading(languages[self.current_language]['document_headings']['ranking_chart'], level=2)
                doc.add_picture(img_path2, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                self.result_label.config(text=languages[self.current_language]["analysis_complete"].format(save_path),
                                         wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]["no_save_path_selected"])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]["analysis_error"].format(str(e)))
            print(str(e))  # 调试用

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])
        self.open_excel_label.config(text=languages[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=languages[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame,
                                               text=languages[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = TOPSISMethodAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()
import tkinter as tk
from tkinter import filedialog
import os
import numpy as np
import pandas as pd
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
import openpyxl
import statistics
from scipy import stats
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
languages = {
    "zh": {
        "title": "描述性统计",
        "select_button": "选择文件",
        "analyze_button": "分析文件",
        "file_not_found": "请选择文件。",
        "analysis_success": "分析完成，结果已保存到 {}\n",
        "no_save_path": "未选择保存路径，结果未保存。",
        "analysis_error": "分析文件时出错: {}",
        "images_saved": "图片已保存到 {}",
        "switch_language": "中/英",
        "x_label": "数值",
        "y_label_freq": "频率",
        "y_label_value": "数值",
        "x_label_index": "索引",
        "x_label_boxplot": "数据",
        "theoretical_quantiles": "理论分位数",
        "sample_quantiles": "样本分位数",
        'open_excel_button_text': "示例数据",
        "file_entry_placeholder": "请输入待分析 Excel 文件的完整路径",
        "explanation": {
            "Mean": "均值是数据集中所有数值的平均值，反映了数据的集中趋势。",
            "Median": "中位数是将数据集按升序排列后位于中间位置的数值，它不受极端值的影响，能更好地反映数据的中间水平。",
            "Standard Deviation": "标准差衡量了数据相对于均值的离散程度，标准差越大，数据越分散。",
            "Minimum": "最小值是数据集中的最小数值。",
            "Maximum": "最大值是数据集中的最大数值。",
            "Range": "极差是最大值与最小值的差值，反映了数据的取值范围。",
            "First Quartile (Q1)": "第一四分位数是将数据集按升序排列后位于 25% 位置的数值，它将数据集分为前 25% 和后 75%。",
            "Third Quartile (Q3)": "第三四分位数是将数据集按升序排列后位于 75% 位置的数值，它将数据集分为前 75% 和后 25%。",
            "Interquartile Range (IQR)": "四分位距是第三四分位数与第一四分位数的差值，反映了数据中间 50% 的分布范围，不受极端值的影响。",
            "Kurtosis": "峰度衡量了数据分布的峰态，反映了数据在均值附近的集中程度和尾部的厚度。",
            "Skewness": "偏度衡量了数据分布的不对称程度，正值表示右偏，负值表示左偏。",
            "Mode": "众数是数据集中出现次数最多的数值。"
        },
        "chart_titles": {
            "histogram": "直方图 - {}",
            "qq_plot": "正态Q-Q图 - {}",
            "box_plot": "箱线图 - {}",
            "scatter_plot": "散点图 - {}"
        },
        "chart_heading": "{}的图表",
        "stats_names": {
            "Mean": "均值",
            "Column Name": "列名",
            "Median": "中位数",
            "Standard Deviation": "标准差",
            "Minimum": "最小值",
            "Maximum": "最大值",
            "Range": "极差",
            "First Quartile (Q1)": "第一四分位数",
            "Third Quartile (Q3)": "第三四分位数",
            "Interquartile Range (IQR)": "四分位距",
            "Kurtosis": "峰度",
            "Skewness": "偏度",
            "Mode": "众数"
        }
    },
    "en": {
        "title": "Descriptive Statistics",
        "select_button": "Select File",
        "analyze_button": "Analyze File",
        "file_not_found": "The file does not exist. Please select again.",
        "analysis_success": "Analysis completed. The results have been saved to {}\n",
        "no_save_path": "No save path selected. The results were not saved.",
        "analysis_error": "An error occurred while analyzing the file: {}",
        "images_saved": "Images have been saved to {}",
        "switch_language": "Chinese/English",
        'open_excel_button_text': "Example data",
        "x_label": "Value",
        "y_label_freq": "Frequency",
        "y_label_value": "Value",
        "x_label_index": "Index",
        "x_label_boxplot": "Data",
        "theoretical_quantiles": "Theoretical quantiles",
        "sample_quantiles": "Sample quantiles",
        "file_entry_placeholder": "Please enter the full path of the Excel file to be analyzed",
        "explanation": {
            "Mean": "The mean is the average of all values in the dataset, reflecting the central tendency of the data.",
            "Median": "The median is the middle value when the dataset is arranged in ascending order. It is not affected by extreme values and better reflects the middle level of the data.",
            "Standard Deviation": "The standard deviation measures the dispersion of data relative to the mean. A larger standard deviation indicates more dispersed data.",
            "Minimum": "The minimum is the smallest value in the dataset.",
            "Maximum": "The maximum is the largest value in the dataset.",
            "Range": "The range is the difference between the maximum and minimum values, reflecting the range of the data.",
            "First Quartile (Q1)": "The first quartile is the value at the 25% position when the dataset is arranged in ascending order. It divides the dataset into the first 25% and the last 75%.",
            "Third Quartile (Q3)": "The third quartile is the value at the 75% position when the dataset is arranged in ascending order. It divides the dataset into the first 75% and the last 25%.",
            "Interquartile Range (IQR)": "The interquartile range is the difference between the third and first quartiles, reflecting the distribution range of the middle 50% of the data and is not affected by extreme values.",
            "Kurtosis": "Kurtosis measures the peakedness of the data distribution, reflecting the concentration of data around the mean and the thickness of the tails.",
            "Skewness": "Skewness measures the asymmetry of the data distribution. A positive value indicates right skewness, and a negative value indicates left skewness.",
            "Mode": "The mode is the value that appears most frequently in the dataset."
        },
        "chart_titles": {
            "histogram": "Histogram - {}",
            "qq_plot": "Normal Q-Q Plot - {}",
            "box_plot": "Box Plot - {}",
            "scatter_plot": "Scatter Plot - {}"
        },
        "chart_heading": "Charts of {}",
        "stats_names": {
            "Mean": "Mean",
            "Median": "Median",
            "Standard Deviation": "Standard Deviation",
            "Minimum": "Minimum",
            "Maximum": "Maximum",
            "Column Name": "Column Name",
            "Range": "Range",
            "First Quartile (Q1)": "First Quartile (Q1)",
            "Third Quartile (Q3)": "Third Quartile (Q3)",
            "Interquartile Range (IQR)": "Interquartile Range (IQR)",
            "Kurtosis": "Kurtosis",
            "Skewness": "Skewness",
            "Mode": "Mode"
        }
    }
}


class DescriptiveStatisticsApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data1.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]['file_not_found'])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件
            workbook = openpyxl.load_workbook(file_path)
            sheet = workbook.active

            rows = sheet.max_row
            columns = sheet.max_column

            data = []
            columns_stats = [languages[self.current_language]["stats_names"][key] for key in
                             ["Mean", "Median", "Standard Deviation", "Minimum", "Maximum", "Range",
                              "First Quartile (Q1)", "Third Quartile (Q3)", "Interquartile Range (IQR)",
                              "Kurtosis", "Skewness", "Mode"]]
            explanations = languages[self.current_language]['explanation']

            for col_idx in range(1, columns + 1):
                column_values = []
                for row_idx in range(2, rows + 1):
                    cell_value = sheet.cell(row=row_idx, column=col_idx).value
                    if isinstance(cell_value, (int, float)):
                        column_values.append(cell_value)

                if column_values:
                    col_name = sheet.cell(row=1, column=col_idx).value
                    # 列名清洗：移除特殊字符、空格标准化、首字母大写
                    if col_name:
                        col_name = str(col_name)
                        # 1. 移除首尾空格
                        col_name = col_name.strip()
                        # 2. 中文标点转英文
                        chinese_punct_map = {'，': ',', '。': '.', '：': ':', '；': ';', '（': '(', '）': ')', '【': '[',
                                             '】': ']', '　': ' '}
                        for c, en_c in chinese_punct_map.items():
                            col_name = col_name.replace(c, en_c)
                        # 3. 移除特殊字符（保留字母、数字、空格、下划线）
                        col_name = ''.join(e for e in col_name if e.isalnum() or e in [' ', '_'])
                        # 4. 多个空格替换为单个下划线
                        col_name = '_'.join(col_name.split())
                        # 5. 处理重复下划线
                        while '__' in col_name:
                            col_name = col_name.replace('__', '_')
                        # 6. 移除首尾下划线
                        col_name = col_name.strip('_')
                        # 7. 统一大小写（示例：全小写）
                        col_name = col_name.lower()
                        # 8. 替换关键字冲突
                        keyword_map = {'class': 'cls', 'def': 'definition'}
                        if col_name in keyword_map:
                            col_name = keyword_map[col_name]
                    mean_val = statistics.mean(column_values)
                    median_val = statistics.median(column_values)
                    try:
                        stdev_val = statistics.stdev(column_values)
                    except statistics.StatisticsError:
                        stdev_val = 0
                    min_val = min(column_values)
                    max_val = max(column_values)
                    range_val = max_val - min_val
                    q1 = np.percentile(column_values, 25)
                    q3 = np.percentile(column_values, 75)
                    iqr = q3 - q1
                    kurtosis_val = stats.kurtosis(column_values, fisher=False)
                    skewness_val = stats.skew(column_values)
                    from scipy.stats import mode

                    mode_result = mode(column_values, keepdims=True)
                    mode_val = mode_result.mode[0] if len(mode_result.mode) else None

                    values = [mean_val, median_val, stdev_val, min_val, max_val, range_val,
                              q1, q3, iqr, kurtosis_val, skewness_val, mode_val]
                    data.append([col_name] + values)

            headers = [languages[self.current_language]["stats_names"].get("Column Name", "Column Name")] + columns_stats
            df = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            # 合并数据和解释说明
            combined_df = df.copy()

            # 转置数据框
            column_name_key = languages[self.current_language]["stats_names"].get("Column Name", "Column Name")
            transposed_df = combined_df.set_index(column_name_key).T.reset_index().rename(
                columns={'index': column_name_key})

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建一个新的 Word 文档
                doc = Document()

                # 添加表格数据到 Word 文档
                table = doc.add_table(rows=transposed_df.shape[0] + 1, cols=transposed_df.shape[1])
                hdr_cells = table.rows[0].cells
                for col_idx, col_name in enumerate(transposed_df.columns):
                    hdr_cells[col_idx].text = col_name

                for row_idx in range(transposed_df.shape[0]):
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx, value in enumerate(transposed_df.iloc[row_idx]):
                        row_cells[col_idx].text = str(value)

                # 获取保存路径的目录
                save_dir = os.path.dirname(save_path)

                for col_idx in range(1, columns + 1):
                    column_values = []
                    for row_idx in range(2, rows + 1):
                        cell_value = sheet.cell(row=row_idx, column=col_idx).value
                        if isinstance(cell_value, (int, float)):
                            column_values.append(cell_value)

                    if column_values:
                        col_name = sheet.cell(row=1, column=col_idx).value

                        # 生成图表
                        fig, axes = plt.subplots(2, 2, figsize=(10, 8))

                        # 直方图
                        axes[0, 0].hist(column_values, bins=20, edgecolor='k')
                        axes[0, 0].set_title(
                            languages[self.current_language]["chart_titles"]["histogram"].format(col_name))
                        axes[0, 0].set_xlabel(languages[self.current_language]["x_label"] if "x_label" in languages[
                            self.current_language] else "Value")
                        axes[0, 0].set_ylabel(
                            languages[self.current_language]["y_label_freq"] if "y_label_freq" in languages[
                                self.current_language] else "Frequency")

                        # 正态 Q-Q 图
                        stats.probplot(column_values, dist="norm", plot=axes[0, 1])
                        axes[0, 1].set_title(languages[self.current_language]["chart_titles"]["qq_plot"].format(col_name))
                        axes[0, 1].set_xlabel(languages[self.current_language]["theoretical_quantiles"])
                        axes[0, 1].set_ylabel(languages[self.current_language]["sample_quantiles"])

                        # 箱线图
                        axes[1, 0].boxplot(column_values)
                        axes[1, 0].set_title(
                            languages[self.current_language]["chart_titles"]["box_plot"].format(col_name))
                        axes[1, 0].set_ylabel(
                            languages[self.current_language]["y_label_value"] if "y_label_value" in languages[
                                self.current_language] else "Value")
                        axes[1, 0].set_xlabel(
                            languages[self.current_language]["x_label_boxplot"] if "x_label_boxplot" in languages[
                                self.current_language] else "Data"
                        )

                        # 散点图（这里简单用索引作为 x 轴）
                        x = np.arange(len(column_values))
                        axes[1, 1].scatter(x, column_values)
                        axes[1, 1].set_title(
                            languages[self.current_language]["chart_titles"]["scatter_plot"].format(col_name))
                        axes[1, 1].set_xlabel(
                            languages[self.current_language]["x_label_index"] if "x_label_index" in languages[
                                self.current_language] else "Index")
                        axes[1, 1].set_ylabel(
                            languages[self.current_language]["y_label_value"] if "y_label_value" in languages[
                                self.current_language] else "Value")

                        plt.tight_layout()

                        # 生成图片保存路径
                        img_name = f"{col_name}_charts.png"
                        img_path = os.path.join(save_dir, img_name)

                        # 保存图片
                        plt.savefig(img_path)
                        plt.close()

                        # 将图片插入到 Word 文档中
                        doc.add_heading(languages[self.current_language]["chart_heading"].format(col_name), level=2)
                        doc.add_picture(img_path, width=Inches(6))

                # 添加统计指标说明
                doc.add_page_break()  # 新起一页
                doc.add_heading(
                    '统计指标说明' if self.current_language == 'zh' else 'Statistical Indicators Explanation', level=1)

                for metric_en, explanation in languages[self.current_language]["explanation"].items():
                    metric_name = languages[self.current_language]["stats_names"][metric_en]
                    doc.add_paragraph(f"{metric_name}: {explanation}")

                # 保存 Word 文档
                doc.save(save_path)

                result_msg = languages[self.current_language]['analysis_success'].format(save_path)
                result_msg += languages[self.current_language]['images_saved'].format(save_dir)
                self.result_label.config(text=result_msg, wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'
        self.root.title(languages[self.current_language]['title'])
        self.select_button.config(text=languages[self.current_language]['select_button'])
        self.analyze_button.config(text=languages[self.current_language]['analyze_button'])
        self.switch_language_label.config(text=languages[self.current_language]['switch_language'])
        self.open_excel_label.config(text=languages[self.current_language]["open_excel_button_text"])
        # 切换语言时更新提示信息
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]['file_entry_placeholder'])
        self.file_entry.config(foreground='gray')

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=languages[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=languages[self.current_language]["switch_language"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = DescriptiveStatisticsApp()
    app.run()


if __name__ == "__main__":
    run_app()
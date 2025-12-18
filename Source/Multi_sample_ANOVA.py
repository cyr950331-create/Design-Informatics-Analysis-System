import tkinter as tk
from tkinter import filedialog
import os
import numpy as np
import pandas as pd
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
from scipy import stats
from docx import Document
from docx.shared import Inches

# 设置支持中文的字体
plt.rcParams["font.family"] = ["SimHei", "WenQuanYi Micro Hei", "Heiti TC", "Arial Unicode MS"]
plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

# 定义语言字典，包含所有需要翻译的文本
languages = {
    'zh': {
        'title': "多样本方差",
        'select_button_text': "选择文件",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'analyze_button_text': "分析文件",
        'no_file_selected': "请选择文件。",
        'file_not_exists': "文件不存在，请重新选择。",
        'analysis_error': "分析文件时出错: {}",
        'analysis_complete': "分析完成，结果已保存到 {}，相关图片已保存。",
        'no_save_path_selected': "未选择保存路径，结果未保存。",
        'open_excel_button_text': "示例数据",
        'switch_language_button_text': "中/英",
        'stats_types': {
            "anova": "方差分析",
            "sample_size": "样本量",
            "mean": "均值"
        },
        'headers': ["统计量类型", "F统计量", "组间自由度", "组内自由度", "p值", "效应量（Eta平方）"],
        'explanation_heading': "解释说明",
        'interpretation_heading': "结果解读",
        'boxplot_title': "箱线图",
        'barplot_title': "柱状图",
        'meanplot_title': "均值图",
        'x_label_groups': "组",
        'y_label_fvalues': "F统计量",
        'y_label_mean': "均值",
        'explanation': {
            "anova": "用于比较三个或更多独立样本的均值是否有显著差异。",
            "sample_size": "每个样本中的观测值数量。",
            "mean": "样本数据的平均值。",
            "f_statistic": "衡量组间差异与组内差异的比值。",
            "df": "用于计算F分布的参数。",
            "p_value": "p值小于显著性水平（通常为0.05）时，拒绝原假设，认为样本之间存在显著差异；否则，接受原假设，认为样本之间无显著差异。",
            "effect_size": "反映自变量对因变量的影响程度。"
        },
        'interpretation': {
            "f_statistic": "F统计量越大，说明组间差异越显著。",
            "p_value": "用于判断样本之间是否存在显著差异。",
            "df": "影响F分布的形状，进而影响p值的计算。",
            "sample_size": "样本量的大小会影响统计检验的功效，较大的样本量通常能提供更准确的结果。",
            "mean": "反映样本数据的平均水平。",
            "effect_size": "效应量越大，说明自变量对因变量的影响越大。"
        }
    },
    'en': {
        'title': "Multi Sample ANOVA",
        'select_button_text': "Select File",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'analyze_button_text': "Analyze File",
        'no_file_selected': "Please select a valid file path.",
        'file_not_exists': "The file does not exist. Please select again.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'analysis_complete': "Analysis completed. The results have been saved to {}, and the relevant images have been saved.",
        'no_save_path_selected': "No save path selected. The results were not saved.",
        'open_excel_button_text': "Example data",
        'switch_language_button_text': "Chinese/English",
        'stats_types': {
            "anova": "ANOVA",
            "sample_size": "Sample Size",
            "mean": "Mean"
        },
        'headers': ["Statistic Type", "F-statistic", "Between Groups DF", "Within Groups DF", "p-value", "Effect Size (Eta Squared)"],
        'explanation_heading': "Explanation",
        'interpretation_heading': "Interpretation",
        'boxplot_title': "Box Plot",
        'barplot_title': "Bar Chart",
        'meanplot_title': "Mean Plot",
        'x_label_groups': "Groups",
        'y_label_fvalues': "F Values",
        'y_label_mean': "Mean",
        'explanation': {
            "anova": "Used to compare whether the means of three or more independent samples are significantly different.",
            "sample_size": "The number of observations in each sample.",
            "mean": "The average value of the sample data.",
            "f_statistic": "Measures the ratio of between-group variance to within-group variance.",
            "df": "Parameters used to calculate the F-distribution.",
            "p_value": "When the p-value is less than the significance level (usually 0.05), the null hypothesis is rejected, indicating a significant difference between samples; otherwise, the null hypothesis is accepted, indicating no significant difference.",
            "effect_size": "Reflects the influence of the independent variable on the dependent variable."
        },
        'interpretation': {
            "f_statistic": "The larger the F-statistic, the more significant the between-group difference.",
            "p_value": "Used to determine whether there is a significant difference between samples.",
            "df": "Affects the shape of the F-distribution, which in turn affects the calculation of the p-value.",
            "sample_size": "The sample size affects the power of the statistical test. A larger sample size usually provides more accurate results.",
            "mean": "Reflects the average level of the sample data.",
            "effect_size": "The larger the effect size, the greater the influence of the independent variable on the dependent variable."
        }
    }
}

class MultiSampleANOVAApp:
    def __init__(self, root=None):
        # 当前语言，默认为中文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data23.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"{languages[self.current_language]['file_not_exists'].split('，')[0]}：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"{languages[self.current_language]['analysis_error'].format('')}{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 检查数据是否为数值类型
            numerical_df = df.select_dtypes(include=[np.number])
            if numerical_df.empty:
                raise ValueError(f"{languages[self.current_language]['analysis_error'].format('')}数据中没有数值列，无法进行方差分析。" if self.current_language == 'zh' else
                                f"{languages[self.current_language]['analysis_error'].format('')}There are no numerical columns in the data for ANOVA analysis.")

            # 进行方差分析
            f_stat, p_value = stats.f_oneway(*numerical_df.T.values)

            # 计算自由度
            df_between = len(numerical_df.columns) - 1
            df_within = numerical_df.size - len(numerical_df.columns)

            # 计算效应量（Eta平方）
            sst = ((numerical_df - numerical_df.stack().mean()) ** 2).sum().sum()  # 使用stack()确保整体均值计算正确
            ssb = sum([len(numerical_df[col]) * (numerical_df[col].mean() - numerical_df.stack().mean()) ** 2 for col in
                       numerical_df.columns])
            eta_squared = ssb / sst

            # 计算样本量和均值
            sample_sizes = numerical_df.count()
            means = numerical_df.mean()

            # 整理数据
            p_value_str = f"{p_value:.6f}"
            if p_value < 0.001:
                p_value_str += "***"
            elif p_value < 0.01:
                p_value_str += "**"
            elif p_value < 0.05:
                p_value_str += "*"

            anova_row = [
                languages[self.current_language]['stats_types']["anova"],
                f"{f_stat:.4f}",
                f"{df_between}",
                f"{df_within}",
                p_value_str,
                f"{eta_squared:.4f}"
            ]

            # 样本量行（合并显示）
            sample_size_str = ", ".join([f"{col}: {cnt}" for col, cnt in sample_sizes.items()])
            sample_row = [languages[self.current_language]['stats_types']["sample_size"], sample_size_str, "", "", "", ""]

            # 均值行（合并显示）
            mean_str = ", ".join([f"{col}: {mean:.4f}" for col, mean in means.items()])
            mean_row = [languages[self.current_language]['stats_types']["mean"], mean_str, "", "", "", ""]

            # 构建结果数据
            headers = languages[self.current_language]['headers']

            # 组合所有行数据
            data = [anova_row, sample_row, mean_row]

            # 添加解释说明
            explanations = languages[self.current_language]['explanation']
            interpretation = languages[self.current_language]['interpretation']

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加分析结果表格
                table = doc.add_table(rows=len(data) + 1, cols=len(headers))  # 动态计算行列数
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(headers):
                    hdr_cells[col_idx].text = header  # 表头严格对应

                for row_idx in range(len(data)):
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx in range(len(data[row_idx])):
                        row_cells[col_idx].text = data[row_idx][col_idx]  # 数据严格对应

                # 添加解释说明（项目符号列表）
                doc.add_heading(languages[self.current_language]['explanation_heading'], level=2)
                explain_paragraph = doc.add_paragraph()
                explain_paragraph.add_run(
                    f"• {languages[self.current_language]['stats_types']['anova']}: {explanations['anova']}\n")  # 此处修改
                explain_paragraph.add_run(
                    f"• {languages[self.current_language]['stats_types']['sample_size']}: {explanations['sample_size']}\n")  # 此处修改
                explain_paragraph.add_run(
                    f"• {languages[self.current_language]['stats_types']['mean']}: {explanations['mean']}\n")  # 此处修改
                explain_paragraph.add_run(f"• {headers[1]}: {explanations['f_statistic']}\n")  # 此处修改
                explain_paragraph.add_run(f"• {headers[2]}/{headers[3]}: {explanations['df']}\n")  # 此处修改
                explain_paragraph.add_run(f"• {headers[4]}: {explanations['p_value']}\n")  # 此处修改
                explain_paragraph.add_run(f"• {headers[5]}: {explanations['effect_size']}\n")  # 此处修改

                # 添加分析结果解读（项目符号列表）
                doc.add_heading(languages[self.current_language]['interpretation_heading'], level=2)
                interpret_paragraph = doc.add_paragraph()
                interpret_paragraph.add_run(f"• {headers[1]}: {interpretation['f_statistic']}\n")
                interpret_paragraph.add_run(f"• {headers[4]}: {interpretation['p_value']}\n")
                interpret_paragraph.add_run(f"• {headers[2]}/{headers[3]}: {interpretation['df']}\n")
                interpret_paragraph.add_run(f"• {languages[self.current_language]['stats_types']['sample_size']}: {interpretation['sample_size']}\n")
                interpret_paragraph.add_run(f"• {languages[self.current_language]['stats_types']['mean']}: {interpretation['mean']}\n")
                interpret_paragraph.add_run(f"• {headers[5]}: {interpretation['effect_size']}\n")

                # 绘制箱线图
                plt.figure(figsize=(10, 6))
                numerical_df.boxplot()
                plt.title(languages[self.current_language]['boxplot_title'])
                plt.xlabel(languages[self.current_language]['x_label_groups'])
                plt.ylabel(languages[self.current_language]['y_label_fvalues'])
                boxplot_path = save_path.replace('.docx', '_boxplot.png')
                plt.savefig(boxplot_path)
                plt.close()

                # 绘制柱状图
                plt.figure(figsize=(10, 6))
                bars = plt.bar(numerical_df.columns, means)
                for bar in bars:
                    height = bar.get_height()
                    plt.annotate(f'{height:.2f}',
                                 xy=(bar.get_x() + bar.get_width() / 2, height),
                                 xytext=(0, 3),  # 3 points vertical offset
                                 textcoords="offset points",
                                 ha='center', va='bottom')
                plt.title(languages[self.current_language]['barplot_title'])
                plt.xlabel(languages[self.current_language]['x_label_groups'])
                plt.ylabel(languages[self.current_language]['y_label_mean'])
                barplot_path = save_path.replace('.docx', '_barplot.png')
                plt.savefig(barplot_path)
                plt.close()

                # 绘制均值图
                plt.figure(figsize=(10, 6))
                plt.plot(numerical_df.columns, means, marker='o')
                plt.title(languages[self.current_language]['meanplot_title'])
                plt.xlabel(languages[self.current_language]['x_label_groups'])
                plt.ylabel(languages[self.current_language]['y_label_mean'])
                meanplot_path = save_path.replace('.docx', '_meanplot.png')
                plt.savefig(meanplot_path)
                plt.close()

                # 在 Word 文档中添加图片
                doc.add_heading(languages[self.current_language]['boxplot_title'], level=2)
                doc.add_picture(boxplot_path, width=Inches(6))
                doc.add_heading(languages[self.current_language]['barplot_title'], level=2)
                doc.add_picture(barplot_path, width=Inches(6))
                doc.add_heading(languages[self.current_language]['meanplot_title'], level=2)
                doc.add_picture(meanplot_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                self.result_label.config(text=languages[self.current_language]['analysis_complete'].format(save_path), wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]['no_save_path_selected'])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        current_entry_text = self.file_entry.get()
        # 只有当输入框显示的是占位文字时才更新
        if current_entry_text == languages['zh' if self.current_language == 'en' else 'en']["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])
        self.open_excel_label.config(text=languages[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=languages[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=languages[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = MultiSampleANOVAApp()
    app.run()

if __name__ == "__main__":
    run_app()
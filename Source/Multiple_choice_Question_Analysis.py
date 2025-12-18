import tkinter as tk
from tkinter import filedialog
import os
import numpy as np
import pandas as pd
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'  # 设置字体为黑体，可根据系统情况修改为其他支持中文的字体
plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

# 定义语言字典
languages = {
    'zh': {
        'title': "问卷多选题",
        'select_button_text': "选择文件",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'analyze_button_text': "分析文件",
        'no_file_selected': "请选择文件。",
        'file_not_exists': "文件不存在，请重新选择。",
        'analysis_error': "分析文件时出错: {}",
        'analysis_complete': "分析完成，结果已保存到 {}，相关图片已保存。",
        'no_save_path_selected': "未选择保存路径，结果未保存。",
        'open_excel_button_text': "示例数据",
        'switch_language_button_text': "中/英",
        'explanation': {
            "选项选择次数": "每个选项被选择的总次数。",
            "样本量": "参与问卷的总人数。"
        },
        'interpretation': {
            "选项选择次数": "选择次数越多，说明该选项在问卷中越受欢迎。",
            "样本量": "样本量的大小会影响统计结果的可靠性，较大的样本量通常能提供更可靠的结果。"
        }
    },
    'en': {
        'title': "Multiple Choice Question of Questionnaire",
        'select_button_text': "Select File",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'analyze_button_text': "Analyze File",
        'no_file_selected': "Please select a valid file path.",
        'file_not_exists': "The file does not exist. Please select again.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'analysis_complete': "Analysis completed. The results have been saved to {}, and the relevant images have been saved.",
        'no_save_path_selected': "No save path selected. The results were not saved.",
        'open_excel_button_text': "Example data",
        'switch_language_button_text': "Chinese/English",
        'explanation': {
            "Number of Choices per Option": "The total number of times each option was selected.",
            "Sample Size": "The total number of people participating in the questionnaire."
        },
        'interpretation': {
            "Number of Choices per Option": "The more times an option is selected, the more popular it is in the questionnaire.",
            "Sample Size": "The sample size affects the reliability of the statistical results. A larger sample size usually provides more reliable results."
        }
    }
}

class MultipleChoiceQuestionAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data16.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def multiple_choice_analysis(self, data):
        # 假设多选题答案以逗号分隔存储在单元格中
        all_choices = {}
        for column in data.columns:
            for cell in data[column].dropna():
                choices = str(cell).split(',')
                for choice in choices:
                    choice = choice.strip()
                    if choice in all_choices:
                        all_choices[choice] += 1
                    else:
                        all_choices[choice] = 1

        sample_size = len(data)
        return all_choices, sample_size

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 进行多选题分析
            choice_counts, sample_size = self.multiple_choice_analysis(df)

            # 根据当前语言定义统计量名称和表头
            if self.current_language == "zh":
                stats_names = {
                    "choice_counts": "选项选择次数",
                    "sample_size": "样本量"
                }
                headers = ["统计量", "统计量值"]
            else:
                stats_names = {
                    "choice_counts": "Number of Choices per Option",
                    "sample_size": "Sample Size"
                }
                headers = ["Statistic", "Value"]

            # 整理多语言数据
            data = [
                [stats_names["choice_counts"], pd.Series(choice_counts).to_dict()],
                [stats_names["sample_size"], sample_size]
            ]
            df_result = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            explanations = languages[self.current_language]['explanation']
            interpretations = languages[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            # 根据当前语言设置列名
            if self.current_language == "zh":
                columns = ["选项选择次数", "样本量"]
            else:
                columns = ["Number of Choices per Option", "Sample Size"]

            explanation_df = explanation_df.reindex(columns=columns)
            explanation_df.insert(0, "统计量_解释说明", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            if self.current_language == "zh":
                columns = ["选项选择次数", "样本量"]
            else:
                columns = ["Number of Choices per Option", "Sample Size"]

            interpretation_df = interpretation_df.reindex(columns=columns)
            interpretation_df.insert(0, "统计量_结果解读", "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加表格
                if self.current_language == "zh":
                    doc.add_heading('分析结果', level=2)
                else:
                    doc.add_heading('Analysis Results', level=2)
                table = doc.add_table(rows=len(df_result) + 1, cols=len(df_result.columns))
                hdr_cells = table.rows[0].cells
                for col_idx, col_name in enumerate(df_result.columns):
                    hdr_cells[col_idx].text = col_name
                for row_idx, row in df_result.iterrows():
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx].text = str(value)

                # 添加解释说明（项目符号列表）
                if self.current_language == "zh":
                    doc.add_heading('解释说明', level=2)
                else:
                    doc.add_heading('Explanations', level=2)
                for key, value in explanations.items():
                    p = doc.add_paragraph(style='List Bullet')  # 使用项目符号样式
                    p.add_run(f"{key}: {value}")

                # 添加结果解读（项目符号列表）
                if self.current_language == "zh":
                    doc.add_heading('结果解读', level=2)
                else:
                    doc.add_heading('Result Interpretation', level=2)
                for key, value in interpretations.items():
                    p = doc.add_paragraph(style='List Bullet')  # 使用项目符号样式
                    p.add_run(f"{key}: {value}")

                # 生成图片（选项选择次数柱状图）
                fig, ax = plt.subplots()
                pd.Series(choice_counts).plot(kind='bar', ax=ax)
                ax.set_title('选项选择次数柱状图' if self.current_language == 'zh' else 'Bar Chart of Option Selection Counts')
                ax.set_xlabel('选项' if self.current_language == 'zh' else 'Options')
                ax.set_ylabel('选择次数' if self.current_language == 'zh' else 'Selection Counts')
                ax.tick_params(axis='x', rotation=0)
                # 保存图片
                img_path = os.path.splitext(save_path)[0] + '.png'
                plt.savefig(img_path)
                plt.close()

                # 在 Word 文档中插入图片
                doc.add_picture(img_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                self.result_label.config(text=languages[self.current_language]["analysis_complete"].format(save_path), wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]["no_save_path_selected"])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]["analysis_error"].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])
        self.open_excel_label.config(text=languages[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=languages[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=languages[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = MultipleChoiceQuestionAnalysisApp()
    app.run()

if __name__ == "__main__":
    run_app()
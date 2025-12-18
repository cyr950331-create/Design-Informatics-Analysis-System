import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
from scipy import stats
import matplotlib.pyplot as plt
import pathlib
import matplotlib
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
matplotlib.rcParams['font.family'] = 'SimHei'
matplotlib.rcParams['axes.unicode_minus'] = False

# 定义语言字典
languages = {
    'zh': {
        'title': "单样本方差",
        'select_button_text': "选择文件",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'analyze_button_text': "分析文件",
        'no_file_selected': "请选择文件。",
        'file_not_exists': "文件不存在，请重新选择。",
        'analysis_error': "分析文件时出错: {}",
        'analysis_complete': "分析完成，结果已保存到 {}，相关图片已保存。",
        'no_save_path_selected': "未选择保存路径，结果未保存。",
        'switch_language_button_text': "中/英",
        'open_excel_button_text': "示例数据",
        'preset_value_label': "预设总体方差:",
        'preset_value_placeholder': "请输入预设的总体方差（默认为1，必须为正数）",
        'invalid_preset_value': "预设方差必须是正数",
        'statistic_explanation_title': "统计量解释说明",
        'statistic_interpretation_title': "统计量结果解读",
        'explanation': {
            "单样本方差分析": "用于检验一个样本的方差是否与某个已知的总体方差存在显著差异。",
            "样本量": "样本中的观测值数量。",
            "样本方差": "样本数据的方差。",
            "样本标准差": "样本数据的标准差（方差的平方根）。",
            "卡方统计量": "衡量样本方差与总体方差之间差异的统计量。",
            "自由度": "用于计算t分布的参数。",
            "p值": "p值小于显著性水平（通常为0.05）时，拒绝原假设，认为样本均值与总体均值存在显著差异；否则，接受原假设，认为样本均值与总体均值无显著差异。",
            "效应量": "反映样本方差与总体方差之间差异的程度（方差比）。",
            "预设值": "作为参考标准的总体方差。",
            "与预设值的差异": "样本均值与预设总体均值之间的差值。"
        },
        'interpretation': {
            "卡方统计量": "卡方统计量的值越偏离自由度，说明样本方差与总体方差之间的差异越显著。",
            "p值": "用于判断样本均值与总体均值之间是否存在显著差异。",
            "自由度": "影响t分布的形状，进而影响p值的计算。",
            "样本量": "样本量的大小会影响统计检验的功效，较大的样本量通常能提供更准确的结果。",
            "样本方差": "反映样本数据的离散程度。",
            "样本标准差": "反映样本数据的离散程度（与原始数据单位一致）。",
            "效应量": "效应量越偏离1，说明样本方差与总体方差之间的差异越大。",
            "预设值": "用于与样本方差进行比较的参考值。",
            "与预设值的差异": "直接反映样本均值与预设值之间的偏离程度，正值表示样本均值高于预设值，负值表示低于预设值。"
        }
    },
    'en': {
        'title': "One Sample ANOVA",
        'select_button_text': "Select File",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'analyze_button_text': "Analyze File",
        'no_file_selected': "Please select a valid file path.",
        'file_not_exists': "The file does not exist. Please select again.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'analysis_complete': "Analysis completed. The results have been saved to {}, and the relevant images have been saved.",
        'no_save_path_selected': "No save path selected. The results were not saved.",
        'switch_language_button_text': "Chinese/English",
        'open_excel_button_text': "Example data",
        'preset_value_label': "Preset Population Variance:",
        'preset_value_placeholder': "Please enter preset population variance (default is 1, must be positive)",
        'invalid_preset_value': "Preset variance must be a positive number",
        'statistic_explanation_title': "Statistic Explanation",
        'statistic_interpretation_title': "Statistic Interpretation",
        'explanation': {
            "One-sample Variance Test": "Used to test whether the variance of a sample is significantly different from a known population variance.",
            "Sample Size": "The number of observations in the sample.",
            "Sample Variance": "The variance of the sample data.",
            "Sample Standard Deviation": "The standard deviation of the sample data (square root of variance).",
            "Chi-square Statistic": "A statistic that measures the difference between the sample variance and the population variance.",
            "Degrees of Freedom": "Parameters used to calculate the chi-square distribution.",
            "p-value": "When the p-value is less than the significance level (usually 0.05), the null hypothesis is rejected, indicating a significant difference between the sample variance and the population variance; otherwise, the null hypothesis is accepted, indicating no significant difference.",
            "Effect Size": "Reflects the degree of difference between the sample variance and the population variance (variance ratio).",
            "Preset Value": "The population variance used as a reference standard.",
            "Difference from Preset": "The difference between the sample variance and the preset population variance."
        },
        'interpretation': {
            "Chi-square Statistic": "The more the chi-square statistic deviates from the degrees of freedom, the more significant the difference between the sample variance and the population variance.",
            "p-value": "Used to determine whether there is a significant difference between the sample variance and the population variance.",
            "Degrees of Freedom": "Affects the shape of the chi-square distribution, which in turn affects the calculation of the p-value.",
            "Sample Size": "The sample size affects the power of the statistical test. A larger sample size usually provides more accurate results.",
            "Sample Variance": "Reflects the degree of dispersion of the sample data.",
            "Sample Standard Deviation": "Reflects the degree of dispersion of the sample data (in the same unit as the original data).",
            "Effect Size": "The more the effect size deviates from 1, the greater the difference between the sample variance and the population variance.",
            "Preset Value": "Reference value used for comparison with the sample variance.",
            "Difference from Preset": "Directly reflects the deviation between the sample variance and the preset value. A positive value indicates the sample variance is higher than the preset value, while a negative value indicates it is lower."
        }
    }
}


class OneSampleANOVAApp:
    def __init__(self, root=None):
        # 当前语言，默认为中文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data22.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        widget = event.widget
        if widget == self.file_entry and self.file_entry.get() == languages[self.current_language][
            "file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')
        elif widget == self.preset_entry and self.preset_entry.get() == languages[self.current_language][
            "preset_value_placeholder"]:
            self.preset_entry.delete(0, tk.END)
            self.preset_entry.config(foreground='black')

    def on_focusout(self, event):
        widget = event.widget
        if widget == self.file_entry and self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')
        elif widget == self.preset_entry and self.preset_entry.get() == "":
            self.preset_entry.insert(0, languages[self.current_language]["preset_value_placeholder"])
            self.preset_entry.config(foreground='gray')

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return

        # 获取并验证预设值
        try:
            preset_text = self.preset_entry.get()
            if preset_text == languages[self.current_language]["preset_value_placeholder"] or not preset_text.strip():
                population_variance = 1  # 默认为1（方差不能为0）
            else:
                population_variance = float(preset_text.strip())
                if population_variance <= 0:  # 方差必须为正数
                    raise ValueError(languages[self.current_language]["invalid_preset_value"])
        except ValueError:
            self.result_label.config(text=languages[self.current_language]["invalid_preset_value"])
            return

        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 检查数据是否为数值类型
            numerical_df = df.select_dtypes(include=[np.number])
            if numerical_df.empty:
                raise ValueError("数据中没有数值列，无法进行方差分析。")

            # 提取第一列数值作为样本
            sample = numerical_df.iloc[:, 0]  # 取第一列数值数据作为分析样本

            # 计算样本方差和标准差
            sample_variance = sample.var(ddof=1)  # 样本方差（自由度n-1）
            sample_std = sample.std(ddof=1)  # 样本标准差

            # 卡方检验（单样本方差检验）
            chi2_stat = (len(sample) - 1) * sample_variance / population_variance  # 卡方统计量
            df_value = len(sample) - 1  # 自由度
            # 双侧检验p值
            p_value = 2 * min(stats.chi2.cdf(chi2_stat, df_value), 1 - stats.chi2.cdf(chi2_stat, df_value))

            # 效应量（方差比）
            effect_size = sample_variance / population_variance

            # 其他统计量
            sample_size = len(sample)
            diff_from_preset = sample_variance - population_variance  # 与预设方差的差异

            # 整理数据
            headers = [
                "统计量" if self.current_language == 'zh' else "Statistic",
                "样本量" if self.current_language == 'zh' else "Sample Size",
                "样本方差" if self.current_language == 'zh' else "Sample Variance",
                "样本标准差" if self.current_language == 'zh' else "Sample Std Dev",
                "预设总体方差" if self.current_language == 'zh' else "Preset Variance",
                "与预设值的差异" if self.current_language == 'zh' else "Difference from Preset",
                "卡方统计量" if self.current_language == 'zh' else "Chi2-statistic",
                "自由度" if self.current_language == 'zh' else "Degrees of Freedom",
                "p值" if self.current_language == 'zh' else "p-value",
                "效应量（方差比）" if self.current_language == 'zh' else "Effect Size (Variance Ratio)"
            ]

            data = [
                "方差分析" if self.current_language == 'zh' else "ANOVA",
                f"{sample_size:.0f}",  # 整数显示
                f"{sample_variance:.4f}",  # 保留4位小数
                f"{sample_std:.4f}",
                f"{population_variance:.4f}",
                f"{diff_from_preset:.4f}",
                f"{chi2_stat:.4f}",
                f"{df_value:.0f}",
                f"{p_value:.6f}",  # p值保留6位小数
                f"{effect_size:.4f}"
            ]

            # 创建结果数据框
            df_result = pd.DataFrame([data], columns=headers)

            # 添加解释说明（列表形式）
            explanations = languages[self.current_language]['explanation']
            interpretation = languages[self.current_language]['interpretation']

            # 绘制箱线图
            plt.figure(figsize=(10, 6))
            numerical_df.iloc[:, 0].plot.box()
            plt.title('样本数据箱线图' if self.current_language == 'zh' else 'Box Plot of Sample Data')
            plt.ylabel('数值' if self.current_language == 'zh' else 'Values')
            box_plot_path = os.path.splitext(file_path)[0] + '_boxplot.png'
            plt.savefig(box_plot_path)
            plt.close()

            # 方差对比图
            plt.figure(figsize=(10, 6))
            bars = plt.bar(
                ['样本方差' if self.current_language == 'zh' else 'Sample Variance',
                 '预设方差' if self.current_language == 'zh' else 'Preset Variance'],
                [sample_variance, population_variance]
            )
            for bar in bars:
                height = bar.get_height()
                plt.annotate(f'{height:.4f}',
                             xy=(bar.get_x() + bar.get_width() / 2, height),
                             xytext=(0, 3),
                             textcoords="offset points",
                             ha='center', va='bottom')
            plt.title('方差比较图' if self.current_language == 'zh' else 'Variance Comparison')
            plt.ylabel('方差值' if self.current_language == 'zh' else 'Variance Value')
            variance_plot_path = os.path.splitext(file_path)[0] + '_varianceplot.png'
            plt.savefig(variance_plot_path)
            plt.close()

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建一个新的 Word 文档
                doc = Document()

                # 添加主标题
                doc.add_heading('单样本方差分析结果' if self.current_language == 'zh' else 'One-sample ANOVA Results',
                                0)

                # 添加分析结果表格部分
                doc.add_heading('分析结果' if self.current_language == 'zh' else 'Analysis Results', level=1)
                table = doc.add_table(rows=1, cols=len(df_result.columns))
                table.style = 'Table Grid'  # 添加网格线

                # 设置表头
                hdr_cells = table.rows[0].cells
                for i, col in enumerate(df_result.columns):
                    hdr_cells[i].text = col
                    # 表头单元格加粗
                    hdr_cells[i].paragraphs[0].runs[0].font.bold = True

                # 添加数据行
                row_cells = table.add_row().cells
                for i, value in enumerate(data):
                    row_cells[i].text = value

                # 调整列宽使表格更紧凑
                for col in table.columns:
                    # 根据内容自动调整列宽（设置最大宽度限制）
                    col.width = Inches(1.1)  # 统一列宽，可根据需要微调

                # 添加解释说明部分（列表形式）
                doc.add_heading(languages[self.current_language]['statistic_explanation_title'], level=1)
                for key, value in explanations.items():
                    para = doc.add_paragraph()
                    para.add_run(f"• {key}: ").bold = True
                    para.add_run(value)

                # 添加结果解读部分（列表形式）
                doc.add_heading(languages[self.current_language]['statistic_interpretation_title'], level=1)
                for key, value in interpretation.items():
                    para = doc.add_paragraph()
                    para.add_run(f"• {key}: ").bold = True
                    para.add_run(value)

                # 添加图片部分
                doc.add_heading('数据可视化' if self.current_language == 'zh' else 'Data Visualization', level=1)
                doc.add_picture(box_plot_path, width=Inches(6))
                doc.add_picture(variance_plot_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                # 设置结果提示
                self.result_label.config(
                    text=languages[self.current_language]["analysis_complete"].format(save_path),
                    wraplength=400
                )

                # 文本自动换行
                self.result_label.config(text=languages[self.current_language]["analysis_complete"].format(save_path),
                                         wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]["no_save_path_selected"])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]["analysis_error"].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])
        self.preset_label.config(text=languages[self.current_language]["preset_value_label"])
        self.preset_entry.delete(0, tk.END)
        self.preset_entry.insert(0, languages[self.current_language]["preset_value_placeholder"])
        self.preset_entry.config(foreground='gray')
        self.open_excel_label.config(text=languages[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建预设值标签和输入框
        self.preset_label = ttk.Label(frame, text=languages[self.current_language]["preset_value_label"])
        self.preset_label.pack(pady=5)

        self.preset_entry = ttk.Entry(frame, width=50)
        self.preset_entry.insert(0, languages[self.current_language]["preset_value_placeholder"])
        self.preset_entry.config(foreground='gray')
        self.preset_entry.bind('<FocusIn>', self.on_entry_click)
        self.preset_entry.bind('<FocusOut>', self.on_focusout)
        self.preset_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=languages[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame,
                                               text=languages[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = OneSampleANOVAApp()
    app.run()


if __name__ == "__main__":
    run_app()
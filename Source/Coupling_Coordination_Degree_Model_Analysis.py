import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
LANGUAGES = {
    'zh': {
        'title': "耦合协调度模型",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "请选择文件。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "中/英",
        'open_excel_button_text': "示例数据",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'explanation': {
            "耦合度": "反映多个系统之间相互作用的强度",
            "耦合协调度": "综合考虑系统发展水平和耦合程度，衡量系统之间的协调发展状况",
            "耦合度分布直方图": "展示耦合度值分布情况的直方图",
            "耦合协调度分布直方图": "展示耦合协调度值分布情况的直方图",
            "系统综合指数折线图": "展示各系统综合发展指数变化趋势的折线图"
        },
        'interpretation': {
            "耦合度": "值越接近 1，系统间相互作用越强",
            "耦合协调度": "值越接近 1，系统间协调发展程度越高",
            "耦合度分布直方图": "直观观察耦合度值的分布特征",
            "耦合协调度分布直方图": "直观观察耦合协调度值的分布特征",
            "系统综合指数折线图": "展示不同系统发展水平的变化趋势及差距"
        },
        'coordination_level': {
            '0-0.2': '极度失调',
            '0.2-0.4': '中度失调',
            '0.4-0.6': '基本协调',
            '0.6-0.8': '中度协调',
            '0.8-1.0': '极度协调'
        },
        'headers_title': '指标表头信息',
        'system_count': '系统数量: {}',
        'system_indices': '系统指标分布: {}',
        'explanation_title': '解释说明',
        'interpretation_title': '结果解读'
    },
    'en': {
        'title': "Coupling Coordination Degree Model",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Chinese/English",
        'open_excel_button_text': "Example data",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'explanation': {
            "Coupling Degree": "Reflects the intensity of interaction between multiple systems",
            "Coupling Coordination Degree": "Comprehensively considers the development level and coupling degree of systems to measure the coordinated development status between systems",
            "Coupling Degree Distribution Histogram": "A histogram showing the distribution of coupling degree values",
            "Coupling Coordination Degree Distribution Histogram": "A histogram showing the distribution of coupling coordination degree values",
            "System Comprehensive Index Line Chart": "Line chart showing the changing trends of comprehensive development indices of each system"
        },
        'interpretation': {
            "Coupling Degree": "The closer the value is to 1, the stronger the interaction between systems",
            "Coupling Coordination Degree": "The closer the value is to 1, the higher the coordinated development degree between systems",
            "Coupling Degree Distribution Histogram": "Visually observe the distribution characteristics of coupling degree values",
            "Coupling Coordination Degree Distribution Histogram": "Visually observe the distribution characteristics of coupling coordination degree values",
            "System Comprehensive Index Line Chart": "Show the changing trends and gaps of development levels of different systems"
        },
        'coordination_level': {
            '0-0.2': 'Extremely uncoordinated',
            '0.2-0.4': 'Moderately uncoordinated',
            '0.4-0.6': 'Basically coordinated',
            '0.6-0.8': 'Moderately coordinated',
            '0.8-1.0': 'Extremely coordinated'
        },
        'headers_title': 'Indicator Header Information',
        'system_count': 'Number of systems: {}',
        'system_indices': 'System index distribution: {}',
        'explanation_title': 'Explanation',
        'interpretation_title': 'Interpretation'
    }
}


class CouplingCoordinationDegreeModelAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data30.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def entropy_weight_method(self, data):
        """熵值法计算指标权重"""
        # 数据标准化
        data = (data - data.min(axis=0)) / (data.max(axis=0) - data.min(axis=0) + 1e-10)
        # 计算比重
        p = data / (data.sum(axis=0) + 1e-10)
        # 计算熵值
        k = 1 / np.log(len(data))
        e = -k * np.sum(p * np.log(p + 1e-10), axis=0)
        # 计算权重
        w = (1 - e) / np.sum(1 - e)
        return w

    def coupling_coordination_degree(self, data, system_ids):
        """
        计算耦合度和耦合协调度（多系统版）
        :param data: 原始数据矩阵，每行代表样本，每列代表一个指标
        :param system_ids: 各指标所属系统的标识列表（0,1,2,3...）
        :return: 各样本的耦合度列表, 各样本的耦合协调度列表, 各系统的综合指数（数组，每行是一个样本的各系统指数）
        """
        # 获取系统数量和每个系统包含的指标索引
        unique_systems = np.unique(system_ids)
        num_systems = len(unique_systems)
        system_indices = {sys_id: np.where(system_ids == sys_id)[0] for sys_id in unique_systems}

        m = data.shape[0]  # 样本数量

        # 计算每个系统的指标权重和综合指数
        system_weights = {}
        system_u = {}

        for sys_id, indices in system_indices.items():
            # 提取该系统的指标数据
            sys_data = data[:, indices]
            # 计算该系统的指标权重
            weights = self.entropy_weight_method(sys_data)
            system_weights[sys_id] = weights
            # 计算该系统的综合指数
            system_u[sys_id] = np.dot(sys_data, weights)

        # 整理各系统综合指数为矩阵（样本数×系统数）
        u_matrix = np.column_stack([system_u[sys_id] for sys_id in sorted(unique_systems)])

        # 计算耦合度和耦合协调度
        C_list = []
        D_list = []

        for i in range(m):
            # 提取当前样本各系统的综合指数
            u_i = u_matrix[i, :]

            # 计算耦合度（多系统公式）
            sum_u = np.sum(u_i)
            product_u = np.prod(u_i)
            if sum_u == 0:
                C = 0.0
            else:
                C = (product_u ** (1 / num_systems)) / (sum_u / num_systems)
            C_list.append(C)

            # 计算综合发展指数（各系统指数的平均值）
            T = np.mean(u_i)

            # 计算耦合协调度
            D = np.sqrt(C * T)
            D_list.append(D)

        return C_list, D_list, u_matrix, unique_systems, system_indices

    def get_coordination_level(self, d_value):
        """根据耦合协调度值判断协调等级"""
        levels = LANGUAGES[self.current_language]['coordination_level']
        if d_value < 0.2:
            return levels['0-0.2']
        elif d_value < 0.4:
            return levels['0.2-0.4']
        elif d_value < 0.6:
            return levels['0.4-0.6']
        elif d_value < 0.8:
            return levels['0.6-0.8']
        else:
            return levels['0.8-1.0']

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        """当用户点击输入框时，清除提示信息"""
        if self.file_entry.get() == LANGUAGES[self.current_language]['file_entry_placeholder']:
            self.file_entry.delete(0, tk.END)
            self.file_entry.configure(style="TEntry")  # 恢复默认样式

    def on_focusout(self, event):
        """当用户离开输入框时，如果没有输入内容，恢复提示信息"""
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
            self.file_entry.configure(style="Gray.TEntry")

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]['file_entry_placeholder']:
            file_path = ""
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 读取Excel文件，第一行为指标名，第二行为系统标识
            # 先读取所有行获取系统标识
            all_rows = pd.read_excel(file_path, header=None).values
            if len(all_rows) < 2:
                raise ValueError("Excel文件至少需要包含表头行和系统标识行")

            # 第一行为指标名，第二行为系统标识
            headers = all_rows[0].tolist()
            system_ids = all_rows[1].astype(int)  # 系统标识转为整数

            # 从第三行开始读取数据
            df = pd.read_excel(file_path, header=2)
            if df.empty:
                raise ValueError("没有有效的数据行，请检查文件内容")

            # 数据清洗：处理缺失值和异常值
            df = df.apply(pd.to_numeric, errors='coerce')  # 转换为数值，非数值转为NaN
            df = df.dropna()  # 删除含缺失值的行
            if df.empty:
                raise ValueError("数据清洗后为空，请检查输入数据")

            # 处理异常值（3σ原则）
            for col in df.columns:
                mean = df[col].mean()
                std = df[col].std()
                df = df[(df[col] >= mean - 3 * std) & (df[col] <= mean + 3 * std)]
            if df.empty:
                raise ValueError("异常值处理后数据为空，请检查输入数据")

            data = df.values.astype(float)

            # 确保指标数量与系统标识数量一致
            if data.shape[1] != len(system_ids):
                raise ValueError(f"数据列数({data.shape[1]})与系统标识数量({len(system_ids)})不匹配")

            # 进行耦合协调度分析
            C_list, D_list, u_matrix, unique_systems, system_indices = self.coupling_coordination_degree(data,
                                                                                                         system_ids)
            num_systems = len(unique_systems)

            # 整理系统分布信息
            system_info = {int(sys_id): len(indices) for sys_id, indices in system_indices.items()}

            # 整理结果数据
            data_rows = []
            for i in range(len(C_list)):
                c = C_list[i]
                d = D_list[i]
                level = self.get_coordination_level(d)
                # 样本行：根据语言显示"样本 N"或"Sample N"
                data_rows.append([
                    f"样本 {i + 1}" if self.current_language == 'zh' else f"Sample {i + 1}",
                    "", ""
                ])
                # 耦合度行：根据语言显示"耦合度"或"Coupling Degree"
                data_rows.append([
                    "耦合度" if self.current_language == 'zh' else "Coupling Degree",
                    round(c, 4), ""
                ])
                # 耦合协调度行：根据语言显示"耦合协调度"或"Coupling Coordination Degree"
                data_rows.append([
                    "耦合协调度" if self.current_language == 'zh' else "Coupling Coordination Degree",
                    round(d, 4), level
                ])
                # 系统综合指数行：根据语言显示"系统 N 综合指数"或"System N Comprehensive Index"
                for sys_idx, sys_id in enumerate(sorted(unique_systems)):
                    data_rows.append([
                        f"系统 {sys_id} 综合指数" if self.current_language == 'zh' else f"System {sys_id} Comprehensive Index",
                        round(u_matrix[i, sys_idx], 4), ""
                    ])
                data_rows.append(["", "", ""])

            # 表头：根据语言显示对应表头文本
            headers_result = [
                "统计量" if self.current_language == 'zh' else "Statistic",
                "统计量值" if self.current_language == 'zh' else "Statistic Value",
                "协调等级/说明" if self.current_language == 'zh' else "Coordination Level/Description"
            ]
            df_result = pd.DataFrame(data_rows, columns=headers_result)

            # 添加表头和系统信息
            headers_info = [
                [LANGUAGES[self.current_language]['headers_title'], ", ".join(map(str, headers)), ""],
                [LANGUAGES[self.current_language]['system_count'].format(num_systems), "", ""],
                [LANGUAGES[self.current_language]['system_indices'], str(system_info), ""]
            ]
            headers_df = pd.DataFrame(headers_info, columns=df_result.columns)

            # 添加解释说明
            explanations = LANGUAGES[self.current_language]['explanation']
            interpretations = LANGUAGES[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(
                columns=["耦合度", "耦合协调度", "耦合度分布直方图", "耦合协调度分布直方图", "系统综合指数折线图"])
            explanation_df.insert(0, "统计量_解释说明", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(
                columns=["耦合度", "耦合协调度", "耦合度分布直方图", "耦合协调度分布直方图", "系统综合指数折线图"])
            interpretation_df.insert(0, "统计量_结果解读",
                                     "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加表头和系统信息（保留表格形式）
                doc.add_heading(LANGUAGES[self.current_language]['headers_title'], level=2)
                headers_table = doc.add_table(rows=len(headers_info), cols=3)
                for row_idx, row_data in enumerate(headers_info):
                    for col_idx, value in enumerate(row_data):
                        headers_table.cell(row_idx, col_idx).text = str(value)
                doc.add_paragraph()  # 空行分隔

                # 添加统计量数据（保留表格形式）
                doc.add_heading("统计量数据" if self.current_language == 'zh' else "Statistical Data", level=2)
                data_table = doc.add_table(rows=1, cols=len(df_result.columns))
                # 添加表头
                hdr_cells = data_table.rows[0].cells
                for col_idx, header in enumerate(df_result.columns):
                    hdr_cells[col_idx].text = header
                # 添加数据行
                for row in df_result.values:
                    row_cells = data_table.add_row().cells
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx].text = str(value)
                doc.add_paragraph()  # 空行分隔

                # 添加解释说明（改为项目符号列表）
                doc.add_heading(LANGUAGES[self.current_language]['explanation_title'], level=2)
                explanation_para = doc.add_paragraph()
                explanations = LANGUAGES[self.current_language]['explanation']
                for key, value in explanations.items():
                    run = explanation_para.add_run(f"• {key}: {value}\n")  # 项目符号+内容
                doc.add_paragraph()  # 空行分隔

                # 添加结果解读（改为项目符号列表）
                doc.add_heading(LANGUAGES[self.current_language]['interpretation_title'], level=2)
                interpretation_para = doc.add_paragraph()
                interpretations = LANGUAGES[self.current_language]['interpretation']
                for key, value in interpretations.items():
                    run = interpretation_para.add_run(f"• {key}: {value}\n")  # 项目符号+内容
                doc.add_paragraph()  # 空行分隔

                # 生成可视化图表
                fig, axes = plt.subplots(3, 1, figsize=(10, 16))

                # 在图表上方添加表头信息
                plt.figtext(0.5, 0.95,
                            f"{LANGUAGES[self.current_language]['headers_title']}: {', '.join(map(str, headers))}\n"
                            f"{LANGUAGES[self.current_language]['system_count'].format(num_systems)}",
                            ha='center', fontsize=10, bbox=dict(facecolor='none', edgecolor='black', pad=8.0))

                # 耦合度直方图
                axes[0].hist(C_list, bins=min(10, len(C_list)), alpha=0.7)
                axes[0].set_title(
                    '耦合度分布直方图' if self.current_language == 'zh' else 'Histogram of Coupling Degree Distribution')
                axes[0].set_xlabel('耦合度' if self.current_language == 'zh' else 'Coupling Degree')
                axes[0].set_ylabel('频数' if self.current_language == 'zh' else 'Frequency')

                # 耦合协调度直方图
                axes[1].hist(D_list, bins=min(10, len(D_list)), alpha=0.7)
                axes[1].set_title(
                    '耦合协调度分布直方图' if self.current_language == 'zh' else 'Histogram of Coupling Coordination Degree Distribution')
                axes[1].set_xlabel('耦合协调度' if self.current_language == 'zh' else 'Coupling Coordination Degree')
                axes[1].set_ylabel('频数' if self.current_language == 'zh' else 'Frequency')

                # 系统综合指数折线图
                for sys_idx, sys_id in enumerate(sorted(unique_systems)):
                    axes[2].plot(u_matrix[:, sys_idx], marker='o',
                                 label=f'系统 {sys_id}' if self.current_language == 'zh' else f'System {sys_id}')
                axes[2].set_title(
                    '系统综合指数折线图' if self.current_language == 'zh' else 'Line Chart of System Comprehensive Index')
                axes[2].set_xlabel('样本序号' if self.current_language == 'zh' else 'Sample Index')
                axes[2].set_ylabel('综合指数' if self.current_language == 'zh' else 'Comprehensive Index')
                axes[2].legend()

                # 保存图片
                img_path = os.path.splitext(save_path)[0] + '_plots.png'
                plt.tight_layout(rect=[0, 0, 1, 0.92])  # 调整布局，为顶部标题留出空间
                plt.savefig(img_path, dpi=300)
                plt.close()

                # 将图片插入到 Word 文档中
                doc.add_picture(img_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(save_path)
                self.result_label.config(text=result_msg, wraplength=400)

            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'
        self.root.title(LANGUAGES[self.current_language]["title"])
        self.select_button.config(text=LANGUAGES[self.current_language]['select_button'])
        self.analyze_button.config(text=LANGUAGES[self.current_language]['analyze_button'])
        self.language_label.config(text=LANGUAGES[self.current_language]['switch_language'])
        self.open_excel_label.config(text=LANGUAGES[self.current_language]["open_excel_button_text"])
        # 切换语言时更新提示信息
        current_text = self.file_entry.get()
        if current_text in [LANGUAGES['zh']['file_entry_placeholder'], LANGUAGES['en']['file_entry_placeholder']]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
            self.file_entry.configure(style="Gray.TEntry")

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建自定义样式
        style = ttk.Style()
        style.configure("Gray.TEntry", foreground="gray")

        # 创建中间容器，用于居中所有元素
        center_frame = ttk.Frame(self.root)
        center_frame.place(relx=0.5, rely=0.5, anchor="center")  # 容器居中

        # 1. 文件选择按钮（放入中间容器）
        self.select_button = ttk.Button(
            center_frame,
            text=LANGUAGES[self.current_language]['select_button'],
            command=self.select_file,
            bootstyle=PRIMARY
        )
        self.select_button.pack(pady=10)

        # 2. 文件路径输入框（放入中间容器）
        self.file_entry = ttk.Entry(center_frame, width=50, style="Gray.TEntry")
        self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
        self.file_entry.pack(pady=5)
        self.file_entry.bind("<FocusIn>", self.on_entry_click)
        self.file_entry.bind("<FocusOut>", self.on_focusout)

        # 3. 分析按钮（放入中间容器）
        self.analyze_button = ttk.Button(
            center_frame,
            text=LANGUAGES[self.current_language]['analyze_button'],
            command=self.analyze_file,
            bootstyle=SUCCESS
        )
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            center_frame,
            text=LANGUAGES[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 4. 语言切换标签（放入中间容器）
        self.language_label = ttk.Label(
            center_frame,
            text=LANGUAGES[self.current_language]['switch_language'],
            foreground="gray",
            cursor="hand2"
        )
        self.language_label.pack(pady=10)
        self.language_label.bind("<Button-1>", self.switch_language)

        # 5. 结果显示标签（放在容器外，底部显示）
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT, wraplength=400)
        self.result_label.pack(side=tk.BOTTOM, pady=20)  # 固定在底部

    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = CouplingCoordinationDegreeModelAnalysisApp()
    app.run()

if __name__ == "__main__":
    run_app()
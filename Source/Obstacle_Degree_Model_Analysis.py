import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
import matplotlib.pyplot as plt
import pathlib
from docx import Document
from docx.shared import Inches


# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
languages = {
    'zh': {
        'title': "障碍度模型",
        'select_button_text': "选择文件",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'analyze_button_text': "分析文件",
        'no_file_selected': "请选择文件。",
        'file_not_exists': "文件不存在，请重新选择。",
        'analysis_error': "分析文件时出错: {}",
        'analysis_complete': "分析完成，结果已保存到 {}，相关图片已保存到桌面。",
        'no_save_path_selected': "未选择保存路径，结果未保存。",
        'switch_language_button_text': "中/英",
        'open_excel_button_text': "示例数据",
        'data_valid_error': "数据格式错误：{}",
        'explanation': {
            "障碍度": "衡量设计因素对方案落地的阻碍程度（设计学专用公式：权重×(1-状态值)×影响系数）",
            "障碍度柱状图": "展示各设计因素障碍度的分布情况",
            "障碍度排序条形图": "按障碍度从高到低展示设计因素",
            "累计占比图": "展示障碍度累计占比情况"
        },
        'interpretation': {
            "障碍度": "障碍度越高，该设计因素对方案落地的阻碍作用越大",
            "障碍度柱状图": "可直观比较不同设计因素障碍度的大小",
            "障碍度排序条形图": "便于快速识别主要障碍因素",
            "累计占比图": "可识别关键少数障碍因素（帕累托法则）"
        }
    },
    'en': {
        'title': "Obstacle Degree Model",
        'select_button_text': "Select File",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'analyze_button_text': "Analyze File",
        'no_file_selected': "Please select a valid file path.",
        'file_not_exists': "The file does not exist. Please select again.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'analysis_complete': "Analysis completed. The results have been saved to {}, and relevant images have been saved to the desktop.",
        'no_save_path_selected': "No save path selected. The results were not saved.",
        'switch_language_button_text': "Chinese/English",
        'open_excel_button_text': "Example data",
        'data_valid_error': "Data format error: {}",
        'explanation': {
            "Obstacle Degree": "Measure the obstruction degree of design factors to scheme implementation (Design-specific formula: weight×(1-status value)×influence coefficient)",
            "Obstacle Degree Bar Chart": "Show the distribution of obstacle degrees of each design factor",
            "Sorted Obstacle Degree Bar Chart": "Show design factors by obstacle degree from high to low",
            "Cumulative Proportion Chart": "Show cumulative proportion of obstacle degrees"
        },
        'interpretation': {
            "Obstacle Degree": "The higher the obstacle degree, the greater the obstructive effect of the design factor on scheme implementation",
            "Obstacle Degree Bar Chart": "It can visually compare the obstacle degrees of different design factors",
            "Sorted Obstacle Degree Bar Chart": "Facilitates quick identification of major obstacle factors",
            "Cumulative Proportion Chart": "Can identify key minority obstacle factors (Pareto principle)"
        }
    }
}


class ObstacleDegreeModelAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        # 获取桌面路径
        self.desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        # 图片直接保存到桌面，不使用子文件夹
        self.img_save_dir = self.desktop_path

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data52.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def validate_data(self, df):
        """验证Excel数据格式是否符合设计学分析要求"""
        # 检查是否有表头
        if df.empty:
            return False, "数据不能为空"

        # 检查列数是否为3列（因素名称、权重、状态值）
        if df.shape[1] != 3:
            return False, f"需包含3列数据（因素名称、权重、状态值），实际为{df.shape[1]}列"

        # 检查第二列（权重）是否为数值且在0-1之间
        try:
            weights = df.iloc[:, 1].astype(float)
            if not ((weights >= 0) & (weights <= 1)).all():
                return False, "权重值必须在0-1之间"
        except ValueError:
            return False, "权重列必须为数值类型"

        # 检查第三列（状态值）是否为数值且在0-1之间
        try:
            status_values = df.iloc[:, 2].astype(float)
            if not ((status_values >= 0) & (status_values <= 1)).all():
                return False, "状态值必须在0-1之间"
        except ValueError:
            return False, "状态值列必须为数值类型"

        return True, "数据验证通过"

    def obstacle_degree_method(self, data):
        weights = data[:, 1].astype(float)
        status_values = data[:, 2].astype(float)

        # 计算影响系数（改进版）
        valid_weights = weights[weights > 0]  # 排除0权重
        if len(valid_weights) == 0:
            # 所有权重都为0时的特殊处理
            influence_coefficients = np.zeros_like(weights)
        else:
            avg_weight = np.mean(valid_weights)
            # 处理平均权重为0的情况，使用最小非零权重替代
            if avg_weight == 0:
                avg_weight = np.min(valid_weights)
            influence_coefficients = np.where(
                weights > 0,  # 条件
                weights / avg_weight,  # 满足条件时的计算
                0  # 权重为0时的影响系数
            )

        # 设计学障碍度计算公式
        obstacle_degrees = weights * (1 - status_values) * influence_coefficients
        return obstacle_degrees

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 打开 Excel 文件（包含表头）
            df = pd.read_excel(file_path)
            # 数据验证
            is_valid, msg = self.validate_data(df)
            if not is_valid:
                self.result_label.config(text=languages[self.current_language]["data_valid_error"].format(msg))
                return

            data = df.values
            # 获取因素名称（从第一列读取）
            factor_names = df.iloc[:, 0].tolist()
            # 计算障碍度
            obstacle_degrees = self.obstacle_degree_method(data)

            # 计算各因素障碍度的百分比贡献
            total_degree = np.sum(obstacle_degrees)
            contributions = obstacle_degrees / total_degree * 100 if total_degree != 0 else np.zeros_like(
                obstacle_degrees)

            # 计算排序后的累积贡献（帕累托分析）
            sorted_indices = np.argsort(obstacle_degrees)[::-1]
            sorted_contributions = contributions[sorted_indices]
            cumulative_contributions = np.cumsum(sorted_contributions)

            # 识别关键障碍因素（累计贡献达到80%的因素）
            key_factor_count = np.argmax(cumulative_contributions >= 80) + 1 if np.any(
                cumulative_contributions >= 80) else len(sorted_contributions)
            key_factors = [factor_names[i] for i in sorted_indices[:key_factor_count]]

            # 准备分析结果数据
            result_data = []
            for i in range(len(factor_names)):
                result_data.append([
                    factor_names[i],
                    round(obstacle_degrees[i], 4)  # 只保留因素名称和障碍度两列
                ])

            # 整理结果表格 headers（修改为两列）
            headers = [
                "因素名称" if self.current_language == 'zh' else "Factor Name",
                "障碍度" if self.current_language == 'zh' else "Obstacle Degree"
            ]
            df_result = pd.DataFrame(result_data, columns=headers)

            # 添加解释说明
            explanations = languages[self.current_language]['explanation']
            interpretation = languages[self.current_language]['interpretation']

            # 使用三元表达式根据语言动态设置列名
            explanation_df = pd.DataFrame([explanations])
            explanation_df.insert(
                0,
                "统计量_解释说明" if self.current_language == 'zh' else "Statistic_Explanation",
                "解释说明" if self.current_language == 'zh' else "Explanation"
            )

            interpretation_df = pd.DataFrame([interpretation])
            interpretation_df.insert(
                0,
                "统计量_结果解读" if self.current_language == 'zh' else "Statistic_Interpretation",
                "结果解读" if self.current_language == 'zh' else "Interpretation"
            )

            # 生成多种可视化图表
            image_paths = []

            # 1. 障碍度柱状图
            plt.figure(figsize=(10, 6))
            plt.bar(factor_names, obstacle_degrees)
            plt.title('障碍度柱状图' if self.current_language == 'zh' else 'Bar Chart of Obstacle Degree')
            plt.xlabel('设计因素' if self.current_language == 'zh' else 'Design Factors')
            plt.ylabel('障碍度' if self.current_language == 'zh' else 'Obstacle Degree')
            plt.xticks(rotation=0, ha='center')
            plt.tight_layout()
            bar_path = os.path.join(self.img_save_dir, 'obstacle_degree_bar.png')
            plt.savefig(bar_path)
            plt.close()
            image_paths.append(bar_path)

            # 2. 排序条形图（从高到低）
            sorted_indices = np.argsort(obstacle_degrees)[::-1]
            sorted_names = [factor_names[i] for i in sorted_indices]
            sorted_degrees = [obstacle_degrees[i] for i in sorted_indices]

            plt.figure(figsize=(10, 6))
            plt.barh(sorted_names, sorted_degrees)
            plt.title('障碍度排序条形图' if self.current_language == 'zh' else 'Sorted Bar Chart of Obstacle Degree')
            plt.xlabel('障碍度' if self.current_language == 'zh' else 'Obstacle Degree')
            plt.ylabel('设计因素' if self.current_language == 'zh' else 'Design Factors')
            plt.tight_layout()
            sorted_path = os.path.join(self.img_save_dir, 'obstacle_degree_sorted.png')
            plt.savefig(sorted_path)
            plt.close()
            image_paths.append(sorted_path)

            # 3. 累计占比图
            total = np.sum(sorted_degrees)
            cumulative = np.cumsum(sorted_degrees) / total * 100

            plt.figure(figsize=(10, 6))
            plt.plot(sorted_names, cumulative, 'o-', color='r')
            plt.axhline(y=80, color='g', linestyle='--', label='80%线')
            plt.title('障碍度累计占比图' if self.current_language == 'zh' else 'Cumulative Proportion Chart')
            plt.xlabel(
                '设计因素（按障碍度排序）' if self.current_language == 'zh' else 'Design Factors (sorted by degree)')
            plt.ylabel('累计占比（%）' if self.current_language == 'zh' else 'Cumulative Proportion (%)')
            plt.xticks(rotation=0, ha='center')
            plt.legend()
            plt.tight_layout()
            cumulative_path = os.path.join(self.img_save_dir, 'obstacle_degree_cumulative.png')
            plt.savefig(cumulative_path)
            plt.close()
            image_paths.append(cumulative_path)

            # 设置Word默认保存路径（桌面）
            default_filename = "障碍度分析结果.docx" if self.current_language == 'zh' else "Obstacle Degree Analysis Results.docx"
            save_path = os.path.join(self.desktop_path, default_filename)

            # 如果文件已存在，添加数字后缀避免覆盖
            counter = 1
            while os.path.exists(save_path):
                base, ext = os.path.splitext(default_filename)
                save_path = os.path.join(self.desktop_path, f"{base}_{counter}{ext}")
                counter += 1

            # 创建Word文档
            doc = Document()
            doc.add_heading(
                '设计学障碍度模型分析结果' if self.current_language == 'zh' else 'Design Obstacle Degree Model Analysis Results',
                0)

            # 添加分析结果表格
            doc.add_heading(
                '障碍度计算结果' if self.current_language == 'zh' else 'Obstacle Degree Calculation Results',
                level=1)
            table = doc.add_table(rows=1, cols=len(df_result.columns))
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(df_result.columns):
                hdr_cells[i].text = col

            for index, row in df_result.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)

            # 添加帕累托分析结果
            doc.add_heading('帕累托分析结果' if self.current_language == 'zh' else 'Pareto Analysis Results', level=1)

            # 准备帕累托分析表格数据
            pareto_data = []
            for i in range(len(sorted_indices)):
                factor_idx = sorted_indices[i]
                pareto_data.append([
                    factor_names[factor_idx],  # 因素名称
                    round(sorted_degrees[i], 4),  # 障碍度（已排序）
                    f"{sorted_contributions[i]:.2f}%",  # 百分比贡献
                    f"{cumulative_contributions[i]:.2f}%"  # 累计贡献
                ])

            # 添加表格
            pareto_headers = [
                "因素名称" if self.current_language == 'zh' else "Factor Name",
                "障碍度" if self.current_language == 'zh' else "Obstacle Degree",
                "百分比贡献" if self.current_language == 'zh' else "Percentage Contribution",
                "累计贡献" if self.current_language == 'zh' else "Cumulative Contribution"
            ]
            pareto_table = doc.add_table(rows=1, cols=4)
            hdr_cells = pareto_table.rows[0].cells
            for i, col in enumerate(pareto_headers):
                hdr_cells[i].text = col

            for row_data in pareto_data:
                row_cells = pareto_table.add_row().cells
                for i, value in enumerate(row_data):
                    row_cells[i].text = str(value)

            # 添加关键障碍因素说明
            key_factors_text = ", ".join(key_factors)
            doc.add_paragraph(
                f"关键障碍因素（累计贡献≥80%）：{key_factors_text}"
                if self.current_language == 'zh'
                else f"Key obstacle factors (cumulative contribution ≥80%): {key_factors_text}"
            )

            # 添加解释说明
            doc.add_heading('指标解释' if self.current_language == 'zh' else 'Indicator Explanation', level=1)
            for idx, row in explanation_df.iterrows():
                for col, val in row.items():
                    if pd.notna(val):
                        doc.add_paragraph(f"{col}: {val}")

            # 添加结果解读
            doc.add_heading('结果解读' if self.current_language == 'zh' else 'Result Interpretation', level=1)
            for idx, row in interpretation_df.iterrows():
                for col, val in row.items():
                    if pd.notna(val):
                        doc.add_paragraph(f"{col}: {val}")

            # 添加可视化图表
            doc.add_heading('可视化分析' if self.current_language == 'zh' else 'Visual Analysis', level=1)
            for img_path in image_paths:
                doc.add_picture(img_path, width=Inches(6))
                img_name = os.path.basename(img_path).replace('_', ' ').replace('.png', '')
                doc.add_paragraph(img_name)

            # 保存Word文档
            doc.save(save_path)
            self.result_label.config(text=languages[self.current_language]["analysis_complete"].format(save_path),
                                     wraplength=400)

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]["analysis_error"].format(str(e)))

    def switch_language(self, event):
        prev_language = self.current_language
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 图片直接保存到桌面，不使用子文件夹
        self.img_save_dir = self.desktop_path

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])
        self.open_excel_label.config(text=languages[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=languages[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame,
                                               text=languages[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = ObstacleDegreeModelAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()
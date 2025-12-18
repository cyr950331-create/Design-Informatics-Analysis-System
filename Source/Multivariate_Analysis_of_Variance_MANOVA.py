import tkinter as tk
from tkinter import filedialog, simpledialog
import os
import numpy as np
import pandas as pd
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
from statsmodels.multivariate.manova import MANOVA
from docx import Document
from docx.shared import Inches

# 设置支持中文的字体
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典，包含所有需要翻译的文本
languages = {
    'zh': {
        'title': "多元方差分析",
        'select_button_text': "选择文件",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'analyze_button_text': "分析文件",
        'no_file_selected': "请选择文件。",
        'file_not_exists': "文件不存在，请重新选择。",
        'analysis_error': "分析文件时出错: {}",
        'analysis_complete': "分析完成，结果已保存到 {}，相关图片已保存。",
        'no_save_path_selected': "未选择保存路径，结果未保存。",
        'switch_language_button_text': "中/英",
        'open_excel_button_text': "示例数据",
        'input_info': "输入信息",
        'input_sample_cols': "请输入样本列（分组变量，用逗号分隔）",
        'input_indicator_cols': "请输入指标列（因变量，用逗号分隔）",
        'input_incomplete': "未输入完整的变量名，分析取消。",
        'invalid_columns': "无效的列名: {}",
        'non_numeric_columns': "以下指标列不是数值类型: {}",
        'missing_values_note': "注意：检测到{}个非数值数据，已自动删除包含非数值的行",
        'explanation': {
            "multivariate_analysis": "用于同时比较多个因变量在不同组之间的均值是否存在显著差异。",
            "sample_size": "每个组中的观测值数量。",
            "mean_value": "每个因变量在每个组中的平均值。",
            "f_statistic": "衡量组间差异与组内差异的比值。",
            "degrees_of_freedom": "用于计算F分布的参数。",
            "p_value": "p值小于显著性水平（通常为0.05）时，拒绝原假设，认为组间存在显著差异；否则，接受原假设，认为组间无显著差异。",
            "effect_size": "反映自变量对因变量的影响程度。"
        },
        'interpretation': {
            "f_statistic": "F统计量越大，说明组间差异越显著。",
            "p_value": "用于判断组间是否存在显著差异。",
            "degrees_of_freedom": "影响F分布的形状，进而影响p值的计算。",
            "sample_size": "样本量的大小会影响统计检验的功效，较大的样本量通常能提供更准确的结果。",
            "mean_value": "反映每个因变量在每个组中的平均水平。",
            "effect_size": "效应量越大，说明自变量对因变量的影响越大。"
        },
        'document_headers': {
            "analysis_parameters": "分析参数",
            "sample_columns": "样本列（分组变量）",
            "indicator_columns": "指标列（因变量）",
            "analysis_results": "分析结果",
            "statistic_explanations": "统计量解释说明",
            "statistic_interpretations": "统计量结果解读",
            "boxplot": "箱线图",
            "barplot": "柱状图"
        },
        'table_headers': {
            "statistic": "统计量",
            "f_statistic": "F统计量",
            "df_between": "组间自由度",
            "df_within": "组内自由度",
            "p_value": "p值",
            "pillai_trace": "Pillai's Trace",
            "partial_eta_squared": "偏Eta平方",
            "cohens_f": "Cohen's f效应量"
        },
        'chart_labels': {
            "boxplot_title": "箱线图",
            "boxplot_xlabel": "因变量",
            "boxplot_ylabel": "值",
            "barplot_title": "各组均值柱状图",
            "barplot_xlabel": "组",
            "barplot_ylabel": "均值"
        },
        'statistic_names': {
            "manova": "多元方差分析（MANOVA）",
            "sample_size": "样本量",
            "mean_value": "均值"
        },
        'explanation_labels': {
            "explanation": "解释说明",
            "interpretation": "结果解读"
        }
    },
    'en': {
        'title': "Multivariate Analysis of Variance",
        'select_button_text': "Select File",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'analyze_button_text': "Analyze File",
        'no_file_selected': "Please select a valid file path.",
        'file_not_exists': "The file does not exist. Please select again.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'analysis_complete': "Analysis completed. The results have been saved to {}, and the relevant images have been saved.",
        'no_save_path_selected': "No save path selected. The results were not saved.",
        'switch_language_button_text': "Chinese/English",
        'open_excel_button_text': "Example data",
        'input_info': "Input Information",
        'input_sample_cols': "Please enter sample columns (group variables, separated by commas)",
        'input_indicator_cols': "Please enter indicator columns (dependent variables, separated by commas)",
        'input_incomplete': "Incomplete variable names entered, analysis canceled.",
        'invalid_columns': "Invalid column names: {}",
        'non_numeric_columns': "The following indicator columns are not numeric: {}",
        'missing_values_note': "Note: {} non-numeric data points detected, rows containing non-numeric values have been automatically removed",
        'explanation': {
            "multivariate_analysis": "Used to simultaneously compare whether the means of multiple dependent variables differ significantly between groups.",
            "sample_size": "The number of observations in each group.",
            "mean_value": "The average value of each dependent variable in each group.",
            "f_statistic": "Measures the ratio of between-group variance to within-group variance.",
            "degrees_of_freedom": "Parameters used to calculate the F-distribution.",
            "p_value": "When the p-value is less than the significance level (usually 0.05), the null hypothesis is rejected, indicating a significant difference between groups; otherwise, the null hypothesis is accepted, indicating no significant difference.",
            "effect_size": "Reflects the influence of the independent variable on the dependent variables."
        },
        'interpretation': {
            "f_statistic": "The larger the F-statistic, the more significant the between-group difference.",
            "p_value": "Used to determine whether there is a significant difference between groups.",
            "degrees_of_freedom": "Affects the shape of the F-distribution, which in turn affects the calculation of the p-value.",
            "sample_size": "The sample size affects the power of the statistical test. A larger sample size usually provides more accurate results.",
            "mean_value": "Reflects the average level of each dependent variable in each group.",
            "effect_size": "The larger the effect size, the greater the influence of the independent variable on the dependent variables."
        },
        'document_headers': {
            "analysis_parameters": "Analysis Parameters",
            "sample_columns": "Sample Column(s) (Group Variable(s))",
            "indicator_columns": "Indicator Columns (Dependent Variables)",
            "analysis_results": "Analysis Results",
            "statistic_explanations": "Statistical Explanations",
            "statistic_interpretations": "Statistical Interpretations",
            "boxplot": "Box Plot",
            "barplot": "Bar Chart"
        },
        'table_headers': {
            "statistic": "Statistic",
            "f_statistic": "F Statistic",
            "df_between": "Between Groups DF",
            "df_within": "Within Groups DF",
            "p_value": "p-value",
            "pillai_trace": "Pillai's Trace",
            "partial_eta_squared": "Partial Eta Squared",
            "cohens_f": "Cohen's f Effect Size"
        },
        'chart_labels': {
            "boxplot_title": "Box Plot",
            "boxplot_xlabel": "Dependent Variables",
            "boxplot_ylabel": "Values",
            "barplot_title": "Mean Values by Group",
            "barplot_xlabel": "Groups",
            "barplot_ylabel": "Mean Value"
        },
        'statistic_names': {
            "manova": "Multivariate Analysis of Variance (MANOVA)",
            "sample_size": "Sample Size",
            "mean_value": "Mean Value"
        },
        'explanation_labels': {
            "explanation": "Explanation",
            "interpretation": "Interpretation"
        }
    }
}

class MultivariateManovaApp:
    def __init__(self, root=None):
        # 当前语言，默认为中文
        self.current_language = "en"
        self.df = None
        self.all_columns = []

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data24.xls")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"{languages[self.current_language]['file_not_exists'].split('，')[0]}：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"{languages[self.current_language]['analysis_error'].format('')}{str(e)}")

    def select_file(self):
        # 确保主窗口在顶层
        self.root.attributes('-topmost', True)

        # 打开文件选择对话框，指定父窗口为当前主窗口，确保对话框在主窗口上方
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 关联到主窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )

        # 可根据需要决定是否保留主窗口始终顶层（如果只需要操作时顶层，这里可以设为False）
        self.root.attributes('-topmost', False)

        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

            # 加载数据并获取列信息
            try:
                self.df = pd.read_excel(file_path, header=0)
                self.df = self.df.reset_index(drop=True)
                self.all_columns = list(self.df.columns)
                self.result_label.config(text="")
            except Exception as e:
                self.result_label.config(text=languages[self.current_language]["analysis_error"].format(str(e)))
                self.df = None
                self.all_columns = []

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        if self.df is None:
            try:
                self.df = pd.read_excel(file_path, header=0)
                self.df = self.df.reset_index(drop=True)
                self.all_columns = list(self.df.columns)
            except Exception as e:
                self.result_label.config(text=languages[self.current_language]["analysis_error"].format(str(e)))
                return

        # 获取变量名（与中介分析保持一致）
        sample_cols_input = simpledialog.askstring(
            languages[self.current_language]['input_info'],
            languages[self.current_language]['input_sample_cols']
        )
        indicator_cols_input = simpledialog.askstring(
            languages[self.current_language]['input_info'],
            languages[self.current_language]['input_indicator_cols']
        )

        # 验证输入完整性
        if not sample_cols_input or not indicator_cols_input:
            self.result_label.config(text=languages[self.current_language]["input_incomplete"])
            return

        # 解析输入的列名
        sample_cols = [col.strip() for col in sample_cols_input.split(',') if col.strip()]
        indicator_cols = [col.strip() for col in indicator_cols_input.split(',') if col.strip()]

        # 验证列名有效性
        invalid_samples = [col for col in sample_cols if col not in self.all_columns]
        invalid_indicators = [col for col in indicator_cols if col not in self.all_columns]
        invalid_cols = invalid_samples + invalid_indicators

        if invalid_cols:
            self.result_label.config(
                text=languages[self.current_language]["invalid_columns"].format(", ".join(invalid_cols)))
            return

        try:
            # 检查指标列是否为数值类型
            non_numeric_cols = []
            for col in indicator_cols:
                if not pd.api.types.is_numeric_dtype(self.df[col]):
                    non_numeric_cols.append(col)

            if non_numeric_cols:
                self.result_label.config(
                    text=languages[self.current_language]["non_numeric_columns"].format(", ".join(non_numeric_cols))
                )
                return

            # 尝试将指标列转换为数值类型
            numeric_df = self.df[indicator_cols].copy()
            for col in indicator_cols:
                # 强制转换为数值类型，无法转换的设为NaN
                numeric_df[col] = pd.to_numeric(numeric_df[col], errors='coerce')

            # 检查并处理缺失值
            missing_count = 0
            if numeric_df.isnull().any().any():
                missing_count = numeric_df.isnull().sum().sum()
                # 先创建分组数据
                group_data = self.df[sample_cols].copy()

                # 再处理缺失值
                valid_indices = numeric_df.dropna().index
                numeric_df = numeric_df.loc[valid_indices]
                group_data = group_data.loc[valid_indices]
                self.result_label.config(text=languages[self.current_language]["missing_values_note"].format(missing_count))

            # 检查转换后是否有非数值
            if numeric_df.isnull().any().any():
                raise ValueError(languages[self.current_language]["non_numeric_columns"].format(""))

            # 创建组合的分组变量
            group_data = self.df[sample_cols].copy()
            # 在处理完缺失值后重新计算组合分组变量
            if len(sample_cols) > 1:
                group_var = group_data.astype(str).agg('||'.join, axis=1)
                group_var_name = '||'.join(sample_cols)
            else:
                group_var = group_data.iloc[:, 0]
                group_var_name = sample_cols[0]

            # 将分组变量转换为虚拟变量（解决object类型问题）
            # 使用get_dummies进行转换，并删除第一列以避免共线性
            group_dummies = pd.get_dummies(group_var, prefix=group_var_name, drop_first=True)

            # 进行多元方差分析（使用虚拟变量作为自变量）
            manova = MANOVA(endog=numeric_df, exog=group_dummies)
            results = manova.mv_test()

            # 提取统计量（使用Pillai's Trace作为替代，更稳健）
            f_stat = results.results['x0']['stat']['F Value']['Pillai\'s trace']
            df_between = results.results['x0']['stat']['Num DF']['Pillai\'s trace']
            df_within = results.results['x0']['stat']['Den DF']['Pillai\'s trace']
            p_value = results.results['x0']['stat']['Pr > F']['Pillai\'s trace']

            # 正确计算Pillai's Trace效应量
            pillai_value = results.results['x0']['stat']['Value']['Pillai\'s trace']
            # 计算偏Eta平方
            partial_eta_squared = df_between * f_stat / (df_between * f_stat + df_within)
            # 计算Cohen's f效应量
            cohens_f = np.sqrt(partial_eta_squared / (1 - partial_eta_squared))

            # 计算样本量和均值
            sample_sizes = self.df.groupby(group_var).size()
            means = numeric_df.groupby(group_var).mean()

            # 整理结果数据
            data = [
                [languages[self.current_language]['statistic_names']['manova'], f_stat, df_between, df_within, p_value, pillai_value, partial_eta_squared, cohens_f],
                [languages[self.current_language]['statistic_names']['sample_size'], sample_sizes.to_dict(), "", "", "", "", "", ""],
                [languages[self.current_language]['statistic_names']['mean_value'], means.to_dict(), "", "", "", "", "", ""]
            ]
            headers = [
                languages[self.current_language]['table_headers']['statistic'],
                languages[self.current_language]['table_headers']['f_statistic'],
                languages[self.current_language]['table_headers']['df_between'],
                languages[self.current_language]['table_headers']['df_within'],
                languages[self.current_language]['table_headers']['p_value'],
                languages[self.current_language]['table_headers']['pillai_trace'],
                languages[self.current_language]['table_headers']['partial_eta_squared'],
                languages[self.current_language]['table_headers']['cohens_f']
            ]

            # 定义df_result
            df_result = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            explanations = languages[self.current_language]['explanation']
            interpretations = languages[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([{
                languages[self.current_language]['statistic_names']['manova']: explanations['multivariate_analysis'],
                languages[self.current_language]['statistic_names']['sample_size']: explanations['sample_size'],
                languages[self.current_language]['statistic_names']['mean_value']: explanations['mean_value'],
                languages[self.current_language]['table_headers']['f_statistic']: explanations['f_statistic'],
                languages[self.current_language]['table_headers']['df_between']: explanations['degrees_of_freedom'],
                languages[self.current_language]['table_headers']['p_value']: explanations['p_value'],
                languages[self.current_language]['table_headers']['cohens_f']: explanations['effect_size']
            }])
            explanation_df.insert(0, languages[self.current_language]['explanation_labels']['explanation'],
                                 languages[self.current_language]['explanation_labels']['explanation'])

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([{
                languages[self.current_language]['table_headers']['f_statistic']: interpretations['f_statistic'],
                languages[self.current_language]['table_headers']['p_value']: interpretations['p_value'],
                languages[self.current_language]['table_headers']['df_between']: interpretations['degrees_of_freedom'],
                languages[self.current_language]['statistic_names']['sample_size']: interpretations['sample_size'],
                languages[self.current_language]['statistic_names']['mean_value']: interpretations['mean_value'],
                languages[self.current_language]['table_headers']['cohens_f']: interpretations['effect_size']
            }])
            interpretation_df.insert(0, languages[self.current_language]['explanation_labels']['interpretation'],
                                    languages[self.current_language]['explanation_labels']['interpretation'])

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加分析参数说明
                doc.add_heading(languages[self.current_language]['document_headers']['analysis_parameters'], level=1)
                param_table = doc.add_table(rows=2, cols=2)
                param_table.cell(0, 0).text = languages[self.current_language]['document_headers']['sample_columns']
                param_table.cell(0, 1).text = ", ".join(sample_cols)
                param_table.cell(1, 0).text = languages[self.current_language]['document_headers']['indicator_columns']
                param_table.cell(1, 1).text = ", ".join(indicator_cols)

                # 添加分析结果表格
                doc.add_heading(languages[self.current_language]['document_headers']['analysis_results'], level=1)
                table = doc.add_table(rows=df_result.shape[0] + 1, cols=df_result.shape[1])
                table.style = 'Table Grid'
                # 调整列宽
                for col in table.columns:
                    col.width = Inches(1.2)
                # 添加表头
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(headers):
                    hdr_cells[col_idx].text = header
                # 添加数据行
                for row_idx in range(df_result.shape[0]):
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx in range(df_result.shape[1]):
                        row_cells[col_idx].text = str(df_result.iloc[row_idx, col_idx])

                # 添加解释说明列表
                doc.add_heading(languages[self.current_language]['document_headers']['statistic_explanations'], level=1)
                doc.add_paragraph(f"{languages[self.current_language]['statistic_names']['manova']}: {explanations['multivariate_analysis']}")
                doc.add_paragraph(f"{languages[self.current_language]['statistic_names']['sample_size']}: {explanations['sample_size']}")
                doc.add_paragraph(f"{languages[self.current_language]['statistic_names']['mean_value']}: {explanations['mean_value']}")
                doc.add_paragraph(f"{languages[self.current_language]['table_headers']['f_statistic']}: {explanations['f_statistic']}")
                doc.add_paragraph(f"{languages[self.current_language]['table_headers']['df_between']}: {explanations['degrees_of_freedom']}")
                doc.add_paragraph(f"{languages[self.current_language]['table_headers']['p_value']}: {explanations['p_value']}")
                doc.add_paragraph(f"{languages[self.current_language]['table_headers']['cohens_f']}: {explanations['effect_size']}")

                # 添加结果解读列表
                doc.add_heading(languages[self.current_language]['document_headers']['statistic_interpretations'], level=1)
                doc.add_paragraph(f"{languages[self.current_language]['table_headers']['f_statistic']}: {interpretations['f_statistic']}")
                doc.add_paragraph(f"{languages[self.current_language]['table_headers']['p_value']}: {interpretations['p_value']}")
                doc.add_paragraph(f"{languages[self.current_language]['table_headers']['df_between']}: {interpretations['degrees_of_freedom']}")
                doc.add_paragraph(f"{languages[self.current_language]['statistic_names']['sample_size']}: {interpretations['sample_size']}")
                doc.add_paragraph(f"{languages[self.current_language]['statistic_names']['mean_value']}: {interpretations['mean_value']}")
                doc.add_paragraph(f"{languages[self.current_language]['table_headers']['cohens_f']}: {interpretations['effect_size']}")

                # 绘制箱线图
                plt.figure(figsize=(10, 6))
                numeric_df.boxplot()
                plt.title(languages[self.current_language]['chart_labels']['boxplot_title'])
                plt.xlabel(languages[self.current_language]['chart_labels']['boxplot_xlabel'])
                plt.ylabel(languages[self.current_language]['chart_labels']['boxplot_ylabel'])
                boxplot_path = save_path.replace('.docx', '_boxplot.png')
                plt.savefig(boxplot_path)
                plt.close()

                # 绘制柱状图
                plt.figure(figsize=(10, 6))
                means.plot(kind='bar')
                plt.title(languages[self.current_language]['chart_labels']['barplot_title'])
                plt.xlabel(languages[self.current_language]['chart_labels']['barplot_xlabel'])
                plt.ylabel(languages[self.current_language]['chart_labels']['barplot_ylabel'])
                plt.xticks(rotation=0, ha='right')
                barplot_path = save_path.replace('.docx', '_barplot.png')
                plt.tight_layout()
                plt.savefig(barplot_path)
                plt.close()

                # 插入图片
                doc.add_heading(languages[self.current_language]['document_headers']['boxplot'], level=1)
                doc.add_picture(boxplot_path, width=Inches(6))
                doc.add_heading(languages[self.current_language]['document_headers']['barplot'], level=1)
                doc.add_picture(barplot_path, width=Inches(6))

                # 保存文档
                doc.save(save_path)
                self.result_label.config(text=languages[self.current_language]["analysis_complete"].format(save_path),
                                         wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]["no_save_path_selected"])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]["analysis_error"].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        current_entry_text = self.file_entry.get()
        # 只在占位符文本时才更新
        if current_entry_text == languages['zh' if self.current_language == 'en' else 'en']["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])
        self.open_excel_label.config(text=languages[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        # 窗口设置
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建主框架
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=languages[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 切换语言标签
        self.switch_language_label = ttk.Label(frame,
                                               text=languages[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        self.root.mainloop()

def run_app():
    app = MultivariateManovaApp()
    app.run()

if __name__ == "__main__":
    run_app()
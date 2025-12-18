import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
from scipy import stats
import matplotlib.pyplot as plt
import pathlib
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = ['Microsoft yahei', 'SimHei', 'SimSun', 'Arial']  # 增加支持上标的字体
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典，包含所有需要翻译的文本
LANGUAGES = {
    'zh': {
        'title': "卡方拟合优度检验",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "请选择文件。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "中/英",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'statistic_labels': {
            'chi2': "卡方值",
            'df': "自由度",
            'p': "p值",
            'alpha': "显著性水平",
            'result': "结果"
        },
        'significance': {
            'significant': "显著",
            'not_significant': "不显著"
        },
        'chart': {
            'observed': "观测频数",
            'expected': "理论频数",
            'x_label': "类别",
            'y_label': "频数",
            'title': "卡方拟合优度检验 - 观测频数 vs 理论频数"
        },
        'docx': {
            'title': "卡方拟合优度检验结果",
            'analysis_results': "分析结果",
            'statistic': "统计量",
            'value': "值",
            'explanation': "解释说明",
            'interpretation': "结果解读"
        },
        'explanation': {
            "卡方拟合优度检验": "用于检验观测数据是否符合某种理论分布。",
        },
        'interpretation': {
            "卡方值": "反映观测频数与理论频数的偏离程度，值越大，偏离越严重。",
            "p值": "若p值小于显著性水平（通常为0.05），则拒绝原假设，认为观测数据与理论分布存在显著差异。",
            "自由度": "自由度 = 类别数 - 约束条件数 - 1，用于确定卡方分布的参考分布。",
            "显著性水平": "判断结果是否显著的阈值，通常设为0.05，表示犯第一类错误的最大可接受概率。"
        }
    },
    'en': {
        'title': "Chi Square Goodness of Fit Test",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "Please select a file.",
        'analysis_success': "Analysis completed. Results saved to {}\n",
        'no_save_path': "No save path selected. Results not saved.",
        'analysis_error': "Error analyzing file: {}",
        'switch_language': "Chinese/English",
        'file_entry_placeholder': "Please enter the full path of the Excel file to analyze",
        'statistic_labels': {
            'chi2': "Chi-Square Value",
            'df': "Degrees of Freedom",
            'p': "p-value",
            'alpha': "Significance Level",
            'result': "Result"
        },
        'significance': {
            'significant': "significant",
            'not_significant': "not significant"
        },
        'chart': {
            'observed': "Observed Frequencies",
            'expected': "Expected Frequencies",
            'x_label': "Categories",
            'y_label': "Frequencies",
            'title': "Chi-Square Goodness-of-Fit Test - Observed vs Expected Frequencies"
        },
        'docx': {
            'title': "Chi-Square Goodness-of-Fit Test Results",
            'analysis_results': "Analysis Results",
            'statistic': "Statistic",
            'value': "Value",
            'explanation': "Explanation",
            'interpretation': "Interpretation"
        },
        'explanation': {
            "Chi-Square Goodness-of-Fit Test": "Used to test whether observed data conforms to a theoretical distribution.",
        },
        'interpretation': {
            "Chi-Square Value": "Reflects the degree of deviation between observed and theoretical frequencies. Larger values indicate greater deviation.",
            "p-value": "If the p-value is less than the significance level (usually 0.05), the null hypothesis is rejected, indicating a significant difference between observed data and the theoretical distribution.",
            "Degrees of Freedom": "Degrees of freedom = number of categories - number of constraints - 1, used to determine the reference distribution for the chi-square test.",
            "Significance Level": "Threshold for judging significance, usually set to 0.05, representing the maximum acceptable probability of a Type I error."
        }
    }
}


class ChiSquareGoodnessOfFitTestApp:
    def __init__(self, root=None):
        # 当前语言，默认为中文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def analyze_file(self):
        lang = LANGUAGES[self.current_language]
        file_path = self.file_entry.get()

        if file_path == lang["file_entry_placeholder"]:
            self.result_label.config(text=lang['file_not_found'])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=lang['file_not_found'])
            return

        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 假设第一列是观测频数，第二列是理论频数
            observed = df.iloc[:, 0]
            expected = df.iloc[:, 1]

            # 进行卡方拟合优度检验
            chi2, p = stats.chisquare(observed, expected)
            # 计算自由度
            df_val = len(observed) - 1

            # 增加显著性水平判断
            alpha = 0.05
            significance = lang['significance']['significant'] if p < alpha else lang['significance']['not_significant']

            # 整理结果
            result_df = pd.DataFrame({
                lang['docx']['statistic']: [
                    lang['statistic_labels']['chi2'],
                    lang['statistic_labels']['df'],
                    lang['statistic_labels']['p'],
                    lang['statistic_labels']['alpha'],
                    lang['statistic_labels']['result']
                ],
                lang['docx']['value']: [
                    f"{chi2:.4f}",
                    f"{df_val:d}",
                    f"{p:.6f}",
                    f"{alpha:.2f}",
                    significance
                ]
            })

            # 生成统计图表
            plt.figure(figsize=(10, 6))
            x = np.arange(len(observed))
            width = 0.35  # 柱子宽度

            plt.bar(x - width / 2, observed, width,
                    label=lang['chart']['observed'], alpha=0.7)
            plt.bar(x + width / 2, expected, width,
                    label=lang['chart']['expected'], alpha=0.7)

            # 添加统计信息到图表
            stat_text = (
                f"χ² = {chi2:.4f}, {lang['statistic_labels']['df']} = {df_val}, {lang['statistic_labels']['p']} = {p:.6f}\n"
                f"{lang['statistic_labels']['alpha']} = {alpha}, {significance}")
            plt.text(0.05, 0.95, stat_text, transform=plt.gca().transAxes,
                     verticalalignment='top', bbox=dict(boxstyle='round', facecolor='white', alpha=0.8))

            plt.xlabel(lang['chart']['x_label'])
            plt.ylabel(lang['chart']['y_label'])
            plt.title(lang['chart']['title'])
            plt.legend()
            plt.tight_layout()

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加标题
                doc.add_heading(lang['docx']['title'], 0)

                # 添加表格
                doc.add_heading(lang['docx']['analysis_results'], level=1)
                stat_table = doc.add_table(rows=1, cols=2)
                hdr_cells = stat_table.rows[0].cells
                hdr_cells[0].text = lang['docx']['statistic']
                hdr_cells[1].text = lang['docx']['value']

                # 添加数据行
                for _, row in result_df.iterrows():
                    row_cells = stat_table.add_row().cells
                    row_cells[0].text = str(row[lang['docx']['statistic']])
                    row_cells[1].text = str(row[lang['docx']['value']])

                # 保存图表到临时文件
                plot_path = "temp_plot.png"
                plt.savefig(plot_path)

                # 将图片插入到 Word 文档中
                doc.add_picture(plot_path, width=Inches(6))

                # 删除临时文件
                os.remove(plot_path)

                # 添加解释说明
                doc.add_heading(lang['docx']['explanation'], level=1)
                explanation_para = doc.add_paragraph()
                explanations = lang['explanation']
                for key, value in explanations.items():
                    explanation_para.add_run(f"• {key}: {value}\n")

                # 添加结果解读
                doc.add_heading(lang['docx']['interpretation'], level=1)
                interpretation_para = doc.add_paragraph()
                interpretations = lang['interpretation']
                for key, value in interpretations.items():
                    interpretation_para.add_run(f"• {key}: {value}\n")

                # 保存 Word 文档
                doc.save(save_path)

                result_msg = lang['analysis_success'].format(save_path)
                self.result_label.config(text=result_msg, wraplength=400)
            else:
                self.result_label.config(text=lang['no_save_path'])

        except Exception as e:
            self.result_label.config(text=lang['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        # 切换当前语言
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'
        lang = LANGUAGES[self.current_language]

        # 更新UI文本
        self.root.title(lang['title'])
        self.select_button.config(text=lang['select_button'])
        self.analyze_button.config(text=lang['analyze_button'])
        self.switch_language_label.config(text=lang['switch_language'])

        # 更新输入框提示文本
        current_entry_text = self.file_entry.get()
        # 只有当输入框是占位文本时才更新，避免覆盖用户输入
        if current_entry_text == LANGUAGES['zh']['file_entry_placeholder'] or current_entry_text == LANGUAGES['en'][
            'file_entry_placeholder']:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, lang['file_entry_placeholder'])
            self.file_entry.config(foreground='gray')

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=LANGUAGES[self.current_language]['select_button'],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=LANGUAGES[self.current_language]['analyze_button'],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=LANGUAGES[self.current_language]['switch_language'],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = ChiSquareGoodnessOfFitTestApp()
    app.run()


if __name__ == "__main__":
    run_app()
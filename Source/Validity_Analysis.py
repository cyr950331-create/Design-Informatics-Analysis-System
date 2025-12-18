import tkinter as tk
from tkinter import filedialog
import os
import pandas as pd
import numpy as np
from factor_analyzer import FactorAnalyzer
from factor_analyzer.factor_analyzer import calculate_kmo, calculate_bartlett_sphericity
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'  # 设置字体为黑体，可根据系统情况修改为其他支持中文的字体
plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

# 定义语言字典
languages = {
    "zh": {
        "title": "效度",
        "select_button_text": "选择文件",
        "file_entry_placeholder": "请输入待分析 Excel 文件的完整路径",
        "analyze_button_text": "分析文件",
        "no_file_selected": "请选择文件。",
        "file_not_exists": "文件不存在，请重新选择。",
        "analysis_error": "分析文件时出错: {}",
        "analysis_complete": "分析完成，结果已保存到 {}，相关图片已保存。",
        "no_save_path_selected": "未选择保存路径，结果未保存。",
        'open_excel_button_text': "示例数据",
        "switch_language_button_text": "中/英",
        "explanation": {
            "KMO检验值": "Kaiser-Meyer-Olkin检验用于衡量数据是否适合进行因子分析，取值范围在0-1之间，越接近1越适合。0.9以上为非常好，0.8-0.9为好，0.7-0.8为一般，0.6-0.7为差，0.6以下不适合。",
            "Bartlett球形检验": "用于检验变量之间是否存在相关性，p值小于0.05表示变量之间存在相关性，适合进行因子分析。",
            "因子数量（Kaiser准则）": "根据Kaiser准则，保留特征值大于1的因子数量。",
            "因子数量（平行分析）": "通过平行分析确定的因子数量，比较实际数据特征值与随机数据特征值的平均值。",
            "最终选择因子数量": "综合Kaiser准则和平行分析结果确定的最终因子数量。",
            "累积方差贡献率": "所有保留因子解释的总方差比例，越高说明因子对数据的解释能力越强，通常建议达到60%以上。",
            "因子载荷矩阵": "反映了每个变量与每个因子之间的相关性，绝对值越大（接近1）表示相关性越强。通常以0.5作为显著相关的临界值。",
            "样本量": "分析中实际使用的样本数量，因子分析建议样本量至少为变量数的5倍，越多越稳定。",
            "各变量样本量": "每个变量的有效观测值数量。",
            "均值±标准差": "样本数据的集中趋势和离散程度，均值反映中心位置，标准差反映数据分散程度。",
            "各项目删除后α系数": "删除对应项目后的Cronbach's α系数，若删除后α系数显著提高，说明该项目可能降低量表一致性。",
        },
        "interpretation": {
            "KMO检验值": "KMO检验值越接近1，说明变量之间的相关性越强，越适合进行因子分析。0.9以上表明非常适合，0.8-0.9为适合，0.7-0.8为基本适合，0.6-0.7为不太适合，0.6以下则不适合进行因子分析。",
            "Bartlett球形检验p值": "若Bartlett球形检验p值小于0.05，则拒绝原假设，表明变量之间存在显著相关性，数据适合进行因子分析；若p值大于等于0.05，则变量间相关性较弱，不适合因子分析。",
            "因子数量（Kaiser准则）": "根据Kaiser准则，仅保留特征值大于1的因子。该数量反映了数据中能够解释大部分变异的主要维度数量，但可能会高估实际因子数量。",
            "因子数量（平行分析）": "平行分析通过比较实际数据与随机数据的特征值确定因子数量，结果通常比Kaiser准则更保守，能有效避免高估因子数量。",
            "最终选择因子数量": "综合Kaiser准则和平行分析结果确定的因子数量，兼顾了数据特征和统计稳定性，是更可靠的因子数量选择。",
            "累积方差贡献率": "累积方差贡献率越高，说明所选因子对原始数据的解释能力越强。通常建议达到60%以上，若能达到70%-80%则表明因子解释效果更佳。",
            "因子载荷矩阵": "因子载荷的绝对值越大（接近1），说明该变量与对应因子的相关性越强。通常以0.5作为临界值，绝对值大于0.5的载荷可认为具有统计学意义。",
            "样本量": "样本量大小直接影响因子分析结果的稳定性。建议样本量至少为变量数的5倍，10倍以上更佳，足够的样本量能减少随机误差对结果的影响。",
            "各变量样本量": "各变量的有效观测值数量应保持一致且充足，若某变量缺失值过多（样本量过小），可能导致该变量在因子分析中表现不稳定。",
            "均值±标准差": "均值反映变量的平均水平，标准差反映数据的离散程度。均值±标准差结合可判断数据的分布特征，标准差越小说明数据越集中。",
            "各项目删除后α系数": "若删除某项目后α系数显著提高，说明该项目可能与其他项目一致性较低，降低了量表的内部一致性；若变化不大或降低，则表明该项目对量表一致性有积极作用。"
        }
    },
    "en": {
        "title": "Validity",
        "select_button_text": "Select File",
        "file_entry_placeholder": "Please enter the full path of the Excel file to be analyzed",
        "analyze_button_text": "Analyze File",
        "no_file_selected": "Please select a valid file path.",
        "file_not_exists": "The file does not exist. Please select again.",
        "analysis_error": "An error occurred while analyzing the file: {}",
        "analysis_complete": "Analysis completed. The results have been saved to {}, and the relevant images have been saved.",
        "no_save_path_selected": "No save path selected. The results were not saved.",
        'open_excel_button_text': "Example data",
        "switch_language_button_text": "Chinese/English",
        "explanation": {
            "KMO Test Value": "Kaiser-Meyer-Olkin test measures suitability for factor analysis, ranging from 0 to 1. Values closer to 1 are better: >0.9 excellent, 0.8-0.9 good, 0.7-0.8 acceptable, 0.6-0.7 poor, <0.6 unsuitable.",
            "Bartlett's Test of Sphericity": "Tests for correlations between variables. A p-value < 0.05 indicates significant correlations, suitable for factor analysis.",
            "Number of Factors (Kaiser Criterion)": "Number of factors retained based on Kaiser's criterion (eigenvalues > 1).",
            "Number of Factors (Parallel Analysis)": "Number of factors determined by parallel analysis, comparing eigenvalues of actual data with random data.",
            "Final Number of Factors Selected": "Final number of factors determined by combining Kaiser's criterion and parallel analysis results.",
            "Cumulative Variance Contribution": "Total variance explained by all retained factors. Higher values indicate better explanatory power, typically recommended to be >60%.",
            "Factor Loading Matrix": "Reflects correlations between variables and factors. Larger absolute values (closer to 1) indicate stronger correlations. Usually, 0.5 is used as the significance threshold.",
            "Sample Size": "Actual sample size used in analysis. Factor analysis recommends at least 5 times more samples than variables for stability.",
            "Sample Sizes per Variable": "Number of valid observations for each variable.",
            "Mean ± Standard Deviation": "Central tendency and dispersion of sample data. Mean reflects central position, standard deviation reflects data spread.",
            "Cronbach's α After Item Removal": "Cronbach's α coefficient after removing the corresponding item. If α increases significantly, the item may reduce scale consistency."
        },
        "interpretation": {
            "KMO Test Value": "The closer the KMO test value is to 1, the stronger the correlation between variables, making them more suitable for factor analysis. Values above 0.9 indicate excellent suitability, 0.8-0.9 good, 0.7-0.8 acceptable, 0.6-0.7 poor, and below 0.6 unsuitable.",
            "Bartlett's Test of Sphericity p-value": "If the p-value of Bartlett's test of sphericity is less than 0.05, the null hypothesis is rejected, indicating significant correlations between variables and suitability for factor analysis. A p-value ≥ 0.05 suggests weak correlations, making factor analysis inappropriate.",
            "Number of Factors (Kaiser Criterion)": "Based on Kaiser's criterion, only factors with eigenvalues > 1 are retained. This number reflects major dimensions explaining most variance but may overestimate the actual number of factors.",
            "Number of Factors (Parallel Analysis)": "Parallel analysis determines the number of factors by comparing eigenvalues of actual and random data. Results are typically more conservative than Kaiser's criterion, effectively avoiding overestimation.",
            "Final Number of Factors Selected": "The final number of factors, determined by combining Kaiser's criterion and parallel analysis, balances data characteristics and statistical stability, providing a more reliable selection.",
            "Cumulative Variance Contribution": "A higher cumulative variance explained indicates stronger explanatory power of selected factors for original data. It is generally recommended to reach over 60%, with 70%-80% indicating better explanatory effect.",
            "Factor Loading Matrix": "Larger absolute factor loadings (closer to 1) indicate stronger correlations between variables and corresponding factors. A threshold of 0.5 is commonly used, with loadings exceeding this considered statistically significant.",
            "Sample Size": "Sample size directly affects the stability of factor analysis results. It is recommended to have at least 5 times more samples than variables, with 10 times or more being better. Sufficient sample size reduces the impact of random errors.",
            "Sample Sizes per Variable": "Valid observations for each variable should be consistent and sufficient. Excessive missing values (small sample size) for a variable may lead to unstable performance in factor analysis.",
            "Mean ± Standard Deviation": "The mean reflects the average level of variables, while the standard deviation indicates data dispersion. Together, they describe data distribution characteristics, with smaller standard deviations indicating more concentrated data.",
            "Cronbach's α After Item Removal": "A significant increase in α coefficient after removing an item suggests poor consistency with other items, reducing scale reliability. Little change or a decrease indicates the item positively contributes to scale consistency."
        }
    }
}

class ValidityAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data14.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def validity_analysis(self, data):
        # 进行KMO检验和Bartlett球形检验
        kmo_all, kmo_model = calculate_kmo(data)
        chi_square_value, p_value = calculate_bartlett_sphericity(data)

        # 确定最佳因子数量 - 使用特征值大于1的标准
        fa_eig = FactorAnalyzer(n_factors=data.shape[1], rotation=None)
        fa_eig.fit(data)
        eigenvalues, _ = fa_eig.get_eigenvalues()
        n_factors_kaiser = sum(eigenvalues > 1)  # 特征值大于1的因子数量

        # 平行分析
        pa_n_factors = self.parallel_analysis(data, eigenvalues)  # 传入特征值参数

        # 综合两种方法确定因子数量
        n_factors = max(1, min(pa_n_factors, n_factors_kaiser))  # 取两种方法的最小值

        # 使用最大方差旋转（科研中常用）
        rotation = 'varimax' if n_factors > 1 else None
        fa = FactorAnalyzer(n_factors=n_factors, rotation=rotation)
        fa.fit(data)
        loadings = fa.loadings_

        # 计算因子方差贡献
        variance_contribution = fa.get_factor_variance()  # 包含三个数组：特征值、方差贡献率、累积方差贡献率

        return kmo_model, p_value, loadings, eigenvalues, n_factors, variance_contribution, chi_square_value, n_factors_kaiser, pa_n_factors

    def parallel_analysis(self, data, eigenvalues, n_iter=100, random_state=42):  # 接收外部传入的特征值
        np.random.seed(random_state)
        n, p = data.shape
        eigen_values = []

        for _ in range(n_iter):
            random_data = np.random.normal(size=(n, p))
            fa_random = FactorAnalyzer(n_factors=p, rotation=None)
            fa_random.fit(random_data)
            ev, _ = fa_random.get_eigenvalues()
            eigen_values.append(ev)

        mean_ev = np.mean(eigen_values, axis=0)
        return sum(eigenvalues > mean_ev)  # 使用传入的特征值进行比较

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 检查数据是否为数值类型
            numerical_df = df.select_dtypes(include=[np.number])
            if numerical_df.empty:
                raise ValueError("数据中没有数值列，无法进行效度分析。")

            # 添加数据预处理步骤
            # 处理缺失值
            missing_values = numerical_df.isnull().sum()
            if missing_values.sum() > 0:
                # 记录缺失值情况
                missing_info = "\n".join(
                    [f"{col}: {cnt}个缺失值" for col, cnt in missing_values.items() if cnt > 0])
                # 使用均值填充缺失值（科研中常用方法）
                numerical_df = numerical_df.fillna(numerical_df.mean())

            # 处理异常值（使用3σ法则）
            for col in numerical_df.columns:
                mean = numerical_df[col].mean()
                std = numerical_df[col].std()
                numerical_df = numerical_df[
                    (numerical_df[col] >= mean - 3 * std) & (numerical_df[col] <= mean + 3 * std)]

            # 检查处理后的数据是否仍然有效
            if len(numerical_df) < 5:  # 样本量过小
                raise ValueError("数据量过小，无法进行有效的因子分析。建议样本量至少为变量数的5倍以上。")

            # 进行效度分析
            kmo, bartlett_p, loadings, eigenvalues, n_factors, variance_contribution, chi_square_value, n_factors_kaiser, pa_n_factors = self.validity_analysis(
                numerical_df)

            # 计算样本量和均值
            sample_sizes = numerical_df.count()
            means = numerical_df.mean()

            # 定义多语言统计量名称和表头
            if self.current_language == "zh":
                # 中文配置
                stats_names = {
                    "kmo": "KMO检验值",
                    "bartlett": "Bartlett球形检验",
                    "kaiser_factors": "因子数量（Kaiser准则）",
                    "pa_factors": "因子数量（平行分析）",
                    "final_factors": "最终选择因子数量",
                    "cumulative_variance": "累积方差贡献率",
                    "loading_matrix": "因子载荷矩阵",
                    "sample_size": "样本量",
                    "var_sample_sizes": "各变量样本量",
                    "mean_std": "均值±标准差"
                }
                sig_text = "显著"
                not_sig_text = "不显著"
                headers = ["统计量", "统计量值", "说明/解释"]
            else:
                # 英文配置
                stats_names = {
                    "kmo": "KMO Test Value",
                    "bartlett": "Bartlett's Test of Sphericity",
                    "kaiser_factors": "Number of Factors (Kaiser Criterion)",
                    "pa_factors": "Number of Factors (Parallel Analysis)",
                    "final_factors": "Final Number of Factors Selected",
                    "cumulative_variance": "Cumulative Variance Contribution",
                    "loading_matrix": "Factor Loading Matrix",
                    "sample_size": "Sample Size",
                    "var_sample_sizes": "Sample Sizes per Variable",
                    "mean_std": "Mean±Standard Deviation"
                }
                sig_text = "significant"
                not_sig_text = "not significant"
                headers = ["Statistic", "Value", "Explanation"]

            # 整理多语言数据
            data = [
                [stats_names["kmo"], f"{kmo:.4f}",
                 languages[self.current_language]["interpretation"].get("KMO检验值", "")],
                [stats_names["bartlett"],
                 f"χ² = {chi_square_value:.4f}, df = {(numerical_df.shape[0] - 1) * numerical_df.shape[1] // 2:.0f}",
                 f"p = {bartlett_p:.6f} ({sig_text if bartlett_p < 0.05 else not_sig_text})"],
                [stats_names["kaiser_factors"], n_factors_kaiser, ""],
                [stats_names["pa_factors"], pa_n_factors, ""],
                [stats_names["final_factors"], n_factors, ""],
                [stats_names["cumulative_variance"], f"{variance_contribution[2][-1]:.2%}", ""],
                [stats_names["loading_matrix"], pd.DataFrame(loadings, index=numerical_df.columns).round(4).to_string(),
                 ""],
                [stats_names["sample_size"], len(numerical_df),
                 f"{'变量数' if self.current_language == 'zh' else 'Number of variables'}: {numerical_df.shape[1]}, "
                 f"{'样本量/变量数' if self.current_language == 'zh' else 'Sample size/variable ratio'}: {len(numerical_df) / numerical_df.shape[1]:.1f}"],
                [stats_names["var_sample_sizes"], sample_sizes.to_string(), ""],
                [stats_names["mean_std"],
                 {k: f"{round(v, 4)}±{round(numerical_df[k].std(), 4)}" for k, v in means.to_dict().items()}, ""]
            ]

            df_result = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            explanations = languages[self.current_language]['explanation']
            interpretations = languages[self.current_language]['interpretation']

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加标题
                doc.add_heading(languages[self.current_language]["title"], 0)

                # 添加统计结果表格
                table = doc.add_table(rows=1, cols=len(headers))
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header

                for _, row in df_result.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        if isinstance(value, dict):
                            value_str = '\n'.join([f"{k}: {v}" for k, v in value.items()])
                            row_cells[i].text = value_str
                        else:
                            row_cells[i].text = str(value)

                # 添加解释说明
                doc.add_heading("解释说明" if self.current_language == 'zh' else "Explanation", 1)
                for key, value in explanations.items():
                    doc.add_paragraph(f"{key}: {value}")

                # 添加分析结果解读
                doc.add_heading("结果解读" if self.current_language == 'zh' else "Interpretation", 1)
                for key, value in interpretations.items():
                    doc.add_paragraph(f"{key}: {value}")

                # 生成图片（均值柱状图）
                fig, ax = plt.subplots()
                means.plot(kind='bar', ax=ax)
                ax.set_title('变量均值柱状图' if self.current_language == 'zh' else 'Bar Chart of Variable Means')
                ax.set_xlabel('变量' if self.current_language == 'zh' else 'Variables')
                ax.set_ylabel('均值' if self.current_language == 'zh' else 'Mean')
                ax.tick_params(axis='x', rotation=0)

                # 保存图片
                img_path = os.path.splitext(save_path)[0] + '.png'
                plt.savefig(img_path)
                plt.close()

                # 添加碎石图（用于确定因子数量）
                fig, ax = plt.subplots()
                ax.plot(range(1, len(eigenvalues) + 1), eigenvalues, 'o-')
                ax.axhline(y=1, color='r', linestyle='--')  # 特征值=1的参考线
                ax.set_title('碎石图' if self.current_language == 'zh' else 'Scree Plot')
                ax.set_xlabel('因子编号' if self.current_language == 'zh' else 'Factor Number')
                ax.set_ylabel('特征值' if self.current_language == 'zh' else 'Eigenvalue')
                plt.tight_layout()
                scree_path = os.path.splitext(save_path)[0] + '_scree.png'
                plt.savefig(scree_path, dpi=300)
                plt.close()

                # 添加因子载荷热图
                fig, ax = plt.subplots(figsize=(10, 8))
                cax = ax.imshow(loadings, cmap='coolwarm', vmin=-1, vmax=1)
                plt.colorbar(cax)
                ax.set_xticks(range(loadings.shape[1]))
                ax.set_yticks(range(loadings.shape[0]))
                ax.set_xticklabels([f'因子{i + 1}' for i in range(loadings.shape[1])]
                                   if self.current_language == 'zh' else [f'Factor {i + 1}' for i in
                                                                          range(loadings.shape[1])])
                ax.set_yticklabels(numerical_df.columns)
                ax.set_title('因子载荷矩阵热图' if self.current_language == 'zh' else 'Factor Loading Heatmap')
                plt.tight_layout()
                loading_path = os.path.splitext(save_path)[0] + '_loadings.png'
                plt.savefig(loading_path, dpi=300)
                plt.close()

                # 将图片插入到 Word 文档中
                doc.add_heading("变量均值柱状图" if self.current_language == 'zh' else 'Bar Chart of Variable Means', 1)
                doc.add_picture(img_path)

                # 保存 Word 文档
                doc.save(save_path)

                self.result_label.config(text=languages[self.current_language]["analysis_complete"].format(save_path),
                                         wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]["no_save_path_selected"])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]["analysis_error"].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])
        self.open_excel_label.config(text=languages[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=languages[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame,
                                               text=languages[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = ValidityAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()
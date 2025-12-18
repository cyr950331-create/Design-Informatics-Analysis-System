import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from tkinter import messagebox
import os
import numpy as np
import pandas as pd
from tkinter import filedialog, simpledialog
import tkinter as tk
import matplotlib.pyplot as plt
import pathlib
from docx import Document
from docx.shared import Inches


# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典，所有键使用英文
LANGUAGES = {
    'zh': {
        'title': "灰色预测模型",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "请选择文件。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "中/英",
        'open_excel_button_text': "示例数据",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'explanation': {
            'original_data': "输入的待分析数据",
            'accumulated_sequence': "对原始数据进行一次累加生成得到的序列",
            'predicted_values': "通过灰色预测模型得到的预测值",
            'prediction_chart': "展示原始数据和预测值的折线图",
            'accuracy_indices': "评估预测模型可靠性的指标"
        },
        'interpretation': {
            'original_data': "作为分析的基础数据",
            'accumulated_sequence': "用于构建灰色预测模型",
            'predicted_values': "反映未来趋势的预测结果",
            'prediction_chart': "直观展示原始数据和预测值的变化趋势",
            'posterior_error_ratio': "越小越好，<0.35为优，<0.5为合格，<0.65为勉强，>0.65为不合格",
            'small_error_probability': "越大越好，>0.95为优，>0.8为合格，>0.7为勉强，<0.7为不合格"
        },
        'data_validity': "数据有效性提示：GM(1,1)模型适合单调增长/衰减序列，当前数据标准差系数为{:.2f}，建议值<0.3",
        'n_pred_prompt': "请输入预测步数（正整数）：",
        'invalid_n_pred': "无效的预测步数，请输入正整数",
        'data_error': "数据格式错误：请确保Excel中包含单列/单行的数值型数据",
        'data_length_error': "数据量不足：至少需要4个数据点才能进行预测",
        'analysis_data': "分析数据",
        'explanation_title': "解释说明",
        'interpretation_title': "结果解读",
        'statistics': "统计量",
        'statistic_value': "统计量值",
        'explanation_label': "解释说明",
        'interpretation_label': "结果解读",
        'original_data_label': "原始数据",
        'predicted_values_label': "预测值",
        'time_step': "时间步",
        'value': "值"
    },
    'en': {
        'title': "Gray Prediction Model",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Chinese/English",
        'open_excel_button_text': "Example data",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'explanation': {
            'original_data': "The input data to be analyzed",
            'accumulated_sequence': "The sequence obtained by accumulating the original data once",
            'predicted_values': "The predicted values obtained through the gray prediction model",
            'prediction_chart': "A line chart showing the original data and predicted values",
            'accuracy_indices': "Indicators for evaluating the reliability of the prediction model"
        },
        'interpretation': {
            'original_data': "As the basic data for analysis",
            'accumulated_sequence': "Used to build the gray prediction model",
            'predicted_values': "The predicted results reflecting future trends",
            'prediction_chart': "Visually display the changing trends of the original data and predicted values",
            'posterior_error_ratio': "Smaller is better, <0.35 is excellent, <0.5 is qualified, <0.65 is marginal, >0.65 is unqualified",
            'small_error_probability': "Larger is better, >0.95 is excellent, >0.8 is qualified, >0.7 is marginal, <0.7 is unqualified"
        },
        'data_validity': "Data validity prompt: GM(1,1) is suitable for monotonic sequences. Current data std coefficient: {:.2f}, recommended <0.3",
        'n_pred_prompt': "Please enter the number of prediction steps (positive integer):",
        'invalid_n_pred': "Invalid number of steps. Please enter a positive integer",
        'data_error': "Data format error: Ensure Excel contains single column/row of numerical data",
        'data_length_error': "Insufficient data: At least 4 data points are required for prediction",
        'analysis_data': "Analysis Data",
        'explanation_title': "Explanation",
        'interpretation_title': "Interpretation",
        'statistics': "Statistics",
        'statistic_value': "Statistic Value",
        'explanation_label': "Explanation",
        'interpretation_label': "Interpretation",
        'original_data_label': "Original Data",
        'predicted_values_label': "Predicted Values",
        'time_step': "Time Step",
        'value': "Value"
    }
}


class GrayPredictionModelAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为中文
        self.current_language = "en"

        # 如果没有提供 root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data53.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        """当用户点击输入框时，清除提示信息"""
        if self.file_entry.get() == LANGUAGES[self.current_language]['file_entry_placeholder']:
            self.file_entry.delete(0, tk.END)
            self.file_entry.configure(style="TEntry")

    def on_focusout(self, event):
        """当用户离开输入框时，如果没有输入内容，恢复提示信息"""
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
            self.file_entry.configure(style="Gray.TEntry")

    def gm11(self, x0, n_pred):
        """
        GM(1,1) 灰色预测模型
        :param x0: 原始数据序列
        :param n_pred: 预测步数
        :return: 预测值序列、后验差比、小误差概率
        """
        x1 = np.cumsum(x0)
        z1 = (x1[:-1] + x1[1:]) / 2
        B = np.vstack([-z1, np.ones_like(z1)]).T
        Y = x0[1:].reshape(-1, 1)
        # 最小二乘法求解参数
        a, b = np.linalg.lstsq(B, Y, rcond=None)[0].flatten()

        # 预测累加序列计算（k从0开始，长度为len(x0)+n_pred）
        x1_pred = [(x0[0] - b / a) * np.exp(-a * k) + b / a for k in range(len(x0) + n_pred)]

        # 还原为原始序列
        x0_pred = [x1_pred[0]] + [x1_pred[k] - x1_pred[k - 1] for k in range(1, len(x1_pred))]

        # 计算精度检验指标
        x0_fitted = x0_pred[:len(x0)]  # 拟合值
        residual = x0 - x0_fitted  # 残差
        s1 = np.std(x0, ddof=1)  # 原始数据标准差
        s2 = np.std(residual, ddof=1)  # 残差标准差
        c = s2 / s1  # 后验差比

        # 小误差概率
        mean_residual = np.mean(np.abs(residual))
        p = np.sum(np.abs(residual - mean_residual) < 0.6745 * s1) / len(residual)

        return x0_pred, c, p

    def center_window(self, window):
        """使弹出窗口居中显示"""
        window.update_idletasks()
        width = window.winfo_width()
        height = window.winfo_height()
        x = (self.root.winfo_width() // 2) + self.root.winfo_x() - (width // 2)
        y = (self.root.winfo_height() // 2) + self.root.winfo_y() - (height // 2)
        window.geometry(f"{width}x{height}+{x}+{y}")

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]['file_entry_placeholder']:
            file_path = ""
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 创建自定义输入对话框确保居中
            dialog = tk.Toplevel(self.root)
            dialog.title(LANGUAGES[self.current_language]['title'])
            dialog.transient(self.root)  # 设置为主窗口的子窗口

            ttk.Label(dialog, text=LANGUAGES[self.current_language]['n_pred_prompt']).pack(padx=20, pady=10)
            entry = ttk.Entry(dialog, width=20)
            entry.pack(padx=20, pady=5)
            entry.focus_set()

            result = [None]  # 用列表存储结果，以便在内部函数中修改

            def on_ok():
                result[0] = entry.get()
                dialog.destroy()

            def on_cancel():
                dialog.destroy()

            button_frame = ttk.Frame(dialog)
            button_frame.pack(pady=10)
            ttk.Button(button_frame, text="OK", command=on_ok).pack(side=tk.LEFT, padx=5)
            ttk.Button(button_frame, text="Cancel", command=on_cancel).pack(side=tk.LEFT, padx=5)

            self.center_window(dialog)  # 居中显示
            self.root.wait_window(dialog)  # 等待对话框关闭

            n_pred_str = result[0]
            if n_pred_str is None or not n_pred_str.isdigit() or int(n_pred_str) <= 0:
                self.result_label.config(text=LANGUAGES[self.current_language]['invalid_n_pred'])
                return
            n_pred = int(n_pred_str)

            # 读取数据（支持Excel和CSV）
            if file_path.endswith(('.xlsx', '.xls')):
                df = pd.read_excel(file_path, header=None)
            elif file_path.endswith('.csv'):
                df = pd.read_csv(file_path, header=None)
            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['data_error'])
                return

            # 提取一维数值数据并处理异常值
            data = df.values.flatten()
            if len(data) < 4:  # GM(1,1)至少需要4个数据点
                self.result_label.config(text=LANGUAGES[self.current_language]['data_length_error'])
                return

            # 过滤非数值数据
            try:
                data = data.astype(float)
                # 移除NaN值
                data = data[~np.isnan(data)]
                if len(data) < 4:
                    self.result_label.config(text=LANGUAGES[self.current_language]['data_length_error'])
                    return
            except (ValueError, TypeError):
                self.result_label.config(text=LANGUAGES[self.current_language]['data_error'])
                return

            # 数据适用性检验（计算标准差系数）
            std_coef = np.std(data) / np.mean(data) if np.mean(data) != 0 else float('inf')
            validity_msg = LANGUAGES[self.current_language]['data_validity'].format(std_coef)

            # 进行灰色预测分析
            pred_values, c, p = self.gm11(data, n_pred)

            # 整理数据，使用英文键
            data_list = [
                [LANGUAGES[self.current_language]['original_data_label'], [float(x) for x in data.tolist()]],
                [LANGUAGES[self.current_language]['explanation']['accumulated_sequence'], [float(x) for x in np.cumsum(data).tolist()]],
                [LANGUAGES[self.current_language]['predicted_values_label'], [float(x) for x in pred_values]],
                [LANGUAGES[self.current_language]['interpretation']['posterior_error_ratio'], f"{c:.4f}"],
                [LANGUAGES[self.current_language]['interpretation']['small_error_probability'], f"{p:.4f}"]
            ]
            headers = [LANGUAGES[self.current_language]['statistics'], LANGUAGES[self.current_language]['statistic_value']]
            df = pd.DataFrame(data_list, columns=headers)

            # 添加解释说明
            explanations = LANGUAGES[self.current_language]['explanation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(
                columns=["original_data", "accumulated_sequence", "predicted_values", "prediction_chart", "accuracy_indices"])
            explanation_df.insert(0, "statistics_explanation", LANGUAGES[self.current_language]['explanation_label'])

            # 添加分析结果解读
            interpretations = LANGUAGES[self.current_language]['interpretation']
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(
                columns=["original_data", "accumulated_sequence", "predicted_values", "prediction_chart", "posterior_error_ratio", "small_error_probability"])
            interpretation_df.insert(0, "statistics_interpretation", LANGUAGES[self.current_language]['interpretation_label'])

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加表格数据
                doc.add_heading(LANGUAGES[self.current_language]['analysis_data'], level=1)
                table = doc.add_table(rows=1, cols=len(headers))
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                for index, row in df.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)

                # 添加解释说明
                doc.add_heading(LANGUAGES[self.current_language]['explanation_title'], level=1)
                explain_paragraph = doc.add_paragraph()
                for key, value in explanations.items():
                    # 添加项目符号
                    run = explain_paragraph.add_run(f"• {LANGUAGES[self.current_language]['explanation'].get(key, key)}: {value}\n")
                    run.font.name = 'SimHei'  # 确保中文显示正常

                # 添加分析结果解读
                doc.add_heading(LANGUAGES[self.current_language]['interpretation_title'], level=1)
                interpret_paragraph = doc.add_paragraph()
                for key, value in interpretations.items():
                    # 添加项目符号
                    run = interpret_paragraph.add_run(f"• {LANGUAGES[self.current_language]['interpretation'].get(key, key)}: {value}\n")
                    run.font.name = 'SimHei'  # 确保中文显示正常

                # 生成预测结果折线图
                plt.figure()
                plt.plot(range(len(data)), data, label=LANGUAGES[self.current_language]['original_data_label'])
                plt.plot(range(len(pred_values)), pred_values,
                         label=LANGUAGES[self.current_language]['predicted_values_label'],
                         linestyle='--')
                plt.title(LANGUAGES[self.current_language]['explanation']['prediction_chart'])
                plt.xlabel(LANGUAGES[self.current_language]['time_step'])
                plt.ylabel(LANGUAGES[self.current_language]['value'])
                plt.legend()

                # 保存图片
                img_path = os.path.splitext(save_path)[0] + '_prediction_chart.png'
                plt.savefig(img_path)
                plt.close()

                # 将图片插入到 Word 文档中
                doc.add_heading(LANGUAGES[self.current_language]['explanation']['prediction_chart'], level=1)
                doc.add_picture(img_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(save_path)
                self.result_label.config(text=result_msg, wraplength=400)

            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'
        self.root.title(LANGUAGES[self.current_language]['title'])
        self.select_button.config(text=LANGUAGES[self.current_language]['select_button'])
        self.analyze_button.config(text=LANGUAGES[self.current_language]['analyze_button'])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]['switch_language'])
        self.open_excel_label.config(text=LANGUAGES[self.current_language]["open_excel_button_text"])

        # 切换语言时更新提示信息
        current_text = self.file_entry.get()
        placeholder = LANGUAGES[self.current_language]['file_entry_placeholder']
        if current_text == LANGUAGES['zh' if self.current_language == 'en' else 'en']['file_entry_placeholder']:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, placeholder)
            self.file_entry.configure(style="Gray.TEntry")

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建主容器，用于将核心元素居中
        main_container = ttk.Frame(self.root)
        main_container.pack(expand=True)  # 让容器扩展以填充空间

        # 创建自定义样式
        style = ttk.Style()
        style.configure("Gray.TEntry", foreground="gray")

        # 创建文件选择按钮（放在主容器中）
        self.select_button = ttk.Button(main_container, text=LANGUAGES[self.current_language]['select_button'],
                                        command=self.select_file,
                                        bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框（放在主容器中）
        self.file_entry = ttk.Entry(main_container, width=50, style="Gray.TEntry")
        self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
        self.file_entry.pack(pady=5)
        self.file_entry.bind("<FocusIn>", self.on_entry_click)
        self.file_entry.bind("<FocusOut>", self.on_focusout)

        # 创建分析按钮（放在主容器中）
        self.analyze_button = ttk.Button(main_container, text=LANGUAGES[self.current_language]['analyze_button'],
                                         command=self.analyze_file,
                                         bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            main_container,
            text=LANGUAGES[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建语言切换标签（放在主容器中）
        self.switch_language_label = ttk.Label(main_container, text=LANGUAGES[self.current_language]['switch_language'],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.pack(pady=10)
        self.switch_language_label.bind("<Button-1>", self.switch_language)

        # 创建结果显示标签（放在主窗口中，保持在底部）
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10, padx=10, side=tk.BOTTOM)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = GrayPredictionModelAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()
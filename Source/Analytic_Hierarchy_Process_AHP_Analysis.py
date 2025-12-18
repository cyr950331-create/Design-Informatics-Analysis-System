import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import warnings

# 忽略python-docx的样式警告
warnings.filterwarnings('ignore', message='style lookup by style_id is deprecated. Use style name as key instead.')

# 设置 matplotlib 支持中文和英文显示
plt.rcParams['font.family'] = ['SimHei', 'WenQuanYi Micro Hei', 'Heiti TC', 'Arial Unicode MS']
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典 - 包含所有需要翻译的文本
languages = {
    'zh': {
        'title': "层次分析法",
        'select_button_text': "选择文件",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'analyze_button_text': "分析文件",
        'no_file_selected': "请选择文件。",
        'file_not_exists': "文件不存在，请重新选择。",
        'analysis_error': "分析文件时出错: {}",
        'analysis_complete': "分析完成，结果已保存到 {}\n",
        'no_save_path_selected': "未选择保存路径，结果未保存。",
        'switch_language_button_text': "中/英",
        'open_excel_button_text': "示例数据",
        'explanation': {
            "特征向量": "反映各因素相对重要性的向量",
            "一致性指标 CI": "衡量判断矩阵一致性的指标",
            "随机一致性指标 RI": "根据矩阵阶数确定的随机一致性指标",
            "一致性比率 CR": "CI 与 RI 的比值，判断矩阵是否具有满意一致性"
        },
        'interpretation': {
            "特征向量": "特征向量值越大，对应因素越重要",
            "一致性指标 CI": "CI 值越小，矩阵一致性越好",
            "随机一致性指标 RI": "不同阶数矩阵有对应标准值",
            "一致性比率 CR": "CR < 0.1 时，矩阵具有满意一致性，结果可信"
        },
        'non_square_matrix': "判断矩阵必须是方阵（行数=列数）",
        'non_numeric_error': "矩阵中包含非数值数据，请检查输入",
        'negative_value_error': "矩阵中包含负数，判断矩阵元素必须为正数",
        'non_reciprocal_error': "判断矩阵不是正互反矩阵，A[i][j] 应等于 1/A[j][i]",
        'high_order_warning': "警告：矩阵阶数为 {}，超过推荐的 9 阶，结果可能不可靠",
        'data_read_error': "无法读取有效的方阵矩阵数据，请检查文件格式",
        'analysis_report': "分析报告",
        'matrix_info': "判断矩阵信息",
        'eigenvector_analysis': "特征向量分析",
        'consistency_analysis': "一致性分析",
        'consistency_result': "一致性检验结果",
        'consistent': "通过 (CR < 0.1) - 结果可信",
        'not_consistent': "未通过 (CR ≥ 0.1) - 结果需谨慎使用",
        'factor_importance': "因素重要性排序",
        'matrix_diagram': "判断矩阵图示",
        'weight_distribution': "权重分布",
        'consistency_trend': "一致性指标趋势",
        'pairwise_comparison': "两两比较强度",
        'comparison_strength': "比较强度",
        'statistics': "统计量",
        'statistic_value': "统计量值",
        'result_explanation': "结果说明",
        'result_interpretation': "结果解释",
        'matrix_order': "矩阵阶数",
        'factor_count': "因素数量",
        'visual_analysis': "可视化分析",
        'max_eigenvalue': "最大特征值"
    },
    'en': {
        'title': "Analytic Hierarchy Process",
        'select_button_text': "Select File",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'analyze_button_text': "Analyze File",
        'no_file_selected': "Please select a valid file path.",
        'file_not_exists': "The file does not exist. Please select again.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'analysis_complete': "Analysis completed. The results have been saved to {}\n",
        'no_save_path_selected': "No save path selected. The results were not saved.",
        'switch_language_button_text': "Chinese/English",
        'open_excel_button_text': "Example data",
        'explanation': {
            "Eigenvector": "A vector reflecting the relative importance of each factor",
            "Consistency Index (CI)": "An indicator to measure the consistency of the judgment matrix",
            "Random Consistency Index (RI)": "A random consistency indicator determined by the order of the matrix",
            "Consistency Ratio (CR)": "The ratio of CI to RI to determine if the matrix has satisfactory consistency"
        },
        "interpretation": {
            "Eigenvector": "The larger the value in the eigenvector, the more important the corresponding factor",
            "Consistency Index (CI)": "The smaller the CI value, the better the consistency of the matrix",
            "Random Consistency Index (RI)": "There are corresponding standard values for matrices of different orders",
            "Consistency Ratio (CR)": "When CR < 0.1, the matrix has satisfactory consistency and the results are reliable"
        },
        'non_square_matrix': "Judgment matrix must be square (rows = columns)",
        'non_numeric_error': "Matrix contains non-numeric data, please check input",
        'negative_value_error': "Matrix contains negative values, judgment matrix elements must be positive",
        'non_reciprocal_error': "Judgment matrix is not positive reciprocal, A[i][j] should equal 1/A[j][i]",
        'high_order_warning': "Warning: Matrix order is {}, exceeding recommended 9 orders, results may be unreliable",
        'data_read_error': "Cannot read valid square matrix data, please check file format",
        'analysis_report': "Analysis Report",
        'matrix_info': "Judgment Matrix Information",
        'eigenvector_analysis': "Eigenvector Analysis",
        'consistency_analysis': "Consistency Analysis",
        'consistency_result': "Consistency Test Result",
        'consistent': "Passed (CR < 0.1) - Results are reliable",
        'not_consistent': "Not Passed (CR ≥ 0.1) - Results should be used with caution",
        'factor_importance': "Factor Importance Ranking",
        'matrix_diagram': "Judgment Matrix Diagram",
        'weight_distribution': "Weight Distribution",
        'consistency_trend': "Consistency Index Trend",
        'pairwise_comparison': "Pairwise Comparison Intensity",
        'comparison_strength': "Comparison Strength",
        'statistics': "Statistics",
        'statistic_value': "Statistic Value",
        'result_explanation': "Result Explanation",
        'result_interpretation': "Interpretation",
        'matrix_order': "Matrix Order",
        'factor_count': "Number of Factors",
        'visual_analysis': "Visual Analysis",
        'max_eigenvalue': "Maximum Eigenvalue"
    }
}

# 随机一致性指标 RI 表
RI_TABLE = {
    1: 0, 2: 0, 3: 0.58, 4: 0.90, 5: 1.12, 6: 1.24, 7: 1.32,
    8: 1.41, 9: 1.45, 10: 1.49, 11: 1.51, 12: 1.54,
    13: 1.56, 14: 1.58, 15: 1.59, 16: 1.60, 17: 1.61, 18: 1.62
}


class AnalyticHierarchyProcessAHPApp:
    def __init__(self, root=None):
        self.current_language = "en"
        self.headers = None  # 存储因素名称
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])
        self.create_ui()

    def validate_matrix(self, data):
        """验证判断矩阵的有效性"""
        if data.shape[0] != data.shape[1]:
            raise ValueError(languages[self.current_language]['non_square_matrix'])
        n = data.shape[0]
        if not np.issubdtype(data.dtype, np.number):
            raise ValueError(languages[self.current_language]['non_numeric_error'])
        if np.any(data <= 0):
            raise ValueError(languages[self.current_language]['negative_value_error'])
        for i in range(n):
            for j in range(n):
                if i != j and not np.isclose(data[i, j], 1 / data[j, i], atol=1e-4):
                    raise ValueError(languages[self.current_language]['non_reciprocal_error'])
        if n > 9:
            Messagebox.show_warning(
                message=languages[self.current_language]['high_order_warning'].format(n)
            )
        return True

    def calculate_eigenvector_sum_method(self, data):
        """使用几何平均法计算特征向量"""
        n = data.shape[0]
        geom_mean = np.prod(data, axis=1) ** (1 / n)
        eigenvector = geom_mean / np.sum(geom_mean)
        return eigenvector

    def ahp_analysis(self, data):
        """AHP分析核心逻辑，返回更详细的分析结果"""
        self.validate_matrix(data)
        n = data.shape[0]

        # 计算特征向量（权重）
        eigenvector = self.calculate_eigenvector_sum_method(data)

        # 计算最大特征值
        weighted_sum = np.dot(data, eigenvector)
        max_eigenvalue = np.dot(eigenvector, weighted_sum) / np.dot(eigenvector, eigenvector)

        # 计算一致性指标
        CI = (max_eigenvalue - n) / (n - 1) if n > 1 else 0
        RI = RI_TABLE.get(n, RI_TABLE[max(RI_TABLE.keys())])
        CR = CI / RI if (RI != 0 and n > 1) else 0

        # 计算排序结果
        sorted_indices = np.argsort(eigenvector)[::-1]  # 降序排序
        sorted_weights = eigenvector[sorted_indices]

        return {
            'eigenvector': eigenvector,
            'max_eigenvalue': max_eigenvalue,
            'CI': CI,
            'RI': RI,
            'CR': CR,
            'sorted_indices': sorted_indices,
            'sorted_weights': sorted_weights,
            'consistent': CR < 0.1 if n > 1 else True
        }

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data46.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def read_matrix_data(self, file_path):
        """读取矩阵数据并提取表头"""
        try:
            # 读取全部数据
            df = pd.read_excel(file_path, header=None)
            print(f"原始数据形状：{df.shape}")

            # 提取表头（第一行和第一列）
            self.headers = df.iloc[1:, 0].astype(str).str.strip().tolist()  # 行名作为因素名称
            col_headers = df.iloc[0, 1:].astype(str).str.strip().tolist()

            # 确保表头一致
            if self.headers != col_headers and len(self.headers) > 0 and len(col_headers) > 0:
                print(f"警告：行表头与列表头不一致 - 行: {self.headers}, 列: {col_headers}")

            # 提取数值区域
            numeric_df = df.iloc[1:, 1:].copy()
            print(f"数值区域形状：{numeric_df.shape}")

            if numeric_df.shape[0] != numeric_df.shape[1]:
                raise ValueError(
                    f"{languages[self.current_language]['non_square_matrix']}：{numeric_df.shape[0]}行，{numeric_df.shape[1]}列")

            # 转换为数值类型
            for col in numeric_df.columns:
                numeric_df[col] = pd.to_numeric(numeric_df[col], errors='coerce')

            # 清除含空值的行和列
            numeric_df = numeric_df.dropna(axis=0, how='any')
            numeric_df = numeric_df.dropna(axis=1, how='any')
            print(f"清洗后数值区域形状：{numeric_df.shape}")

            if numeric_df.shape[0] != numeric_df.shape[1]:
                raise ValueError(
                    f"{languages[self.current_language]['non_square_matrix']}：{numeric_df.shape[0]}行，{numeric_df.shape[1]}列")

            # 如果没有表头信息，创建默认表头
            if not self.headers or len(self.headers) != numeric_df.shape[0]:
                self.headers = [f"{languages[self.current_language]['factor_importance'].split()[0]}{i + 1}"
                                for i in range(numeric_df.shape[0])]

            return numeric_df.values.astype(float)

        except Exception as e:
            raise ValueError(f"{languages[self.current_language]['data_read_error']}：{str(e)}")

    def generate_visualizations(self, data, results, img_dir):
        """生成多种可视化图表，所有文字随语言变化"""
        n = data.shape[0]
        eigenvector = results['eigenvector']
        factor_names = self.headers

        # 创建保存图表的目录
        os.makedirs(img_dir, exist_ok=True)

        # 权重分布柱状图
        plt.figure(figsize=(10, 6))
        bars = plt.bar(factor_names, eigenvector)
        plt.title(languages[self.current_language]['weight_distribution'])
        plt.xticks(rotation=0, ha='right')
        plt.tight_layout()
        for bar in bars:
            height = bar.get_height()
            plt.text(bar.get_x() + bar.get_width() / 2., height + 0.01,
                     f'{height:.4f}', ha='center', va='bottom')
        weight_img = os.path.join(img_dir, 'weight_distribution.png')
        plt.savefig(weight_img)
        plt.close()

        # 判断矩阵热力图
        plt.figure(figsize=(10, 8))
        plt.imshow(data, cmap='YlOrRd')
        plt.colorbar(label=languages[self.current_language]['comparison_strength'])
        plt.xticks(range(n), factor_names, rotation=0, ha='right')
        plt.yticks(range(n), factor_names)
        plt.title(languages[self.current_language]['matrix_diagram'])
        # 添加数值标签
        for i in range(n):
            for j in range(n):
                plt.text(j, i, f'{data[i, j]:.2f}', ha='center', va='center',
                         color='black' if data[i, j] < np.max(data) / 2 else 'white')
        plt.tight_layout()
        heatmap_img = os.path.join(img_dir, 'matrix_heatmap.png')
        plt.savefig(heatmap_img)
        plt.close()

        # 因素重要性排序图
        sorted_indices = results['sorted_indices']
        sorted_names = [factor_names[i] for i in sorted_indices]
        sorted_weights = results['sorted_weights']

        plt.figure(figsize=(10, 6))
        plt.barh(sorted_names, sorted_weights, color='skyblue')
        plt.title(languages[self.current_language]['factor_importance'])
        for i, v in enumerate(sorted_weights):
            plt.text(v + 0.01, i, f'{v:.4f}', va='center')
        plt.tight_layout()
        ranking_img = os.path.join(img_dir, 'factor_ranking.png')
        plt.savefig(ranking_img)
        plt.close()

        # 一致性指标图
        if n > 1:
            plt.figure(figsize=(8, 5))
            plt.bar(
                # 从当前语言的explanation中获取后3个键（CI、RI、CR）
                list(languages[self.current_language]['explanation'].keys())[1:4],
                [results['CI'], results['RI'], results['CR']],
                color=['#4CAF50', '#2196F3', '#ff9800']
            )
            plt.axhline(y=0.1, color='r', linestyle='--',
                label=f"{list(languages[self.current_language]['explanation'].keys())[3]} {languages[self.current_language]['result_explanation']} (0.1)"
            )
            plt.title(languages[self.current_language]['consistency_trend'])
            plt.legend()
            for i, v in enumerate([results['CI'], results['RI'], results['CR']]):
                plt.text(i, v + 0.02, f'{v:.4f}', ha='center')
            plt.tight_layout()
            consistency_img = os.path.join(img_dir, 'consistency_indices.png')
            plt.savefig(consistency_img)
            plt.close()
        else:
            consistency_img = None

        return {
            'weight': weight_img,
            'heatmap': heatmap_img,
            'ranking': ranking_img,
            'consistency': consistency_img
        }

    def analyze_file(self):
        """增强的分析文件逻辑，所有分析结果文字随语言变化"""
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 读取数据
            data = self.read_matrix_data(file_path)
            results = self.ahp_analysis(data)

            # 获取当前语言的解释键列表（动态适配中英文）
            expl_keys = list(languages[self.current_language]['explanation'].keys())

            result_data = [
                [expl_keys[0],  # 对应：中文"特征向量" / 英文"Eigenvector"
                 [f"{v:.4f}" for v in results['eigenvector']], ""],
                [languages[self.current_language]['max_eigenvalue'],  # 最大特征值/Maximum Eigenvalue
                 round(results['max_eigenvalue'], 4), ""],
                [expl_keys[1],  # 对应：中文"一致性指标 CI" / 英文"Consistency Index (CI)"
                 round(results['CI'], 4), ""],
                [expl_keys[2],  # 对应：中文"随机一致性指标 RI" / 英文"Random Consistency Index (RI)"
                 results['RI'], ""],
                [expl_keys[3],  # 对应：中文"一致性比率 CR" / 英文"Consistency Ratio (CR)"
                 round(results['CR'], 4),
                 languages[self.current_language]['consistent'] if results['consistent']
                 else languages[self.current_language]['not_consistent']]
            ]
            headers = [
                languages[self.current_language]['statistics'],
                languages[self.current_language]['statistic_value'],
                languages[self.current_language]['result_explanation']
            ]
            df = pd.DataFrame(result_data, columns=headers)

            # 获取保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                                     filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建临时目录保存图片
                img_dir = os.path.splitext(save_path)[0] + "_images"
                os.makedirs(img_dir, exist_ok=True)

                # 生成可视化图表
                images = self.generate_visualizations(data, results, img_dir)

                # 创建Word文档
                doc = Document()
                doc.add_heading(languages[self.current_language]['analysis_report'], 0)

                # 添加基本信息
                doc.add_heading(languages[self.current_language]['matrix_info'], level=1)
                p = doc.add_paragraph(
                    f"{languages[self.current_language]['matrix_order']}: {data.shape[0]}x{data.shape[1]}")
                p = doc.add_paragraph(f"{languages[self.current_language]['factor_count']}: {len(self.headers)}")

                # 添加分析结果表格
                doc.add_heading(languages[self.current_language]['eigenvector_analysis'], level=1)
                for item in result_data:
                    stats_name, stats_value, explanation = item
                    if isinstance(stats_value, list):
                        values_text = ", ".join(stats_value)
                    else:
                        values_text = str(stats_value)
                    p = doc.add_paragraph(f"• {stats_name}: {values_text}", style='ListBullet')
                    if explanation:
                        p = doc.add_paragraph(f"  - {explanation}", style='ListBullet2')

                # 添加解释说明
                doc.add_heading(languages[self.current_language]['result_interpretation'], level=1)
                for key, value in languages[self.current_language]['explanation'].items():
                    p = doc.add_paragraph(f"• {key}: {value}", style='ListBullet')

                # 添加一致性分析
                doc.add_heading(languages[self.current_language]['consistency_analysis'], level=1)
                consistency_text = (
                        f"{list(languages[self.current_language]['explanation'].keys())[3]} = {results['CR']:.4f} → " +
                        (languages[self.current_language]['consistent'] if results['consistent']
                         else languages[self.current_language]['not_consistent'])
                )
                p = doc.add_paragraph(f"{languages[self.current_language]['consistency_result']}: " + consistency_text)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # 添加图表
                doc.add_heading(languages[self.current_language]['visual_analysis'], level=1)

                doc.add_paragraph(languages[self.current_language]['weight_distribution'])
                doc.add_picture(images['weight'], width=Inches(6))

                doc.add_paragraph(languages[self.current_language]['matrix_diagram'])
                doc.add_picture(images['heatmap'], width=Inches(6))

                doc.add_paragraph(languages[self.current_language]['factor_importance'])
                doc.add_picture(images['ranking'], width=Inches(6))

                if images['consistency']:
                    doc.add_paragraph(languages[self.current_language]['consistency_trend'])
                    doc.add_picture(images['consistency'], width=Inches(6))

                # 保存文档
                doc.save(save_path)
                self.result_label.config(text=languages[self.current_language]['analysis_complete'].format(save_path))
            else:
                self.result_label.config(text=languages[self.current_language]['no_save_path_selected'])
        except Exception as e:
            self.result_label.config(text=languages[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        """切换语言，更新所有UI文本"""
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'
        self.root.title(languages[self.current_language]['title'])
        self.select_button.config(text=languages[self.current_language]['select_button_text'])
        self.analyze_button.config(text=languages[self.current_language]['analyze_button_text'])
        self.switch_language_label.config(text=languages[self.current_language]['switch_language_button_text'])
        self.open_excel_label.config(text=languages[self.current_language]["open_excel_button_text"])

        # 更新输入框占位符
        current_text = self.file_entry.get()
        if current_text == languages['zh' if self.current_language == 'en' else 'en']['file_entry_placeholder']:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, languages[self.current_language]['file_entry_placeholder'])

    def create_ui(self):
        """创建UI"""
        screen_width, screen_height = self.root.winfo_screenwidth(), self.root.winfo_screenheight()
        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        x, y = (screen_width - window_width) // 2, (screen_height - window_height) // 2
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        style = ttk.Style()
        style.configure("Gray.TEntry", foreground="gray")

        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        self.file_entry = ttk.Entry(frame, width=50, style="Gray.TEntry")
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=languages[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        self.switch_language_label = ttk.Label(frame,
                                               text=languages[self.current_language]["switch_language_button_text"],
                                               cursor="hand2", foreground="Gray")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=5)

        self.result_label = ttk.Label(frame, text="", wraplength=400)
        self.result_label.pack(pady=10)

    def run(self):
        self.root.mainloop()


if __name__ == "__main__":
    app = AnalyticHierarchyProcessAHPApp()
    app.run()
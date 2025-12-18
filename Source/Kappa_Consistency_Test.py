import tkinter as tk
from tkinter import filedialog
import os
import pandas as pd
import numpy as np
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from docx import Document
from sklearn.metrics import cohen_kappa_score
from sklearn.utils import resample


# 定义语言字典
languages = {
    'zh': {
        'title': "Kappa 一致性检验",
        'select_button_text': "选择文件",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'analyze_button_text': "分析文件",
        'no_file_selected': "请选择文件。",
        'file_not_exists': "文件不存在，请重新选择。",
        'analysis_error': "分析文件时出错: {}",
        'analysis_success': "分析完成，结果已保存到 {}",
        'no_save_path': "未选择保存路径，结果未保存。",
        'open_excel_button_text': "示例数据",
        'switch_language_button_text': "中/英",
        'explanation': {
            "未加权Kappa": "不考虑分类错误严重程度的一致性指标",
            "线性加权Kappa": "考虑分类错误线性严重程度的一致性指标",
            "二次加权Kappa": "考虑分类错误二次严重程度的一致性指标",
            "95%置信区间": "用于估计统计量真实值的范围（95%置信水平）"
        },
        'interpretation': {
            "未加权Kappa": "适用于无序分类的一致性评估",
            "线性加权Kappa": "适用于有序分类，错误权重随差异线性增加",
            "二次加权Kappa": "适用于有序分类，错误权重随差异平方增加"
        },
        'confusion_matrix': "混淆矩阵",
        'conclusion': "结论",
        'strong_agreement': "Kappa值 > 0.8，表示几乎完全一致",
        'moderate_agreement': "0.6 - 0.8，表示高度一致",
        'fair_agreement': "0.4 - 0.6，表示中度一致",
        'slight_agreement': "0.2 - 0.4，表示一般一致",
        'poor_agreement': "< 0.2，表示一致性较差"

    },
    'en': {
        'title': "Kappa Consistency Test",
        'select_button_text': "Select File",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'analyze_button_text': "Analyze File",
        'no_file_selected': "Please select a valid file path.",
        'file_not_exists': "The file does not exist. Please select again.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'analysis_success': "Analysis completed. The results have been saved to {}",
        'no_save_path': "No save path selected. The results were not saved.",
        'open_excel_button_text': "Example data",
        'switch_language_button_text': "Chinese/English",
        'explanation': {
            "Unweighted Kappa": "Consistency index without considering classification error severity",
            "Linear Weighted Kappa": "Consistency index with linear weight on classification errors",
            "Quadratic Weighted Kappa": "Consistency index with quadratic weight on classification errors",
            "95% Confidence Interval": "Range for estimating the true value of the statistic (95% confidence level)"
        },
        'interpretation': {
            "Unweighted Kappa": "Suitable for unordered classification",
            "Linear Weighted Kappa": "Suitable for ordered classification with linear error weights",
            "Quadratic Weighted Kappa": "Suitable for ordered classification with quadratic error weights"
        },
        'confusion_matrix': "Confusion Matrix",
        'conclusion': "Conclusion",
        'strong_agreement': "Kappa > 0.8: Almost perfect agreement",
        'moderate_agreement': "0.6 - 0.8: Substantial agreement",
        'fair_agreement': "0.4 - 0.6: Moderate agreement",
        'slight_agreement': "0.2 - 0.4: Fair agreement",
        'poor_agreement': "< 0.2: Slight agreement"
    }
}

class KappaConsistencyTestApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data19.xls")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 检查数据是否有两列
            if df.shape[1] != 2:
                raise ValueError("数据必须包含两列，用于计算 Kappa 一致性检验。")

            # 提取两列数据
            r1 = df.iloc[:, 0].astype(str).values
            r2 = df.iloc[:, 1].astype(str).values
            labels_union = list(set(r1).union(set(r2)))  # 所有类别

            # 计算基础与加权 Kappa
            kappa_unweighted = cohen_kappa_score(r1, r2, labels=labels_union)
            kappa_linear = cohen_kappa_score(r1, r2, labels=labels_union, weights="linear")
            kappa_quadratic = cohen_kappa_score(r1, r2, labels=labels_union, weights="quadratic")

            # Bootstrap 置信区间（科研报告必需）
            def bootstrap_kappa(r1, r2, n_boot=1000, weights=None):
                rng = np.random.default_rng(0)
                n = len(r1)
                vals = []
                for _ in range(n_boot):
                    idx = rng.integers(0, n, n)
                    try:
                        val = cohen_kappa_score(r1[idx], r2[idx], labels=labels_union, weights=weights)
                        vals.append(val)
                    except Exception:
                        continue
                if len(vals) == 0:
                    return (np.nan, np.nan)
                return (np.percentile(vals, 2.5), np.percentile(vals, 97.5))

            ci_unw = bootstrap_kappa(r1, r2)
            ci_lin = bootstrap_kappa(r1, r2, weights="linear")
            ci_quad = bootstrap_kappa(r1, r2, weights="quadratic")

            # 混淆矩阵
            confusion = pd.crosstab(pd.Series(r1, name="rater1"), pd.Series(r2, name="rater2"))

            # 计算样本量
            sample_size = len(df)

            # 整理数据
            # 根据当前语言定义统计量名称
            if self.current_language == "zh":
                stat_names = {
                    "sample_size": "样本量",
                    "unweighted_kappa": "未加权Kappa",
                    "linear_kappa": "线性加权Kappa",
                    "quadratic_kappa": "二次加权Kappa",
                    "ci_label": "95% CI: ({:.4f}, {:.4f})"  # 置信区间格式（中文保持原样）
                }
            else:
                stat_names = {
                    "sample_size": "Sample Size",
                    "unweighted_kappa": "Unweighted Kappa",
                    "linear_kappa": "Linear Weighted Kappa",
                    "quadratic_kappa": "Quadratic Weighted Kappa",
                    "ci_label": "95% CI: ({:.4f}, {:.4f})"  # 置信区间格式（英文通用表达）
                }

            # 跟随语言切换的数据列表
            data = [
                [stat_names["sample_size"], sample_size, ""],
                [
                    stat_names["unweighted_kappa"],
                    f"{kappa_unweighted:.4f}",
                    stat_names["ci_label"].format(ci_unw[0], ci_unw[1])
                ],
                [
                    stat_names["linear_kappa"],
                    f"{kappa_linear:.4f}",
                    stat_names["ci_label"].format(ci_lin[0], ci_lin[1])
                ],
                [
                    stat_names["quadratic_kappa"],
                    f"{kappa_quadratic:.4f}",
                    stat_names["ci_label"].format(ci_quad[0], ci_quad[1])
                ]
            ]
            headers = ["统计量", "统计量值", "p值"] if self.current_language == 'zh' else ["Statistic",
                                                                                           "Statistic Value", "p-value"]
            df_result = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            explanations = languages[self.current_language]['explanation']
            interpretations = languages[self.current_language]['interpretation']

            # 根据当前语言设置解释表格列名
            if self.current_language == 'zh':
                expl_columns = ["Kappa一致性检验", "样本量", "Kappa值"]
            else:
                expl_columns = ["Kappa Consistency Test", "Sample Size", "Kappa Value"]
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(columns=expl_columns)
            explanation_df.insert(0,
                                  "统计量_解释说明" if self.current_language == 'zh' else "Statistic_Explanation",
                                  "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            if self.current_language == 'zh':
                interp_columns = ["统计量", "p值", "样本量", "Kappa值"]
            else:
                interp_columns = ["Statistic", "p-value", "Sample Size", "Kappa Value"]
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(columns=interp_columns)
            interpretation_df.insert(0,
                                     "统计量_结果解读" if self.current_language == 'zh' else "Statistic_Interpretation",
                                     "结果解读" if self.current_language == 'zh' else "Interpretation")
            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建一个新的 Word 文档
                doc = Document()

                # 添加结果表格
                table = doc.add_table(rows=len(df_result) + 1, cols=len(df_result.columns))
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(headers):
                    hdr_cells[col_idx].text = header
                for row_idx, row in df_result.iterrows():
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx].text = str(value)

                # 添加解释说明（改进版）
                doc.add_heading("统计量解释" if self.current_language == 'zh' else "Statistic Explanation", level=2)
                explanation_data = [
                    [
                        "未加权Kappa" if self.current_language == 'zh' else "Unweighted Kappa",
                        explanations["未加权Kappa" if self.current_language == 'zh' else "Unweighted Kappa"]
                    ],
                    [
                        "线性加权Kappa" if self.current_language == 'zh' else "Linear Weighted Kappa",
                        explanations["线性加权Kappa" if self.current_language == 'zh' else "Linear Weighted Kappa"]
                    ],
                    [
                        "二次加权Kappa" if self.current_language == 'zh' else "Quadratic Weighted Kappa",
                        explanations["二次加权Kappa" if self.current_language == 'zh' else "Quadratic Weighted Kappa"]
                    ],
                    [
                        "95%置信区间" if self.current_language == 'zh' else "95% Confidence Interval",
                        explanations["95%置信区间" if self.current_language == 'zh' else "95% Confidence Interval"]
                    ]
                ]
                expl_table = doc.add_table(rows=len(explanation_data) + 1, cols=2)
                # 设置解释表格标题
                expl_table.cell(0, 0).text = "统计量" if self.current_language == 'zh' else "Statistic"
                expl_table.cell(0, 1).text = "解释" if self.current_language == 'zh' else "Explanation"
                # 填充解释表格内容
                for i, (stat, desc) in enumerate(explanation_data, 1):
                    expl_table.cell(i, 0).text = stat
                    expl_table.cell(i, 1).text = desc

                # 添加结果解读（改进版）
                doc.add_paragraph()
                doc.add_heading("结果解读" if self.current_language == 'zh' else "Result Interpretation", level=2)
                # 根据当前语言动态生成解读数据
                interpretation_data = [
                    [
                        "未加权Kappa" if self.current_language == 'zh' else "Unweighted Kappa",
                        interpretations["未加权Kappa" if self.current_language == 'zh' else "Unweighted Kappa"]
                    ],
                    [
                        "线性加权Kappa" if self.current_language == 'zh' else "Linear Weighted Kappa",
                        interpretations["线性加权Kappa" if self.current_language == 'zh' else "Linear Weighted Kappa"]
                    ],
                    [
                        "二次加权Kappa" if self.current_language == 'zh' else "Quadratic Weighted Kappa",
                        interpretations[
                            "二次加权Kappa" if self.current_language == 'zh' else "Quadratic Weighted Kappa"]
                    ]
                ]
                interp_table = doc.add_table(rows=len(interpretation_data) + 1, cols=2)
                # 设置解读表格标题
                interp_table.cell(0, 0).text = "统计量" if self.current_language == 'zh' else "Statistic"
                interp_table.cell(0, 1).text = "解读" if self.current_language == 'zh' else "Interpretation"
                # 填充解读表格内容
                for i, (stat, desc) in enumerate(interpretation_data, 1):
                    interp_table.cell(i, 0).text = stat
                    interp_table.cell(i, 1).text = desc

                # 添加混淆矩阵
                doc.add_paragraph()
                doc.add_heading(languages[self.current_language]['confusion_matrix'], level=2)  # 添加标题
                # 创建混淆矩阵表格
                confusion_table = doc.add_table(rows=len(confusion) + 1, cols=len(confusion.columns) + 1)
                # 填充表头（第一行和第一列）
                confusion_table.cell(0, 0).text = ""  # 左上角空白
                for col_idx, col_name in enumerate(confusion.columns):
                    confusion_table.cell(0, col_idx + 1).text = str(col_name)  # 列标题（评估者2）
                for row_idx, row_name in enumerate(confusion.index):
                    confusion_table.cell(row_idx + 1, 0).text = str(row_name)  # 行标题（评估者1）
                    for col_idx, value in enumerate(confusion.iloc[row_idx]):
                        confusion_table.cell(row_idx + 1, col_idx + 1).text = str(value)  # 填充数据

                # 添加结论标题
                doc.add_paragraph()
                doc.add_heading(languages[self.current_language]['conclusion'], level=2)

                # 根据未加权Kappa值判断一致性等级
                kappa = kappa_unweighted
                if kappa > 0.8:
                    conclusion_text = languages[self.current_language]['strong_agreement']
                elif 0.6 <= kappa <= 0.8:
                    conclusion_text = languages[self.current_language]['moderate_agreement']
                elif 0.4 <= kappa < 0.6:
                    conclusion_text = languages[self.current_language]['fair_agreement']
                elif 0.2 <= kappa < 0.4:
                    conclusion_text = languages[self.current_language]['slight_agreement']
                else:
                    conclusion_text = languages[self.current_language]['poor_agreement']

                # 添加结论内容
                doc.add_paragraph(conclusion_text)

                # 保存 Word 文档
                doc.save(save_path)

                result_msg = languages[self.current_language]['analysis_success'].format(save_path)
                self.result_label.config(text=result_msg, wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])
        self.open_excel_label.config(text=languages[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=languages[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=languages[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = KappaConsistencyTestApp()
    app.run()

if __name__ == "__main__":
    run_app()
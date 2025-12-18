import tkinter as tk
from tkinter import filedialog
import os
import pandas as pd
import numpy as np
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
from sklearn.linear_model import Ridge, RidgeCV
from sklearn.metrics import mean_squared_error, r2_score
from sklearn.preprocessing import StandardScaler
from sklearn.model_selection import cross_val_score, train_test_split
from scipy.stats import t, shapiro, normaltest, levene, bartlett
from statsmodels.stats.stattools import durbin_watson
from statsmodels.graphics.gofplots import qqplot
from docx import Document
from docx.shared import Inches

# 定义语言字典（所有键使用英文）
LANGUAGES = {
    'zh': {
        'title': "岭回归",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "请选择文件。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'images_saved': "图片已保存到 {}",
        'switch_language': "中/英",
        'open_excel_button_text': "示例数据",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'explanation': {
            "coefficients": "回归系数，表示每个自变量对因变量的影响程度。",
            "intercept": "截距，是当所有自变量为 0 时因变量的预测值。",
            "mse": "均方误差，衡量预测值与真实值之间的平均误差。",
            "r_squared": "决定系数，取值范围在 0 到 1 之间，越接近 1 表示模型拟合效果越好。",
            "adjusted_r_squared": "调整决定系数，考虑了模型中自变量的数量，对模型的拟合优度进行了调整。",
            "f_value": "F 统计量，用于检验整个回归模型的显著性。",
            "t_value": "近似 t 统计量，用于参考每个自变量的相对重要性。",
            "p_value": "近似 p 值，用于参考自变量的相对显著性，值越小参考意义越大。",
            "optimal_alpha": "最优正则化参数，通过交叉验证确定的最佳alpha值",
            "regression_equation": "回归方程，用于预测因变量的数学表达式。",
            "95_ci": "95%置信区间，表示系数可能存在的范围"
        },
        'chart_labels': {
            'actual': '实际值',
            'predicted': '预测值',
            'actual_vs_predicted': '实际值 vs 预测值',
            'alpha': 'Alpha值',
            'coefficients': '系数值',
            'ridge_trace_title': '岭回归系数随正则化参数的变化',
            'cv_mse_title': '交叉验证均方误差与Alpha值的关系'
        },
        'table_headers': {
            'indicator': '指标',
            'value': '值',
            'model': '模型',
            'optimal_alpha': '最优Alpha',
            'intercept': '截距',
            'regression_equation': '回归方程',
            'train_mse': '训练集MSE',
            'test_mse': '测试集MSE',
            'train_r_squared': '训练集R²',
            'test_r_squared': '测试集R²',
            'train_adjusted_r_squared': '训练集调整后R²',
            'test_adjusted_r_squared': '测试集调整后R²',
            'f_value': 'F值'
        },
        'residual_analysis': {
            'title': '残差分析',
            'residuals': '残差',
            'frequency': '频率',
            'histogram_title': '残差分布直方图',
            'qq_plot_title': '残差Q-Q图',
            'predicted_values': '预测值',
            'residuals_vs_predicted': '残差 vs 预测值',
            'normality_test': '残差正态性检验: Shapiro-Wilk统计量={:.4f}, p值={:.4f}',
            'd_agostino_test': "D'Agostino-Pearson统计量={:.4f}, p值={:.4f}\n",
            'small_sample': "样本量不足，未执行D'Agostino-Pearson检验\n",
            'homoscedasticity_test': '同方差性检验: Levene统计量={:.4f}, p值={:.4f}',
            'independence_test': '残差独立性检验: Durbin-Watson统计量={:.4f}',
            'min_mse': '最小MSE: {:.4f}',
            'optimal_alpha_label': '最优Alpha: {:.4f}',
            'training_set': '训练集',
            'test_set': '测试集'
        }
    },
    'en': {
        'title': "Ridge Regression",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'images_saved': "Images have been saved to {}",
        'switch_language': "Chinese/English",
        'open_excel_button_text': "Example data",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'explanation': {
            "coefficients": "Regression coefficients, indicating the influence of each independent variable on the dependent variable.",
            "intercept": "Intercept, which is the predicted value of the dependent variable when all independent variables are 0.",
            "mse": "Mean squared error, measuring the average error between the predicted and actual values.",
            "r_squared": "Coefficient of determination, ranging from 0 to 1. A value closer to 1 indicates a better fit of the model.",
            "adjusted_r_squared": "Adjusted coefficient of determination, which takes into account the number of independent variables in the model and adjusts the goodness of fit of the model.",
            "f_value": "F statistic, used to test the significance of the entire regression model.",
            "t_value": "Approximate t-statistic, used to reference the relative importance of each independent variable.",
            "p_value": "Approximate p-value, used to reference the relative significance of the independent variable. The smaller the value, the more meaningful the reference.",
            "optimal_alpha": "Optimal regularization parameter, the best alpha value determined by cross-validation",
            "regression_equation": "Regression equation, a mathematical expression used to predict the dependent variable.",
            "95_ci": "95% confidence interval, representing the range where the coefficient is likely to exist"
        },
        'chart_labels': {
            'actual': 'Actual Values',
            'predicted': 'Predicted Values',
            'actual_vs_predicted': 'Actual vs Predicted Values',
            'alpha': 'Alpha',
            'coefficients': 'Coefficients',
            'ridge_trace_title': 'Ridge coefficients as a function of regularization',
            'cv_mse_title': 'Cross-validation MSE vs Alpha'
        },
        'table_headers': {
            'indicator': 'Indicator',
            'value': 'Value',
            'model': 'Model',
            'optimal_alpha': 'Optimal Alpha',
            'intercept': 'Intercept',
            'regression_equation': 'Regression Equation',
            'train_mse': 'Train MSE',
            'test_mse': 'Test MSE',
            'train_r_squared': 'Train R-squared (R²)',
            'test_r_squared': 'Test R-squared (R²)',
            'train_adjusted_r_squared': 'Train Adjusted R-squared',
            'test_adjusted_r_squared': 'Test Adjusted R-squared',
            'f_value': 'F-value'
        },
        'residual_analysis': {
            'title': 'Residual Analysis',
            'residuals': 'Residuals',
            'frequency': 'Frequency',
            'histogram_title': 'Histogram of Residuals',
            'qq_plot_title': 'Q-Q Plot of Residuals',
            'predicted_values': 'Predicted Values',
            'residuals_vs_predicted': 'Residuals vs Predicted Values',
            'normality_test': 'Residual normality test: Shapiro-Wilk statistic={:.4f}, p-value={:.4f}',
            'd_agostino_test': "D'Agostino-Pearson statistic={:.4f}, p-value={:.4f}\n",
            'small_sample': "Sample size is insufficient, D'Agostino-Pearson test not performed\n",
            'homoscedasticity_test': 'Homoscedasticity test: Levene statistic={:.4f}, p-value={:.4f}',
            'independence_test': 'Residual independence test: Durbin-Watson statistic={:.4f}',
            'min_mse': 'Min MSE: {:.4f}',
            'optimal_alpha_label': 'Optimal Alpha: {:.4f}',
            'training_set': 'Training Set',
            'test_set': 'Test Set'
        }
    }
}

plt.rcParams["font.family"] = ["SimHei", "WenQuanYi Micro Hei", "Heiti TC", "Arial Unicode MS", "DejaVu Sans"]
plt.rcParams["axes.unicode_minus"] = False  # 解决负号显示问题

class RidgeRegressionAnalysisApp:
    def __init__(self, root=None):
        self.current_language = "en"  # 当前语言，默认为中文

        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data32.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"{LANGUAGES[self.current_language]['file_not_found']}：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"{LANGUAGES[self.current_language]['analysis_error'].format(str(e))}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 读取Excel文件（第一行是表头，第一列是因素名称列）
            df = pd.read_excel(file_path)
            # 提取变量名称（第一行表头）
            var_names = df.columns[1:].tolist()  # 从第二列开始是变量名
            factor_names = df.iloc[:, 0].tolist()  # 第一列是因素名称

            # 自变量为第一列之后的数据，因变量为最后一列
            X = df.iloc[:, 1:-1].values  # 自变量：第二列到倒数第二列
            y = df.iloc[:, -1].values  # 因变量：最后一列
            dependent_var_name = var_names[-1]  # 因变量名称

            # 拆分训练集(80%)和测试集(20%)
            X_train, X_test, y_train, y_test = train_test_split(
                X, y, test_size=0.2, random_state=42, stratify=y if len(np.unique(y)) <= 10 else None
            )

            # 交叉验证选择最优alpha值（仅用训练集）
            alphas = np.logspace(-4, 4, 100)  # 更广泛的alpha范围
            ridge_cv = RidgeCV(alphas=alphas, cv=5, scoring='neg_mean_squared_error')
            ridge_cv.fit(X_train, y_train)  # 仅用训练集
            optimal_alpha = ridge_cv.alpha_

            # 使用最优alpha进行岭回归（仅用训练集）
            ridge = Ridge(alpha=optimal_alpha)
            ridge.fit(X_train, y_train)

            # 预测需同时在训练集和测试集上进行
            y_pred_train = ridge.predict(X_train)
            y_pred_test = ridge.predict(X_test)

            # 计算基本指标
            coefficients = ridge.coef_
            intercept = ridge.intercept_

            # 评估指标需同时计算训练集和测试集
            mse_train = mean_squared_error(y_train, y_pred_train)
            mse_test = mean_squared_error(y_test, y_pred_test)
            r2_train = r2_score(y_train, y_pred_train)
            r2_test = r2_score(y_test, y_pred_test)

            # 调整R²计算
            n_train, p = X_train.shape
            adjusted_r2_train = 1 - (1 - r2_train) * (n_train - 1) / (n_train - p - 1)
            n_test = X_test.shape[0]
            adjusted_r2_test = 1 - (1 - r2_test) * (n_test - 1) / (n_test - p - 1)

            # 改进的自助法计算系数显著性
            np.random.seed(42)
            n_boot = 1000
            boot_coefs = []
            # 使用分层抽样的自助法
            for _ in range(n_boot):
                weights = np.ones(n_train) / n_train  # 等权重
                idx = np.random.choice(n_train, size=n_train, replace=True, p=weights)
                ridge_boot = Ridge(alpha=optimal_alpha)
                ridge_boot.fit(X_train[idx], y_train[idx])
                boot_coefs.append(ridge_boot.coef_)

            boot_coefs = np.array(boot_coefs)
            se_boot = np.std(boot_coefs, axis=0, ddof=1)  # 使用样本标准差（n-1自由度）

            # 计算t值和经验p值
            t_values = coefficients / se_boot
            p_values = []
            for i in range(len(coefficients)):
                # 计算系数绝对值大于观测值的比例
                extreme = np.sum(np.abs(boot_coefs[:, i]) >= np.abs(coefficients[i]))
                p_value = (extreme + 1) / (n_boot + 1)  # 加1校正
                p_values.append(p_value)
            p_values = np.array(p_values)

            # 添加95%置信区间
            ci_lower = np.percentile(boot_coefs, 2.5, axis=0)
            ci_upper = np.percentile(boot_coefs, 97.5, axis=0)

            # 计算标准化系数（同时标准化自变量和因变量）
            scaler_X = StandardScaler()
            scaler_y = StandardScaler()
            X_scaled = scaler_X.fit_transform(X_train)
            y_scaled = scaler_y.fit_transform(y_train.reshape(-1, 1)).flatten()
            ridge_scaled = Ridge(alpha=optimal_alpha)
            ridge_scaled.fit(X_scaled, y_scaled)
            standardized_coefficients = ridge_scaled.coef_

            # 计算F值（使用交叉验证结果）
            cv_scores = cross_val_score(ridge, X_train, y_train, cv=5, scoring='r2')
            mean_r2 = np.mean(cv_scores)
            f_value = (mean_r2 / p) / ((1 - mean_r2) / (n_train - p - 1))

            # 构建回归方程
            equation_parts = [f"{dependent_var_name} = {intercept:.4f}"]
            for i, var_name in enumerate(var_names[:-1]):  # 排除因变量
                coef = coefficients[i]
                if coef >= 0 and i > 0:
                    equation_parts.append(f"+ {coef:.4f} × {var_name}")
                else:
                    equation_parts.append(f"{coef:.4f} × {var_name}")
            regression_equation = " ".join(equation_parts)

            # 准备结果数据（使用实际变量名称）
            headers = [
                LANGUAGES[self.current_language]['table_headers']['model'],
                LANGUAGES[self.current_language]['table_headers']['optimal_alpha'],
                LANGUAGES[self.current_language]['table_headers']['intercept'],
                LANGUAGES[self.current_language]['table_headers']['regression_equation'],
                LANGUAGES[self.current_language]['table_headers']['train_mse'],
                LANGUAGES[self.current_language]['table_headers']['test_mse'],
                LANGUAGES[self.current_language]['table_headers']['train_r_squared'],
                LANGUAGES[self.current_language]['table_headers']['test_r_squared'],
                LANGUAGES[self.current_language]['table_headers']['train_adjusted_r_squared'],
                LANGUAGES[self.current_language]['table_headers']['test_adjusted_r_squared'],
                LANGUAGES[self.current_language]['table_headers']['f_value']
            ]

            # 添加自变量相关列（使用实际变量名）
            for name in var_names[:-1]:  # 排除因变量名称
                headers.extend([
                    f"{name} {LANGUAGES[self.current_language]['explanation']['coefficients']}",
                    f"{name} {LANGUAGES[self.current_language]['explanation']['coefficients']} (Standardized)",
                    f"{name} {LANGUAGES[self.current_language]['explanation']['t_value']}",
                    f"{name} {LANGUAGES[self.current_language]['explanation']['p_value']}",
                    f"{name} {LANGUAGES[self.current_language]['explanation']['95_ci']}"
                ])

            # 构建结果行
            result_row = [
                "Ridge Regression",
                optimal_alpha,
                intercept,
                regression_equation,
                mse_train,
                mse_test,
                r2_train,
                r2_test,
                adjusted_r2_train,
                adjusted_r2_test,
                f_value
            ]
            for i in range(len(coefficients)):
                result_row.extend([
                    coefficients[i],
                    standardized_coefficients[i],
                    t_values[i],
                    p_values[i],
                    f"[{ci_lower[i]:.4f}, {ci_upper[i]:.4f}]"
                ])

            # 创建结果数据框并转置
            df_result = pd.DataFrame([result_row], columns=headers)
            df_result_transposed = df_result.transpose()  # 转置表格，行变列，列变行
            df_result_transposed.columns = [LANGUAGES[self.current_language]['table_headers']['value']]  # 转置后设置列名

            # 添加解释说明
            explanation_rows = []
            for col in headers[1:]:  # 跳过Model列
                # 提取指标基础名称
                if LANGUAGES[self.current_language]['table_headers']['regression_equation'] in col:
                    base_name = "regression_equation"
                elif LANGUAGES[self.current_language]['table_headers']['optimal_alpha'] in col:
                    base_name = "optimal_alpha"
                elif LANGUAGES[self.current_language]['table_headers']['intercept'] in col:
                    base_name = "intercept"
                elif "MSE" in col or LANGUAGES[self.current_language]['table_headers']['train_mse'] in col:
                    base_name = "mse"
                elif LANGUAGES[self.current_language]['explanation']['coefficients'] in col:
                    base_name = "coefficients"
                elif LANGUAGES[self.current_language]['explanation']['95_ci'] in col:
                    base_name = "95_ci"
                elif "R-squared" in col or "R²" in col or "决定系数" in col:
                    base_name = "r_squared"
                elif "Adjusted" in col or "调整后" in col:
                    base_name = "adjusted_r_squared"
                elif "F-value" in col or "F值" in col:
                    base_name = "f_value"
                elif LANGUAGES[self.current_language]['explanation']['t_value'] in col:
                    base_name = "t_value"
                elif LANGUAGES[self.current_language]['explanation']['p_value'] in col:
                    base_name = "p_value"
                else:
                    base_name = None

                explanation = LANGUAGES[self.current_language]['explanation'].get(base_name, "") if base_name else ""
                explanation_rows.append([col, explanation])

            explanation_df = pd.DataFrame(
                explanation_rows,
                columns=[
                    LANGUAGES[self.current_language]['table_headers']['indicator'],
                    LANGUAGES[self.current_language]['table_headers']['indicator'] + " " +
                    ("解释" if self.current_language == 'zh' else "Explanation")
                ]
            )

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建Word文档
                doc = Document()

                # 添加模型结果表格（使用转置后的表格）
                doc.add_heading(
                    '岭回归分析结果' if self.current_language == 'zh' else 'Ridge Regression Results',
                    level=1
                )
                # 转置后的表格行数为原列数，列数为原行数+1（包含指标名称）
                table = doc.add_table(rows=df_result_transposed.shape[0] + 1, cols=df_result_transposed.shape[1] + 1)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = LANGUAGES[self.current_language]['table_headers']['indicator']
                hdr_cells[1].text = LANGUAGES[self.current_language]['table_headers']['value']

                for row_idx, (index, row) in enumerate(df_result_transposed.iterrows()):
                    row_cells = table.rows[row_idx + 1].cells
                    row_cells[0].text = str(index)  # 指标名称
                    value = row.iloc[0]
                    # 格式化数值显示
                    if isinstance(value, float):
                        row_cells[1].text = f"{value:.4f}"
                    else:
                        row_cells[1].text = str(value)

                # 添加解释项目列表
                doc.add_heading(
                    LANGUAGES[self.current_language]['table_headers']['indicator'] +
                    ("解释" if self.current_language == 'zh' else " Explanation"),
                    level=1
                )
                # 创建一个段落用于存放项目符号列表
                bullet_paragraph = doc.add_paragraph()
                for row_idx, row in explanation_df.iterrows():
                    # 每个指标作为一个项目符号，格式为“• 指标: 解释”
                    indicator = str(row.iloc[0])  # 使用iloc按位置访问第一列
                    explanation = str(row.iloc[1])  # 使用iloc按位置访问第二列
                    # 添加项目符号和内容
                    bullet_paragraph.add_run(f"• {indicator}: {explanation}\n")

                # 获取保存路径的目录
                save_dir = os.path.dirname(save_path)

                # 生成散点图（支持多语言）
                plt.figure(figsize=(10, 6))
                plt.scatter(
                    y_test, y_pred_test, alpha=0.6,
                    label=LANGUAGES[self.current_language]['residual_analysis']['test_set']
                )
                plt.scatter(
                    y_train, y_pred_train, alpha=0.3,
                    label=LANGUAGES[self.current_language]['residual_analysis']['training_set']
                )
                plt.plot([y.min(), y.max()], [y.min(), y.max()], 'r--', lw=2)
                plt.xlabel(LANGUAGES[self.current_language]['chart_labels']['actual'])
                plt.ylabel(LANGUAGES[self.current_language]['chart_labels']['predicted'])
                plt.title(LANGUAGES[self.current_language]['chart_labels']['actual_vs_predicted'])
                plt.legend()
                plt.grid(alpha=0.3)
                plt.tight_layout()
                img_name = "ridge_regression_scatter.png"
                img_path = os.path.join(save_dir, img_name)
                plt.savefig(img_path, dpi=300, bbox_inches='tight')
                plt.close()

                # 生成岭迹图
                alphas_plot = np.logspace(-4, 4, 50)
                coefs_plot = []
                for a in alphas_plot:
                    ridge_plot = Ridge(alpha=a)
                    ridge_plot.fit(X_train, y_train)
                    coefs_plot.append(ridge_plot.coef_)

                plt.figure(figsize=(10, 6))
                ax = plt.gca()
                # 为每条系数曲线添加标签
                for i, label in enumerate(var_names[:-1]):
                    ax.plot(alphas_plot, np.array(coefs_plot)[:, i], label=label)
                ax.set_xscale('log')
                plt.axvline(
                    x=optimal_alpha, color='red', linestyle='--',
                    label=LANGUAGES[self.current_language]['residual_analysis']['optimal_alpha_label'].format(optimal_alpha)
                )
                plt.xlabel(LANGUAGES[self.current_language]['chart_labels']['alpha'])
                plt.ylabel(LANGUAGES[self.current_language]['chart_labels']['coefficients'])
                plt.title(LANGUAGES[self.current_language]['chart_labels']['ridge_trace_title'])
                plt.legend(bbox_to_anchor=(1.05, 1), loc='upper left')
                plt.grid(alpha=0.3)
                plt.tight_layout()
                img_name_ridge_trace = "ridge_regression_trace.png"
                img_path_ridge_trace = os.path.join(save_dir, img_name_ridge_trace)
                plt.savefig(img_path_ridge_trace, dpi=300, bbox_inches='tight')
                plt.close()

                # 生成交叉验证MSE图
                mse_values = []
                for a in alphas_plot:
                    ridge_cv = Ridge(alpha=a)
                    scores = cross_val_score(ridge_cv, X_train, y_train, cv=5, scoring='neg_mean_squared_error')
                    mse_values.append(-np.mean(scores))

                plt.figure(figsize=(10, 6))
                plt.plot(alphas_plot, mse_values, 'b-', linewidth=2)
                plt.xscale('log')
                plt.axvline(
                    x=optimal_alpha, color='red', linestyle='--',
                    label=LANGUAGES[self.current_language]['residual_analysis']['optimal_alpha_label'].format(optimal_alpha)
                )
                # 添加最优MSE标注
                min_mse = np.min(mse_values)
                plt.scatter(
                    [optimal_alpha], [min_mse], color='green', s=100, zorder=5,
                    label=LANGUAGES[self.current_language]['residual_analysis']['min_mse'].format(min_mse)
                )
                plt.xlabel(LANGUAGES[self.current_language]['chart_labels']['alpha'])
                plt.ylabel('均方误差 (MSE)' if self.current_language == 'zh' else 'Mean Squared Error (MSE)')
                plt.title(LANGUAGES[self.current_language]['chart_labels']['cv_mse_title'])
                plt.legend()
                plt.grid(alpha=0.3)
                plt.tight_layout()
                img_name_cv = "ridge_cv_mse.png"
                img_path_cv = os.path.join(save_dir, img_name_cv)
                plt.savefig(img_path_cv, dpi=300, bbox_inches='tight')
                plt.close()

                # 模型假设检验（残差分析）
                residuals = y_test - y_pred_test

                # 1. 残差正态性检验
                shapiro_stat, shapiro_p = shapiro(residuals)  # Shapiro-Wilk检验
                k2_stat, normality_p = normaltest(residuals)  # D'Agostino-Pearson检验

                # 2. 同方差性检验
                levene_stat, levene_p = levene(y_pred_test, residuals)  # Levene检验
                bartlett_stat, bartlett_p = bartlett(y_pred_test, residuals)  # Bartlett检验

                # 3. 残差独立性检验（Durbin-Watson检验）
                dw_stat = durbin_watson(residuals)

                # 添加残差诊断图
                # 残差直方图（正态性）
                plt.figure(figsize=(10, 6))
                plt.hist(residuals, bins=15, alpha=0.6, edgecolor='black')
                plt.xlabel(LANGUAGES[self.current_language]['residual_analysis']['residuals'])
                plt.ylabel(LANGUAGES[self.current_language]['residual_analysis']['frequency'])
                plt.title(LANGUAGES[self.current_language]['residual_analysis']['histogram_title'])
                residual_hist_path = os.path.join(save_dir, "residual_histogram.png")
                plt.savefig(residual_hist_path, bbox_inches='tight', dpi=300)
                plt.close()

                # 残差Q-Q图
                plt.figure(figsize=(10, 6))
                qqplot(residuals, line='s', ax=plt.gca())
                plt.title(LANGUAGES[self.current_language]['residual_analysis']['qq_plot_title'])
                qq_plot_path = os.path.join(save_dir, "residual_qqplot.png")
                plt.savefig(qq_plot_path, bbox_inches='tight', dpi=300)
                plt.close()

                # 残差 vs 预测值图（同方差性）
                plt.figure(figsize=(10, 6))
                plt.scatter(y_pred_test, residuals, alpha=0.6)
                plt.axhline(y=0, color='r', linestyle='--')
                plt.xlabel(LANGUAGES[self.current_language]['residual_analysis']['predicted_values'])
                plt.ylabel(LANGUAGES[self.current_language]['residual_analysis']['residuals'])
                plt.title(LANGUAGES[self.current_language]['residual_analysis']['residuals_vs_predicted'])
                residual_vs_pred_path = os.path.join(save_dir, "residuals_vs_predicted.png")
                plt.savefig(residual_vs_pred_path, bbox_inches='tight', dpi=300)
                plt.close()

                # 在Word文档中插入图片
                doc.add_heading(LANGUAGES[self.current_language]['chart_labels']['actual_vs_predicted'], level=2)
                doc.add_picture(img_path, width=Inches(6))

                doc.add_heading(LANGUAGES[self.current_language]['chart_labels']['ridge_trace_title'], level=2)
                doc.add_picture(img_path_ridge_trace, width=Inches(6))

                doc.add_heading(LANGUAGES[self.current_language]['chart_labels']['cv_mse_title'], level=2)
                doc.add_picture(img_path_cv, width=Inches(6))

                # 添加残差分析结果
                doc.add_heading(LANGUAGES[self.current_language]['residual_analysis']['title'], level=1)
                doc.add_paragraph(f'''
                {LANGUAGES[self.current_language]['residual_analysis']['normality_test'].format(shapiro_stat, shapiro_p)}
                {LANGUAGES[self.current_language]['residual_analysis']['d_agostino_test'].format(k2_stat, normality_p) if not np.isnan(k2_stat) else LANGUAGES[self.current_language]['residual_analysis']['small_sample']}
                {LANGUAGES[self.current_language]['residual_analysis']['homoscedasticity_test'].format(levene_stat, levene_p)}
                {LANGUAGES[self.current_language]['residual_analysis']['independence_test'].format(dw_stat)}
                ''')

                # 添加残差图表到文档
                doc.add_picture(residual_hist_path, width=Inches(6))
                doc.add_picture(qq_plot_path, width=Inches(6))
                doc.add_picture(residual_vs_pred_path, width=Inches(6))

                # 保存Word文档
                doc.save(save_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(save_path)
                result_msg += LANGUAGES[self.current_language]['images_saved'].format(save_dir)
                self.result_label.config(text=result_msg, wraplength=400)
            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(LANGUAGES[self.current_language]["title"])
        self.select_button.config(text=LANGUAGES[self.current_language]["select_button"])
        current_entry = self.file_entry.get()
        self.file_entry.delete(0, tk.END)
        if current_entry == LANGUAGES['zh']["file_entry_placeholder"] or current_entry == LANGUAGES['en'][
            "file_entry_placeholder"]:
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')
        else:
            self.file_entry.insert(0, current_entry)
        self.analyze_button.config(text=LANGUAGES[self.current_language]["analyze_button"])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]["switch_language"])
        self.open_excel_label.config(text=LANGUAGES[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        # 窗口居中设置
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建界面组件
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        self.select_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["select_button"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        self.analyze_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["analyze_button"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=LANGUAGES[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        self.switch_language_label = ttk.Label(frame, text=LANGUAGES[self.current_language]["switch_language"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        self.root.mainloop()


def run_app():
    app = RidgeRegressionAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()
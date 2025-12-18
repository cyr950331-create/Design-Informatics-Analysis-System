import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog, simpledialog
import tkinter as tk
from scipy import stats
import matplotlib.pyplot as plt
import pathlib
from sklearn.cross_decomposition import CCA
from sklearn.preprocessing import StandardScaler
from docx import Document
from docx.shared import Inches
import traceback

# 支持中文字体显示（如系统中无 SimHei，可改为其他字体）
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 多语言字典（新增了关于载荷、核心程度、符号含义、图例等文本）
languages = {
    'zh': {
        'title': "典型相关分析",
        'select_button_text': "选择文件",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'analyze_button_text': "分析文件",
        'no_file_selected': "请选择文件。",
        'file_not_exists': "文件不存在，请重新选择。",
        'analysis_error': "分析文件时出错: {}",
        'analysis_complete': "分析完成，结果已保存到 {}",
        'images_saved': "结果图片已保存到 {}",
        'no_save_path_selected': "未选择保存路径，结果未保存。",
        'switch_language_button_text': "中/英",
        'open_excel_button_text': "示例数据",
        'input_x_columns': "请输入 X 组变量的列名（以英文逗号分隔）：\n例如：A,B,C",
        'invalid_columns': "输入的列名无效或不存在，请检查。",
        'info_auto_y': "未输入的列将自动作为 Y 组变量。",
        'cancel_msg': "分析已取消。",
        'samples_error': "样本数必须大于每组变量数，否则 CCA 不稳定。",
        'no_numeric': "文件中没有有效的数值列。",
        'core_section_title': "核心设计指标 / 典型载荷与权重",
        'core_section_desc': "下表显示各原始指标在各典型变量上的载荷（相关系数）与权重，并按绝对载荷值排序，便于定位核心指标。",
        'loading_table_x_title': "X 组载荷（按 |载荷| 降序）",
        'loading_table_y_title': "Y 组载荷（按 |载荷| 降序）",
        'weights_x_title': "X 组权重表",
        'weights_y_title': "Y 组权重表",
        'loading_cols': ["变量", "载荷", "|载荷|", "核心程度"],
        'loading_degree_absolute_core': "绝对核心",
        'loading_degree_core': "核心",
        'loading_degree_secondary': "次要",
        'loading_degree_negligible': "几乎无关",
        'loading_sign_explain_title': "载荷符号含义说明",
        'loading_sign_explain': "载荷为正表示该原始指标与典型变量同方向变化（正相关）；载荷为负表示反方向变化（负相关）。在实际解释中，正/负符号反映该指标对典型方向的贡献方向，需结合业务含义解读。",
        'top_k_label': "主指标",
        'axis_label_format': "（主指标: {}）",  # will format with comma-separated list like Var1 0.78, Var2 -0.65...
    },
    'en': {
        'title': "Canonical Correlation Analysis",
        'select_button_text': "Select File",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'analyze_button_text': "Analyze File",
        'no_file_selected': "Please select a valid file path.",
        'file_not_exists': "The file does not exist. Please select again.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'analysis_complete': "Analysis completed. Results saved to {}",
        'images_saved': "The result image has been saved to {}",
        'no_save_path_selected': "No save path selected. The results were not saved.",
        'switch_language_button_text': "Chinese/English",
        'open_excel_button_text': "Example data",
        'input_x_columns': "Enter X variable column names (comma-separated):\nExample: A,B,C",
        'invalid_columns': "Invalid or non-existing column names. Please check.",
        'info_auto_y': "Unlisted columns will automatically be used as Y variables.",
        'cancel_msg': "Analysis canceled.",
        'samples_error': "Number of samples must be greater than variables in each set; CCA unstable otherwise.",
        'no_numeric': "No numeric columns found in the file.",
        'core_section_title': "Core design indicators / Canonical loadings & weights",
        'core_section_desc': "Below tables show each original variable's loading (correlation) with canonical variables and the weights, ordered by absolute loading to help locate core indicators.",
        'loading_table_x_title': "X-group loadings (sorted by |loading| desc)",
        'loading_table_y_title': "Y-group loadings (sorted by |loading| desc)",
        'weights_x_title': "X-group weights table",
        'weights_y_title': "Y-group weights table",
        'loading_cols': ["Variable", "Loading", "|Loading|", "Core level"],
        'loading_degree_absolute_core': "Absolute core",
        'loading_degree_core': "Core",
        'loading_degree_secondary': "Secondary",
        'loading_degree_negligible': "Negligible",
        'loading_sign_explain_title': "Meaning of loading sign",
        'loading_sign_explain': "A positive loading means the original variable varies in the same direction as the canonical variable (positive correlation); a negative loading means opposite direction (negative correlation). Interpret sign in the context of your domain.",
        'top_k_label': "Top indicators",
        'axis_label_format': "(Top: {})",
    }
}

def classify_loading_degree(abs_loading, lang_map):
    """根据绝对载荷给出核心程度（阈值可调整）"""
    # 阈值（可根据领域调整）
    if abs_loading >= 0.8:
        return lang_map['loading_degree_absolute_core']
    elif abs_loading >= 0.6:
        return lang_map['loading_degree_core']
    elif abs_loading >= 0.3:
        return lang_map['loading_degree_secondary']
    else:
        return lang_map['loading_degree_negligible']

class CanonicalCorrelationAnalysisApp:
    def __init__(self, root=None):
        self.current_language = "en"
        if root is None:
            self.root = ttk.Window(themename="flatly")
        else:
            self.root = root
        self.root.title(languages[self.current_language]["title"])
        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data21.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def analyze_file(self):
        lang = languages[self.current_language]
        file_path = self.file_entry.get()

        if file_path == lang["file_entry_placeholder"]:
            self.result_label.config(text=lang["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=lang["file_not_exists"])
            return

        try:
            # 读取并保留数值列，删除含 NA 的行
            df = pd.read_excel(file_path)
            numeric_df = df.select_dtypes(include=[np.number]).dropna()
            if numeric_df.empty:
                Messagebox.show_error(lang['no_numeric'], title="错误" if self.current_language == "zh" else "Error")
                return

            columns = list(numeric_df.columns)

            # 弹窗输入 X 列名（允许大小写/空格不严格匹配并去重）
            input_str = simpledialog.askstring(
                title=lang["title"],
                prompt=lang["input_x_columns"] + f"\n\n可选列：{', '.join(columns)}\n\n" + lang["info_auto_y"],
                parent=self.root
            )
            if not input_str:
                msg = lang['cancel_msg']
                Messagebox.show_info(msg, title="信息" if self.current_language == "zh" else "Info")
                self.result_label.config(text=msg)
                return

            # 构造大小写不敏感的映射并去重用户输入
            cols_map = {c.strip().lower(): c for c in columns}
            # preserve order and dedupe
            input_tokens = [tok.strip() for tok in input_str.split(',') if tok.strip()]
            seen = set()
            x_cols = []
            for tok in input_tokens:
                key = tok.lower()
                if key in cols_map and key not in seen:
                    x_cols.append(cols_map[key])
                    seen.add(key)

            if len(x_cols) == 0:
                Messagebox.show_error(lang["invalid_columns"], title="错误" if self.current_language == "zh" else "Error")
                return

            # Y 为剩余列
            y_cols = [c for c in columns if c not in x_cols]
            if len(y_cols) == 0:
                Messagebox.show_error(lang["invalid_columns"], title="错误" if self.current_language == "zh" else "Error")
                return

            # 提取数据（已删除 NA）
            df_clean = numeric_df.reset_index(drop=True)
            X = df_clean[x_cols].values
            Y = df_clean[y_cols].values

            # --- 样本数检查（防止 CCA 崩溃） ---
            if X.shape[0] <= max(X.shape[1], Y.shape[1]):
                Messagebox.show_error(lang['samples_error'], title="错误" if self.current_language == "zh" else "Error")
                self.result_label.config(text=lang['samples_error'])
                return

            # 标准化
            scaler_x = StandardScaler()
            scaler_y = StandardScaler()
            Xs = scaler_x.fit_transform(X)
            Ys = scaler_y.fit_transform(Y)

            # 自动设定 n_components
            n_components = min(Xs.shape[1], Ys.shape[1])
            cca = CCA(n_components=n_components)
            cca.fit(Xs, Ys)
            X_c, Y_c = cca.transform(Xs, Ys)

            # 计算每对典型相关系数（按输出顺序）
            canonical_correlations = []
            for i in range(n_components):
                corr = np.corrcoef(X_c[:, i], Y_c[:, i])[0, 1]
                canonical_correlations.append(corr)

            # 控制台打印（便于调试）
            print("Canonical correlations:", canonical_correlations)

            # 计算 Wilks' lambda + Bartlett 近似检验（修正自由度和编号，防护 lambda 值）
            n, p, q = Xs.shape[0], Xs.shape[1], Ys.shape[1]
            wilks_results = []
            for j in range(n_components):
                lambda_j = np.prod(1 - np.square(canonical_correlations[j:]))
                # 防止数值误差导致 lambda_j <= 0 或 >1
                lambda_j = max(min(lambda_j, 0.9999999999), np.finfo(float).eps)
                df_chi = (p - j) * (q - j)            # 修正自由度
                s = (n - 1 - (p + q + 1) / 2.0)
                chi_stat = -s * np.log(lambda_j)
                # 防护 chi_stat 非负
                if chi_stat < 0:
                    chi_stat = 0.0
                p_value = 1 - stats.chi2.cdf(chi_stat, df_chi)
                wilks_results.append((j + 1, lambda_j, chi_stat, df_chi, p_value))

            print("Wilks results:", wilks_results)

            # 让用户选择保存路径（Word）
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if not save_path:
                self.result_label.config(text=lang["no_save_path_selected"])
                return

            save_dir = pathlib.Path(save_path).parent

            # 写 Word 文档（标题和说明按语言）
            doc = Document()
            doc.add_heading(lang["title"], level=1)
            doc.add_paragraph(f"{'文件' if self.current_language=='zh' else 'File'}: {file_path}")
            doc.add_paragraph(f"{'样本数' if self.current_language=='zh' else 'Samples'}: {n}, X dims: {p}, Y dims: {q}, n_components: {n_components}")
            doc.add_paragraph(f"{'X 变量' if self.current_language=='zh' else 'X variables'}: {', '.join(x_cols)}")
            doc.add_paragraph(f"{'Y 变量' if self.current_language=='zh' else 'Y variables'}: {', '.join(y_cols)}")

            # === 新增：核心设计指标区（载荷 + 权重） ===
            doc.add_heading(lang['core_section_title'], level=2)
            doc.add_paragraph(lang['core_section_desc'])

            # 计算载荷矩阵：每个原始标准化变量与每个典型变量的相关系数
            # Xs: (n_samples, p), X_c: (n_samples, n_components)
            loadings_x = np.zeros((p, n_components))
            for i_var in range(p):
                for j_comp in range(n_components):
                    # 若数值异常或常数列，可能出现 nan，防护处理
                    try:
                        load_val = np.corrcoef(Xs[:, i_var], X_c[:, j_comp])[0, 1]
                        if np.isnan(load_val):
                            load_val = 0.0
                    except Exception:
                        load_val = 0.0
                    loadings_x[i_var, j_comp] = load_val

            loadings_y = np.zeros((q, n_components))
            for i_var in range(q):
                for j_comp in range(n_components):
                    try:
                        load_val = np.corrcoef(Ys[:, i_var], Y_c[:, j_comp])[0, 1]
                        if np.isnan(load_val):
                            load_val = 0.0
                    except Exception:
                        load_val = 0.0
                    loadings_y[i_var, j_comp] = load_val

            # 输出权重表（cca.x_weights_ / cca.y_weights_）
            doc.add_heading(lang['weights_x_title'], level=3)
            table = doc.add_table(rows=1, cols=n_components + 1)
            hdr = table.rows[0].cells
            hdr[0].text = "X变量" if self.current_language == "zh" else "X Variable"
            for i in range(n_components):
                hdr[i + 1].text = f"典型变量{i+1}" if self.current_language == "zh" else f"Canonical{i+1}"
            for i, var in enumerate(x_cols):
                row = table.add_row().cells
                row[0].text = var
                for j in range(n_components):
                    row[j + 1].text = f"{cca.x_weights_[i, j]:.4f}"

            doc.add_heading(lang['weights_y_title'], level=3)
            table = doc.add_table(rows=1, cols=n_components + 1)
            hdr = table.rows[0].cells
            hdr[0].text = "Y变量" if self.current_language == "zh" else "Y Variable"
            for i in range(n_components):
                hdr[i + 1].text = f"典型变量{i+1}" if self.current_language == "zh" else f"Canonical{i+1}"
            for i, var in enumerate(y_cols):
                row = table.add_row().cells
                row[0].text = var
                for j in range(n_components):
                    row[j + 1].text = f"{cca.y_weights_[i, j]:.4f}"

            # 输出载荷表：按每个典型变量，对应原始变量按 |载荷| 降序
            for comp_idx in range(n_components):
                # X 组载荷表（排序）
                doc.add_heading(f"{lang['loading_table_x_title']} - {('第' + str(comp_idx+1) + '个') if self.current_language == 'zh' else 'Canonical ' + str(comp_idx+1)}", level=3)
                # 构造排序数据
                rows = []
                for i_var, varname in enumerate(x_cols):
                    loading = float(loadings_x[i_var, comp_idx])
                    rows.append((varname, loading, abs(loading), classify_loading_degree(abs(loading), lang)))
                # 按绝对值降序
                rows_sorted = sorted(rows, key=lambda x: x[2], reverse=True)
                # 创建表格（列：变量, 载荷, |载荷|, 核心程度）
                table = doc.add_table(rows=1, cols=4)
                hdr = table.rows[0].cells
                for cidx, colname in enumerate(lang['loading_cols']):
                    hdr[cidx].text = colname
                for r in rows_sorted:
                    row = table.add_row().cells
                    row[0].text = str(r[0])
                    row[1].text = f"{r[1]:.4f}"
                    row[2].text = f"{r[2]:.4f}"
                    row[3].text = str(r[3])

                # Y 组载荷表（排序）
                doc.add_heading(f"{lang['loading_table_y_title']} - {('第' + str(comp_idx+1) + '个') if self.current_language == 'zh' else 'Canonical ' + str(comp_idx+1)}", level=3)
                rows = []
                for i_var, varname in enumerate(y_cols):
                    loading = float(loadings_y[i_var, comp_idx])
                    rows.append((varname, loading, abs(loading), classify_loading_degree(abs(loading), lang)))
                rows_sorted = sorted(rows, key=lambda x: x[2], reverse=True)
                table = doc.add_table(rows=1, cols=4)
                hdr = table.rows[0].cells
                for cidx, colname in enumerate(lang['loading_cols']):
                    hdr[cidx].text = colname
                for r in rows_sorted:
                    row = table.add_row().cells
                    row[0].text = str(r[0])
                    row[1].text = f"{r[1]:.4f}"
                    row[2].text = f"{r[2]:.4f}"
                    row[3].text = str(r[3])

            # 添加载荷符号含义解释
            doc.add_heading(lang['loading_sign_explain_title'], level=3)
            doc.add_paragraph(lang['loading_sign_explain'])

            # === 原结果表格部分（典型相关系数、Wilks） ===
            doc.add_heading("典型相关系数" if self.current_language == "zh" else "Canonical correlations", level=2)
            doc.add_paragraph("典型相关系数 r 的绝对值越接近 1，说明 X 组与 Y 组之间的线性关系越强。" if self.current_language == "zh" else "The closer the absolute value of canonical correlation r is to 1, the stronger the linear relationship between X and Y.")
            table = doc.add_table(rows=1, cols=2)
            if self.current_language == "zh":
                table.rows[0].cells[0].text = "对"
                table.rows[0].cells[1].text = "相关系数 (r)"
            else:
                table.rows[0].cells[0].text = "Pair"
                table.rows[0].cells[1].text = "Correlation Coefficient (r)"
            for i, r in enumerate(canonical_correlations, start=1):
                row = table.add_row().cells
                row[0].text = str(i)
                row[1].text = f"{r:.6f}"

            doc.add_heading("Wilks' lambda 检验" if self.current_language == "zh" else "Wilks' lambda test", level=2)
            table = doc.add_table(rows=1, cols=5)
            hdr = table.rows[0].cells
            hdr[0].text = "起始对" if self.current_language == "zh" else "Start at"
            hdr[1].text = "Lambda"
            hdr[2].text = "Chi2"
            hdr[3].text = "df"
            hdr[4].text = "p-value"
            for j, lmbd, chi2, dfc, pv in wilks_results:
                row = table.add_row().cells
                row[0].text = str(j)
                row[1].text = f"{lmbd:.6e}"
                row[2].text = f"{chi2:.4f}"
                row[3].text = str(int(dfc))
                sig = ("显著" if pv < 0.05 else "不显著") if self.current_language == "zh" else ("Significant" if pv < 0.05 else "Not significant")
                row[4].text = f"{pv:.4g} ({sig})"

            # === 图形部分：散点图标注优化（坐标轴标签包含每个典型变量上 |载荷| 最大的 3 个指标与其载荷值） ===
            for i, r in enumerate(canonical_correlations, start=1):
                comp_idx = i - 1

                # 取 X 组前3（按 |载荷|）
                x_rows = []
                for i_var, varname in enumerate(x_cols):
                    loading = float(loadings_x[i_var, comp_idx])
                    x_rows.append((varname, loading, abs(loading)))
                x_top3 = sorted(x_rows, key=lambda x: x[2], reverse=True)[:3]
                x_top3_labels = [f"{name} {val:+.3f}" for name, val, _ in x_top3]
                x_axis_extra = lang['axis_label_format'].format(", ".join(x_top3_labels))

                # 取 Y 组前3
                y_rows = []
                for i_var, varname in enumerate(y_cols):
                    loading = float(loadings_y[i_var, comp_idx])
                    y_rows.append((varname, loading, abs(loading)))
                y_top3 = sorted(y_rows, key=lambda x: x[2], reverse=True)[:3]
                y_top3_labels = [f"{name} {val:+.3f}" for name, val, _ in y_top3]
                y_axis_extra = lang['axis_label_format'].format(", ".join(y_top3_labels))

                fig = plt.figure(figsize=(10, 8))
                plt.scatter(X_c[:, comp_idx], Y_c[:, comp_idx], alpha=0.7)
                # 坐标轴标签包含典型变量及其主要指标
                xlabel_base = f"X canonical var {i}" if self.current_language == "en" else f"X典型变量 {i}"
                ylabel_base = f"Y canonical var {i}" if self.current_language == "en" else f"Y典型变量 {i}"
                plt.xlabel(f"{xlabel_base} {x_axis_extra}")
                plt.ylabel(f"{ylabel_base} {y_axis_extra}")
                plt.title(f"Canonical pair {i} (r={r:.4f})" if self.current_language == "en" else f"典型变量对 {i} (r={r:.4f})")
                plot_file = save_dir / f"canonical_pair_{i}.png"
                plt.tight_layout()
                plt.savefig(str(plot_file.resolve()), dpi=300, bbox_inches='tight')
                plt.close(fig)
                doc.add_paragraph(f"{'Pair' if self.current_language=='en' else '典型变量对'} {i}: r = {r:.6f}")
                try:
                    doc.add_picture(str(plot_file.resolve()), width=Inches(7.5))
                except Exception:
                    doc.add_paragraph(f"[Could not embed image; saved at {plot_file}]")

            # 保存 Word
            doc.save(save_path)

            result_msg = lang["analysis_complete"].format(save_path)
            result_msg += "\n" + lang["images_saved"].format(save_dir)
            self.result_label.config(text=result_msg, wraplength=480)

        except Exception as e:
            # 显示友好错误，同时在控制台输出堆栈，便于调试
            self.result_label.config(text=lang["analysis_error"].format(str(e)))
            print(traceback.format_exc())

    def switch_language(self, event=None):
        self.current_language = "en" if self.current_language == "zh" else "zh"
        lang = languages[self.current_language]
        self.root.title(lang["title"])
        self.select_button.config(text=lang["select_button_text"])
        self.analyze_button.config(text=lang["analyze_button_text"])
        self.open_excel_label.config(text=languages[self.current_language]["open_excel_button_text"])
        self.switch_language_label.config(text=lang["switch_language_button_text"])
        if self.file_entry.get() == "" or self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, lang["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def create_ui(self):
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=languages[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        self.switch_language_label = ttk.Label(frame, text=languages[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        self.root.mainloop()

def run_app():
    app = CanonicalCorrelationAnalysisApp()
    app.run()

if __name__ == "__main__":
    run_app()

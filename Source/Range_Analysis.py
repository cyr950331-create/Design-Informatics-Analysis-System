import tkinter as tk
from tkinter import filedialog, messagebox
import os
import pandas as pd
import numpy as np
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
import openpyxl
import pathlib
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import datetime

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
languages = {
    "zh": {
        "title": "极差分析",
        "select_button_text": "选择文件",
        "file_entry_placeholder": "请输入待分析 Excel 文件的完整路径",
        "analyze_button_text": "分析文件",
        "no_file_selected": "请选择文件。",
        "file_not_exists": "文件不存在，请重新选择。",
        "analysis_error": "分析文件时出错: {}",
        "analysis_complete": "分析完成，结果已保存到 {}，相关图片已保存。",
        "no_save_path_selected": "未选择保存路径，结果未保存。",
        "result_column": "结果列",
        'open_excel_button_text': "示例数据",
        "switch_language_button_text": "中/英",
        "explanation": {
            "极差": "极差反映了数据的离散程度，在极差分析中，极差越大说明该因素对试验结果的影响越大。"
        },
        "interpretation": {
            "极差": "极差越大，表明该因素对试验结果的影响越显著。",
            "均值": "各水平下试验结果的平均值，用于比较不同水平对试验结果的影响。（包含重复实验数据的合并计算）"
        },
        "chart": {
            "title": "极差分析 - 各因素水平均值",
            "x_label": "水平",
            "y_label": "均值"
        }
    },
    "en": {
        "title": "Range Analysis",
        "select_button_text": "Select File",
        "file_entry_placeholder": "Please enter the full path of the Excel file to be analyzed",
        "analyze_button_text": "Analyze File",
        "no_file_selected": "Please select a valid file path.",
        "file_not_exists": "The file does not exist. Please select again.",
        "analysis_error": "An error occurred while analyzing the file: {}",
        "analysis_complete": "Analysis completed. The results have been saved to {}, and the relevant images have been saved.",
        "no_save_path_selected": "No save path selected. The results were not saved.",
        "result_column": "Result Column",
        'open_excel_button_text': "Example data",
        "switch_language_button_text": "Chinese/English",
        "explanation": {
            "Range": "The range reflects the dispersion of the data. In range analysis, a larger range indicates that the factor has a greater influence on the test results."
        },
        "interpretation": {
            "Range": "A larger range indicates that the factor has a more significant influence on the test results.",
            "Average": "The average value of the test results at each level (including merged calculation of repeated experiment data), used to compare the influence of different levels on the test results."
        },
        "chart": {
            "title": "Range Analysis - Mean Values by Factor and Level",
            "x_label": "Level",
            "y_label": "Mean Value"
        }
    }
}


class RangeAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data31.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return

        try:
            # 使用 pandas 读取数据（将首行作为表头）
            df_raw = pd.read_excel(file_path, header=0, engine='openpyxl')
            # 删除全空行与全空列
            df_raw = df_raw.dropna(how='all')
            df_raw = df_raw.dropna(axis=1, how='all')

            if df_raw.shape[1] < 2:
                raise ValueError("数据列数不足，至少需要 1 个因素列和 1 个结果列（最后一列为结果）。")

            # 获取表头信息
            headers = df_raw.columns.tolist()
            # 保留原始 DataFrame 的副本以便在报告中显示水平标签
            df = df_raw.copy()

            # 将最后一列视为结果列，尝试转换为数值并丢弃无法转换的行
            result_col = df.columns[-1]
            df[result_col] = pd.to_numeric(df[result_col], errors='coerce')

            before_rows = df.shape[0]
            df = df.dropna(subset=[result_col])  # 删除结果为空的行
            after_rows = df.shape[0]
            dropped = before_rows - after_rows

            # 如果丢弃了所有行，则报错
            if df.shape[0] == 0:
                raise ValueError("结果列（最后一列）中没有可用的数值数据。")

            num_factors = df.shape[1] - 1  # 最后一列为结果列
            results = df[result_col].to_numpy()

            # 结果数据就绪，开始对每个因素进行水平编码与均值/极差计算
            factor_means = []
            factor_ranges = []
            factor_levels_list = []  # 每个因素的原始水平标签
            factor_level_counts = []
            factor_names = headers[:-1]  # 因素名称（除最后一列）
            result_name = headers[-1]    # 结果列名称

            # 用于生成报告表格
            report_records = []

            for f in range(num_factors):
                col = df.iloc[:, f]
                # 将水平全部转换为字符串以便统一比较（保留原始标签用于报告）
                col_str = col.astype(str)

                # 获取唯一水平（保持出现顺序）
                uniques = pd.unique(col_str)

                # 尝试判断这些水平是否可以按数值排序（例如 ['1','2','3']）
                try:
                    uniques_numeric = pd.to_numeric(uniques, errors='coerce')
                    if not np.any(pd.isna(uniques_numeric)):
                        # 可以全部转换为数值 -> 按数值排序
                        order_idx = np.argsort(uniques_numeric)
                        uniques = uniques[order_idx]
                except Exception:
                    pass

                level_means = []
                valid_levels = []
                for lvl in uniques:
                    mask = (col_str == lvl)
                    lvl_results = results[mask]
                    if len(lvl_results) == 0:
                        # 跳过没有数据的水平
                        continue
                    lvl_mean = float(np.mean(lvl_results))
                    level_means.append(lvl_mean)
                    valid_levels.append(lvl)

                if len(level_means) == 0:
                    # 该因素没有可用水平数据，记录并继续
                    factor_means.append([])
                    factor_ranges.append(0.0)
                    factor_levels_list.append([])
                    factor_level_counts.append(0)
                    report_records.append([f"{factor_names[f]} 无可用水平数据", ""])
                    continue

                rng = float(np.max(level_means) - np.min(level_means))
                factor_means.append(level_means)
                factor_ranges.append(rng)
                factor_levels_list.append(valid_levels)
                factor_level_counts.append(len(valid_levels))

                # 将结果写入 report_records（按“因素 - 水平 - 均值”）
                for idx, mean_val in enumerate(level_means):
                    report_records.append([f"{factor_names[f]} 水平{idx+1} ({valid_levels[idx]}) 均值", round(mean_val, 6)])
                report_records.append([f"{factor_names[f]} 极差", round(rng, 6)])

            # 将 report_records 转为 DataFrame 以便调试/导出
            report_df = pd.DataFrame(report_records, columns=["统计量", "值"])

            # 生成图表（单图，多个因素的不同水平均值曲线）
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            base_plot_path = os.path.splitext(file_path)[0] + f"_range_plot_{timestamp}.png"
            plt.figure(figsize=(10, 6))

            max_levels = max([len(lv) for lv in factor_levels_list]) if len(factor_level_counts) > 0 else 0
            for i in range(len(factor_means)):
                if len(factor_means[i]) == 0:
                    continue
                x = np.arange(1, len(factor_means[i]) + 1)
                # 使用因素表头作为图例
                plt.plot(x, factor_means[i], marker='o', label=factor_names[i])
                # 在每个点上标注原始水平标签
                for xi, yi, lbl in zip(x, factor_means[i], factor_levels_list[i]):
                    plt.text(xi, yi, str(lbl), fontsize=8, ha='center', va='bottom', rotation=0)

            plt.title(f"{languages[self.current_language]['chart']['title']} ({result_name})")
            plt.xlabel(languages[self.current_language]["chart"]["x_label"])
            plt.ylabel(languages[self.current_language]["chart"]["y_label"])
            plt.grid(True)
            plt.legend()
            plt.tight_layout()
            plt.savefig(base_plot_path, bbox_inches='tight')
            plt.close()

            # 让用户选择保存路径（Word 文档）
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if not save_path:
                self.result_label.config(text=languages[self.current_language]["no_save_path_selected"])
                return

            # 创建 Word 文档并写入结果
            doc = Document()
            doc.add_heading(languages[self.current_language]["title"], level=1)
            doc.add_paragraph(f"{languages[self.current_language]['result_column']}: {result_name}")

            # 如果在读取时丢弃了部分行，写明提示
            if dropped > 0:
                doc.add_paragraph(f"注意: 输入数据中有 {dropped} 行由于结果列无法转换为数值而被忽略。")

            # 写入每个因素的水平均值与极差（分表）
            for i in range(len(factor_means)):
                # 使用因素表头作为标题
                doc.add_heading(factor_names[i], level=2)
                levels = factor_levels_list[i]
                means = factor_means[i]
                if len(levels) == 0:
                    doc.add_paragraph("无可用水平数据。")
                    continue

                table = doc.add_table(rows=1 + len(levels) + 1, cols=2)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                # 表头
                table.cell(0,0).text = ( "水平" if self.current_language=='zh' else "Level" )
                table.cell(0,1).text = ( "均值" if self.current_language=='zh' else "Mean" )
                # 内容
                for r, (lvl, mv) in enumerate(zip(levels, means), start=1):
                    table.cell(r,0).text = str(lvl)
                    table.cell(r,1).text = str(round(mv,6))
                    table.cell(r,0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    table.cell(r,1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                # 极差行
                table.cell(len(levels)+1, 0).text = ( "极差" if self.current_language=='zh' else "Range" )
                table.cell(len(levels)+1, 1).text = str(round(factor_ranges[i],6))
                table.cell(len(levels)+1,0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                table.cell(len(levels)+1,1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                # 尝试设置列宽
                try:
                    for col in table.columns:
                        for cell in col.cells:
                            cell.width = Cm(4)
                except Exception:
                    pass

            # 添加总体解释与结果解读
            doc.add_heading(( "解释说明" if self.current_language=='zh' else "Explanation" ), level=2)
            exp_text = languages[self.current_language]['explanation'].get("极差", "")
            doc.add_paragraph(exp_text)

            doc.add_heading(( "结果解读" if self.current_language=='zh' else "Interpretation" ), level=2)
            interp = languages[self.current_language]['interpretation']
            # 将每一条解读作为段落写入
            for k, v in interp.items():
                doc.add_paragraph(f"{k} : {v}")

            # 插入图表
            doc.add_heading(( "图表" if self.current_language=='zh' else "Plot" ), level=2)
            try:
                doc.add_picture(base_plot_path, width=Inches(6))
            except Exception as e:
                doc.add_paragraph(f"无法插入图像: {e}")

            # 保存文档
            doc.save(save_path)

            # 最终通知用户（界面上）
            msg = languages[self.current_language]['analysis_complete'].format(save_path)
            self.result_label.config(text=msg, wraplength=400)

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])
        self.open_excel_label.config(text=languages[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=languages[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame,
                                               text=languages[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = RangeAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()
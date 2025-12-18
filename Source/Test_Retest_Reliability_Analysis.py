import tkinter as tk
from tkinter import filedialog
import os
import pandas as pd
import numpy as np
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import pingouin as pg

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'  # 设置字体为黑体，可根据系统情况修改为其他支持中文的字体
plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

# 定义语言字典
languages = {
    'zh': {
        'title': "重测信度",
        'select_button_text': "选择文件",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'analyze_button_text': "分析文件",
        'no_file_selected': "请选择文件。",
        'file_not_exists': "文件不存在，请重新选择。",
        'analysis_error': "分析文件时出错: {}",
        'analysis_complete': "分析完成，结果已保存到 {}，相关图片已保存。",
        'no_save_path_selected': "未选择保存路径，结果未保存。",
        'open_excel_button_text': "示例数据",
        'switch_language_button_text': "中/英",
        'explanation': {
            "ICC (2,1) 重测信度系数": "用同一种测验在不同时间对同一组被试进行两次测量，两次测量结果的相关系数。",
            "样本量": "每个样本中的观测值数量。",
            "均值": "样本数据的平均值。"
        },
        'interpretation': {
            "ICC (2,1) 重测信度系数": "重测信度系数越接近1，表示测验结果越稳定；越接近0，表示测验结果的稳定性越差。",
            "样本量": "样本量的大小会影响统计检验的稳定性，较大的样本量通常能提供更可靠的结果。",
            "均值": "均值反映了数据的平均水平，可用于比较不同变量的集中趋势。"
        },
        'plot_titles': {
            'scatter': '两次测量散点图',
            'bland_altman': 'Bland-Altman图',
            'boxplot': '两次测量箱线图对比',
            'mean_bar': '变量均值柱状图'
        },
        'axis_labels': {
            'x_scatter': '第一次测量值',
            'y_scatter': '第二次测量值',
            'x_bland': '两次测量均值',
            'y_bland': '两次测量差值',
            'y_boxplot': '测量值',
            'x_mean': '变量',
            'y_mean': '均值'
        },
        'table_headers': {
            'overall': '总体分析结果',
            'icc': '各变量组内相关系数 ICC(2,1)',
            'explanation': '解释说明',
            'interpretation': '结果解读'
        }
    },
    'en': {
        'title': "Test Retest Reliability",
        'select_button_text': "Select File",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'analyze_button_text': "Analyze File",
        'no_file_selected': "Please select a valid file path.",
        'file_not_exists': "The file does not exist. Please select again.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'analysis_complete': "Analysis completed. The results have been saved to {}, and the relevant images have been saved.",
        'no_save_path_selected': "No save path selected. The results were not saved.",
        'open_excel_button_text': "Example data",
        'switch_language_button_text': "Chinese/English",
        'explanation': {
            "Test-Retest Reliability Coefficient": "The correlation coefficient between the results of two measurements of the same group of subjects using the same test at different times.",
            "Sample Size": "The number of observations in each sample.",
            "Mean": "The average value of the sample data."
        },
        'interpretation': {
            "Test-Retest Reliability Coefficient": "The closer the test-retest reliability coefficient is to 1, the more stable the test results are; the closer it is to 0, the worse the stability of the test results.",
            "Sample Size": "The sample size affects the stability of the statistical test. A larger sample size usually provides more reliable results.",
            "Mean": "The mean reflects the average level of the data and can be used to compare the central tendencies of different variables."
        },
        'plot_titles': {
            'scatter': 'Test-Retest Scatter Plot',
            'bland_altman': 'Bland-Altman Plot',
            'boxplot': 'Boxplot Comparison of Two Measurements',
            'mean_bar': 'Bar Chart of Variable Means'
        },
        'axis_labels': {
            'x_scatter': 'First Measurement',
            'y_scatter': 'Second Measurement',
            'x_bland': 'Mean of Two Measurements',
            'y_bland': 'Difference Between Measurements',
            'y_boxplot': 'Measurement Value',
            'x_mean': 'Variables',
            'y_mean': 'Mean'
        },
        'table_headers': {
            'overall': 'Overall Analysis Results',
            'icc': 'Intraclass Correlation Coefficient ICC(2,1) for Each Variable',
            'explanation': 'Explanation',
            'interpretation': 'Interpretation'
        }
    }
}


class TestRetestReliabilityAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data12.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def test_retest_reliability(self, data1, data2):
        """
        使用组内相关系数 ICC(2,1) 计算重测信度。
        假设前半部分列为第一次测量，后半部分列为第二次测量。
        """
        results = []
        n_vars = min(data1.shape[1], data2.shape[1])

        for i in range(n_vars):
            var1 = data1.columns[i]
            var2 = data2.columns[i]
            # 新代码：用实际索引作为样本ID，确保两次测量行匹配
            # 提取当前变量的有效数据（过滤单列缺失值）
            t1_score = data1[var1].dropna()
            t2_score = data2[var2].dropna()
            # 取当前变量的共同样本（避免单测有值、复测空值）
            common_subj = t1_score.index.intersection(t2_score.index)
            if len(common_subj) < 2:
                results.append((var1, np.nan, "有效样本<2", len(common_subj)))
                continue

            # 构建数据框（长度必然一致，因用共同索引）
            df_t1 = pd.DataFrame({
                "subject": common_subj,  # 用原始索引，确保样本对应
                "rater": "T1",
                "score": t1_score[common_subj]
            })
            df_t2 = pd.DataFrame({
                "subject": common_subj,
                "rater": "T2",
                "score": t2_score[common_subj]
            })
            df_pair = pd.concat([df_t1, df_t2], ignore_index=True)  # 合并后长度一致

            if df_pair["score"].nunique() < 2:
                # 无方差，无法计算 ICC
                results.append((var1, np.nan, np.nan, len(df_pair) // 2))
                continue

            try:
                icc = pg.intraclass_corr(data=df_pair, targets="subject", raters="rater",
                                         ratings="score").round(4)
                icc_row = icc.loc[icc['Type'] == 'ICC2']
                icc_value = float(icc_row["ICC"].values[0])
                ci95 = f"[{icc_row['CI95%'].values[0]}]"
                results.append((var1, icc_value, ci95, len(df_pair) // 2))
            except Exception as e:
                results.append((var1, np.nan, str(e), len(df_pair) // 2))

        if self.current_language == "zh":
            # 中文列名
            columns = ["变量名", "ICC(2,1)", "95%置信区间", "样本量(n)"]
        else:
            # 英文列名
            columns = ["Variable Name", "ICC(2,1)", "95% Confidence Interval", "Sample Size (n)"]

        # 创建结果DataFrame（使用语言对应的列名）
        result_df = pd.DataFrame(results, columns=columns)
        return result_df

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 检查数据是否为数值类型
            numerical_df = df.select_dtypes(include=[np.number])
            if numerical_df.empty:
                raise ValueError("数据中没有数值列，无法进行重测信度分析。")

            # 步骤1：校验列数为偶数（适配“前后半列”格式）
            total_cols = len(numerical_df.columns)
            if total_cols % 2 != 0:
                raise ValueError(f"数据列数为{total_cols}（奇数），无法按“前后半列”分割，请确保列数为偶数（如图片中8列）。")

            # 步骤2：按前后半列分割（保留原逻辑，匹配图片格式）
            half = total_cols // 2
            data1 = numerical_df.iloc[:, :half].copy()  # 前半列：测试1-4
            data2 = numerical_df.iloc[:, half:].copy()  # 后半列：重测1-4

            # 步骤3：强制对齐行数（核心！解决长度不一致）
            # 删除两次测量数据中的全空行（避免空行导致行数差异）
            data1 = data1.dropna(how="all")  # 仅删除“所有列都为空”的行
            data2 = data2.dropna(how="all")
            # 取两次测量的共同索引（确保行数完全一致）
            common_index = data1.index.intersection(data2.index)
            data1 = data1.loc[common_index]  # 仅保留两次测量都有数据的行
            data2 = data2.loc[common_index]

            # 步骤4：校验最终行数（避免无有效数据）
            if len(data1) == 0 or len(data2) == 0:
                raise ValueError("数据中无有效行（可能全为空），无法进行重测信度分析。")
            if len(data1) != len(data2):
                raise ValueError(
                    f"第一次测量有效行数（{len(data1)}）与第二次测量有效行数（{len(data2)}）仍不相等，请手动删除空行后重试。")

            # 计算样本量和均值（用于描述和绘图）
            sample_sizes = numerical_df.count()
            means = numerical_df.mean()

            # 进行 ICC(2,1) 重测信度分析
            reliability_df = self.test_retest_reliability(data1, data2)

            # 计算总体平均 ICC
            mean_icc = reliability_df["ICC(2,1)"].mean()

            # 根据当前语言定义文本
            if self.current_language == "zh":
                # 中文文本
                stats_name1 = "平均ICC(2,1)"
                stats_name2 = "平均样本量"
                headers = ["统计量", "统计量值"]
            else:
                # 英文文本
                stats_name1 = "Average ICC(2,1)"
                stats_name2 = "Average Sample Size"
                headers = ["Statistic", "Value"]

            # 整理数据并创建DataFrame
            sample_size_col = "样本量(n)" if self.current_language == "zh" else "Sample Size (n)"
            data = [
                [stats_name1, round(mean_icc, 4)],
                [stats_name2, reliability_df[sample_size_col].mean()]
            ]
            df_result = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            explanations = languages[self.current_language]['explanation']
            interpretations = languages[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            if self.current_language == "zh":
                # 中文列名
                columns = ["ICC (2,1) 重测信度系数", "样本量", "均值"]
            else:
                # 英文列名
                columns = ["ICC (2,1) Test-Retest Reliability Coefficient", "Sample Size", "Mean"]

            # 重新索引列名
            explanation_df = explanation_df.reindex(columns=columns)
            explanation_df.insert(0, "统计量_解释说明", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(columns=["ICC (2,1) 重测信度系数", "样本量", "均值"])
            interpretation_df.insert(0, "统计量_结果解读", "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加总体统计结果
                doc.add_heading(languages[self.current_language]['table_headers']['overall'], level=1)
                table = doc.add_table(rows=1, cols=len(headers))
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                for _, row in df_result.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(round(value, 4) if isinstance(value, float) else str(value))

                # 添加逐变量的重测信度结果表
                doc.add_heading(languages[self.current_language]['table_headers']['icc'], level=1)
                table = doc.add_table(rows=1, cols=len(reliability_df.columns))
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(reliability_df.columns):
                    hdr_cells[i].text = header
                for _, row in reliability_df.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(round(value, 4) if isinstance(value, float) else str(value))

                # 添加解释说明表格
                doc.add_heading(languages[self.current_language]['table_headers']['explanation'], level=1)
                explanation_para = doc.add_paragraph()
                for key, value in explanations.items():
                    explanation_para.add_run(f'• {key}: {value}\n')

                # 结果解读部分（改为项目符号列表）
                doc.add_heading(languages[self.current_language]['table_headers']['interpretation'], level=1)
                interpretation_para = doc.add_paragraph()
                for key, value in interpretations.items():
                    interpretation_para.add_run(f'• {key}: {value}\n')

                # 生成图片（均值柱状图）
                fig, ax = plt.subplots(figsize=(10, 8))
                means.plot(kind='bar', ax=ax)
                ax.set_title(languages[self.current_language]['plot_titles']['mean_bar'])
                ax.set_xlabel(languages[self.current_language]['axis_labels']['x_mean'])
                ax.set_ylabel(languages[self.current_language]['axis_labels']['y_mean'])
                ax.tick_params(axis='x', rotation=0)  # 关键修改：0度即水平方向
                # 保存图片
                img_path = os.path.splitext(save_path)[0] + '.png'
                plt.savefig(img_path)
                plt.close()

                # 在 Word 文档中插入图片
                doc.add_heading(languages[self.current_language]['plot_titles']['mean_bar'], level=1)
                doc.add_picture(img_path, width=Inches(8))

                # 生成两次测量的散点图
                plt.figure(figsize=(10, 6))
                for i in range(min(data1.shape[1], data2.shape[1])):
                    if self.current_language == "zh":
                        label_text = f'变量{i + 1}'
                    else:
                        label_text = f'Variable {i + 1}'
                    plt.scatter(data1.iloc[:, i], data2.iloc[:, i], label=label_text)
                plt.xlabel(languages[self.current_language]['axis_labels']['x_scatter'])
                plt.ylabel(languages[self.current_language]['axis_labels']['y_scatter'])
                plt.title(languages[self.current_language]['plot_titles']['scatter'])
                plt.legend()
                scatter_img_path = os.path.splitext(save_path)[0] + '_scatter.png'
                plt.savefig(scatter_img_path)
                plt.close()

                # 在Word中插入
                doc.add_heading(languages[self.current_language]['plot_titles']['scatter'], level=1)
                doc.add_picture(scatter_img_path, width=Inches(8))

                # 生成Bland-Altman图
                plt.figure(figsize=(10, 6))
                for i in range(min(data1.shape[1], data2.shape[1])):
                    mean_vals = (data1.iloc[:, i] + data2.iloc[:, i]) / 2
                    diffs = data1.iloc[:, i] - data2.iloc[:, i]
                    if self.current_language == "zh":
                        label_text = f'变量{i + 1}'  # 中文：变量1、变量2...
                    else:
                        label_text = f'Variable {i + 1}'  # 英文：Variable 1、Variable 2...

                    plt.scatter(mean_vals, diffs, label=label_text)
                    mean_diff = np.mean(diffs)
                    if self.current_language == "zh":
                        # 中文标签：平均差值: X.XX
                        label_text = f'平均差值: {mean_diff:.2f}'
                    else:
                        # 英文标签：Mean Difference: X.XX
                        label_text = f'Mean Difference: {mean_diff:.2f}'

                    plt.axhline(mean_diff, color='r', linestyle='--', label=label_text)
                    plt.axhline(mean_diff + 1.96 * np.std(diffs), color='g', linestyle=':')
                    plt.axhline(mean_diff - 1.96 * np.std(diffs), color='g', linestyle=':')
                plt.xlabel(languages[self.current_language]['axis_labels']['x_bland'])
                plt.ylabel(languages[self.current_language]['axis_labels']['y_bland'])
                plt.title(languages[self.current_language]['plot_titles']['bland_altman'])
                plt.legend()
                bland_altman_path = os.path.splitext(save_path)[0] + '_bland_altman.png'
                plt.savefig(bland_altman_path)
                plt.close()

                # 在Word中插入
                doc.add_heading(languages[self.current_language]['plot_titles']['bland_altman'], level=1)
                doc.add_picture(bland_altman_path, width=Inches(8))

                # 生成箱线图
                plt.figure(figsize=(12, 10))
                combined_data = []
                labels = []
                for i in range(min(data1.shape[1], data2.shape[1])):
                    combined_data.append(data1.iloc[:, i].dropna())
                    combined_data.append(data2.iloc[:, i].dropna())
                    if self.current_language == "zh":
                        var_text = "变量"
                        first_text = "第一次"
                        second_text = "第二次"
                    else:
                        var_text = "Variable"
                        first_text = "First"
                        second_text = "Second"
                    labels.append(f'{var_text}{i + 1}_{first_text}')
                    labels.append(f'{var_text}{i + 1}_{second_text}')
                plt.boxplot(combined_data, tick_labels=labels, patch_artist=True)
                plt.xticks(rotation=45)
                plt.ylabel(languages[self.current_language]['axis_labels']['y_boxplot'])
                plt.title(languages[self.current_language]['plot_titles']['boxplot'])
                boxplot_path = os.path.splitext(save_path)[0] + '_boxplot.png'
                plt.savefig(boxplot_path)
                plt.close()

                # 在Word中插入
                doc.add_heading(languages[self.current_language]['plot_titles']['boxplot'], level=1)
                doc.add_picture(boxplot_path, width=Inches(8))

                # 保存 Word 文档
                doc.save(save_path)

                self.result_label.config(text=languages[self.current_language]["analysis_complete"].format(save_path),
                                         wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]["no_save_path_selected"])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]["analysis_error"].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])
        self.open_excel_label.config(text=languages[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=languages[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame,
                                               text=languages[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = TestRetestReliabilityAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()
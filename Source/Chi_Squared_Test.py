import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
from scipy import stats
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches


# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
LANGUAGES = {
    'zh': {
        'title': "卡方检验",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "请选择文件。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "中/英",
        'open_excel_button_text': "示例数据",
        'data_type_error': "数据中包含非数值类型，请检查Excel文件。",
        'insufficient_data': "数据不足，至少需要2行2列（包括列名和行标签）。",
        'non_integer_error': "列联表中的值必须是整数频数。",
        'negative_values_error': "列联表中不能有负数。",
        'all_zero_error': "列联表中不能全为零。",
        'fisher_test': 'Fisher卡方',
        'test_type_label': '检验类型',
        'sample_size_label': '样本量',
        'chi2_value_label': '卡方值',
        'dof_label': '自由度',
        'p_value_label': 'p值',
        'significance_label': '显著性（α=0.05）',
        'significant': '显著',
        'not_significant': '不显著',
        'odds_ratio_label': '优势比',
        'odds_ratio_explanation': '优势比大于1表示暴露因素与结果正相关，小于1表示负相关。',
        'test_types': {
            'pearson': "Pearson卡方",
            'yates': "Yates校正卡方",
            'likelihood': "似然比卡方",
            'fisher': "Fisher卡方"
        },
        'explanation': {
            "Pearson卡方": "当2*2列联表中n>=40且所有期望频数E>=5时使用，衡量实际频数与理论频数的差异。",
            "Yates校正卡方": "当2*2列联表中n>=40但有一个格子的期望频数满足1<=E<5时使用，对Pearson卡方的校正。",
            "似然比卡方": "当R*C列联表中期望频数不满足使用Pearson卡方的条件时使用。",
            "Fisher卡方": "当2*2列联表中任何一格子出现E<1或n<40时使用。",
            "Phi系数": "用于衡量2*2列联表的效应量。",
            "Cramer's V": "用于衡量R*C列联表的效应量。"
        },
        'interpretation': {
            "卡方值": "卡方值越大，说明实际频数与理论频数之间的差异越大。",
            "p值": "p值小于显著性水平（通常为0.05）时，拒绝原假设，认为变量之间存在显著关联；否则，接受原假设，认为变量之间无显著关联。",
            "自由度": "自由度反映了数据的独立变化程度，用于计算卡方分布的临界值。",
            "显著性（α=0.05）": "表示在0.05的显著性水平下，变量之间是否存在显著关联。",
            "Phi系数": "Phi系数的绝对值越接近1，说明2*2列联表中两个变量之间的关联越强。",
            "Cramer's V": "Cramer's V的值越接近1，说明R*C列联表中两个变量之间的关联越强。"
        }
    },
    'en': {
        'title': "Chi Square Test",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Chinese/English",
        'open_excel_button_text': "Example data",
        'data_type_error': "The data contains non-numeric types, please check the Excel file.",
        'insufficient_data': "Insufficient data, at least 2 rows and 2 columns (including headers and row labels) are required.",
        'non_integer_error': "Values in the contingency table must be integer frequencies.",
        'negative_values_error': "Contingency table cannot contain negative values.",
        'all_zero_error': "Contingency table cannot contain all zeros.",
        'fisher_test': "Fisher's Exact Test",
        'test_type_label': 'Test Type',
        'sample_size_label': 'Sample Size',
        'chi2_value_label': 'Chi-square Value',
        'dof_label': 'Degrees of Freedom',
        'p_value_label': 'p-value',
        'significance_label': 'Significance (α=0.05)',
        'significant': 'Significant',
        'not_significant': 'Not Significant',
        'odds_ratio_label': 'Odds Ratio',
        'odds_ratio_explanation': 'An odds ratio greater than 1 indicates a positive correlation between the exposure factor and the outcome; less than 1 indicates a negative correlation.',
        'test_types': {
            'pearson': "Pearson Chi-square",
            'yates': "Yates' Corrected Chi-square",
            'likelihood': "Likelihood Ratio Chi-square",
            'fisher': "Fisher's Exact Test"
        },
        'explanation': {
            "Pearson Chi-square": "Used when n>=40 and all expected frequencies E>=5 in a 2*2 contingency table, measuring the difference between observed and expected frequencies.",
            "Yates' Corrected Chi-square": "Used when n>=40 but there is one cell with 1<=E<5 in a 2*2 contingency table, a correction to the Pearson chi-square.",
            "Likelihood Ratio Chi-square": "Used when the expected frequencies in an R*C contingency table do not meet the conditions for using the Pearson chi-square.",
            "Fisher's Exact Test": "Used when any cell has E<1 or n<40 in a 2*2 contingency table.",
            "Phi Coefficient": "Used to measure the effect size of a 2*2 contingency table.",
            "Cramer's V": "Used to measure the effect size of an R*C contingency table."
        },
        'interpretation': {
            "Chi-square Value": "A larger chi-square value indicates a greater difference between the observed and expected frequencies.",
            "p-value": "When the p-value is less than the significance level (usually 0.05), the null hypothesis is rejected, indicating a significant association between variables; otherwise, the null hypothesis is accepted, indicating no significant association.",
            "Degrees of Freedom": "The degrees of freedom reflect the independent variation of the data and are used to calculate the critical value of the chi-square distribution.",
            "Significance (α=0.05)": "Indicates whether there is a significant association between variables at the 0.05 significance level.",
            "Phi Coefficient": "The closer the absolute value of the Phi coefficient is to 1, the stronger the association between the two variables in the 2*2 contingency table.",
            "Cramer's V": "The closer the value of Cramer's V is to 1, the stronger the association between the two variables in the R*C contingency table."
        }
    }
}


class ChiSquaredTestApp:
    def __init__(self, root=None):
        # 当前语言
        self.current_language = 'en'

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]['title'])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]['title'])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data40.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def calculate_effect_size(self, observed, chi2, n):
        """计算效应量（Phi系数或Cramer's V），确保类型名称与语言字典匹配"""
        r, c = observed.shape  # 获取列联表的行数和列数
        df = min(r - 1, c - 1)  # 用于Cramer's V计算的自由度

        if r == 2 and c == 2:
            # 2x2列联表：计算Phi系数，类型名称严格匹配中文/英文字典的键
            phi = np.sqrt(chi2 / n)
            effect_type = "Phi系数" if self.current_language == "zh" else "Phi Coefficient"
            return {"类型": effect_type, "值": phi}
        else:
            # 大于2x2列联表：计算Cramer's V，类型名称严格匹配字典的键
            cramer_v = np.sqrt(chi2 / (n * df))
            effect_type = "Cramer's V"
            return {"类型": effect_type, "值": cramer_v}

    def select_test_type(self, observed):
        """根据数据特征自动选择合适的检验类型（使用语言字典中的键）"""
        r, c = observed.shape
        n = observed.sum()
        chi2, _, _, expected = stats.chi2_contingency(observed, correction=False)

        # 检查期望频数
        has_small_expected = (expected < 1).any()
        has_moderate_expected = ((expected >= 1) & (expected < 5)).any()
        total_small = ((expected < 5).sum() / expected.size) > 0.2

        # 从语言字典获取检验类型名称（确保与当前语言匹配）
        lang = LANGUAGES[self.current_language]['explanation']
        pearson = list(lang.keys())[0]  # Pearson卡方/Pearson Chi-square
        yates = list(lang.keys())[1]  # Yates校正卡方/Yates' Corrected Chi-square
        likelihood = list(lang.keys())[2]  # 似然比卡方/Likelihood Ratio Chi-square
        fisher = LANGUAGES[self.current_language]['fisher_test']

        # 2x2列联表的情况
        if r == 2 and c == 2:
            if n < 40 or has_small_expected:
                return fisher
            elif has_moderate_expected:
                return yates
            else:
                return pearson

        # RxC列联表的情况
        else:
            if has_small_expected or total_small:
                return likelihood
            else:
                return pearson

    def perform_test(self, observed, test_type):
        """执行选定的卡方检验（使用语言字典中的键进行判断）"""
        n = observed.sum()
        # 从语言字典动态获取各检验类型的键名（替代硬编码）
        lang_test_types = LANGUAGES[self.current_language]['test_types']

        # 动态赋值（中英文自动匹配）
        pearson = lang_test_types['pearson']  # 中文："Pearson卡方"；英文："Pearson Chi-square"
        yates = lang_test_types['yates']  # 中文："Yates校正卡方"；英文："Yates' Corrected Chi-square"
        likelihood = lang_test_types['likelihood']  # 中文："似然比卡方"；英文："Likelihood Ratio Chi-square"
        fisher = lang_test_types['fisher']  # 中文："Fisher卡方"；英文："Fisher's Exact Test"

        # 根据当前语言的检验类型键进行判断
        if test_type == pearson:
            chi2, p, dof, expected = stats.chi2_contingency(observed, correction=False)
            return {
                "test_type": test_type,
                "chi2_value": chi2,
                "p_value": p,
                "dof": dof,
                "expected": expected,
                "sample_size": n
            }
        elif test_type == yates:
            chi2, p, dof, expected = stats.chi2_contingency(observed, correction=True)
            return {
                "test_type": test_type,
                "chi2_value": chi2,
                "p_value": p,
                "dof": dof,
                "expected": expected,
                "sample_size": n
            }
        elif test_type == likelihood:
            chi2, p, dof, expected = stats.chi2_contingency(observed, lambda_="log-likelihood")
            return {
                "test_type": test_type,
                "chi2_value": chi2,
                "p_value": p,
                "dof": dof,
                "expected": expected,
                "sample_size": n
            }
        elif test_type == fisher:
            odds_ratio, p = stats.fisher_exact(observed)
            return {
                "test_type": test_type,
                "p_value": p,
                "odds_ratio": odds_ratio,
                "sample_size": n
            }

    # 生成图表函数
    def generate_plots(self, observed, expected, row_labels, column_names, save_path):
        """生成实际频数与期望频数对比图"""
        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 6))
        fig.suptitle('实际频数与期望频数对比' if self.current_language == 'zh' else 'Observed vs Expected Frequencies')

        # 实际频数热图
        im1 = ax1.imshow(observed, cmap='Blues')
        ax1.set_title('实际频数' if self.current_language == 'zh' else 'Observed Frequencies')
        ax1.set_xticks(range(len(column_names)))
        ax1.set_yticks(range(len(row_labels)))
        ax1.set_xticklabels(column_names, rotation=45)
        ax1.set_yticklabels(row_labels)

        # 添加数值标签
        for i in range(len(row_labels)):
            for j in range(len(column_names)):
                text = ax1.text(j, i, observed[i, j], ha='center', va='center', color='black')

        # 期望频数热图
        im2 = ax2.imshow(expected, cmap='Greens')
        ax2.set_title('期望频数' if self.current_language == 'zh' else 'Expected Frequencies')
        ax2.set_xticks(range(len(column_names)))
        ax2.set_yticks(range(len(row_labels)))
        ax2.set_xticklabels(column_names, rotation=45)
        ax2.set_yticklabels(row_labels)

        # 添加数值标签
        for i in range(len(row_labels)):
            for j in range(len(column_names)):
                text = ax2.text(j, i, f"{expected[i, j]:.1f}", ha='center', va='center', color='black')

        plt.tight_layout()
        plot_path = os.path.splitext(save_path)[0] + '.png'
        plt.savefig(plot_path, dpi=300, bbox_inches='tight')
        plt.close()
        return plot_path

    def analyze_file(self):
        file_path = self.file_entry.get()
        placeholder_zh = "请输入待分析 Excel 文件的完整路径"
        placeholder_en = "Please enter the full path of the Excel file to be analyzed"
        if not file_path or file_path == placeholder_zh or file_path == placeholder_en:
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return

        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return

        try:
            # 读取Excel文件，第一行为列名
            df = pd.read_excel(file_path, header=0)

            # 提取行标签（第一列）和列名
            row_labels = df.iloc[:, 0].tolist()  # 第一列作为行标签
            column_names = df.columns.tolist()[1:]  # 从第二列开始作为列名

            # 提取数据部分（从第二列开始）
            data_df = df.iloc[:, 1:].copy()

            # 尝试转换数据为数值类型
            try:
                data_df = data_df.apply(pd.to_numeric)
            except ValueError:
                self.result_label.config(text=LANGUAGES[self.current_language]['data_type_error'])
                return

            observed = data_df.values

            # 验证数据有效性
            if observed.ndim != 2:
                raise ValueError(LANGUAGES[self.current_language]['insufficient_data'])
            if observed.shape[0] < 1 or observed.shape[1] < 2:
                raise ValueError(LANGUAGES[self.current_language]['insufficient_data'])
            if (observed < 0).any():
                raise ValueError(LANGUAGES[self.current_language]['negative_values_error'])
            if observed.sum() == 0:
                raise ValueError(LANGUAGES[self.current_language]['all_zero_error'])
            # 确保所有值都是整数（频数不能是小数）
            if not np.allclose(observed, observed.astype(int)):
                raise ValueError(LANGUAGES[self.current_language]['non_integer_error'])

            # 自动选择检验类型
            test_type = self.select_test_type(observed)

            # 执行检验
            test_result = self.perform_test(observed, test_type)

            # 计算效应量（Fisher检验除外）
            effect_size = None
            current_fisher = LANGUAGES[self.current_language]['fisher_test']
            if test_type != current_fisher:
                effect_size = self.calculate_effect_size(
                    observed,
                    test_result["chi2_value"],
                    test_result["sample_size"]
                )

            # 保存结果
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建Word文档并保存结果
                doc = Document()
                doc.add_heading('卡方检验结果' if self.current_language == 'zh' else 'Chi-Square Test Results', 0)

                # 添加变量名称信息
                doc.add_heading('变量名称' if self.current_language == 'zh' else 'Variable Names', level=1)
                row_label_text = "行标签" if self.current_language == 'zh' else "Row Labels"
                col_label_text = "列名" if self.current_language == 'zh' else "Column Names"
                doc.add_paragraph(f"{row_label_text}: {', '.join(map(str, row_labels))}")
                doc.add_paragraph(f"{col_label_text}: {', '.join(column_names)}")

                # 添加原始数据
                doc.add_heading('原始数据' if self.current_language == 'zh' else 'Original Data', level=1)
                table = doc.add_table(rows=observed.shape[0] + 1, cols=observed.shape[1] + 1)

                # 添加表头（第一列是行标签标题）
                table.cell(0, 0).text = row_label_text
                for j, col_name in enumerate(column_names):
                    if j < observed.shape[1]:
                        table.cell(0, j + 1).text = str(col_name)

                # 添加数据（第一列是行标签）
                for i in range(observed.shape[0]):
                    table.cell(i + 1, 0).text = str(row_labels[i])
                    for j in range(observed.shape[1]):
                        table.cell(i + 1, j + 1).text = str(observed[i, j])

                # 添加检验结果
                doc.add_heading('检验结果' if self.current_language == 'zh' else 'Test Results', level=1)
                lang = LANGUAGES[self.current_language]
                # 使用统一的键名访问检验结果，修复语言切换导致的键名不匹配问题
                doc.add_paragraph(f"{lang['test_type_label']}: {test_result['test_type']}")
                doc.add_paragraph(f"{lang['sample_size_label']}: {test_result['sample_size']}")

                # 非Fisher检验结果
                if test_type != current_fisher:
                    doc.add_paragraph(f"{lang['chi2_value_label']}: {test_result['chi2_value']:.4f}")
                    doc.add_paragraph(f"{lang['dof_label']}: {test_result['dof']}")
                    doc.add_paragraph(f"{lang['p_value_label']}: {test_result['p_value']:.6f}")
                    # 显著性判断
                    significance_text = lang['significant'] if test_result['p_value'] < 0.05 else lang['not_significant']
                    doc.add_paragraph(f"{lang['significance_label']}: {significance_text}")

                    # 添加效应量
                    if effect_size:
                        doc.add_paragraph(f"{effect_size['类型']}: {effect_size['值']:.4f}")
                        doc.add_paragraph(LANGUAGES[self.current_language]['explanation'][effect_size['类型']])

                    # 添加期望频数
                    doc.add_heading('期望频数' if self.current_language == 'zh' else 'Expected Frequencies', level=2)
                    exp_table = doc.add_table(
                        rows=test_result['expected'].shape[0] + 1,
                        cols=test_result['expected'].shape[1] + 1
                    )

                    # 期望频数表格表头
                    exp_table.cell(0, 0).text = row_label_text
                    for j, col_name in enumerate(column_names):
                        if j < test_result['expected'].shape[1]:
                            exp_table.cell(0, j + 1).text = str(col_name)

                    # 添加期望频数数据
                    for i in range(test_result['expected'].shape[0]):
                        exp_table.cell(i + 1, 0).text = str(row_labels[i])
                        for j in range(test_result['expected'].shape[1]):
                            exp_table.cell(i + 1, j + 1).text = f"{test_result['expected'][i, j]:.2f}"

                # Fisher检验结果
                else:
                    doc.add_paragraph(f"{lang['p_value_label']}: {test_result['p_value']:.6f}")
                    doc.add_paragraph(f"{lang['odds_ratio_label']}: {test_result['odds_ratio']:.4f}")
                    significance_text = lang['significant'] if test_result['p_value'] < 0.05 else lang['not_significant']
                    doc.add_paragraph(f"{lang['significance_label']}: {significance_text}")

                # 添加检验解释
                doc.add_heading('检验解释' if self.current_language == 'zh' else 'Test Explanation', level=1)
                doc.add_paragraph(LANGUAGES[self.current_language]['explanation'][test_type])

                # 添加统计量解释
                if test_type != current_fisher:
                    # 动态获取统计量名称（适配中英文）
                    stats_keys = {
                        'zh': ["卡方值", "p值", "自由度", "显著性（α=0.05）"],
                        'en': ["Chi-square Value", "p-value", "Degrees of Freedom", "Significance (α=0.05)"]
                    }
                    for key in stats_keys[self.current_language]:
                        doc.add_paragraph(f"{key}: {LANGUAGES[self.current_language]['interpretation'][key]}")
                else:
                    doc.add_paragraph(f"{lang['p_value_label']}: {LANGUAGES[self.current_language]['interpretation'][lang['p_value_label']]}")
                    doc.add_paragraph(f"{lang['odds_ratio_label']}: {lang['odds_ratio_explanation']}")

                # 保存文档
                doc.save(save_path)

                self.result_label.config(text=LANGUAGES[self.current_language]['analysis_success'].format(save_path))
            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        self.current_language = 'zh' if self.current_language == 'en' else 'en'
        self.root.title(LANGUAGES[self.current_language]['title'])
        self.select_button.config(text=LANGUAGES[self.current_language]['select_button'])
        self.analyze_button.config(text=LANGUAGES[self.current_language]['analyze_button'])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]['switch_language'])
        self.open_excel_label.config(text=LANGUAGES[self.current_language]["open_excel_button_text"])

        # 更新输入框提示文字
        current_text = self.file_entry.get()
        placeholder = "请输入待分析 Excel 文件的完整路径" if self.current_language == 'zh' else "Please enter the full path of the Excel file to be analyzed"
        if current_text in ["请输入待分析 Excel 文件的完整路径",
                            "Please enter the full path of the Excel file to be analyzed"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, placeholder)

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=LANGUAGES[self.current_language]['select_button'],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        placeholder = "请输入待分析 Excel 文件的完整路径" if self.current_language == 'zh' else "Please enter the full path of the Excel file to be analyzed"
        self.file_entry.insert(0, placeholder)
        self.file_entry.pack(pady=5)
        self.file_entry.config(foreground='gray')

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=LANGUAGES[self.current_language]['analyze_button'],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=LANGUAGES[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=LANGUAGES[self.current_language]['switch_language'],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = ChiSquaredTestApp()
    app.run()


if __name__ == "__main__":
    run_app()
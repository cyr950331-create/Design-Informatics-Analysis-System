import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
from tkinter import messagebox
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典（包含所有需要翻译的文本）
LANGUAGES = {
    'zh': {
        'title': "独立性权重法",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "请选择文件。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "中/英",
        'open_excel_button_text': "示例数据",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'preprocessing': "数据预处理：已移除包含非数值/空值的行，共保留 {} 行有效数据",
        'all_positive': "所有指标均为正向指标（越大越好）",
        'result_heading': "独立性权重法分析结果",
        'table_heading': "分析结果表格",
        'explanations_heading': "解释说明",
        'interpretation_heading': "结果解读",
        'visualization_heading': "指标权重可视化",
        'chart_title': "指标权重柱状图",
        'chart_xlabel': "指标名称",
        'chart_ylabel': "指标权重",
        'stats': {
            "original_data": "原始数据矩阵",
            "processed_data": "预处理后数据",
            "std_matrix": "标准差矩阵",
            "weights": "指标权重",
            "weight_calc": "权重计算过程"
        },
        'explanation': {
            "original_data": "从 Excel 文件中读取的原始数据矩阵",
            "processed_data": "经过清洗处理后的数据",
            "std_matrix": "各指标的标准差矩阵，反映指标的变异程度",
            "weights": "根据标准差矩阵计算得到的各指标权重",
            "weight_calc": "标准差及权重的计算步骤"
        },
        'interpretation': {
            "original_data": "用于后续分析的基础数据",
            "processed_data": "移除无效值后的数据",
            "std_matrix": "标准差越大，该指标的变异程度越大，在综合评价中越重要",
            "weights": "各指标在综合评价中的相对重要程度，权重越大越重要",
            "weight_calc": "权重 = 指标标准差 / 所有指标标准差之和"
        },
        'table_headers': {
            "statistic": "统计量",
            "value": "统计量值"
        }
    },
    'en': {
        'title': "Independence Weight Method",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Chinese/English",
        'open_excel_button_text': "Example data",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'preprocessing': "Data preprocessing: Rows with non-numeric/empty values removed, {} valid rows retained",
        'all_positive': "All indicators are positive indicators (larger is better)",
        'result_heading': "Independence Weighting Method Analysis Results",
        'table_heading': "Analysis Results Table",
        'explanations_heading': "Explanations",
        'interpretation_heading': "Result Interpretation",
        'visualization_heading': "Indicator Weight Visualization",
        'chart_title': "Bar Chart of Indicator Weights",
        'chart_xlabel': "Indicator Name",
        'chart_ylabel': "Indicator Weight",
        'stats': {
            "original_data": "Original Data Matrix",
            "processed_data": "Processed Data",
            "std_matrix": "Standard Deviation Matrix",
            "weights": "Indicator Weights",
            "weight_calc": "Weight Calculation Process"
        },
        'explanation': {
            "original_data": "The original data matrix read from the Excel file",
            "processed_data": "Data after cleaning processing",
            "std_matrix": "The standard deviation matrix of each indicator, reflecting the degree of variation",
            "weights": "The weight of each indicator calculated based on the standard deviation matrix",
            "weight_calc": "Calculation steps of standard deviation and weight"
        },
        'interpretation': {
            "original_data": "The basic data for subsequent analysis",
            "processed_data": "Data after removing invalid values",
            "std_matrix": "The larger the standard deviation, the greater the variation degree and importance",
            "weights": "The relative importance of each indicator in comprehensive evaluation",
            "weight_calc": "Weight = indicator standard deviation / sum of all indicator standard deviations"
        },
        'table_headers': {
            "statistic": "Statistic",
            "value": "Statistic Value"
        }
    }
}


class IndependenceWeightingMethodAnalysisApp:
    def __init__(self, root=None):
        self.current_language = "en"
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])
        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data28.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"{LANGUAGES[self.current_language]['file_not_found']}：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"{LANGUAGES[self.current_language]['analysis_error'].format(str(e))}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def preprocess_data(self, data):
        """数据预处理：移除包含非数值和空值的行"""
        # 转换为DataFrame便于处理
        df = pd.DataFrame(data)
        # 保留所有值都是数值的行
        df = df.apply(pd.to_numeric, errors='coerce').dropna()
        return df.values, len(df)

    def independence_weight_method(self, data):
        """实现独立性权重法"""
        std_matrix = np.std(data, axis=0)
        weights = std_matrix / np.sum(std_matrix)
        return std_matrix, weights

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 默认文件包含表头
            has_header = True
            header = 0 if has_header else None

            # 读取数据
            df = pd.read_excel(file_path, header=header)
            original_data = df.values
            headers = df.columns.tolist() if has_header else [f"{LANGUAGES[self.current_language]['table_headers']['statistic']}{i + 1}" for i in range(df.shape[1])]

            # 数据预处理
            processed_data, valid_rows = self.preprocess_data(original_data)
            if processed_data.size == 0:
                raise ValueError(f"{LANGUAGES[self.current_language]['analysis_error'].format('No valid numeric data after preprocessing')}")

            # 计算权重
            std_matrix, weights = self.independence_weight_method(processed_data)

            # 准备权重计算过程数据
            weight_calc = []
            total_std = np.sum(std_matrix)
            for i in range(len(std_matrix)):
                calc_str = f"{std_matrix[i]:.4f} / {total_std:.4f} = {weights[i]:.4f}"
                weight_calc.append(calc_str)

            # 整理结果数据
            data = [
                [LANGUAGES[self.current_language]['stats']['original_data'], original_data.tolist()],
                [LANGUAGES[self.current_language]['stats']['processed_data'], processed_data.tolist()],
                [LANGUAGES[self.current_language]['stats']['std_matrix'], [f"{x:.4f}" for x in std_matrix.tolist()]],
                [LANGUAGES[self.current_language]['stats']['weights'], [f"{x:.4f}" for x in weights.tolist()]],
                [LANGUAGES[self.current_language]['stats']['weight_calc'], weight_calc]
            ]
            headers_table = [
                LANGUAGES[self.current_language]['table_headers']['statistic'],
                LANGUAGES[self.current_language]['table_headers']['value']
            ]

            # 创建结果表格
            result_df = pd.DataFrame(data, columns=headers_table)

            # 保存结果
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                doc = Document()
                doc.add_heading(LANGUAGES[self.current_language]['result_heading'], 0)
                # 添加预处理信息
                doc.add_paragraph(LANGUAGES[self.current_language]['preprocessing'].format(valid_rows))
                doc.add_paragraph(LANGUAGES[self.current_language]['all_positive'])

                # 添加表格
                doc.add_heading(LANGUAGES[self.current_language]['table_heading'], level=1)
                table = doc.add_table(rows=1, cols=len(result_df.columns))
                hdr_cells = table.rows[0].cells
                for col, header in enumerate(result_df.columns):
                    hdr_cells[col].text = header

                for _, row in result_df.iterrows():
                    row_cells = table.add_row().cells
                    for col, value in enumerate(row):
                        # 处理列表类型数据的显示
                        if isinstance(value, list):
                            row_cells[col].text = "\n".join([str(item) for item in value[:5]]) + (
                                "..." if len(value) > 5 else "")
                        else:
                            row_cells[col].text = str(value)

                # 添加解释说明
                doc.add_heading(LANGUAGES[self.current_language]['explanations_heading'], level=1)
                for key, value in LANGUAGES[self.current_language]['explanation'].items():
                    doc.add_paragraph(f"{LANGUAGES[self.current_language]['stats'][key]}: {value}", style='List Bullet')

                # 添加结果解读
                doc.add_heading(LANGUAGES[self.current_language]['interpretation_heading'], level=1)
                for key, value in LANGUAGES[self.current_language]['interpretation'].items():
                    doc.add_paragraph(f"{LANGUAGES[self.current_language]['stats'][key]}: {value}", style='List Bullet')

                # 生成权重柱状图
                fig, ax = plt.subplots(figsize=(12, 8))
                ax.bar(headers, weights)
                ax.set_title(LANGUAGES[self.current_language]['chart_title'])
                ax.set_xlabel(LANGUAGES[self.current_language]['chart_xlabel'])
                ax.set_ylabel(LANGUAGES[self.current_language]['chart_ylabel'])
                plt.xticks(rotation=0, ha='right')
                plt.tight_layout()

                img_path = os.path.splitext(save_path)[0] + '_indicator_weights.png'
                plt.savefig(img_path, dpi=300)
                plt.close()

                doc.add_heading(LANGUAGES[self.current_language]['visualization_heading'], level=1)
                doc.add_picture(img_path, width=Inches(6))
                doc.save(save_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(save_path)
                self.result_label.config(text=result_msg, wraplength=400)
            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'
        self.root.title(LANGUAGES[self.current_language]['title'])
        self.select_button.config(text=LANGUAGES[self.current_language]['select_button'])
        self.analyze_button.config(text=LANGUAGES[self.current_language]['analyze_button'])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]['switch_language'])
        current_entry_text = self.file_entry.get()
        # 只有当输入框显示的是占位符时才更新
        if current_entry_text == LANGUAGES['zh']['file_entry_placeholder'] or current_entry_text == LANGUAGES['en']['file_entry_placeholder']:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
            self.file_entry.config(foreground='gray')
        self.open_excel_label.config(text=LANGUAGES[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        style = ttk.Style()
        style.configure("Gray.TEntry", foreground="gray")

        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        self.select_button = ttk.Button(frame, text=LANGUAGES[self.current_language]['select_button'],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        self.analyze_button = ttk.Button(frame, text=LANGUAGES[self.current_language]['analyze_button'],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=LANGUAGES[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        self.switch_language_label = ttk.Label(frame, text=LANGUAGES[self.current_language]['switch_language'],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        self.root.mainloop()


def run_app():
    app = IndependenceWeightingMethodAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()
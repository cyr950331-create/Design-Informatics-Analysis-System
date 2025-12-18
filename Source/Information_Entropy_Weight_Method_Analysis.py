import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典，所有可翻译文本均包含在内
LANGUAGES = {
    'zh': {
        'title': "信息量权重法",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "请选择文件。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "中/英",
        'open_excel_button_text': "示例数据",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'statistic': "统计量",
        'statistic_value': "统计量值",
        'explanation_title': "解释说明",
        'interpretation_title': "结果解读",
        'original_data': "原始数据",
        'indicator_entropy': "指标熵值",
        'indicator_redundancy': "指标冗余度",
        'information_weight': "信息量权重",
        'explanation': {
            'original_data': "输入的待分析数据",
            'indicator_entropy': "反映各指标信息无序程度的统计量",
            'indicator_redundancy': "指标熵值的互补量，反映指标提供的有效信息",
            'information_weight': "根据指标冗余度计算得到的各指标权重"
        },
        'interpretation': {
            'original_data': "作为分析的基础数据",
            'indicator_entropy': "值越大，指标信息越无序，提供的有效信息越少",
            'indicator_redundancy': "值越大，指标提供的有效信息越多",
            'information_weight': "权重越大，该指标在综合评价中越重要"
        },
        'statistical_results': "统计结果",
        'pie_chart_title': "信息量权重分布饼图",
        'bar_chart_title': "信息量权重柱状图",
        'indicators': "指标",
        'weights': "权重"
    },
    'en': {
        'title': "Information Entropy Weight Method",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Chinese/English",
        'open_excel_button_text': "Example data",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'statistic': "Statistic",
        'statistic_value': "Statistic Value",
        'explanation_title': "Explanation",
        'interpretation_title': "Interpretation",
        'original_data': "Original Data",
        'indicator_entropy': "Indicator Entropy",
        'indicator_redundancy': "Indicator Redundancy",
        'information_weight': "Information Weight",
        'explanation': {
            'original_data': "The input data to be analyzed",
            'indicator_entropy': "A statistic reflecting the degree of disorder of information for each indicator",
            'indicator_redundancy': "The complementary quantity of the indicator entropy, reflecting the effective information provided by the indicator",
            'information_weight': "The weight of each indicator calculated based on the indicator redundancy"
        },
        'interpretation': {
            'original_data': "As the basic data for analysis",
            'indicator_entropy': "The larger the value, the more disordered the indicator information and the less effective information it provides",
            'indicator_redundancy': "The larger the value, the more effective information the indicator provides",
            'information_weight': "The larger the weight, the more important the indicator is in the comprehensive evaluation"
        },
        'statistical_results': "Statistical Results",
        'pie_chart_title': "Pie Chart of Information Entropy Weights",
        'bar_chart_title': "Bar Chart of Information Entropy Weights",
        'indicators': "Indicators",
        'weights': "Weights"
    }
}


class InformationEntropyWeightMethodAnalysisApp:
    def __init__(self, root=None):
        self.current_language = "en"
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])
        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data28.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"{LANGUAGES[self.current_language]['file_not_found']}：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"{LANGUAGES[self.current_language]['analysis_error'].format(str(e))}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def information_entropy_weight_method(self, data):
        """实现信息量权重法（优化版）"""
        # 数据标准化优化：支持正负数据的min-max标准化
        min_vals = data.min(axis=0)
        max_vals = data.max(axis=0)
        ranges = max_vals - min_vals

        # 处理极差为0的列（所有值相同）
        ranges[ranges == 0] = 1e-8  # 避免除零错误
        standardized_data = (data - min_vals) / ranges

        # 确保标准化后数据非负且在合理范围
        standardized_data = np.clip(standardized_data, 1e-8, 1 - 1e-8)

        # 计算指标熵值（优化异常处理）
        n_samples = data.shape[0]
        entropy = np.zeros(standardized_data.shape[1])

        for i in range(standardized_data.shape[1]):
            col_data = standardized_data[:, i]
            # 检查列数据是否接近常数（方差极小）
            if np.var(col_data) < 1e-10:
                entropy[i] = 1.0  # 常数指标熵值为1（最大无序）
            else:
                entropy[i] = -np.sum(col_data * np.log(col_data)) / np.log(n_samples)

        # 计算指标冗余度
        redundancy = 1 - entropy

        # 处理所有冗余度为0的极端情况
        if np.sum(redundancy) == 0:
            weights = np.ones_like(redundancy) / len(redundancy)
        else:
            weights = redundancy / np.sum(redundancy)

        # 确保权重非负（解决饼图绘制问题）
        weights = np.clip(weights, 0, None)
        # 处理权重总和为0的极端情况
        if np.sum(weights) == 0:
            weights = np.ones_like(weights) / len(weights)
        else:
            # 重新归一化确保权重和为1
            weights = weights / np.sum(weights)

        return entropy, redundancy, weights

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 关键修改：将第一行作为表头（header=0表示第一行为表头）
            df = pd.read_excel(file_path, header=0)
            # 获取表头（指标名称）
            headers = df.columns.tolist()
            # 获取数据部分（排除表头）
            data = df.values.astype(float)

            # 进行信息量权重法分析
            entropy, redundancy, weights = self.information_entropy_weight_method(data)

            # 整理数据
            data = [
                [LANGUAGES[self.current_language]['original_data'], ""],  # 占位行
                [LANGUAGES[self.current_language]['indicator_entropy'], entropy.tolist()],
                [LANGUAGES[self.current_language]['indicator_redundancy'], redundancy.tolist()],
                [LANGUAGES[self.current_language]['information_weight'], weights.tolist()]
            ]
            stats_headers = [
                LANGUAGES[self.current_language]['statistic'],
                LANGUAGES[self.current_language]['statistic_value']
            ]
            stats_df = pd.DataFrame(data, columns=stats_headers)

            # 添加解释说明
            explanations = {
                LANGUAGES[self.current_language]['original_data']: LANGUAGES[self.current_language]['explanation'][
                    'original_data'],
                LANGUAGES[self.current_language]['indicator_entropy']: LANGUAGES[self.current_language]['explanation'][
                    'indicator_entropy'],
                LANGUAGES[self.current_language]['indicator_redundancy']:
                    LANGUAGES[self.current_language]['explanation']['indicator_redundancy'],
                LANGUAGES[self.current_language]['information_weight']: LANGUAGES[self.current_language]['explanation'][
                    'information_weight']
            }

            # 添加分析结果解读
            interpretations = {
                LANGUAGES[self.current_language]['original_data']: LANGUAGES[self.current_language]['interpretation'][
                    'original_data'],
                LANGUAGES[self.current_language]['indicator_entropy']:
                    LANGUAGES[self.current_language]['interpretation']['indicator_entropy'],
                LANGUAGES[self.current_language]['indicator_redundancy']:
                    LANGUAGES[self.current_language]['interpretation']['indicator_redundancy'],
                LANGUAGES[self.current_language]['information_weight']:
                    LANGUAGES[self.current_language]['interpretation']['information_weight']
            }

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加原始数据表格（使用实际表头）
                doc.add_heading(LANGUAGES[self.current_language]['original_data'], level=2)
                data_rows, data_cols = df.shape
                data_table = doc.add_table(rows=data_rows + 1, cols=data_cols)
                # 添加表头（使用Excel中的实际表头）
                for col_idx in range(data_cols):
                    data_table.rows[0].cells[col_idx].text = str(headers[col_idx])
                # 填充原始数据
                for row_idx in range(data_rows):
                    for col_idx in range(data_cols):
                        data_table.rows[row_idx + 1].cells[col_idx].text = f"{df.iloc[row_idx, col_idx]:.4f}"

                # 添加其他统计量表格
                doc.add_heading(LANGUAGES[self.current_language]['statistical_results'], level=2)
                stats_table = doc.add_table(rows=stats_df.shape[0] + 1, cols=stats_df.shape[1])
                hdr_cells = stats_table.rows[0].cells
                for col_idx, header in enumerate(stats_headers):
                    hdr_cells[col_idx].text = header
                for row_idx in range(stats_df.shape[0]):
                    row_cells = stats_table.rows[row_idx + 1].cells
                    for col_idx in range(stats_df.shape[1]):
                        row_cells[col_idx].text = str(stats_df.iloc[row_idx, col_idx])

                # 添加解释说明
                doc.add_heading(LANGUAGES[self.current_language]['explanation_title'], level=2)
                for key, value in explanations.items():
                    doc.add_paragraph(f"{key}: {value}")

                # 添加分析结果解读
                doc.add_heading(LANGUAGES[self.current_language]['interpretation_title'], level=2)
                for key, value in interpretations.items():
                    doc.add_paragraph(f"{key}: {value}")

                # 生成信息量权重分布饼图（使用实际表头作为标签）
                fig, ax = plt.subplots(figsize=(12, 8))
                ax.pie(weights, labels=headers, autopct='%1.1f%%')
                ax.set_title(LANGUAGES[self.current_language]['pie_chart_title'])
                # 保存图片
                img_path = os.path.splitext(save_path)[0] + '_weights_pie_chart.png'
                plt.savefig(img_path, dpi=300)
                plt.close()

                # 新增：生成权重柱状图
                fig, ax = plt.subplots(figsize=(12, 8))
                ax.bar(headers, weights)
                ax.set_title(LANGUAGES[self.current_language]['bar_chart_title'])
                ax.set_xlabel(LANGUAGES[self.current_language]['indicators'])
                ax.set_ylabel(LANGUAGES[self.current_language]['weights'])
                plt.xticks(rotation=0, ha='center')
                plt.tight_layout()
                bar_img_path = os.path.splitext(save_path)[0] + '_weights_bar_chart.png'
                plt.savefig(bar_img_path, dpi=300)
                plt.close()

                # 将图片插入到 Word 文档中
                doc.add_heading(LANGUAGES[self.current_language]['pie_chart_title'], level=2)
                doc.add_picture(img_path, width=Inches(6))
                # 将柱状图插入到Word文档
                doc.add_heading(LANGUAGES[self.current_language]['bar_chart_title'], level=2)
                doc.add_picture(bar_img_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(save_path)
                self.result_label.config(text=result_msg, wraplength=400)

            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'
        self.root.title(LANGUAGES[self.current_language]['title'])
        self.select_button.config(text=LANGUAGES[self.current_language]['select_button'])
        self.analyze_button.config(text=LANGUAGES[self.current_language]['analyze_button'])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]['switch_language'])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
        self.file_entry.config(foreground='gray')
        self.open_excel_label.config(text=LANGUAGES[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        style = ttk.Style()
        style.configure("Gray.TEntry", foreground="gray")

        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        self.select_button = ttk.Button(frame, text=LANGUAGES[self.current_language]['select_button'],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        self.analyze_button = ttk.Button(frame, text=LANGUAGES[self.current_language]['analyze_button'],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=LANGUAGES[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        self.switch_language_label = ttk.Label(frame, text=LANGUAGES[self.current_language]['switch_language'],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        self.root.mainloop()


def run_app():
    app = InformationEntropyWeightMethodAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()
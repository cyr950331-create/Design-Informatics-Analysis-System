import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox, Dialog
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
from sklearn.metrics import accuracy_score, confusion_matrix, cohen_kappa_score
import matplotlib.pyplot as plt
from statsmodels.miscmodels.ordinal_model import OrderedModel
from docx import Document
from docx.shared import Cm
from statsmodels.stats.outliers_influence import variance_inflation_factor

# 设置中文字体，确保中文正常显示
plt.rcParams["font.family"] = ["SimHei", "WenQuanYi Micro Hei", "Heiti TC"]
plt.rcParams["axes.unicode_minus"] = False  # 解决负号显示问题

def to_numpy_safe(x):

    if isinstance(x, np.ndarray):
        return x
    if hasattr(x, "to_numpy"):
        try:
            return x.to_numpy()
        except Exception:
            pass
    if hasattr(x, "values"):
        try:
            return np.asarray(x.values)
        except Exception:
            pass
    return np.asarray(x)

LANGUAGES = {
    'zh': {
        'title': "有序Logit回归",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "请选择文件。",
        'analysis_success': "分析完成，结果已保存到 {}\n图片已保存到同目录下的images文件夹",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'images_saved': "图片已保存到 {}",
        'switch_language': "中/英",
        'open_excel_button_text': "示例数据",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'missing_value_title': "缺失值处理",
        'missing_value_message': "检测到缺失值，请选择处理方式：",
        'delete_rows': "删除含缺失值的行",
        'mean_median': "均值/中位数填充",
        'vif_warning': "警告：以下自变量存在较高的多重共线性（VIF > 10）：\n{}\n建议检查并处理这些变量。",
        'vif_handle_title': "多重共线性处理",
        'vif_handle_message': "检测到高VIF变量，是否自动移除？",
        'vif_removed': "已自动移除高VIF变量：{}",
        'outlier_warning': "警告：以下自变量检测到极端值（超过3个标准差）：\n{}\n建议检查并处理这些异常值。",
        'convergence_warning': "警告：模型未收敛！已尝试增加迭代次数和调整优化方法，结果可能不可靠。",
        'dv_not_ordered': "错误：最后一列不是有效的有序分类变量，请检查数据。",
        'dv_check_failed': "错误：因变量必须是有序分类变量（至少2个类别且有明确顺序）。",
        'regression_equation': "回归方程",
        'equation_explanation': "有序Logit模型通过累积概率建模，公式如下：",
        'cumulative_prob': "累积概率公式: P(Y ≤ k) = 1 / (1 + exp(-(α_k - (β₁X₁ + β₂X₂ + ... + βₚXₚ))))",
        'model_convergence': "模型收敛状态",
        'optimization_method': "优化方法",
        'iterations_used': "使用迭代次数",
        'data_processing_notes': "数据处理说明",
        'vif_removed_note': "自动移除的高VIF变量（VIF>10）: {}",
        'no_high_vif_vars': "未检测到需要移除的高VIF变量",
        'factorized_vars_note': "注意: 对下列自变量执行了类别编码（factorize）以便用于模型（原始为非数值型）: {}",
        'linear_combination': "线性组合: Z = {}",
        'equation_terms_explanation': "其中，α_k 代表各个阈值参数，β代表回归系数，X代表自变量",
        'true_category': "真实类别: {}",
        'predicted_category': "预测类别: {}",
        'category_label': "类别 {}",
        'threshold_label': "阈值 {}",
        'analysis_results_title': "有序Logit回归分析结果",
        'model_summary_heading': '模型摘要',
        'confusion_matrix_heading': '混淆矩阵',
        'coefficients_heading': '系数表格',
        'coef_significance_heading': '系数显著性（含置信区间）',
        'thresholds_heading': '阈值参数',
        'threshold_visualization_heading': '阈值可视化',
        'pred_prob_distribution_heading': '各真实类别下的预测概率分布',
        'explanations_heading': '指标解释',
        'stat_names': {
            "Coefficients": "回归系数",
            "Intercept": "截距项",
            "Accuracy": "准确率",
            "Weighted Kappa": "加权Kappa系数",
            "z-value": "z统计量",
            "p-value": "p值",
            "Thresholds": "阈值参数",
            "Pseudo R-squared": "伪R²（McFadden）",
            "Metric": "指标",
            "Value": "数值",
            "Variable": "变量",
            "Significance": "显著性",
            "No. Observations": "观测值数量",
            "Log-Likelihood": "对数似然值",
            "LL-Null": "零模型对数似然值",
            "LLR p-value": "似然比检验p值",
            "Threshold {i}": "阈值 {i}"
        },
        'explanation': {
            "Coefficients": "表示每个自变量对因变量的影响程度。",
            "Intercept": "是有序Logit模型阈值的组成部分。",
            "Accuracy": "衡量模型预测正确的比例。",
            "Weighted Kappa": "考虑类别顺序的一致性度量，范围-1到1，值越大越好。",
            "z-value": "用于检验每个自变量的显著性。",
            "p-value": "用于判断自变量的显著性，p 值越小，自变量越显著。",
            "Thresholds": "用于确定有序分类的边界。Threshold k 对应类别 k 和 k+1 的边界。",
            "Pseudo R-squared": "衡量模型拟合优度，范围0-1，值越大表示拟合越好。"
        },
        'interpretation': {
            "Coefficients": "正值表示该自变量增加时，因变量取值有增大的趋势；负值表示该自变量增加时，因变量取值有减小的趋势。",
            "Intercept": "与其他阈值共同构成有序分类的边界点，单独解释意义有限。",
            "Accuracy": "模型预测准确率越高，说明模型对数据的拟合效果越好。",
            "Weighted Kappa": "考虑类别顺序的一致性指标，0.8以上表示几乎完美，0.6-0.8表示高度一致，0.4-0.6表示中度一致。",
            "z-value": "绝对值越大，表示该自变量对因变量的影响越显著。",
            "p-value": "通常以0.05为阈值，小于0.05表示该自变量对因变量有显著影响。",
            "Thresholds": "决定了有序分类的临界点。例如，若因变量有3个类别，阈值1是类别1和2的边界，阈值2是类别2和3的边界。",
            "Pseudo R-squared": "用于比较模型与空模型的拟合效果，值接近1表示模型拟合较好。",
            "LLR p-value": "如果p值小于0.05，表示模型显著优于空模型"
        }
    },
    'en': {
        'title': "Ordered Logit Regression",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\nImages saved to 'images' folder in the same directory",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'images_saved': "Images have been saved to {}",
        'switch_language': "Chinese/English",
        'open_excel_button_text': "Example data",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'missing_value_title': "Missing Value Handling",
        'missing_value_message': "Missing values detected. Please choose a handling method:",
        'delete_rows': "Delete rows with missing values",
        'mean_median': "Mean/median imputation",
        'vif_warning': "Warning: The following variables have high multicollinearity (VIF > 10):\n{}\nIt is recommended to check and handle these variables.",
        'vif_handle_title': "Multicollinearity Handling",
        'vif_handle_message': "High VIF variables detected. Automatically remove them?",
        'vif_removed': "Automatically removed high VIF variables: {}",
        'outlier_warning': "Warning: Outliers detected in the following variables (exceeding 3 standard deviations):\n{}\nIt is recommended to check and handle these outliers.",
        'convergence_warning': "Warning: Model did not converge! Increased iterations and adjusted optimization method, results may be unreliable.",
        'dv_not_ordered': "Error: The last column is not a valid ordered categorical variable. Please check your data.",
        'dv_check_failed': "Error: Dependent variable must be an ordered categorical variable (at least 2 categories with clear order).",
        'regression_equation': "Regression Equation",
        'equation_explanation': "Ordered Logit models use cumulative probabilities with the following formula:",
        'cumulative_prob': "Cumulative probability formula: P(Y ≤ k) = 1 / (1 + exp(-(α_k - (β₁X₁ + β₂X₂ + ... + βₚXₚ))))",
        'model_convergence': "Model Convergence Status",
        'optimization_method': "Optimization Method",
        'iterations_used': "Iterations Used",
        'data_processing_notes': "Data Processing Notes",
        'vif_removed_note': "Automatically removed high VIF variables (VIF>10): {}",
        'no_high_vif_vars': "No high VIF variables detected for removal",
        'factorized_vars_note': "Note: The following independent variables were factorized for model compatibility (originally non-numeric): {}",
        'linear_combination': "Linear combination: Z = {}",
        'equation_terms_explanation': "Where α_k represents each threshold parameter, β represents regression coefficients, and X represents independent variables",
        'true_category': "True Category: {}",
        'predicted_category': "Predicted Category: {}",
        'category_label': "Category {}",
        'threshold_label': "Threshold {}",
        'analysis_results_title': "Ordered Logit Regression Analysis Results",
        'model_summary_heading': 'Model Summary',
        'confusion_matrix_heading': 'Confusion Matrix',
        'coefficients_heading': 'Coefficients',
        'coef_significance_heading': 'Coefficient Significance with Confidence Intervals',
        'thresholds_heading': 'Thresholds',
        'threshold_visualization_heading': 'Threshold Visualization',
        'pred_prob_distribution_heading': 'Predicted Probability Distribution by True Category',
        'explanations_heading': 'Explanations',
        'stat_names': {
            "Coefficients": "Coefficients",
            "Intercept": "Intercept",
            "Accuracy": "Accuracy",
            "Weighted Kappa": "Weighted Kappa",
            "z-value": "z-value",
            "p-value": "p-value",
            "Thresholds": "Thresholds",
            "Pseudo R-squared": "Pseudo R-squared",
            "Metric": "Metric",
            "Value": "Value",
            "Variable": "Variable",
            "Significance": "Significance",
            "No. Observations": "No. Observations",
            "Log-Likelihood": "Log-Likelihood",
            "LL-Null": "LL-Null",
            "LLR p-value": "LLR p-value",
            "Threshold {i}": "Threshold {i}"
        },
        'explanation': {
            "Coefficients": "Indicating the influence of each independent variable on the dependent variable.",
            "Intercept": "Which is part of the threshold structure in ordered Logit models.",
            "Accuracy": "Measuring the proportion of correct predictions of the model.",
            "Weighted Kappa": "A measure of agreement considering category order, ranging from -1 to 1, higher values are better.",
            "z-value": "Used to test the significance of each independent variable.",
            "p-value": "Used to determine the significance of the independent variable. The smaller the p value，the more significant the independent variable.",
            "Thresholds": "Threshold k separates category k and k+1.",
            "Pseudo R-squared": "Measuring model fit，ranging from 0-1，higher values indicate better fit."
        },
        'interpretation': {
            "Coefficients": "A positive value indicates that as the independent variable increases，the dependent variable tends to increase；a negative value indicates the opposite trend.",
            "Intercept": "Part of the threshold structure that defines category boundaries，has limited standalone interpretation.",
            "Accuracy": "The higher the accuracy，the better the model fits the data.",
            "Weighted Kappa": "Order-aware agreement metric: >0.8 nearly perfect, 0.6-0.8 substantial, 0.4-0.6 moderate.",
            "z-value": "The larger the absolute value，the more significant the impact of the independent variable on the dependent variable.",
            "p-value": "Typically using a threshold of 0.05，values smaller than 0.05 indicate a significant impact of the independent variable on the dependent variable.",
            "Thresholds": "Define critical points for ordered classification. For example, with 3 categories，Threshold 1 separates categories 1 and 2，Threshold 2 separates categories 2 and 3.",
            "Pseudo R-squared": "Compares model fit with a null model，values close to 1 indicate better fit.",
            "LLR p-value": "If p-value < 0.05, the model is significantly better than the null model"
        }
    }
}


# 对话框类
class MissingValueDialog(Dialog):
    def __init__(self, parent, language):
        self.language = language
        self.result = None
        super().__init__(parent, title=LANGUAGES[language]['missing_value_title'])

    def body(self, master):
        ttk.Label(master, text=LANGUAGES[self.language]['missing_value_message']).pack(pady=10)

        frame = ttk.Frame(master)
        frame.pack(pady=5)

        self.var = tk.IntVar(value=0)

        ttk.Radiobutton(
            frame,
            text=LANGUAGES[self.language]['delete_rows'],
            variable=self.var,
            value=0
        ).pack(anchor=W, pady=2)

        ttk.Radiobutton(
            frame,
            text=LANGUAGES[self.language]['mean_median'],
            variable=self.var,
            value=1
        ).pack(anchor=W, pady=2)

        return frame

    def apply(self):
        self.result = self.var.get()

class VIFHandlingDialog(Dialog):
    def __init__(self, parent, language):
        self.language = language
        self.result = None
        super().__init__(parent, title=LANGUAGES[language]['vif_handle_title'])

    def body(self, master):
        ttk.Label(master, text=LANGUAGES[self.language]['vif_handle_message']).pack(pady=10)
        return super().body(master)

    def apply(self):
        self.result = True  # 确认删除

# 主类
class OrderedLogitRegressionAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data36.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def handle_missing_values(self, df):
        # 检查是否有缺失值
        if not df.isnull().any().any():
            return df

        # 显示缺失值处理对话框（阻塞等待用户选择）
        dialog = MissingValueDialog(self.root, self.current_language)
        # wait until dialog closed so dialog.result is set by apply()
        self.root.wait_window(dialog)

        # 默认选择为删除（如果用户未选择，dialog.result 可能为 None）
        choice = dialog.result if dialog.result is not None else 0

        if choice == 0:  # 删除行
            df = df.dropna()
            if df.shape[0] == 0:
                raise ValueError("数据中所有行都包含缺失值，无法进行分析。")
        else:  # 均值/中位数填充
            # 对数值型列用均值填充
            numeric_cols = df.select_dtypes(include=['number']).columns
            df[numeric_cols] = df[numeric_cols].fillna(df[numeric_cols].mean())

            # 对分类型列用众数填充
            categorical_cols = df.select_dtypes(exclude=['number']).columns
            for col in categorical_cols:
                if not df[col].mode().empty:
                    df[col] = df[col].fillna(df[col].mode()[0])
                else:
                    df[col] = df[col].fillna("")

        return df

    def check_multicollinearity(self, X):
        # 计算VIF值检测多重共线性
        vif_data = pd.DataFrame()
        vif_data["Variable"] = X.columns
        if X.shape[1] == 0:
            return X, [], vif_data
        try:
            vif_data["VIF"] = [variance_inflation_factor(X.values, i) for i in range(X.shape[1])]
        except Exception as e:
            raise RuntimeError(f"计算VIF时出错: {e}")

        high_vif_vars = vif_data[vif_data["VIF"] > 10]["Variable"].tolist()
        removed_vars = []

        if high_vif_vars:
            Messagebox.show_info(
                "Warning",
                LANGUAGES[self.current_language]['vif_warning'].format(", ".join(high_vif_vars))
            )

            # 使用阻塞式 yes/no 确认（ttkbootstrap Messagebox 的 yesno）
            try:
                user_choice = Messagebox.yesno("Confirm", LANGUAGES[self.current_language]['vif_handle_message'])
            except Exception:
                # 如果没有 yesno 方法，使用 VIFHandlingDialog 阻塞式对话框
                dialog = VIFHandlingDialog(self.root, self.current_language)
                self.root.wait_window(dialog)
                user_choice = bool(dialog.result)

            if user_choice:
                X = X.drop(columns=high_vif_vars)
                removed_vars = high_vif_vars
                Messagebox.show_info(
                    "Info",
                    LANGUAGES[self.current_language]['vif_removed'].format(", ".join(removed_vars))
                )
                if X.shape[1] == 0:
                    raise ValueError("所有自变量均因高VIF被移除，无法构建模型。")

        return X, removed_vars, vif_data

    def check_outliers(self, X):
        # 检测极端值（超过3个标准差）
        outlier_vars = []
        for col in X.columns:
            mean = X[col].mean()
            std = X[col].std()
            if pd.isna(std) or std == 0:
                continue
            outliers = (X[col] < (mean - 3 * std)) | (X[col] > (mean + 3 * std))
            if outliers.any():
                outlier_vars.append(col)

        if outlier_vars:
            Messagebox.show_info(
                "Warning",
                LANGUAGES[self.current_language]['outlier_warning'].format(", ".join(outlier_vars))
            )

    def validate_dependent_variable(self, y):
        """验证因变量是否为有效的有序分类变量并返回排序后的类别列表"""
        unique_values = pd.Series(y.dropna().unique())
        if unique_values.size < 2:
            return False, "因变量类别数不足（至少需要2个类别）", None

        # 兼容新版本 pandas
        if isinstance(getattr(y, "dtype", None), pd.CategoricalDtype):
            if y.cat.ordered:
                return True, "", list(y.cat.categories)
            else:
                try:
                    categories = list(y.cat.categories)
                    return True, "", categories
                except Exception:
                    return False, "因变量无法自动排序", None

        if pd.api.types.is_numeric_dtype(y):
            try:
                sorted_unique = sorted(unique_values.tolist())
                return True, "", sorted_unique
            except Exception:
                return False, "因变量值无法排序，不是有效的有序分类", None

        try:
            sorted_unique = sorted(unique_values.tolist())
            return True, "", sorted_unique
        except TypeError:
            return False, "因变量值无法排序，不是有效的有序分类", None

    def get_significance_mark(self, p_value):
        if p_value < 0.001:
            return "***"
        elif p_value < 0.01:
            return "**"
        elif p_value < 0.05:
            return "*"
        else:
            return ""

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 处理缺失值（会弹窗阻塞等待选择）
            df = self.handle_missing_values(df)

            # 检查数据是否包含足够的列
            if df.shape[1] < 2:
                raise ValueError("数据至少需要两列（自变量和因变量）。")

            # 假设最后一列是因变量，其余列是自变量
            X = df.iloc[:, :-1].copy()
            y = df.iloc[:, -1].copy()

            # 验证因变量是否为有序分类变量并获得排序后的类别列表
            is_valid, msg, categories = self.validate_dependent_variable(y)
            if not is_valid:
                raise ValueError(f"{LANGUAGES[self.current_language]['dv_check_failed']} {msg}")

            # 转换因变量为有序分类（使用 validate 返回的 categories，确保顺序明确）
            y = pd.Categorical(y, categories=categories, ordered=True)

            # 检查自变量是否为数值型：尝试自动转换非数值列
            non_numeric_cols = X.select_dtypes(exclude=['number']).columns.tolist()
            converted_cols = []
            factorized_cols = []

            if non_numeric_cols:
                for col in non_numeric_cols:
                    # 尝试用 to_numeric 转换（适用于字符串数字）
                    X[col] = pd.to_numeric(X[col], errors='coerce')
                    nan_ratio = X[col].isna().sum() / len(X)
                    if nan_ratio > 0.5:
                        # 使用 factorize 将类别编码为整数
                        codes, uniques = pd.factorize(df.iloc[:, :-1].iloc[:, X.columns.get_loc(col)])
                        X[col] = codes
                        factorized_cols.append(col)
                    else:
                        if X[col].isna().any():
                            X[col] = X[col].fillna(X[col].mean())
                        converted_cols.append(col)

            still_non_numeric = X.select_dtypes(exclude=['number']).columns.tolist()
            if still_non_numeric:
                raise ValueError(f"自变量包含无法转换为数值的列，请处理后重试：{', '.join(still_non_numeric)}")

            # 检查并处理多重共线性（VIF对话框为阻塞式）
            X, removed_vars, vif_data = self.check_multicollinearity(X)

            # 检查极端值（仅提示）
            self.check_outliers(X)

            # 构建模型并处理收敛问题
            logit_model = OrderedModel(y, X, distr='logit')

            # 尝试不同优化方法和迭代次数解决收敛问题
            methods = ['bfgs', 'lbfgs', 'newton-cg']
            maxiter_values = [1000, 2000, 5000]
            result = None
            converged = False
            used_method = ""
            used_iterations = 0

            for method in methods:
                for maxiter in maxiter_values:
                    try:
                        result = logit_model.fit(method=method, maxiter=maxiter, disp=False)
                        used_method = method
                        used_iterations = maxiter
                        mle_retvals = getattr(result, 'mle_retvals', {})
                        converged = bool(mle_retvals.get('converged', True))
                        break
                    except Exception:
                        continue
                if result is not None:
                    break

            if result is None:
                raise RuntimeError("模型无法拟合，请检查数据或简化模型")

            if not converged:
                Messagebox.show_warning(
                    "Convergence Warning",
                    LANGUAGES[self.current_language]['convergence_warning']
                )

            # 计算McFadden伪R²并约束在[0,1]
            try:
                mcfadden_r2 = 1 - (result.llf / result.llnull)
                if pd.isna(mcfadden_r2) or np.isinf(mcfadden_r2):
                    mcfadden_r2 = np.nan
                else:
                    if not np.isnan(mcfadden_r2) and mcfadden_r2 < 0:
                        print("Warning: 模型目前没有实际价值，需要从数据、变量选择或模型设定上重新排查问题")
            except Exception:
                mcfadden_r2 = np.nan

            # 预测类别（处理NaN概率与零概率行）
            pred_probs = result.predict()

            # 统一安全转换预测概率为 numpy 数组
            if isinstance(pred_probs, pd.DataFrame):
                try:
                    col_names = pred_probs.columns.astype(str)
                    if all([str(c) in col_names for c in categories]):
                        col_order = [str(c) for c in categories]
                        try:
                            pred_probs = pred_probs[col_order]
                        except Exception:
                            # 如果列对不上则忽略
                            pass
                except Exception:
                    pass
                pred_probs_np = to_numpy_safe(pred_probs).copy()
            else:
                pred_probs_np = to_numpy_safe(pred_probs).copy()

            # 填充NaN值为0
            pred_probs_np = np.nan_to_num(pred_probs_np)

            # 处理零概率行：使用最可能的类别（全局最频繁）
            if pred_probs_np.ndim == 1:
                pred_probs_np = pred_probs_np.reshape(-1, 1)

            row_sums = pred_probs_np.sum(axis=1)
            zero_sum_rows = row_sums == 0

            if np.any(zero_sum_rows):
                global_mode = int(np.argmax(pred_probs_np.sum(axis=0)))
                pred_probs_np[zero_sum_rows, :] = 0.0
                pred_probs_np[zero_sum_rows, global_mode] = 1.0

            # 计算预测类别索引并映射回 categories（使用安全映射）
            y_pred_idx = np.argmax(pred_probs_np, axis=1)

            categories_list = list(categories)
            # 如果列数与类别数不一致，做最小调整
            if pred_probs_np.shape[1] != len(categories_list):
                if pred_probs_np.shape[1] > len(categories_list):
                    # 截断多余列（保守策略）
                    pred_probs_np = pred_probs_np[:, :len(categories_list)]
                else:
                    pad_cols = len(categories_list) - pred_probs_np.shape[1]
                    pred_probs_np = np.hstack([pred_probs_np, np.zeros((pred_probs_np.shape[0], pad_cols))])
                y_pred_idx = np.argmax(pred_probs_np, axis=1)

            y_pred_vals = [categories_list[idx] if 0 <= idx < len(categories_list) else categories_list[0] for idx in y_pred_idx]
            y_pred_categorical = pd.Categorical(y_pred_vals, categories=categories_list, ordered=True)

            # 计算评估指标
            coefficients = result.params[:len(X.columns)]
            z_values = result.tvalues[:len(X.columns)]
            p_values = result.pvalues[:len(X.columns)]
            conf_int = result.conf_int()
            accuracy = accuracy_score(y, y_pred_categorical)
            weighted_kappa = cohen_kappa_score(y, y_pred_categorical, weights='linear')
            thresholds = result.params[len(X.columns):]
            threshold_z = result.tvalues[len(X.columns):] if len(result.tvalues) > len(X.columns) else pd.Series([])
            threshold_p = result.pvalues[len(X.columns):] if len(result.pvalues) > len(X.columns) else pd.Series([])
            conf_matrix = confusion_matrix(y, y_pred_categorical)

            # 准备系数数据
            coef_data = []
            for i, col in enumerate(X.columns):
                sig_mark = self.get_significance_mark(p_values.iloc[i])

                coef_data.append([
                    col,
                    coefficients.iloc[i],
                    z_values.iloc[i],
                    p_values.iloc[i],
                    sig_mark
                ])

            coef_df = pd.DataFrame(coef_data, columns=[
                LANGUAGES[self.current_language]['stat_names']["Variable"],
                LANGUAGES[self.current_language]['stat_names']["Coefficients"],
                LANGUAGES[self.current_language]['stat_names']["z-value"],
                LANGUAGES[self.current_language]['stat_names']["p-value"],
                LANGUAGES[self.current_language]['stat_names']["Significance"],
            ])

            # 创建阈值表格
            threshold_data = []
            num_categories = len(y.categories)
            for i, threshold in enumerate(thresholds):
                threshold_name = LANGUAGES[self.current_language]['stat_names']["Threshold {i}"].format(i=i + 1)
                if i < len(thresholds):
                    if i + 1 < num_categories:
                        boundary_desc = f"对应类别 {y.categories[i]} 和 {y.categories[i + 1]} 的边界"
                    else:
                        boundary_desc = f"对应类别 {y.categories[i]} 和更高类别的边界"
                sig_mark = self.get_significance_mark(threshold_p.iloc[i]) if i < len(threshold_p) else ""

                threshold_data.append([
                    threshold_name,
                    threshold,
                    threshold_z.iloc[i] if i < len(threshold_z) else "",
                    threshold_p.iloc[i] if i < len(threshold_p) else "",
                    sig_mark
                ])

            threshold_df = pd.DataFrame(threshold_data, columns=[
                LANGUAGES[self.current_language]['stat_names']["Variable"],
                LANGUAGES[self.current_language]['stat_names']["Coefficients"],
                LANGUAGES[self.current_language]['stat_names']["z-value"],
                LANGUAGES[self.current_language]['stat_names']["p-value"],
                LANGUAGES[self.current_language]['stat_names']["Significance"],
            ])

            # 模型汇总表格
            model_summary = pd.DataFrame({
                # 使用多语言指标名
                "Metric": [
                    LANGUAGES[self.current_language]['stat_names']["Accuracy"],
                    LANGUAGES[self.current_language]['stat_names']["Weighted Kappa"],
                    LANGUAGES[self.current_language]['stat_names']["Pseudo R-squared"],
                    LANGUAGES[self.current_language]['stat_names']["No. Observations"],
                    LANGUAGES[self.current_language]['stat_names']["Log-Likelihood"],
                    LANGUAGES[self.current_language]['stat_names']["LL-Null"],
                    LANGUAGES[self.current_language]['stat_names']["LLR p-value"],
                    LANGUAGES[self.current_language]['model_convergence'],
                    LANGUAGES[self.current_language]['optimization_method'],
                    LANGUAGES[self.current_language]['iterations_used']
                ],
                # 对应的值保持不变
                "Value": [
                    accuracy,
                    weighted_kappa,
                    mcfadden_r2,
                    len(y),
                    result.llf,
                    result.llnull,
                    getattr(result, 'llr_pvalue', np.nan),
                    "Converged" if converged else "Not converged",
                    used_method,
                    used_iterations
                ],
            })
            # 修改表格表头为多语言
            model_summary.columns = [
                LANGUAGES[self.current_language]['stat_names']["Metric"],  # 新增表头翻译
                LANGUAGES[self.current_language]['stat_names']["Value"]  # 新增表头翻译
            ]

            # 构建回归方程字符串
            equation_terms = []
            for i, col in enumerate(X.columns):
                coef = coefficients.iloc[i]
                if i == 0:
                    term = f"{coef:.4f}×{col}" if coef >= 0 else f"({coef:.4f})×{col}"
                else:
                    term = f"+ {coef:.4f}×{col}" if coef >= 0 else f"- {abs(coef):.4f}×{col}"
                equation_terms.append(term)

            linear_combination = "".join(equation_terms)

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建图片保存目录
                save_dir = os.path.dirname(save_path)
                image_dir = os.path.join(save_dir, "images")
                os.makedirs(image_dir, exist_ok=True)

                # 生成可视化图表（注意：这里传入 pred_probs_np）
                self._generate_visualizations(result, X, pred_probs_np, y, y_pred_categorical,
                                              conf_int, thresholds, categories_list, conf_matrix, image_dir)

                # 创建Word文档
                doc = Document()

                # 添加标题
                doc.add_heading(LANGUAGES[self.current_language]['analysis_results_title'], 0)

                # 添加数据处理说明（包括移除的VIF变量与自动编码说明）
                doc.add_heading(LANGUAGES[self.current_language]['data_processing_notes'], 1)
                if removed_vars:
                    doc.add_paragraph(
                        LANGUAGES[self.current_language]['vif_removed_note'].format(", ".join(removed_vars)))
                else:
                    doc.add_paragraph(LANGUAGES[self.current_language]['no_high_vif_vars'])

                if factorized_cols:
                    doc.add_paragraph(
                        LANGUAGES[self.current_language]['factorized_vars_note'].format(", ".join(factorized_cols)))

                if not converged:
                    doc.add_paragraph("警告：模型未完全收敛，结果仅供参考")

                # 添加回归方程
                doc.add_heading(LANGUAGES[self.current_language]['regression_equation'], 1)
                doc.add_paragraph(LANGUAGES[self.current_language]['equation_explanation'])
                doc.add_paragraph(LANGUAGES[self.current_language]['cumulative_prob'])
                doc.add_paragraph(LANGUAGES[self.current_language]['linear_combination'].format(linear_combination))
                doc.add_paragraph(LANGUAGES[self.current_language]['equation_terms_explanation'])

                # 添加模型摘要
                doc.add_heading(LANGUAGES[self.current_language]['model_summary_heading'], 1)
                self._add_dataframe_to_doc(doc, model_summary)

                # 添加混淆矩阵
                doc.add_heading(LANGUAGES[self.current_language]['confusion_matrix_heading'], 1)
                conf_matrix_df = pd.DataFrame(
                    conf_matrix,
                    index=[LANGUAGES[self.current_language]['true_category'].format(c) for c in categories_list],
                    columns=[LANGUAGES[self.current_language]['predicted_category'].format(c) for c in categories_list]
                )
                self._add_dataframe_to_doc(doc, conf_matrix_df)

                # 添加系数表格
                doc.add_heading(LANGUAGES[self.current_language]['coefficients_heading'], 1)
                self._add_dataframe_to_doc(doc, coef_df)

                # 添加系数显著性条形图
                doc.add_heading(LANGUAGES[self.current_language]['coef_significance_heading'], 2)
                doc.add_picture(os.path.join(image_dir, "coefficients.png"), width=Cm(15))

                # 添加阈值表格
                doc.add_heading(LANGUAGES[self.current_language]['thresholds_heading'], 1)
                self._add_dataframe_to_doc(doc, threshold_df)

                # 添加阈值可视化
                doc.add_heading(LANGUAGES[self.current_language]['threshold_visualization_heading'], 2)
                doc.add_picture(os.path.join(image_dir, "thresholds.png"), width=Cm(15))

                # 添加预测概率分布图
                doc.add_heading(LANGUAGES[self.current_language]['pred_prob_distribution_heading'], 2)
                doc.add_picture(os.path.join(image_dir, "pred_prob.png"), width=Cm(15))

                # 添加混淆矩阵热图
                doc.add_heading(LANGUAGES[self.current_language]['confusion_matrix_heading'], 2)
                doc.add_picture(os.path.join(image_dir, "confusion_matrix.png"), width=Cm(15))

                # 添加解释说明
                doc.add_heading(LANGUAGES[self.current_language]['explanations_heading'], 1)
                explanations = LANGUAGES[self.current_language]['explanation']
                for stat_key, explanation in explanations.items():
                    # 使用本地化的统计量名称
                    stat_name = LANGUAGES[self.current_language]['stat_names'].get(stat_key, stat_key)
                    doc.add_paragraph(f"{stat_name}: {explanation}")

                # 保存Word文档
                doc.save(save_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(save_path)
                self.result_label.config(text=result_msg, wraplength=400)
            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def _generate_visualizations(self, result, X, pred_probs, y, y_pred, conf_int, thresholds, categories, conf_matrix,
                                 image_dir):
        """生成可视化图表并保存（优化版本，已做 numpy/pandas 兼容）"""
        # 系数显著性条形图（带置信区间）
        plt.figure(figsize=(12, 8))
        coefs = result.params[:len(X.columns)]
        coefs = coefs.drop(index=[i for i in coefs.index if 'threshold' in i.lower()])

        # 获取对应的置信区间
        try:
            ci_lower = conf_int.loc[coefs.index, 0]
            ci_upper = conf_int.loc[coefs.index, 1]
        except Exception:
            se = result.bse[:len(X.columns)]
            ci_lower = coefs - 1.96 * se
            ci_upper = coefs + 1.96 * se

        x_pos = np.arange(len(coefs))
        plt.bar(x_pos, coefs.values, yerr=[coefs.values - ci_lower, ci_upper - coefs.values],
                capsize=5, error_kw={'ecolor': 'darkred', 'capthick': 2})
        plt.axhline(y=0, color='r', linestyle='-', alpha=0.3)
        plt.xticks(x_pos, coefs.index, rotation=0, ha='center')
        plt.title(
            '回归系数及其95%置信区间' if self.current_language == 'zh' else 'Regression Coefficients with 95% Confidence Intervals')
        plt.tight_layout()
        plt.savefig(os.path.join(image_dir, "coefficients.png"))
        plt.close()

        # 阈值可视化
        plt.figure(figsize=(10, 6))
        try:
            thr_vals = np.array(thresholds)
        except Exception:
            thr_vals = np.array([])

        if thr_vals.size > 0:
            plt.vlines(thr_vals, ymin=0, ymax=1, color='blue', linestyle='--',
                       label='阈值' if self.current_language == 'zh' else 'Thresholds')
            for i, thr in enumerate(thr_vals):
                plt.text(thr, 1.02, LANGUAGES[self.current_language]['threshold_label'].format(i + 1),
                    rotation=90)

        # 添加类别标签（根据阈值位置估计）
        if len(categories) > 1:
            label_positions = []
            if thr_vals.size == 0:
                positions = np.linspace(-1, 1, len(categories))
                label_positions = positions
            else:
                label_positions.append(thr_vals[0] - 0.5)
                for i in range(len(thr_vals) - 1):
                    label_positions.append((thr_vals[i] + thr_vals[i + 1]) / 2)
                label_positions.append(thr_vals[-1] + 0.5)

            for pos, cat in zip(label_positions[:len(categories)], categories):
                plt.text(
                    pos, 0.5,
                    LANGUAGES[self.current_language]['category_label'].format(cat),
                    ha='center', va='center',
                    bbox=dict(facecolor='white', edgecolor='gray')
                )

        plt.ylim(0, 1.1)
        plt.yticks([])
        plt.title('有序分类阈值' if self.current_language == 'zh' else 'Ordered Classification Thresholds')
        plt.xlabel('潜变量值' if self.current_language == 'zh' else 'Latent Variable Value')
        plt.legend()
        plt.tight_layout()
        plt.savefig(os.path.join(image_dir, "thresholds.png"))
        plt.close()

        # 预测概率分布图（按真实类别分组）
        plt.figure(figsize=(12, 8))
        num_categories = len(categories)
        colors = plt.cm.Set3(np.linspace(0, 1, num_categories))

        # 安全转换 pred_probs 为 ndarray
        pred_probs_arr = to_numpy_safe(pred_probs)

        for true_idx, true_cat in enumerate(categories):
            mask = to_numpy_safe(y == true_cat)
            if mask.ndim > 1:
                mask = mask.ravel()
            if not np.any(mask):
                continue

            mean_probs = pred_probs_arr[mask].mean(axis=0)
            x = np.arange(len(mean_probs))

            plt.bar(
                x + true_idx * 0.15,
                mean_probs,
                width=0.15,
                label=LANGUAGES[self.current_language]['true_category'].format(true_cat)
            )

        plt.xticks(np.arange(num_categories) + 0.15 * (num_categories - 1) / 2,
            [LANGUAGES[self.current_language]['predicted_category'].format(c) for c in categories])
        plt.title('按真实类别分组的平均预测概率分布' if self.current_language == 'zh' else
                  'Average Predicted Probability Distribution by True Category')
        plt.xlabel('预测类别' if self.current_language == 'zh' else 'Predicted Category')
        plt.ylabel('平均概率' if self.current_language == 'zh' else 'Average Probability')
        plt.legend(bbox_to_anchor=(1.05, 1), loc='upper left')
        plt.tight_layout()
        plt.savefig(os.path.join(image_dir, "pred_prob.png"))
        plt.close()

        # 混淆矩阵热图
        plt.figure(figsize=(10, 8))
        cm = to_numpy_safe(conf_matrix)
        plt.imshow(cm, cmap='Blues', interpolation='nearest')
        plt.colorbar(label='数量' if self.current_language == 'zh' else 'Count')

        for i in range(cm.shape[0]):
            for j in range(cm.shape[1]):
                plt.text(j, i, str(int(cm[i, j])),
                         horizontalalignment='center',
                         verticalalignment='center',
                         color='white' if cm[i, j] > cm.max() / 2 else 'black')

        plt.xticks(np.arange(len(categories)), categories)
        plt.yticks(np.arange(len(categories)), categories)
        plt.xlabel('预测类别' if self.current_language == 'zh' else 'Predicted Category')
        plt.ylabel('真实类别' if self.current_language == 'zh' else 'True Category')
        plt.title('混淆矩阵' if self.current_language == 'zh' else 'Confusion Matrix')
        plt.tight_layout()
        plt.savefig(os.path.join(image_dir, "confusion_matrix.png"))
        plt.close()

    def _add_dataframe_to_doc(self, doc, df):
        """将DataFrame添加到Word文档中"""
        table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
        hdr_cells = table.rows[0].cells
        for col_idx, header in enumerate(df.columns):
            hdr_cells[col_idx].text = header

        for row_idx, row in enumerate(df.values):
            row_cells = table.rows[row_idx + 1].cells
            for col_idx, value in enumerate(row):
                if isinstance(value, float):
                    if pd.isna(value):
                        row_cells[col_idx].text = ""
                    else:
                        row_cells[col_idx].text = f"{value:.4f}"
                else:
                    row_cells[col_idx].text = str(value)

    def switch_language(self, event=None):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(LANGUAGES[self.current_language]["title"])
        self.select_button.config(text=LANGUAGES[self.current_language]["select_button"])
        current_text = self.file_entry.get()
        if current_text == LANGUAGES['zh']["file_entry_placeholder"] or current_text == LANGUAGES['en'][
            "file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=LANGUAGES[self.current_language]["analyze_button"])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]["switch_language"])
        self.open_excel_label.config(text=LANGUAGES[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        self.select_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["select_button"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        self.analyze_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["analyze_button"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=LANGUAGES[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        self.switch_language_label = ttk.Label(frame, text=LANGUAGES[self.current_language]["switch_language"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = OrderedLogitRegressionAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()

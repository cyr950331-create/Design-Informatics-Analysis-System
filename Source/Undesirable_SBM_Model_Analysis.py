import tkinter as tk
from tkinter import filedialog
import tkinter.messagebox as messagebox
import os
import pandas as pd
import numpy as np
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from scipy.optimize import linprog

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
languages = {
    'zh': {
        'title': "非期望 SBM 模型",
        'select_button_text': "选择文件",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'analyze_button_text': "分析文件",
        'no_file_selected': "请选择文件。",
        'file_not_exists': "文件不存在，请重新选择。",
        'analysis_error': "分析文件时出错: {}",
        'analysis_complete': "分析完成，结果已保存到 {}，相关图片已保存。",
        'no_save_path_selected': "未选择保存路径，结果未保存。",
        'switch_language_button_text': "中/英",
        'open_excel_button_text': "示例数据",
        'explanation': {
            "效率值": "各决策单元的非期望SBM效率得分",
        },
        'interpretation': {
            "效率值": "效率值越接近1，决策单元的效率越高",
        },
        'data_valid': "数据验证通过",
        'missing_cols': "数据缺少必要列，请确保包含以下类型的列：{}",
        'non_numeric': "以下列包含非数值数据：{}",
        'missing_values': "数据中存在缺失值，请检查：{}",
        'invalid_column_prefix': "列名前缀错误，正确格式应为：Input_*, Output_*, UndesirableOutput_*（区分大小写）",
        'extreme_values': "检测到极端值，请检查以下列：{}",
        'param_dialog_title': "模型参数设置",
        'orientation_label': "导向性:",
        'returns_label': "规模报酬:",
        'input_oriented': "投入导向",
        'output_oriented': "产出导向",
        'constant': "不变规模报酬",
        'variable': "可变规模报酬",
        'confirm': "确认",
        'invalid_dmu': "无效决策单元分析",
        'improvement_suggestion': "改进建议",
        'reduce_input': "应减少投入: {}",
        'adjust_output': "应增加期望产出: {}",
        'adjust_undesirable': "应减少非期望产出: {}",
        'efficiency_value': "效率值",
        'efficiency_bar_chart': "效率值柱状图",
        'improvement_direction': "改进方向",
        'statistics': "统计量",
        'statistic_value': "统计量值",
        'p_value': "p值",
        'explanation_label': "解释说明",
        'interpretation_label': "结果解读",
        'decision_units': "决策单元",
        'model_parameters': "模型参数",
        'efficiency_results': "效率值结果",
        'efficiency_frontier': "效率边界",
        'solve_failed': "决策单元 {} 求解失败: {}"
    },
    'en': {
        'title': "Undesirable SBM Model",
        'select_button_text': "Select File",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'analyze_button_text': "Analyze File",
        'no_file_selected': "Please select a valid file path.",
        'file_not_exists': "The file does not exist. Please select again.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'analysis_complete': "Analysis completed. The results have been saved to {}, and the relevant images have been saved.",
        'no_save_path_selected': "No save path selected. The results were not saved.",
        'switch_language_button_text': "Chinese/English",
        'open_excel_button_text': "Example data",
        'explanation': {
            "efficiency": "The efficiency scores of each decision unit based on the undesirable SBM model",
        },
        'interpretation': {
            "efficiency": "The closer the efficiency value is to 1, the higher the efficiency of the decision unit",
        },
        'data_valid': "Data validation passed",
        'missing_cols': "Data lacks necessary columns types. Please ensure it contains: {}",
        'non_numeric': "The following columns(s) contain non-numeric data: {}",
        'missing_values': "There are missing values in: {}",
        'invalid_column_prefix': "Invalid column prefixes. Correct formats: Input_*, Output_*, UndesirableOutput_* (case-sensitive)",
        'extreme_values': "Extreme values detected in columns: {}",
        'param_dialog_title': "Model Parameter Settings",
        'orientation_label': "Orientation:",
        'returns_label': "Returns to scale:",
        'input_oriented': "Input oriented",
        'output_oriented': "Output oriented",
        'constant': "Constant returns to scale",
        'variable': "Variable returns to scale",
        'confirm': "Confirm",
        'invalid_dmu': "Inefficient DMU Analysis",
        'improvement_suggestion': "Improvement Suggestions",
        'reduce_input': "Should reduce input by: {}",
        'adjust_output': "Should increase desirable output by: {}",
        'adjust_undesirable': "Should reduce undesirable output by: {}",
        'efficiency_value': "Efficiency Value",
        'efficiency_bar_chart': "Bar Chart of Efficiency Values",
        'improvement_direction': "Improvement Direction",
        'statistics': "Statistics",
        'statistic_value': "Statistic Value",
        'p_value': "p-value",
        'explanation_label': "Explanation",
        'interpretation_label': "Interpretation",
        'decision_units': "Decision Units",
        'model_parameters': "Model Parameters",
        'efficiency_results': "Efficiency Results",
        'efficiency_frontier': "Efficiency Frontier",
        'solve_failed': "DMU {} solving failed: {}"
    }
}


# 改进的非期望SBM模型实现（修复无界问题）
class ImprovedSBMModel:
    def __init__(self, orientation, returns_to_scale, input_cols, output_cols, undesirable_cols):
        self.n_inputs = len(input_cols)
        self.n_outputs = len(output_cols)
        self.n_undesirables = len(undesirable_cols)
        self.orientation = orientation  # "Input_Oriented" 或 "Output_Oriented"
        self.returns_to_scale = returns_to_scale  # "constant" 或 "variable"
        self.input_max = None  # 新增：保存输入最大值用于还原松弛变量
        self.input_min = None  # 新增：保存输入最小值用于还原松弛变量
        self.output_max = None
        self.output_min = None
        self.undesirable_max = None
        self.undesirable_min = None
        self.data = None
        self.dmu_names = None
        self.original_data = None  # 保存原始数据用于计算改进量
        self.slacks = {}  # 保存各DMU的松弛变量: {dmu_name: (input_slacks, output_slacks, undesirable_slacks)}
        self.epsilon = 1e-8  # 防止除以零的小值

    def fit(self, dmu_names, data, original_data):
        """拟合模型数据，新增original_data参数保存原始数据"""
        self.dmu_names = dmu_names
        self.data = data  # 标准化后的数据
        self.original_data = original_data  # 原始数据（未标准化）
        self.input_min = np.min(original_data[:, :self.n_inputs], axis=0)
        self.input_max = np.max(original_data[:, :self.n_inputs], axis=0)
        self.output_min = np.min(original_data[:, self.n_inputs:self.n_inputs + self.n_outputs], axis=0)
        self.output_max = np.max(original_data[:, self.n_inputs:self.n_inputs + self.n_outputs], axis=0)
        self.undesirable_min = np.min(original_data[:, self.n_inputs + self.n_outputs:], axis=0)
        self.undesirable_max = np.max(original_data[:, self.n_inputs + self.n_outputs:], axis=0)

    # 添加松弛变量还原方法
    def denormalize_slacks(self, normalized_slacks, min_vals, max_vals):
        """将标准化的松弛变量还原为原始数据尺度"""
        range_vals = max_vals - min_vals
        range_vals[range_vals == 0] = 1.0  # 处理极差为0的情况
        # 还原公式：基于标准化公式的逆运算
        return normalized_slacks * 0.9 * range_vals / 0.9

    def run(self):
        """计算各决策单元效率值 - 修复无界问题"""
        if self.data is None or self.dmu_names is None or self.original_data is None:
            raise ValueError("Model not fitted with data")

        n_dmus = len(self.dmu_names)
        efficiency_scores = []
        solve_errors = []  # 记录求解失败信息

        for i in range(n_dmus):
            # 提取当前决策单元数据（原始数据用于松弛变量计算）
            x0_original = self.original_data[i, :self.n_inputs]  # 原始输入
            y0_original = self.original_data[i, self.n_inputs:self.n_inputs + self.n_outputs]  # 原始期望输出
            b0_original = self.original_data[i, self.n_inputs + self.n_outputs:]  # 原始非期望输出

            # 处理原始数据中的零值，避免除以零
            x0_original = np.where(x0_original < self.epsilon, self.epsilon, x0_original)
            y0_original = np.where(y0_original < self.epsilon, self.epsilon, y0_original)
            b0_original = np.where(b0_original < self.epsilon, self.epsilon, b0_original)

            # 提取标准化数据用于建模
            x0 = self.data[i, :self.n_inputs]
            y0 = self.data[i, self.n_inputs:self.n_inputs + self.n_outputs]
            b0 = self.data[i, self.n_inputs + self.n_outputs:]

            # 构建线性规划问题
            n_vars = 1 + n_dmus + self.n_inputs + self.n_outputs + self.n_undesirables

            # 目标函数系数
            c = np.zeros(n_vars)
            c[0] = 1.0  # θ的系数为1

            # 松弛变量系数
            for k in range(self.n_inputs):
                c[1 + n_dmus + k] = 1.0 / self.n_inputs / x0_original[k]

            sum_output_undesirable = self.n_outputs + self.n_undesirables
            for k in range(self.n_outputs):
                c[1 + n_dmus + self.n_inputs + k] = 1.0 / sum_output_undesirable / y0_original[k]

            for k in range(self.n_undesirables):
                c[1 + n_dmus + self.n_inputs + self.n_outputs + k] = 1.0 / sum_output_undesirable / b0_original[k]

            # 约束条件构建
            A_eq = []
            b_eq = []
            A_ub = []
            b_ub = []

            # 输入约束: Σλ_j x_jk + s_k^- = θ x0k (k=1..n_inputs)
            for k in range(self.n_inputs):
                row = np.zeros(n_vars)
                row[0] = -x0[k]  # -θ x0k
                row[1:n_dmus + 1] = self.data[:, k]  # λ_j x_jk
                row[1 + n_dmus + k] = 1  # s_k^-
                A_eq.append(row)
                b_eq.append(0.0)

            # 期望输出约束: Σλ_j y_jk - s_k^+ = y0k (k=1..n_outputs)
            for k in range(self.n_outputs):
                row = np.zeros(n_vars)
                row[1:n_dmus + 1] = self.data[:, self.n_inputs + k]  # λ_j y_jk
                row[1 + n_dmus + self.n_inputs + k] = -1  # -s_k^+
                A_eq.append(row)
                b_eq.append(y0[k])

            # 非期望输出约束: Σλ_j b_jk + s_k^b = b0k (k=1..n_undesirables)
            for k in range(self.n_undesirables):
                row = np.zeros(n_vars)
                row[1:n_dmus + 1] = self.data[:, self.n_inputs + self.n_outputs + k]  # λ_j b_jk
                row[1 + n_dmus + self.n_inputs + self.n_outputs + k] = 1  # s_k^b
                A_eq.append(row)
                b_eq.append(b0[k])

            # 规模报酬约束
            if self.returns_to_scale == "constant":
                row = np.zeros(n_vars)
                row[1:n_dmus + 1] = 1.0  # sum(lambda) = 1
                A_eq.append(row)
                b_eq.append(1.0)
            else:
                # 为VRS添加lambda总和的上下界约束，防止无界
                row = np.zeros(n_vars)
                row[1:n_dmus + 1] = 1.0  # sum(lambda) <= 1 + 1e-6
                A_ub.append(row)
                b_ub.append(1.0 + 1e-6)

                row = np.zeros(n_vars)
                row[1:n_dmus + 1] = -1.0  # sum(lambda) >= 0 + 1e-6
                A_ub.append(row)
                b_ub.append(-1e-6)

            # 导向性约束
            if self.orientation == "Input_Oriented":
                row = np.zeros(n_vars)
                row[0] = 1.0  # θ ≤ 1
                A_ub.append(row)
                b_ub.append(1.0)
            else:
                # 产出导向添加theta的下界约束
                row = np.zeros(n_vars)
                row[0] = -1.0  # θ >= 1e-6
                A_ub.append(row)
                b_ub.append(-1e-6)

            # 变量边界: 为所有变量添加合理边界，防止无界
            bounds = [(1e-6, 1.0 + 1e-6)]  # θ的边界：投入导向<=1，产出导向>=小值
            # lambda变量边界：0到n_dmus（合理上限）
            bounds += [(0, n_dmus) for _ in range(n_dmus)]
            # 松弛变量边界：0到原始值的2倍（防止无限大）
            input_bounds = [(0, 2 * max_val) for max_val in np.max(self.original_data[:, :self.n_inputs], axis=0)]
            output_bounds = [(0, 2 * max_val) for max_val in
                             np.max(self.original_data[:, self.n_inputs:self.n_inputs + self.n_outputs], axis=0)]
            undesirable_bounds = [(0, 2 * max_val) for max_val in
                                  np.max(self.original_data[:, self.n_inputs + self.n_outputs:], axis=0)]
            bounds += input_bounds + output_bounds + undesirable_bounds

            # 求解线性规划，使用HiGHS求解器的正确参数名称
            result = linprog(
                c,
                A_eq=A_eq,
                b_eq=b_eq,
                A_ub=A_ub,
                b_ub=b_ub,
                bounds=bounds,
                method='highs',
                options={
                    'max_iteration': 100000,
                    'dual_feasibility_tolerance': 1e-9,
                    'primal_feasibility_tolerance': 1e-9,
                    'presolve': True,
                    'time_limit': 30.0
                }
            )

            if result.success:
                # 提取结果变量
                theta = result.x[0]
                input_slacks = result.x[1 + n_dmus: 1 + n_dmus + self.n_inputs]
                output_slacks = result.x[1 + n_dmus + self.n_inputs: 1 + n_dmus + self.n_inputs + self.n_outputs]
                undesirable_slacks = result.x[1 + n_dmus + self.n_inputs + self.n_outputs:]

                # 计算SBM效率值
                input_sum = 0.0
                for s, x in zip(input_slacks, x0_original):
                    input_sum += s / x

                output_sum = 0.0
                for s, y in zip(output_slacks, y0_original):
                    output_sum += s / y

                undesirable_sum = 0.0
                for s, b in zip(undesirable_slacks, b0_original):
                    undesirable_sum += s / b

                # 根据导向性计算不同的效率值
                if self.orientation == "Input_Oriented":
                    efficiency = theta / (1 + (input_sum / self.n_inputs))
                else:  # Output_Oriented
                    efficiency = 1 / (theta * (1 + ((output_sum + undesirable_sum) / sum_output_undesirable)))

                # 限制效率值范围在0到1之间
                efficiency = max(0.0001, min(1.0, round(efficiency, 4)))
                efficiency_scores.append(efficiency)
                # 保存松弛变量
                self.slacks[self.dmu_names[i]] = (input_slacks, output_slacks, undesirable_slacks)

            else:

                # 求解失败时尝试使用替代方法
                efficiency = np.nan
                solve_errors.append(languages[self.current_language]["solve_failed"].format(
                    self.dmu_names[i], result.message))
                # 尝试使用简化模型重新求解
                try:
                    # 使用简化的目标函数重新求解
                    c_simple = np.zeros(n_vars)
                    c_simple[0] = 1.0  # 仅优化theta
                    result_simple = linprog(
                        c_simple,
                        A_eq=A_eq,
                        b_eq=b_eq,
                        A_ub=A_ub,
                        b_ub=b_ub,
                        bounds=bounds,
                        method='highs'
                    )

                    if result_simple.success:
                        theta_simple = result_simple.x[0]
                        efficiency = max(0.0001, min(1.0, round(theta_simple, 4)))
                        solve_errors[-1] += " - 使用简化模型求解成功"
                except:
                    pass
                efficiency_scores.append(efficiency)

        # 提示求解失败信息
        if solve_errors:
            messagebox.showwarning("求解警告", "\n".join(solve_errors))

        return efficiency_scores


# 自定义结果包装类
class SBMResult:
    def __init__(self, dmu_list, scores, input_cols, output_cols, undesirable_cols, slacks):
        self.dmu_list = dmu_list
        self.scores = scores
        self.input_categories = input_cols
        self.output_categories = output_cols
        self.undesirable_output_categories = undesirable_cols
        self.slacks = slacks  # 保存松弛变量

    def get_efficiency_score(self, dmu):
        idx = self.dmu_list.index(dmu)
        return self.scores[idx]

    def get_slacks(self, dmu):
        return self.slacks.get(dmu, (np.zeros(len(self.input_categories)),
                                     np.zeros(len(self.output_categories)),
                                     np.zeros(len(self.undesirable_output_categories))))


# 自定义数据读取函数（改进数据标准化）
def custom_read_data(file_path):
    """替代pyDEA的read_data，返回数据结构与原库兼容的结果，改进数据标准化处理"""
    df = pd.read_excel(file_path)

    # 提取决策单元名称（第一列）
    dmu_names = df.iloc[:, 0].tolist()

    # 严格区分列名前缀（区分大小写）
    input_cols = [col for col in df.columns if col.startswith('Input_')]
    output_cols = [col for col in df.columns if col.startswith('Output_')]
    undesirable_cols = [col for col in df.columns if col.startswith('UndesirableOutput_')]

    # 提取原始数值数据
    input_data = df[input_cols].values
    output_data = df[output_cols].values
    undesirable_data = df[undesirable_cols].values
    original_data = np.hstack((input_data, output_data, undesirable_data))

    # 数据标准化（使用Min-Max标准化替代Z-score，更适合效率分析）
    def normalize_data(data, is_input=True):
        """改进的标准化方法，对投入和产出使用不同策略"""
        min_val = np.min(data, axis=0)
        max_val = np.max(data, axis=0)
        range_val = max_val - min_val
        range_val[range_val == 0] = 1.0  # 处理极差为0的情况

        if is_input:
            # 投入指标：越大越不利，标准化到[0.1, 1.0]
            return 0.1 + 0.9 * (data - min_val) / range_val
        else:
            # 产出指标：越大越有利，标准化到[0.1, 1.0]，但保持相对比例
            return 0.1 + 0.9 * (data - min_val) / range_val

    # 在custom_read_data中使用改进的标准化
    normalized_input = normalize_data(input_data, is_input=True)
    normalized_output = normalize_data(output_data, is_input=False)
    normalized_undesirable = normalize_data(undesirable_data, is_input=True)  # 非期望产出按投入处理

    # 合并标准化数据
    combined_data = np.hstack((normalized_input, normalized_output, normalized_undesirable))

    return dmu_names, input_cols, output_cols, undesirable_cols, combined_data, original_data


class UndesirableSBMModelAnalysisApp:
    def __init__(self, root=None):
        self.current_language = "en"
        self.orientation = "Input_Oriented"
        self.returns_to_scale = "variable"
        self.original_data = None  # 保存原始数据用于改进建议计算

        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data51.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def validate_data_format(self, file_path):
        try:
            df = pd.read_excel(file_path)
            columns = df.columns.tolist()

            # 严格检查列名前缀（区分大小写）
            has_input = any(col.startswith('Input_') for col in columns)
            has_output = any(col.startswith('Output_') for col in columns)
            has_undesirable = any(col.startswith('UndesirableOutput_') for col in columns)

            if not (has_input and has_output and has_undesirable):
                missing = []
                if not has_input:
                    missing.append("Input_*")
                if not has_output:
                    missing.append("Output_*")
                if not has_undesirable:
                    missing.append("UndesirableOutput_*")
                return False, languages[self.current_language]["missing_cols"].format(", ".join(missing))

            # 检查是否有不符合格式的列（排除第一列和正确前缀的列）
            valid_prefixes = {'Input_', 'Output_', 'UndesirableOutput_'}
            invalid_cols = []
            for col in columns[1:]:  # 跳过第一列（DMU名称）
                prefix = col.split('_')[0] + '_' if '_' in col else ''
                if prefix not in valid_prefixes:
                    invalid_cols.append(col)
            if invalid_cols:
                return False, languages[self.current_language]["invalid_column_prefix"]

            # 检查非数值数据
            non_numeric_cols = []
            for col in columns[1:]:
                if not pd.api.types.is_numeric_dtype(df[col]):
                    non_numeric_cols.append(col)
            if non_numeric_cols:
                return False, languages[self.current_language]["non_numeric"].format(", ".join(non_numeric_cols))

            # 检查缺失值
            missing_values = df.isnull().sum()
            missing_cols = [col for col, count in missing_values.items() if count > 0]
            if missing_cols:
                return False, languages[self.current_language]["missing_values"].format(", ".join(missing_cols))

            # 检查极端值（使用IQR方法）
            extreme_cols = []
            for col in columns[1:]:
                data = df[col].dropna()
                q1 = data.quantile(0.25)
                q3 = data.quantile(0.75)
                iqr = q3 - q1
                lower_bound = q1 - 1.5 * iqr
                upper_bound = q3 + 1.5 * iqr
                if (data < lower_bound).any() or (data > upper_bound).any():
                    extreme_cols.append(col)
            if extreme_cols:
                # 弹出框相对于主窗口居中
                messagebox.showwarning(
                    "极端值警告",
                    languages[self.current_language]["extreme_values"].format(", ".join(extreme_cols)),
                    parent=self.root
                )

            return True, languages[self.current_language]["data_valid"]
        except Exception as e:
            return False, str(e)

    def show_parameter_dialog(self):
        dialog = ttk.Toplevel(self.root)
        dialog.title(languages[self.current_language]["param_dialog_title"])
        dialog.geometry("300x200")
        dialog.transient(self.root)  # 设置为主窗口的子窗口
        dialog.grab_set()

        # 计算屏幕中央位置并设置对话框位置
        dialog.update_idletasks()  # 确保获取正确的窗口尺寸
        width = dialog.winfo_width()
        height = dialog.winfo_height()
        x = (self.root.winfo_width() // 2) + self.root.winfo_x() - (width // 2)
        y = (self.root.winfo_height() // 2) + self.root.winfo_y() - (height // 2)
        dialog.geometry(f"+{x}+{y}")

        ttk.Label(dialog, text=languages[self.current_language]["orientation_label"]).pack(pady=5)
        orientation_frame = ttk.Frame(dialog)
        orientation_frame.pack(pady=5)

        orientation_var = tk.StringVar(value=self.orientation)
        ttk.Radiobutton(
            orientation_frame,
            text=languages[self.current_language]["input_oriented"],
            variable=orientation_var,
            value="Input_Oriented"
        ).pack(side=LEFT, padx=5)
        ttk.Radiobutton(
            orientation_frame,
            text=languages[self.current_language]["output_oriented"],
            variable=orientation_var,
            value="Output_Oriented"
        ).pack(side=LEFT, padx=5)

        ttk.Label(dialog, text=languages[self.current_language]["returns_label"]).pack(pady=5)
        returns_frame = ttk.Frame(dialog)
        returns_frame.pack(pady=5)

        returns_var = tk.StringVar(value=self.returns_to_scale)
        ttk.Radiobutton(
            returns_frame,
            text=languages[self.current_language]["variable"],
            variable=returns_var,
            value="variable"
        ).pack(side=LEFT, padx=5)
        ttk.Radiobutton(
            returns_frame,
            text=languages[self.current_language]["constant"],
            variable=returns_var,
            value="constant"
        ).pack(side=LEFT, padx=5)

        def confirm():
            self.orientation = orientation_var.get()
            self.returns_to_scale = returns_var.get()
            dialog.destroy()

        ttk.Button(dialog, text=languages[self.current_language]["confirm"], command=confirm).pack(pady=10)
        self.root.wait_window(dialog)

    def get_improvement_suggestions(self, results, original_data_df, model):
        """改进建议计算，使用还原后的松弛变量"""
        suggestions = {}
        input_cols = results.input_categories
        output_cols = results.output_categories
        undesirable_cols = results.undesirable_output_categories

        for dmu in results.dmu_list:
            score = results.get_efficiency_score(dmu)
            if score < 1 - 1e-9:  # 考虑浮点数精度问题
                input_slacks, output_slacks, undesirable_slacks = results.get_slacks(dmu)

                # 还原松弛变量到原始数据尺度
                input_slacks_original = model.denormalize_slacks(
                    input_slacks, model.input_min, model.input_max)
                output_slacks_original = model.denormalize_slacks(
                    output_slacks, model.output_min, model.output_max)
                undesirable_slacks_original = model.denormalize_slacks(
                    undesirable_slacks, model.undesirable_min, model.undesirable_max)

                # 计算具体调整量（基于原始数据）
                input_improve = {}
                for col, slack, idx in zip(input_cols, input_slacks_original, range(len(input_cols))):
                    original_value = original_data_df.loc[original_data_df.iloc[:, 0] == dmu, col].values[0]
                    adjust_amount = round(slack, 4)
                    # 计算调整百分比
                    pct_change = round((adjust_amount / original_value) * 100, 2) if original_value != 0 else 0
                    input_improve[col] = f"{adjust_amount} ({pct_change}%)"

                output_improve = {}
                for col, slack in zip(output_cols, output_slacks):
                    original_value = original_data_df.loc[original_data_df.iloc[:, 0] == dmu, col].values[0]
                    adjust_amount = round(slack, 4)
                    output_improve[col] = f"{adjust_amount} (原始值: {original_value})"

                undesirable_improve = {}
                for col, slack in zip(undesirable_cols, undesirable_slacks):
                    original_value = original_data_df.loc[original_data_df.iloc[:, 0] == dmu, col].values[0]
                    adjust_amount = round(slack, 4)
                    undesirable_improve[col] = f"{adjust_amount} (原始值: {original_value})"

                suggestions[dmu] = {
                    'efficiency': round(score, 4),
                    'input_reduction': input_improve,
                    'output_increase': output_improve,
                    'undesirable_reduction': undesirable_improve
                }
        return suggestions

    def undesirable_sbm_analysis(self, data):
        """使用改进的SBM模型计算"""
        # 读取数据（包含原始数据）
        dmu_names, input_categories, output_categories, undesirable_output_categories, combined_data, original_data = custom_read_data(
            data)
        # 保存原始数据DataFrame用于改进建议
        original_data_df = pd.read_excel(data)

        # 构建并运行改进的SBM模型
        model = ImprovedSBMModel(
            self.orientation,
            self.returns_to_scale,
            input_categories,
            output_categories,
            undesirable_output_categories
        )
        model.current_language = self.current_language  # 传递语言设置用于错误提示
        model.fit(dmu_names, combined_data, original_data)
        efficiency_scores = model.run()

        # 包装结果（包含松弛变量）
        results = SBMResult(
            dmu_list=dmu_names,
            scores=efficiency_scores,
            input_cols=input_categories,
            output_cols=output_categories,
            undesirable_cols=undesirable_output_categories,
            slacks=model.slacks
        )

        # 提取结果
        dmu_names = results.dmu_list
        efficiency_values = [round(results.get_efficiency_score(dmu), 4) for dmu in dmu_names]
        improvement_suggestions = self.get_improvement_suggestions(results, original_data_df, model)

        return efficiency_values, dmu_names, improvement_suggestions

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            is_valid, msg = self.validate_data_format(file_path)
            if not is_valid:
                # 错误弹出框相对于主窗口居中
                messagebox.showerror("Error", msg, parent=self.root)
                self.result_label.config(text=msg)
                return

            self.show_parameter_dialog()

            efficiency_values, dmu_names, improvement_suggestions = self.undesirable_sbm_analysis(file_path)

            # 构建结果数据
            data = []
            for dmu, eff in zip(dmu_names, efficiency_values):
                data.append([f"{dmu} {languages[self.current_language]['efficiency_value']}", eff, ""])

            headers = [
                languages[self.current_language]['statistics'],
                languages[self.current_language]['statistic_value'],
                languages[self.current_language]['p_value']
            ]
            df_result = pd.DataFrame(data, columns=headers)

            explanations = languages[self.current_language]['explanation']
            interpretations = languages[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(columns=["efficiency", "efficiency_bar", "improvement"])
            explanation_df.insert(0, "statistic_explanation", languages[self.current_language]['explanation_label'])

            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(columns=["efficiency", "efficiency_bar", "improvement"])
            interpretation_df.insert(0, "statistic_interpretation",
                                     languages[self.current_language]['interpretation_label'])

            combined_df = pd.concat([df_result, explanation_df, interpretation_df], ignore_index=True)

            # 保存文件对话框相对于主窗口居中
            save_path = filedialog.asksaveasfilename(
                parent=self.root,
                defaultextension=".docx",
                filetypes=[("Word files", "*.docx")]
            )
            if save_path:
                doc = Document()
                doc.add_heading(languages[self.current_language]["title"], level=1)

                doc.add_heading(languages[self.current_language]["model_parameters"], level=2)
                param_table = doc.add_table(rows=2, cols=2)
                param_table.cell(0, 0).text = languages[self.current_language]["orientation_label"]
                param_table.cell(0, 1).text = languages[self.current_language][
                    "input_oriented"] if self.orientation == "Input_Oriented" else languages[self.current_language][
                    "output_oriented"]
                param_table.cell(1, 0).text = languages[self.current_language]["returns_label"]
                param_table.cell(1, 1).text = languages[self.current_language][
                    "constant"] if self.returns_to_scale == "constant" else languages[self.current_language]["variable"]

                doc.add_heading(languages[self.current_language]["efficiency_results"], level=2)
                table = doc.add_table(rows=len(data) + 1, cols=2)
                table.cell(0, 0).text = languages[self.current_language]["decision_units"]
                table.cell(0, 1).text = languages[self.current_language]["efficiency_value"]
                for row_idx, (dmu, eff) in enumerate(zip(dmu_names, efficiency_values)):
                    table.cell(row_idx + 1, 0).text = dmu
                    table.cell(row_idx + 1, 1).text = str(eff)

                # 添加解释说明部分
                doc.add_heading(languages[self.current_language]["explanation_label"], level=2)
                for key, value in explanations.items():
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(f"{key}: ").bold = True
                    p.add_run(value)

                # 添加结果解读部分
                doc.add_heading(languages[self.current_language]["interpretation_label"], level=2)
                for key, value in interpretations.items():
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(f"{key}: ").bold = True
                    p.add_run(value)

                fig, ax = plt.subplots(figsize=(12, 8))
                ax.bar(dmu_names, efficiency_values)
                ax.set_title(languages[self.current_language]["efficiency_bar_chart"])
                ax.set_xlabel(languages[self.current_language]["decision_units"])
                ax.set_ylabel(languages[self.current_language]["efficiency_value"])
                ax.axhline(y=1.0, color='r', linestyle='--',
                           label=languages[self.current_language]["efficiency_frontier"])
                ax.legend()
                plt.xticks(rotation=0)
                plt.tight_layout()
                img_path = "efficiency_bar.png"
                plt.savefig(img_path, dpi=300)
                plt.close()
                doc.add_picture(img_path, width=Inches(6))

                if improvement_suggestions:
                    doc.add_heading(languages[self.current_language]["invalid_dmu"], level=2)
                    for dmu, suggestions in improvement_suggestions.items():
                        doc.add_heading(f"{dmu} ({languages[self.current_language]['improvement_suggestion']})",
                                        level=3)
                        p = doc.add_paragraph()
                        p.add_run(languages[self.current_language]["reduce_input"].format(
                            suggestions['input_reduction'])).bold = True
                        p = doc.add_paragraph()
                        p.add_run(languages[self.current_language]["adjust_output"].format(
                            suggestions['output_increase'])).bold = True
                        p = doc.add_paragraph()
                        p.add_run(languages[self.current_language]["adjust_undesirable"].format(
                            suggestions['undesirable_reduction'])).bold = True


                # 保存图片
                img_path = "efficiency_bar.png"
                plt.tight_layout()
                plt.savefig(img_path)
                plt.close()

                # 插入Word
                doc.add_picture(img_path, width=Inches(6))

                doc.save(save_path)
                self.result_label.config(text=languages[self.current_language]["analysis_complete"].format(save_path),
                                         wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]["no_save_path_selected"])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]["analysis_error"].format(str(e)))

    def switch_language(self, event):
        self.current_language = "en" if self.current_language == "zh" else "zh"
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])
        self.open_excel_label.config(text=languages[self.current_language]["open_excel_button_text"])


    def create_ui(self):
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=languages[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        self.switch_language_label = ttk.Label(frame,
                                               text=languages[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        self.root.mainloop()


def run_app():
    app = UndesirableSBMModelAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()
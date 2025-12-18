import tkinter as tk
from tkinter import filedialog, ttk, messagebox
import os
import numpy as np
import pandas as pd
pd.set_option('future.no_silent_downcasting', True)
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
from docx import Document
from sklearn.cluster import KMeans, DBSCAN
from sklearn.preprocessing import StandardScaler, MinMaxScaler
from sklearn.metrics import silhouette_score, calinski_harabasz_score, davies_bouldin_score
from sklearn.decomposition import PCA
from scipy.cluster.hierarchy import dendrogram, linkage
import tempfile
import openpyxl

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 语言字典（保留原样）
languages = {
    "zh": {
        "title": "二阶聚类",
        "select_button_text": "选择文件",
        "file_entry_placeholder": "请输入待分析 Excel 文件的完整路径",
        "analyze_button_text": "分析文件",
        "no_file_selected": "请选择文件。",
        "file_not_exists": "文件不存在，请重新选择。",
        "invalid_file": "无效的 Excel 文件，请选择正确的文件。",
        "file_corrupted": "Excel 文件损坏或无法读取",
        "permission_error": "没有权限访问文件或保存结果",
        "non_numeric_data": "数据中包含非数值类型，请清理数据后重试",
        "empty_data": "数据为空或格式不正确",
        "analysis_error": "分析文件时出错: {}",
        "analysis_complete": "分析完成，结果已保存到 {}\n",
        "no_save_path_selected": "未选择保存路径，结果未保存。",
        "switch_language_button_text": "中/英",
        'open_excel_button_text': "示例数据",
        "processing_data": "正在处理数据...",
        "clustering": "正在进行二阶聚类分析...",
        "generating_results": "正在生成结果...",
        "linkage_method": "连接方式:",
        "preprocessing": "数据预处理:",
        "pca_option": "PCA降维:",
        "none": "无",
        "standardization": "标准化 (Z-score)",
        "normalization": "归一化 (0-1)",
        "ward": "ward (最小方差法)",
        "single": "single (单连接)",
        "complete": "complete (全连接)",
        "average": "average (平均连接)",
        "explanation": {
            "cluster_result": "每个样本所属的聚类类别",
            "cluster_count": "聚类的总数量",
            "silhouette": "评估聚类质量的指标，范围[-1,1]，越接近1越好",
            "calinski_harabasz": "值越大表示聚类效果越好",
            "davies_bouldin": "值越小表示聚类效果越好",
            "dendrogram": "展示预聚类中心的二阶聚类关系"
        },
        "interpretation": {
            "cluster_result": "可用于区分不同样本所属的类别",
            "cluster_count": "根据轮廓系数自动确定的最佳聚类数量",
            "silhouette": "值越高表示聚类效果越好，样本与其自身聚类的相似度高于与其他聚类的相似度",
            "calinski_harabasz": "值越大表明簇内越紧凑，簇间越分散",
            "davies_bouldin": "值越小表明聚类效果越好，簇内样本相似度高且簇间差异大",
            "dendrogram": "直观展示预聚类中心的二阶聚类结构，帮助理解聚类形成过程"
        },
        "stats": {
            "cluster_result": "聚类结果",
            "cluster_count": "聚类数量",
            "precluster_count": "预聚类数量",
            "silhouette": "轮廓系数",
            "calinski_harabasz": "Calinski-Harabasz指数",
            "davies_bouldin": "Davies-Bouldin指数",
            "dendrogram": "聚类树状图"
        },
        "table_headers": {
            "statistic": "统计量",
            "value": "统计量值",
            "p_value": "p值",
            "explanation_header": "统计量_解释说明",
            "interpretation_header": "统计量_结果解读",
            "explanation_text": "解释说明",
            "interpretation_text": "结果解读"
        },
        "interpretations": {
            "silhouette_excellent": "，当前聚类质量优秀，样本划分清晰",
            "silhouette_good": "，当前聚类质量良好，样本划分较合理",
            "silhouette_fair": "，当前聚类质量一般，部分样本可能划分不够合理",
            "silhouette_poor": "，当前聚类质量较差，建议尝试其他参数或预处理方法",
            "silhouette_unavailable": "，指标不可用（簇数或簇样本数不足）",
            "precluster_explanation": "根据数据分布自动确定的初始分组数，影响最终聚类精度"
        },
        "cluster_stats": "聚类统计信息",
        "cluster_id": "聚类ID",
        "sample_count": "样本数量",
        "feature_mean": "特征均值",
        "feature_std": "特征标准差",
        "feature_median": "特征中位数",
        "missing_values": "数据中包含空值，已进行智能填充",
        "optimal_clusters": "自动推荐的最佳聚类数为: {}",
        "pca_explained": "PCA解释方差比例: {}"
    },
    "en": {
        "title": "Two Step Clustering",
        "select_button_text": "Select File",
        "file_entry_placeholder": "Please enter the full path of the Excel file to be analyzed",
        "analyze_button_text": "Analyze File",
        "no_file_selected": "Please select a valid file path.",
        "file_not_exists": "The file does not exist. Please select again.",
        "invalid_file": "Invalid Excel file. Please select a correct file.",
        "file_corrupted": "Excel file is corrupted or unreadable",
        "permission_error": "No permission to access file or save results",
        "non_numeric_data": "Data contains non-numeric values, please clean data and try again",
        "empty_data": "Data is empty or in incorrect format",
        "analysis_error": "An error occurred while analyzing the file: {}",
        "analysis_complete": "Analysis completed. The results have been saved to {}\n",
        "no_save_path_selected": "No save path selected. The results were not saved.",
        "switch_language_button_text": "Chinese/English",
        'open_excel_button_text': "Example data",
        "processing_data": "Processing data...",
        "clustering": "Performing two-step clustering analysis...",
        "generating_results": "Generating results...",
        "linkage_method": "Linkage method:",
        "preprocessing": "Data preprocessing:",
        "pca_option": "PCA dimensionality reduction:",
        "none": "None",
        "standardization": "Standardization (Z-score)",
        "normalization": "Normalization (0-1)",
        "ward": "ward (minimum variance)",
        "single": "single (nearest neighbor)",
        "complete": "complete (farthest neighbor)",
        "average": "average (average distance)",
        "explanation": {
            "cluster_result": "The cluster label to which each sample belongs",
            "cluster_count": "Total number of clusters",
            "silhouette": "Metric to evaluate clustering quality, ranging from [-1,1], closer to 1 is better",
            "calinski_harabasz": "Higher values indicate better clustering quality",
            "davies_bouldin": "Lower values indicate better clustering quality",
            "dendrogram": "Show the two-step clustering relationship of pre-cluster centers"
        },
        "interpretation": {
            "cluster_result": "Can be used to distinguish the categories to which different samples belong",
            "cluster_count": "Optimal number of clusters determined by silhouette score",
            "silhouette": "Higher values indicate better clustering, where samples are more similar to their own cluster than to others",
            "calinski_harabasz": "Higher values indicate tighter clusters and better separation between clusters",
            "davies_bouldin": "Lower values indicate better clustering with high intra-cluster similarity and distinct inter-cluster differences",
            "dendrogram": "Visually show the two-step structure of pre-cluster centers, helping understand the clustering process"
        },
        "stats": {
            "cluster_result": "Cluster Result",
            "cluster_count": "Number of Clusters",
            "precluster_count": "Number of Pre-clusters",
            "silhouette": "Silhouette Score",
            "calinski_harabasz": "Calinski-Harabasz Index",
            "davies_bouldin": "Davies-Bouldin Index",
            "dendrogram": "Cluster Dendrogram"
        },
        "table_headers": {
            "statistic": "Statistic",
            "value": "Value",
            "p_value": "p-value",
            "explanation_header": "Statistic_Explanation",
            "interpretation_header": "Statistic_Interpretation",
            "explanation_text": "Explanation",
            "interpretation_text": "Interpretation"
        },
        "interpretations": {
            "silhouette_excellent": ", current clustering quality is excellent with clear sample division",
            "silhouette_good": ", current clustering quality is good with reasonable sample division",
            "silhouette_fair": ", current clustering quality is average, some samples may be unreasonably divided",
            "silhouette_poor": ", current clustering quality is poor, suggest trying other parameters",
            "silhouette_unavailable": ", indicator unavailable (insufficient number of clusters or samples)",
            "precluster_explanation": "Automatically determined initial groups based on data distribution"
        },
        "cluster_stats": "Cluster Statistics",
        "cluster_id": "Cluster ID",
        "sample_count": "Sample Count",
        "feature_mean": "Feature Mean",
        "feature_std": "Feature Std",
        "feature_median": "Feature Median",
        "missing_values": "Data contains missing values, which have been intelligently imputed",
        "optimal_clusters": "Automatically recommended optimal number of clusters: {}",
        "pca_explained": "PCA explained variance ratio: {}"
    }
}


class SecondOrderClusteringAnalysisApp:
    def __init__(self, root=None):
        self.current_language = "en"
        self.file_path = ""

        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data37.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()
        self.file_path = file_path

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        current_text = self.file_entry.get()
        if current_text == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')
        else:
            self.file_path = current_text

    def validate_excel_file(self, file_path):
        try:
            if not (file_path.endswith('.xlsx') or file_path.endswith('.xls')):
                return False
            # 简单尝试用 pandas 读取一小块以验证
            pd.read_excel(file_path, nrows=1)
            return True
        except Exception:
            return False

    def handle_merged_cells(self, df):
        # 如果列名中有NaN，使用前一个非空列名附加_sub
        if df.columns.isnull().any():
            new_headers = []
            prev_header = None
            for header in df.columns:
                if pd.notnull(header):
                    prev_header = header
                    new_headers.append(header)
                else:
                    # 避免 None
                    suffix = "_sub" if prev_header is not None else "col"
                    new_headers.append(f"{prev_header}{suffix}" if prev_header is not None else suffix)
            df.columns = new_headers
        return df

    def smart_fill_missing_values(self, df):
        missing_cols = df.columns[df.isnull().any()].tolist()
        if not missing_cols:
            return df, False

        has_missing = False
        for col in missing_cols:
            missing_ratio = df[col].isnull().mean()
            if missing_ratio > 0.5:
                messagebox.showwarning("警告", f"列 {col} 缺失值超过50%，可能影响分析结果")
                has_missing = True

            # 对分类少值整数列用众数填充，否则用均值填充
            if pd.api.types.is_integer_dtype(df[col]) and df[col].nunique() < 10:
                mode_val = df[col].mode()
                if not mode_val.empty:
                    df[col] = df[col].fillna(mode_val[0]).infer_objects(copy=False)
                else:
                    df[col] = df[col].fillna(0).infer_objects(copy=False)
            else:
                # 如果列不可转换为数值的均值会失败，先尝试转为数值
                try:
                    df[col] = pd.to_numeric(df[col], errors='coerce')
                    mean_val = df[col].mean()
                    if pd.isna(mean_val):
                        df[col] = df[col].fillna(0)
                    else:
                        df[col] = df[col].fillna(mean_val)
                except Exception:
                    df[col] = df[col].fillna(method='ffill').fillna(method='bfill').fillna(0)

        return df, has_missing

    def preprocess_data(self, data, method):
        if method == "standardization":
            scaler = StandardScaler()
            return scaler.fit_transform(data)
        elif method == "normalization":
            scaler = MinMaxScaler()
            return scaler.fit_transform(data)
        return data

    def apply_pca(self, data, n_components=0.95):
        if data.shape[1] <= 1:
            return data, 1.0

        pca = PCA(n_components=n_components)
        transformed_data = pca.fit_transform(data)
        explained_variance = sum(pca.explained_variance_ratio_)
        return transformed_data, explained_variance

    def find_optimal_preclusters(self, data, max_clusters=20):
        """
        使用轮廓系数（silhouette）寻找最佳 K（若失败再回退到肘部法）
        """
        n_samples = len(data)
        max_possible = min(max_clusters, max(2, n_samples // 3))
        if max_possible < 2:
            return 2

        best_k = 2
        best_score = -1
        # silhouette requires at least 2 clusters and each cluster should have >1 sample ideally;
        # 但是我们只用来作参考
        for k in range(2, max_possible + 1):
            try:
                labels = KMeans(n_clusters=k, random_state=42).fit_predict(data)
                # 需要至少2个簇
                if len(set(labels)) < 2:
                    continue
                score = silhouette_score(data, labels)
                if score > best_score:
                    best_score = score
                    best_k = k
            except Exception:
                continue

        # 如果 silhouette 结果不可靠（best_score仍为-1），回退到肘部法（inertia 二阶导数）
        if best_score == -1:
            distortions = []
            for k in range(2, max_possible + 1):
                try:
                    kmeans = KMeans(n_clusters=k, random_state=42)
                    kmeans.fit(data)
                    distortions.append(kmeans.inertia_)
                except Exception:
                    distortions.append(np.nan)

            distortions = np.array(distortions)
            # 过滤 nan
            valid_idx = ~np.isnan(distortions)
            if valid_idx.sum() >= 2:
                d = distortions[valid_idx]
                if len(d) >= 3:
                    second_derivatives = np.diff(np.diff(d))
                    if len(second_derivatives) > 0:
                        best_idx = np.argmax(second_derivatives)
                        # best_idx 对应的 k = 2 + offset_of_valid_idx + best_idx
                        ks = np.arange(2, max_possible + 1)[valid_idx]
                        if best_idx < len(ks):
                            best_k = ks[best_idx] if len(ks) > 0 else 2
                else:
                    # 少数情况退到最小 k
                    best_k = 2
            else:
                best_k = 2

        return int(best_k)

    def two_step_clustering(self, data):
        """
        二阶聚类实现（KMeans 预聚类 + DBSCAN 对预聚类中心聚合）
        包含多处鲁棒性与边界条件处理
        返回：original_labels, cluster_centers, n_clusters, pre_clusters
        """
        # 防护：若数据为空
        if data is None or len(data) == 0:
            return np.array([]), np.array([]), 0, 0

        # 第一步：预聚类（K-Means）
        pre_clusters = self.find_optimal_preclusters(data)
        pre_clusters = max(2, min(pre_clusters, len(data)))  # 保证合理范围
        pre_model = KMeans(n_clusters=pre_clusters, random_state=42)
        pre_labels = pre_model.fit_predict(data)
        cluster_centers = pre_model.cluster_centers_

        n_centers = len(cluster_centers)
        if n_centers < 2:
            # 无法做二次聚类，直接把所有样本归为单一类
            original_labels = np.zeros(len(data), dtype=int)
            return original_labels, cluster_centers, 1, pre_clusters

        # 第二步：对预聚类中心做 DBSCAN
        try:
            from sklearn.neighbors import NearestNeighbors
            # 选择 k 为 min(5, n_centers-1)，确保不会越界
            k_nn = min(5, max(1, n_centers - 1))
            neighbors = NearestNeighbors(n_neighbors=k_nn)
            neighbors_fit = neighbors.fit(cluster_centers)
            distances_all, indices = neighbors_fit.kneighbors(cluster_centers)
            # 使用 k_nn-1 列（0-based index）即第 k_nn 个最近距离
            # 排序并用高分位数（90%）作为 eps，较稳健
            distances_k = np.sort(distances_all[:, k_nn - 1], axis=0)
            # 如果中心少，percentile 仍可使用
            eps = float(np.percentile(distances_k, 90)) if len(distances_k) > 0 else 0.5
            if eps <= 0:
                eps = float(np.median(distances_k)) if len(distances_k) > 0 else 0.5
        except Exception:
            # 回退默认值
            eps = 0.5

        # min_samples 至少1或2，按中心数取比例
        min_samples = max(2, int(max(1, n_centers * 0.1)))
        # DBSCAN fit
        try:
            dbscan = DBSCAN(eps=eps, min_samples=min_samples)
            final_labels = dbscan.fit_predict(cluster_centers)
        except Exception:
            # 如果 DBSCAN 失败（例如 eps 太小或其他异常），把每个预聚类中心作为独立聚类
            final_labels = np.arange(n_centers, dtype=int)

        # 处理噪声点标签为 -1
        noise_mask = final_labels == -1
        if np.any(noise_mask):
            # 对每个噪声中心逐个计算距离并分配到最近的非噪声簇
            non_noise_idx = np.where(~noise_mask)[0]
            for i in np.where(noise_mask)[0]:
                if len(non_noise_idx) == 0:
                    # 没有非噪声簇，给一个新簇号（0）
                    final_labels[i] = 0
                else:
                    # 计算当前噪声中心到非噪声中心的距离（按索引）
                    distances = np.linalg.norm(cluster_centers[i] - cluster_centers[non_noise_idx], axis=1)
                    nearest_idx = non_noise_idx[np.argmin(distances)]
                    final_labels[i] = final_labels[nearest_idx]

        # 确保存在至少一个聚类标签
        unique_final_labels = np.unique(final_labels)
        if unique_final_labels.size == 0:
            final_labels = np.zeros(n_centers, dtype=int)

        # 建立从预聚类中心索引到最终簇标签的映射（避免映射失真）
        cluster_label_map = dict(zip(range(n_centers), final_labels.tolist()))
        # 将原始样本映射到预聚类中心，再通过映射表得到最终标签
        original_labels = np.array([cluster_label_map[int(pre_labels[i])] for i in range(len(pre_labels))])

        n_clusters = len(np.unique(original_labels))
        return original_labels, cluster_centers, n_clusters, pre_clusters

    def plot_cluster_relationship(self, centers, labels):
        """绘制预聚类中心与最终聚类的关系图"""
        plt.figure(figsize=(10, 6))
        unique_labels = np.unique(labels)
        colors = plt.cm.rainbow(np.linspace(0, 1, len(unique_labels)))

        # 如果 centers 只有 1 维，绘制直方图或散点
        if centers.shape[1] == 1:
            for label, color in zip(unique_labels, colors):
                mask = labels == label
                plt.scatter(np.arange(np.sum(mask)), centers[mask, 0], c=[color], label=f'Cluster {label}')
            plt.xlabel('Index within label')
            plt.ylabel('Feature 1')
        else:
            for label, color in zip(unique_labels, colors):
                mask = labels == label
                plt.scatter(centers[mask, 0], centers[mask, 1], c=[color], label=f'Cluster {label}')

        plt.title('预聚类中心的二阶聚类分布' if self.current_language == 'zh' else
                  'Distribution of Pre-cluster Centers in Two-Step Clustering')
        plt.xlabel('特征1' if self.current_language == 'zh' else 'Feature 1')
        plt.ylabel('特征2' if self.current_language == 'zh' else 'Feature 2')
        plt.legend()
        plt.grid(True, alpha=0.3)

    def update_status(self, message):
        self.result_label.config(text=message)
        self.root.update_idletasks()

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"] or not file_path:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return

        # 参数设置（如需可改为 UI 控件）
        preprocessing_method = "none"
        use_pca = False

        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return

        if not self.validate_excel_file(file_path):
            self.result_label.config(text=languages[self.current_language]["invalid_file"])
            return

        try:
            self.update_status(languages[self.current_language]["processing_data"])

            # 使用 pandas 读取 Excel，效率更高
            df = pd.read_excel(file_path, header=0)

            # 处理合并表头（如果存在）
            df = self.handle_merged_cells(df)

            # 填充缺失值（智能填充）
            df, has_missing = self.smart_fill_missing_values(df)
            if has_missing:
                messagebox.showinfo("提示", languages[self.current_language]["missing_values"])

            # 检查并保留数值列
            numeric_df = df.select_dtypes(include=[np.number])
            non_numeric_cols = set(df.columns) - set(numeric_df.columns)
            if non_numeric_cols:
                self.result_label.config(
                    text=f"{languages[self.current_language]['non_numeric_data']}: {', '.join(non_numeric_cols)}"
                )
                return

            if numeric_df.empty:
                self.result_label.config(text=languages[self.current_language]["empty_data"])
                return

            data = numeric_df.values
            self.update_status(languages[self.current_language]["clustering"])

            # 数据预处理
            processed_data = self.preprocess_data(data, preprocessing_method)

            # PCA 可选
            pca_explained = 0.0
            if use_pca:
                processed_data, pca_explained = self.apply_pca(processed_data)
                self.update_status(
                    f"{languages[self.current_language]['pca_explained']}".format(f"{pca_explained:.2%}"))

            # 执行二阶聚类
            labels, cluster_centers, n_clusters, pre_clusters = self.two_step_clustering(processed_data)

            # 保护性评估指标计算：需要至少 2 个簇且每簇至少2样本（某些指标要求）
            silhouette_avg = calinski_harabasz = davies_bouldin = np.nan
            try:
                unique_labels, counts = np.unique(labels, return_counts=True)
                if len(unique_labels) >= 2 and np.all(counts >= 2):
                    silhouette_avg = silhouette_score(processed_data, labels)
                    calinski_harabasz = calinski_harabasz_score(processed_data, labels)
                    davies_bouldin = davies_bouldin_score(processed_data, labels)
                else:
                    # 无法可靠计算指标，保持为 NaN
                    silhouette_avg = np.nan
                    calinski_harabasz = np.nan
                    davies_bouldin = np.nan
            except Exception:
                silhouette_avg = np.nan
                calinski_harabasz = np.nan
                davies_bouldin = np.nan

            self.update_status(languages[self.current_language]["generating_results"])
            self.update_status(languages[self.current_language]["optimal_clusters"].format(n_clusters))

            # 整理聚类统计信息（按最终标签分组）
            cluster_stats = []
            if len(labels) > 0:
                for i in range(n_clusters):
                    cluster_samples = data[labels == i]
                    sample_count = len(cluster_samples)
                    if sample_count > 0:
                        feature_means = cluster_samples.mean(axis=0).tolist()
                        feature_stds = cluster_samples.std(axis=0).tolist()
                        feature_medians = np.median(cluster_samples, axis=0).tolist()
                    else:
                        feature_means = [np.nan] * data.shape[1]
                        feature_stds = [np.nan] * data.shape[1]
                        feature_medians = [np.nan] * data.shape[1]
                    cluster_stats.append([i, sample_count, feature_means, feature_stds, feature_medians])

            # 整理结果数据
            result_data = [
                [languages[self.current_language]["stats"]["cluster_result"],
                 labels.tolist() if len(labels) > 0 else []],
                [languages[self.current_language]["stats"]["cluster_count"], n_clusters],
                [languages[self.current_language]["stats"]["precluster_count"], pre_clusters],
                [languages[self.current_language]["stats"]["silhouette"],
                 f"{silhouette_avg:.4f}" if not pd.isna(silhouette_avg) else "N/A"],
                [languages[self.current_language]["stats"]["calinski_harabasz"],
                 f"{calinski_harabasz:.4f}" if not pd.isna(calinski_harabasz) else "N/A"],
                [languages[self.current_language]["stats"]["davies_bouldin"],
                 f"{davies_bouldin:.4f}" if not pd.isna(davies_bouldin) else "N/A"]
            ]

            headers = [
                languages[self.current_language]["table_headers"]["statistic"],
                languages[self.current_language]["table_headers"]["value"]
            ]
            result_df = pd.DataFrame(result_data, columns=headers)

            # 添加解释说明
            explanations = {
                languages[self.current_language]["stats"]["cluster_result"]:
                    languages[self.current_language]["explanation"]["cluster_result"],
                languages[self.current_language]["stats"]["cluster_count"]:
                    languages[self.current_language]["explanation"]["cluster_count"],
                languages[self.current_language]["stats"]["precluster_count"]:
                    languages[self.current_language]["interpretations"]["precluster_explanation"],
                languages[self.current_language]["stats"]["silhouette"]:
                    languages[self.current_language]["explanation"]["silhouette"],
                languages[self.current_language]["stats"]["calinski_harabasz"]:
                    languages[self.current_language]["explanation"]["calinski_harabasz"],
                languages[self.current_language]["stats"]["davies_bouldin"]:
                    languages[self.current_language]["explanation"]["davies_bouldin"],
                languages[self.current_language]["stats"]["dendrogram"]:
                    languages[self.current_language]["explanation"]["dendrogram"]
            }

            # 同时删除后续的 display_explanations 转换，直接使用：
            explanation_df = pd.DataFrame([explanations])
            # 确保列顺序
            cols = [
                languages[self.current_language]["stats"]["cluster_result"],
                languages[self.current_language]["stats"]["cluster_count"],
                languages[self.current_language]["stats"]["precluster_count"],
                languages[self.current_language]["stats"]["silhouette"],
                languages[self.current_language]["stats"]["calinski_harabasz"],
                languages[self.current_language]["stats"]["davies_bouldin"],
                languages[self.current_language]["stats"]["dendrogram"]
            ]
            for c in cols:
                if c not in explanation_df.columns:
                    explanation_df[c] = ""
            explanation_df = explanation_df.reindex(columns=cols)
            explanation_df.insert(0,
                                  languages[self.current_language]["table_headers"]["explanation_header"],
                                  languages[self.current_language]["table_headers"]["explanation_text"]
                                  )

            # 添加结果解读
            interpretations = {
                languages[self.current_language]["stats"]["cluster_result"]:
                    languages[self.current_language]["interpretation"]["cluster_result"],
                languages[self.current_language]["stats"]["cluster_count"]:
                    languages[self.current_language]["interpretation"]["cluster_count"],
                languages[self.current_language]["stats"]["precluster_count"]:
                    languages[self.current_language]["interpretations"]["precluster_explanation"],
                languages[self.current_language]["stats"]["silhouette"]:
                    languages[self.current_language]["interpretation"]["silhouette"],
                languages[self.current_language]["stats"]["calinski_harabasz"]:
                    languages[self.current_language]["interpretation"]["calinski_harabasz"],
                languages[self.current_language]["stats"]["davies_bouldin"]:
                    languages[self.current_language]["interpretation"]["davies_bouldin"],
                languages[self.current_language]["stats"]["dendrogram"]:
                    languages[self.current_language]["interpretation"]["dendrogram"]
            }

            interpretation_df = pd.DataFrame([interpretations])
            for c in cols:
                if c not in interpretation_df.columns:
                    interpretation_df[c] = ""
            interpretation_df = interpretation_df.reindex(columns=cols)
            interpretation_df.insert(0,
                                     languages[self.current_language]["table_headers"]["interpretation_header"],
                                     languages[self.current_language]["table_headers"]["interpretation_text"]
                                     )
            # 保存结果到 Word
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_img:
                    img_path = temp_img.name

                # 生成聚类关系图（若有中心坐标）
                try:
                    if cluster_centers is not None and cluster_centers.shape[1] >= 1:
                        if cluster_centers.shape[1] >= 2:
                            center_labels = None
                            try:
                                if cluster_centers.shape[1] > 2:
                                    from sklearn.decomposition import PCA as _PCA
                                    pca = _PCA(n_components=2)
                                    centers_2d = pca.fit_transform(cluster_centers)
                                else:
                                    centers_2d = cluster_centers[:, :2]
                                plt.figure(figsize=(10, 6))
                                plt.scatter(centers_2d[:, 0], centers_2d[:, 1])
                                plt.title('预聚类中心分布' if self.current_language == 'zh' else 'Pre-cluster center distribution')
                            except Exception:
                                plt.figure(figsize=(10, 6))
                                plt.scatter(cluster_centers[:, 0], np.zeros(len(cluster_centers)))
                                plt.title('预聚类中心分布（1D）' if self.current_language == 'zh' else 'Pre-cluster center distribution(1D)')
                        else:
                            plt.figure(figsize=(10, 6))
                            plt.hist(labels if len(labels) > 0 else [0], bins=max(1, n_clusters))
                            plt.title('聚类分布直方图' if self.current_language == 'zh' else 'Cluster Distribution Histogram')
                    else:
                        plt.figure(figsize=(10, 6))
                        plt.hist(labels if len(labels) > 0 else [0], bins=max(1, n_clusters))
                        plt.title('聚类分布直方图' if self.current_language == 'zh' else 'Cluster Distribution Histogram')
                    plt.savefig(img_path, bbox_inches='tight')
                    plt.close()
                except Exception:
                    # 回退：生成空白图
                    plt.figure(figsize=(6, 4))
                    plt.text(0.5, 0.5, 'Visualization failed', horizontalalignment='center', verticalalignment='center')
                    plt.axis('off')
                    plt.savefig(img_path)
                    plt.close()

                # 创建 Word 文档并插入表格与图像
                doc = Document()
                doc.add_heading('分析结果' if self.current_language == 'zh' else 'Analysis Results', level=1)
                table = doc.add_table(rows=result_df.shape[0] + 1, cols=result_df.shape[1])
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(result_df.columns):
                    hdr_cells[col_idx].text = header
                for row_idx, row in result_df.iterrows():
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx].text = str(value)

                # 聚类统计信息
                doc.add_heading(languages[self.current_language]["cluster_stats"], level=1)
                stats_table = doc.add_table(rows=len(cluster_stats) + 1, cols=5)
                stats_hdr = stats_table.rows[0].cells
                stats_hdr[0].text = languages[self.current_language]["cluster_id"]
                stats_hdr[1].text = languages[self.current_language]["sample_count"]
                stats_hdr[2].text = languages[self.current_language]["feature_mean"]
                stats_hdr[3].text = languages[self.current_language]["feature_std"]
                stats_hdr[4].text = languages[self.current_language]["feature_median"]

                for row_idx, stats in enumerate(cluster_stats):
                    row_cells = stats_table.rows[row_idx + 1].cells
                    row_cells[0].text = str(stats[0])
                    row_cells[1].text = str(stats[1])
                    row_cells[2].text = ", ".join([f"{x:.4f}" for x in stats[2]])
                    row_cells[3].text = ", ".join([f"{x:.4f}" for x in stats[3]])
                    row_cells[4].text = ", ".join([f"{x:.4f}" for x in stats[4]])

                # 替换原解释说明表格部分代码
                doc.add_heading('解释说明' if self.current_language == 'zh' else 'Explanation', level=1)
                for item in explanations.items():
                    doc.add_paragraph(f"• {item[0]}: {item[1]}", style='ListBullet')

                # 替换原结果解读表格部分代码
                doc.add_heading('结果解读' if self.current_language == 'zh' else 'Interpretation', level=1)
                for item in interpretations.items():
                    doc.add_paragraph(f"• {item[0]}: {item[1]}", style='ListBullet')

                # 插入图像
                doc.add_heading('二阶聚类分布' if self.current_language == 'zh' else
                                'Two-Step Clustering Distribution', level=1)
                try:
                    doc.add_picture(img_path)
                except Exception:
                    pass

                doc.save(save_path)

                # 清理临时文件
                try:
                    os.unlink(img_path)
                except Exception:
                    pass

                result_msg = languages[self.current_language]['analysis_complete'].format(save_path)
                self.result_label.config(text=result_msg, wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]["no_save_path_selected"])

        except PermissionError:
            self.result_label.config(text=languages[self.current_language]["permission_error"])
        except openpyxl.utils.exceptions.InvalidFileException:
            self.result_label.config(text=languages[self.current_language]["file_corrupted"])
        except Exception as e:
            self.result_label.config(text=languages[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event):
        current_text = self.file_entry.get()
        if current_text != languages[self.current_language]["file_entry_placeholder"]:
            self.file_path = current_text

        self.current_language = "en" if self.current_language == "zh" else "zh"

        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)

        if self.file_path and os.path.exists(self.file_path):
            self.file_entry.insert(0, self.file_path)
            self.file_entry.config(foreground='black')
        else:
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])
        self.open_excel_label.config(text=languages[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的40%，与第一个程序一致）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小（与第一个程序一致）
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口居中位置（与第一个程序一致）
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建主框架（无额外内边距，与第一个程序一致）
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮（统一间距为pady=10，移除顶部90像素的额外间距）
        self.select_button = ttk.Button(
            frame,
            text=languages[self.current_language]["select_button_text"],
            command=self.select_file,
            bootstyle=PRIMARY,
            cursor="hand2"  # 与第一个程序一致，添加手型光标
        )
        self.select_button.pack(pady=10)

        # 创建文件路径输入框（与第一个程序保持一致）
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮（统一间距为pady=10，添加手型光标）
        self.analyze_button = ttk.Button(
            frame,
            text=languages[self.current_language]["analyze_button_text"],
            command=self.analyze_file,
            bootstyle=SUCCESS,
            cursor="hand2"  # 与第一个程序一致
        )
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=languages[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签（与第一个程序保持一致的间距和样式）
        self.switch_language_label = ttk.Label(
            frame,
            text=languages[self.current_language]["switch_language_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签（挂载到root，不设置固定换行长度，与第一个程序一致）
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        self.root.mainloop()


def run_app():
    app = SecondOrderClusteringAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()

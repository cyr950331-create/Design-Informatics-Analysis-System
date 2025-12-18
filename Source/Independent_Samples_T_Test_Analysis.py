import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
from scipy import stats
import matplotlib.pyplot as plt
import pathlib
from docx import Document
from docx.shared import Inches
from docx.shared import Cm

# 设置支持中文的字体
plt.rcParams['font.family'] = 'SimHei'  # 使用黑体字体，可根据系统情况修改
plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

# 定义语言字典 - 所有显示文本均包含在内
LANGUAGES = {
    'zh': {
        'title': "独立样本 t 检验",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "请选择文件。",
        'analysis_success': "分析完成，结果已保存到 {}\n图表已保存到 {}",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "中/英",
        'open_excel_button_text': "示例数据",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        # 图表相关文本
        'chart': {
            'bar_title': '柱状图',
            'error_bar_title': '误差线图',
            'boxplot_title': '箱线图',
            'line_title': '折线图',
            'x_label_samples': '样本',
            'x_label_index': '索引',
            'y_label_mean': '均值',
            'y_label_value': '数值',
            'sample1': '样本 1',
            'sample2': '样本 2',
            'legend_sample1': '样本 1',
            'legend_sample2': '样本 2'
        },
        # Word文档相关文本
        'doc': {
            'main_heading': '独立样本 t 检验分析结果',
            'results_table_heading': '分析结果表格',
            'explanation_heading': '解释说明',
            'interpretation_heading': '结果解读',
            'charts_heading': '数据分析图表',
            'explanation_label': '解释说明',
            'interpretation_label': '结果解读'
        },
        # 统计量相关文本
        'stats': {
            'independent_t_test': '独立样本 t 检验',
            'sample_size': '样本量',
            'mean': '均值',
            'std_dev': '标准差',
            't_statistic': 't 统计量',
            'degrees_of_freedom': '自由度',
            'p_value': 'p 值',
            'confidence_interval': '置信区间',
            'mean_difference': '均值差异'
        },
        # 表格标题
        'table_headers': {
            'statistic': '统计量',
            't_stat': 't 统计量',
            'df': '自由度',
            'p_val': 'p 值',
            'conf_int': '置信区间'
        },
        # 解释说明
        'explanation': {
            "independent_t_test": "用于比较两个独立样本的均值是否有显著差异。",
            "sample_size": "每个样本中的观测值数量。",
            "mean": "样本数据的平均值。",
            "std_dev": "样本数据的离散程度。",
            "t_statistic": "用于衡量两个样本均值差异的程度。",
            "degrees_of_freedom": "在统计分析中能够自由取值的变量个数。",
            "p_value": "用于判断两个样本均值是否有显著差异的指标。",
            "confidence_interval": "均值差异的可能范围。"
        },
        # 结果解读
        'interpretation': {
            "t_statistic": "t 统计量的绝对值越大，说明两个样本均值的差异越显著。",
            "degrees_of_freedom": "自由度越大，t 分布越接近正态分布。",
            "p_value": "p 值小于显著性水平（通常为 0.05）时，拒绝原假设，认为两个样本均值存在显著差异；否则，接受原假设，认为两个样本均值无显著差异。",
            "confidence_interval": "如果置信区间不包含 0，说明两个样本均值存在显著差异。"
        }
    },
    'en': {
        'title': "Independent Samples T Test",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "Please select a file.",
        'analysis_success': "Analysis completed. Results saved to {}\nCharts saved to {}",
        'no_save_path': "No save path selected. Results not saved.",
        'analysis_error': "Error analyzing file: {}",
        'switch_language': "Chinese/English",
        'open_excel_button_text': "Example Data",
        'file_entry_placeholder': "Please enter the full path of the Excel file to analyze",
        # 图表相关文本
        'chart': {
            'bar_title': 'Bar Chart',
            'error_bar_title': 'Error Bar Chart',
            'boxplot_title': 'Box Plot',
            'line_title': 'Line Chart',
            'x_label_samples': 'Samples',
            'x_label_index': 'Index',
            'y_label_mean': 'Mean',
            'y_label_value': 'Value',
            'sample1': 'Sample 1',
            'sample2': 'Sample 2',
            'legend_sample1': 'Sample 1',
            'legend_sample2': 'Sample 2'
        },
        # Word文档相关文本
        'doc': {
            'main_heading': 'Independent Samples T-Test Analysis Results',
            'results_table_heading': 'Analysis Results Table',
            'explanation_heading': 'Explanations',
            'interpretation_heading': 'Interpretations',
            'charts_heading': 'Data Analysis Charts',
            'explanation_label': 'Explanation',
            'interpretation_label': 'Interpretation'
        },
        # 统计量相关文本
        'stats': {
            'independent_t_test': 'Independent Samples T-Test',
            'sample_size': 'Sample Size',
            'mean': 'Mean',
            'std_dev': 'Standard Deviation',
            't_statistic': 't Statistic',
            'degrees_of_freedom': 'Degrees of Freedom',
            'p_value': 'p Value',
            'confidence_interval': 'Confidence Interval',
            'mean_difference': 'Mean Difference'
        },
        # 表格标题
        'table_headers': {
            'statistic': 'Statistic',
            't_stat': 't Statistic',
            'df': 'Degrees of Freedom',
            'p_val': 'p Value',
            'conf_int': 'Confidence Interval'
        },
        # 解释说明
        'explanation': {
            "independent_t_test": "Used to compare whether the means of two independent samples are significantly different.",
            "sample_size": "The number of observations in each sample.",
            "mean": "The average value of the sample data.",
            "std_dev": "The degree of dispersion of the sample data.",
            "t_statistic": "Used to measure the degree of difference between the means of two samples.",
            "degrees_of_freedom": "The number of variables that can take on independent values in a statistical analysis.",
            "p_value": "An indicator used to determine whether the means of two samples are significantly different.",
            "confidence_interval": "The possible range of the difference in means."
        },
        # 结果解读
        'interpretation': {
            "t_statistic": "The larger the absolute value of the t statistic, the more significant the difference between the means of the two samples.",
            "degrees_of_freedom": "The larger the degrees of freedom, the closer the t distribution is to the normal distribution.",
            "p_value": "When the p-value is less than the significance level (usually 0.05), the null hypothesis is rejected, indicating a significant difference between the means of the two samples; otherwise, the null hypothesis is accepted, indicating no significant difference.",
            "confidence_interval": "If the confidence interval does not contain 0, it indicates a significant difference between the means of the two samples."
        }
    }
}


class IndependentSamplesTTestAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为中文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data41.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"{LANGUAGES[self.current_language]['file_not_found']}：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"{LANGUAGES[self.current_language]['analysis_error']}：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def plot_results(self, sample1, sample2, save_path=None):
        # 获取当前语言的图表文本
        chart_text = LANGUAGES[self.current_language]['chart']

        # 柱状图
        plt.figure(figsize=(12, 8))
        plt.subplot(2, 2, 1)
        means = [sample1.mean(), sample2.mean()]
        labels = [chart_text['sample1'], chart_text['sample2']]
        plt.bar(labels, means)
        plt.title(chart_text['bar_title'])
        plt.xlabel(chart_text['x_label_samples'])
        plt.ylabel(chart_text['y_label_mean'])

        # 误差线图
        plt.subplot(2, 2, 2)
        stds = [sample1.std(), sample2.std()]
        plt.errorbar(labels, means, yerr=stds, fmt='o')
        plt.title(chart_text['error_bar_title'])
        plt.xlabel(chart_text['x_label_samples'])
        plt.ylabel(chart_text['y_label_mean'])

        # 箱线图
        plt.subplot(2, 2, 3)
        plt.boxplot([sample1, sample2])
        plt.title(chart_text['boxplot_title'])
        plt.xlabel(chart_text['x_label_samples'])
        plt.ylabel(chart_text['y_label_value'])

        # 折线图
        plt.subplot(2, 2, 4)
        plt.plot(sample1, label=chart_text['legend_sample1'])
        plt.plot(sample2, label=chart_text['legend_sample2'])
        plt.title(chart_text['line_title'])
        plt.xlabel(chart_text['x_label_index'])
        plt.ylabel(chart_text['y_label_value'])
        plt.legend()

        plt.tight_layout()

        # 如果提供了保存路径，则保存图表
        chart_path = None
        if save_path:
            # 生成图表保存路径
            base_name = os.path.splitext(save_path)[0]
            chart_path = f"{base_name}_charts.png"
            plt.savefig(chart_path, dpi=300, bbox_inches='tight')

            plt.close()
        return chart_path

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 检查数据是否为数值类型
            numerical_df = df.select_dtypes(include=[np.number])
            if numerical_df.empty:
                error_msg = "数据中没有数值列，无法进行独立样本 t 检验。" if self.current_language == 'zh' else "No numerical columns in data, cannot perform independent samples t-test."
                raise ValueError(error_msg)
            if len(numerical_df.columns) != 2:
                error_msg = "数据必须包含两列数值数据，用于独立样本 t 检验。" if self.current_language == 'zh' else "Data must contain two columns of numerical data for independent samples t-test."
                raise ValueError(error_msg)

            # 进行独立样本 t 检验
            sample1 = numerical_df.iloc[:, 0]
            sample2 = numerical_df.iloc[:, 1]
            t_stat, p_value = stats.ttest_ind(sample1, sample2)
            df_value = len(sample1) + len(sample2) - 2
            mean_diff = sample1.mean() - sample2.mean()
            std_err = stats.sem(sample1 - sample2)
            # 计算置信区间并转换为普通浮点数
            conf_int = stats.t.interval(0.95, df_value, loc=mean_diff, scale=std_err)
            # 将numpy浮点数转换为普通Python浮点数
            conf_int = (float(conf_int[0]), float(conf_int[1]))

            # 计算样本量、均值和标准差
            sample_sizes = numerical_df.count()
            means = numerical_df.mean()
            stds = numerical_df.std()

            # 获取当前语言的统计文本
            stats_text = LANGUAGES[self.current_language]['stats']
            table_headers = LANGUAGES[self.current_language]['table_headers']

            # 整理数据
            data = [
                [stats_text['independent_t_test'], t_stat, df_value, p_value, conf_int],
                [stats_text['sample_size'], sample_sizes.to_dict(), "", "", ""],
                [stats_text['mean'], means.to_dict(), "", "", ""],
                [stats_text['std_dev'], stds.to_dict(), "", "", ""]
            ]
            headers = [
                table_headers['statistic'],
                table_headers['t_stat'],
                table_headers['df'],
                table_headers['p_val'],
                table_headers['conf_int']
            ]
            df_result = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            explanations = LANGUAGES[self.current_language]['explanation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(
                columns=[
                    "independent_t_test",
                    "sample_size",
                    "mean",
                    "std_dev",
                    "t_statistic",
                    "degrees_of_freedom",
                    "p_value",
                    "confidence_interval"
                ])
            explanation_df.insert(0, "统计量_解释说明",
                                  LANGUAGES[self.current_language]['doc']['explanation_label'])

            # 添加分析结果解读
            interpretations = LANGUAGES[self.current_language]['interpretation']
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(
                columns=["t_statistic", "degrees_of_freedom", "p_value", "confidence_interval"])
            interpretation_df.insert(0, "统计量_结果解读",
                                     LANGUAGES[self.current_language]['doc']['interpretation_label'])

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                                     filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建一个新的Word文档
                doc = Document()
                doc_text = LANGUAGES[self.current_language]['doc']
                doc.add_heading(doc_text['main_heading'], 0)

                # 添加表格
                doc.add_heading(doc_text['results_table_heading'], level=1)
                table = doc.add_table(rows=df_result.shape[0] + 1, cols=df_result.shape[1])
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(df_result.columns):
                    hdr_cells[col_idx].text = header

                for row_idx in range(df_result.shape[0]):
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx, value in enumerate(df_result.iloc[row_idx]):
                        row_cells[col_idx].text = str(value)

                # 添加解释说明（项目符号列表）
                doc.add_heading(doc_text['explanation_heading'], level=1)
                explanations = LANGUAGES[self.current_language]['explanation']
                for key, value in explanations.items():
                    # 将键转换为当前语言的显示文本
                    display_key = stats_text[key] if key in stats_text else key
                    doc.add_paragraph(value, style='List Bullet').add_run(f' ({display_key})')

                # 添加结果解读（项目符号列表）
                doc.add_heading(doc_text['interpretation_heading'], level=1)
                interpretations = LANGUAGES[self.current_language]['interpretation']
                for key, value in interpretations.items():
                    # 将键转换为当前语言的显示文本
                    display_key = stats_text[key] if key in stats_text else key
                    doc.add_paragraph(value, style='List Bullet').add_run(f' ({display_key})')

                # 添加图表
                doc.add_heading(doc_text['charts_heading'], level=1)
                chart_path = self.plot_results(sample1, sample2, save_path)
                if chart_path and os.path.exists(chart_path):
                    doc.add_picture(chart_path, width=Cm(15))

                # 保存文档
                doc.save(save_path)
                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(save_path, chart_path)
                self.result_label.config(text=result_msg, wraplength=400)
            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        # 切换当前语言
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'

        # 更新UI文本
        self.root.title(LANGUAGES[self.current_language]['title'])
        self.select_button.config(text=LANGUAGES[self.current_language]['select_button'])
        self.analyze_button.config(text=LANGUAGES[self.current_language]['analyze_button'])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]['switch_language'])
        self.open_excel_label.config(text=LANGUAGES[self.current_language]["open_excel_button_text"])

        # 更新输入框提示文本
        current_text = self.file_entry.get()
        placeholder = LANGUAGES[self.current_language]['file_entry_placeholder']
        if current_text == LANGUAGES['zh']['file_entry_placeholder'] or current_text == LANGUAGES['en'][
            'file_entry_placeholder']:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, placeholder)
            self.file_entry.config(foreground='gray')

        # 清空结果标签
        self.result_label.config(text="")

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建自定义样式
        style = ttk.Style()
        style.configure("Gray.TEntry", foreground="gray")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame,
                                        text=LANGUAGES[self.current_language]['select_button'],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame,
                                         text=LANGUAGES[self.current_language]['analyze_button'],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=LANGUAGES[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame,
                                               text=LANGUAGES[self.current_language]['switch_language'],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = IndependentSamplesTTestAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()
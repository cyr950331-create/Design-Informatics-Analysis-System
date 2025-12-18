import tkinter as tk
from tkinter import filedialog
import os
import numpy as np
import pandas as pd
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
from sklearn.manifold import MDS
from sklearn.metrics.pairwise import euclidean_distances
from docx import Document
from docx.shared import Inches

# 设置支持中文的字体
plt.rcParams['font.family'] = ['SimHei', 'WenQuanYi Micro Hei', 'Heiti TC']
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
languages = {
    'zh': {
        'title': "多维尺度分析",
        'select_button_text': "选择文件",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'analyze_button_text': "分析文件",
        'has_header_checkbox': "文件包含列名行",
        'index_column_label': "样本名所在列(可选):",
        'no_file_selected': "请选择文件。",
        'file_not_exists': "文件不存在，请重新选择。",
        'analysis_error': "分析文件时出错: {}",
        'analysis_complete': "分析完成，结果已保存到 {}，相关图片已保存。",
        'no_save_path_selected': "未选择保存路径，结果未保存。",
        'switch_language_button_text': "中/英",
        'open_excel_button_text': "示例数据",
        'missing_values_title': "缺失值处理",
        'missing_values_message': "数据中存在缺失值，是否使用均值填充？\n取消将删除含缺失值的行",
        'invalid_index_column': "无效的样本名列索引",
        'no_numeric_data': "数据中没有可分析的数值型列",
        'insufficient_samples': "样本数量不足，至少需要3个样本",
        'explanation': {
            "MDS坐标": "经过多维尺度分析后得到的低维坐标",
            "MDS散点图": "展示样本在低维空间中的分布情况"
        },
        'interpretation': {
            "MDS坐标": "可用于观察样本在低维空间中的相对位置关系",
            "MDS散点图": "直观展示样本之间的相似性，距离近的样本更相似"
        }
    },
    'en': {
        'title': "Multidimensional Scaling Analysis",
        'select_button_text': "Select File",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'analyze_button_text': "Analyze File",
        'has_header_checkbox': "File contains header row",
        'index_column_label': "Sample name column (optional):",
        'no_file_selected': "Please select a valid file path.",
        'file_not_exists': "The file does not exist. Please select again.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'analysis_complete': "Analysis completed. The results have been saved to {}, and the relevant images have been saved.",
        'no_save_path_selected': "No save path selected. The results were not saved.",
        'switch_language_button_text': "Chinese/English",
        'open_excel_button_text': "Example data",
        'missing_values_title': "Missing Value Handling",
        'missing_values_message': "There are missing values in the data. Do you want to use the mean to fill them？\n Cancel to delete rows with missing values",
        'invalid_index_column': "Invalid sample column index",
        'no_numeric_data': "There are no numeric columns that can be analyzed in the data",
        'insufficient_samples': "Insufficient sample size. At least 3 samples are required",
        'explanation': {
            "MDS Coordinates": "Low-dimensional coordinates obtained after multidimensional scaling analysis",
            "MDS Scatter Plot": "Show the distribution of samples in the low-dimensional space"
        },
        'interpretation': {
            "MDS Coordinates": "Can be used to observe the relative positional relationship of samples in the low-dimensional space",
            "MDS Scatter Plot": "Visually show the similarity between samples. Samples closer in distance are more similar"
        }
    }
}


class MultidimensionalScalingMDSApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data15.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def mds_method(self, data):
        """使用默认参数的MDS分析方法"""
        # 固定默认参数：2维、欧氏距离、度量MDS
        mds = MDS(
            n_components=2,  # 默认2维
            metric=True,  # 默认度量MDS
            dissimilarity='precomputed' if self._is_distance_matrix(data) else 'euclidean',  # 默认欧氏距离
            random_state=42  # 固定随机种子，保证结果可复现
        )
        mds_coords = mds.fit_transform(data)
        self.stress_value = mds.stress_  # 保存压力值（可选，用于结果展示）
        return mds_coords

    # 添加距离矩阵判断辅助方法（约第128行）
    def _is_distance_matrix(self, data):
        """判断输入是否为距离矩阵"""
        if data.shape[0] != data.shape[1]:
            return False
        if not np.allclose(np.diag(data), 0, atol=1e-6):
            return False
        if not np.allclose(data, data.T, atol=1e-6):
            return False
        return True

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 获取用户设置：是否有列名行和样本名所在列
            has_header = True  # 固定有列名行
            index_col = 0  # 固定第一列（0索引）为样本名

            # 读取Excel文件，区分列名和文本信息
            df = pd.read_excel(
                file_path,
                header=0 if has_header else None,  # 如果有列名行，将第一行作为列名
                index_col=index_col  # 设置样本名所在列（如果有）
            )

            # 处理缺失值
            if df.isnull().any().any():
                # 提示用户选择缺失值处理方式
                from tkinter import messagebox
                response = messagebox.askyesnocancel(
                    title=languages[self.current_language]["missing_values_title"],
                    message=languages[self.current_language]["missing_values_message"]
                )
                if response is None:  # 取消操作
                    return
                elif response:  # 填充均值
                    df = df.fillna(df.mean(numeric_only=True))
                else:  # 删除含缺失值的行
                    df = df.dropna()

            # 保存样本名（如果有）
            if index_col is not None:
                # 验证样本名列是否存在且为字符串类型
                if index_col >= len(df.columns):
                    raise ValueError(languages[self.current_language]["invalid_index_column"])
                sample_names = df.index.astype(str).tolist()
            else:
                sample_names = [f"样本{i + 1}" for i in range(len(df))]

            # 提取纯数值数据用于分析
            numeric_df = df.select_dtypes(include=[np.number])
            if numeric_df.empty:
                raise ValueError(languages[self.current_language]["no_numeric_data"])
            data = numeric_df.values.astype(float)

            # 验证样本量
            if len(data) < 3:
                raise ValueError(languages[self.current_language]["insufficient_samples"])

            # 将数据转换为浮点类型
            data = data.astype(float)

            # 自动判断并计算欧氏距离
            is_square = data.shape[0] == data.shape[1]
            has_zero_diagonal = np.allclose(np.diag(data), 0, atol=1e-6) if is_square else False

            if not (is_square and has_zero_diagonal):
                # 计算样本间欧氏距离（每行一个样本，每列一个特征）
                data = euclidean_distances(data)

            # 进行 MDS 分析
            mds_coords = self.mds_method(data)

            # 根据当前语言定义统计量名称和参数文本
            if self.current_language == "zh":
                stats_names = {
                    "mds_coords": "MDS坐标",
                    "parameters": "使用参数"
                }
                params_text = "维度: 2D, 距离度量: 欧氏距离, 类型: 度量MDS"
                headers = ["统计量", "统计量值"]
            else:
                stats_names = {
                    "mds_coords": "MDS Coordinates",
                    "parameters": "Parameters Used"
                }
                params_text = "Dimensions: 2D, Distance Metric: Euclidean Distance, Type: Metric MDS"
                headers = ["Statistic", "Value"]

            # 整理多语言数据
            data = [
                [stats_names["mds_coords"], mds_coords.tolist()],
                [stats_names["parameters"], params_text]
            ]
            df_result = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            explanations = languages[self.current_language]['explanation']
            interpretations = languages[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            if self.current_language == "zh":
                columns = ["MDS坐标", "MDS散点图"]
            else:
                columns = ["MDS Coordinates", "MDS Scatter Plot"]

            explanation_df = explanation_df.reindex(columns=columns)
            explanation_df.insert(0, "统计量_解释说明", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            if self.current_language == "zh":
                columns = ["MDS坐标", "MDS散点图"]
            else:
                columns = ["MDS Coordinates", "MDS Scatter Plot"]

            interpretation_df = interpretation_df.reindex(columns=columns)
            interpretation_df.insert(0, "统计量_结果解读",
                                     "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建一个新的 Word 文档
                doc = Document()

                # 添加统计结果表格
                # 根据当前语言切换一级标题文本
                if self.current_language == "zh":
                    doc.add_heading("统计结果", level=1)
                else:
                    doc.add_heading("Statistical Results", level=1)
                table = doc.add_table(rows=1, cols=len(headers))
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                for index, row in df_result.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)

                # 添加解释说明（改为项目符号列表）
                # 根据当前语言切换一级标题和二级段落标题
                if self.current_language == "zh":
                    doc.add_heading("解释说明", level=1)
                    doc.add_paragraph("统计量解释：", style='Heading 2')
                else:
                    doc.add_heading("Explanations", level=1)
                    doc.add_paragraph("Statistic Explanations:", style='Heading 2')
                for key, value in explanations.items():
                    para = doc.add_paragraph(style='List Bullet')
                    para.add_run(f"{key}：").bold = True
                    para.add_run(value)

                # 添加结果解读（改为项目符号列表）
                # 根据当前语言切换标题文本
                if self.current_language == "zh":
                    doc.add_heading("结果解读", level=1)
                    doc.add_paragraph("统计量解读：", style='heading 2')
                else:
                    doc.add_heading("Result Interpretation", level=1)
                    doc.add_paragraph("Statistic Interpretation:", style='heading 2')
                for key, value in interpretations.items():
                    para = doc.add_paragraph(style='List Bullet')
                    para.add_run(f"{key}：").bold = True
                    para.add_run(value)

                plt.figure(figsize=(10, 8))  # 增大画布尺寸
                ax = plt.gca()

                # 绘制散点图（使用随机颜色增强区分度）
                scatter = ax.scatter(
                    mds_coords[:, 0],
                    mds_coords[:, 1],
                    c=np.random.rand(len(mds_coords)),  # 随机颜色
                    cmap='viridis',
                    s=100,  # 点大小
                    edgecolors='black',
                    alpha=0.7
                )

                # 为每个点添加样本名标签（优化位置）
                for i, name in enumerate(sample_names):
                    ax.annotate(
                        name,
                        (mds_coords[i, 0], mds_coords[i, 1]),
                        fontsize=10,
                        ha='right',
                        xytext=(5, 5),  # 文本偏移量
                        textcoords='offset points',
                        bbox=dict(boxstyle="round,pad=0.3", fc="white", ec="gray", alpha=0.7)  # 添加白色背景框
                    )

                # 设置标题和轴标签
                ax.set_title(
                    'MDS散点图' if self.current_language == 'zh' else 'Scatter Plot of MDS',
                    fontsize=16
                )
                ax.set_xlabel(
                    '维度1' if self.current_language == 'zh' else 'Dimension 1',
                    fontsize=12
                )
                ax.set_ylabel(
                    '维度2' if self.current_language == 'zh' else 'Dimension 2',
                    fontsize=12
                )

                # 添加网格线
                ax.grid(True, linestyle='--', alpha=0.6)

                # 调整坐标轴范围（留出更多空间给标签）
                x_min, x_max = mds_coords[:, 0].min(), mds_coords[:, 0].max()
                y_min, y_max = mds_coords[:, 1].min(), mds_coords[:, 1].max()
                x_range = x_max - x_min
                y_range = y_max - y_min
                ax.set_xlim(x_min - 0.1 * x_range, x_max + 0.1 * x_range)
                ax.set_ylim(y_min - 0.1 * y_range, y_max + 0.1 * y_range)

                # 保存图片前先删除旧图（确保更新）
                img_path = os.path.splitext(save_path)[0] + '_mds_scatter.png'
                if os.path.exists(img_path):
                    os.remove(img_path)

                # 保存图片（确保完整显示所有元素）
                plt.tight_layout()
                plt.savefig(img_path, dpi=300, bbox_inches='tight')
                plt.close('all')  # 关闭所有图形，释放资源

                # 在 Word 文档中插入图片
                if self.current_language == "zh":
                    doc.add_heading("MDS 散点图", level=1)
                else:
                    doc.add_heading("MDS Scatter Plot", level=1)
                doc.add_picture(img_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                self.result_label.config(text=languages[self.current_language]["analysis_complete"].format(save_path),
                                         wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]["no_save_path_selected"])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]["analysis_error"].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])
        self.open_excel_label.config(text=languages[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建主框架，使用fill和expand让框架充满窗口
        main_frame = ttk.Frame(self.root)
        main_frame.pack(fill=tk.BOTH, expand=True)  # 关键：让主框架充满窗口

        # 创建内容框架，用于居中放置所有元素
        content_frame = ttk.Frame(main_frame, padding=20)
        content_frame.place(relx=0.5, rely=0.5, anchor=tk.CENTER)

        # 创建文件选择按钮
        self.select_button = ttk.Button(content_frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(content_frame, width=60)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(content_frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            content_frame,
            text=languages[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(content_frame,
                                               text=languages[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=5)

        # 创建结果显示标签（放在主框架底部，保持原有布局）
        self.result_label = ttk.Label(main_frame, text="", justify=tk.LEFT, wraplength=500)
        self.result_label.pack(pady=10, padx=20, side=tk.BOTTOM)  # 放在底部

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = MultidimensionalScalingMDSApp()
    app.run()


if __name__ == "__main__":
    run_app()

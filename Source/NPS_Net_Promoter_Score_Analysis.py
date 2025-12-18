import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
import tkinter.simpledialog
import matplotlib.pyplot as plt
import pathlib
from docx import Document
from docx.shared import Inches


# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
languages = {
    'zh': {
        'title': "NPS净推荐值",
        'select_button_text': "选择文件",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'analyze_button_text': "分析文件",
        'no_file_selected': "请选择文件。",
        'file_not_exists': "文件不存在，请重新选择。",
        'analysis_error': "分析文件时出错: {}",
        'analysis_complete': "分析完成，结果已保存到 {}，相关图片已保存。",
        'no_save_path_selected': "未选择保存路径，结果未保存。",
        'open_excel_button_text': "示例数据",
        'switch_language_button_text': "中/英",
        'explanation': {
            "推荐者": "给出9 - 10分的客户，是产品或服务的忠实拥护者，会积极推荐给他人。",
            "被动者": "给出7 - 8分的客户，对产品或服务基本满意，但不会主动推荐。",
            "贬损者": "给出0 - 6分的客户，对产品或服务不满意，可能会向他人抱怨。",
            "NPS净推荐值": "NPS = 推荐者比例 - 贬损者比例，反映了客户对产品或服务的整体态度。"
        },
        'interpretation': {
            "推荐者": "应关注推荐者的需求，提供更好的服务，鼓励他们继续推荐。",
            "被动者": "可以通过改进产品或服务，将被动者转化为推荐者。",
            "贬损者": "及时了解贬损者的不满原因，采取措施改进，避免负面影响扩大。",
            "NPS净推荐值": "NPS值越高，说明客户对产品或服务越满意，忠诚度越高。"
        },
        'chart_titles': {
            'type_percentage': '各类型占比情况',
            'score_distribution': '各分数段分布情况'
        },
        'chart_labels': {
            'percentage': '比例 (%)',
            'score': '分数',
            'count': '数量'
        },
        'table_headers': {
            'indicator': '指标',
            'value': '数值'
        },
        'sections': {
            'analysis_results': '分析结果',
            'explanation': '解释说明',
            'interpretation': '结果解读'
        }
    },
    'en': {
        'title': "NPS Net Promoter Score",
        'select_button_text': "Select Files",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'analyze_button_text': "Analyze Files",
        'no_file_selected': "Please select a valid file path.",
        'file_not_exists': "The file does not exist. Please select again.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'analysis_complete': "Analysis completed. The results have been saved to {}, and the relevant images have been saved.",
        'no_save_path_selected': "No save path selected. The results were not saved.",
        'open_excel_button_text': "Example data",
        'switch_language_button_text': "Chinese/English",
        'explanation': {
            "Promoters": "Customers who give a score of 9 - 10 are loyal advocates of the product or service and will actively recommend it to others.",
            "Passives": "Customers who give a score of 7 - 8 are generally satisfied with the product or service but will not actively recommend it.",
            "Detractors": "Customers who give a score of 0 - 6 are dissatisfied with the product or service and may complain to others.",
            "NPS Net Promoter Score": "NPS = Percentage of Promoters - Percentage of Detractors, which reflects the overall attitude of customers towards the product or service."
        },
        'interpretation': {
            "Promoters": "Pay attention to the needs of promoters, provide better services, and encourage them to continue recommending.",
            "Passives": "Improve the product or service to convert passives into promoters.",
            "Detractors": "Understand the reasons for detractors' dissatisfaction in a timely manner, take measures to improve, and avoid the expansion of negative impacts.",
            "NPS Net Promoter Score": "The higher the NPS value, the more satisfied and loyal the customers are with the product or service."
        },
        'chart_titles': {
            'type_percentage': 'Percentage of Each Type',
            'score_distribution': 'Score Distribution'
        },
        'chart_labels': {
            'percentage': 'Percentage (%)',
            'score': 'Score',
            'count': 'Count'
        },
        'table_headers': {
            'indicator': 'Indicator',
            'value': 'Value'
        },
        'sections': {
            'analysis_results': 'Analysis Results',
            'explanation': 'Explanation',
            'interpretation': 'Interpretation'
        }
    }
}

class NPSNetPromoterScoreAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"
        
        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])
            
        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data6.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")


    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def nps_analysis(self, data, question_columns):
        all_results = []
        for question_column in question_columns:
            if question_column not in data.columns:
                print(f"列名 {question_column} 不在文件中，跳过该列分析。")
                continue
            responses = data[question_column]
            promoters = (responses >= 9).sum()
            passives = ((responses >= 7) & (responses <= 8)).sum()
            detractors = (responses <= 6).sum()
            total = len(responses)
            promoter_percentage = promoters / total * 100
            passive_percentage = passives / total * 100
            detractor_percentage = detractors / total * 100
            nps = promoter_percentage - detractor_percentage

            # 各分数段占比情况
            score_bins = np.arange(-0.5, 10.6, 1)  # 每个整数分一个区间
            score_counts = pd.cut(responses, bins=score_bins, right=False, labels=range(0, 11)).value_counts().sort_index()
            score_percentages = score_counts / total * 100

            results = {
                f"{question_column}_推荐者数量": promoters,
                f"{question_column}_推荐者比例": promoter_percentage,
                f"{question_column}_被动者数量": passives,
                f"{question_column}_被动者比例": passive_percentage,
                f"{question_column}_贬损者数量": detractors,
                f"{question_column}_贬损者比例": detractor_percentage,
                f"{question_column}_NPS净推荐值": nps,
                f"{question_column}_各分数段占比情况": score_percentages
            }
            all_results.append(results)
        return all_results

    def analyze_file(self):
        file_paths = self.file_entry.get().split(", ")
        if not file_paths or file_paths[0] == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        for file_path in file_paths:
            if not os.path.exists(file_path):
                self.result_label.config(text=languages[self.current_language]["file_not_exists"])
                return
        try:
            question_columns = []
            while True:
                question_column = tkinter.simpledialog.askstring(
                    "输入信息",
                    "请输入NPS问题的列名（点击取消结束输入）",
                    parent=self.root,  # 指定父窗口
                )
                if question_column is None:
                    break
                if question_column.strip():
                    question_columns.append(question_column.strip())
                else:
                    print("输入的列名不能为空，请重新输入。")

            if not question_columns:
                self.result_label.config(text="未输入有效的问题列名，分析取消。")
                return

            all_results = []
            all_score_percentages = []
            file_names = []
            for file_path in file_paths:
                # 打开 Excel 文件
                df = pd.read_excel(file_path)

                # 进行NPS分析
                nps_results = self.nps_analysis(df, question_columns)
                all_results.extend(nps_results)
                for result in nps_results:
                    for key, value in result.items():
                        if "_各分数段占比情况" in key:
                            all_score_percentages.append(value)
                file_names.extend([os.path.basename(file_path)] * len([res for res in nps_results if res]))

            # 整理数据
            all_data = []
            for i, results in enumerate(all_results):
                if results:
                    data = []
                    for key, value in results.items():
                        if "_各分数段占比情况" not in key:
                            data.append([f"{file_names[i]}_{key}", value])
                    all_data.extend(data)
            headers = [
                languages[self.current_language]['table_headers']['indicator'],
                languages[self.current_language]['table_headers']['value']
            ]
            df_result = pd.DataFrame(all_data, columns=headers)

            # 添加解释说明
            explanations = languages[self.current_language]['explanation']
            interpretations = languages[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(
                columns=["推荐者" if self.current_language == 'zh' else "Promoters", 
                         "被动者" if self.current_language == 'zh' else "Passives", 
                         "贬损者" if self.current_language == 'zh' else "Detractors", 
                         "NPS净推荐值" if self.current_language == 'zh' else "NPS Net Promoter Score"])
            explanation_df.insert(0, "指标_解释说明",
                              "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(
                columns=["推荐者" if self.current_language == 'zh' else "Promoters", 
                         "被动者" if self.current_language == 'zh' else "Passives", 
                         "贬损者" if self.current_language == 'zh' else "Detractors", 
                         "NPS净推荐值" if self.current_language == 'zh' else "NPS Net Promoter Score"])
            interpretation_df.insert(0, "指标_结果解读",
                                 "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 合并数据、解释说明和结果解读
            combined_df = pd.concat(
                [df_result, explanation_df, interpretation_df], ignore_index=True)

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建一个新的 Word 文档
                doc = Document()

                # 添加标题
                doc.add_heading(
                    'NPS净推荐值分析结果' if self.current_language == 'zh' else 'NPS Net Promoter Score Analysis Results',
                    0)

                # 添加分析结果表格
                doc.add_heading(languages[self.current_language]['sections']['analysis_results'], level=1)
                table = doc.add_table(rows=1, cols=len(df_result.columns))
                hdr_cells = table.rows[0].cells
                # 使用语言字典中的表头
                hdr_cells[0].text = languages[self.current_language]['table_headers']['indicator']
                hdr_cells[1].text = languages[self.current_language]['table_headers']['value']
                for i, col in enumerate(df_result.columns):  # 表头仅包含分析结果列
                    hdr_cells[i].text = col

                # 添加数据到表格（仅添加分析结果数据，不包含解释和解读）
                for index, row in df_result.iterrows():
                    row_cells = table.add_row().cells
                    # 只处理前两列数据
                    for i in range(min(2, len(row))):  # 确保只取前两列
                        row_cells[i].text = str(row.iloc[i])

                # 添加解释说明（项目符号列表）
                doc.add_heading(languages[self.current_language]['sections']['explanation'], level=1)
                expl_paragraph = doc.add_paragraph()
                for key, value in explanations.items():
                    run = expl_paragraph.add_run()
                    run.add_text(f'• {key}: {value}\n')

                # 添加结果解读（项目符号列表）
                doc.add_heading(languages[self.current_language]['sections']['interpretation'], level=1)
                inter_paragraph = doc.add_paragraph()
                for key, value in interpretations.items():
                    run = inter_paragraph.add_run()
                    run.add_text(f'• {key}: {value}\n')

                # 生成各类型占比情况柱状图（保持不变）
                categories = ["推荐者" if self.current_language == 'zh' else "Promoters",
                              "被动者" if self.current_language == 'zh' else "Passives",
                              "贬损者" if self.current_language == 'zh' else "Detractors"]
                percentages_list = []
                valid_results = [res for res in all_results if res]
                for result in valid_results:
                    base_key = list(result.keys())[0].split('_')[0]
                    keys = [f"{base_key}_推荐者比例", f"{base_key}_被动者比例", f"{base_key}_贬损者比例"]
                    if all(key in result for key in keys):
                        percentages = [result[key] for key in keys]
                        percentages_list.append(percentages)

                # 绘制柱状图
                plt.figure(figsize=(10, 6))
                x = np.arange(len(categories))
                width = 0.8 / len(percentages_list) if percentages_list else 0.8
                for i, percentages in enumerate(percentages_list):
                    plt.bar(x + i * width, percentages, width, label=f"File {i + 1}")

                plt.title(languages[self.current_language]['chart_titles']['type_percentage'])
                plt.ylabel(languages[self.current_language]['chart_labels']['percentage'])
                plt.xticks(x + width * (len(percentages_list) - 1) / 2 if percentages_list else x, categories)


                # 保存图片并添加到文档
                img_path = os.path.splitext(save_path)[0] + '_types.png'
                plt.savefig(img_path)
                plt.close()
                doc.add_picture(img_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                self.result_label.config(text=languages[self.current_language]["analysis_complete"].format(save_path), wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]["no_save_path_selected"])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]["analysis_error"].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])
        self.open_excel_label.config(text=languages[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"], 
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"], 
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=languages[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=languages[self.current_language]["switch_language_button_text"],
                                              foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)
        
    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = NPSNetPromoterScoreAnalysisApp()
    app.run()

if __name__ == "__main__":
    run_app()
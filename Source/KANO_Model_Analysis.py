import tkinter as tk
from tkinter import filedialog
import tkinter.simpledialog
import os
import pandas as pd
import numpy as np
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import pathlib

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'  # 设置字体为黑体，可根据系统情况修改为其他支持中文的字体
plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

# 定义语言字典
languages = {
    'zh': {
        'title': "KANO模型分析",
        'select_button_text': "选择文件",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'analyze_button_text': "分析文件",
        'no_file_selected': "请选择文件。",
        'file_not_exists': "文件不存在，请重新选择。",
        'analysis_error': "分析文件时出错: {}",
        'analysis_success': "分析完成，结果已保存到 {}",
        'no_save_path': "未选择保存路径，结果未保存。",
        'open_excel_button_text': "示例数据",
        'switch_language_button_text': "中/英",
        'explanation': {
            "基本型需求（M）": "用户认为产品必须具备的功能，缺乏这些功能会导致用户不满。",
            "期望型需求（O）": "用户的满意度随该需求的满足程度而线性增加。",
            "兴奋型需求（A）": "用户没有预期到的需求，满足这些需求会极大提高用户满意度。",
            "无差异型需求（I）": "用户对该需求的满足与否不太关心。",
            "反向型需求（R）": "满足该需求会导致用户不满。",
            "可疑结果（Q）": "回答存在矛盾，结果不可靠。"
        },
        'interpretation': {
            "基本型需求（M）": "应确保产品满足基本型需求，以避免用户不满。",
            "期望型需求（O）": "可根据资源情况逐步提升期望型需求的满足程度，以提高用户满意度。",
            "兴奋型需求（A）": "挖掘和满足兴奋型需求可以使产品脱颖而出，吸引更多用户。",
            "无差异型需求（I）": "可以适当减少在无差异型需求上的投入。",
            "反向型需求（R）": "应避免满足反向型需求，以免引起用户不满。",
            "可疑结果（Q）": "需要重新确认用户回答，确保结果可靠性。"
        },
        'better_worse': {
            'better': 'Better系数',
            'worse': 'Worse系数'
        }
    },
    'en': {
        'title': "KANO Model Analysis",
        'select_button_text': "Select File",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'analyze_button_text': "Analyze File",
        'no_file_selected': "Please select a valid file path.",
        'file_not_exists': "The file does not exist. Please select again.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'analysis_success': "Analysis completed. The results have been saved to {}",
        'no_save_path': "No save path selected. The results were not saved.",
        'open_excel_button_text': "Example data",
        'switch_language_button_text': "Chinese/English",
        'explanation': {
            "Must Must-be Requirements (M)": "Basic requirements that users expect the product to have. Lack of these features will lead to user dissatisfaction.",
            "One-dimensional Requirements (O)": "Expected requirements where user satisfaction increases linearly with the degree of fulfillment.",
            "Attractive Requirements (A)": "Exciting requirements that users do not expect. Meeting these requirements can greatly improve user satisfaction.",
            "Indifferent Requirements (I)": "Indifferent requirements that users do not care much about whether they are met or not.",
            "Reverse Requirements (R)": "Reverse requirements where meeting them will lead to user dissatisfaction.",
            "Questionable Results (Q)": "The responses are contradictory, and the results are unreliable."
        },
        'interpretation': {
            "Must-be Requirements (M)": "Ensure that the product meets basic requirements to avoid user dissatisfaction.",
            "One-dimensional Requirements (O)": "Gradually improve the fulfillment of expected requirements according to available resources to enhance user satisfaction.",
            "Attractive Requirements (A)": "Discover and meet exciting requirements to make the product stand out and attract more users.",
            "Indifferent Requirements (I)": "Reduce investment in indifferent requirements appropriately.",
            "Reverse Requirements (R)": "Avoid meeting reverse requirements to prevent user dissatisfaction.",
            "Questionable Results (Q)": "Reconfirm user responses to ensure result reliability."
        },
        'better_worse': {
            'better': 'Better Coefficient',
            'worse': 'Worse Coefficient'
        }
    }
}


class KANOModelAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data5.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def kano_analysis(self, data, positive_question_columns, negative_question_columns):
        kano_results = {}
        better_worse_results = {}
        for i in range(len(positive_question_columns)):
            positive_responses = data[positive_question_columns[i]]
            negative_responses = data[negative_question_columns[i]]
            category = self.classify_kano(positive_responses, negative_responses)
            kano_results[positive_question_columns[i]] = category
            better, worse = self.calculate_better_worse(positive_responses, negative_responses)
            better_worse_results[positive_question_columns[i]] = (better, worse)
        return kano_results, better_worse_results

    def classify_kano(self, positive_responses, negative_responses):
        counts = {
            'A': 0, 'O': 0, 'M': 0, 'I': 0, 'R': 0, 'Q': 0
        }
        kano_matrix = {
            (1, 1): 'Q', (1, 2): 'A', (1, 3): 'A', (1, 4): 'A', (1, 5): 'O',
            (2, 1): 'R', (2, 2): 'I', (2, 3): 'I', (2, 4): 'I', (2, 5): 'M',
            (3, 1): 'R', (3, 2): 'I', (3, 3): 'I', (3, 4): 'I', (3, 5): 'M',
            (4, 1): 'R', (4, 2): 'I', (4, 3): 'I', (4, 4): 'I', (4, 5): 'M',
            (5, 1): 'R', (5, 2): 'R', (5, 3): 'R', (5, 4): 'R', (5, 5): 'Q',
        }

        counts = {k: 0 for k in ['A', 'O', 'M', 'I', 'R', 'Q']}
        for pos, neg in zip(positive_responses, negative_responses):
            if pd.isna(pos) or pd.isna(neg):
                continue
            key = (int(pos), int(neg))
            category = kano_matrix.get(key, 'Q')
            counts[category] += 1
        max_count_category = max(counts, key=counts.get)
        if self.current_language == 'zh':
            category_mapping = {
                'A': "兴奋型需求（A）",
                'O': "期望型需求（O）",
                'M': "基本型需求（M）",
                'I': "无差异型需求（I）",
                'R': "反向型需求（R）",
                'Q': "可疑结果（Q）"
            }
        else:
            category_mapping = {
                'A': "Attractive Requirements (A)",
                'O': "One-dimensional Requirements (O)",
                'M': "Must-be Requirements (M)",
                'I': "Indifferent Requirements (I)",
                'R': "Reverse Requirements (R)",
                'Q': "Questionable Results (Q)"
            }
        return category_mapping[max_count_category]

    def calculate_better_worse(self, positive_responses, negative_responses):
        a_count = 0
        o_count = 0
        m_count = 0
        i_count = 0
        r_count = 0
        total_count = len(positive_responses)
        for pos, neg in zip(positive_responses, negative_responses):
            if pos == 5 and neg == 1:
                a_count += 1
            elif pos == 5 and neg == 2:
                a_count += 1
            elif pos == 5 and neg == 3:
                o_count += 1
            elif pos == 4 and neg == 1:
                a_count += 1
            elif pos == 4 and neg == 2:
                o_count += 1
            elif pos == 4 and neg == 3:
                o_count += 1
            elif pos == 3 and neg == 1:
                o_count += 1
            elif pos == 3 and neg == 2:
                o_count += 1
            elif pos == 3 and neg == 3:
                i_count += 1
            elif pos == 2 and neg == 1:
                i_count += 1
            elif pos == 2 and neg == 2:
                i_count += 1
            elif pos == 2 and neg == 3:
                m_count += 1
            elif pos == 1 and neg == 1:
                r_count += 1
            elif pos == 1 and neg == 2:
                m_count += 1
            elif pos == 1 and neg == 3:
                m_count += 1
        denominator = a_count + o_count + m_count + i_count
        if denominator == 0:
            better, worse = 0, 0
        else:
            better = (a_count + o_count) / denominator
            worse = - (o_count + m_count) / denominator
        return better, worse

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def generate_better_worse_plot(self, df_result, save_path):
        better = df_result[languages[self.current_language]['better_worse']['better']]
        worse = df_result[languages[self.current_language]['better_worse']['worse']]
        question_col = "问题" if self.current_language == 'zh' else "Question"
        labels = df_result[question_col]

        plt.figure(figsize=(10, 8))
        plt.scatter(better, worse)
        for i, label in enumerate(labels):
            plt.annotate(label, (better[i], worse[i]), textcoords="offset points", xytext=(0, 10), ha='center')
        plt.axhline(y=0, color='k')
        plt.axvline(x=0, color='k')
        plt.xlabel(languages[self.current_language]['better_worse']['better'])
        plt.ylabel(languages[self.current_language]['better_worse']['worse'])
        plt.title('Better - Worse 象限图' if self.current_language == 'zh' else 'Better - Worse Quadrant Plot')
        img_path = os.path.splitext(save_path)[0] + '_better_worse.png'
        plt.savefig(img_path)
        plt.close()
        return img_path

    def generate_kano_plot(self, df_result, save_path):
        category_col = "KANO分类" if self.current_language == 'zh' else "KANO Category"
        category_counts = df_result[category_col].value_counts()
        plt.figure(figsize=(10, 8))
        plt.bar(category_counts.index, category_counts.values)
        plt.xlabel('KANO分类' if self.current_language == 'zh' else 'KANO Category')
        plt.ylabel('数量' if self.current_language == 'zh' else 'Count')
        plt.title('KANO模型分析结果' if self.current_language == 'zh' else 'KANO Model Analysis Results')
        img_path = os.path.splitext(save_path)[0] + '_kano.png'
        plt.savefig(img_path)
        plt.close()
        return img_path

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 让用户输入正向问题和负向问题的列名
            positive_question_columns = tkinter.simpledialog.askstring("输入信息", "请输入正向问题的列名，用逗号分隔").split(',')
            negative_question_columns = tkinter.simpledialog.askstring("输入信息", "请输入负向问题的列名，用逗号分隔").split(',')

            if not positive_question_columns or not negative_question_columns:
                self.result_label.config(text="未输入完整的问题列名，分析取消。")
                return

            # 进行KANO模型分析
            kano_results, better_worse_results = self.kano_analysis(df, positive_question_columns, negative_question_columns)

            # 整理数据
            data = []
            for question, category in kano_results.items():
                better, worse = better_worse_results[question]
                data.append([question, category, better, worse])
            headers = [
                "问题" if self.current_language == 'zh' else "Question",
                "KANO分类" if self.current_language == 'zh' else "KANO Category",
                languages[self.current_language]['better_worse']['better'],
                languages[self.current_language]['better_worse']['worse']
            ]
            df_result = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            explanations = languages[self.current_language]['explanation']
            interpretations = languages[self.current_language]['interpretation']

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(
                defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加分析结果表格
                doc.add_heading('KANO模型分析结果' if self.current_language == 'zh' else 'KANO Model Analysis Results', level=1)
                table = doc.add_table(rows=1, cols=len(headers))
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                for row in df_result.values.tolist():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)

                # 添加解释说明
                doc.add_heading('KANO分类解释说明' if self.current_language == 'zh' else 'Explanation of KANO Categories', level=1)
                for category, explanation in explanations.items():
                    doc.add_paragraph(f'{category}: {explanation}')

                # 添加分析结果解读
                doc.add_heading('KANO分类结果解读' if self.current_language == 'zh' else 'Interpretation of KANO Results', level=1)
                for category, interpretation in interpretations.items():
                    doc.add_paragraph(f'{category}: {interpretation}')

                # 生成 Better 和 Worse 象限图
                better_worse_path = self.generate_better_worse_plot(df_result, save_path)
                # 生成 KANO 图
                kano_path = self.generate_kano_plot(df_result, save_path)

                # 添加图片到 Word 文档
                doc.add_heading('Better - Worse 象限图' if self.current_language == 'zh' else 'Better - Worse Quadrant Chart', level=1)
                doc.add_picture(better_worse_path, width=Inches(6))
                doc.add_heading('KANO 分析图' if self.current_language == 'zh' else 'KANO Analysis Chart', level=1)
                doc.add_picture(kano_path, width=Inches(6))

                if self.current_language == 'zh':
                    doc.add_paragraph("象限解释：\n右上象限（兴奋型）→ 满足后用户满意度高；\n"
                                      "左上象限（基本型）→ 不满足则不满；\n"
                                      "右下象限（无差异型）→ 满足与否差别不大；\n"
                                      "左下象限（反向型）→ 满足反而导致不满。")
                else:
                    doc.add_paragraph(
                        "Quadrant interpretation:\nTop-right (Attractive): greatly increases satisfaction;\n"
                        "Top-left (Must-be): dissatisfaction if not met;\n"
                        "Bottom-right (Indifferent): little effect;\n"
                        "Bottom-left (Reverse): satisfaction decreases if met.")

                # 保存 Word 文档
                doc.save(save_path)

                result_msg = languages[self.current_language]['analysis_success'].format(save_path)
                result_msg += f"\nBetter-Worse 象限图已保存到 {better_worse_path}"
                result_msg += f"\nKANO 分析图已保存到 {kano_path}"
                self.result_label.config(text=result_msg, wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])
        self.open_excel_label.config(text=languages[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=languages[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=languages[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = KANOModelAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()
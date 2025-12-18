import tkinter as tk
from tkinter import filedialog, messagebox
import os
import numpy as np
import pandas as pd
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import accuracy_score, roc_auc_score, roc_curve, classification_report, confusion_matrix, f1_score
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import label_binarize, OneHotEncoder
from sklearn.compose import ColumnTransformer
from sklearn.pipeline import Pipeline
from sklearn.impute import SimpleImputer
import statsmodels.api as sm
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import warnings

# 设置支持中文的字体
plt.rcParams['font.family'] = ['SimHei', 'WenQuanYi Micro Hei', 'Heiti TC']
plt.rcParams['axes.unicode_minus'] = False

# 忽略特定警告
warnings.filterwarnings("ignore", category=FutureWarning, message="'multi_class' was deprecated in version 1.5")
warnings.filterwarnings("ignore", category=RuntimeWarning, message="overflow encountered in exp")
warnings.filterwarnings("ignore", category=RuntimeWarning, message="invalid value encountered in divide")

# 定义语言字典
languages = {
    'zh': {
        'title': "多分类Logit回归",
        'select_button_text': "选择文件",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'analyze_button_text': "分析文件",
        'no_file_selected': "请选择文件。",
        'file_not_exists': "文件不存在，请重新选择。",
        'analysis_error': "分析文件时出错: {}",
        'analysis_complete': "分析完成，结果已保存到 {}（Word）和 {}（Excel），相关图片已保存。",
        'no_save_path_selected': "未选择保存路径，结果未保存。",
        'switch_language_button_text': "中/英",
        'open_excel_button_text': "示例数据",
        'images_saved': "图片已保存到 {}",
        'data_preprocessing': "数据预处理完成：\n- 缺失值已填充\n- 分类变量已编码\n- 异常值已处理",
        'train_test_split': "数据集划分：训练集{:.0f}%，测试集{:.0f}%",
        'select_save_location': "请选择保存结果的位置",
        'explanation_heading': '指标说明',
        'explanation': {
            "Coefficients": "回归系数，表示每个自变量对因变量的影响程度。",
            "Intercept": "截距，是当所有自变量为 0 时因变量的预测值。",
            "Accuracy": "准确率，衡量模型预测正确的比例。",
            "ROC-AUC": "ROC曲线下面积，衡量模型的分类能力。"
        },
        'performance_metrics': {
            "Model": "模型",
            "Accuracy": "准确率",
            "F1 Score (Macro)": "F1分数（宏平均）",
            "F1 Score (Weighted)": "F1分数（加权平均）",
            "ROC-AUC (Macro)": "ROC-AUC（宏平均）",
            "ROC-AUC (Weighted)": "ROC-AUC（加权平均）",
            "Solver Used": "使用的求解器",
            "Regularization (C)": "正则化参数(C)",
            "Note": "说明"
        },
        'classification_report': {
            'precision': '精确率',
            'recall': '召回率',
            'f1-score': 'F1分数',
            'support': '支持数',
            'class': '类别',
            'macro_avg': '宏平均',
            'weighted_avg': '加权平均',
            'accuracy': '准确率'
        },
        'progress_loading_data': "正在加载数据...",
        'progress_preprocessing': "正在进行数据预处理...",
        'progress_training_model': "正在训练模型...",
        'progress_generating_results': "正在生成结果...",
        'regression_equations': "回归方程",
        'feature_limit_warning': "特征数量过多（{}个），已限制为前{}个特征以避免文件问题"
    },
    'en': {
        'title': "Multinomial Logit Regression",
        'select_button_text': "Select File",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'analyze_button_text': "Analyze File",
        'no_file_selected': "Please select a valid file path.",
        'file_not_exists': "The file does not exist. Please select again.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'analysis_complete': "Analysis completed. Results saved to {} (Word) and {} (Excel). Related images saved.",
        'no_save_path_selected': "No save path selected. The results were not saved.",
        'switch_language_button_text': "Chinese/English",
        'open_excel_button_text': "Example data",
        'images_saved': "Images have been saved to {}",
        'data_preprocessing': "Data preprocessing completed:\n- Missing values imputed\n- Categorical variables encoded\n- Outliers handled",
        'train_test_split': "Dataset split: Train {:.0f}%, Test {:.0f}%",
        'select_save_location': "Please select where to save the results",
        'explanation_heading': 'Explanations',
        'explanation': {
            "Coefficients": "Regression coefficients, indicating the influence of each independent variable on the dependent variable.",
            "Intercept": "Intercept, which is the predicted value of the dependent variable when all independent variables are 0.",
            "Accuracy": "Accuracy, measuring the proportion of correct predictions of the model.",
            "ROC-AUC": "Area under the ROC curve, measuring the classification ability of the model."
        },
        'performance_metrics': {
            "Model": "Model",
            "Accuracy": "Accuracy",
            "F1 Score (Macro)": "F1 Score (Macro)",
            "F1 Score (Weighted)": "F1 Score (Weighted)",
            "ROC-AUC (Macro)": "ROC-AUC (Macro)",
            "ROC-AUC (Weighted)": "ROC-AUC (Weighted)",
            "Solver Used": "Solver Used",
            "Regularization (C)": "Regularization (C)",
            "Note": "Note"
        },
        'classification_report': {
            'precision': 'Precision',
            'recall': 'Recall',
            'f1-score': 'F1-Score',
            'support': 'Support',
            'class': 'Class',
            'macro_avg': 'Macro Avg',
            'weighted_avg': 'Weighted Avg',
            'accuracy': 'Accuracy'
        },
        'progress_loading_data': "Loading data...",
        'progress_preprocessing': "Performing data preprocessing...",
        'progress_training_model': "Training model...",
        'progress_generating_results': "Generating results...",
        'regression_equations': "Regression Equations",
        'feature_limit_warning': "Too many features ({}), limited to first {} features to avoid file issues"
    }
}


class MultinomialLogitRegressionApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"
        # 最大特征数量限制（防止Word表格过大）
        self.MAX_FEATURES = 30

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data35.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def detect_categorical_columns(self, df, threshold=10):
        """改进的分类变量检测逻辑"""
        categorical_cols = []
        for col in df.columns:
            # 明确的非数值类型直接判定为分类变量
            if not pd.api.types.is_numeric_dtype(df[col]):
                categorical_cols.append(col)
                continue

            unique_count = df[col].nunique()
            total_count = len(df)

            # 数值型但唯一值少且为整数（可能是编码分类）
            if (unique_count < threshold and
                    np.issubdtype(df[col].dtype, np.integer) and
                    unique_count / total_count < 0.2):
                categorical_cols.append(col)

        return categorical_cols

    def handle_outliers(self, df):
        """处理数值型变量的异常值（使用IQR方法）"""
        numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
        # 排除因变量和已识别的分类变量
        categorical_cols = self.detect_categorical_columns(df)
        if df.columns[-1] in numeric_cols:
            numeric_cols.remove(df.columns[-1])
        numeric_cols = [col for col in numeric_cols if col not in categorical_cols]

        for col in numeric_cols:
            q1 = df[col].quantile(0.25)
            q3 = df[col].quantile(0.75)
            iqr = q3 - q1
            lower_bound = q1 - 1.5 * iqr
            upper_bound = q3 + 1.5 * iqr
            # 截断异常值
            df[col] = df[col].clip(lower_bound, upper_bound)
        return df

    def get_significance(self, p_value):
        """根据p值判断显著性"""
        if p_value < 0.001:
            return "***"
        elif p_value < 0.01:
            return "**"
        elif p_value < 0.05:
            return "*"
        else:
            return ""

    def train_logistic_regression(self, X_train, y_train):
        """训练逻辑回归模型，自动处理不收敛情况并提供正则化参数选择"""
        solvers = ['lbfgs', 'newton-cg', 'sag', 'saga']
        C_values = [0.01, 0.1, 1, 10, 100]  # 正则化参数选项

        for C in C_values:
            for solver in solvers:
                try:
                    # 移除multi_class参数以避免FutureWarning
                    logit = LogisticRegression(
                        solver=solver,
                        C=C,  # 正则化参数，较小的值表示更强的正则化
                        max_iter=1000,
                        random_state=42
                    )
                    logit.fit(X_train, y_train)
                    self.result_label.config(
                        text=f"Model trained with solver: {solver}, C: {C}"
                    )
                    self.root.update()
                    return logit, solver, C
                except Exception as e:
                    continue

        # 如果所有求解器都失败，尝试增加迭代次数
        logit = LogisticRegression(
            solver='saga',  # saga通常对收敛问题更稳健
            C=1,
            max_iter=5000,
            random_state=42
        )
        logit.fit(X_train, y_train)
        return logit, 'saga', 1

    def generate_regression_equations(self, unique_classes, intercept, coefficients, feature_names):
        """生成所有非基准类别之间的概率对比回归方程"""
        equations = []
        n_classes = len(unique_classes)

        # 生成所有类别之间的对比方程（i > j）
        for i in range(n_classes):
            for j in range(i):
                cls_i = unique_classes[i]
                cls_j = unique_classes[j]

                # 计算两个类别相对于基准类别的对数几率差
                intercept_diff = intercept[i] - intercept[j]
                eq_parts = [f"ln(P({cls_i})/P({cls_j})) = {intercept_diff:.4f}"]

                # 确保特征索引不越界
                max_feature_index = min(len(feature_names), coefficients.shape[1])
                for k in range(max_feature_index):
                    feature = feature_names[k]
                    # 系数差 = 类别i的系数 - 类别j的系数
                    coef_diff = coefficients[i, k] - coefficients[j, k]

                    if coef_diff >= 0:
                        eq_parts.append(f"+ {coef_diff:.4f}×{feature}")
                    else:
                        eq_parts.append(f"- {abs(coef_diff):.4f}×{feature}")

                equations.append(f"Class {cls_i} vs Class {cls_j}: {' '.join(eq_parts)}")

        return equations

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 显示进度提示
            self.result_label.config(text=languages[self.current_language]["progress_loading_data"])
            self.root.update()

            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 显示进度提示
            self.result_label.config(text=languages[self.current_language]["progress_preprocessing"])
            self.root.update()

            # 处理异常值
            df = self.handle_outliers(df)

            # 清洗列名（去除空格和隐藏字符）
            df.columns = df.columns.astype(str).str.strip()  # 仅去除首尾空格

            # 检测分类变量
            categorical_cols = self.detect_categorical_columns(df)
            # 确保分类列都存在于数据框中
            categorical_cols = [col for col in categorical_cols if col in df.columns]
            # 确保数值列都存在于数据框中（排除因变量列df.columns[-1]）
            numeric_cols = [col for col in df.columns[:-1] if col in df.columns and col not in categorical_cols]

            # 数据预处理管道
            numeric_transformer = Pipeline(steps=[
                ('imputer', SimpleImputer(strategy='mean'))
            ])

            categorical_transformer = Pipeline(steps=[
                ('imputer', SimpleImputer(strategy='most_frequent')),
                ('onehot', OneHotEncoder(drop='first', sparse_output=False))
            ])

            # 确保 numeric_cols 和 categorical_cols 均存在于 X 的列中
            numeric_cols = [col for col in numeric_cols if col in df.columns]
            categorical_cols = [col for col in categorical_cols if col in df.columns]

            # 防止没有分类或数值列时报错
            if not numeric_cols and not categorical_cols:
                raise ValueError("No valid numeric or categorical columns found in the data.")

            transformers = []
            # 再次验证数值列存在性
            valid_numeric = [col for col in numeric_cols if col in df.columns]
            if valid_numeric:
                transformers.append(('num', numeric_transformer, valid_numeric))
            # 再次验证分类列存在性
            valid_categorical = [col for col in categorical_cols if col in df.columns]
            if valid_categorical:
                transformers.append(('cat', categorical_transformer, valid_categorical))

            preprocessor = ColumnTransformer(transformers=transformers)

            # 划分自变量和因变量
            X = df.iloc[:, :-1].copy()
            y = df.iloc[:, -1].values

            # 重新基于 X 的列名筛选 numeric/categorical（避免名称不一致）
            numeric_cols = [col for col in numeric_cols if col in X.columns]
            categorical_cols = [col for col in categorical_cols if col in X.columns]

            # 如果两者都为空，抛错
            if not numeric_cols and not categorical_cols:
                raise ValueError(
                    "No valid numeric or categorical columns found in the data. Check your input file's columns.")

            # 构造 transformers
            transformers = []
            if numeric_cols:
                numeric_transformer = Pipeline(steps=[('imputer', SimpleImputer(strategy='mean'))])
                transformers.append(('num', numeric_transformer, numeric_cols))

            if categorical_cols:
                categorical_transformer = Pipeline(steps=[
                    ('imputer', SimpleImputer(strategy='most_frequent')),
                    ('onehot', OneHotEncoder(drop='first', sparse_output=False))
                ])
                transformers.append(('cat', categorical_transformer, categorical_cols))

            # 如果 transformers 列表为空，说明没有任何能被处理的列
            if not transformers:
                raise ValueError("No transformers to apply: check numeric_cols / categorical_cols after filtering.")

            preprocessor = ColumnTransformer(transformers=transformers, remainder='drop')

            # 划分训练集和测试集（7:3）
            X_train, X_test, y_train, y_test = train_test_split(
                X, y, test_size=0.3, random_state=42, stratify=y)
            self.result_label.config(
                text=languages[self.current_language]["train_test_split"].format(70, 30)
            )
            self.root.update()

            # 在 fit_transform 前再做一次列存在性校验，若不通过则给出明确错误信息
            requested_cols = []
            for name, _, cols in transformers:
                requested_cols.extend(cols)

            missing_cols = [c for c in requested_cols if c not in X_train.columns]
            if missing_cols:
                # 这个错误信息会帮助你定位到底哪些列没找到
                raise ValueError(
                    f"The following columns specified for preprocessing are not in the DataFrame: {missing_cols}. DataFrame columns: {list(X_train.columns)}")

            # 预处理数据
            X_train_processed = preprocessor.fit_transform(X_train)
            X_test_processed = preprocessor.transform(X_test)

            # 获取处理后的特征名称（安全访问 'cat' transformer）
            cat_feature_names = []
            if any(t[0] == 'cat' for t in transformers):
                ohe = preprocessor.named_transformers_['cat'].named_steps['onehot']
                cat_feature_names = list(ohe.get_feature_names_out(categorical_cols))
            feature_names = numeric_cols + cat_feature_names
            # 替换结束

            # 显示进度提示
            self.result_label.config(text=languages[self.current_language]["progress_training_model"])
            self.root.update()

            # 多分类Logit回归（自动选择合适的求解器和正则化参数）
            logit, used_solver, used_C = self.train_logistic_regression(X_train_processed, y_train)

            # 在测试集上预测
            y_pred = logit.predict(X_test_processed)
            y_pred_proba = logit.predict_proba(X_test_processed)

            # 计算评估指标（使用测试集）
            coefficients = logit.coef_
            intercept = logit.intercept_
            accuracy = accuracy_score(y_test, y_pred)
            class_report = classification_report(y_test, y_pred, output_dict=True)
            f1_macro = f1_score(y_test, y_pred, average='macro')
            f1_weighted = f1_score(y_test, y_pred, average='weighted')
            conf_matrix = confusion_matrix(y_test, y_pred)

            # 计算多分类ROC-AUC（提供两种平均方式）
            unique_classes = np.unique(y)
            n_classes = len(unique_classes)
            y_test_binarized = label_binarize(y_test, classes=unique_classes)

            roc_auc_macro = roc_auc_score(
                y_test_binarized, y_pred_proba,
                multi_class='ovr', average='macro'
            )
            roc_auc_weighted = roc_auc_score(
                y_test_binarized, y_pred_proba,
                multi_class='ovr', average='weighted'
            )

            # 使用statsmodels进行多分类分析（基准类别为第一个类别）
            X_train_sm = sm.add_constant(X_train_processed)
            sm_model = sm.MNLogit(y_train, X_train_sm).fit(maxiter=1000)
            z_values = sm_model.tvalues
            p_values = sm_model.pvalues
            coefficients = logit.coef_  # 补充定义coefficients

            # 特征数量限制（防止表格过大导致Word无法打开）
            original_feature_count = len(feature_names)
            if original_feature_count > self.MAX_FEATURES:
                # 截取特征名（前MAX_FEATURES个）
                feature_names = feature_names[:self.MAX_FEATURES]
                # 系数矩阵列数对应特征数，直接截取
                coefficients = coefficients[:, :self.MAX_FEATURES]
                # z_values和p_values第一列是常数项，从第二列开始截取特征对应的值
                z_values = z_values[:, [0] + list(range(1, min(self.MAX_FEATURES + 1, z_values.shape[1])))]
                p_values = p_values[:, [0] + list(range(1, min(self.MAX_FEATURES + 1, p_values.shape[1])))]
                # 显示警告信息
                self.result_label.config(
                    text=languages[self.current_language]["feature_limit_warning"].format(
                        original_feature_count, self.MAX_FEATURES)
                )
                self.root.update()

            # 确保z_values和p_values是numpy数组
            if not isinstance(z_values, np.ndarray):
                z_values = z_values.to_numpy()
            if not isinstance(p_values, np.ndarray):
                p_values = p_values.to_numpy()

            significance = [self.get_significance(p) for p in p_values.flatten()]

            # 准备结果数据
            result_data = []
            # 限制循环范围，防止索引越界
            max_class_idx = min(len(unique_classes), len(intercept), z_values.shape[0], p_values.shape[0])
            for i in range(max_class_idx):
                cls = unique_classes[i]
                row = {
                    "Model": f"Class {cls} vs others",
                    "Intercept": intercept[i],
                    # 检查列索引是否有效
                    "Intercept_z": z_values[i, 0] if z_values.shape[1] > 0 else np.nan,
                    "Intercept_p": p_values[i, 0] if p_values.shape[1] > 0 else np.nan,
                    "Intercept_sig": self.get_significance(p_values[i, 0]) if (
                            p_values.shape[1] > 0 and i < len(p_values)) else ""
                }
                # 添加特征系数，检查索引范围
                max_feat_idx = min(len(feature_names), coefficients.shape[1])
                for j in range(max_feat_idx):
                    feature = feature_names[j]
                    row[f"Coeff_{feature}"] = coefficients[i, j]
                    # 检查z_values和p_values的列索引是否有效
                    if z_values.shape[1] > j + 1:
                        row[f"z_{feature}"] = z_values[i, j + 1]
                    else:
                        row[f"z_{feature}"] = np.nan

                    if p_values.shape[1] > j + 1:
                        row[f"p_{feature}"] = p_values[i, j + 1]
                        row[f"sig_{feature}"] = self.get_significance(p_values[i, j + 1])
                    else:
                        row[f"p_{feature}"] = np.nan
                        row[f"sig_{feature}"] = ""
                result_data.append(row)

            # 创建结果DataFrame
            result_df = pd.DataFrame(result_data)

            # 模型性能指标，添加使用的求解器和正则化参数
            metrics_df = pd.DataFrame([{
                languages[self.current_language]['performance_metrics'][
                    "Model"]: "性能指标" if self.current_language == 'zh' else "Performance Metrics",
                languages[self.current_language]['performance_metrics']["Accuracy"]: accuracy,
                languages[self.current_language]['performance_metrics']["F1 Score (Macro)"]: f1_macro,
                languages[self.current_language]['performance_metrics']["F1 Score (Weighted)"]: f1_weighted,
                languages[self.current_language]['performance_metrics']["ROC-AUC (Macro)"]: roc_auc_macro,
                languages[self.current_language]['performance_metrics']["ROC-AUC (Weighted)"]: roc_auc_weighted,
                languages[self.current_language]['performance_metrics']["Solver Used"]: used_solver,
                languages[self.current_language]['performance_metrics']["Regularization (C)"]: used_C,
                languages[self.current_language]['performance_metrics'][
                    "Note"]: "基于测试集计算的指标" if self.current_language == 'zh' else "Metrics calculated on test set"
            }])

            # 生成回归方程
            regression_equations = self.generate_regression_equations(
                unique_classes, intercept, coefficients, feature_names
            )

            # 解释说明
            explanations = languages[self.current_language]['explanation']
            explanation_df = pd.DataFrame([{
                "Model": "Explanation" if self.current_language == 'en' else "解释说明", **explanations
            }])

            # 显示进度提示
            self.result_label.config(text=languages[self.current_language]["progress_generating_results"])
            self.root.update()

            # 分析完成后，提示用户选择保存位置（修改部分）
            self.result_label.config(text=languages[self.current_language]["select_save_location"])
            self.root.update()

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建Word文档
                doc = Document()
                doc.add_heading(languages[self.current_language]["title"], level=1)

                # 添加数据预处理说明
                doc.add_heading("Data Preprocessing" if self.current_language == 'en' else "数据预处理", level=2)
                preprocess_text = (
                    f"{'- 分类列编码: ' if self.current_language == 'zh' else '- Categorical columns encoded: '}"
                    f"{', '.join(categorical_cols) if categorical_cols else ('无' if self.current_language == 'zh' else 'None')}\n"
                    f"{'- 数值列: ' if self.current_language == 'zh' else '- Numeric columns: '}"
                    f"{', '.join(numeric_cols) if numeric_cols else ('无' if self.current_language == 'zh' else 'None')}\n"
                    f"{'- 训练集-测试集划分: 70%-30%' if self.current_language == 'zh' else '- Train-test split: 70%-30%'}"
                )
                doc.add_paragraph(preprocess_text)

                # 添加特征数量限制说明（如果有）
                if original_feature_count > self.MAX_FEATURES:
                    doc.add_paragraph(
                        languages[self.current_language]["feature_limit_warning"].format(
                            original_feature_count, self.MAX_FEATURES)
                    )

                # 添加回归方程
                doc.add_heading(languages[self.current_language]["regression_equations"], level=2)
                for equation in regression_equations:
                    doc.add_paragraph(equation)

                # 添加回归系数表（行列互换版本）
                doc.add_heading("Regression Coefficients" if self.current_language == 'en' else "回归系数", level=2)

                # 构建转置后的表格列（第一列是特征名，后续列是每个模型）
                table_columns = [
                                    "特征/模型" if self.current_language == 'zh' else "Feature/Model"
                                ] + [
                                    f"类别 {cls} 与其他类别对比" if self.current_language == 'zh' else f"Class {cls} vs others"
                                    for cls in unique_classes[:max_class_idx]
                                ]

                # 创建表格
                table = doc.add_table(rows=1, cols=len(table_columns))
                hdr_cells = table.rows[0].cells
                for i, col in enumerate(table_columns):
                    hdr_cells[i].text = col

                # 添加截距行
                intercept_row = table.add_row().cells
                intercept_row[0].text = "截距" if self.current_language == 'zh' else "Intercept"
                for i in range(max_class_idx):
                    intercept_row[
                        i + 1].text = f"{result_df.iloc[i]['Intercept']:.4f} {result_df.iloc[i]['Intercept_sig']}"

                # 添加特征系数行
                for feature in feature_names:
                    feature_row = table.add_row().cells
                    feature_row[0].text = feature
                    for i in range(max_class_idx):
                        coeff_val = result_df.iloc[i].get(f"Coeff_{feature}", "")
                        sig_val = result_df.iloc[i].get(f"sig_{feature}", "")
                        feature_row[i + 1].text = f"{coeff_val:.4f} {sig_val}" if pd.notna(coeff_val) else ""

                # 添加模型性能指标表格
                doc.add_heading("Model Performance" if self.current_language == 'en' else "模型性能", level=2)
                metrics_table = doc.add_table(rows=1, cols=len(metrics_df.columns))
                # 使用多语言表头
                for i, col in enumerate(metrics_df.columns):
                    metrics_table.rows[0].cells[i].text = col
                for _, row in metrics_df.iterrows():
                    metrics_row = metrics_table.add_row().cells
                    for i, col in enumerate(metrics_df.columns):
                        value = row[col]
                        metrics_row[i].text = f"{value:.4f}" if isinstance(value, float) else str(value)

                # 添加混淆矩阵
                doc.add_heading("Confusion Matrix" if self.current_language == 'en' else "混淆矩阵", level=2)
                # 创建混淆矩阵表格
                conf_matrix_table = doc.add_table(rows=len(conf_matrix) + 1, cols=len(conf_matrix[0]) + 1)
                # 添加表头（第一行和第一列表示类别）
                # 设置混淆矩阵表格的行列标题（Actual/Predicted 多语言切换）
                conf_matrix_table.rows[0].cells[
                    0].text = "实际/预测" if self.current_language == 'zh' else "Actual/Predicted"

                for i, cls in enumerate(unique_classes):
                    # 列标题（预测类别）：中文显示“类别 X”，英文显示“Class X”
                    conf_matrix_table.rows[0].cells[
                        i + 1].text = f"类别 {cls}" if self.current_language == 'zh' else f"Class {cls}"
                    # 行标题（实际类别）：同上，保持术语一致
                    conf_matrix_table.rows[i + 1].cells[
                        0].text = f"类别 {cls}" if self.current_language == 'zh' else f"Class {cls}"
                # 填充混淆矩阵数据
                for i in range(len(conf_matrix)):
                    for j in range(len(conf_matrix[i])):
                        conf_matrix_table.rows[i + 1].cells[j + 1].text = str(conf_matrix[i][j])

                # 添加分类报告
                doc.add_heading("Classification Report" if self.current_language == 'en' else "分类报告", level=2)

                # 转换分类报告字典为DataFrame
                report_df = pd.DataFrame(class_report).transpose()

                # 获取当前语言的分类报告翻译
                cr_trans = languages[self.current_language]['classification_report']

                # 创建分类报告表格（多语言表头）
                report_table = doc.add_table(rows=1, cols=len(report_df.columns) + 1)

                # 第一行表头：第一列为"类别"，其余为指标翻译
                report_table.rows[0].cells[0].text = cr_trans['class']
                for i, col in enumerate(report_df.columns):
                    # 匹配指标翻译（如将"precision"转为"精确率"）
                    trans_col = cr_trans.get(col, col)
                    report_table.rows[0].cells[i + 1].text = trans_col

                # 填充分类报告数据（多语言第一列）
                for idx, row in report_df.iterrows():
                    table_row = report_table.add_row().cells

                    # 处理行标题的多语言转换（所有行都适配语言切换）
                    try:
                        # 尝试将索引转换为整数（适用于普通类别行）
                        class_num = int(idx)
                        row_title = f"{cr_trans['class']} {class_num}"
                    except (ValueError, TypeError):
                        # 非数字索引（特殊行）直接使用翻译字典中的对应值
                        row_title = cr_trans.get(idx, str(idx))  # 自动适配宏平均/加权平均/准确率等特殊行

                    table_row[0].text = row_title

                    # 填充后续指标列
                    for i in range(len(row)):
                        table_row[i + 1].text = f"{row.iloc[i]:.4f}" if isinstance(row.iloc[i], float) else str(
                            row.iloc[i])

                # 添加解释说明
                doc.add_heading(
                    languages[self.current_language]['explanation_heading'] if 'explanation_heading' in languages[
                        self.current_language] else ("Explanations" if self.current_language == 'en' else "指标说明"),
                    level=2)
                expl_table = doc.add_table(rows=1, cols=2)
                expl_table.rows[0].cells[0].text = "Metric" if self.current_language == 'en' else "指标"
                expl_table.rows[0].cells[1].text = "Explanation" if self.current_language == 'en' else "说明"

                # 定义指标名的多语言映射
                metric_names = {
                    'zh': {
                        "Coefficients": "回归系数",
                        "Intercept": "截距",
                        "Accuracy": "准确率",
                        "ROC-AUC": "ROC曲线下面积"
                    },
                    'en': {
                        "Coefficients": "Coefficients",
                        "Intercept": "Intercept",
                        "Accuracy": "Accuracy",
                        "ROC-AUC": "ROC-AUC"
                    }
                }

                for metric_key, expl in explanations.items():
                    expl_row = expl_table.add_row().cells
                    # 根据当前语言获取对应的指标名
                    expl_row[0].text = metric_names[self.current_language][metric_key]
                    expl_row[1].text = expl

                # 生成ROC曲线
                save_dir = os.path.dirname(save_path)
                plt.figure(figsize=(10, 6))
                for i in range(n_classes):
                    fpr, tpr, _ = roc_curve(y_test_binarized[:, i], y_pred_proba[:, i])
                    plt.plot(
                        fpr, tpr, lw=2,
                        label=(
                            f'类别 {unique_classes[i]} (AUC = {roc_auc_score(y_test_binarized[:, i], y_pred_proba[:, i]):.2f})'
                            if self.current_language == 'zh'
                            else f'Class {unique_classes[i]} (AUC = {roc_auc_score(y_test_binarized[:, i], y_pred_proba[:, i]):.2f})')
                    )

                plt.plot([0, 1], [0, 1], 'k--', lw=2)
                plt.xlim([0.0, 1.0])
                plt.ylim([0.0, 1.05])
                plt.xlabel('假阳性率' if self.current_language == "zh" else 'False Positive Rate')
                plt.ylabel('真阳性率' if self.current_language == "zh" else 'True Positive Rate')
                plt.title('ROC Curves for Each Class' if self.current_language == 'en' else '各类别的ROC曲线')
                plt.legend(loc="lower right")
                img_name = "multinomial_logit_roc_curves.png"
                img_path = os.path.join(save_dir, img_name)
                plt.savefig(img_path)
                plt.close()

                # 添加ROC曲线到文档
                doc.add_heading("ROC Curves" if self.current_language == 'en' else "ROC曲线", level=2)
                doc.add_picture(img_path, width=Inches(6))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

                # 生成系数热力图
                plt.figure(figsize=(12, 8))
                # 截取有效特征和系数（避免索引越界）
                plot_coefficients = coefficients[:, :len(feature_names)]  # 系数矩阵
                # 绘制热力图
                im = plt.imshow(plot_coefficients, cmap='coolwarm', aspect='auto')
                plt.colorbar(im, label='系数值' if self.current_language == "zh" else 'Coefficient Value')
                plt.yticks(
                    range(len(unique_classes)),
                    [f'{"类别" if self.current_language == "zh" else "Class"} {cls}' for cls in unique_classes]
                )
                plt.xticks(range(len(feature_names)), feature_names, rotation=30, ha='center')
                plt.title('Coefficient Heatmap' if self.current_language == 'en' else '系数热力图')
                plt.tight_layout()
                # 保存图片
                coef_heatmap_name = "multinomial_coefficient_heatmap.png"
                coef_heatmap_path = os.path.join(save_dir, coef_heatmap_name)
                plt.savefig(coef_heatmap_path)
                plt.close()

                # 添加到Word文档
                doc.add_heading("Coefficient Heatmap" if self.current_language == 'en' else "系数热力图", level=2)
                doc.add_picture(coef_heatmap_path, width=Inches(6))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

                # 生成特征重要性排序图（以系数绝对值的平均值为指标）
                if len(feature_names) > 0:
                    # 计算每个特征的平均绝对系数（跨所有类别）
                    feature_importance = np.mean(np.abs(plot_coefficients), axis=0)
                    # 按重要性排序
                    sorted_idx = np.argsort(feature_importance)[::-1]
                    sorted_features = [feature_names[i] for i in sorted_idx]
                    sorted_importance = feature_importance[sorted_idx]

                    plt.figure(figsize=(10, 6))
                    plt.barh(range(len(sorted_features)), sorted_importance, color='skyblue')
                    plt.yticks(range(len(sorted_features)), sorted_features)
                    plt.xlabel('Mean Absolute Coefficient' if self.current_language == 'en' else '平均绝对系数')
                    plt.title('Feature Importance Ranking' if self.current_language == 'en' else '特征重要性排序')
                    plt.gca().invert_yaxis()  # 重要性高的在上方
                    plt.tight_layout()
                    # 保存图片
                    feat_importance_name = "multinomial_feature_importance.png"
                    feat_importance_path = os.path.join(save_dir, feat_importance_name)
                    plt.savefig(feat_importance_path)
                    plt.close()

                    # 添加到Word文档
                    doc.add_heading("Feature Importance" if self.current_language == 'en' else "特征重要性", level=2)
                    doc.add_picture(feat_importance_path, width=Inches(6))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

                # 保存文档
                doc.save(save_path)

                # 生成Excel备用文件
                excel_path = os.path.splitext(save_path)[0] + ".xlsx"
                with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
                    # 保存预处理信息
                    pd.DataFrame({
                        "Categorical Columns": [', '.join(categorical_cols) if categorical_cols else 'None'],
                        "Numeric Columns": [', '.join(numeric_cols) if numeric_cols else 'None'],
                        "Train-test Split": ["70%-30%"],
                        "Feature Limit Note": [
                            languages[self.current_language]["feature_limit_warning"].format(
                                original_feature_count, self.MAX_FEATURES)
                            if original_feature_count > self.MAX_FEATURES else "No limit applied"
                        ]
                    }).to_excel(writer, sheet_name="Preprocessing", index=False)

                    # 保存回归方程
                    pd.DataFrame({
                        "Regression Equations": regression_equations
                    }).to_excel(writer, sheet_name="Equations", index=False)

                    # 保存回归系数
                    result_df.to_excel(writer, sheet_name="Coefficients", index=False)

                    # 保存模型性能
                    metrics_df.to_excel(writer, sheet_name="Performance", index=False)

                    # 保存指标说明
                    explanation_df.to_excel(writer, sheet_name="Explanations", index=False)

                    # 保存混淆矩阵
                    conf_matrix_df = pd.DataFrame(
                        conf_matrix,
                        index=[f"Actual Class {cls}" for cls in unique_classes],
                        columns=[f"Predicted Class {cls}" for cls in unique_classes]
                    )
                    conf_matrix_df.to_excel(writer, sheet_name="Confusion Matrix")

                    # 保存分类报告
                    report_df = pd.DataFrame(class_report).transpose()
                    report_df.to_excel(writer, sheet_name="Classification Report")

                self.result_label.config(
                    text=languages[self.current_language]["analysis_complete"].format(
                        save_path, excel_path)
                )
            else:
                self.result_label.config(text=languages[self.current_language]["no_save_path_selected"])
        except Exception as e:
            self.result_label.config(
                text=languages[self.current_language]["analysis_error"].format(str(e))
                )

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])
        self.open_excel_label.config(text=languages[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=languages[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame,
                                               text=languages[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = MultinomialLogitRegressionApp()
    app.run()

if __name__ == "__main__":
    run_app()
import tkinter as tk
from tkinter import filedialog
import os
import pandas as pd
import numpy as np
from scipy import stats
import matplotlib.pyplot as plt
import pathlib
from statsmodels.stats.multicomp import pairwise_tukeyhsd
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from docx import Document
from docx.shared import Pt
import matplotlib.patches as mpatches


# 设置支持中文的字体
plt.rcParams["font.family"] = ["SimHei", "WenQuanYi Micro Hei", "Heiti TC", "Arial Unicode MS"]
plt.rcParams["axes.unicode_minus"] = False  # 解决负号显示问题

# 定义语言字典，包含所有需要翻译的文本
LANGUAGES = {
    'zh': {
        'title': "事后多重比较",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "请选择文件。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "中/英",
        'open_excel_button_text': "示例数据",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'explanation': {
            "Tukey's HSD": "用于方差分析后，检验多组均值之间的差异是否显著。",
        },
        'interpretation': {
            "p-adj": "经过校正后的 p 值，小于显著性水平（通常为 0.05）时，拒绝原假设，认为两组均值之间存在显著差异。",
            "meandiff": "两组均值的差值，反映了两组之间的差异大小。",
            "lower": "差异的置信区间下限。",
            "upper": "差异的置信区间上限。",
            "reject": "是否拒绝原假设，True 表示拒绝，认为两组均值有显著差异。"
        },
        'columns': {
            'group1': '组1',
            'group2': '组2',
            'meandiff': '均值差',
            'p-adj': '校正后p值',
            'lower': '下限',
            'upper': '上限',
            'reject': '是否拒绝'
        },
        'explanation_heading': "解释说明",
        'interpretation_heading': "结果解读",
        'statistic': "统计量",
        'mean_difference_ci': "均值差 (95% 置信区间)",
        'tukey_title': "Tukey事后多重比较结果",
        'significant_legend': "p < 0.05 (显著)",
        'not_significant_legend': "p ≥ 0.05 (不显著)",
        'zero_std_message': "以下列的标准差为零，将被移除: {}"
    },
    'en': {
        'title': "Post Hoc Multiple Comparisons",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Chinese/English",
        'open_excel_button_text': "Example data",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'explanation': {
            "Tukey's HSD": "Used after ANOVA to test whether the differences between the means of multiple groups are significant.",
        },
        'interpretation': {
            "p-adj": "The adjusted p-value. When it is less than the significance level (usually 0.05), the null hypothesis is rejected, indicating a significant difference between the means of two groups.",
            "meandiff": "The difference between the means of two groups, reflecting the magnitude of the difference between the two groups.",
            "lower": "The lower limit of the confidence interval for the difference.",
            "upper": "The upper limit of the confidence interval for the difference.",
            "reject": "Whether to reject the null hypothesis. True indicates rejection, suggesting a significant difference between the means of two groups."
        },
        'columns': {
            'group1': 'group1',
            'group2': 'group2',
            'meandiff': 'meandiff',
            'p-adj': 'p-adj',
            'lower': 'lower',
            'upper': 'upper',
            'reject': 'reject'
        },
        'explanation_heading': "Explanation",
        'interpretation_heading': "Interpretation",
        'statistic': "Statistic",
        'mean_difference_ci': "Mean Difference (95% CI)",
        'tukey_title': "Tukey's HSD Post-hoc Comparisons",
        'significant_legend': "p < 0.05 (Significant)",
        'not_significant_legend': "p ≥ 0.05 (Not Significant)",
        'zero_std_message': "The following columns have zero standard deviation and will be removed: {}"
    }
}

class PostHocMultipleComparisonsApp:
    def __init__(self, root=None):
        # 当前语言，默认为中文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data44.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 检查数据是否有标准差为零的列
            std_values = df.iloc[:, 1:].std()
            zero_std_columns = std_values[std_values == 0].index
            if len(zero_std_columns) > 0:
                print(LANGUAGES[self.current_language]['zero_std_message'].format(zero_std_columns))
                df = df.drop(zero_std_columns, axis=1)

            # 假设数据的第一列为分组变量，其余列为观测值
            group = df.iloc[:, 0]
            values = df.iloc[:, 1:].values.flatten()
            groups = np.repeat(group, df.shape[1] - 1)

            # 进行 Tukey's HSD 事后多重比较
            tukey = pairwise_tukeyhsd(endog=values, groups=groups, alpha=0.05)
            tukey_df = pd.DataFrame(data=tukey._results_table.data[1:], columns=tukey._results_table.data[0])

            # 添加解释说明
            explanations = LANGUAGES[self.current_language]['explanation']
            interpretations = LANGUAGES[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(columns=["Tukey's HSD"])
            explanation_df.insert(0, LANGUAGES[self.current_language]['statistic'],
                                 LANGUAGES[self.current_language]['explanation_heading'])

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            col_keys = ["p-adj", "meandiff", "lower", "upper", "reject"]
            interpretation_df = interpretation_df.reindex(columns=col_keys)
            # 使用翻译后的列名
            interpretation_df.columns = [LANGUAGES[self.current_language]['columns'][key] for key in col_keys]
            interpretation_df.insert(0, LANGUAGES[self.current_language]['statistic'],
                                    LANGUAGES[self.current_language]['interpretation_heading'])

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加 Tukey's HSD 结果表格
                table = doc.add_table(rows=tukey_df.shape[0] + 1, cols=tukey_df.shape[1])
                hdr_cells = table.rows[0].cells
                # 获取当前语言的列名翻译
                col_translations = LANGUAGES[self.current_language]['columns']
                # 按原列顺序映射翻译后的列名
                translated_columns = [col_translations[col] for col in tukey_df.columns]
                for col_idx, col_name in enumerate(translated_columns):
                    hdr_cells[col_idx].text = col_name
                for row_idx, row in tukey_df.iterrows():
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx].text = str(value)

                # 添加解释说明
                doc.add_paragraph()
                doc.add_heading(LANGUAGES[self.current_language]['explanation_heading'], level=2)
                for idx, row in explanation_df.iterrows():
                    for col in explanation_df.columns[1:]:
                        doc.add_paragraph(f"{col}: {row[col]}")

                # 添加结果解读
                doc.add_paragraph()
                doc.add_heading(LANGUAGES[self.current_language]['interpretation_heading'], level=2)
                for idx, row in interpretation_df.iterrows():
                    for col in interpretation_df.columns[1:]:
                        doc.add_paragraph(f"{col}: {row[col]}")

                # 保存 Word 文档
                doc.save(save_path)

                # 生成森林图
                desktop_path = pathlib.Path.home() / 'Desktop'
                plot_path = desktop_path / 'tukey_forest_plot.png'

                # 准备数据
                tukey_data = tukey._results_table.data[1:]  # 获取比较结果
                comparisons = [f"{row[0]} vs {row[1]}" for row in tukey_data]  # 组间比较标签
                meandiffs = [float(row[2]) for row in tukey_data]  # 均值差
                lower = [float(row[4]) for row in tukey_data]  # 置信区间下限
                upper = [float(row[5]) for row in tukey_data]  # 置信区间上限
                p_adj = [float(row[3]) for row in tukey_data]  # 校正后p值

                # 设置图形大小
                plt.figure(figsize=(10, len(comparisons) * 0.5 + 1))
                y_pos = np.arange(len(comparisons))

                # 绘制置信区间
                plt.errorbar(meandiffs, y_pos, xerr=[[meandiffs[i] - lower[i] for i in range(len(meandiffs))],
                                                     [upper[i] - meandiffs[i] for i in range(len(meandiffs))]],
                             fmt='o', color='gray', ecolor='lightgray', elinewidth=2, capsize=4, markersize=5)

                # 绘制均值差点（根据p值设置颜色）
                colors = ['red' if p < 0.05 else 'blue' for p in p_adj]
                plt.scatter(meandiffs, y_pos, color=colors, s=50, zorder=3)

                # 添加零线（无差异线）
                plt.axvline(x=0, color='black', linestyle='--', alpha=0.5)

                # 设置Y轴标签（组间比较）
                plt.yticks(y_pos, comparisons, fontsize=9)
                plt.gca().invert_yaxis()  # 翻转Y轴，让第一个比较显示在顶部

                # 设置X轴标签和标题
                plt.xlabel(LANGUAGES[self.current_language]['mean_difference_ci'], fontsize=10)
                plt.title(LANGUAGES[self.current_language]['tukey_title'], fontsize=12, pad=20)

                # 添加网格线
                plt.grid(axis='x', linestyle='--', alpha=0.3)

                # 添加图例
                sig_patch = mpatches.Patch(color='red', label=LANGUAGES[self.current_language]['significant_legend'])
                ns_patch = mpatches.Patch(color='blue', label=LANGUAGES[self.current_language]['not_significant_legend'])
                plt.legend(handles=[sig_patch, ns_patch], loc='upper right', fontsize=8)

                # 调整布局
                plt.tight_layout()

                # 保存图片
                plt.savefig(plot_path, dpi=300, bbox_inches='tight')
                plt.close()

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(
                    save_path) + f"结果图片已保存到 {plot_path}" if self.current_language == 'zh' else \
                             LANGUAGES[self.current_language]['analysis_success'].format(
                    save_path) + f"Result image saved to {plot_path}"
                self.result_label.config(text=result_msg, wraplength=400)
            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(LANGUAGES[self.current_language]["title"])
        self.select_button.config(text=LANGUAGES[self.current_language]["select_button"])
        current_text = self.file_entry.get()
        # 只有当输入框显示的是占位符时才更新
        if current_text in [LANGUAGES['zh']["file_entry_placeholder"], LANGUAGES['en']["file_entry_placeholder"]]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=LANGUAGES[self.current_language]["analyze_button"])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]["switch_language"])
        self.open_excel_label.config(text=LANGUAGES[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["select_button"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["analyze_button"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=LANGUAGES[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=LANGUAGES[self.current_language]["switch_language"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = PostHocMultipleComparisonsApp()
    app.run()


if __name__ == "__main__":
    run_app()
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
from scipy.stats import chi2
from docx import Document
from docx.shared import Inches


# 手动实现Cochran's Q检验
def cochrans_q_test(data):
    k = data.shape[1]  # 样本数量（列数）
    n = data.shape[0]  # 观测数量（行数）

    # 检查样本数量是否满足要求（至少3个）
    if k < 3:
        raise ValueError("Cochran's Q检验需要至少3个样本（列）")

    # 计算每行总和（每个观测对象的阳性结果总数）
    row_sums = data.sum(axis=1)
    # 计算每列总和（每个样本的阳性结果总数）
    col_sums = data.sum(axis=0)
    # 计算总阳性结果数
    total = data.sum()

    # 计算Q统计量分子和分母
    numerator = (k - 1) * (k * np.sum(col_sums ** 2) - total ** 2)
    denominator = k * total - np.sum(row_sums ** 2)

    # 避免除以零
    if denominator == 0:
        raise ValueError("无法计算Q统计量，分母为零（可能所有观测结果完全一致）")

    q_stat = numerator / denominator
    # 计算p值（基于自由度为k-1的卡方分布）
    p_value = 1 - chi2.cdf(q_stat, df=k - 1)

    return q_stat, p_value

# 定义语言字典
LANGUAGES = {
    'zh': {
        'title': "Cochran's Q 检验",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "请选择文件。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "中/英",
        'open_excel_button_text': "示例数据",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'explanation': {
            "Cochran's Q 检验": "用于比较三个或更多相关的二分变量样本的分布是否有显著差异。",
            "样本量": "每个样本中的观测值数量。",
            "阳性比例": "每个样本中阳性结果的比例。"
        },
        'interpretation': {
            "统计量": "Cochran's Q 检验的统计量值，用于衡量样本之间的差异程度。",
            "p值": "p值小于显著性水平（通常为0.05）时，拒绝原假设，认为样本之间存在显著差异；否则，接受原假设，认为样本之间无显著差异。",
            "样本量": "样本量的大小会影响统计检验的功效，较大的样本量通常能提供更准确的结果。",
            "阳性比例": "阳性比例反映了样本中阳性结果的相对数量，可用于比较不同样本的阳性情况。"
        }
    },
    'en': {
        'title': "Cochran's Q Test",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Chinese/English",
        'open_excel_button_text': "Sample Data",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'explanation': {
            "Cochran's Q Test": "Used to compare whether the distributions of three or more related binary variable samples are significantly different.",
            "Sample Size": "The number of observations in each sample.",
            "Positive Proportion": "The proportion of positive results in each sample."
        },
        'interpretation': {
            "Statistic": "The test statistic value of Cochran's Q test, used to measure the degree of difference between samples.",
            "p-value": "When the p-value is less than the significance level (usually 0.05), the null hypothesis is rejected, indicating a significant difference between samples; otherwise, the null hypothesis is accepted, indicating no significant difference.",
            "Sample Size": "The sample size affects the power of the statistical test. A larger sample size usually provides more accurate results.",
            "Positive Proportion": "The positive proportion reflects the relative number of positive results in the sample and can be used to compare the positive situations of different samples."
        }
    }
}


class CochransQTestApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data18.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 检查数据是否为数值类型
            numerical_df = df.select_dtypes(include=[np.number])
            if numerical_df.empty:
                raise ValueError("数据中没有数值列，无法进行 Cochran's Q 检验。")

            # 使用手动实现的Cochran's Q检验
            stat, p_value = cochrans_q_test(numerical_df.values)

            # 计算样本量和阳性比例
            sample_sizes = numerical_df.count()
            positive_proportions = numerical_df.mean()

            # 让用户选择保存路径（修改为Word格式）
            save_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word files", "*.docx")]
            )
            if save_path:
                # 创建Word文档并写入内容
                doc = Document()

                # 添加标题
                doc.add_heading(LANGUAGES[self.current_language]["title"], level=1)

                # 根据当前语言定义文本内容
                if self.current_language == "zh":
                    heading_title = "检验结果"
                    table_headers = ["统计量", "统计量值", "p值"]
                    test_name = "Cochran's Q 检验"
                    sample_size_label = "样本量"
                    positive_ratio_label = "阳性比例"
                    explanation_heading = "解释说明"
                    interpretation_heading = "结果解读"
                else:
                    heading_title = "Test Results"
                    table_headers = ["Statistic", "Statistic Value", "p-value"]
                    test_name = "Cochran's Q Test"
                    sample_size_label = "Sample Size"
                    positive_ratio_label = "Positive Proportion"
                    explanation_heading = "Explanations"
                    interpretation_heading = "Interpretations"

                # 添加检验结果表格
                doc.add_heading(heading_title, level=2)
                table = doc.add_table(rows=4, cols=3)
                table.style = 'Table Grid'

                # 表格标题行
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = table_headers[0]
                hdr_cells[1].text = table_headers[1]
                hdr_cells[2].text = table_headers[2]

                # 填充检验结果
                row1_cells = table.rows[1].cells
                row1_cells[0].text = test_name
                row1_cells[1].text = f"{stat:.4f}"
                row1_cells[2].text = f"{p_value:.4f}"

                row2_cells = table.rows[2].cells
                row2_cells[0].text = sample_size_label
                row2_cells[1].text = str(sample_sizes.to_dict())

                row3_cells = table.rows[3].cells
                row3_cells[0].text = positive_ratio_label
                row3_cells[1].text = str({k: f"{v:.4f}" for k, v in positive_proportions.to_dict().items()})

                # 添加解释说明
                doc.add_heading(explanation_heading, level=2)
                explanations = LANGUAGES[self.current_language]['explanation']
                for key, value in explanations.items():
                    doc.add_paragraph(f"{key}: {value}")

                # 添加结果解读
                doc.add_heading(interpretation_heading, level=2)
                interpretations = LANGUAGES[self.current_language]['interpretation']
                for key, value in interpretations.items():
                    doc.add_paragraph(f"{key}: {value}")

                # 保存Word文档
                doc.save(save_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(
                    save_path)
                self.result_label.config(text=result_msg, wraplength=400)
            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except ValueError as ve:
            self.result_label.config(text=f"数据错误: {ve}")
        except FileNotFoundError:
            self.result_label.config(text="未找到文件，请重新选择。")
        except Exception as e:
            self.result_label.config(text=f"分析文件时出现未知错误: {e}")

    def switch_language(self, event=None):
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'
        self.root.title(LANGUAGES[self.current_language]['title'])
        self.select_button.config(text=LANGUAGES[self.current_language]['select_button'])
        self.analyze_button.config(text=LANGUAGES[self.current_language]['analyze_button'])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]['switch_language'])
        self.open_excel_label.config(text=LANGUAGES[self.current_language]["open_excel_button_text"])
        # 切换语言时更新提示信息
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
        self.file_entry.config(foreground='gray')

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=LANGUAGES[self.current_language]['select_button'],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=LANGUAGES[self.current_language]['analyze_button'],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=LANGUAGES[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=LANGUAGES[self.current_language]['switch_language'],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = CochransQTestApp()
    app.run()


if __name__ == "__main__":
    run_app()
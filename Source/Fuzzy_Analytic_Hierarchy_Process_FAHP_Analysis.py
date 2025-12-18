import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典 - 所有显示文本都在这里
LANGUAGES = {
    'zh': {
        'title': "模糊层次分析法",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "请选择文件。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "中/英",
        'open_excel_button_text': "示例数据",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'terms': {
            "fuzzy_weights": "模糊权重",
            "defuzzified_weights": "去模糊化权重",
            "CI": "一致性指标 CI",
            "RI": "随机一致性指标 RI",
            "CR": "一致性比率 CR"
        },
        'explanation': {
            "fuzzy_weights": "反映各因素相对重要性的模糊权重",
            "defuzzified_weights": "将模糊权重转换为精确值的结果",
            "CI": "衡量判断矩阵一致性的指标",
            "RI": "根据矩阵阶数确定的随机一致性指标",
            "CR": "CI 与 RI 的比值，判断矩阵是否具有满意一致性"
        },
        'interpretation': {
            "fuzzy_weights": "三角模糊数表示的各因素重要性范围",
            "defuzzified_weights": "值越大，对应因素越重要",
            "CI": "CI 值越小，矩阵一致性越好",
            "RI": "不同阶数矩阵有对应标准值",
            "CR": "CR < 0.1 时，矩阵具有满意一致性，结果可信"
        },
        'non_square_matrix': "输入矩阵必须是方阵（行数=列数）",
        'not_reciprocal_matrix': "输入矩阵不是正互反矩阵，请检查对角线元素是否为1且a_ij=1/a_ji",
        'matrix_order_error': "矩阵阶数必须大于0",
        'consistency_check_na': "1-2阶矩阵无需进行一致性检验",
        'data_empty_error': "有效数据区域为空，请检查文件内容",
        'data_conversion_error': "数据转换错误，请确保数据区域只包含数值或分数(如1/3)",
        'matrix_cleanup_info': "已自动清理数据：移除了{}行空行和{}列空列",
        'analysis_results': "FAHP 分析结果",
        'explanation_title': "解释说明",
        'interpretation_title': "结果解读",
        'statistics': "统计量",
        'statistics_value': "统计量值",
        'factor': "因素",
        'weight': "权重",
        'defuzzified_weights_chart': "去模糊化权重柱状图",
        'factor_default': "因素{}",
        'unsupported_matrix_order': "不支持大于15阶的矩阵"
    },
    'en': {
        'title': "Fuzzy Analytic Hierarchy Process",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Chinese/English",
        'open_excel_button_text': "Example data",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'terms': {
            "fuzzy_weights": "Fuzzy Weights",
            "defuzzified_weights": "Defuzzified Weights",
            "CI": "Consistency Index (CI)",
            "RI": "Random Consistency Index (RI)",
            "CR": "Consistency Ratio (CR)"
        },
        'explanation': {
            "fuzzy_weights": "Fuzzy weights reflecting the relative importance of each factor",
            "defuzzified_weights": "Defuzzified weights converted from fuzzy weights",
            "CI": "Indicator to measure the consistency of the judgment matrix",
            "RI": "Random consistency index determined by matrix order",
            "CR": "Ratio of CI to RI to judge satisfactory consistency"
        },
        'interpretation': {
            "fuzzy_weights": "Range of factor importance expressed by triangular fuzzy numbers",
            "defuzzified_weights": "Larger values indicate more important factors",
            "CI": "Smaller CI values indicate better matrix consistency",
            "RI": "Standard values for matrices of different orders",
            "CR": "When CR < 0.1, the matrix has satisfactory consistency"
        },
        'non_square_matrix': "Input matrix must be square (rows = columns)",
        'not_reciprocal_matrix': "Input matrix is not a positive reciprocal matrix. Check that diagonal elements are 1 and a_ij=1/a_ji",
        'matrix_order_error': "Matrix order must be greater than 0",
        'consistency_check_na': "Consistency check is not required for 1-2 order matrices",
        'data_empty_error': "Valid data area is empty, please check file content",
        'data_conversion_error': "Data conversion error, please ensure data area contains only numbers or fractions (e.g., 1/3)",
        'matrix_cleanup_info': "Automatically cleaned data: removed {} empty rows and {} empty columns",
        'analysis_results': "FAHP Analysis Results",
        'explanation_title': "Explanations",
        'interpretation_title': "Result Interpretation",
        'statistics': "Statistic",
        'statistics_value': "Statistic Value",
        'factor': "Factors",
        'weight': "Weights",
        'defuzzified_weights_chart': "Bar Chart of Defuzzified Weights",
        'factor_default': "Factor{}",
        'unsupported_matrix_order': "Does not support matrices larger than 15th order"
    }
}

# 随机一致性指标 RI 表 (扩展到15阶)
RI_TABLE = {
    1: 0, 2: 0, 3: 0.58, 4: 0.90, 5: 1.12, 6: 1.24, 7: 1.32, 8: 1.41, 9: 1.45,
    10: 1.49, 11: 1.51, 12: 1.54, 13: 1.56, 14: 1.58, 15: 1.59
}


class FuzzyAnalyticHierarchyProcessFAHPApp:
    def __init__(self, root=None):
        self.current_language = "en"  # 默认中文

        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data46.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == LANGUAGES[self.current_language]['file_entry_placeholder']:
            self.file_entry.delete(0, tk.END)
            self.file_entry.configure(style="TEntry")

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
            self.file_entry.configure(style="Gray.TEntry")

    def parse_fraction(self, value):
        """将分数字符串（如"1/3"）转换为浮点数"""
        if isinstance(value, str):
            # 去除所有空格
            value = value.strip().replace(" ", "")
            if '/' in value:
                parts = value.split('/')
                if len(parts) == 2:
                    try:
                        numerator = float(parts[0])
                        denominator = float(parts[1])
                        if denominator != 0:
                            return numerator / denominator
                    except:
                        pass
            # 尝试直接转换为浮点数
            try:
                return float(value)
            except:
                return None
        elif isinstance(value, (int, float)):
            return float(value)
        return None

    def validate_matrix(self, data):
        """验证矩阵是否为方阵和正互反矩阵"""
        # 检查是否为方阵
        if data.shape[0] != data.shape[1]:
            raise ValueError(LANGUAGES[self.current_language]['non_square_matrix'])

        n = data.shape[0]
        if n <= 0:
            raise ValueError(LANGUAGES[self.current_language]['matrix_order_error'])

        # 检查是否为正互反矩阵 (允许微小浮点误差)
        for i in range(n):
            for j in range(n):
                # 对角线元素必须为1
                if i == j and not np.isclose(data[i, j], 1.0, atol=1e-6):
                    raise ValueError(
                        f"{LANGUAGES[self.current_language]['not_reciprocal_matrix']}\n问题位置: 第{i + 1}行第{j + 1}列 (应为1)")
                # 互反性检查 a_ij = 1/a_ji
                if i != j and not np.isclose(data[i, j], 1.0 / data[j, i], atol=1e-6):
                    raise ValueError(
                        f"{LANGUAGES[self.current_language]['not_reciprocal_matrix']}\n问题位置: 第{i + 1}行第{j + 1}列与第{j + 1}行第{i + 1}列 (应互为倒数)")

        return True

    def convert_to_triangular_fuzzy(self, data):
        """将传统AHP标度转换为三角模糊数"""
        n = data.shape[0]
        fuzzy_matrix = np.zeros((n, n, 3))

        for i in range(n):
            for j in range(n):
                if i == j:
                    fuzzy_matrix[i, j] = [1, 1, 1]
                else:
                    val = data[i, j]
                    if val < 1:
                        inv_val = 1 / val
                        l, m, u = self._get_fuzzy_scale(inv_val)
                        fuzzy_matrix[i, j] = [1 / u, 1 / m, 1 / l]
                    else:
                        fuzzy_matrix[i, j] = self._get_fuzzy_scale(val)
        return fuzzy_matrix

    def _get_fuzzy_scale(self, value):
        """根据AHP标度获取对应的三角模糊数"""
        scales = {
            1: [1, 1, 1],
            2: [1, 2, 3],
            3: [2, 3, 4],
            4: [3, 4, 5],
            5: [4, 5, 6],
            6: [5, 6, 7],
            7: [6, 7, 8],
            8: [7, 8, 9],
            9: [8, 9, 9]
        }
        return scales.get(int(round(value)), [1, 1, 1])

    def calculate_fuzzy_weights(self, fuzzy_matrix):
        """计算模糊权重"""
        n = fuzzy_matrix.shape[0]

        fuzzy_geometric = np.zeros((n, 3))
        for i in range(n):
            product_l = 1.0
            product_m = 1.0
            product_u = 1.0
            for j in range(n):
                product_l *= fuzzy_matrix[i, j, 0]
                product_m *= fuzzy_matrix[i, j, 1]
                product_u *= fuzzy_matrix[i, j, 2]

            fuzzy_geometric[i, 0] = product_l **(1 / n)
            fuzzy_geometric[i, 1] = product_m** (1 / n)
            fuzzy_geometric[i, 2] = product_u **(1 / n)

        sum_l = np.sum(fuzzy_geometric[:, 0])
        sum_m = np.sum(fuzzy_geometric[:, 1])
        sum_u = np.sum(fuzzy_geometric[:, 2])

        fuzzy_weights = np.zeros((n, 3))
        for i in range(n):
            fuzzy_weights[i, 0] = fuzzy_geometric[i, 0] / sum_u
            fuzzy_weights[i, 1] = fuzzy_geometric[i, 1] / sum_m
            fuzzy_weights[i, 2] = fuzzy_geometric[i, 2] / sum_l

        return fuzzy_weights

    def defuzzify(self, fuzzy_weights):
        """使用重心法去模糊化"""
        defuzzified = []
        for w in fuzzy_weights:
            l, m, u = w
            defuzzified_value = (l + 2 * m + u) / 4
            defuzzified.append(defuzzified_value)

        total = sum(defuzzified)
        return [v / total for v in defuzzified]

    def power_method(self, matrix):
        """幂法求最大特征值和特征向量"""
        n = matrix.shape[0]
        v = np.ones(n)
        v = v / np.linalg.norm(v)

        max_iter = 1000
        tol = 1e-8
        for _ in range(max_iter):
            v_new = np.dot(matrix, v)
            v_new = v_new / np.linalg.norm(v_new)

            if np.linalg.norm(v_new - v) < tol:
                v = v_new
                break
            v = v_new

        max_eigenvalue = np.dot(v.T, np.dot(matrix, v)) / np.dot(v.T, v)
        return max_eigenvalue, v / np.sum(v)

    def fahp_analysis(self, data):
        """执行FAHP分析"""
        self.validate_matrix(data)
        n = data.shape[0]

        fuzzy_matrix = self.convert_to_triangular_fuzzy(data)
        fuzzy_weights = self.calculate_fuzzy_weights(fuzzy_matrix)
        defuzzified_weights = self.defuzzify(fuzzy_weights)

        CI = None
        RI = None
        CR = None
        if n >= 3:
            max_eigenvalue, _ = self.power_method(data)
            CI = (max_eigenvalue - n) / (n - 1)
            RI = RI_TABLE.get(n, None)
            if RI is None:
                raise ValueError(LANGUAGES[self.current_language]['unsupported_matrix_order'])
            CR = CI / RI if RI != 0 else None

        return fuzzy_weights, defuzzified_weights, CI, RI, CR

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]['file_entry_placeholder']:
            file_path = ""
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 读取Excel数据，保留所有行和列
            df = pd.read_excel(file_path, header=None)

            # 清理空行和空列
            initial_rows, initial_cols = df.shape
            df = df.dropna(how='all')  # 删除全为空的行
            df = df.dropna(axis=1, how='all')  # 删除全为空的列
            cleaned_rows = initial_rows - df.shape[0]
            cleaned_cols = initial_cols - df.shape[1]

            # 提取数据区域（跳过第一行和第一列作为表头）
            if df.shape[0] < 2 or df.shape[1] < 2:
                raise ValueError(LANGUAGES[self.current_language]['data_empty_error'])

            # 提取表头信息（第一行从第二列开始，作为因素名称）
            headers = df.iloc[0, 1:].tolist()
            # 确保表头数量与数据列数一致
            if len(headers) != df.shape[1] - 1:
                # 如果表头不完整，使用默认编号
                headers = [LANGUAGES[self.current_language]['factor_default'].format(i+1)
                          for i in range(df.shape[1] - 1)]

            data_df = df.iloc[1:, 1:]  # 从第二行第二列开始取数据

            # 转换分数和清理空格
            data_numeric = []
            for _, row in data_df.iterrows():
                converted_row = []
                for value in row:
                    parsed = self.parse_fraction(value)
                    if parsed is None:
                        raise ValueError(
                            f"{LANGUAGES[self.current_language]['data_conversion_error']}\n问题值: {value}")
                    converted_row.append(parsed)
                data_numeric.append(converted_row)

            data = np.array(data_numeric, dtype=float)

            # 显示清理信息
            if cleaned_rows > 0 or cleaned_cols > 0:
                self.result_label.config(
                    text=LANGUAGES[self.current_language]['matrix_cleanup_info'].format(cleaned_rows, cleaned_cols)
                )
                self.root.update()

            # 执行FAHP分析
            fuzzy_weights, defuzzified_weights, CI, RI, CR = self.fahp_analysis(data)

            # 整理结果数据（确保表头与结果一一对应）
            data_rows = [
                [LANGUAGES[self.current_language]['terms']['fuzzy_weights'],
                 [f"[{w[0]:.4f}, {w[1]:.4f}, {w[2]:.4f}]" for w in fuzzy_weights]],
                [LANGUAGES[self.current_language]['terms']['defuzzified_weights'],
                 [f"{w:.6f}" for w in defuzzified_weights]],
                [LANGUAGES[self.current_language]['terms']['CI'],
                 f"{CI:.6f}" if CI is not None else LANGUAGES[self.current_language]['consistency_check_na']],
                [LANGUAGES[self.current_language]['terms']['RI'],
                 f"{RI:.6f}" if RI is not None else "-"],
                [LANGUAGES[self.current_language]['terms']['CR'],
                 f"{CR:.6f}" if CR is not None else "-"]
            ]
            headers_table = [
                LANGUAGES[self.current_language]['statistics'],
                LANGUAGES[self.current_language]['statistics_value']
            ]

            # 获取解释说明
            explanations = LANGUAGES[self.current_language]['explanation']
            interpretations = LANGUAGES[self.current_language]['interpretation']

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建Word文档
                doc = Document()
                doc.add_heading(LANGUAGES[self.current_language]['analysis_results'], 0)

                # 添加分析结果表格（确保表头与内容严格对应）
                table = doc.add_table(rows=1, cols=len(headers_table))
                hdr_cells = table.rows[0].cells
                for col, header in enumerate(headers_table):
                    hdr_cells[col].text = header
                for row_data in data_rows:
                    row_cells = table.add_row().cells
                    row_cells[0].text = row_data[0]
                    row_cells[1].text = ", ".join([f"{h}: {v}" for h, v in zip(headers, row_data[1])]) if isinstance(row_data[1], list) else str(row_data[1])

                # 添加解释说明
                doc.add_heading(LANGUAGES[self.current_language]['explanation_title'], level=1)
                for key, value in explanations.items():
                    doc.add_paragraph(f"{LANGUAGES[self.current_language]['terms'][key]}: {value}")

                # 添加结果解读
                doc.add_heading(LANGUAGES[self.current_language]['interpretation_title'], level=1)
                for key, value in interpretations.items():
                    doc.add_paragraph(f"{LANGUAGES[self.current_language]['terms'][key]}: {value}")

                # 生成权重柱状图，使用表头作为横轴标签
                fig, ax = plt.subplots()
                ax.bar(range(len(defuzzified_weights)), defuzzified_weights)
                ax.set_title(LANGUAGES[self.current_language]['defuzzified_weights_chart'])
                ax.set_xlabel(LANGUAGES[self.current_language]['factor'])
                ax.set_ylabel(LANGUAGES[self.current_language]['weight'])
                # 设置横轴标签为表头信息
                ax.set_xticks(range(len(headers)))
                ax.set_xticklabels(headers, rotation=0, ha='right')  # 旋转标签避免重叠
                plt.tight_layout()  # 调整布局

                # 保存图片
                img_path = os.path.splitext(save_path)[0] + '_weights.png'
                plt.savefig(img_path, bbox_inches='tight')
                plt.close()

                # 插入图片到Word
                doc.add_heading(LANGUAGES[self.current_language]['defuzzified_weights_chart'], level=1)
                doc.add_picture(img_path, width=Inches(6))

                # 保存文档
                doc.save(save_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(save_path)
                self.result_label.config(text=result_msg, wraplength=400)

            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)),
                                     wraplength=400)

    def switch_language(self, event=None):
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'
        self.root.title(LANGUAGES[self.current_language]['title'])
        self.select_button.config(text=LANGUAGES[self.current_language]['select_button'])
        self.analyze_button.config(text=LANGUAGES[self.current_language]['analyze_button'])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]['switch_language'])
        self.open_excel_label.config(text=LANGUAGES[self.current_language]["open_excel_button_text"])
        # 更新输入框提示
        current_text = self.file_entry.get()
        placeholder = LANGUAGES[self.current_language]['file_entry_placeholder']
        if current_text == LANGUAGES['zh' if self.current_language == 'en' else 'en']['file_entry_placeholder']:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, placeholder)
            self.file_entry.configure(style="Gray.TEntry")

    def create_ui(self):
        # 窗口设置
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建主框架并居中
        main_frame = ttk.Frame(self.root)
        main_frame.place(relx=0.5, rely=0.5, anchor="center")  # 居中放置主框架

        # 自定义样式
        style = ttk.Style()
        style.configure("Gray.TEntry", foreground="gray")

        # 选择文件按钮
        self.select_button = ttk.Button(
            main_frame,
            text=LANGUAGES[self.current_language]['select_button'],
            command=self.select_file,
            bootstyle=PRIMARY
        )
        self.select_button.pack(pady=10)

        # 文件路径输入框
        self.file_entry = ttk.Entry(main_frame, width=60, style="Gray.TEntry")
        self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
        self.file_entry.pack(pady=5)
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)

        # 分析按钮
        self.analyze_button = ttk.Button(
            main_frame,
            text=LANGUAGES[self.current_language]['analyze_button'],
            command=self.analyze_file,
            bootstyle=SUCCESS
        )
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            main_frame,
            text=LANGUAGES[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 切换语言标签（作为按钮使用）
        self.switch_language_label = ttk.Label(
            main_frame,
            text=LANGUAGES[self.current_language]['switch_language'],
            foreground="gray",
            cursor="hand2"
        )
        self.switch_language_label.pack(pady=5)
        self.switch_language_label.bind("<Button-1>", self.switch_language)

        # 结果显示标签
        self.result_label = ttk.Label(main_frame, text="", wraplength=400)
        self.result_label.pack(pady=10)


if __name__ == "__main__":
    app = FuzzyAnalyticHierarchyProcessFAHPApp()
    app.root.mainloop()
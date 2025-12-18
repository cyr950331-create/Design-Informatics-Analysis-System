import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import os
import numpy as np
import pandas as pd
from tkinter import filedialog, simpledialog, Checkbutton, IntVar
import tkinter as tk
import matplotlib.pyplot as plt
from sklearn.cluster import KMeans, AgglomerativeClustering
from sklearn.preprocessing import StandardScaler
from sklearn.metrics import silhouette_score
from docx import Document
from docx.shared import Inches
import seaborn as sns
from scipy import stats

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典，所有键使用英文
LANGUAGES = {
    'zh': {
        'title': "K均值聚类分析",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'auto_cluster': "自动确定最佳聚类数",
        'file_not_found': "请选择文件。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "中/英",
        'open_excel_button_text': "示例数据",
        'file_entry_placeholder': "请输入待分析文件的完整路径",
        'cluster_count': "请输入聚类数量（1-10）",
        'invalid_cluster': "聚类数量必须是1-10之间的整数",
        'feature_select': "请选择可视化特征（用逗号分隔列索引，如1,2,3）",
        'invalid_feature': "无效的特征选择，请检查输入",
        'optimal_clusters': "自动计算的最佳聚类数为: {}",
        'sample_index': "样本索引",
        'cluster_label_column': "聚类标签",
        'explanation': {
            "cluster_labels": "每个样本所属的聚类类别",
            "cluster_centers": "每个聚类的中心位置（标准化后的值）",
            "parallel_coordinates": "展示不同聚类在多个特征上的分布",
            "radar_chart": "展示各聚类中心的特征分布（标准化后的值）",
            "heatmap": "展示特征间的相关性",
            "dendrogram": "展示样本的层次聚类关系（与K-Means结果无直接关联）"
        },
        'interpretation': {
            "cluster_labels": "可用于区分不同样本所属的类别",
            "cluster_centers": "代表每个聚类的典型特征，已标准化处理",
            "parallel_coordinates": "直观比较不同聚类在多个特征上的差异",
            "radar_chart": "展示各聚类中心在不同特征上的相对强弱，基于标准化值",
            "heatmap": "显示特征之间的相关性强度",
            "dendrogram": "展示样本间的相似度和层次结构，与K-Means结果独立",
            "outlier_note": "注：聚类结果对异常值敏感，已自动移除Z-score>3的极端值"
        },
        'heading_cluster_labels': "聚类标签",
        'heading_cluster_centers': "聚类中心",
        'heading_explanation': "解释说明",
        'heading_interpretation': "结果解读",
        'heading_parallel_coordinates': "平行坐标图",
        'heading_radar_chart': "雷达图",
        'heading_heatmap': "特征相关性热图",
        'heading_dendrogram': "聚类树状图",
        'cluster_id': "聚类ID",
        'statistics_explanation': "统计量_解释说明",
        'statistics_interpretation': "统计量_结果解读",
        'radar_chart_title': "雷达图（原始特征值）",
        'parallel_coordinates_title': "平行坐标图",
        'heatmap_title': "特征相关性热图",
        'dendrogram_title': "聚类树状图（层次关系）",
        'cluster_label': "聚类 {}"
    },
    'en': {
        'title': "K Means Clustering",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'auto_cluster': "Auto-determine optimal clusters",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Chinese/English",
        'open_excel_button_text': "Example data",
        'file_entry_placeholder': "Please enter the full path of the file to be analyzed",
        'cluster_count': "Please enter the number of clusters (1-10)",
        'invalid_cluster': "Number of clusters must be an integer between 1-10",
        'feature_select': "Please select visualization features (column indices separated by commas, e.g. 1,2,3)",
        'invalid_feature': "Invalid feature selection, please check input",
        'optimal_clusters': "Automatically calculated optimal clusters: {}",
        'sample_index': "Sample Index",
        'cluster_label_column': "Cluster Label",
        'explanation': {
            "cluster_labels": "The cluster label to which each sample belongs",
            "cluster_centers": "The center position of each cluster (standardized values)",
            "parallel_coordinates": "Show distribution of different clusters across multiple features",
            "radar_chart": "Show feature distribution of each cluster center (standardized values)",
            "heatmap": "Show correlation between features",
            "dendrogram": "Show hierarchical clustering relationship (not directly related to K-Means results)"
        },
        'interpretation': {
            "cluster_labels": "Can be used to distinguish the categories to which different samples belong",
            "cluster_centers": "Represents the typical characteristics of each cluster, standardized",
            "parallel_coordinates": "Intuitively compare differences between clusters across multiple features",
            "radar_chart": "Show relative strength of each cluster center across features, based on standardized values",
            "heatmap": "Show strength of correlation between features",
            "dendrogram": "Show similarity and hierarchy between samples, independent of K-Means results",
            "outlier_note": "Note: Clustering results are sensitive to outliers. Extreme values with Z-score>3 have been removed"
        },
        'heading_cluster_labels': "Cluster Labels",
        'heading_cluster_centers': "Cluster Centers",
        'heading_explanation': "Explanation",
        'heading_interpretation': "Interpretation",
        'heading_parallel_coordinates': "Parallel Coordinates Plot",
        'heading_radar_chart': "Radar Chart",
        'heading_heatmap': "Feature Correlation Heatmap",
        'heading_dendrogram': "Clustering Dendrogram",
        'cluster_id': "Cluster ID",
        'statistics_explanation': "Statistics_Explanation",
        'statistics_interpretation': "Statistics_Interpretation",
        'radar_chart_title': "Radar Chart (Original Feature Values)",
        'parallel_coordinates_title': "Parallel Coordinates Plot",
        'heatmap_title': "Feature Correlation Heatmap",
        'dendrogram_title': "Clustering Dendrogram (Hierarchical Relationship)",
        'cluster_label': "Cluster {}"
    }
}


class KMeansApp:
    def __init__(self, root=None):
        # 先初始化主窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES["en"]["title"])  # 先使用默认语言
        else:
            self.root = root
            self.root.title(LANGUAGES["en"]["title"])

        # 在主窗口初始化之后再创建IntVar变量
        self.current_language = "en"
        self.auto_cluster_var = IntVar()  # 现在主窗口已存在，可以安全创建变量

        # 更新标题为当前语言
        self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data37.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def clustering_analysis(self, data, n_clusters=3):
        """进行聚类分析"""
        kmeans = KMeans(n_clusters=n_clusters, random_state=42)
        kmeans.fit(data)
        labels = kmeans.labels_
        centers = kmeans.cluster_centers_
        return labels, centers

    def calculate_optimal_clusters(self, data, max_clusters=10):
        """使用肘部法和轮廓系数确定最佳聚类数"""
        distortions = []
        silhouette_scores = []

        for k in range(2, max_clusters + 1):  # 轮廓系数至少需要2个聚类
            kmeans = KMeans(n_clusters=k, random_state=42)
            labels = kmeans.fit_predict(data)
            distortions.append(kmeans.inertia_)
            silhouette_scores.append(silhouette_score(data, labels))

        # 找到轮廓系数最大的聚类数
        optimal_k = np.argmax(silhouette_scores) + 2  # +2因为从k=2开始

        # 如果只有1个特征或样本太少，返回默认值
        if data.shape[1] == 1 or data.shape[0] < 3:
            return min(3, data.shape[0])

        return optimal_k

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return

        try:
            # 读取数据（默认第一行为表头）
            file_ext = os.path.splitext(file_path)[1].lower()
            if file_ext in ['.xlsx', '.xls']:
                df = pd.read_excel(file_path)
            elif file_ext == '.csv':
                df = pd.read_csv(file_path)
            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
                return

            # 数据预处理
            # 处理缺失值
            df = df.dropna()
            # 检查是否有非数值列
            numeric_df = df.select_dtypes(include=[np.number])
            if numeric_df.empty:
                raise ValueError("No numeric columns found in the data")

            # 异常值处理（Z-score法）
            z_scores = stats.zscore(numeric_df)
            abs_z_scores = np.abs(z_scores)
            filtered_entries = (abs_z_scores < 3).all(axis=1)
            numeric_df = numeric_df[filtered_entries]

            # 数据标准化
            scaler = StandardScaler()
            scaled_data = scaler.fit_transform(numeric_df)
            original_data = scaled_data  # 使用标准化后的数据进行聚类
            feature_names = numeric_df.columns.tolist()

            # 确定聚类数量
            if self.auto_cluster_var.get() == 1:
                # 自动确定最佳聚类数
                n_clusters = self.calculate_optimal_clusters(original_data)
                # 显示自动计算的聚类数
                Messagebox.show_info(
                    title=LANGUAGES[self.current_language]['title'],
                    message=LANGUAGES[self.current_language]['optimal_clusters'].format(n_clusters)
                )
            else:
                # 获取用户输入的聚类数量
                cluster_input = simpledialog.askstring(
                    title=LANGUAGES[self.current_language]['title'],
                    prompt=LANGUAGES[self.current_language]['cluster_count'],
                    parent=self.root
                )
                if not cluster_input:
                    return
                try:
                    n_clusters = int(cluster_input)
                    if not (1 <= n_clusters <= 10):
                        raise ValueError
                except ValueError:
                    self.result_label.config(text=LANGUAGES[self.current_language]['invalid_cluster'])
                    return

            # 进行聚类分析
            labels, centers = self.clustering_analysis(original_data, n_clusters)

            # 整理数据（改进表格结构）
            # 聚类标签表格
            label_df = pd.DataFrame({
                LANGUAGES[self.current_language]['sample_index']: range(len(labels)),
                LANGUAGES[self.current_language]['cluster_label_column']: labels
            })

            # 聚类中心表格
            center_df = pd.DataFrame(centers, columns=feature_names)
            center_df.index.name = LANGUAGES[self.current_language]['cluster_id']

            # 添加解释说明
            explanations = LANGUAGES[self.current_language]['explanation']
            interpretations = LANGUAGES[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(
                columns=["cluster_labels", "cluster_centers", "parallel_coordinates", "radar_chart", "heatmap", "dendrogram"])
            explanation_df.insert(0, LANGUAGES[self.current_language]['statistics_explanation'],
                                LANGUAGES[self.current_language]['heading_explanation'])

            # 添加分析结果解读，包含异常值说明
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(
                columns=["cluster_labels", "cluster_centers", "parallel_coordinates", "radar_chart", "heatmap", "dendrogram", "outlier_note"])
            interpretation_df.insert(0, LANGUAGES[self.current_language]['statistics_interpretation'],
                                     LANGUAGES[self.current_language]['heading_interpretation'])

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if not save_path:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])
                return

            # 创建 Word 文档
            doc = Document()

            # 添加聚类标签表格
            doc.add_heading(LANGUAGES[self.current_language]['heading_cluster_labels'], 1)
            table = doc.add_table(rows=1, cols=len(label_df.columns))
            hdr_cells = table.rows[0].cells
            for col_idx, col_name in enumerate(label_df.columns):
                hdr_cells[col_idx].text = str(col_name)
            for _, row in label_df.iterrows():
                row_cells = table.add_row().cells
                for col_idx, value in enumerate(row):
                    row_cells[col_idx].text = str(value)

            # 添加聚类中心表格
            doc.add_heading(LANGUAGES[self.current_language]['heading_cluster_centers'], 1)
            # 增加一列用于存放聚类ID
            table = doc.add_table(rows=1, cols=len(center_df.columns) + 1)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = LANGUAGES[self.current_language]['cluster_id']
            for col_idx, col_name in enumerate(center_df.columns):
                hdr_cells[col_idx + 1].text = str(col_name)
            for idx, row in center_df.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(idx)  # 聚类ID
                for col_idx, value in enumerate(row):
                    row_cells[col_idx + 1].text = f"{value:.4f}"

            # 添加解释说明（项目列表）
            doc.add_heading(LANGUAGES[self.current_language]['heading_explanation'], 2)
            explanation_para = doc.add_paragraph()
            for value in explanations.values():
                explanation_para.add_run(f"• {value}\n")

            # 添加结果解读（项目列表）
            doc.add_heading(LANGUAGES[self.current_language]['heading_interpretation'], 2)
            interpretation_para = doc.add_paragraph()
            for value in interpretations.values():
                interpretation_para.add_run(f"• {value}\n")

            # 生成可视化图片
            if original_data.shape[1] >= 2:
                # 处理特征选择
                feat_input = simpledialog.askstring(
                    title=LANGUAGES[self.current_language]['title'],
                    prompt=LANGUAGES[self.current_language]['feature_select'],
                    parent=self.root,
                    initialvalue="1,2,3"
                )
                if not feat_input:
                    max_feats = min(3, original_data.shape[1])
                    feat_indices = list(range(max_feats))
                else:
                    try:
                        # 将用户输入的1开始索引转为0开始（减1处理）
                        feat_indices = [int(x.strip()) - 1 for x in feat_input.split(',')]
                        # 验证索引有效性并添加保护
                        max_valid_index = numeric_df.shape[1] - 1
                        feat_indices = [i for i in feat_indices if 0 <= i <= max_valid_index]
                        if len(feat_indices) < 2:
                            raise ValueError
                    except ValueError:
                        self.result_label.config(text=LANGUAGES[self.current_language]['invalid_feature'])
                        return

                # 平行坐标图
                plt.figure(figsize=(10, 6))
                for i in range(n_clusters):
                    # 只选择feat_indices指定的特征
                    cluster_data = original_data[labels == i][:, feat_indices]
                    plt.plot(cluster_data.T, c=plt.cm.viridis(i / n_clusters), alpha=0.3)
                plt.xticks(range(len(feat_indices)), [feature_names[i] for i in feat_indices])
                plt.title(LANGUAGES[self.current_language]['parallel_coordinates_title'])
                plt.tight_layout()
                parallel_path = os.path.splitext(save_path)[0] + '_parallel_coordinates.png'
                plt.savefig(parallel_path, bbox_inches='tight')
                plt.close()

                # 雷达图
                plt.figure(figsize=(8, 8))
                angles = np.linspace(0, 2 * np.pi, len(feat_indices), endpoint=False).tolist()
                angles += angles[:1]  # 闭合雷达图

                ax = plt.subplot(111, polar=True)
                for i in range(n_clusters):
                    full_center = np.zeros(original_data.shape[1])
                    full_center[feat_indices] = centers[i, feat_indices]
                    original_center = scaler.inverse_transform([full_center])[0]
                    values = original_center[feat_indices].tolist()
                    values += values[:1]
                    ax.plot(angles, values, 'o-', linewidth=2,
                            label=LANGUAGES[self.current_language]['cluster_label'].format(i + 1))
                    ax.fill(angles, values, alpha=0.25)

                ax.set_thetagrids(np.degrees(angles[:-1]), [feature_names[i] for i in feat_indices])
                ax.set_title(LANGUAGES[self.current_language]['radar_chart_title'])
                ax.legend(loc='upper right', bbox_to_anchor=(1.1, 1.1))
                plt.tight_layout()
                radar_path = os.path.splitext(save_path)[0] + '_radar_chart.png'
                plt.savefig(radar_path, bbox_inches='tight')
                plt.close()

                # 热图
                plt.figure(figsize=(12, 10))
                valid_feats = [i for i in feat_indices if i < numeric_df.shape[1]]
                if not valid_feats:
                    valid_feats = list(range(min(3, numeric_df.shape[1])))
                corr_matrix = numeric_df.iloc[:, valid_feats].corr()
                sns.heatmap(corr_matrix, annot=True, cmap='coolwarm', fmt='.2f', linewidths=0.5)
                plt.xticks(rotation=0, ha='center')
                plt.title(LANGUAGES[self.current_language]['heatmap_title'])
                plt.tight_layout()
                heatmap_path = os.path.splitext(save_path)[0] + '_heatmap.png'
                plt.savefig(heatmap_path, bbox_inches='tight', dpi=300)
                plt.close()

                # 聚类树状图
                plt.figure(figsize=(12, 8))
                from scipy.cluster.hierarchy import dendrogram, linkage
                linkage_matrix = linkage(original_data[:, feat_indices], 'ward')
                dendrogram(linkage_matrix, truncate_mode='lastp', p=30, leaf_rotation=0, leaf_font_size=10)
                plt.title(LANGUAGES[self.current_language]['dendrogram_title'])
                plt.tight_layout()
                dendrogram_path = os.path.splitext(save_path)[0] + '_dendrogram.png'
                plt.savefig(dendrogram_path, bbox_inches='tight')
                plt.close()

                # 将图片插入到 Word 文档中
                doc.add_heading(LANGUAGES[self.current_language]['heading_parallel_coordinates'], 2)
                doc.add_picture(parallel_path, width=Inches(6))

                doc.add_heading(LANGUAGES[self.current_language]['heading_radar_chart'], 2)
                doc.add_picture(radar_path, width=Inches(6))

                doc.add_heading(LANGUAGES[self.current_language]['heading_heatmap'], 2)
                doc.add_picture(heatmap_path, width=Inches(6))

                doc.add_heading(LANGUAGES[self.current_language]['heading_dendrogram'], 2)
                doc.add_picture(dendrogram_path, width=Inches(6))

            # 保存 Word 文档
            doc.save(save_path)

            result_msg = LANGUAGES[self.current_language]['analysis_success'].format(save_path)
            self.result_label.config(text=result_msg, wraplength=400)

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'
        self.root.title(LANGUAGES[self.current_language]['title'])
        self.select_button.config(text=LANGUAGES[self.current_language]['select_button'])
        self.analyze_button.config(text=LANGUAGES[self.current_language]['analyze_button'])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]['switch_language'])
        self.auto_cluster_check.config(text=LANGUAGES[self.current_language]['auto_cluster'])
        self.open_excel_label.config(text=LANGUAGES[self.current_language]["open_excel_button_text"])
        # 切换语言时更新提示信息
        current_text = self.file_entry.get()
        if current_text in [LANGUAGES['zh']['file_entry_placeholder'], LANGUAGES['en']['file_entry_placeholder']]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
            self.file_entry.config(foreground='gray')

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=LANGUAGES[self.current_language]['select_button'],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 添加自动确定聚类数的复选框
        self.auto_cluster_check = Checkbutton(
            frame,
            text=LANGUAGES[self.current_language]['auto_cluster'],
            variable=self.auto_cluster_var
        )
        self.auto_cluster_check.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=LANGUAGES[self.current_language]['analyze_button'],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=LANGUAGES[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=LANGUAGES[self.current_language]['switch_language'],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        self.root.mainloop()


def run_app():
    app = KMeansApp()
    app.run()


if __name__ == "__main__":
    run_app()
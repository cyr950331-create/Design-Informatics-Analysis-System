import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog, simpledialog, Checkbutton, IntVar
import tkinter as tk
import matplotlib.pyplot as plt
import pathlib
from factor_analyzer import FactorAnalyzer
from factor_analyzer.factor_analyzer import calculate_bartlett_sphericity
from factor_analyzer.factor_analyzer import calculate_kmo
from docx import Document
from docx.shared import Inches
from scipy import stats

# 定义常量（替换魔法数值）
CONDITION_NUMBER_THRESHOLD = 1e10  # 矩阵条件数阈值
EPSILON_FACTOR = 1e-6  # 正则化扰动因子
EIGENVALUE_THRESHOLD = 1e-8  # 特征值最小值
CORRELATION_THRESHOLD = 0.95  # 高度相关变量阈值

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
LANGUAGES = {
    'zh': {
        'title': "因子分析",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "请选择文件。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "中/英",
        'open_excel_button_text': "示例数据",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'explanation': {
            "因子载荷矩阵": "显示每个变量在各个因子上的载荷，反映变量与因子的相关性",
            "共同度": "表示每个变量被因子所解释的方差比例",
            "特征值和方差贡献率": "特征值表示每个因子解释的总方差，方差贡献率表示每个因子解释的方差占总方差的比例",
            "Bartlett球形检验": "检验变量之间是否存在相关性",
            "KMO检验": "衡量变量之间的偏相关性，判断数据是否适合进行因子分析",
            "碎石图": "展示特征值随因子数量的变化情况，帮助确定因子的数量"
        },
        'interpretation': {
            "因子载荷矩阵": "绝对值越大，说明变量与因子的相关性越强",
            "共同度": "值越接近1，说明变量被因子解释的程度越高",
            "特征值和方差贡献率": "特征值大于1的因子通常被保留，方差贡献率越高，说明该因子越重要",
            "Bartlett球形检验": "p值小于0.05时，拒绝原假设，表明变量之间存在相关性，适合进行因子分析",
            "KMO检验": "KMO值大于0.6时，适合进行因子分析",
            "碎石图": "曲线的拐点处通常表示合适的因子数量"
        },
        'rotate_method': "请选择旋转方法:\n1. varimax (正交旋转)\n2. promax (斜交旋转)",
        'factor_num_title': "确定因子数量",
        'factor_num_msg': "建议因子数量: {}\n特征值大于1的因子数: {}\n请输入因子数量:",
        'preprocess_info': "{}",
        'strong_corr': "强相关 (>0.7)",
        'moderate_corr': "中等相关 (0.5-0.7)",
        'weak_corr': "弱相关 (<0.5)",
        'singular_matrix_warning': "检测到数据矩阵接近奇异，已移除高度相关的变量以提高矩阵稳定性",
        'high_corr_removed': "已移除{}个高度相关变量",
        'options_title': "分析选项",
        'rotation_option1': "varimax (正交旋转)",
        'rotation_option2': "promax (斜交旋转)"
    },
    'en': {
        'title': "Factor Analysis",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Chinese/English",
        'open_excel_button_text': "Example data",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'explanation': {
            "Factor Loadings": "Shows the loadings of each variable on each factor, reflecting the correlation between variables and factors",
            "Communalities": "Represents the proportion of variance of each variable explained by the factors",
            "Eigenvalues & Variance Contribution": "The eigenvalue represents the total variance explained by each factor, and the variance contribution rate represents the proportion of variance explained by each factor to the total variance",
            "Bartlett's Test": "Tests whether there is a correlation between variables",
            "KMO Test": "Measures the partial correlation between variables to determine whether the data is suitable for factor analysis",
            "Scree Plot": "Shows the change of eigenvalues with the number of factors, helping to determine the number of factors"
        },
        'interpretation': {
            "Factor Loadings": "The larger the absolute value, the stronger the correlation between the variable and the factor",
            "Communalities": "The closer the value is to 1, the higher the degree to which the variable is explained by the factor",
            "Eigenvalues & Variance Contribution": "Factors with eigenvalues greater than 1 are usually retained. The higher the variance contribution rate, the more important the factor",
            "Bartlett's Test": "When the p-value is less than 0.05, the null hypothesis is rejected, indicating that there is a correlation between variables and factor analysis is suitable",
            "KMO Test": "When the KMO value is greater than 0.6, factor analysis is suitable",
            "Scree Plot": "The inflection point of the curve usually indicates the appropriate number of factors"
        },
        'rotate_method': "Please select rotation method:\n1. varimax (orthogonal)\n2. promax (oblique)",
        'factor_num_title': "Determine Number of Factors",
        'factor_num_msg': "Recommended factors: {}\nEigenvalue >1 factors: {}\nPlease enter number of factors:",
        'preprocess_info': "{}",
        'strong_corr': "Strong correlation (>0.7)",
        'moderate_corr': "Moderate correlation (0.5-0.7)",
        'weak_corr': "Weak correlation (<0.5)",
        'singular_matrix_warning': "Data matrix is nearly singular, removed highly correlated variables to improve stability",
        'high_corr_removed': "Removed {} highly correlated variables",
        'options_title': "Analysis Options",
        'rotation_option1': "varimax (orthogonal)",
        'rotation_option2': "promax (oblique)"
    }
}

class FactorAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data29.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def preprocess_data(self, df):
        """数据预处理：移除非数值列、处理缺失值和异常值，并添加奇异矩阵处理"""
        # 移除非数值列
        numeric_df = df.select_dtypes(include=['number'])
        non_numeric_cols = set(df.columns) - set(numeric_df.columns)
        if non_numeric_cols:
            print(f"移除非数值列: {', '.join(non_numeric_cols)}")

        # 处理缺失值（用列均值填充）
        missing_count = numeric_df.isnull().sum().sum()
        if missing_count > 0:
            numeric_df = numeric_df.fillna(numeric_df.mean())

        # 处理异常值（3σ准则）
        z_scores = np.abs(stats.zscore(numeric_df))
        outlier_count = (z_scores > 3).sum().sum()
        if outlier_count > 0:
            # 用3σ范围内的最大值/最小值替换异常值
            for col in numeric_df.columns:
                col_data = numeric_df[col]
                mean = col_data.mean()
                std = col_data.std()
                upper_limit = mean + 3 * std
                lower_limit = mean - 3 * std
                numeric_df[col] = np.where(col_data > upper_limit, upper_limit, col_data)
                numeric_df[col] = np.where(col_data < lower_limit, lower_limit, numeric_df[col])

        # 检查并处理高度相关的变量以避免奇异矩阵
        corr_removed_msg = ""
        if len(numeric_df.columns) > 1:
            corr_matrix = numeric_df.corr().abs()
            # 找到高度相关的变量对 (相关系数 > 0.95)
            upper = corr_matrix.where(np.triu(np.ones(corr_matrix.shape), k=1).astype(bool))
            to_drop = [column for column in upper.columns if any(upper[column] > 0.95)]

            if to_drop:
                numeric_df = numeric_df.drop(to_drop, axis=1)
                removed_count = len(to_drop)
                corr_removed_msg = LANGUAGES[self.current_language]['high_corr_removed'].format(removed_count)
            else:
                corr_removed_msg = LANGUAGES[self.current_language]['singular_matrix_warning']

        # 添加正则化处理以确保矩阵非奇异
        # 计算协方差矩阵
        cov_matrix = numeric_df.cov()

        # 检查矩阵是否接近奇异（条件数很大）
        cond = np.linalg.cond(cov_matrix)
        if cond > 1e10:  # 条件数阈值，超过此值认为矩阵接近奇异
            # 添加小的扰动到对角线以正则化矩阵
            epsilon = 1e-6 * np.mean(np.diag(cov_matrix))
            cov_matrix += epsilon * np.eye(cov_matrix.shape[0])

            # 使用正则化后的协方差矩阵进行数据转换
            # 计算Cholesky分解
            try:
                L = np.linalg.cholesky(cov_matrix)
                # 对数据进行白化处理
                numeric_df = numeric_df @ np.linalg.inv(L.T)
            except np.linalg.LinAlgError:
                # 如果Cholesky分解失败，使用特征值分解
                eigenvalues, eigenvectors = np.linalg.eigh(cov_matrix)
                # 替换接近零的特征值
                min_eig = np.min(eigenvalues)
                if min_eig < 1e-8:
                    eigenvalues[eigenvalues < 1e-8] = 1e-8
                # 重构协方差矩阵
                cov_matrix = eigenvectors @ np.diag(eigenvalues) @ eigenvectors.T
                # 再次尝试Cholesky分解
                L = np.linalg.cholesky(cov_matrix)
                numeric_df = numeric_df @ np.linalg.inv(L.T)

        return numeric_df, corr_removed_msg

    def suggest_factor_number(self, ev):
        """根据特征值和碎石图拐点建议因子数量"""
        # 特征值大于1的因子数
        kaiser_num = sum(ev > 1)

        # 寻找碎石图拐点（特征值变化率最大的点）
        diffs = np.diff(ev)
        abs_diffs = np.abs(diffs)
        # 排除最后一个点，因为没有下一个点的差值
        if len(abs_diffs) > 1:
            elbow_idx = np.argmax(abs_diffs[:-1]) + 1  # 拐点索引
            elbow_num = elbow_idx + 1  # 因子数量从1开始计数
        else:
            elbow_num = 1

        # 综合建议（取两种方法的较小值，避免因子过多）
        return min(kaiser_num, elbow_num), kaiser_num

    def factor_analysis(self, data, num_factors, rotation_method):
        # Bartlett球形检验
        chi_square_value, p_value = calculate_bartlett_sphericity(data)

        # KMO检验
        kmo_all, kmo_model = calculate_kmo(data)

        # 创建因子分析对象
        fa = FactorAnalyzer(n_factors=num_factors, rotation=rotation_method)
        fa.fit(data)

        # 获取因子载荷矩阵
        loadings = fa.loadings_

        # 获取共同度
        communalities = fa.get_communalities()

        # 计算特征值和方差贡献率
        ev, v = fa.get_eigenvalues()

        # 计算因子得分（移到前面，避免被提前return跳过）
        factor_scores = fa.transform(data)

        # 计算反图像相关矩阵（修正变量名+异常处理）
        from factor_analyzer.factor_analyzer import corr
        corr_matrix = corr(data)
        try:
            inv_corr = np.linalg.inv(corr_matrix)  # 尝试求逆
        except np.linalg.LinAlgError:
            inv_corr = np.linalg.pinv(corr_matrix)  # 奇异矩阵用伪逆
        anti_image_matrix = np.eye(corr_matrix.shape[0]) - inv_corr  # 英文变量名

        # 统一返回所有结果（删除中间return）
        return loadings, communalities, ev, v, (chi_square_value, p_value), kmo_model, factor_scores, anti_image_matrix

    def generate_loading_interpretation(self, loadings_df):
        """生成因子载荷矩阵的针对性解释"""
        lang = LANGUAGES[self.current_language]
        interpretation = []

        for factor in loadings_df.columns:
            factor_loadings = loadings_df[factor].abs().sort_values(ascending=False)
            strong_vars = factor_loadings[factor_loadings > 0.7].index.tolist()
            moderate_vars = factor_loadings[(factor_loadings >= 0.5) & (factor_loadings <= 0.7)].index.tolist()
            weak_vars = factor_loadings[factor_loadings < 0.5].index.tolist()

            parts = [f"{factor}:"]
            if strong_vars:
                parts.append(f"- {lang['strong_corr']}: {', '.join(strong_vars)}")
            if moderate_vars:
                parts.append(f"- {lang['moderate_corr']}: {', '.join(moderate_vars)}")
            if weak_vars:
                parts.append(f"- {lang['weak_corr']}: {', '.join(weak_vars)}")

            interpretation.append("\n".join(parts))

        return "\n\n".join(interpretation)

    def plot_scree_plot(self, ev, save_path=None):
        plt.figure(figsize=(10, 5))
        plt.plot(range(1, len(ev) + 1), ev, marker='o')

        # 添加表头信息
        title = '碎石图' if self.current_language == 'zh' else 'Scree Plot'
        xlabel = '因子数量' if self.current_language == 'zh' else 'Number of Factors'
        ylabel = '特征值' if self.current_language == 'zh' else 'Eigenvalues'

        plt.title(title, fontsize=14)
        plt.xlabel(xlabel, fontsize=12)
        plt.ylabel(ylabel, fontsize=12)

        # 如果没有提供保存路径，让用户选择
        if save_path is None:
            img_path = filedialog.asksaveasfilename(
                defaultextension=".png",
                filetypes=[("PNG files", "*.png"), ("All files", "*.*")],
                title="保存碎石图" if self.current_language == 'zh' else "Save Scree Plot"
            )
        else:
            img_path = os.path.splitext(save_path)[0] + '_scree_plot.png'

        if img_path:
            plt.savefig(img_path, bbox_inches='tight')
            plt.close()
            return img_path
        else:
            plt.close()
            return None

    def plot_loading_heatmap(self, loadings_df, save_path=None):
        """绘制因子载荷矩阵热力图"""
        plt.figure(figsize=(12, 8))
        import seaborn as sns
        sns.heatmap(loadings_df, annot=True, cmap='coolwarm', center=0, fmt='.4f',
                    annot_kws={"size": 10},
                    cbar_kws={"label": "载荷值" if self.current_language == 'zh' else "Loading Value"})

        title = '因子载荷矩阵热力图' if self.current_language == 'zh' else 'Factor Loading Heatmap'
        plt.title(title, fontsize=14)
        plt.xlabel('因子' if self.current_language == 'zh' else 'Factors', fontsize=12)
        plt.ylabel('变量' if self.current_language == 'zh' else 'Variables', fontsize=12)
        plt.tight_layout()

        if save_path is None:
            img_path = filedialog.asksaveasfilename(
                defaultextension=".png",
                title="保存因子载荷热力图" if self.current_language == 'zh' else "Save Loading Heatmap"
            )
        else:
            img_path = os.path.splitext(save_path)[0] + '_loading_heatmap.png'

        if img_path:
            plt.savefig(img_path, bbox_inches='tight')
            plt.close()
            return img_path
        plt.close()
        return None

    def plot_factor_scores(self, factor_scores, save_path=None):
        """绘制因子得分散点图（前两个因子）"""
        if factor_scores.shape[1] < 2:
            return None  # 至少需要2个因子才能绘制散点图

        plt.figure(figsize=(10, 8))
        plt.scatter(factor_scores[:, 0], factor_scores[:, 1], alpha=0.7, edgecolors='k')

        title = '因子得分散点图' if self.current_language == 'zh' else 'Factor Score Scatter Plot'
        plt.title(title, fontsize=14)
        plt.xlabel('因子1' if self.current_language == 'zh' else 'Factor 1', fontsize=12)
        plt.ylabel('因子2' if self.current_language == 'zh' else 'Factor 2', fontsize=12)
        plt.grid(linestyle='--', alpha=0.7)
        plt.tight_layout()

        if save_path is None:
            img_path = filedialog.asksaveasfilename(
                defaultextension=".png",
                title="保存因子得分散点图" if self.current_language == 'zh' else "Save Score Scatter Plot"
            )
        else:
            img_path = os.path.splitext(save_path)[0] + '_factor_scores.png'

        if img_path:
            plt.savefig(img_path, bbox_inches='tight')
            plt.close()
            return img_path
        plt.close()
        return None

    def plot_communalities(self, communalities_df, save_path=None):
        """绘制共同度条形图"""
        plt.figure(figsize=(14, 10))
        communalities_df.sort_values(
            by='共同度' if self.current_language == 'zh' else 'Communality',
            ascending=False
        ).plot(kind='bar', legend=False, color='skyblue')

        title = '变量共同度条形图' if self.current_language == 'zh' else 'Variable Communalities'
        plt.title(title, fontsize=12)
        plt.xlabel('变量' if self.current_language == 'zh' else 'Variables', fontsize=12)
        plt.ylabel('共同度' if self.current_language == 'zh' else 'Communality', fontsize=12)
        plt.xticks(rotation=30)
        plt.ylim(0, 1.1)  # 共同度范围在0-1之间
        plt.grid(axis='y', linestyle='--', alpha=0.7)
        plt.tight_layout()

        if save_path is None:
            img_path = filedialog.asksaveasfilename(
                defaultextension=".png",
                title="保存共同度条形图" if self.current_language == 'zh' else "Save Communalities Plot"
            )
        else:
            img_path = os.path.splitext(save_path)[0] + '_communalities.png'

        if img_path:
            plt.savefig(img_path, bbox_inches='tight')
            plt.close()
            return img_path
        plt.close()
        return None

    def show_options_dialog(self, suggested_num, kaiser_num):
        """显示包含因子数量输入和旋转方法选择的综合对话框"""
        dialog = tk.Toplevel(self.root)
        dialog.title(LANGUAGES[self.current_language]['options_title'])
        dialog.geometry("400x300")
        dialog.transient(self.root)
        dialog.grab_set()

        # 新增：计算并设置对话框位置到屏幕中央
        dialog.update_idletasks()  # 确保获取正确的窗口尺寸
        screen_width = dialog.winfo_screenwidth()
        screen_height = dialog.winfo_screenheight()
        window_width = dialog.winfo_width()
        window_height = dialog.winfo_height()
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2
        dialog.geometry(f"+{x}+{y}")  # 设置位置

        # 因子数量部分
        ttk.Label(dialog,
                  text=LANGUAGES[self.current_language]['factor_num_msg'].format(suggested_num, kaiser_num)).pack(
            pady=10)

        factor_frame = ttk.Frame(dialog)
        factor_frame.pack(pady=5)
        ttk.Label(factor_frame, text="因子数量:").pack(side=tk.LEFT)
        factor_entry = ttk.Entry(factor_frame, width=10)
        factor_entry.pack(side=tk.LEFT, padx=5)
        factor_entry.insert(0, str(suggested_num))

        # 旋转方法选择部分
        ttk.Label(dialog, text="选择旋转方法:").pack(pady=10)

        rotation_var = IntVar(value=1)  # 默认选择varimax

        rotation_frame = ttk.Frame(dialog)
        rotation_frame.pack(pady=5)

        ttk.Radiobutton(
            rotation_frame,
            text=LANGUAGES[self.current_language]['rotation_option1'],
            variable=rotation_var,
            value=1
        ).pack(anchor=tk.W, pady=2)

        ttk.Radiobutton(
            rotation_frame,
            text=LANGUAGES[self.current_language]['rotation_option2'],
            variable=rotation_var,
            value=2
        ).pack(anchor=tk.W, pady=2)

        # 确认按钮
        result = [None, None]  # 存储因子数量和旋转方法

        def on_confirm():
            try:
                num = int(factor_entry.get())
                # 限制因子数量在1到变量数之间
                if 1 <= num <= self.max_factors:
                    result[0] = num
                    result[1] = 'varimax' if rotation_var.get() == 1 else 'promax'
                    dialog.destroy()
                else:
                    Messagebox.show_error(
                        f"因子数量必须在1到{self.max_factors}之间",
                        "错误" if self.current_language == 'zh' else "Error"
                    )
            except ValueError:
                Messagebox.show_error("请输入有效的整数", "错误" if self.current_language == 'zh' else "Error")

        ttk.Button(dialog, text="确认", command=on_confirm).pack(pady=20)

        self.root.wait_window(dialog)
        return result[0], result[1]

    # 添加模型拟合指标计算
    def calculate_fit_indices(self, data, loadings, communalities):
        """计算因子分析模型拟合指标"""
        # 计算再生相关矩阵

        再生矩阵 = np.dot(loadings, loadings.T)
        np.fill_diagonal(再生矩阵, communalities)

        # 原始相关矩阵
        原始矩阵 = data.corr().values

        # 计算拟合指标
        n = data.shape[0]
        p = data.shape[1]
        df = (p * (p + 1) / 2) - (p + loadings.shape[1] - p)

        # 残差平方和
        sse = np.sum((原始矩阵 - 再生矩阵) ** 2)

        # 均方根误差近似值
        rmsea = np.sqrt(sse / (df * n))

        # 标准化残差均方根
        srmr = np.sqrt(sse / (p * p))

        return {
            'RMSEA': rmsea,
            'SRMR': srmr,
            '残差平方和': sse
        }

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            file_path = ""
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 数据预处理
            df, corr_removed_msg = self.preprocess_data(df)
            if len(df.columns) == 0:
                raise ValueError("预处理后数据为空，请检查输入文件")

            # 显示预处理信息（仅保留高度相关变量移除信息）
            preprocess_msg = LANGUAGES[self.current_language]['preprocess_info'].format(corr_removed_msg)

            # 初步分析获取特征值（用于确定因子数量）
            fa_initial = FactorAnalyzer()
            fa_initial.fit(df)
            ev_initial, _ = fa_initial.get_eigenvalues()

            # 建议因子数量
            suggested_num, kaiser_num = self.suggest_factor_number(ev_initial)
            self.max_factors = len(df.columns)  # 用于验证用户输入

            # 显示综合选项对话框
            num_factors, rotate_method = self.show_options_dialog(suggested_num, kaiser_num)

            if num_factors is None or rotate_method is None:  # 用户取消
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])
                return

            # 进行因子分析
            loadings, communalities, ev, v, bartlett_result, kmo_result, factor_scores, anti_image_matrix = self.factor_analysis(
                df, num_factors, rotate_method
            )

            # 计算模型拟合指标
            fit_indices = self.calculate_fit_indices(df, loadings, communalities)

            # 整理数据
            # 根据当前语言动态生成列名
            factor_names = [f'因子{i + 1}' if self.current_language == 'zh' else f'Factor {i + 1}' for i in
                            range(len(loadings[0]))]
            loadings_df = pd.DataFrame(loadings, index=df.columns, columns=factor_names)
            communalities_df = pd.DataFrame(communalities, index=df.columns,
                                            columns=['共同度' if self.current_language == 'zh' else 'Communality'])
            ev_df = pd.DataFrame(ev[:num_factors],
                                 columns=['特征值' if self.current_language == 'zh' else 'Eigenvalue'])
            v_df = pd.DataFrame(v[:num_factors], columns=[
                '方差贡献率' if self.current_language == 'zh' else 'Variance Contribution Rate'])
            bartlett_df = pd.DataFrame(
                [bartlett_result],
                columns=['卡方值' if self.current_language == 'zh' else 'Chi-square Value',
                         'p值' if self.current_language == 'zh' else 'p-value'],
                index=['Bartlett球形检验' if self.current_language == 'zh' else "Bartlett's Test of Sphericity"]
            )
            kmo_df = pd.DataFrame(
                [kmo_result],
                columns=['KMO值' if self.current_language == 'zh' else 'KMO Value'],
                index=['KMO检验' if self.current_language == 'zh' else 'KMO Test']
            )

            # 添加解释说明
            explanations = LANGUAGES[self.current_language]['explanation']
            interpretations = LANGUAGES[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(
                columns=["因子载荷矩阵", "共同度", "特征值和方差贡献率", "Bartlett球形检验", "KMO检验", "碎石图"]
                if self.current_language == "zh"
                else ["Factor Loadings", "Communalities", "Eigenvalues & Variance Contribution", "Bartlett's Test",
                      "KMO Test", "Scree Plot"]
            )
            explanation_df.insert(0, "统计量_解释说明", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(
                columns=["因子载荷矩阵", "共同度", "特征值和方差贡献率", "Bartlett球形检验", "KMO检验", "碎石图"]
                if self.current_language == "zh"
                else ["Factor Loadings", "Communalities", "Eigenvalues & Variance Contribution", "Bartlett's Test",
                      "KMO Test", "Scree Plot"]
            )
            interpretation_df.insert(0, "统计量_结果解读",
                                     "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 创建 Word 文档
            doc = Document()

            # 添加标题
            doc.add_heading('因子分析结果' if self.current_language == "zh" else 'Factor Analysis Results', 0)

            # 添加预处理信息
            doc.add_heading('数据预处理信息' if self.current_language == "zh" else 'Data Preprocessing Information', 1)
            doc.add_paragraph(preprocess_msg)

            # 添加因子载荷矩阵
            doc.add_heading('因子载荷矩阵' if self.current_language == "zh" else 'Factor Loading Matrix', 1)
            table = doc.add_table(rows=1, cols=len(loadings_df.columns) + 1)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '变量' if self.current_language == "zh" else 'Variable'
            for col_idx, col_name in enumerate(loadings_df.columns):
                hdr_cells[col_idx + 1].text = col_name
            for row_idx, row in loadings_df.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = row_idx
                for col_idx, value in enumerate(row):
                    row_cells[col_idx + 1].text = f"{value:.4f}"

            # 添加因子得分到结果
            doc.add_heading('因子得分' if self.current_language == "zh" else 'Factor Scores', 1)
            scores_df = pd.DataFrame(factor_scores, columns=factor_names)
            # 添加原始数据索引
            scores_df.index = df.index
            table = doc.add_table(rows=1, cols=len(scores_df.columns) + 1)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '样本' if self.current_language == "zh" else 'Sample'
            for col_idx, col_name in enumerate(scores_df.columns):
                hdr_cells[col_idx + 1].text = col_name
            for row_idx, row in scores_df.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(row_idx)
                for col_idx, value in enumerate(row):
                    row_cells[col_idx + 1].text = f"{value:.4f}"

            # 添加模型拟合指标到报告
            doc.add_heading('模型拟合指标' if self.current_language == "zh" else 'Model Fit Indices', 1)
            fit_table = doc.add_table(rows=1, cols=2)
            fit_hdr = fit_table.rows[0].cells
            fit_hdr[0].text = '指标' if self.current_language == "zh" else 'Index'
            fit_hdr[1].text = '值' if self.current_language == "zh" else 'Value'

            # 添加RMSEA
            row = fit_table.add_row().cells
            row[0].text = 'RMSEA'
            row[1].text = f"{fit_indices['RMSEA']:.4f}"
            doc.add_paragraph(
                "RMSEA < 0.05 表示良好拟合，< 0.08 表示可接受拟合"
                if self.current_language == "zh"
                else "RMSEA < 0.05 indicates good fit, < 0.08 indicates acceptable fit"
            )

            # 添加SRMR
            row = fit_table.add_row().cells
            row[0].text = 'SRMR'
            row[1].text = f"{fit_indices['SRMR']:.4f}"
            doc.add_paragraph(
                "SRMR < 0.08 表示良好拟合"
                if self.current_language == "zh"
                else "SRMR < 0.08 indicates good fit"
            )

            # 添加反图像相关矩阵检验
            doc.add_heading(
                '反图像相关矩阵检验' if self.current_language == "zh" else 'Anti-image Correlation Matrix Test', 1)
            # 提取反图像矩阵对角线元素（MSA值）
            msa_values = np.diag(anti_image_matrix)
            msa_df = pd.DataFrame(msa_values, index=df.columns, columns=['MSA值'])

            table = doc.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '变量' if self.current_language == "zh" else 'Variable'
            hdr_cells[1].text = 'MSA值' if self.current_language == "zh" else 'MSA Value'
            for row_idx, row in msa_df.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = row_idx
                row_cells[1].text = f"{row.iloc[0]:.4f}"

            doc.add_paragraph(
                "MSA值 > 0.5 表示变量适合因子分析，值越高越好"
                if self.current_language == "zh"
                else "MSA value > 0.5 indicates the variable is suitable for factor analysis; higher values are better"
            )

            # 添加因子载荷解释
            doc.add_heading('因子载荷解读' if self.current_language == "zh" else 'Interpretation of Factor Loadings', 1)
            doc.add_paragraph(self.generate_loading_interpretation(loadings_df))

            # 添加共同度
            doc.add_heading('共同度' if self.current_language == "zh" else 'Communality', 1)
            table = doc.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '变量' if self.current_language == "zh" else 'Variable'
            hdr_cells[1].text = '共同度' if self.current_language == "zh" else 'Communality'
            for row_idx, row in communalities_df.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = row_idx
                row_cells[1].text = f"{row.iloc[0]:.4f}"

            # 添加特征值和方差贡献率
            doc.add_heading(
                '特征值和方差贡献率' if self.current_language == "zh" else 'Eigenvalues and Variance Contribution Rate',1)
            table = doc.add_table(rows=1, cols=3)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '因子' if self.current_language == "zh" else 'Factor'
            hdr_cells[1].text = '特征值' if self.current_language == "zh" else 'Eigenvalue'
            hdr_cells[2].text = '方差贡献率' if self.current_language == "zh" else 'Variance Contribution Rate'
            for i in range(len(ev_df)):
                row_cells = table.add_row().cells
                row_cells[0].text = f'因子{i + 1}' if self.current_language == "zh" else f'Factor {i + 1}'
                row_cells[1].text = f"{ev_df.iloc[i, 0]:.4f}"
                row_cells[2].text = f"{v_df.iloc[i, 0]:.4f}"

            # 添加 Bartlett 球形检验
            doc.add_heading('Bartlett球形检验' if self.current_language == "zh" else "Bartlett's Test of Sphericity", 1)
            table = doc.add_table(rows=1, cols=3)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '检验名称' if self.current_language == "zh" else 'Test Name'
            hdr_cells[1].text = '卡方值' if self.current_language == "zh" else 'Chi-square Value'
            hdr_cells[2].text = 'p值' if self.current_language == "zh" else 'p-value'
            row_cells = table.add_row().cells
            row_cells[0].text = 'Bartlett球形检验' if self.current_language == "zh" else "Bartlett's Test of Sphericity"
            row_cells[1].text = f"{bartlett_df.iloc[0, 0]:.4f}"
            row_cells[2].text = f"{bartlett_df.iloc[0, 1]:.4e}"

            # 添加 KMO 检验
            doc.add_heading('KMO检验' if self.current_language == "zh" else 'KMO Test', 1)
            table = doc.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '检验名称' if self.current_language == "zh" else 'Test Name'
            hdr_cells[1].text = 'KMO值' if self.current_language == "zh" else 'KMO Value'
            row_cells = table.add_row().cells
            row_cells[0].text = 'KMO检验' if self.current_language == "zh" else 'KMO Test'
            row_cells[1].text = f"{kmo_df.iloc[0, 0]:.4f}"

            # 在添加解释说明部分，将原来的表格替换为项目符号列表
            doc.add_heading('解释说明' if self.current_language == "zh" else 'Explanations', 1)
            for item in explanation_df.columns[1:]:  # 跳过第一列标题
                doc.add_paragraph(f"- {item}: {explanations[item]}", style='ListBullet')

            # 在添加分析结果解读部分，将原来的表格替换为项目符号列表
            doc.add_heading('结果解读' if self.current_language == "zh" else 'Interpretation of Results', 1)
            for item in interpretation_df.columns[1:]:  # 跳过第一列标题
                doc.add_paragraph(f"- {item}: {interpretations[item]}", style='ListBullet')

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:

                # 生成碎石图
                img_paths = []  # 初始化图片路径列表
                scree_img = None
                if 'ev_initial' in locals() and len(ev_initial) > 0:
                    scree_img = self.plot_scree_plot(ev_initial, save_path)
                    # 将碎石图添加到图片列表
                    if scree_img:
                        img_paths.append(('碎石图' if self.current_language == "zh" else 'Scree Plot', scree_img))

                # 因子载荷热力图（新增）
                loading_img = self.plot_loading_heatmap(loadings_df, save_path)
                if loading_img:
                    img_paths.append(
                        ('因子载荷热力图' if self.current_language == "zh" else 'Factor Loading Heatmap', loading_img))

                # 因子得分散点图（新增）
                score_img = self.plot_factor_scores(factor_scores, save_path)
                if score_img:
                    img_paths.append(
                        ('因子得分散点图' if self.current_language == "zh" else 'Factor Score Scatter Plot', score_img))

                # 共同度条形图（新增）
                communal_img = self.plot_communalities(communalities_df, save_path)
                if communal_img:
                    img_paths.append(
                        ('变量共同度条形图' if self.current_language == "zh" else 'Variable Communality Bar Chart',
                         communal_img))

                # 将所有图表添加到Word
                for img_title, img_path in img_paths:
                    doc.add_heading(img_title, 1)
                    doc.add_picture(img_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                # 显示预处理信息和保存结果
                result_msg = f"{preprocess_msg}\n{LANGUAGES[self.current_language]['analysis_success'].format(save_path)}"
                self.result_label.config(text=result_msg, wraplength=400)

            else:
                # 用户取消保存Word，但仍允许保存碎石图
                self.plot_scree_plot(ev_initial)
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'
        self.root.title(LANGUAGES[self.current_language]['title'])
        self.select_button.config(text=LANGUAGES[self.current_language]['select_button'])
        self.analyze_button.config(text=LANGUAGES[self.current_language]['analyze_button'])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]['switch_language'])
        self.open_excel_label.config(text=LANGUAGES[self.current_language]["open_excel_button_text"])
        # 切换语言时更新提示信息
        current_text = self.file_entry.get()
        placeholder = LANGUAGES[self.current_language]['file_entry_placeholder']
        if current_text == LANGUAGES['zh' if self.current_language == 'en' else 'en']['file_entry_placeholder']:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, placeholder)
            self.file_entry.config(foreground='gray')

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["select_button"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["analyze_button"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=LANGUAGES[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=LANGUAGES[self.current_language]["switch_language"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = FactorAnalysisApp()
    app.run()

if __name__ == "__main__":
    run_app()
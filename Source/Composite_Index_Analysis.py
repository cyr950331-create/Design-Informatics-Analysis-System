import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches


# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典（所有显示文本均在此处管理）
LANGUAGES = {
    'zh': {
        'title': "综合指数",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "请选择文件。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "中/英",
        'open_excel_button_text': "示例数据",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'explanation': {
            "composite_index": "将多个指标综合成一个单一指数的值",
            "bar_chart": "展示各个样本的综合指数分布情况",
            "indicator_weights": "各指标在综合指数计算中的重要性占比，权重越高对结果影响越大"
        },
        'interpretation': {
            "composite_index": "综合指数越高，代表该样本在多个指标上的综合表现越好",
            "bar_chart": "可直观比较不同样本的综合指数大小",
            "indicator_weights": "权重分配直接影响综合指数结果，需根据实际业务场景合理设置"
        },
        'weight_input_title': "输入指标权重（用逗号分隔，总和应为1）",
        'weight_validation': "权重输入错误！请确保输入正确的数值且总和为1",
        'insufficient_data': "数据不足！至少需要1个样本数据",
        'weight_sum_error': "权重之和不等于1，请检查第二行权重数据",
        'weight_count_error': "权重数量与指标数量不匹配",
        'analysis_indicators': "分析指标",
        'indicator_weights_label': "指标权重",
        'analysis_results': "分析结果",
        'explanation_title': "解释说明",
        'interpretation_title': "结果解读",
        'bar_chart_title': "综合指数柱状图",
        'samples': "样本",
        'composite_index_label': "综合指数",
        'statistics': "统计量",
        'statistics_value': "统计量值"
    },
    'en': {
        'title': "Composite Index",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Chinese/English",
        'open_excel_button_text': "Example data",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'explanation': {
            "composite_index": "The value obtained by synthesizing multiple indicators into a single index",
            "bar_chart": "Show the distribution of composite indices for each sample",
            "indicator_weights": "The proportion of importance of each indicator in the calculation of the composite index; higher weight means greater impact on the result"
        },
        'interpretation': {
            "composite_index": "The higher the composite index, the better the overall performance of the sample in multiple indicators",
            "bar_chart": "It can visually compare the composite index values of different samples",
            "indicator_weights": "Weight allocation directly affects the composite index results and should be set reasonably according to actual business scenarios"
        },
        'weight_input_title': "Enter indicator weights (comma separated, sum should be 1)",
        'weight_validation': "Weight input error! Please ensure correct values with sum 1",
        'insufficient_data': "Insufficient data! At least 1 sample is required",
        'weight_sum_error': "Sum of weights is not equal to 1, please check the second row",
        'weight_count_error': "Number of weights does not match number of indicators",
        'analysis_indicators': "Analysis Indicators",
        'indicator_weights_label': "Indicator Weights",
        'analysis_results': "Analysis Results",
        'explanation_title': "Explanation",
        'interpretation_title': "Interpretation",
        'bar_chart_title': "Bar Chart of Composite Index",
        'samples': "Samples",
        'composite_index_label': "Composite Index",
        'statistics': "Statistic",
        'statistics_value': "Statistic Value"
    }
}


class CompositeIndexAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为中文
        self.current_language = "en"
        self.file_path = ""  # 保存当前文件路径
        self.indicator_names = []  # 保存指标名称
        self.sample_names = []  # 保存样本名称
        self.weights = None  # 保存权重

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data48.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        current_text = self.file_entry.get()
        if current_text == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        current_text = self.file_entry.get()
        if not current_text:
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')
        else:
            self.file_path = current_text

    def normalize_data(self, data):
        """数据标准化：将每个指标归一化到[0,1]范围"""
        min_vals = np.min(data, axis=0)
        max_vals = np.max(data, axis=0)
        # 处理最大值等于最小值的情况（避免除零）
        ranges = max_vals - min_vals
        ranges[ranges == 0] = 1
        return (data - min_vals) / ranges

    def get_weights_from_data(self, weights_data, num_indicators):
        """从数据中获取权重（第二行）"""
        try:
            # 转换为数值
            weights = np.array(weights_data, dtype=float)

            # 验证权重数量是否与指标数量匹配
            if len(weights) != num_indicators:
                return None

            # 验证权重之和是否为1（允许微小误差）
            if not np.isclose(np.sum(weights), 1.0, atol=0.01):
                return None

            return weights
        except:
            return None

    def composite_index_method(self, data):
        """
        计算综合指数（带标准化和自定义权重）
        :param data: 原始数据矩阵，每行代表一个样本，每列代表一个指标
        :return: 综合指数数组
        """
        # 数据标准化
        normalized_data = self.normalize_data(data)
        # 计算加权综合指数
        composite_indices = np.dot(normalized_data, self.weights)
        return composite_indices

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 读取Excel文件，第一行为表头，第二行为权重
            # 先读取所有数据，包括表头和权重行
            all_data = pd.read_excel(file_path, header=None)

            # 检查数据是否足够
            if len(all_data) < 3:  # 至少需要表头、权重行和一行样本数据
                self.result_label.config(text=LANGUAGES[self.current_language]['insufficient_data'])
                return

            # 提取表头（第一行）
            self.indicator_names = all_data.iloc[0].tolist()
            # 提取权重（第二行）
            weights_data = all_data.iloc[1].tolist()
            # 提取样本数据（从第三行开始）
            sample_data = all_data.iloc[2:].copy()

            # 识别第一列为因素列（样本名称）
            if not sample_data.empty:
                # 获取第一列作为样本名称
                self.sample_names = sample_data.iloc[:, 0].astype(str).tolist()
                # 移除第一列，保留其余列作为数值数据列
                sample_data = sample_data.iloc[:, 1:]
                # 更新指标名称（移除第一列的指标名）
                self.indicator_names = self.indicator_names[1:]
            else:
                # 如果没有样本数据，使用默认命名
                self.sample_names = []

            # 检查并处理非数值数据列
            numeric_columns = []
            non_numeric_columns = []

            for col in sample_data.columns:
                try:
                    # 尝试转换为数值
                    sample_data[col] = pd.to_numeric(sample_data[col])
                    numeric_columns.append(col)
                except:
                    non_numeric_columns.append(col)

            # 移除非数值列
            if non_numeric_columns:
                sample_data = sample_data.drop(columns=non_numeric_columns)
                # 更新指标名称（只保留数值列的名称）
                self.indicator_names = [self.indicator_names[i] for i in
                                       [idx for idx, col in enumerate(sample_data.columns)
                                        if col in numeric_columns]]

            # 如果没有样本名称，使用默认命名
            if not self.sample_names or len(self.sample_names) != len(sample_data):
                self.sample_names = [f"{LANGUAGES[self.current_language]['samples'].lower()}{i + 1}"
                                    for i in range(len(sample_data))]

            # 确保还有数值列
            if sample_data.empty:
                raise ValueError(LANGUAGES[self.current_language]['no_valid_numeric_data'])

            num_indicators = len(self.indicator_names)

            # 处理权重（权重应对应剩余的指标列）
            self.weights = self.get_weights_from_data(weights_data[1:], num_indicators)  # 跳过第一列权重

            # 如果从数据中获取权重失败，抛出错误
            if self.weights is None:
                # 检查是数量不匹配还是总和不等于1
                try:
                    weights = np.array(weights_data[1:], dtype=float)  # 跳过第一列权重
                    if len(weights) != num_indicators:
                        raise ValueError(LANGUAGES[self.current_language]['weight_count_error'])
                    else:
                        raise ValueError(LANGUAGES[self.current_language]['weight_sum_error'])
                except:
                    raise ValueError(LANGUAGES[self.current_language]['weight_validation'])

            # 转换为数值矩阵
            data = sample_data.values.astype(float)

            # 进行综合指数计算
            composite_indices = self.composite_index_method(data)

            # 整理数据
            data = [
                [LANGUAGES[self.current_language]['composite_index_label'], composite_indices.tolist()],
                [LANGUAGES[self.current_language]['indicator_weights_label'],
                 [f"{name}: {weight:.4f}" for name, weight in zip(self.indicator_names, self.weights)]]
            ]
            headers = [LANGUAGES[self.current_language]['statistics'],
                      LANGUAGES[self.current_language]['statistics_value']]
            result_df = pd.DataFrame(data, columns=headers)


            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加标题
                doc.add_heading(LANGUAGES[self.current_language]['title'], 0)

                # 添加指标和权重信息
                doc.add_paragraph(f"{LANGUAGES[self.current_language]['analysis_indicators']:<10}: {', '.join(self.indicator_names)}")
                doc.add_paragraph(
                    f"{LANGUAGES[self.current_language]['indicator_weights_label']:<10}: {', '.join([f'{name}: {w:.4f}' for name, w in zip(self.indicator_names, self.weights)])}")
                doc.add_paragraph("")  # 空行

                # 添加表格数据
                doc.add_heading(LANGUAGES[self.current_language]['analysis_results'], level=1)
                table = doc.add_table(rows=1, cols=len(result_df.columns))
                hdr_cells = table.rows[0].cells
                for col_idx, col_name in enumerate(result_df.columns):
                    hdr_cells[col_idx].text = col_name
                for _, row in result_df.iterrows():
                    row_cells = table.add_row().cells
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx].text = str(value)

                # 添加解释说明部分
                doc.add_heading(LANGUAGES[self.current_language]['explanation_title'], level=1)
                explanations = LANGUAGES[self.current_language]['explanation']
                for key, value in explanations.items():
                    # 主项：项目名称（加粗）
                    main_para = doc.add_paragraph(style='List Bullet')
                    # 根据键获取对应的显示名称
                    display_name = (LANGUAGES[self.current_language]['composite_index_label']
                                  if key == 'composite_index'
                                  else LANGUAGES[self.current_language]['bar_chart_title']
                                  if key == 'bar_chart'
                                  else LANGUAGES[self.current_language]['indicator_weights_label'])
                    main_run = main_para.add_run(display_name)
                    main_run.bold = True
                    # 子项：解释内容
                    doc.add_paragraph(value, style='List Bullet 2')

                # 添加结果解读部分
                doc.add_heading(LANGUAGES[self.current_language]['interpretation_title'], level=1)
                interpretations = LANGUAGES[self.current_language]['interpretation']
                for key, value in interpretations.items():
                    # 主项：项目名称（加粗）
                    main_para = doc.add_paragraph(style='List Bullet')
                    # 根据键获取对应的显示名称
                    display_name = (LANGUAGES[self.current_language]['composite_index_label']
                                  if key == 'composite_index'
                                  else LANGUAGES[self.current_language]['bar_chart_title']
                                  if key == 'bar_chart'
                                  else LANGUAGES[self.current_language]['indicator_weights_label'])
                    main_run = main_para.add_run(display_name)
                    main_run.bold = True
                    # 子项：解读内容
                    doc.add_paragraph(value, style='List Bullet 2')

                # 生成综合指数柱状图
                fig, ax = plt.subplots(figsize=(12, 8))
                bars = ax.bar(self.sample_names, composite_indices, color='steelblue', edgecolor='black')

                # 添加数值标签
                for bar in bars:
                    height = bar.get_height()
                    ax.text(bar.get_x() + bar.get_width() / 2., height + 0.01,
                            f'{height:.4f}', ha='center', va='bottom', fontsize=9)

                ax.set_title(
                    LANGUAGES[self.current_language]['bar_chart_title'],
                    fontsize=14, pad=20
                )
                ax.set_xlabel(LANGUAGES[self.current_language]['samples'], fontsize=12, labelpad=10)
                ax.set_ylabel(LANGUAGES[self.current_language]['composite_index_label'], fontsize=12,
                              labelpad=10)
                plt.xticks(rotation=0, ha='center')

                # 添加网格线
                ax.yaxis.grid(True, linestyle='--', alpha=0.7)

                # 调整布局
                plt.tight_layout()

                # 保存图片
                img_path = os.path.splitext(save_path)[0] + '_composite_index_bar.png'
                plt.savefig(img_path, dpi=300)
                plt.close()

                # 将图片添加到 Word 文档
                doc.add_heading(LANGUAGES[self.current_language]['bar_chart_title'], level=1)
                doc.add_picture(img_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                # 删除临时图片
                if os.path.exists(img_path):
                    os.remove(img_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(save_path)
                self.result_label.config(text=result_msg, wraplength=400)

            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        # 保存当前文件路径
        current_file = self.file_entry.get()
        if current_file == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            current_file = self.file_path

        # 切换语言
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'

        # 更新UI文本
        self.root.title(LANGUAGES[self.current_language]['title'])
        self.select_button.config(text=LANGUAGES[self.current_language]['select_button'])
        self.analyze_button.config(text=LANGUAGES[self.current_language]['analyze_button'])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]['switch_language'])
        self.open_excel_label.config(text=LANGUAGES[self.current_language]["open_excel_button_text"])

        # 恢复文件路径
        self.file_entry.delete(0, tk.END)
        if current_file and current_file != LANGUAGES['zh']["file_entry_placeholder"] and current_file != \
                LANGUAGES['en']["file_entry_placeholder"]:
            self.file_entry.insert(0, current_file)
            self.file_entry.config(foreground='black')
        else:
            self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
            self.file_entry.config(foreground='gray')

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=LANGUAGES[self.current_language]['select_button'],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=LANGUAGES[self.current_language]['analyze_button'],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=LANGUAGES[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=LANGUAGES[self.current_language]['switch_language'],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = CompositeIndexAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()
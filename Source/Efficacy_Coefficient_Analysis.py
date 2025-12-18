import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
import matplotlib.pyplot as plt
import pathlib
from docx import Document
from docx.shared import Inches


# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典，使用英文键名
LANGUAGES = {
    'zh': {
        'title': "功效系数",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "请选择文件。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "中/英",
        'open_excel_button_text': "示例数据",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'explanation': {
            'actual_values': "每个评价指标的实际测量值",
            'unacceptable_values': "每个评价指标的最低可接受值",
            'satisfactory_values': "每个评价指标的理想值",
            'efficacy_coefficients': "根据各指标实际值、不允许值和满意值计算得到的功效系数",
            'comprehensive_efficacy': "所有指标功效系数的加权平均值"
        },
        'interpretation': {
            'actual_values': "反映各指标的实际表现",
            'unacceptable_values': "作为指标表现的下限参考",
            'satisfactory_values': "作为指标表现的上限参考",
            'efficacy_coefficients': "值越高，说明该指标表现越好",
            'comprehensive_efficacy': "综合反映所有指标的整体表现，值越高越好"
        },
        'weight_warning': "警告：权重总和为 {:.2f}，建议调整为 1.0",
        'div_zero_error': "错误：指标 '{}' 的满意值与不允许值相等，无法计算",
        'outlier_warning': "警告：指标 '{}' 的实际值超出范围，已进行修正",
        'statistic_explanation': "解释说明",
        'statistic_interpretation': "结果解读",
        'analysis_results': "功效系数分析结果",
        'analysis_warnings': "分析警告",
        'detailed_results': "详细结果",
        'indicator_statistic': "指标/统计量",
        'details': "详细信息",
        'result_value': "结果值",
        'indicator_explanation': "指标解释",
        'result_interpretation': "结果解读",
        'visualization_results': "可视化结果",
        'bar_chart_title': "功效系数向量柱状图",
        'radar_chart_title': "功效系数雷达图",
        'x_label': "指标",
        'y_label': "功效系数",
        'comprehensive_efficacy_label': "综合功效系数",
        'all_indicators_weighted_average': "所有指标加权平均"
    },
    'en': {
        'title': "Efficacy Coefficient",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Chinese/English",
        'open_excel_button_text': "Example data",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'explanation': {
            'actual_values': "The actual measured values of each evaluation indicator",
            'unacceptable_values': "The minimum acceptable values of each evaluation indicator",
            'satisfactory_values': "The ideal values of each evaluation indicator",
            'efficacy_coefficients': "The efficacy coefficients calculated based on the actual values, unacceptable values, and satisfactory values of each indicator",
            'comprehensive_efficacy': "The weighted average of the efficacy coefficients of all indicators"
        },
        'interpretation': {
            'actual_values': "Reflects the actual performance of each indicator",
            'unacceptable_values': "Serves as the lower limit reference for indicator performance",
            'satisfactory_values': "Serves as the upper limit reference for indicator performance",
            'efficacy_coefficients': "The higher the value, the better the performance of the indicator",
            'comprehensive_efficacy': "Comprehensively reflects the overall performance of all indicators. The higher the value, the better"
        },
        'weight_warning': "Warning: The sum of weights is {:.2f}, it's recommended to adjust to 1.0",
        'div_zero_error': "Error: The satisfactory value equals the unacceptable value for indicator '{}', cannot calculate",
        'outlier_warning': "Warning: The actual value of indicator '{}' is out of range and has been corrected",
        'statistic_explanation': "Explanation",
        'statistic_interpretation': "Interpretation",
        'analysis_results': "Efficacy Coefficient Analysis Results",
        'analysis_warnings': "Analysis Warnings",
        'detailed_results': "Detailed Results",
        'indicator_statistic': "Indicator/Statistic",
        'details': "Details",
        'result_value': "Result Value",
        'indicator_explanation': "Indicator Explanation",
        'result_interpretation': "Result Interpretation",
        'visualization_results': "Visualization Results",
        'bar_chart_title': "Bar Chart of Efficacy Coefficient Vector",
        'radar_chart_title': "Radar Chart of Efficacy Coefficients",
        'x_label': "Indicators",
        'y_label': "Efficacy Coefficient",
        'comprehensive_efficacy_label': "Comprehensive Efficacy Coefficient",
        'all_indicators_weighted_average': "Weighted average of all indicators"
    }
}


class EfficacyCoefficientAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为中文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data49.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def efficacy_coefficient_analysis(self, actual_values, unacceptable_values, satisfactory_values, weights,
                                      indicators):
        """
        进行功效系数分析
        :param actual_values: 各指标实际值
        :param unacceptable_values: 各指标不允许值
        :param satisfactory_values: 各指标满意值
        :param weights: 各指标权重
        :param indicators: 指标名称列表
        :return: 功效系数向量、综合功效系数和警告信息
        """
        warnings = []
        efficacy_coefficients = []

        # 检查权重总和
        weight_sum = np.sum(weights)
        if not np.isclose(weight_sum, 1.0, atol=0.01):
            warnings.append(LANGUAGES[self.current_language]['weight_warning'].format(weight_sum))

        # 计算每个指标的功效系数
        for i in range(len(actual_values)):
            # 检查满意值和不允许值是否相等
            if np.isclose(satisfactory_values[i], unacceptable_values[i]):
                raise ValueError(LANGUAGES[self.current_language]['div_zero_error'].format(indicators[i]))

            # 处理异常值（实际值超出范围）
            actual = actual_values[i]
            if actual < unacceptable_values[i]:
                warnings.append(LANGUAGES[self.current_language]['outlier_warning'].format(indicators[i]))
                actual = unacceptable_values[i]  # 修正为不允许值
            elif actual > satisfactory_values[i]:
                warnings.append(LANGUAGES[self.current_language]['outlier_warning'].format(indicators[i]))
                actual = satisfactory_values[i]  # 修正为满意值

            # 计算功效系数
            coeff = (actual - unacceptable_values[i]) / (satisfactory_values[i] - unacceptable_values[i]) * 40 + 60
            efficacy_coefficients.append(coeff)

        # 计算综合功效系数
        comprehensive_efficacy_coefficient = np.dot(efficacy_coefficients, weights)
        return np.array(efficacy_coefficients), comprehensive_efficacy_coefficient, warnings

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            file_path = ""
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件，第一行为表头，第一列为因素列
            df = pd.read_excel(file_path, header=0)  # 第一行为表头
            if df.empty:
                raise ValueError("Excel文件内容为空")

            # 提取指标名称（第一列）
            indicators = df.iloc[:, 0].tolist()

            # 提取数据（实际值、不允许值、满意值、权重）
            # 假设表头包含这些关键词（不区分大小写）
            actual_col = None
            unacceptable_col = None
            satisfactory_col = None
            weight_col = None

            for col in df.columns:
                col_lower = str(col).lower()
                if '实际值' in col_lower or 'actual' in col_lower:
                    actual_col = col
                elif '不允许值' in col_lower or 'unacceptable' in col_lower:
                    unacceptable_col = col
                elif '满意值' in col_lower or 'satisfactory' in col_lower:
                    satisfactory_col = col
                elif '权重' in col_lower or 'weight' in col_lower:
                    weight_col = col

            # 检查是否找到所有必要的列
            if None in [actual_col, unacceptable_col, satisfactory_col, weight_col]:
                raise ValueError("Excel文件必须包含实际值、不允许值、满意值和权重列")

            # 提取数据
            actual_values = df[actual_col].values.astype(float)
            unacceptable_values = df[unacceptable_col].values.astype(float)
            satisfactory_values = df[satisfactory_col].values.astype(float)
            weights = df[weight_col].values.astype(float)

            # 进行功效系数分析
            efficacy_coefficients, comprehensive_efficacy_coefficient, warnings = self.efficacy_coefficient_analysis(
                actual_values, unacceptable_values, satisfactory_values, weights, indicators)

            # 整理数据
            data = []
            for i, indicator in enumerate(indicators):
                data.append([
                    indicator,
                    f"{LANGUAGES[self.current_language]['explanation']['actual_values'].split('的')[0]}: {actual_values[i]}, {LANGUAGES[self.current_language]['explanation']['unacceptable_values'].split('的')[0]}: {unacceptable_values[i]}, {LANGUAGES[self.current_language]['explanation']['satisfactory_values'].split('的')[0]}: {satisfactory_values[i]}, 权重: {weights[i]}",
                    f"{LANGUAGES[self.current_language]['explanation']['efficacy_coefficients'].split('的')[0]}: {efficacy_coefficients[i]:.2f}"
                ])

            # 添加综合功效系数
            data.append([
                LANGUAGES[self.current_language]['comprehensive_efficacy_label'],
                LANGUAGES[self.current_language]['all_indicators_weighted_average'],
                f"{comprehensive_efficacy_coefficient:.2f}"
            ])

            # 添加警告信息
            for warn in warnings:
                data.append(["警告" if self.current_language == 'zh' else "Warning", warn, ""])

            headers = [
                LANGUAGES[self.current_language]['indicator_statistic'],
                LANGUAGES[self.current_language]['details'],
                LANGUAGES[self.current_language]['result_value']
            ]
            result_df = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            explanations = LANGUAGES[self.current_language]['explanation']
            interpretations = LANGUAGES[self.current_language]['interpretation']

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建一个新的 Word 文档
                doc = Document()

                # 添加标题
                doc.add_heading(LANGUAGES[self.current_language]['analysis_results'], 0)

                # 添加警告信息（如果有）
                if warnings:
                    doc.add_heading(LANGUAGES[self.current_language]['analysis_warnings'], level=1)
                    for warn in warnings:
                        doc.add_paragraph(warn)
                    doc.add_paragraph("")  # 空行分隔

                # 添加结果表格
                doc.add_heading(LANGUAGES[self.current_language]['detailed_results'], level=1)
                table = doc.add_table(rows=1, cols=len(headers))
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header

                # 添加数据行
                for index, row in result_df.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)

                # 添加解释说明
                doc.add_heading(LANGUAGES[self.current_language]['indicator_explanation'], level=1)
                expl_paragraph = doc.add_paragraph()
                expl_paragraph.add_run(f"{LANGUAGES[self.current_language]['statistic_explanation']}: ").bold = True

                # 创建解释说明项目符号列表
                expl_list = doc.add_paragraph(style='List Bullet')
                expl_list.add_run(f"{LANGUAGES[self.current_language]['explanation']['actual_values'].split('的')[0]}: ").bold = True
                expl_list.add_run(explanations["actual_values"])

                expl_list = doc.add_paragraph(style='List Bullet')
                expl_list.add_run(f"{LANGUAGES[self.current_language]['explanation']['unacceptable_values'].split('的')[0]}: ").bold = True
                expl_list.add_run(explanations["unacceptable_values"])

                expl_list = doc.add_paragraph(style='List Bullet')
                expl_list.add_run(f"{LANGUAGES[self.current_language]['explanation']['satisfactory_values'].split('的')[0]}: ").bold = True
                expl_list.add_run(explanations["satisfactory_values"])

                expl_list = doc.add_paragraph(style='List Bullet')
                expl_list.add_run(f"{LANGUAGES[self.current_language]['explanation']['efficacy_coefficients'].split('的')[0]}: ").bold = True
                expl_list.add_run(explanations["efficacy_coefficients"])

                expl_list = doc.add_paragraph(style='List Bullet')
                expl_list.add_run(f"{LANGUAGES[self.current_language]['comprehensive_efficacy_label']}: ").bold = True
                expl_list.add_run(explanations["comprehensive_efficacy"])

                # 添加结果解读
                doc.add_heading(LANGUAGES[self.current_language]['result_interpretation'], level=1)
                interp_paragraph = doc.add_paragraph()
                interp_paragraph.add_run(f"{LANGUAGES[self.current_language]['statistic_interpretation']}: ").bold = True

                # 创建结果解读项目符号列表
                interp_list = doc.add_paragraph(style='List Bullet')
                interp_list.add_run(f"{LANGUAGES[self.current_language]['explanation']['actual_values'].split('的')[0]}: ").bold = True
                interp_list.add_run(interpretations["actual_values"])

                interp_list = doc.add_paragraph(style='List Bullet')
                interp_list.add_run(f"{LANGUAGES[self.current_language]['explanation']['unacceptable_values'].split('的')[0]}: ").bold = True
                interp_list.add_run(interpretations["unacceptable_values"])

                interp_list = doc.add_paragraph(style='List Bullet')
                interp_list.add_run(f"{LANGUAGES[self.current_language]['explanation']['satisfactory_values'].split('的')[0]}: ").bold = True
                interp_list.add_run(interpretations["satisfactory_values"])

                interp_list = doc.add_paragraph(style='List Bullet')
                interp_list.add_run(f"{LANGUAGES[self.current_language]['explanation']['efficacy_coefficients'].split('的')[0]}: ").bold = True
                interp_list.add_run(interpretations["efficacy_coefficients"])

                interp_list = doc.add_paragraph(style='List Bullet')
                interp_list.add_run(f"{LANGUAGES[self.current_language]['comprehensive_efficacy_label']}: ").bold = True
                interp_list.add_run(interpretations["comprehensive_efficacy"])

                # 生成功效系数向量柱状图
                fig, ax = plt.subplots(figsize=(12, 8), dpi=300)  # 提高分辨率
                bars = ax.bar(indicators, efficacy_coefficients, color='skyblue', edgecolor='black')

                # 添加数据标签
                for bar in bars:
                    height = bar.get_height()
                    ax.text(bar.get_x() + bar.get_width() / 2., height + 0.5,
                            f'{height:.2f}', ha='center', va='bottom')

                ax.set_title(
                    LANGUAGES[self.current_language]['bar_chart_title'],
                    fontsize=14, pad=20)
                ax.set_xlabel(LANGUAGES[self.current_language]['x_label'], fontsize=12, labelpad=10)
                ax.set_ylabel(LANGUAGES[self.current_language]['y_label'], fontsize=12,
                              labelpad=10)
                plt.xticks(rotation=0, ha='center')
                plt.grid(axis='y', linestyle='--', alpha=0.7)  # 添加网格线
                plt.tight_layout()

                # 生成综合功效系数雷达图
                fig2, ax2 = plt.subplots(figsize=(8, 8), subplot_kw=dict(polar=True), dpi=300)
                angles = np.linspace(0, 2 * np.pi, len(indicators), endpoint=False).tolist()
                values = efficacy_coefficients.tolist()
                values += values[:1]
                angles += angles[:1]
                indicators_radar = indicators + [indicators[0]]

                ax2.fill(angles, values, color='skyblue', alpha=0.4)
                ax2.plot(angles, values, color='blue', linewidth=2, marker='o', markersize=5)
                ax2.set_thetagrids(np.degrees(angles), indicators_radar)
                ax2.set_ylim(0, 100)  # 设置合理的Y轴范围
                ax2.set_title(
                    LANGUAGES[self.current_language]['radar_chart_title'],
                    fontsize=14, pad=20)

                # 添加网格线
                ax2.grid(True, linestyle='--', alpha=0.7)
                # 保存图片
                img_path = os.path.splitext(save_path)[0] + '_bar.png'
                radar_path = os.path.splitext(save_path)[0] + '_radar.png'
                plt.savefig(radar_path)
                plt.close(fig2)
                plt.savefig(img_path)
                plt.close(fig)

                # 将图片插入到 Word 文档中
                doc.add_heading(LANGUAGES[self.current_language]['visualization_results'], level=1)
                doc.add_picture(img_path, width=Inches(6))
                doc.add_picture(radar_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                # 构建结果消息
                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(save_path)
                if warnings:
                    result_msg += "\n".join(warnings)
                self.result_label.config(text=result_msg, wraplength=400)

            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'
        self.root.title(LANGUAGES[self.current_language]['title'])
        self.select_button.config(text=LANGUAGES[self.current_language]['select_button'])
        self.analyze_button.config(text=LANGUAGES[self.current_language]['analyze_button'])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]['switch_language'])
        self.open_excel_label.config(text=LANGUAGES[self.current_language]["open_excel_button_text"])

        # 切换语言时更新提示信息
        current_text = self.file_entry.get()
        if current_text == LANGUAGES['zh' if self.current_language == 'en' else 'en']["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["select_button"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["analyze_button"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=LANGUAGES[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=LANGUAGES[self.current_language]["switch_language"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = EfficacyCoefficientAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()
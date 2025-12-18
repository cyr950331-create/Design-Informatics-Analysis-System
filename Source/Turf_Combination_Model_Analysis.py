import tkinter as tk
from tkinter import filedialog, simpledialog
import os
import pandas as pd
import numpy as np
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import itertools

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
languages = {
    'zh': {
        'title': "Turf 组合模型",
        'select_button_text': "选择文件",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'analyze_button_text': "分析文件",
        'no_file_selected': "请选择文件。",
        'file_not_exists': "文件不存在，请重新选择。",
        'analysis_error': "分析文件时出错: {}",
        'analysis_complete': "分析完成，结果已保存到 {}，相关图片已保存。",
        'no_save_path_selected': "未选择保存路径，结果未保存。",
        'open_excel_button_text': "示例数据",
        'switch_language_button_text': "中/英",
        'combo_count_prompt': "请输入需要的组合数量 (1-6):",
        'invalid_combo_count': "请输入1-6之间的整数",
        'explanation': {
            "Turf组合得分": "Turf组合模型中每个组合的得分，反映了该组合的吸引力。",
            "最优组合": "得分最高的组合，代表了最有吸引力的产品或服务组合。",
            "样本量": "每个样本中的观测值数量。",
            "均值": "样本数据的平均值。"
        },
        'interpretation': {
            "Turf组合得分": "Turf组合得分越高，说明该组合越受用户欢迎。",
            "最优组合": "最优组合是最能满足用户需求的组合，可作为产品或服务的推荐组合。",
            "样本量": "样本量的大小会影响统计检验的稳定性，较大的样本量通常能提供更可靠的结果。",
            "均值": "均值反映了数据的平均水平，可用于比较不同变量的集中趋势。"
        },
        'statistics': {
            'turf_score': 'Turf组合得分',
            'top_combinations': '最优组合（前3名）',
            'sample_sizes': '各列样本量',
            'means': '各列均值',
            'statistic': '统计量',
            'statistic_value': '统计量值'
        }
    },
    'en': {
        'title': "Turf Combination Model",
        'select_button_text': "Select File",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'analyze_button_text': "Analyze File",
        'no_file_selected': "Please select a valid file path.",
        'file_not_exists': "The file does not exist. Please select again.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'analysis_complete': "Analysis completed. The results have been saved to {}, and the relevant images have been saved.",
        'no_save_path_selected': "No save path selected. The results were not saved.",
        'open_excel_button_text': "Example data",
        'switch_language_button_text': "Chinese/English",
        'combo_count_prompt': "Please enter the number of combinations (1-6):",
        'invalid_combo_count': "Please enter an integer between 1 and 6",
        'explanation': {
            "Turf Combination Score": "The score of each combination in the Turf combination model, reflecting the attractiveness of the combination.",
            "Optimal Combination": "The combination with the highest score, representing the most attractive product or service combination.",
            "Sample Size": "The number of observations in each sample.",
            "Mean Value": "The average value of the sample data."
        },
        'interpretation': {
            "Turf Combination Score": "The higher the Turf combination score, the more popular the combination is among users.",
            "Optimal Combination": "The optimal combination is the one that best meets the needs of users and can be recommended as a product or service combination.",
            "Sample Size": "The sample size affects the stability of the statistical test. A larger sample size usually provides more reliable results.",
            "Mean Value": "The mean reflects the average level of the data and can be used to compare the central tendencies of different variables."
        },
        'statistics': {
            'turf_score': 'Turf Combination Score',
            'top_combinations': 'Top 3 Optimal Combinations',
            'sample_sizes': 'Column Sample Sizes',
            'means': 'Column Means',
            'statistic': 'Statistic',
            'statistic_value': 'Statistic Value'
        }
    }
}


class TurfCombinationModelAnalysisApp:
    def __init__(self, root=None):
        self.current_language = "en"
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])
        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data7.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    # 新增：安全处理NaN值的工具函数（核心修复）
    def safe_isna(self, value):
        """处理不同类型数据的NaN判断，避免布尔歧义"""
        # 如果是字典、列表、元组等可迭代对象，直接返回False（不视为NaN）
        if isinstance(value, (dict, list, tuple, set)):
            return False
        # 如果是numpy数组，用any()判断是否包含NaN
        if isinstance(value, np.ndarray):
            return np.isnan(value).any()
        # 普通数值/字符串，直接用pd.isna判断
        return pd.isna(value)

    def turf_analysis(self, data, combo_size):
        columns = data.columns.tolist()
        combinations = list(itertools.combinations(columns, combo_size))
        combination_scores = {}
        for combo in combinations:
            # 计算得分时忽略NaN，并用round避免浮点数精度问题
            score = float(round(((data[list(combo)].sum(axis=1) > 0).mean() * 100), 2))
            combination_scores[combo] = score
        sorted_combinations = sorted(combination_scores.items(), key=lambda x: x[1], reverse=True)
        return combination_scores, sorted_combinations

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 获取组合数量
            combo_count_str = simpledialog.askstring(
                "组合数量",
                languages[self.current_language]["combo_count_prompt"],
                parent=self.root
            )
            if not combo_count_str or not combo_count_str.isdigit():
                self.result_label.config(text=languages[self.current_language]["invalid_combo_count"])
                return
            combo_count = int(combo_count_str)
            if combo_count < 1 or combo_count > 6:
                self.result_label.config(text=languages[self.current_language]["invalid_combo_count"])
                return

            # 1. 读取Excel并预处理NaN值（核心修复：明确处理逻辑）
            df = pd.read_excel(file_path)

            # 处理数值列的NaN：用均值填充
            numerical_cols = df.select_dtypes(include=[np.number]).columns
            for col in numerical_cols:
                # 使用直接赋值方式替代inplace=True
                df[col] = df[col].fillna(round(df[col].mean(skipna=True), 2))

            # 处理非数值列的NaN：删除整行
            non_numerical_cols = df.select_dtypes(exclude=[np.number]).columns
            if not non_numerical_cols.empty:
                df.dropna(subset=non_numerical_cols, inplace=True)

            # 筛选数值列用于分析
            numerical_df = df[numerical_cols]
            if numerical_df.empty:
                raise ValueError("数据中没有数值列，无法进行Turf组合模型分析。")
            if len(numerical_df.columns) < combo_count:
                raise ValueError(f"数据列数量不足（仅{len(numerical_df.columns)}列），无法生成{combo_count}个元素的组合")

            # 2. 执行Turf分析
            combination_scores, sorted_combinations = self.turf_analysis(numerical_df, combo_count)

            # 3. 计算统计量
            sample_sizes = numerical_df.count().to_dict()
            means = {col: float(round(numerical_df[col].mean(skipna=True), 2)) for col in numerical_df.columns}

            # 整理结果数据
            formatted_scores = {", ".join(combo): score for combo, score in combination_scores.items()}
            formatted_top_combinations = [", ".join(combo) for combo, _ in
                                          sorted_combinations[:3]] if sorted_combinations else []

            # 只保留"统计量"和"统计量值"两列
            stats = languages[self.current_language]['statistics']
            result_data = [
                [stats['turf_score'], str(formatted_scores)],
                [stats['top_combinations'], str(formatted_top_combinations)],
                [stats['sample_sizes'], str(sample_sizes)],
                [stats['means'], str(means)]
            ]
            headers = [stats['statistic'], stats['statistic_value']]
            df_result = pd.DataFrame(result_data, columns=headers)

            # 保存到Word文档
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                doc = Document()
                doc.add_heading('Turf组合模型分析结果' if self.current_language == 'zh' else 'Turf Combination Model Analysis Results',0)

                # 添加结果表格
                doc.add_heading('分析结果' if self.current_language == 'zh' else 'Analysis Results', level=1)
                table = doc.add_table(rows=1, cols=len(df_result.columns))
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(df_result.columns):
                    hdr_cells[i].text = str(header)

                for _, row in df_result.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        if self.safe_isna(value):
                            row_cells[i].text = ""
                        else:
                            row_cells[i].text = str(value)

                # 添加解释说明（改为项目符号列表）
                doc.add_heading('解释说明' if self.current_language == 'zh' else 'Explanations', level=1)
                explanations = languages[self.current_language]['explanation']
                explanation_para = doc.add_paragraph()
                for key, value in explanations.items():
                    explanation_para.add_run(f'• {key}: {value}\n')

                # 添加结果解读（改为项目符号列表）
                doc.add_heading('结果解读' if self.current_language == 'zh' else 'Interpretations', level=1)
                interpretations = languages[self.current_language]['interpretation']
                interpretation_para = doc.add_paragraph()
                for key, value in interpretations.items():
                    interpretation_para.add_run(f'• {key}: {value}\n')

                # 生成柱状图（只显示前10个组合）
                if sorted_combinations:
                    fig, ax = plt.subplots(figsize=(12, 6))
                    top_10 = sorted_combinations[:10]
                    combo_labels = [", ".join(combo) for combo, _ in top_10]
                    combo_values = [score for _, score in top_10]

                    ax.bar(combo_labels, combo_values, color='#3498db')
                    ax.set_title(
                        f'{combo_count}元素组合得分TOP10' if self.current_language == 'zh' else f'TOP 10 {combo_count}-element Combination Scores',
                        fontsize=14)
                    ax.set_xlabel('组合' if self.current_language == 'zh' else 'Combinations', fontsize=12)
                    ax.set_ylabel('覆盖率 (%)' if self.current_language == 'zh' else 'Coverage (%)', fontsize=12)
                    plt.xticks(rotation=45, ha='right')
                    plt.tight_layout()

                    img_path = os.path.splitext(save_path)[0] + '_scores.png'
                    plt.savefig(img_path, dpi=300, bbox_inches='tight')
                    plt.close()
                    doc.add_picture(img_path, width=Inches(6))

                # 保存文档
                doc.save(save_path)
                self.result_label.config(text=languages[self.current_language]["analysis_complete"].format(save_path),
                                         wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]["no_save_path_selected"])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]["analysis_error"].format(str(e)))

    def switch_language(self, event):
        self.current_language = "zh" if self.current_language == "en" else "en"
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])
        self.open_excel_label.config(text=languages[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=languages[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)
        self.switch_language_label = ttk.Label(frame,
                                               text=languages[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        self.root.mainloop()


def run_app():
    app = TurfCombinationModelAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()
import tkinter as tk
from tkinter import filedialog
import os
import pandas as pd
import matplotlib.pyplot as plt
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from docx import Document
from docx.shared import Inches
import numpy as np

# 设置支持中文的字体
plt.rcParams['font.family'] = 'SimHei'  # 使用黑体字体，可根据系统情况修改
plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

# 定义语言字典
languages = {
    "zh": {
        "title": "价格敏感度测试模型",
        "select_button_text": "选择文件",
        "file_entry_placeholder": "请输入待分析 Excel 文件的完整路径",
        "analyze_button_text": "分析文件",
        "no_file_selected": "请选择文件。",
        "file_not_exists": "文件不存在，请重新选择。",
        "analysis_error": "分析文件时出错: {}",
        "analysis_complete": "分析完成，结果已保存到 {}，PSM 图已保存。",
        "no_save_path_selected": "未选择保存路径，结果未保存。",
        "columns_stats": ["价格点", "太便宜比例", "便宜比例", "贵比例", "太贵比例"],
        "switch_language_button_text": "中/英",
        'open_excel_button_text': "示例数据",
        "column_name_hint": "列名应包含 TooCheap_, Cheap_, Expensive_, TooExpensive_",
        "warning_price_order": "警告: 价格点未按升序排列，已自动排序"
    },
    "en": {
        "title": "Price Sensitivity Meter",
        "select_button_text": "Select File",
        "file_entry_placeholder": "Please enter the full path of the Excel file to be analyzed",
        "analyze_button_text": "Analyze File",
        "no_file_selected": "Please select a valid file path.",
        "file_not_exists": "The file does not exist. Please select again.",
        "analysis_error": "An error occurred while analyzing the file: {}",
        "analysis_complete": "Analysis completed. The results have been saved to {}, and the PSM plot has been saved.",
        "no_save_path_selected": "No save path selected. The results were not saved.",
        "columns_stats": ["Price Point", "Too Cheap Ratio", "Cheap Ratio", "Expensive Ratio", "Too Expensive Ratio"],
        "switch_language_button_text": "Chinese/English",
        'open_excel_button_text': "Example data",
        "column_name_hint": "Column names should include TooCheap_, Cheap_, Expensive_, TooExpensive_",
        "warning_price_order": "Warning: Price points were not sorted, automatically sorted"
    }
}


def find_intersection(x, y1, y2):
    """
    改进版交点检测函数
    采用线性插值 + 边界检测，返回所有交点坐标
    """
    x = np.array(x)
    y1 = np.array(y1)
    y2 = np.array(y2)
    intersections = []

    for i in range(len(x) - 1):
        diff1 = y1[i] - y2[i]
        diff2 = y1[i + 1] - y2[i + 1]

        # 检测符号变化或端点相等
        if diff1 == 0:
            intersections.append(x[i])
        elif diff1 * diff2 < 0:
            # 线性插值求交点
            t = abs(diff1) / (abs(diff1) + abs(diff2))
            x_cross = x[i] + t * (x[i + 1] - x[i])
            intersections.append(x_cross)

    # 去重与排序
    intersections = sorted(list(set([round(v, 2) for v in intersections])))
    return intersections


class PriceSensitivityMeterAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data38.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # === Step 1. 读取 Excel 文件 ===
            df = pd.read_excel(file_path)

            # === Step 2. 检查并设定价格点 ===
            # 如果存在 "Price" 列，就用它作为价格点索引
            if 'Price' in df.columns:
                df = df.set_index('Price')
            elif '价格' in df.columns:
                df = df.set_index('价格')
            else:
                # 如果用户没提供列名，则提示
                raise ValueError("未检测到价格列，请确保 Excel 中包含 'Price' 或 '价格' 列。")

            # 确认索引是数值类型
            if not pd.api.types.is_numeric_dtype(df.index):
                raise ValueError("价格点必须为数值类型（Price 列需为数字）")

            # === Step 3. 检查必要列 ===
            required_columns = ['TooCheap', 'Cheap', 'Expensive', 'TooExpensive']
            if not set(required_columns).issubset(df.columns):
                missing = set(required_columns) - set(df.columns)
                raise ValueError(f"缺少必要的列: {', '.join(missing)}")

            # === Step 4. 检查并排序价格 ===
            sorted_prices = sorted(df.index)
            if not np.array_equal(sorted_prices, df.index.values):
                self.result_label.config(text=languages[self.current_language]["warning_price_order"])
                df = df.reindex(sorted_prices)
            price_points = df.index

            sorted_prices = sorted(df.index)
            if not np.array_equal(sorted_prices, df.index.values):
                self.result_label.config(text=languages[self.current_language]["warning_price_order"])
                df = df.reindex(sorted_prices)
                price_points = df.index  # 重新定义price_points
            else:
                price_points = df.index

            # 检查数据是否为非负整数（计数数据）
            for col in required_columns:
                if not pd.api.types.is_integer_dtype(df[col]) or (df[col] < 0).any():
                    raise ValueError(f"{col}列必须包含非负整数")

            # 计算每个价格点的总受访者数（每行总和）
            total_respondents = df[['TooCheap', 'Cheap', 'Expensive', 'TooExpensive']].sum(axis=1)

            # 计算各选项占比（基于每个价格点的样本量）
            too_cheap_ratio = df['TooCheap'] / total_respondents
            cheap_ratio = df['Cheap'] / total_respondents
            expensive_ratio = df['Expensive'] / total_respondents
            too_expensive_ratio = df['TooExpensive'] / total_respondents

            too_cheap_cum = df['TooCheap'][::-1].cumsum()[::-1] / df['TooCheap'].sum()
            cheap_cum = df['Cheap'][::-1].cumsum()[::-1] / df['Cheap'].sum()
            expensive_cum = df['Expensive'].cumsum() / df['Expensive'].sum()
            too_expensive_cum = df['TooExpensive'].cumsum() / df['TooExpensive'].sum()

            # 计算关键交点（使用改进的累积比例）
            indifference_points = find_intersection(price_points, cheap_cum, expensive_cum)
            optimal_price_points = find_intersection(price_points, too_cheap_cum, too_expensive_cum)
            lower_bounds = find_intersection(price_points, too_cheap_cum, cheap_cum)
            upper_bounds = find_intersection(price_points, expensive_cum, too_expensive_cum)

            # 取第一个有效交点
            indifference_point = indifference_points[0] if indifference_points else None
            optimal_price_point = optimal_price_points[0] if optimal_price_points else None
            lower_bound = lower_bounds[0] if lower_bounds else None
            upper_bound = upper_bounds[0] if upper_bounds else None

            # 绘制 PSM 图（改进版）
            plt.figure(figsize=(12, 8))
            if self.current_language == "zh":
                plt.plot(price_points, too_cheap_cum, label='太便宜', color='green', linewidth=2)
                plt.plot(price_points, cheap_cum, label='便宜', color='blue', linewidth=2)
                plt.plot(price_points, expensive_cum, label='贵', color='orange', linewidth=2)
                plt.plot(price_points, too_expensive_cum, label='太贵', color='red', linewidth=2)
            else:
                plt.plot(price_points, too_cheap_cum, label='Too Cheap', color='green', linewidth=2)
                plt.plot(price_points, cheap_cum, label='Cheap', color='blue', linewidth=2)
                plt.plot(price_points, expensive_cum, label='Expensive', color='orange', linewidth=2)
                plt.plot(price_points, too_expensive_cum, label='Too Expensive', color='red', linewidth=2)

            # 添加可接受价格区间填充
            if lower_bound and upper_bound:
                plt.fill_betweenx([0, 1], lower_bound, upper_bound, color='gray', alpha=0.2,
                                  label='Acceptable Price Range')

            # 改进图表样式
            plt.title(
                '价格敏感度测试（PSM）分析' if self.current_language == "zh" else 'Price Sensitivity Meter (PSM) Analysis',
                fontsize=14,fontweight='bold')
            plt.xlabel('价格' if self.current_language == "zh" else 'Price', fontsize=12)
            plt.ylabel('比例' if self.current_language == "zh" else 'Percentage', fontsize=12)
            plt.ylim(0, 1)
            plt.grid(True, linestyle='--', alpha=0.7)
            plt.legend(fontsize=10)
            plt.tight_layout()

            psm_plot_path = os.path.splitext(file_path)[0] + '_psm_plot.png'
            plt.savefig(psm_plot_path)
            plt.close()

            # 保存结果到 DataFrame
            columns_stats = languages[self.current_language]["columns_stats"]
            data = {
                columns_stats[0]: price_points,
                columns_stats[1]: too_cheap_ratio,
                columns_stats[2]: cheap_ratio,
                columns_stats[3]: expensive_ratio,
                columns_stats[4]: too_expensive_ratio
            }
            result_df = pd.DataFrame(data)

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加标题
                doc.add_heading(
                    '价格敏感度测试（PSM）分析结果' if self.current_language == "zh" else 'Price Sensitivity Meter (PSM) Analysis Results',0)

                # 添加比例数据表
                table = doc.add_table(rows=1, cols=len(result_df.columns))
                hdr_cells = table.rows[0].cells
                for i, col in enumerate(result_df.columns):
                    hdr_cells[i].text = col

                for index, row in result_df.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = f"{value:.4f}" if isinstance(value, float) else str(value)

                # 添加关键统计结果表格
                doc.add_heading("关键统计结果" if self.current_language == "zh" else "Key Statistical Results", level=1)
                stats_table = doc.add_table(rows=1, cols=2)
                stats_hdr = stats_table.rows[0].cells
                stats_hdr[0].text = "指标" if self.current_language == "zh" else "Metric"
                stats_hdr[1].text = "数值" if self.current_language == "zh" else "Value"

                if self.current_language == "zh":
                    stats_data = [
                        ["无差异点 (Indifference Point)", indifference_point],
                        ["最优价格点 (Optimal Price Point)", optimal_price_point],
                        ["价格下限 (Lower Bound)", lower_bound],
                        ["价格上限 (Upper Bound)", upper_bound]
                    ]
                else:
                    stats_data = [
                        ["Indifference Point", indifference_point],
                        ["Optimal Price Point", optimal_price_point],
                        ["Lower Bound", lower_bound],
                        ["Upper Bound", upper_bound]
                    ]

                for label, value in stats_data:
                    cells = stats_table.add_row().cells
                    cells[0].text = label
                    cells[1].text = f"{value:.2f}" if value is not None else "N/A"

                # 添加 PSM 图
                doc.add_picture(psm_plot_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                # 设置 wraplength 属性让文本自动换行
                self.result_label.config(text=languages[self.current_language]["analysis_complete"].format(save_path),
                                         wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]["no_save_path_selected"])

        except ValueError as ve:
            self.result_label.config(text=str(ve))
        except Exception as e:
            self.result_label.config(text=languages[self.current_language]["analysis_error"].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])
        self.column_name_hint_label.config(text=languages[self.current_language]["column_name_hint"])
        self.open_excel_label.config(text=languages[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建列名提示标签
        self.column_name_hint_label = ttk.Label(frame, text=languages[self.current_language]["column_name_hint"],
                                                foreground="gray")
        self.column_name_hint_label.pack(pady=5)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=languages[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame,
                                               text=languages[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = PriceSensitivityMeterAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()
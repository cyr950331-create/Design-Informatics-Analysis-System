import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import os
import tempfile
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
from scipy import stats
import matplotlib.pyplot as plt
import pathlib
from docx import Document
from docx.shared import Pt
from docx.shared import Inches

# 设置支持中文的字体
plt.rcParams["font.family"] = ["SimHei", "WenQuanYi Micro Hei", "Heiti TC", "Arial Unicode MS"]

# 定义语言字典 - 包含所有需要翻译的文本
LANGUAGES = {
    'zh': {
        'title': "Friedman 检验",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "请选择文件。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "中/英",
        'open_excel_button_text': "示例数据",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'explanation': {
            "friedman_test": "用于比较三个或更多相关样本的分布是否有显著差异。",
            "sample_size": "每个样本中的观测值数量。",
            "median": "样本数据的中间值，将数据分为上下两部分。",
            "effect_size": "衡量组间差异的实际显著性，取值范围0-1，值越大表示差异越明显。"
        },
        'interpretation': {
            "statistic": "Friedman检验的统计量值，用于衡量样本之间的差异程度。",
            "p_value": "p值小于显著性水平（通常为0.05）时，拒绝原假设，认为样本之间存在显著差异；否则，接受原假设，认为样本之间无显著差异。",
            "sample_size": "样本量的大小会影响统计检验的功效，较大的样本量通常能提供更准确的结果。",
            "median": "中位数反映了数据的中心位置，可用于比较不同样本的集中趋势。",
            "effect_size": "不受样本量影响，0.1以下为微小效应，0.1-0.3为小效应，0.3-0.5为中等效应，0.5以上为大效应。"
        },
        'results': {
            "friedman_stat": "Friedman检验统计量",
            "p_value": "p值",
            "effect_size": "效应量 (Kendall's W)",
            "effect_interpretation": "效应解释",
            "sample_size": "样本量",
            "median": "中位数"
        },
        'document': {
            "title": "Friedman检验分析结果",
            "boxplot_heading": "数据分布箱线图",
            "boxplot_title": "各组数据分布",
            "boxplot_ylabel": "数值",
            "stats_heading": "统计分析结果",
            "explanation_heading": "统计量解释说明",
            "interpretation_heading": "统计量结果解读",
            "table_headers": ["统计指标", "数值"]
        },
        'effect_levels': {
            "negligible": "微小效应",
            "small": "小效应",
            "medium": "中等效应",
            "large": "大效应"
        }
    },
    'en': {
        'title': "Friedman Test",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Chinese/English",
        'open_excel_button_text': "Example data",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'explanation': {
            "friedman_test": "Used to compare whether the distributions of three or more related samples are significantly different.",
            "sample_size": "The number of observations in each sample.",
            "median": "The middle value of the sample data, dividing the data into two parts.",
            "effect_size": "Measures the practical significance of differences between groups, ranging from 0-1, with larger values indicating more significant differences."
        },
        'interpretation': {
            "statistic": "The test statistic value of the Friedman test, used to measure the degree of difference between samples.",
            "p_value": "When the p-value is less than the significance level (usually 0.05), the null hypothesis is rejected, indicating a significant difference between samples; otherwise, the null hypothesis is accepted, indicating no significant difference.",
            "sample_size": "The sample size affects the power of the statistical test. A larger sample size usually provides more accurate results.",
            "median": "The median reflects the central position of the data and can be used to compare the central tendencies of different samples.",
            "effect_size": "Not affected by sample size. <0.1: negligible effect, 0.1-0.3: small effect, 0.3-0.5: medium effect, >0.5: large effect."
        },
        'results': {
            "friedman_stat": "Friedman Test Statistic",
            "p_value": "p-value",
            "effect_size": "Effect Size (Kendall's W)",
            "effect_interpretation": "Effect Interpretation",
            "sample_size": "Sample Size",
            "median": "Median"
        },
        'document': {
            "title": "Friedman Test Analysis Results",
            "boxplot_heading": "Data Distribution Boxplot",
            "boxplot_title": "Distribution of Groups",
            "boxplot_ylabel": "Value",
            "stats_heading": "Statistical Analysis Results",
            "explanation_heading": "Statistic Explanations",
            "interpretation_heading": "Statistic Interpretations",
            "table_headers": ["Statistical Indicator", "Value"]
        },
        'effect_levels': {
            "negligible": "Negligible effect",
            "small": "Small effect",
            "medium": "Medium effect",
            "large": "Large effect"
        }
    }
}


class FriedmanTestApp:
    def __init__(self, root=None):
        # 当前语言，默认为中文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data23.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            file_path = ""
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 检查数据是否为数值类型
            numerical_df = df.select_dtypes(include=[np.number])
            if numerical_df.empty:
                raise ValueError("数据中没有数值列，无法进行Friedman检验。" if self.current_language == 'zh'
                                 else "There are no numerical columns in the data, Friedman test cannot be performed.")

            # 进行Friedman检验
            stat, p_value = stats.friedmanchisquare(*numerical_df.T.values)

            # 计算效应量Kendall's W
            k = len(numerical_df.columns)  # 组数
            n = len(numerical_df)  # 样本数
            if n * (k - 1) == 0:
                w = 0.0  # 避免除以零
            else:
                w = stat / (n * (k - 1))  # 计算Kendall's W

            # 效应量解释
            def interpret_w(w_val):
                if w_val < 0.1:
                    return LANGUAGES[self.current_language]['effect_levels']['negligible']
                elif w_val < 0.3:
                    return LANGUAGES[self.current_language]['effect_levels']['small']
                elif w_val < 0.5:
                    return LANGUAGES[self.current_language]['effect_levels']['medium']
                else:
                    return LANGUAGES[self.current_language]['effect_levels']['large']

            effect_interpretation = interpret_w(w)

            # 计算样本量和中位数
            sample_sizes = numerical_df.count()
            medians = numerical_df.median()

            # 整理数据（使用语言相关的键）
            stats_data = [
                [LANGUAGES[self.current_language]['results']['friedman_stat'], f"{stat:.4f}"],
                [LANGUAGES[self.current_language]['results']['p_value'], f"{p_value:.4f}"],
                [LANGUAGES[self.current_language]['results']['effect_size'], f"{w:.4f}"],
                [LANGUAGES[self.current_language]['results']['effect_interpretation'], effect_interpretation],
                [LANGUAGES[self.current_language]['results']['sample_size'], str(sample_sizes.to_dict())],
                [LANGUAGES[self.current_language]['results']['median'], str(medians.to_dict())]
            ]

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建一个新的 Word 文档
                doc = Document()

                # 添加标题
                doc.add_heading(LANGUAGES[self.current_language]['document']['title'], 0)

                # 添加箱线图
                doc.add_heading(LANGUAGES[self.current_language]['document']['boxplot_heading'], level=1)

                # 绘制箱线图（使用语言相关的文本）
                plt.figure(figsize=(10, 6))
                numerical_df.boxplot()
                plt.title(LANGUAGES[self.current_language]['document']['boxplot_title'])
                plt.ylabel(LANGUAGES[self.current_language]['document']['boxplot_ylabel'])
                plt.tight_layout()

                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmpfile:
                    plt.savefig(tmpfile, dpi=300, bbox_inches='tight')
                    tmpfile_path = tmpfile.name

                # 将图表添加到Word文档
                doc.add_picture(tmpfile_path, width=Inches(6))
                os.unlink(tmpfile_path)  # 删除临时文件

                # 添加分析结果表格
                doc.add_heading(LANGUAGES[self.current_language]['document']['stats_heading'], level=1)
                # 创建表格（行数=数据行数+表头行）
                table = doc.add_table(rows=1, cols=2)
                hdr_cells = table.rows[0].cells
                # 设置表头（使用语言相关的文本）
                headers = LANGUAGES[self.current_language]['document']['table_headers']
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                    hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
                    hdr_cells[i].paragraphs[0].runs[0].bold = True

                # 添加数据行
                for row_data in stats_data:
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row_data):
                        row_cells[i].text = str(value)
                        row_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
                    # 自动调整列宽
                    for cell in row_cells:
                        for para in cell.paragraphs:
                            para.paragraph_format.space_after = Pt(0)

                # 添加解释说明
                doc.add_heading(LANGUAGES[self.current_language]['document']['explanation_heading'], level=1)
                explanations = LANGUAGES[self.current_language]['explanation']
                doc.add_paragraph(
                    f"{LANGUAGES[self.current_language]['results']['friedman_stat']}: {explanations['friedman_test']}",
                    style='ListBullet')
                doc.add_paragraph(
                    f"{LANGUAGES[self.current_language]['results']['sample_size']}: {explanations['sample_size']}",
                    style='ListBullet')
                doc.add_paragraph(f"{LANGUAGES[self.current_language]['results']['median']}: {explanations['median']}",
                                  style='ListBullet')
                doc.add_paragraph(
                    f"{LANGUAGES[self.current_language]['results']['effect_size']}: {explanations['effect_size']}",
                    style='ListBullet')

                # 添加结果解读
                doc.add_heading(LANGUAGES[self.current_language]['document']['interpretation_heading'], level=1)
                interpretations = LANGUAGES[self.current_language]['interpretation']
                doc.add_paragraph(
                    f"{LANGUAGES[self.current_language]['results']['friedman_stat']}: {interpretations['statistic']}",
                    style='ListBullet')
                doc.add_paragraph(
                    f"{LANGUAGES[self.current_language]['results']['p_value']}: {interpretations['p_value']}",
                    style='ListBullet')
                doc.add_paragraph(
                    f"{LANGUAGES[self.current_language]['results']['sample_size']}: {interpretations['sample_size']}",
                    style='ListBullet')
                doc.add_paragraph(
                    f"{LANGUAGES[self.current_language]['results']['median']}: {interpretations['median']}",
                    style='ListBullet')
                doc.add_paragraph(
                    f"{LANGUAGES[self.current_language]['results']['effect_size']}: {interpretations['effect_size']}",
                    style='ListBullet')

                # 保存 Word 文档
                doc.save(save_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(
                    save_path)
                self.result_label.config(text=result_msg, wraplength=400)
            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        # 切换当前语言
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'
        # 更新所有UI文本
        self.root.title(LANGUAGES[self.current_language]['title'])
        self.select_button.config(text=LANGUAGES[self.current_language]['select_button'])
        self.analyze_button.config(text=LANGUAGES[self.current_language]['analyze_button'])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]['switch_language'])
        self.open_excel_label.config(text=LANGUAGES[self.current_language]["open_excel_button_text"])

        # 保存当前输入框内容（如果不是占位符）
        current_text = self.file_entry.get()
        is_placeholder = current_text == LANGUAGES['zh']['file_entry_placeholder'] or current_text == LANGUAGES['en'][
            'file_entry_placeholder']

        self.file_entry.delete(0, tk.END)
        if is_placeholder:
            self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
            self.file_entry.config(foreground='gray')
        else:
            self.file_entry.insert(0, current_text)
            self.file_entry.config(foreground='black')

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["select_button"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["analyze_button"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=LANGUAGES[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=LANGUAGES[self.current_language]["switch_language"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = FriedmanTestApp()
    app.run()


if __name__ == "__main__":
    run_app()
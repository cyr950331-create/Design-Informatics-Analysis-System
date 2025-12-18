import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
from scipy import stats
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# 设置支持中文的字体
plt.rcParams['font.family'] = 'SimHei'  # 使用黑体字体，可根据系统情况修改
plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

# 定义语言字典 - 包含所有需要翻译的文本
LANGUAGES = {
    'zh': {
        'title': "单样本 Wilcoxon 检验",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "请选择文件。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "中/英",
        'open_excel_button_text': "示例数据",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'hypothesized_value_label': "预设值:",
        'hypothesized_value_placeholder': "请输入假设的中位数（默认为0）",
        'invalid_hypothesized_value': "无效的预设值，请输入数字",
        'stat_labels': {
            "test_statistic": "检验统计量（W）",
            "sample_size": "样本量",
            "valid_sample_size": "有效样本量",
            "median": "中位数",
            "hypothesized_value": "预设值",
            "median_diff": "与预设值的中位数差异",
            "effect_size": "效应量 r",
            "p_value": "p值",
            "ci_95": "中位数差异的95%置信区间"
        },
        'explanation': {
            "test_statistic": "用于检验样本中位数是否与给定的假设中位数存在显著差异。",
            "sample_size": "样本中的观测值数量。",
            "valid_sample_size": "去除缺失值后的观测值数量。",
            "median": "样本数据的中间值，将数据分为上下两部分。",
            "hypothesized_value": "用于与样本数据进行比较的假设值，通常为假设的中位数。",
            "median_diff": "样本中位数与预设值之间的差值，反映实际数据与假设值的偏离程度。",
            "effect_size": "衡量效应大小的指标，值越大表示效应越明显。",
            "p_value": "p值小于显著性水平（通常为0.05）时，拒绝原假设，认为样本中位数与假设中位数存在显著差异；否则，接受原假设，认为无显著差异。",
            "ci_95": "包含真实中位数差异的一个区间，反映了估计的不确定性。"
        },
        'interpretation': {
            "test_statistic": "W值越大，说明样本与假设中位数的差异越明显。",
            "p_value": "用于判断样本与假设中位数之间是否存在显著差异的依据。",
            "sample_size": "样本量的大小会影响统计检验的功效，较大的样本量通常能提供更准确的结果。",
            "valid_sample_size": "实际参与分析的样本数量，排除了缺失值。",
            "median": "中位数反映了数据的中心位置，可用于比较样本与假设中位数的差异。",
            "ci_95": "如果置信区间不包含0，说明样本与假设中位数存在显著差异。",
            "hypothesized_value": "作为比较基准的假设值，是统计检验的参考点。",
            "median_diff": "正值表示样本中位数大于预设值，负值表示样本中位数小于预设值，绝对值越大差异越明显。",
            "effect_size": "通常0.1为小效应，0.3为中等效应，0.5为大效应。"
        },
        'document': {
            "explanation_heading": "解释说明",
            "interpretation_heading": "结果解读",
            "conclusion_heading": "分析结论",
            "charts_heading": "图表",
            "bar_chart_title": "柱状图",
            "error_bar_chart_title": "误差线图",
            "box_plot_title": "数据分布箱线图",
            "qq_plot_title": "Q-Q图",
            "median_label": "中位数",
            "value_label": "数值",
            "hypothesized_line_label": "预设值: {}"
        },
        'charts':{
            'qq_plot_x': '理论分位数',
            'qq_plot_y': '样本分位数',
        },
        'conclusion': {
            "reject_null": "在显著性水平α={}下，拒绝原假设，样本中位数与预设值{}存在显著差异。样本中位数为{:.4f}，与预设值的差异为{:.4f}。",
            "accept_null": "在显著性水平α={}下，未能拒绝原假设，没有足够证据表明样本中位数与预设值{}存在显著差异。样本中位数为{:.4f}，与预设值的差异为{:.4f}。"
        }
    },
    'en': {
        'title': "One Sample Wilcoxon Test",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Chinese/English",
        'open_excel_button_text': "Example data",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'hypothesized_value_label': "Hypothesized Value:",
        'hypothesized_value_placeholder': "Please enter hypothesized median (default 0)",
        'invalid_hypothesized_value': "Invalid hypothesized value, please enter a number",
        'stat_labels': {
            "test_statistic": "Test Statistic (W)",
            "sample_size": "Sample Size",
            "valid_sample_size": "Valid Sample Size",
            "median": "Median",
            "hypothesized_value": "Hypothesized Value",
            "median_diff": "Median Difference from Hypothesized",
            "effect_size": "Effect Size r",
            "p_value": "p-value",
            "ci_95": "95% Confidence Interval for Median Difference"
        },
        'explanation': {
            "test_statistic": "Used to test whether the median of a sample is significantly different from a given hypothesized median.",
            "sample_size": "The number of observations in the sample.",
            "valid_sample_size": "The number of observations after removing missing values.",
            "median": "The middle value of the sample data, dividing the data into two parts.",
            "hypothesized_value": "The hypothesized value used for comparison with sample data, typically the hypothesized median.",
            "median_diff": "The difference between the sample median and the hypothesized value, reflecting the deviation between actual data and the hypothesized value.",
            "effect_size": "A measure of effect size; larger values indicate more obvious effects.",
            "p_value": "When the p-value is less than the significance level (usually 0.05), the null hypothesis is rejected, indicating a significant difference between the sample median and the hypothesized median; otherwise, the null hypothesis is accepted, indicating no significant difference.",
            "ci_95": "An interval containing the true median difference, reflecting the uncertainty of the estimate."
        },
        'interpretation': {
            "test_statistic": "A larger W value indicates a more obvious difference between the sample and the hypothesized median.",
            "p_value": "The basis for determining whether there is a significant difference between the sample and the hypothesized median.",
            "sample_size": "The sample size affects the power of the statistical test. A larger sample size usually provides more accurate results.",
            "valid_sample_size": "The actual number of samples involved in the analysis, excluding missing values.",
            "median": "The median reflects the central position of the data and can be used to compare the difference between the sample and the hypothesized median.",
            "ci_95": "If the confidence interval does not contain 0, it indicates a significant difference between the sample and the hypothesized median.",
            "hypothesized_value": "The hypothesized value used as a comparison benchmark, serving as the reference point for statistical testing.",
            "median_diff": "A positive value indicates the sample median is greater than the hypothesized value, a negative value indicates it is smaller, and a larger absolute value indicates a more significant difference.",
            "effect_size": "Typically, 0.1 is a small effect, 0.3 is a medium effect, and 0.5 is a large effect."
        },
        'document': {
            "explanation_heading": "Explanation",
            "interpretation_heading": "Interpretation",
            "conclusion_heading": "Analysis Conclusion",
            "charts_heading": "Charts",
            "bar_chart_title": "Bar Chart",
            "error_bar_chart_title": "Error Bar Chart",
            "box_plot_title": "Box Plot of Data Distribution",
            "qq_plot_title": "Q-Q Plot",
            "median_label": "Median",
            "value_label": "Value",
            "hypothesized_line_label": "Hypothesized Value: {}"
        },

        'charts':{
            'qq_plot_x': 'Theoretical Quantiles',
            'qq_plot_y': 'Sample Quantiles',
        },
        'conclusion': {
            "reject_null": "At the significance level α={}, the null hypothesis is rejected. There is a significant difference between the sample median and the hypothesized value {}. The sample median is {:.4f}, and the difference from the hypothesized value is {:.4f}.",
            "accept_null": "At the significance level α={}, the null hypothesis cannot be rejected. There is insufficient evidence to indicate a significant difference between the sample median and the hypothesized value {}. The sample median is {:.4f}, and the difference from the hypothesized value is {:.4f}."
        }
    }
}


class OneSampleWilcoxonTestAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为中文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data25.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"{LANGUAGES[self.current_language]['file_not_found']}：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"{LANGUAGES[self.current_language]['analysis_error'].format(str(e))}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event, entry_widget, placeholder):
        if entry_widget.get() == placeholder:
            entry_widget.delete(0, tk.END)
            entry_widget.config(foreground='black')

    def on_focusout(self, event, entry_widget, placeholder):
        if entry_widget.get() == "":
            entry_widget.insert(0, placeholder)
            entry_widget.config(foreground='gray')

    def analyze_file(self):
        lang = LANGUAGES[self.current_language]
        file_path = self.file_entry.get()
        if file_path == lang["file_entry_placeholder"]:
            self.result_label.config(text=lang['file_not_found'])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=lang['file_not_found'])
            return

        # 获取并验证预设值
        hypothesized_value_text = self.hypothesized_value_entry.get()
        if hypothesized_value_text == lang["hypothesized_value_placeholder"]:
            hypothesized_median = 0  # 默认值
        else:
            try:
                hypothesized_median = float(hypothesized_value_text)
            except ValueError:
                self.result_label.config(text=lang['invalid_hypothesized_value'])
                return

        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 检查数据是否为数值类型
            numerical_df = df.select_dtypes(include=[np.number])
            if numerical_df.empty:
                raise ValueError(f"{lang['analysis_error'].format('数据中没有数值列，无法进行单样本Wilcoxon检验。')}")

            # 进行单样本Wilcoxon检验
            statistic, p_value = stats.wilcoxon(numerical_df.squeeze() - hypothesized_median)

            # 提取数据序列
            data_series = numerical_df.squeeze()
            missing_count = data_series.isna().sum()  # 计算缺失值数量

            # 计算统计量
            sample_size = len(data_series)
            valid_sample_size = len(data_series) - missing_count if missing_count > 0 else len(data_series)
            median = np.median(data_series)
            data_diff = data_series - hypothesized_median
            median_diff = np.median(data_diff)

            # 计算中位数差异的置信区间
            def bootstrap_ci(data, func=np.median, alpha=0.05, n_boot=1000):
                boots = np.random.choice(data, size=(n_boot, len(data)), replace=True)
                stats_boot = np.apply_along_axis(func, 1, boots)
                return np.percentile(stats_boot, [100 * alpha / 2, 100 * (1 - alpha / 2)])

            confidence_interval = bootstrap_ci(data_diff)

            # 计算效应量
            effect_size = abs(statistic) / np.sqrt(len(data_series))

            # 整理数据 - 使用语言相关的统计标签
            stats_data = {
                "test_statistic": f"{statistic:.4f}",
                "sample_size": sample_size,
                "valid_sample_size": valid_sample_size,
                "median": f"{median:.4f}",
                "hypothesized_value": f"{hypothesized_median:.4f}",
                "median_diff": f"{median_diff:.4f}",
                "effect_size": f"{effect_size:.4f}",
                "p_value": f"{p_value:.6f}",
                "ci_95": f"[{confidence_interval[0]:.4f}, {confidence_interval[1]:.4f}]"
            }

            # 定义表头（使用语言相关的标签）
            headers = [
                lang['stat_labels']["test_statistic"],
                lang['stat_labels']["sample_size"],
                lang['stat_labels']["valid_sample_size"],
                lang['stat_labels']["median"],
                lang['stat_labels']["hypothesized_value"],
                lang['stat_labels']["median_diff"],
                lang['stat_labels']["effect_size"],
                lang['stat_labels']["p_value"],
                lang['stat_labels']["ci_95"]
            ]

            # 生成单行数据
            data_row = [stats_data[key] for key in stats_data.keys()]

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()
                explanations = lang["explanation"]
                interpretations = lang["interpretation"]
                doc_text = lang["document"]

                # 创建Word表格（1行数据 + 1行表头）
                table = doc.add_table(rows=2, cols=len(headers))

                # 填充表头
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(headers):
                    hdr_cells[col_idx].text = header

                # 填充数据行
                data_cells = table.rows[1].cells
                for col_idx, value in enumerate(data_row):
                    data_cells[col_idx].text = str(value)

                # 调整表格样式
                for row in table.rows:
                    for cell in row.cells:
                        cell.width = Inches(1.2)

                # 添加解释说明
                doc.add_heading(doc_text["explanation_heading"], level=2)
                for key, explanation in explanations.items():
                    doc.add_paragraph(f"{lang['stat_labels'][key]}: {explanation}")

                # 添加分析结果解读
                doc.add_heading(doc_text["interpretation_heading"], level=2)
                for key, interpretation in interpretations.items():
                    doc.add_paragraph(f"{lang['stat_labels'][key]}: {interpretation}")

                # 添加分析结论
                doc.add_heading(doc_text["conclusion_heading"], level=2)
                alpha = 0.05
                if p_value < alpha:
                    conclusion = lang['conclusion']['reject_null'].format(
                        alpha, hypothesized_median, median, median_diff)
                else:
                    conclusion = lang['conclusion']['accept_null'].format(
                        alpha, hypothesized_median, median, median_diff)
                doc.add_paragraph(conclusion)

                # 绘制柱状图
                plt.figure(figsize=(8, 6))
                plt.bar([doc_text["median_label"]], [median])
                plt.axhline(
                    y=hypothesized_median,
                    color='r',
                    linestyle='--',
                    label=doc_text["hypothesized_line_label"].format(hypothesized_median)
                )
                plt.title(doc_text["bar_chart_title"])
                plt.ylabel(doc_text["median_label"])
                plt.legend()
                bar_chart_path = save_path.replace('.docx', '_bar_chart.png')
                plt.savefig(bar_chart_path)
                plt.close()

                # 绘制误差线图
                q1, q3 = np.percentile(data_series.dropna(), [25, 75])
                iqr = q3 - q1
                plt.figure(figsize=(8, 6))
                plt.errorbar(
                    [doc_text["median_label"]],
                    [median],
                    yerr=[iqr / 2],
                    fmt='o'
                )
                plt.axhline(
                    y=hypothesized_median,
                    color='r',
                    linestyle='--',
                    label=doc_text["hypothesized_line_label"].format(hypothesized_median)
                )
                plt.title(doc_text["error_bar_chart_title"])
                plt.ylabel(doc_text["median_label"])
                plt.legend()
                error_bar_path = save_path.replace('.docx', '_error_bar_chart.png')
                plt.savefig(error_bar_path)
                plt.close()

                # 绘制箱线图
                plt.figure(figsize=(8, 6))
                data_series.plot(kind='box')
                plt.axhline(
                    y=hypothesized_median,
                    color='r',
                    linestyle='--',
                    label=doc_text["hypothesized_line_label"].format(hypothesized_median)
                )
                plt.title(doc_text["box_plot_title"])
                plt.ylabel(doc_text["value_label"])
                plt.text(
                    0.95, 0.95,
                    f"n={len(data_series)}",
                    transform=plt.gca().transAxes,
                    ha='right', va='top',
                    bbox=dict(facecolor='white', alpha=0.8)
                )
                plt.legend()
                box_plot_path = save_path.replace('.docx', '_box_plot.png')
                plt.savefig(box_plot_path)
                plt.close()

                # 绘制Q-Q图
                plt.figure(figsize=(8, 6))
                stats.probplot(data_series, plot=plt)
                plt.title(doc_text["qq_plot_title"])
                # 添加X轴和Y轴标签（使用语言相关文本）
                plt.xlabel(lang['charts']['qq_plot_x'])  # 新增X轴标签
                plt.ylabel(lang['charts']['qq_plot_y'])  # 新增Y轴标签
                qq_plot_path = save_path.replace('.docx', '_qq_plot.png')
                plt.savefig(qq_plot_path)
                plt.close()

                # 将图表插入 Word 文档
                doc.add_heading(doc_text["charts_heading"], level=2)
                doc.add_picture(bar_chart_path, width=Inches(6))
                doc.add_picture(error_bar_path, width=Inches(6))
                doc.add_picture(box_plot_path, width=Inches(6))
                doc.add_picture(qq_plot_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                result_msg = lang['analysis_success'].format(save_path)
                self.result_label.config(text=result_msg, wraplength=400)

            else:
                self.result_label.config(text=lang['no_save_path'])

        except Exception as e:
            self.result_label.config(text=lang['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        # 切换当前语言
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        lang = LANGUAGES[self.current_language]

        # 更新界面文字
        self.root.title(lang["title"])
        self.select_button.config(text=lang["select_button"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, lang["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=lang["analyze_button"])
        self.switch_language_label.config(text=lang["switch_language"])
        self.hypothesized_value_label.config(text=lang["hypothesized_value_label"])
        self.hypothesized_value_entry.delete(0, tk.END)
        self.hypothesized_value_entry.insert(0, lang["hypothesized_value_placeholder"])
        self.hypothesized_value_entry.config(foreground='gray')
        self.open_excel_label.config(text=lang["open_excel_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(
            frame,
            text=LANGUAGES[self.current_language]["select_button"],
            command=self.select_file,
            bootstyle=PRIMARY
        )
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind(
            '<FocusIn>',
            lambda e: self.on_entry_click(
                e,
                self.file_entry,
                LANGUAGES[self.current_language]["file_entry_placeholder"]
            )
        )
        self.file_entry.bind(
            '<FocusOut>',
            lambda e: self.on_focusout(
                e,
                self.file_entry,
                LANGUAGES[self.current_language]["file_entry_placeholder"]
            )
        )
        self.file_entry.pack(pady=5)

        # 创建预设值标签和输入框
        hypothesized_frame = ttk.Frame(frame)
        hypothesized_frame.pack(pady=5)

        self.hypothesized_value_label = ttk.Label(
            hypothesized_frame,
            text=LANGUAGES[self.current_language]["hypothesized_value_label"]
        )
        self.hypothesized_value_label.pack(side=LEFT, padx=5)

        self.hypothesized_value_entry = ttk.Entry(hypothesized_frame, width=35)
        self.hypothesized_value_entry.insert(
            0,
            LANGUAGES[self.current_language]["hypothesized_value_placeholder"]
        )
        self.hypothesized_value_entry.config(foreground='gray')
        self.hypothesized_value_entry.bind(
            '<FocusIn>',
            lambda e: self.on_entry_click(
                e,
                self.hypothesized_value_entry,
                LANGUAGES[self.current_language]["hypothesized_value_placeholder"]
            )
        )
        self.hypothesized_value_entry.bind(
            '<FocusOut>',
            lambda e: self.on_focusout(
                e,
                self.hypothesized_value_entry,
                LANGUAGES[self.current_language]["hypothesized_value_placeholder"]
            )
        )
        self.hypothesized_value_entry.pack(side=LEFT)

        # 创建分析按钮
        self.analyze_button = ttk.Button(
            frame,
            text=LANGUAGES[self.current_language]["analyze_button"],
            command=self.analyze_file,
            bootstyle=SUCCESS
        )
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=LANGUAGES[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(
            frame,
            text=LANGUAGES[self.current_language]["switch_language"],
            foreground="gray",
            cursor="hand2"
        )
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = OneSampleWilcoxonTestAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()
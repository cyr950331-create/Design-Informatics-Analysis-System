import tkinter as tk
from tkinter import filedialog
import openpyxl
import os
import pandas as pd
import numpy as np
from statsmodels.stats.outliers_influence import variance_inflation_factor
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
import matplotlib
from docx import Document
from docx.shared import Inches
import statsmodels.api as sm


# 设置支持中文的字体
matplotlib.rcParams['font.family'] = 'SimHei'
# 解决负号显示问题
matplotlib.rcParams['axes.unicode_minus'] = False

# 定义语言字典
languages = {
    "zh": {
        "title": "共线性分析",
        "select_button_text": "选择文件",
        "file_entry_placeholder": "请输入待分析 Excel 文件的完整路径",
        "analyze_button_text": "分析文件",
        "no_file_selected": "请选择文件。",
        "file_not_exists": "文件不存在，请重新选择。",
        "analysis_error": "分析文件时出错: {}",
        "analysis_complete": "分析完成，结果已保存到 {}，相关图片已保存。",
        "no_save_path_selected": "未选择保存路径，结果未保存。",
        "columns_stats": ["变量名", "方差膨胀因子（VIF）", "结果解读"],
        "interpretation_low_vif": "方差膨胀因子（VIF）小于 5，表明该变量与其他变量之间不存在严重的共线性。",
        "interpretation_medium_vif": "方差膨胀因子（VIF）在 5 到 10 之间，表明该变量与其他变量之间可能存在一定的共线性。",
        "interpretation_high_vif": "方差膨胀因子（VIF）大于 10，表明该变量与其他变量之间存在严重的共线性。",
        'open_excel_button_text': "示例数据",
        "switch_language_button_text": "中/英",
        "x_label": "变量名称",
        "y_label": "方差膨胀因子（VIF）",
        "plot_title": "各变量的方差膨胀因子（VIF）"
    },
    "en": {
        "title": "Collinearity Analysis",
        "select_button_text": "Select File",
        "file_entry_placeholder": "Please enter the full path of the Excel file to be analyzed",
        "analyze_button_text": "Analyze File",
        "no_file_selected": "Please select a valid file path.",
        "file_not_exists": "The file does not exist. Please select again.",
        "analysis_error": "An error occurred while analyzing the file: {}",
        "analysis_complete": "Analysis completed. The results have been saved to {}, and the relevant images have been saved.",
        "no_save_path_selected": "No save path selected. The results were not saved.",
        "columns_stats": ["Variable Name", "Variance Inflation Factor (VIF)", "Result Interpretation"],
        "interpretation_low_vif": "The Variance Inflation Factor (VIF) is less than 5, indicating that there is no severe collinearity between this variable and other variables.",
        "interpretation_medium_vif": "The Variance Inflation Factor (VIF) is between 5 and 10, indicating that there may be some collinearity between this variable and other variables.",
        "interpretation_high_vif": "The Variance Inflation Factor (VIF) is greater than 10, indicating that there is severe collinearity between this variable and other variables.",
        'open_excel_button_text': "Example data",
        "switch_language_button_text": "Chinese/English",
        "x_label": "Variable Name",
        "y_label": "Variance Inflation Factor (VIF)",
        "plot_title": "Variance Inflation Factor (VIF) for Each Variable"
    }
}

class CollinearityAnalysisVIFApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"
        
        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])
            
        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data2.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    # 计算方差膨胀因子（VIF）的函数
    def calculate_vif(self, X):
        # 新增：清洗列名（去除特殊字符、空格，转换为小写）
        X.columns = X.columns.str.replace(r'[^\w\s]', '', regex=True)  # 去除特殊字符
        X.columns = X.columns.str.replace(r'\s+', '_', regex=True)  # 空格替换为下划线
        X.columns = X.columns.str.lower()  # 转为小写

        # 去除缺失值并选择数值列
        X = X.select_dtypes(include=[np.number]).dropna()
        if X.empty:
            raise ValueError("数据中不包含可计算 VIF 的数值型变量。")

        # 添加常数列用于计算 VIF
        X = sm.add_constant(X)

        vif_data = pd.DataFrame()
        vif_data["Variable Name"] = X.columns
        vif_data["Variance Inflation Factor (VIF)"] = [
            variance_inflation_factor(X.values, i) for i in range(X.shape[1])
        ]

        # 删除常数列结果
        vif_data = vif_data[vif_data["Variable Name"] != "const"]
        return vif_data

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 假设最后一列为因变量，其余为自变量
            X = df.iloc[:, :-1]

            # 计算方差膨胀因子（VIF）
            vif_data = self.calculate_vif(X)

            # 根据 VIF 值进行结果解读
            interpretations = []
            for vif in vif_data["Variance Inflation Factor (VIF)"]:
                if vif < 5:
                    interpretations.append(languages[self.current_language]["interpretation_low_vif"])
                elif 5 <= vif < 10:
                    interpretations.append(languages[self.current_language]["interpretation_medium_vif"])
                else:
                    interpretations.append(languages[self.current_language]["interpretation_high_vif"])

            vif_data["Result Interpretation"] = interpretations

            # 绘制 VIF 值的柱状图
            plt.figure(figsize=(8, 5))
            vif_data_sorted = vif_data.sort_values(by="Variance Inflation Factor (VIF)", ascending=False)
            plt.bar(vif_data_sorted["Variable Name"], vif_data_sorted["Variance Inflation Factor (VIF)"])
            plt.xlabel(languages[self.current_language]["x_label"])  # 变量名称
            plt.ylabel(languages[self.current_language]["y_label"])  # 方差膨胀因子（VIF）
            plt.title(languages[self.current_language]["plot_title"])  # 各变量的方差膨胀因子（VIF）
            plt.xticks(rotation=60, ha='right')
            plt.tight_layout()

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 获取保存目录和文件名（不含扩展名）
                save_dir = os.path.dirname(save_path)
                file_name = os.path.splitext(os.path.basename(save_path))[0]

                # 生成图片路径（与Word文件同目录）
                if self.current_language == "en":
                    image_path = os.path.join(save_dir, f"{file_name}_vif_plot_en.png")
                else:
                    image_path = os.path.join(save_dir, f"{file_name}_vif_plot.png")
                plt.savefig(image_path, bbox_inches='tight')
                plt.close()

                # 创建一个新的 Word 文档
                doc = Document()

                # 添加标题
                doc.add_heading('共线性分析 (VIF) 结果', 0)

                # 添加表格
                table = doc.add_table(rows=1, cols=len(languages[self.current_language]["columns_stats"]))
                hdr_cells = table.rows[0].cells
                for i, col in enumerate(languages[self.current_language]["columns_stats"]):
                    hdr_cells[i].text = col

                # 添加数据到表格
                for index, row in vif_data.iterrows():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(row['Variable Name'])
                    row_cells[1].text = str(round(row['Variance Inflation Factor (VIF)'], 3))
                    row_cells[2].text = str(row['Result Interpretation'])

                # 添加图片到文档
                doc.add_picture(image_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                # 设置 wraplength 属性让文本自动换行
                self.result_label.config(text=languages[self.current_language]["analysis_complete"].format(save_path), wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]["no_save_path_selected"])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]["analysis_error"].format(str(e)))

    def switch_language(self, event=None):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])
        self.open_excel_label.config(text=languages[self.current_language]["open_excel_button_text"])
        
    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"], 
                                       command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"], 
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=languages[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=languages[self.current_language]["switch_language_button_text"],
                                              foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)
        
    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = CollinearityAnalysisVIFApp()
    app.run()

if __name__ == "__main__":
    run_app()
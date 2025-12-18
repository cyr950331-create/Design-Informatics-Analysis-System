import tkinter as tk
from tkinter import filedialog, simpledialog, messagebox
import os
import numpy as np
import pandas as pd
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
import statsmodels.api as sm
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典（保持原样）
languages = {
    'zh': {
        'title': "调节中介作用",
        'select_button_text': "选择文件",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'analyze_button_text': "分析文件",
        'no_file_selected': "请选择文件。",
        'file_not_exists': "文件不存在，请重新选择。",
        'analysis_error': "分析文件时出错: {}",
        'analysis_complete': "分析完成，结果已保存到 {}，相关图片已保存。",
        'no_save_path_selected': "未选择保存路径，结果未保存。",
        'open_excel_button_text': "示例数据",
        'switch_language_button_text': "中/英",
        'input_info': "输入信息",
        'input_ind_var': "请输入自变量的列名",
        'input_med_var': "请输入中介变量的列名（多个用逗号分隔）",
        'input_dep_var': "请输入因变量的列名",
        'input_mod_var': "请输入调节变量的列名",
        'input_incomplete': "未输入完整的变量名，分析取消。",
        'explanation': {
            "自变量对因变量的总效应": "自变量直接对因变量产生的影响。",
            "自变量对中介变量的效应": "自变量对中介变量产生的影响。",
            "中介变量对因变量的效应（控制自变量）": "在控制自变量的情况下，中介变量对因变量产生的影响。",
            "调节变量对自变量 - 中介变量关系的调节效应": "调节变量对自变量与中介变量之间关系的影响。",
            "调节变量对中介变量 - 因变量关系的调节效应": "调节变量对中介变量与因变量之间关系的影响。",
            "中介效应": "自变量通过中介变量对因变量产生的间接影响。",
            "样本量": "参与分析的样本数量。"
        },
        'interpretation': {
            "自变量对因变量的总效应": "总效应显著表示自变量对因变量有直接影响。",
            "自变量对中介变量的效应": "该效应显著表示自变量能够影响中介变量。",
            "中介变量对因变量的效应（控制自变量）": "此效应显著表示中介变量在控制自变量后仍对因变量有影响。",
            "调节变量对自变量 - 中介变量关系的调节效应": "该调节效应显著表示调节变量会影响自变量与中介变量之间的关系。",
            "调节变量对中介变量 - 因变量关系的调节效应": "该调节效应显著表示调节变量会影响中介变量与因变量之间的关系。",
            "中介效应": "中介效应显著表示自变量通过中介变量对因变量产生了间接影响。",
            "样本量": "样本量的大小会影响统计结果的可靠性，较大的样本量通常能提供更可靠的结果。"
        },
        'level_terms': {'low': '低水平', 'mean': '中水平', 'high': '高水平'},
        'stat_terms': {'point_estimate': '点估计'}
    },
    'en': {
        'title': "Moderated Mediation",
        'select_button_text': "Select File",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'analyze_button_text': "Analyze File",
        'no_file_selected': "Please select a valid file path.",
        'file_not_exists': "The file does not exist. Please select again.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'analysis_complete': "Analysis completed. The results have been saved to {}, and the relevant images have been saved.",
        'no_save_path_selected': "No save path selected. The results were not saved.",
        'open_excel_button_text': "Example data",
        'switch_language_button_text': "Chinese/English",
        'input_info': "Input Information",
        'input_ind_var': "Please enter the column name of the independent variable",
        'input_med_var': "Please enter column names of mediator variables (separated by commas)",
        'input_dep_var': "Please enter the column name of the dependent variable",
        'input_mod_var': "Please enter the column name of the moderator variable",
        'input_incomplete': "Incomplete variable names entered, analysis canceled.",
        'explanation': {
            "Total Effect of Independent Variable on Dependent Variable": "The total effect of the independent variable on the dependent variable.",
            "Effect of Independent Variable on Mediator Variable": "The effect of the independent variable on the mediator variable.",
            "Effect of Mediator Variable on Dependent Variable (Controlling for Independent Variable)": "The effect of the mediator variable on the dependent variable while controlling for the independent variable.",
            "Moderating Effect of Moderator on IV-MV Relationship": "The moderating effect of the moderator variable on the relationship between the independent variable and the mediator variable.",
            "Moderating Effect of Moderator on MV-DV Relationship": "The moderating effect of the moderator variable on the relationship between the mediator variable and the dependent variable.",
            "Mediation Effect": "The indirect effect of the independent variable on the dependent variable through the mediator variable.",
            "Sample Size": "The number of samples involved in the analysis."
        },
        'interpretation': {
            "Total Effect of Independent Variable on Dependent Variable": "A significant total effect indicates that the independent variable has a direct impact on the dependent variable.",
            "Effect of Independent Variable on Mediator Variable": "A significant effect indicates that the independent variable can influence the mediator variable.",
            "Effect of Mediator Variable on Dependent Variable (Controlling for Independent Variable)": "A significant effect indicates that the mediator variable still has an impact on the dependent variable after controlling for the independent variable.",
            "Moderating Effect of Moderator on IV-MV Relationship": "A significant moderating effect indicates that the moderator variable affects the relationship between the independent variable and the mediator variable.",
            "Moderating Effect of Moderator on MV-DV Relationship": "A significant moderating effect indicates that the moderator variable affects the relationship between the mediator variable and the dependent variable.",
            "Mediation Effect": "A significant mediation effect indicates that the independent variable has an indirect impact on the dependent variable through the mediator variable.",
            "Sample Size": "The sample size affects the reliability of the statistical results. A larger sample size usually provides more reliable results."
        },
        'level_terms': {'low': 'Low Level', 'mean': 'Mean Level', 'high': 'High Level'},
        'stat_terms': {'point_estimate': 'Point Estimate'}
    }
}


class ModeratedMediationAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data10.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")


    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def moderated_mediation_analysis(self, data, ind_var, med_vars, dep_var, mod_var):
        """
        说明：
        - 对自变量（ind_var）、中介变量（每个 med_var）与调节变量（mod_var）进行中心化（_c 后缀）。
        - 在回归中使用中心化变量并构造交互项 ind_mod_c, med_mod_c。
        - 计算各路径系数及 p 值。
        - 计算条件间接效应（W = Mean ± 1SD）。
        - 使用 bootstrap（在每次重采样中拟合带交互项模型）估计每个 W 水平下的间接效应的 95% CI。
        """
        results = []
        total_effect = None
        p_value_total = None

        # 数据清洗与中心化
        # 只保留需要的列（避免意外列）
        cols_needed = [ind_var, dep_var, mod_var] + med_vars
        data = data[cols_needed].copy()

        # 丢弃含缺失值的行（你也可以选择插补策略，这里先删除以确保稳健）
        data = data.dropna(axis=0, subset=cols_needed)

        # 强制转换为数值（有非数值会抛出异常）
        for c in cols_needed:
            data[c] = pd.to_numeric(data[c], errors='raise')

        # 中心化（mean-centering）
        data[ind_var + '_c'] = data[ind_var] - data[ind_var].mean()
        data[mod_var + '_c'] = data[mod_var] - data[mod_var].mean()
        for m in med_vars:
            data[m + '_c'] = data[m] - data[m].mean()

        sample_size = len(data)

        # 第一步：总效应（使用中心化后的自变量）
        iv = ind_var + '_c'
        # 使用常数项和中心化自变量计算总效应
        X1 = sm.add_constant(data[[iv]])
        model1 = sm.OLS(data[dep_var], X1).fit()
        total_effect = model1.params.get(iv, np.nan)
        p_value_total = model1.pvalues.get(iv, np.nan)

        # 对每个中介逐一分析
        for med_var in med_vars:
            mv = med_var + '_c'
            wv = mod_var + '_c'
            iv = ind_var + '_c'  # 保证局部一致

            #  第二步：自变量对中介的效应（a 路径）
            X2 = sm.add_constant(data[[iv]])
            model2 = sm.OLS(data[mv], X2).fit()
            effect_ind_med = model2.params.get(iv, np.nan)
            p_value_ind_med = model2.pvalues.get(iv, np.nan)

            # 第三步：调节对自变量-中介关系的影响（a 路径是否受 W 调节）
            data['ind_mod_c'] = data[iv] * data[wv]
            X3 = sm.add_constant(data[[iv, wv, 'ind_mod_c']])
            model3 = sm.OLS(data[mv], X3).fit()
            effect_mod_ind_med = model3.params.get('ind_mod_c', 0.0)
            p_value_mod_ind_med = model3.pvalues.get('ind_mod_c', np.nan)

            # 第四步：中介对因变量的效应（b 路径），控制自变量与调节变量主效应
            X4 = sm.add_constant(data[[iv, mv, wv]])
            model4 = sm.OLS(data[dep_var], X4).fit()
            effect_med_dep = model4.params.get(mv, np.nan)
            p_value_med_dep = model4.pvalues.get(mv, np.nan)

            # 第五步：调节对中介-因变量关系的影响（b 路径是否受 W 调节）
            data['med_mod_c'] = data[mv] * data[wv]
            X5 = sm.add_constant(data[[iv, mv, wv, 'med_mod_c']])
            model5 = sm.OLS(data[dep_var], X5).fit()
            effect_mod_med_dep = model5.params.get('med_mod_c', 0.0)
            p_value_mod_med_dep = model5.pvalues.get('med_mod_c', np.nan)

            # 第六步：点估计的中介效应（简单乘积）
            # 这是在“平均”条件下（若使用中心化变量，a*b 可代表在均值处的间接效应）
            mediation_effect = effect_ind_med * effect_med_dep

            # 在不同调节水平下计算条件间接效应（W = Mean ± 1SD）
            mean_W = data[mod_var].mean()  # 注意是未中心化的均值（用于直观意义），但 calculations 使用 center-level value
            std_W = data[mod_var].std()
            # 这里为了和中心化后的系数匹配，我们需要把未中心化的 W 转换成中心化值（w_value - mean）
            # 但 data[wv] 是已经中心化（w - mean），因此当我们想要在 "原始均值 ±1SD" 下计算 a(w) * b(w)，
            # 我们需要把这些原始值减去 mean，得到中心化值，再代入公式。
            moderation_levels_raw = {
                "低水平 (-1SD)": mean_W - std_W,
                "中水平 (Mean)": mean_W,
                "高水平 (+1SD)": mean_W + std_W
            }
            # 转换为中心化值
            moderation_levels = {}
            for k, raw_val in moderation_levels_raw.items():
                moderation_levels[k] = raw_val - mean_W  # center value

            indirect_effects_by_W = {}
            for level_name, w_centered in moderation_levels.items():
                # a(W) = a + a_mod * W_centered
                a_w = effect_ind_med + effect_mod_ind_med * w_centered
                # b(W) = b + b_mod * W_centered
                b_w = effect_med_dep + effect_mod_med_dep * w_centered
                indirect_effects_by_W[level_name] = a_w * b_w

            # Bootstrap（在每次重采样中拟合含交互项的模型）
            n_boot = 1000  # 可调整为 500 以加速测试
            rng = np.random.default_rng(42)
            # 为每个 level 存储 bootstrap 样本的间接效应
            boot_indirects = {level: [] for level in moderation_levels.keys()}
            boot_failures = 0

            # 增加进度提示变量
            progress_interval = max(1, n_boot // 10)  # 每10%进度提示一次

            for i in range(n_boot):
                # 每完成10%的抽样就更新一次进度
                if i % progress_interval == 0:
                    progress = (i / n_boot) * 100
                    self.result_label.config(
                        text=f"分析中：Bootstrap抽样完成 {progress:.0f}%..." if self.current_language == 'zh'
                        else f"Analyzing: Bootstrap sampling {progress:.0f}% complete..."
                    )
                    self.root.update_idletasks()  # 刷新界面

                try:
                    sample = data.sample(frac=1, replace=True,
                                         random_state=rng.integers(0, 2 ** 31 - 1)).copy()
                    # 重新构造必要列（中心化列在 data 已存在）
                    # 使用中心化列名：iv, mv, wv
                    sample['ind_mod_c'] = sample[iv] * sample[wv]
                    sample['med_mod_c'] = sample[mv] * sample[wv]

                    # a-path：X -> M (含X, W, X*W)
                    Xa = sm.add_constant(sample[[iv, wv, 'ind_mod_c']])
                    model_a = sm.OLS(sample[mv], Xa).fit()
                    a_coef = model_a.params.get(iv, 0.0)
                    a_mod_coef = model_a.params.get('ind_mod_c', 0.0)

                    # b-path：M -> Y (含 M, W, M*W，同时控制 X)
                    Xb = sm.add_constant(sample[[iv, mv, wv, 'med_mod_c']])
                    model_b = sm.OLS(sample[dep_var], Xb).fit()
                    b_coef = model_b.params.get(mv, 0.0)
                    b_mod_coef = model_b.params.get('med_mod_c', 0.0)

                    # 计算每个 W 水平的间接效应
                    for level_name, w_centered in moderation_levels.items():
                        indirect = (a_coef + a_mod_coef * w_centered) * (b_coef + b_mod_coef * w_centered)
                        boot_indirects[level_name].append(indirect)
                except Exception:
                    # 某些重采样可能因为共线性或其他数值问题失败，记录并继续
                    boot_failures += 1
                    continue

                # 抽样完成后更新提示：告知用户正在整理结果
                self.result_label.config(
                    text="Bootstrapstrap抽样完成，正在整理结果..." if self.current_language == 'zh'
                    else "Bootstrapstrap sampling completed, organizing results..."
                )
                self.root.update_idletasks()  # 立即刷新界面显示

            # 计算 95% CI（如果可用）
            boot_CI_by_W = {}
            for level_name, arr in boot_indirects.items():
                if len(arr) > 0:
                    lower, upper = np.percentile(arr, [2.5, 97.5])
                else:
                    lower, upper = np.nan, np.nan
                boot_CI_by_W[level_name] = (lower, upper)

            # 保存本中介变量的所有结果
            results.append({
                'med_var': med_var,
                'total_effect': total_effect,
                'p_value_total': p_value_total,
                'effect_ind_med': effect_ind_med,
                'p_value_ind_med': p_value_ind_med,
                'effect_mod_ind_med': effect_mod_ind_med,
                'p_value_mod_ind_med': p_value_mod_ind_med,
                'effect_med_dep': effect_med_dep,
                'p_value_med_dep': p_value_med_dep,
                'effect_mod_med_dep': effect_mod_med_dep,
                'p_value_mod_med_dep': p_value_mod_med_dep,
                'mediation_effect': mediation_effect,
                'indirect_effects_by_W': indirect_effects_by_W,
                'boot_CI_by_W': boot_CI_by_W,
                'boot_failures': boot_failures
            })

            # 在控制台打印该中介的关键结果（便于调试）
            print(f"\n== 中介变量 [{med_var}] 分析结果 ==")
            print(f"样本量: {sample_size}")
            print(f"总效应 (X->{dep_var}) 点估计: {total_effect:.4f}, p={p_value_total:.4f}")
            print(f"a (X->M) 点估计: {effect_ind_med:.4f}, p={p_value_ind_med:.4f}")
            print(f"b (M->Y) 点估计: {effect_med_dep:.4f}, p={p_value_med_dep:.4f}")
            print("不同调节水平下的间接效应（点估计）:")
            for level_name, val in indirect_effects_by_W.items():
                ci = boot_CI_by_W.get(level_name, (np.nan, np.nan))
                print(f"  {level_name}: 点估计={val:.4f}, 95% CI=[{ci[0]:.4f}, {ci[1]:.4f}]")
            if boot_failures:
                print(f"注意：bootstrap 中有 {boot_failures} 次重采样失败（被跳过）。")

        return results, total_effect, p_value_total, sample_size

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 让用户输入变量名
            ind_var = simpledialog.askstring(
                languages[self.current_language]['input_info'],
                languages[self.current_language]['input_ind_var'],
                parent=self.root  # 添加父窗口参数
            )
            med_var_input = simpledialog.askstring(
                languages[self.current_language]['input_info'],
                languages[self.current_language]['input_med_var'],
                parent=self.root  # 添加父窗口参数
            )
            dep_var = simpledialog.askstring(
                languages[self.current_language]['input_info'],
                languages[self.current_language]['input_dep_var'],
                parent=self.root  # 添加父窗口参数
            )
            mod_var = simpledialog.askstring(
                languages[self.current_language]['input_info'],
                languages[self.current_language]['input_mod_var'],
                parent=self.root  # 添加父窗口参数
            )

            # 处理中介变量（去空格并分割）
            if not all([ind_var, med_var_input, dep_var, mod_var]):
                self.result_label.config(text=languages[self.current_language]['input_incomplete'])
                return

            med_vars = [med.strip() for med in med_var_input.split(',') if med.strip()]
            if not med_vars:
                self.result_label.config(text=languages[self.current_language]['input_incomplete'])
                return

            self.result_label.config(
                text="正在进行分析，请稍候..." if self.current_language == 'zh'
                else "Analyzing, please wait..."
            )
            self.root.update_idletasks()  # 刷新界面显示提示信息

            # 进行调节中介作用分析
            results, total_effect, p_value_total, sample_size = self.moderated_mediation_analysis(
                df, ind_var, med_vars, dep_var, mod_var)

            # 整理数据（用于表格）
            data = []
            # 添加总效应（只显示一次）
            data.append([
                "自变量对因变量的总效应" if self.current_language == 'zh' else "Total Effect of Independent Variable on Dependent Variable",
                total_effect,
                p_value_total
            ])

            # 添加每个中介变量的分析结果
            for result in results:
                med_var = result['med_var']
                data.append([
                    f"自变量对中介变量 [{med_var}] 的效应" if self.current_language == 'zh'
                    else f"Effect of Independent Variable on Mediator [{med_var}]",
                    result['effect_ind_med'],
                    result['p_value_ind_med']
                ])
                data.append([
                    f"调节变量对自变量 - 中介变量 [{med_var}] 关系的调节效应" if self.current_language == 'zh'
                    else f"Moderating Effect on IV-MV [{med_var}] Relationship",
                    result['effect_mod_ind_med'],
                    result['p_value_mod_ind_med']
                ])
                data.append([
                    f"中介变量 [{med_var}] 对因变量的效应（控制自变量）" if self.current_language == 'zh'
                    else f"Effect of Mediator [{med_var}] on Dependent Variable (Controlling IV)",
                    result['effect_med_dep'],
                    result['p_value_med_dep']
                ])
                data.append([
                    f"调节变量对中介变量 [{med_var}] - 因变量关系的调节效应" if self.current_language == 'zh'
                    else f"Moderating Effect on MV [{med_var}]-DV Relationship",
                    result['effect_mod_med_dep'],
                    result['p_value_mod_med_dep']
                ])
                data.append([
                    f"中介变量 [{med_var}] 的中介效应 (点估计)" if self.current_language == 'zh'
                    else f"Mediation Effect of [{med_var}] (Point Estimate)",
                    result['mediation_effect'],
                    ""
                ])
                # 也可在 Word 中输出间接效应的点估计与 CI（见下）

            data.append([
                "样本量" if self.current_language == 'zh' else "Sample Size", sample_size, ""])

            if self.current_language == 'zh':
                headers = ["统计量", "统计量值", "p值"]
            else:
                headers = ["Statistic", "Statistic Value", "p-value"]
            df_result = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            explanations = languages[self.current_language]['explanation']
            interpretations = languages[self.current_language]['interpretation']

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(
                defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加标题
                doc.add_heading(
                    '调节中介作用分析结果' if self.current_language == 'zh' else 'Moderated Mediation Analysis Results',0)

                # 添加分析结果表格
                table = doc.add_table(rows=1, cols=len(headers))
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                for index, row in df_result.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)

                # 将按 W 水平的点估计与 Bootstrap CI 写入 Word
                doc.add_heading(
                    '不同调节水平下的间接效应（包含 95% Bootstrap CI）' if self.current_language == 'zh'
                    else 'Indirect Effects at Different Moderator Levels (with 95% Bootstrap CI)',
                    level=1
                )

                # 获取当前语言的术语（提前获取，避免重复查询）
                lang_levels = languages[self.current_language]['level_terms']
                lang_stats = languages[self.current_language]['stat_terms']

                for result in results:
                    med_var = result['med_var']
                    # 中介变量标题（已处理，保持不变）
                    doc.add_paragraph(
                        f"中介变量 [{med_var}]：" if self.current_language == 'zh'
                        else f"Mediator Variable [{med_var}]:"
                    )

                    # 循环生成各水平的效应描述
                    for level_text, point_est in result['indirect_effects_by_W'].items():
                        ci = result.get('boot_CI_by_W', {}).get(level_text, (np.nan, np.nan))

                        # 关键：从原始level_text中识别水平类型（解决中文键问题）
                        if '低' in level_text:
                            level_type = 'low'
                            sd_tag = '(-1SD)'
                        elif '中' in level_text or 'Mean' in level_text:
                            level_type = 'mean'
                            sd_tag = '(Mean)'
                        elif '高' in level_text:
                            level_type = 'high'
                            sd_tag = '(+1SD)'
                        else:
                            level_type = 'unknown'
                            sd_tag = ''

                        # 根据level_key（如'low'/'mean'/'high'）匹配对应语言的水平名称
                        level_name = lang_levels.get(level_type, level_text)   # 默认为原key，防止匹配失败

                        # 根据语言生成完整文本
                        if self.current_language == 'zh':
                            para_text = f"    {level_name} {sd_tag}: {lang_stats['point_estimate']}={point_est:.4f}，95% CI=[{ci[0]:.4f}, {ci[1]:.4f}]"
                        else:
                            para_text = f"    {level_name} {sd_tag}: {lang_stats['point_estimate']}={point_est:.4f}, 95% CI=[{ci[0]:.4f}, {ci[1]:.4f}]"

                        doc.add_paragraph(para_text)
                # 添加解释说明
                doc.add_heading(
                    '统计量解释说明' if self.current_language == 'zh' else 'Statistical Explanations', level=1)
                for key, value in explanations.items():
                    doc.add_paragraph(f"{key}: {value}")

                # 添加分析结果解读
                doc.add_heading(
                    '统计量结果解读' if self.current_language == 'zh' else 'Statistical Interpretations', level=1)
                for key, value in interpretations.items():
                    doc.add_paragraph(f"{key}: {value}")

                # 生成图片（中介效应柱状图）
                fig, ax = plt.subplots(figsize=(12, 6))
                effects = []
                labels = []

                # 添加总效应
                effects.append(total_effect)
                labels.append("自变量对因变量总效应" if self.current_language == 'zh' else "Total Effect of Independent on Dependent")

                # 添加每个中介变量的效应（按顺序：a, a_mod, b, b_mod, indirect）
                p_values = [p_value_total]  # 用于后面在图上标注 p 值
                for result in results:
                    med_var = result['med_var']
                    effects.extend([
                        result['effect_ind_med'],
                        result['effect_mod_ind_med'],
                        result['effect_med_dep'],
                        result['effect_mod_med_dep'],
                        result['mediation_effect']
                    ])

                    if self.current_language == 'zh':
                        labels.extend([
                            f"自变量对中介[{med_var}]效应",
                            f"调节对自变量-中介[{med_var}]效应",
                            f"中介[{med_var}]对因变量效应",
                            f"调节对中介[{med_var}]-因变量效应",
                            f"中介[{med_var}]的中介效应"
                        ])
                    else:
                        labels.extend([
                            f"Ind on Mediator [{med_var}]",
                            f"Mod on Ind-Mediator [{med_var}]",
                            f"Mediator [{med_var}] on Dep",
                            f"Mod on Mediator [{med_var}]-Dep",
                            f"Mediation by [{med_var}]"
                        ])

                    # 对应 p 值顺序要和上面 effects 顺序一致
                    p_values.extend([
                        result.get('p_value_ind_med', np.nan),
                        result.get('p_value_mod_ind_med', np.nan),
                        result.get('p_value_med_dep', np.nan),
                        result.get('p_value_mod_med_dep', np.nan),
                        np.nan  # 中介效应点估计没有直接的 p 值（使用 Bootstrap CI 判定显著性）
                    ])

                ax.bar(labels, effects)
                ax.set_title('调节中介作用分析结果' if self.current_language == 'zh' else 'Moderated Mediation Analysis Results')
                ax.set_ylabel('效应值' if self.current_language == 'zh' else 'Effect Value')
                plt.xticks(rotation=45, ha='right')

                # 在每根柱上标注 p 值（若存在）
                for i, (v, p) in enumerate(zip(effects, p_values)):
                    if not (p is None or (isinstance(p, float) and np.isnan(p))):
                        ax.text(i, v, f"p={p:.3f}", ha='center', va='bottom', fontsize=8)

                plt.tight_layout()

                # 保存图片
                img_path = os.path.splitext(save_path)[0] + '.png'
                plt.savefig(img_path)
                plt.close()

                # 将图片插入 Word 文档
                doc.add_heading(
                    '中介效应柱状图' if self.current_language == 'zh' else 'Mediation Effect Bar Chart', level=1)
                doc.add_picture(img_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                self.result_label.config(text=languages[self.current_language]['analysis_complete'].format(save_path),
                                         wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]['no_save_path_selected'])

        except Exception as e:
            # 显示更友好的错误信息
            self.result_label.config(text=languages[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])
        self.open_excel_label.config(text=languages[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个主框架来包含所有控件
        main_frame = ttk.Frame(self.root)
        main_frame.pack(expand=True)  # 让框架在窗口中居中

        # 创建文件选择按钮
        self.select_button = ttk.Button(main_frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(main_frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(main_frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            main_frame,
            text=languages[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建语言切换标签（放入主框架中）
        self.switch_language_label = ttk.Label(main_frame,
                                               text=languages[self.current_language]["switch_language_button_text"],
                                               cursor="hand2", foreground="gray")  # 改为蓝色更像可点击链接
        self.switch_language_label.pack(pady=10)
        self.switch_language_label.bind("<Button-1>", self.switch_language)

        # 创建结果显示标签
        self.result_label = ttk.Label(main_frame, text="", wraplength=400)  # 修正这里的错误
        self.result_label.pack(pady=10)


if __name__ == "__main__":
    try:
        app = ModeratedMediationAnalysisApp()
        app.root.mainloop()
    except KeyboardInterrupt:
        print("程序已被用户中断。")

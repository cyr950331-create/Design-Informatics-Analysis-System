import tkinter as tk
from tkinter import filedialog
import openpyxl
import os
import pandas as pd
import numpy as np
from scipy import stats
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
import matplotlib
from docx import Document
from docx.shared import Inches

# 设置支持中文的字体
matplotlib.rcParams['font.family'] = 'SimHei'
# 解决负号显示问题
matplotlib.rcParams['axes.unicode_minus'] = False

# 定义语言字典
languages = {
    "zh": {
        "title": "偏相关分析",
        "select_button_text": "选择文件",
        "file_entry_placeholder": "请输入待分析 Excel 文件的完整路径",
        "analyze_button_text": "分析文件",
        "no_file_selected": "请选择文件。",
        "file_not_exists": "请选择文件。",
        "analysis_error": "分析文件时出错: {}",
        "analysis_complete": "分析完成，结果已保存到 {}，相关图片已保存。",
        "no_save_path_selected": "未选择保存路径，结果未保存。",
        "columns_stats": ["变量对", "偏相关系数", "p 值"],
        "interpretation_low_p": "p 值小于 0.05，表明该变量对之间的偏相关性显著。",
        "interpretation_high_p": "p 值大于等于 0.05，表明该变量对之间的偏相关性不显著。",
        'open_excel_button_text': "示例数据",
        "switch_language_button_text": "中/英",
        'partial_corr': '偏相关系数',
        'degree_of_freedom': '自由度'
    },
    "en": {
        "title": "Partial Correlation Analysis",
        "select_button_text": "Select File",
        "file_entry_placeholder": "Please enter the full path of the Excel file to be analyzed",
        "analyze_button_text": "Analyze File",
        "no_file_selected": "Please select a valid file path.",
        "file_not_exists": "The file does not exist. Please select again.",
        "analysis_error": "An error occurred while analyzing the file: {}",
        "analysis_complete": "Analysis completed. The results have been saved to {}, and the relevant images have been saved.",
        "no_save_path_selected": "No save path selected. The results were not saved.",
        "columns_stats": ["Variable Pair", "Partial Correlation Coefficient", "p-value", "Result Interpretation"],
        "interpretation_low_p": "The p-value is less than 0.05, indicating that the partial correlation between this variable pair is significant.",
        "interpretation_high_p": "The p-value is greater than or equal to 0.05, indicating that the partial correlation between this variable pair is not significant.",
        'open_excel_button_text': "Example data",
        "switch_language_button_text": "Chinese/English",
        'partial_corr': 'Partial correlation coefficient',
        'degree_of_freedom': 'Degrees of freedom'
    }
}


class PartialCorrelationAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data20.xls")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    # 计算偏相关系数的函数
    def partial_corr(self, data, x, y, z):

        sub_data = data[[x, y] + z].dropna()
        n = len(sub_data)

        # 当控制变量存在时，需要正确构建回归模型
        if z:
            # 对x进行回归：控制z变量
            X = sub_data[z]
            X = np.column_stack((np.ones(n), X))  # 添加常数项
            y_x = sub_data[x]
            beta_x = np.linalg.lstsq(X, y_x, rcond=None)[0]
            resid_x = y_x - np.dot(X, beta_x)

            # 对y进行回归：控制z变量
            y_y = sub_data[y]
            beta_y = np.linalg.lstsq(X, y_y, rcond=None)[0]
            resid_y = y_y - np.dot(X, beta_y)
        else:
            # 无控制变量时退化为简单相关
            resid_x = sub_data[x] - sub_data[x].mean()
            resid_y = sub_data[y] - sub_data[y].mean()

        # 计算残差的相关系数
        corr, p = stats.pearsonr(resid_x, resid_y)

        # 计算自由度和显著性
        dof = n - len(z) - 2  # 自由度 = 样本量 - 控制变量数 - 2
        return corr, p, dof

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 获取所有变量名
            variables = df.columns.tolist()

            # 存储结果的列表
            results = []

            for i in range(len(variables)):
                for j in range(i + 1, len(variables)):
                    x = variables[i]
                    y = variables[j]
                    other_vars = [var for var in variables if var not in [x, y]]
                    corr, p, dof = self.partial_corr(df, x, y, other_vars)  # 获取自由度
                    pair = f"{x} - {y}"

                    # 更详细的结果解读
                    if p < 0.001:
                        significance = "p < 0.001"
                    elif p < 0.01:
                        significance = "p < 0.01"
                    elif p < 0.05:
                        significance = "p < 0.05"
                    else:
                        significance = f"p = {p:.4f}"

                    # 增加统计显著性的完整解读（支持语言切换）
                    interpretation = (
                        f"{languages[self.current_language]['interpretation_low_p']} "
                        f"{languages[self.current_language]['partial_corr']} = {corr:.4f}, "
                        f"{significance}, "
                        f"{languages[self.current_language]['degree_of_freedom']} = {dof}"
                    ) if p < 0.05 else (
                        f"{languages[self.current_language]['interpretation_high_p']} "
                        f"{languages[self.current_language]['partial_corr']} = {corr:.4f}, "
                        f"{significance}, "
                        f"{languages[self.current_language]['degree_of_freedom']} = {dof}"
                    )

                    results.append([pair, corr, p, dof])

            # 修改结果DataFrame的列定义
            result_df = pd.DataFrame(
                results,
                columns=[
                    languages[self.current_language]["columns_stats"][0],  # 变量对
                    languages[self.current_language]["columns_stats"][1],  # 偏相关系数
                    languages[self.current_language]["columns_stats"][2],  # p值
                    languages[self.current_language]["degree_of_freedom"]
                ]
            )

            # 绘制偏相关系数的柱状图，增加误差线和显著性标记
            plt.figure(figsize=(12, 8), dpi=300)
            # 替换原柱状图绘制代码
            # 获取当前语言下的列名
            corr_col = languages[self.current_language]["columns_stats"][1]  # 偏相关系数/Partial Correlation Coefficient
            p_col = languages[self.current_language]["columns_stats"][2]  # p 值/p-value

            bars = plt.bar(
                result_df[languages[self.current_language]["columns_stats"][0]],  # 变量对列
                result_df[corr_col],  # 偏相关系数列（动态获取）
                yerr=1.96 * result_df[corr_col].std() / np.sqrt(len(result_df)),  # 修正误差线计算
                capsize=5
            )

            # 显著性标记部分也需要修改
            for i, (_, row) in enumerate(result_df.iterrows()):
                if row[p_col] < 0.001:  # 使用动态获取的p值列名
                    plt.text(i, row[corr_col] + 0.05, '***', ha='center')
                elif row[p_col] < 0.01:
                    plt.text(i, row[corr_col] + 0.05, '**', ha='center')
                elif row[p_col] < 0.05:
                    plt.text(i, row[corr_col] + 0.05, '*', ha='center')

            plt.xlabel('变量对' if self.current_language == "zh" else 'Variable Pair')
            plt.ylabel('偏相关系数' if self.current_language == "zh" else 'Partial Correlation Coefficient')
            plt.title('偏相关系数分析结果' if self.current_language == "zh" else 'Partial Correlation Analysis Results')
            plt.xticks(rotation=45, ha='right')  # 旋转标签并右对齐
            plt.axhline(y=0, color='gray', linestyle='--')  # 添加水平线指示零值
            plt.tight_layout()  # 自动调整布局

            # 保存图片，使用更高的分辨率
            image_path = os.path.splitext(file_path)[0] + '_partial_corr_plot.png'
            plt.savefig(image_path, dpi=300, bbox_inches='tight')
            plt.close()

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建一个新的Word文档
                doc = Document()

                # 添加标题
                doc.add_heading(
                    '偏相关分析结果' if self.current_language == "zh" else 'Partial Correlation Analysis Results', 0)

                # 添加表格
                table = doc.add_table(rows=1, cols=len(result_df.columns))
                hdr_cells = table.rows[0].cells
                for col_idx, col_name in enumerate(result_df.columns):
                    hdr_cells[col_idx].text = col_name

                for index, row in result_df.iterrows():
                    row_cells = table.add_row().cells
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx].text = str(value)

                # 添加解释说明部分（新增代码）
                doc.add_heading(
                    '结果解释说明' if self.current_language == "zh" else 'Result Interpretation', level=1)
                explanation = doc.add_paragraph()

                # 添加中文解释列表
                if self.current_language == "zh":
                    explanation.add_run('• ').bold = True
                    explanation.add_run(
                        '偏相关系数：用于衡量两个变量在控制其他变量影响后的相关性强度，取值范围为[-1, 1]。\n')

                    explanation.add_run('• ').bold = True
                    explanation.add_run('p值：用于判断相关性的统计显著性，p < 0.05 表示显著相关。\n')

                    explanation.add_run('• ').bold = True
                    explanation.add_run('显著性标记：***表示p < 0.001，** 表示p < 0.01，*表示p < 0.05。\n')

                    explanation.add_run('• ').bold = True
                    explanation.add_run('自由度：样本量减去控制变量数量再减2，影响统计检验的临界值。')

                # 添加英文解释列表
                else:
                    explanation.add_run('• ').bold = True
                    explanation.add_run(
                        'Partial correlation coefficient: Measures the strength of association between two variables after controlling other variables, ranging from [-1, 1].\n')

                    explanation.add_run('• ').bold = True
                    explanation.add_run(
                        'p-value: Used to determine statistical significance, p < 0.05 indicates significant correlation.\n')

                    explanation.add_run('• ').bold = True
                    explanation.add_run(
                        'Significance markers: ***indicates p < 0.001,** indicates p < 0.01, *indicates p < 0.05.\n')

                    explanation.add_run('• ').bold = True
                    explanation.add_run(
                        'Degrees of freedom: Sample size minus number of control variables minus 2, affecting critical values for statistical tests.')

                # 添加图片
                doc.add_picture(image_path, width=Inches(6))

                # 保存Word文档
                doc.save(save_path)

                # 设置 wraplength 属性让文本自动换行
                self.result_label.config(text=languages[self.current_language]["analysis_complete"].format(save_path),
                                         wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]["no_save_path_selected"])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]["analysis_error"].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])
        self.open_excel_label.config(text=languages[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=languages[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame,
                                               text=languages[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = PartialCorrelationAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()
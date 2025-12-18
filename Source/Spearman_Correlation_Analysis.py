import tkinter as tk
from tkinter import filedialog
import os
import numpy as np
import pandas as pd
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
import pathlib
import pandas.plotting as pd_plotting
from scipy import stats
from docx import Document
from docx.shared import Inches
from statsmodels.stats.multitest import multipletests

# 设置Matplotlib支持中文显示
plt.rcParams["font.family"] = ["SimHei", "WenQuanYi Micro Hei", "Heiti TC", "Arial Unicode MS"]
plt.rcParams["axes.unicode_minus"] = False  # 正确显示负号

# 定义语言字典
languages = {
    "zh": {
        "title": "Spearman 相关性分析",
        "select_button_text": "选择文件",
        "file_entry_placeholder": "请输入待分析 Excel 文件的完整路径",
        "analyze_button_text": "分析文件",
        "no_file_selected": "请选择文件。",
        "file_not_exists": "文件不存在不存在，请重新选择。",
        "analysis_error": "分析文件时出错: {}",
        "analysis_complete": "分析完成，结果已保存到 {}\n",
        "no_save_path_selected": "未选择保存路径，结果未保存。",
        'open_excel_button_text': "示例数据",
        "switch_language_button_text": "中/英",
        "correlation_types": ["Spearman相关系数"],
        "explanation": {
            "Spearman相关系数": "基于变量的秩次而非实际值，用于衡量两个变量之间的单调关系，不受变量分布的影响。"
        },
        "interpretation": {
            "相关系数": "相关系数的绝对值越接近1，说明两个变量之间的相关性越强；接近0则表示相关性较弱。",
            "p值": "p值小于显著性水平（通常为0.05）时，拒绝原假设，认为两个变量之间存在显著相关性；否则，接受原假设，认为两个变量之间无显著相关性。"
        }
    },
    "en": {
        "title": "Spearman Correlation Analysis",
        "select_button_text": "Select File",
        "file_entry_placeholder": "Please enter the full path of the Excel file to be analyzed",
        "analyze_button_text": "Analyze File",
        "no_file_selected": "Please select a valid file path.",
        "file_not_exists": "The file does not exist. Please select again.",
        "analysis_error": "An error occurred while analyzing the file: {}",
        "analysis_complete": "Analysis completed. The results have been saved to {}\n",
        "no_save_path_selected": "No save path selected. The results were not saved.",
        'open_excel_button_text': "Example data",
        "switch_language_button_text": "Chinese/English",
        "correlation_types": ["Spearman's Rank Correlation Coefficient"],
        "explanation": {
            "Spearman's Rank Correlation Coefficient": "Based on the ranks of variables rather than the actual values, used to measure the monotonic relationship between two variables, regardless of the variable distribution."
        },
        "interpretation": {
            "Correlation Coefficient": "The closer the absolute value of the correlation coefficient is to 1, the stronger the correlation between the two variables; close to 0 indicates a weak correlation.",
            "p-value": "When the p-value is less than the significance level (usually 0.05), the null hypothesis is rejected, indicating a significant correlation between the two variables; otherwise, the null hypothesis is accepted, indicating no significant correlation."
        }
    }
}


class SpearmanCorrelationAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data20.xls")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(
                file_path,
                header=0  # 明确指定第一行为表头
            )

            # 手动删除空行
            df = df.dropna(how='all')

            # 清理列名
            df.columns = df.columns.str.strip().str.replace(r'[^\w\s]', '', regex=True)

            # 检查数据是否为数值类型
            numerical_df = df.select_dtypes(include=[np.number])
            if numerical_df.empty:
                raise ValueError("数据中没有数值列，无法进行相关性分析。")

            # 确保有至少两列数据用于分析
            if len(numerical_df.columns) < 2:
                raise ValueError("数据中至少需要两列数值列才能进行相关性分析。")

            # 计算Spearman相关性
            spearman_corr = numerical_df.corr(method='spearman')

            # 验证相关性矩阵有效性
            if spearman_corr.empty:
                raise ValueError("相关性矩阵计算为空，请检查输入数据")
            if len(spearman_corr.columns) != len(spearman_corr.index):
                raise ValueError(
                    f"相关性矩阵行列数不匹配: 列数={len(spearman_corr.columns)}, 行数={len(spearman_corr.index)}")

            # 计算p值
            def calculate_pvalues(df):
                df = df.dropna()._get_numeric_data()
                dfcols = pd.DataFrame(columns=df.columns)
                pvalues = dfcols.transpose().join(dfcols, how='outer')
                stats_dict = {}  # 新增：存储统计量

                for r in df.columns:
                    for c in df.columns:
                        if r == c:
                            # 对角线元素，自相关
                            pvalues.loc[r, c] = 1.0  # p值为1.0
                            stats_dict[(r, c)] = {
                                'correlation': 1.0,
                                'pvalue': 1.0,
                                'n': len(df.dropna(subset=[r])),
                                'significant': False
                            }
                        else:
                            # 计算Spearman相关系数、p值和样本量
                            corr, p = stats.spearmanr(df[r], df[c], nan_policy='omit')
                            n = len(df[[r, c]].dropna())  # 有效样本量
                            significant = p < 0.05  # 显著性判断
                            pvalues.loc[r, c] = p
                            stats_dict[(r, c)] = {
                                'correlation': corr,
                                'pvalue': p,
                                'n': n,
                                'significant': significant
                            }
                return pvalues, stats_dict

            # 修改调用处（约第126行）
            spearman_pvalues, stats_dict = calculate_pvalues(numerical_df)

            # 整理数据
            data = []
            correlation_types = languages[self.current_language]["correlation_types"]
            explanations = languages[self.current_language]['explanation']
            interpretations = languages[self.current_language]['interpretation']

            # 获取并验证列名
            columns = spearman_corr.columns.tolist()
            if not columns:
                raise ValueError("无法获取有效的列名，无法进行相关性分析")
            num_columns = len(columns)
            if num_columns < 2:
                raise ValueError(f"有效数值列不足: 仅发现{num_columns}列，至少需要2列")

            # 多重比较校正（Bonferroni方法）
            p_values = []
            indices = []
            for i in range(num_columns):
                for j in range(i + 1, num_columns):
                    p_values.append(spearman_pvalues.iloc[i, j])
                    indices.append((i, j))

            # 执行Bonferroni校正（此时p_values和indices已填充数据）
            corrected_p = multipletests(p_values, method='bonferroni')[1]

            # 创建校正后的p值矩阵（保持对称性）
            corrected_pvalues = spearman_pvalues.copy()
            for idx, (i, j) in enumerate(indices):
                corrected_pvalues.iloc[i, j] = corrected_p[idx]
                corrected_pvalues.iloc[j, i] = corrected_p[idx]  # 对称位置赋值

            # 循环处理相关性结果（增加索引边界检查）
            for i, (corr, pvalues) in enumerate(zip([spearman_corr], [spearman_pvalues])):
                # 验证当前循环的相关性矩阵
                current_num_cols = len(corr.columns)
                if current_num_cols != num_columns:
                    raise ValueError(f"相关性矩阵列数不匹配: 预期{num_columns}, 实际{current_num_cols}")

                # 只处理上三角部分（i < j）
                for col1_idx in range(num_columns):
                    if col1_idx < 0 or col1_idx >= num_columns:
                        print(f"跳过无效的col1_idx: {col1_idx}")
                        continue

                    col1 = columns[col1_idx]
                    max_col2_idx = num_columns - 1
                    start_col2_idx = col1_idx + 1

                    if start_col2_idx > max_col2_idx:
                        continue

                    for col2_idx in range(start_col2_idx, num_columns):
                        if col2_idx < 0 or col2_idx >= num_columns:
                            print(f"跳过无效的col2_idx: {col2_idx}")
                            continue

                        col2 = columns[col2_idx]

                        try:
                            # 从stats_dict获取完整统计信息
                            stats_info = stats_dict[(col1, col2)]
                            corr_value = round(stats_info['correlation'], 4)
                            p_value = round(stats_info['pvalue'], 4)
                            n = stats_info['n']
                            significant = stats_info['significant']

                            # 添加显著性标记
                            sig_symbol = "*" if significant else ""

                            data.append([
                                f"{correlation_types[i]} ({col1} vs {col2})",
                                corr_value,
                                p_value,
                                round(corrected_pvalues.iloc[col1_idx, col2_idx], 4)  # 新增：校正后的p值
                            ])
                        except IndexError as e:
                            print(f"索引越界错误: {e}")
                            continue
                        except Exception as e:
                            print(f"处理 {col1} vs {col2} 时出错: {e}")
                            continue

                # 根据当前语言设置表头
                if self.current_language == "zh":
                    headers = ["统计量", "相关系数", "p值", "校正后p值（Bonferroni）"]
                else:
                    headers = ["Statistic", "Correlation Coefficient", "p-value", "Corrected p-value (Bonferroni)"]

            # 合并数据、解释说明和结果解读
            df_results = pd.DataFrame(data, columns=headers)  # 补充此句

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建一个新的 Word 文档
                doc = Document()

                # 添加表格标题
                if self.current_language == "zh":
                    result_heading = "分析结果"
                else:
                    result_heading = "Analysis Results"
                doc.add_heading(result_heading, level=1)

                # 添加结果表格 - 确保列数正确
                num_cols = len(headers)
                table = doc.add_table(rows=1, cols=num_cols)
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    if i < len(hdr_cells):  # 确保不超出表头单元格范围
                        hdr_cells[i].text = header

                # 添加数据行
                for index, row in df_results.iterrows():  # 只添加分析结果数据
                    row_cells = table.add_row().cells
                    # 确保只访问存在的单元格
                    for i in range(min(len(row), len(row_cells))):
                        row_cells[i].text = str(row.iloc[i])  # 使用iloc明确指定按位置访问

                # 添加解释说明部分（列表形式）
                # 根据当前语言设置解释说明标题
                if self.current_language == "zh":
                    explanation_heading = "解释说明"
                else:
                    explanation_heading = "Explanation"
                doc.add_heading(explanation_heading, level=1)
                explanation_list = doc.add_paragraph()

                # 添加相关系数解释
                if self.current_language == "zh":
                    explanation_list.add_run("• 相关系数: ").bold = True
                    explanation_list.add_run(explanations.get("Spearman相关系数", "") + "\n")

                    explanation_list.add_run("• 校正后p值（Bonferroni）: ").bold = True
                    explanation_list.add_run("经Bonferroni多重比较校正后的p值，用于控制I类错误\n")
                else:
                    explanation_list.add_run("• Correlation Coefficient: ").bold = True
                    explanation_list.add_run(explanations.get("Spearman's Rank Correlation Coefficient", "") + "\n")

                    explanation_list.add_run("• Corrected p-value (Bonferroni): ").bold = True
                    explanation_list.add_run(
                        "p-value corrected by Bonferroni multiple comparison method to control Type I error\n")

                # 添加结果解读部分（列表形式）
                # 根据当前语言设置结果解读标题
                if self.current_language == "zh":
                    interpretation_heading = "结果解读"
                else:
                    interpretation_heading = "Result Interpretation"
                doc.add_heading(interpretation_heading, level=1)
                interpretation_list = doc.add_paragraph()

                if self.current_language == "zh":
                    interpretation_list.add_run("• 相关系数: ").bold = True
                    interpretation_list.add_run(interpretations.get("相关系数", "") + "\n")

                    interpretation_list.add_run("• p值: ").bold = True
                    interpretation_list.add_run(interpretations.get("p值", "") + "\n")

                    interpretation_list.add_run("• 校正后p值（Bonferroni）: ").bold = True
                    interpretation_list.add_run("校正后p值 < 0.05 表示在多重比较中仍具有统计学显著性\n")
                else:
                    interpretation_list.add_run("• Correlation Coefficient: ").bold = True
                    interpretation_list.add_run(interpretations.get("Correlation Coefficient", "") + "\n")

                    interpretation_list.add_run("• p-value: ").bold = True
                    interpretation_list.add_run(interpretations.get("p-value", "") + "\n")

                    interpretation_list.add_run("• Corrected p-value (Bonferroni): ").bold = True
                    interpretation_list.add_run(
                        "Corrected p-value < 0.05 indicates statistical significance remains in multiple comparisons\n")

                # 生成相关性热力图
                try:
                    desktop_path = pathlib.Path.home() / 'Desktop'
                    desktop_path.mkdir(exist_ok=True)

                    plot_path = desktop_path / 'correlation_heatmap.png'
                    plt.figure(figsize=(10, 8))
                    im = plt.imshow(spearman_corr, cmap='coolwarm', interpolation='nearest', vmin=-1, vmax=1)
                    plt.colorbar(im, label='相关系数' if self.current_language == 'zh' else 'Correlation Coefficient')
                    plt.xticks(range(num_columns), columns, rotation=0, ha='right')
                    plt.yticks(range(num_columns), columns)
                    for i in range(num_columns):
                        for j in range(num_columns):
                            corr_val = spearman_corr.iloc[i, j]
                            p_val = spearman_pvalues.iloc[i, j]
                            sig = "*" if p_val < 0.05 else ""
                            plt.text(j, i, f'{corr_val:.2f}{sig}', ha='center', va='center',
                                     color='white' if abs(corr_val) > 0.5 else 'black')

                    plt.title(
                        'Spearman相关性热力图' if self.current_language == 'zh' else 'Spearman Correlation Heatmap')
                    plt.tight_layout()
                    plt.savefig(plot_path, dpi=300)  # 提高分辨率
                    plt.close()
                except Exception as e:
                    raise RuntimeError(f"保存热力图时出错: {str(e)}")

                # 生成散点图矩阵
                try:
                    scatter_matrix_path = desktop_path / 'scatter_matrix.png'
                    # 生成散点图矩阵并获取所有子图的坐标轴对象
                    axes = pd_plotting.scatter_matrix(numerical_df, alpha=0.8, figsize=(10, 10), diagonal='hist')
                    # 遍历所有子图，设置X轴刻度标签横向显示
                    for ax in axes.flatten():
                        # 设置X轴刻度标签旋转0度（横向），并居中对齐
                        plt.setp(ax.get_xticklabels(), rotation=0, ha='center')
                    plt.suptitle('散点图矩阵' if self.current_language == 'zh' else 'Scatter Matrix')
                    plt.tight_layout(rect=[0, 0, 1, 0.96])
                    plt.savefig(scatter_matrix_path)
                    plt.close()
                except Exception as e:
                    raise RuntimeError(f"保存散点图矩阵时出错: {str(e)}")

                # 生成相关性柱状图
                try:
                    if num_columns > 0:  # 再次检查列数
                        selected_variable = columns[0]
                        correlation_column = spearman_corr[selected_variable]
                        bar_plot_path = desktop_path / 'correlation_bar_plot.png'
                        plt.figure(figsize=(10, 6))
                        correlation_column.plot(kind='bar')
                        plt.title(
                            f'与{selected_variable}的相关性' if self.current_language == 'zh' else f'Correlation with {selected_variable}')
                        plt.xlabel('变量' if self.current_language == 'zh' else 'Variables')
                        plt.ylabel('相关系数' if self.current_language == 'zh' else 'Correlation Coefficient')
                        plt.xticks(rotation=0)
                        plt.tight_layout()
                        plt.savefig(bar_plot_path)
                        plt.close()
                    else:
                        raise ValueError("没有可用的变量生成柱状图")
                except Exception as e:
                    raise RuntimeError(f"保存柱状图时出错: {str(e)}")

                # 在 Word 文档中添加图片
                doc.add_heading(
                    'Spearman相关性热力图' if self.current_language == 'zh' else 'Spearman Correlation Heatmap',
                    level=2)
                doc.add_picture(str(plot_path), width=Inches(6))
                doc.add_heading('散点图矩阵' if self.current_language == 'zh' else 'Scatter Matrix', level=2)
                doc.add_picture(str(scatter_matrix_path), width=Inches(6))
                doc.add_heading(
                    f'与{selected_variable}的相关性' if self.current_language == 'zh' else f'Correlation with {selected_variable}',
                    level=2)
                doc.add_picture(str(bar_plot_path), width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                result_msg = languages[self.current_language]['analysis_complete'].format(
                    save_path) + f"相关性热力图已保存到 {plot_path}"
                result_msg += f"\n散点图矩阵已保存到 {scatter_matrix_path}"
                result_msg += f"\n相关性柱状图已保存到 {bar_plot_path}"
                self.result_label.config(text=result_msg, wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]['no_save_path_selected'])

        except Exception as e:
            # 输出详细错误信息到控制台，便于调试
            import traceback
            print("详细错误信息:")
            traceback.print_exc()
            self.result_label.config(text=languages[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])
        self.open_excel_label.config(text=languages[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=languages[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame,
                                               text=languages[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = SpearmanCorrelationAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()

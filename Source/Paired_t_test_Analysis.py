import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
from scipy import stats
import matplotlib.pyplot as plt
import pathlib
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'  # 指定中文字体，SimHei 是黑体
plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

# 定义语言字典，包含所有需要翻译的键和内容
LANGUAGES = {
    'zh': {
        'title': "配对 t 检验",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "请选择文件。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "中/英",
        'open_excel_button_text': "示例数据",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'statistic': "统计量",
        't_statistic': "t 统计量",
        'degrees_of_freedom': "自由度",
        'p_value': "p 值",
        'confidence_interval': "置信区间",
        'paired_t_test': "配对 t 检验",
        'sample_size': "样本量",
        'mean': "均值",
        'explanation_heading': "解释说明",
        'interpretation_heading': "结果解读",
        'chart_heading': "图表",
        'bar_chart': "柱状图",
        'error_bar_chart': "误差线图",
        'box_plot': "箱线图",
        'line_chart': "折线图",
        'sample': "样本",
        'value': "数值",
        'observation': "观测值",
        'explanation': {
            "paired_t_test": "用于比较两个相关样本的均值是否有显著差异。",
            "sample_size": "每个样本中的观测值数量。",
            "mean": "样本数据的平均值。",
            "t_statistic": "用于衡量两组样本均值差异的统计量。",
            "degrees_of_freedom": "在统计计算中能够自由取值的变量个数。",
            "p_value": "用于判断两组样本均值是否有显著差异的概率值。",
            "confidence_interval": "均值差异可能存在的区间范围。"
        },
        'interpretation': {
            "t_statistic": "t 统计量的绝对值越大，说明两组样本均值差异越显著。",
            "p_value": "p 值小于显著性水平（通常为 0.05）时，拒绝原假设，认为两组样本均值存在显著差异；否则，接受原假设，认为两组样本均值无显著差异。",
            "degrees_of_freedom": "自由度影响 t 分布的形状，自由度越大，t 分布越接近正态分布。",
            "confidence_interval": "如果置信区间不包含 0，说明两组样本均值存在显著差异。"
        }
    },
    'en': {
        'title': "Paired T Test",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Chinese/English",
        'open_excel_button_text': "Example data",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'statistic': "Statistic",
        't_statistic': "t Statistic",
        'degrees_of_freedom': "Degrees of Freedom",
        'p_value': "p Value",
        'confidence_interval': "Confidence Interval",
        'paired_t_test': "Paired T Test",
        'sample_size': "Sample Size",
        'mean': "Mean",
        'explanation_heading': "Explanation",
        'interpretation_heading': "Interpretation",
        'chart_heading': "Charts",
        'bar_chart': "Bar Chart",
        'error_bar_chart': "Error Bar Chart",
        'box_plot': "Box Plot",
        'line_chart': "Line Chart",
        'sample': "Sample",
        'value': "Value",
        'observation': "Observation",
        'explanation': {
            "paired_t_test": "Used to compare whether the means of two related samples are significantly different.",
            "sample_size": "The number of observations in each sample.",
            "mean": "The average value of the sample data.",
            "t_statistic": "A statistic used to measure the difference between the means of two groups of samples.",
            "degrees_of_freedom": "The number of variables that can take on independent values in a statistical calculation.",
            "p_value": "A probability value used to determine whether there is a significant difference between the means of two groups of samples.",
            "confidence_interval": "The range within which the difference in means may lie."
        },
        'interpretation': {
            "t_statistic": "The larger the absolute value of the t statistic, the more significant the difference between the means of the two groups of samples.",
            "p_value": "When the p-value is less than the significance level (usually 0.05), the null hypothesis is rejected, indicating a significant difference between the means of the two groups of samples; otherwise, the null hypothesis is accepted, indicating no significant difference.",
            "degrees_of_freedom": "The degrees of freedom affect the shape of the t-distribution. The larger the degrees of freedom, the closer the t-distribution is to the normal distribution.",
            "confidence_interval": "If the confidence interval does not contain 0, it indicates a significant difference between the means of the two groups of samples."
        }
    }
}


class PairedTTestAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为中文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data42.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 检查数据是否为数值类型
            numerical_df = df.select_dtypes(include=[np.number])
            if numerical_df.empty:
                raise ValueError("数据中没有数值列，无法进行配对t检验。")
            if numerical_df.shape[1] != 2:
                raise ValueError("数据必须包含且仅包含两列数值数据，才能进行配对t检验。")

            # 进行配对 t 检验
            t_stat, p_value = stats.ttest_rel(*numerical_df.T.values)
            df_value = len(numerical_df) - 1  # 自由度
            mean_diff = numerical_df.iloc[:, 0].mean() - numerical_df.iloc[:, 1].mean()
            std_err = stats.sem(numerical_df.iloc[:, 0] - numerical_df.iloc[:, 1])
            conf_int = stats.t.interval(0.95, df_value, loc=mean_diff, scale=std_err)

            # 格式化置信区间，去除npfloat字样
            conf_int_str = f"({conf_int[0]:.4f}, {conf_int[1]:.4f})"

            # 计算样本量和均值
            sample_sizes = numerical_df.count()
            means = numerical_df.mean()

            # 整理数据
            data = [
                [LANGUAGES[self.current_language]['paired_t_test'], f"{t_stat:.4f}", f"{df_value}", f"{p_value:.6f}",
                 conf_int_str],
                [LANGUAGES[self.current_language]['sample_size'], sample_sizes.to_dict(), "", "", ""],
                [LANGUAGES[self.current_language]['mean'], means.to_dict(), "", "", ""]
            ]
            headers = [
                LANGUAGES[self.current_language]['statistic'],
                LANGUAGES[self.current_language]['t_statistic'],
                LANGUAGES[self.current_language]['degrees_of_freedom'],
                LANGUAGES[self.current_language]['p_value'],
                LANGUAGES[self.current_language]['confidence_interval']
            ]
            result_df = pd.DataFrame(data, columns=headers)

            # 添加解释说明（列表形式）
            explanations = LANGUAGES[self.current_language]['explanation']
            interpretations = LANGUAGES[self.current_language]['interpretation']

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加分析结果表格
                doc.add_heading(LANGUAGES[self.current_language]['title'], level=1)
                table = doc.add_table(rows=result_df.shape[0] + 1, cols=result_df.shape[1])
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(result_df.columns):
                    hdr_cells[col_idx].text = header
                for row_idx, row in result_df.iterrows():
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx].text = str(value)

                # 添加解释说明（列表形式）
                doc.add_heading(LANGUAGES[self.current_language]['explanation_heading'], level=1)
                explanation_items = [
                    f"{LANGUAGES[self.current_language]['paired_t_test']}: {explanations['paired_t_test']}",
                    f"{LANGUAGES[self.current_language]['sample_size']}: {explanations['sample_size']}",
                    f"{LANGUAGES[self.current_language]['mean']}: {explanations['mean']}",
                    f"{LANGUAGES[self.current_language]['t_statistic']}: {explanations['t_statistic']}",
                    f"{LANGUAGES[self.current_language]['degrees_of_freedom']}: {explanations['degrees_of_freedom']}",
                    f"{LANGUAGES[self.current_language]['p_value']}: {explanations['p_value']}",
                    f"{LANGUAGES[self.current_language]['confidence_interval']}: {explanations['confidence_interval']}"
                ]
                for item in explanation_items:
                    para = doc.add_paragraph()
                    para.add_run("• ").bold = True
                    para.add_run(item)

                # 添加结果解读（列表形式）
                doc.add_heading(LANGUAGES[self.current_language]['interpretation_heading'], level=1)
                interpretation_items = [
                    f"{LANGUAGES[self.current_language]['t_statistic']}: {interpretations['t_statistic']}",
                    f"{LANGUAGES[self.current_language]['p_value']}: {interpretations['p_value']}",
                    f"{LANGUAGES[self.current_language]['degrees_of_freedom']}: {interpretations['degrees_of_freedom']}",
                    f"{LANGUAGES[self.current_language]['confidence_interval']}: {interpretations['confidence_interval']}"
                ]
                for item in interpretation_items:
                    para = doc.add_paragraph()
                    para.add_run("• ").bold = True
                    para.add_run(item)

                # 绘制图表
                plot_path = self.plot_results(numerical_df, save_path)
                if plot_path:
                    doc.add_heading(LANGUAGES[self.current_language]['chart_heading'], level=1)
                    doc.add_picture(plot_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(
                    save_path)
                self.result_label.config(text=result_msg, wraplength=400)
            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def plot_results(self, numerical_df, save_path):
        # 柱状图
        plt.figure(figsize=(12, 8))
        plt.subplot(2, 2, 1)
        means = numerical_df.mean()
        bars = plt.bar(means.index, means)
        for bar in bars:
            height = bar.get_height()
            plt.annotate(f'{height:.2f}',
                         xy=(bar.get_x() + bar.get_width() / 2, height),
                         xytext=(0, 3),  # 3 points vertical offset
                         textcoords="offset points",
                         ha='center', va='bottom')
        plt.title(LANGUAGES[self.current_language]['bar_chart'])
        plt.xlabel(LANGUAGES[self.current_language]['sample'])
        plt.ylabel(LANGUAGES[self.current_language]['mean'])

        # 误差线图
        plt.subplot(2, 2, 2)
        means = numerical_df.mean()
        stds = numerical_df.std()
        plt.errorbar(means.index, means, yerr=stds, fmt='o')
        plt.title(LANGUAGES[self.current_language]['error_bar_chart'])
        plt.xlabel(LANGUAGES[self.current_language]['sample'])
        plt.ylabel(LANGUAGES[self.current_language]['mean'])

        # 箱线图
        plt.subplot(2, 2, 3)
        numerical_df.boxplot()
        plt.title(LANGUAGES[self.current_language]['box_plot'])
        plt.xlabel(LANGUAGES[self.current_language]['sample'])
        plt.ylabel(LANGUAGES[self.current_language]['value'])

        # 折线图
        plt.subplot(2, 2, 4)
        plt.plot(numerical_df)
        plt.title(LANGUAGES[self.current_language]['line_chart'])
        plt.xlabel(LANGUAGES[self.current_language]['observation'])
        plt.ylabel(LANGUAGES[self.current_language]['value'])
        plt.legend(numerical_df.columns)

        plt.tight_layout()
        plot_path = save_path.replace('.docx', '_plots.png')
        plt.savefig(plot_path)
        plt.close()
        return plot_path

    def switch_language(self, event=None):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(LANGUAGES[self.current_language]["title"])
        self.select_button.config(text=LANGUAGES[self.current_language]["select_button"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=LANGUAGES[self.current_language]["analyze_button"])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]["switch_language"])
        self.open_excel_label.config(text=LANGUAGES[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["select_button"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["analyze_button"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=LANGUAGES[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=LANGUAGES[self.current_language]["switch_language"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = PairedTTestAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()
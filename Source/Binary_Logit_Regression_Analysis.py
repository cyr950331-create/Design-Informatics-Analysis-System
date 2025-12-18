import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import accuracy_score, roc_auc_score, roc_curve, precision_score, recall_score, f1_score, confusion_matrix
from sklearn.model_selection import train_test_split
import matplotlib.pyplot as plt
import statsmodels.api as sm
from docx import Document
from docx.shared import Inches


# 设置支持中文的字体
plt.rcParams['font.family'] = 'SimHei'  # 使用黑体字体，可根据系统情况修改
plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

# 定义语言字典（所有键使用英文）
languages = {
    'zh': {
        'title': "二元Logit回归",
        'main_heading': "二元Logit回归分析结果",
        'select_button_text': "选择文件",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'analyze_button_text': "分析文件",
        'no_file_selected': "请选择文件。",
        'file_not_exists': "文件不存在，请重新选择。",
        'analysis_error': "分析文件时出错: {}",
        'analysis_complete': "分析完成，结果已保存到 {}",
        'images_saved': "图片已保存到 {}",
        'no_save_path_selected': "未选择保存路径，结果未保存。",
        'open_excel_button_text': "示例数据",
        'switch_language_button_text': "中/英",
        'explanation_row': '解释说明',
        'explanation_coefficients': '回归系数',
        'explanation_intercept': '截距',
        'explanation_accuracy': '准确率',
        'explanation_roc_auc': 'ROC曲线下面积',
        'explanation_z_value': 'z统计量',
        'explanation_p_value': 'p值',
        'explanation_regression_equation': '回归方程',
        'explanation': {
            "coefficients": "回归系数，表示每个自变量对因变量的影响程度。",
            "intercept": "截距，是当所有自变量为 0 时因变量的预测值。",
            "accuracy": "准确率，衡量模型预测正确的比例。",
            "roc_auc": "ROC曲线下面积，衡量模型的分类能力。",
            "z_value": "z 统计量，用于检验每个自变量的显著性。",
            "p_value": "p 值，用于判断自变量的显著性，p 值越小，自变量越显著。",
            "regression_equation": "回归方程，展示因变量与自变量之间的关系。"
        },
        'table_headers': {
            'model': '模型',
            'intercept': '截距',
            'accuracy': '准确率',
            'precision': '精确率',
            'recall': '召回率',
            'f1_score': 'F1分数',
            'roc_auc': 'ROC曲线下面积',
            'regression_equation': '回归方程',
            'z_value_prefix': 'z值（',
            'p_value_prefix': 'p值（',
            'z_value_suffix': '）',
            'p_value_suffix': '）'
        },
        'dependent_var_not_binary': "因变量必须是二元数据（只能包含0和1），请检查数据。",
        'missing_values': "数据中存在缺失值，请先处理缺失值后再分析。",
        'non_numeric_vars': "以下自变量包含非数值类型数据：{}，请检查数据。",
        'roc_curve_title': '受试者工作特征曲线',
        'roc_x_label': '假阳性率',
        'roc_y_label': '真阳性率',
        'test_size_info': '（使用80%数据训练，20%数据测试）',
        'regression_equation_text': '回归方程',
        'analysis_results_heading': '1. 分析结果',
        'result_interpretation_heading': '2. 结果解读',
        'explanations_heading': '3. 解释说明',
        'confusion_matrix_heading': '3. 混淆矩阵',
        'roc_curve_heading': '4. ROC曲线',
        'predicted_0': '预测为0',
        'predicted_1': '预测为1',
        'actual_0': '实际为0',
        'actual_1': '实际为1',
        'roc_curve_legend': 'ROC曲线 (面积 = {})',
        'template_file_not_found': '模板文件不存在：{}',
        'open_file_failed': '打开文件失败：{}'
    },
    'en': {
        'title': "Binary Logit Regression",
        'main_heading': "Binary Logit Regression Analysis Results",
        'select_button_text': "Select File",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'analyze_button_text': "Analyze File",
        'no_file_selected': "Please select a valid file path.",
        'file_not_exists': "The file does not exist. Please select again.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'analysis_complete': "Analysis completed. The results have been saved to {}",
        'images_saved': "Images have been saved to {}",
        'no_save_path_selected': "No save path selected. The results were not saved.",
        'open_excel_button_text': "Example data",
        'switch_language_button_text': "Chinese/English",
        'explanation_row': 'Explanation',
        'explanation_coefficients': 'Coefficients',
        'explanation_intercept': 'Intercept',
        'explanation_accuracy': 'Accuracy',
        'explanation_roc_auc': 'ROC-AUC',
        'explanation_z_value': 'Z-value',
        'explanation_p_value': 'P-value',
        'explanation_regression_equation': 'Regression Equation',
        'explanation': {
            "coefficients": "Regression coefficients, indicating the influence of each independent variable on the dependent variable.",
            "intercept": "Intercept, which is the predicted value of the dependent variable when all independent variables are 0.",
            "accuracy": "Accuracy, measuring the proportion of correct predictions of the model.",
            "roc_auc": "Area under the ROC curve, measuring the classification ability of the model.",
            "z_value": "z statistic, used to test the significance of each independent variable.",
            "p_value": "p value, used to determine the significance of the independent variable. The smaller the p value, the more significant the independent variable.",
            "regression_equation": "Regression equation, showing the relationship between dependent and independent variables."
        },
        'table_headers': {
            'model': 'Model',
            'intercept': 'Intercept',
            'accuracy': 'Accuracy',
            'precision': 'Precision',
            'recall': 'Recall',
            'f1_score': 'F1-Score',
            'roc_auc': 'ROC-AUC',
            'regression_equation': 'Regression Equation',
            'z_value_prefix': 'z-value (',
            'p_value_prefix': 'p-value (',
            'z_value_suffix': ')',
            'p_value_suffix': ')'
        },
        'dependent_var_not_binary': "The dependent variable must be binary (only 0 and 1), please check the data.",
        'missing_values': "There are missing values in the data. Please handle them before analysis.",
        'non_numeric_vars': "The following independent variables contain non-numeric data: {}, please check the data.",
        'roc_curve_title': 'Receiver Operating Characteristic',
        'roc_x_label': 'False Positive Rate',
        'roc_y_label': 'True Positive Rate',
        'test_size_info': '（Using 80% data for training, 20% for testing）',
        'regression_equation_text': 'Regression Equation',
        'analysis_results_heading': '1. Analysis Results',
        'result_interpretation_heading': '2. Result Interpretation',
        'explanations_heading': '3. Explanations',
        'confusion_matrix_heading': '3. Confusion Matrix',
        'roc_curve_heading': '4. ROC Curve',
        'predicted_0': 'Predicted 0',
        'predicted_1': 'Predicted 1',
        'actual_0': 'Actual 0',
        'actual_1': 'Actual 1',
        'roc_curve_legend': 'ROC curve (area = {})',
        'template_file_not_found': 'Template file not found: {}',
        'open_file_failed': 'Failed to open file: {}'
    }
}


class BinaryLogitRegressionAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为中文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data34.xlsx")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=languages[self.current_language]["template_file_not_found"].format(excel_path))
        except Exception as e:
            self.result_label.config(text=languages[self.current_language]["open_file_failed"].format(str(e)))

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 检查数据中是否有缺失值
            if df.isnull().any().any():
                self.result_label.config(text=languages[self.current_language]["missing_values"])
                return

            # 检查自变量是否都是数值类型
            independent_vars = df.columns[:-1]
            non_numeric_vars = []
            for var in independent_vars:
                if not pd.api.types.is_numeric_dtype(df[var]):
                    non_numeric_vars.append(var)
            if non_numeric_vars:
                self.result_label.config(
                    text=languages[self.current_language]["non_numeric_vars"].format(", ".join(non_numeric_vars)))
                return

            # 提取自变量和因变量
            X = df.iloc[:, :-1]
            y = df.iloc[:, -1]
            feature_names = X.columns.tolist()
            dependent_var_name = y.name  # 获取因变量名称

            # 检查因变量是否为二元数据（只能是0或1）
            if not set(y.unique()).issubset({0, 1}):
                self.result_label.config(text=languages[self.current_language]["dependent_var_not_binary"])
                return

            # 划分训练集和测试集（80%训练，20%测试）
            X_train, X_test, y_train, y_test = train_test_split(
                X, y, test_size=0.2, random_state=42
            )

            # 使用statsmodels进行Logit回归（带截距），统一模型逻辑
            X_train_with_const = sm.add_constant(X_train)
            logit_model = sm.Logit(y_train, X_train_with_const).fit(
                disp=0,
                method='bfgs',  # 更稳定的优化算法
                maxiter=1000  # 只保留支持的参数
            )

            # 在测试集上进行预测
            X_test_with_const = sm.add_constant(X_test)
            y_pred_proba = logit_model.predict(X_test_with_const)
            y_pred = (y_pred_proba >= 0.5).astype(int)  # 以0.5为阈值

            # 计算评估指标
            accuracy = accuracy_score(y_test, y_pred)
            roc_auc = roc_auc_score(y_test, y_pred_proba)
            precision = precision_score(y_test, y_pred)
            recall = recall_score(y_test, y_pred)
            f1 = f1_score(y_test, y_pred)
            cm = confusion_matrix(y_test, y_pred)

            # 获取模型参数
            coefficients = logit_model.params.iloc[1:]  # 排除截距（使用iloc按位置访问）
            intercept = logit_model.params.iloc[0]  # 截距（使用iloc按位置访问）
            z_values = logit_model.tvalues
            p_values = logit_model.pvalues

            # 生成回归方程
            # 逻辑回归方程形式：logit(P) = ln(P/(1-P)) = 截距 + 系数1*变量1 + 系数2*变量2 + ...
            equation_parts = [f"{intercept:.4f}"]
            for name, coef in zip(feature_names, coefficients):
                if coef >= 0:
                    equation_parts.append(f"+ {coef:.4f}*{name}")
                else:
                    equation_parts.append(f"- {abs(coef):.4f}*{name}")

            logit_equation = f"logit(P({dependent_var_name}=1)) = " + "".join(equation_parts)

            # 准备表格数据
            headers = [
                          languages[self.current_language]['table_headers']['model']
                      ] + feature_names + [
                          languages[self.current_language]['table_headers']['intercept'],
                          languages[self.current_language]['table_headers']['accuracy'],
                          languages[self.current_language]['table_headers']['precision'],
                          languages[self.current_language]['table_headers']['recall'],
                          languages[self.current_language]['table_headers']['f1_score'],
                          languages[self.current_language]['table_headers']['roc_auc'],
                          languages[self.current_language]['table_headers']['regression_equation']
                      ] + [
                          f"{languages[self.current_language]['table_headers']['z_value_prefix']}{name}{languages[self.current_language]['table_headers']['z_value_suffix']}"
                          for name in [languages[self.current_language]['table_headers']['intercept']] + feature_names
                      ] + [
                          f"{languages[self.current_language]['table_headers']['p_value_prefix']}{name}{languages[self.current_language]['table_headers']['p_value_suffix']}"
                          for name in [languages[self.current_language]['table_headers']['intercept']] + feature_names
                      ]

            data_row = [
                           "二元Logit回归" if self.current_language == 'zh' else "Binary Logit Regression"
                       ] + coefficients.tolist() + [
                           intercept, accuracy, precision, recall, f1, roc_auc, logit_equation
                       ] + z_values.tolist() + p_values.tolist()

            df_results = pd.DataFrame([data_row], columns=headers)

            # 添加解释说明
            explanations = languages[self.current_language]['explanation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(
                columns=["coefficients", "intercept", "accuracy", "roc_auc", "z_value", "p_value",
                         "regression_equation"])
            # 这里使用语言字典中的翻译替代硬编码的"解释说明"和"Explanation"
            explanation_df.insert(0, "Model", languages[self.current_language]['explanation_row'])

            # 直接使用结果数据框，不添加解释说明
            if 'Model' in df_results.columns:
                transposed_df = df_results.set_index('Model').T.reset_index().rename(columns={'index': 'Model'})
            else:
                # 使用第一列作为索引（假设第一列是模型名称）
                transposed_df = df_results.set_index(df_results.columns[0]).T.reset_index().rename(
                    columns={'index': df_results.columns[0]})

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 获取保存目录
                save_dir = os.path.dirname(save_path)

                # 创建 Word 文档
                doc = Document()

                # 添加标题
                doc.add_heading(languages[self.current_language]['main_heading'], level=1)
                doc.add_paragraph(languages[self.current_language]['test_size_info'])

                # 添加分析结果表格
                doc.add_heading(languages[self.current_language]['analysis_results_heading'], level=2)
                table = doc.add_table(rows=transposed_df.shape[0] + 1, cols=transposed_df.shape[1])
                hdr_cells = table.rows[0].cells
                for col_idx, col_name in enumerate(transposed_df.columns):
                    hdr_cells[col_idx].text = col_name

                for row_idx in range(transposed_df.shape[0]):
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx, value in enumerate(transposed_df.iloc[row_idx]):
                        if isinstance(value, float):
                            row_cells[col_idx].text = f"{value:.4f}"
                        else:
                            row_cells[col_idx].text = str(value)

                # 添加结果解读（回归方程单独列出）
                doc.add_heading(languages[self.current_language]['result_interpretation_heading'], level=2)
                interpretation = doc.add_paragraph()
                interpretation.add_run(languages[self.current_language]['regression_equation_text'] + ": ").bold = True
                interpretation.add_run(logit_equation)

                # 添加解释说明（项目符号列表）
                doc.add_heading(languages[self.current_language]['explanations_heading'], level=2)
                explanations = languages[self.current_language]['explanation']
                # 遍历解释项，使用翻译后的键名显示
                for key, value in explanations.items():
                    expl_list = doc.add_paragraph(style='List Bullet')
                    # 使用翻译后的键名替代原英文键名
                    translated_key = languages[self.current_language].get(f'explanation_{key}', key.capitalize())
                    expl_list.add_run(f"{translated_key}: {value}")

                # 添加混淆矩阵
                doc.add_heading(languages[self.current_language]['confusion_matrix_heading'], level=2)
                cm_table = doc.add_table(rows=cm.shape[0] + 1, cols=cm.shape[1] + 1)
                # 添加混淆矩阵表头
                cm_table.cell(0, 0).text = ""
                cm_table.cell(0, 1).text = languages[self.current_language]['predicted_0']
                cm_table.cell(0, 2).text = languages[self.current_language]['predicted_1']
                cm_table.cell(1, 0).text = languages[self.current_language]['actual_0']
                cm_table.cell(2, 0).text = languages[self.current_language]['actual_1']
                # 填充混淆矩阵数据
                for i in range(cm.shape[0]):
                    for j in range(cm.shape[1]):
                        cm_table.cell(i + 1, j + 1).text = str(cm[i, j])

                # 生成ROC曲线
                fpr, tpr, thresholds = roc_curve(y_test, y_pred_proba)
                plt.figure(figsize=(10, 6))
                plt.plot(fpr, tpr,
                         label=languages[self.current_language]['roc_curve_legend'].format(f"{roc_auc:.2f}"))
                plt.plot([0, 1], [0, 1], 'k--')
                plt.xlim([0.0, 1.0])
                plt.ylim([0.0, 1.05])
                plt.xlabel(languages[self.current_language]['roc_x_label'])
                plt.ylabel(languages[self.current_language]['roc_y_label'])
                plt.title(languages[self.current_language]['roc_curve_title'])
                plt.legend(loc="lower right")
                img_name = "logit_regression_roc.png"
                img_path = os.path.join(save_dir, img_name)
                plt.savefig(img_path)
                plt.close()

                # 在 Word 文档中插入图片
                doc.add_heading(languages[self.current_language]['roc_curve_heading'], level=2)
                doc.add_picture(img_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                result_msg = languages[self.current_language]['analysis_complete'].format(save_path)
                result_msg += "\n" + languages[self.current_language]['images_saved'].format(save_dir)
                self.result_label.config(text=result_msg, wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]['no_save_path_selected'])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])
        self.open_excel_label.config(text=languages[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=languages[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame,
                                               text=languages[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = BinaryLogitRegressionAnalysisApp()
    app.run()

if __name__ == "__main__":
    run_app()
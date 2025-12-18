import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
from scipy.stats import pearsonr
import matplotlib.pyplot as plt
import pathlib
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'  # 设置字体为黑体，可根据系统情况修改为其他支持中文的字体
plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

# 定义语言字典
LANGUAGES = {
    'zh': {
        'title': "信度",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "请选择文件。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "中/英",
        'open_excel_button_text': "示例数据",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'explanation': {
            "Cronbach's Alpha系数": "Cronbach's Alpha系数用于衡量量表的内部一致性信度，取值范围在0 - 1之间，越接近1表示信度越高。",
            "样本量": "每个样本中的观测值数量。",
            "均值": "样本数据的平均值。"
        },
        'interpretation': {
            "Cronbach's Alpha系数": "Cronbach's Alpha系数越接近1，说明量表的内部一致性越好，信度越高。",
            "样本量": "样本量的大小会影响统计检验的稳定性，较大的样本量通常能提供更可靠的结果。",
            "均值": "均值反映了数据的平均水平，可用于比较不同变量的集中趋势。"
        }
    },
    'en': {
        'title': "Reliability",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Chinese/English",
        'open_excel_button_text': "Example data",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'explanation': {
            "Cronbach's Alpha Coefficient": "Cronbach's Alpha coefficient is used to measure the internal consistency reliability of a scale. The value ranges from 0 to 1, and the closer it is to 1, the higher the reliability.",
            "Sample Size": "The number of observations in each sample.",
            "Mean": "The average value of the sample data."
        },
        "interpretation": {
            "Cronbach's Alpha Coefficient": "The closer the Cronbach's Alpha coefficient is to 1, the better the internal consistency of the scale and the higher the reliability.",
            "Sample Size": "The sample size affects the stability of the statistical test. A larger sample size usually provides more reliable results.",
            "Mean": "The mean reflects the average level of the data and can be used to compare the central tendencies of different variables."
        }
    }
}


class ReliabilityAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def open_template_excel(self):
        try:
            # 获取当前脚本所在目录的上级目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)  # 上级目录

            excel_path = os.path.join(parent_dir, "Sample_data", "Data11.xls")

            if os.path.exists(excel_path):
                # 打开Excel文件（使用系统默认程序）
                os.startfile(excel_path)  # Windows系统
            else:
                self.result_label.config(text=f"模板文件不存在：{excel_path}")
        except Exception as e:
            self.result_label.config(text=f"打开文件失败：{str(e)}")

    def select_file(self):
        # 设置主窗口为对话框的父窗口，并保持在顶层
        self.root.transient()  # 将窗口设置为临时窗口，保持在应用程序顶层
        file_path = filedialog.askopenfilename(
            parent=self.root,  # 指定父窗口
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')
        # 确保主窗口在对话框关闭后仍保持顶层
        self.root.lift()

    def on_entry_click(self, event):
        if self.file_entry.get() == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def cronbach_alpha(self, data):
        """
        - 不修改原始 data（先拷贝）
        - 自动检测并反转高度负相关项（阈值 -0.3，可调整）
        - 防止 total_variance 为 0 导致除零
        """
        if data.shape[1] < 2:
            return float('nan')

        df = data.copy()  # 不要修改原始数据

        # 相关矩阵
        corr_matrix = df.corr()

        # 检测高度负相关的变量（任意一列与其他列存在负相关）
        negative_items = (corr_matrix < -0.3).any()
        neg_cols = list(negative_items[negative_items].index)  # 列名列表

        # 反向编码这些列（若确实存在反向题）
        for col in neg_cols:
            df[col] = -df[col]

        item_variances = df.var(axis=0, ddof=1)
        total_variance = df.sum(axis=1).var(ddof=1)
        num_items = df.shape[1]

        # 防护：若总方差为0，则无法计算alpha
        if total_variance == 0:
            return float('nan')

        alpha = (num_items / (num_items - 1)) * (1 - (item_variances.sum() / total_variance))
        return round(float(alpha), 4)

    def reliability_analysis(self, data):
        # 计算Cronbach's Alpha系数
        alpha = self.cronbach_alpha(data)
        return alpha

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=LANGUAGES[self.current_language]["file_not_found"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]["file_not_found"])
            return
        try:
            # 读取数据
            df = pd.read_excel(file_path)

            # 仅保留数值型列
            numerical_df = df.select_dtypes(include=[np.number])
            if numerical_df.empty:
                raise ValueError("数据中没有数值列，无法进行信度分析。")

            # 删除缺失值和常数列
            numerical_df = numerical_df.dropna()
            numerical_df = numerical_df.loc[:, numerical_df.std() > 0]
            if numerical_df.shape[1] < 2:
                raise ValueError("变量数量不足，无法计算信度。")

            # 标准化（Z-score 标准化，消除量纲影响）
            numerical_df = (numerical_df - numerical_df.mean()) / numerical_df.std()

            # 样本描述统计
            # 注意： numerical_df 已经被标准化为 Z-score
            desc_stats = numerical_df.describe().T[['mean', 'std', 'min', 'max']]
            desc_stats.rename(columns={'mean': '均值', 'std': '标准差', 'min': '最小值', 'max': '最大值'}, inplace=True)

            # 保留均值用于后续绘图
            means = desc_stats['均值']

            # 计算每个变量的样本量（非NA计数）
            sample_sizes = numerical_df.count()

            # 进行信度分析
            alpha = self.reliability_analysis(numerical_df)

            # α 系数等级解释
            if self.current_language == "zh":
                # 中文等级描述
                if np.isnan(alpha):
                    reliability_level = "不适用 (N/A)"
                else:
                    if alpha >= 0.9:
                        reliability_level = "极高信度 (Excellent)"
                    elif alpha >= 0.8:
                        reliability_level = "良好信度 (Good)"
                    elif alpha >= 0.7:
                        reliability_level = "可接受信度 (Acceptable)"
                    elif alpha >= 0.6:
                        reliability_level = "较低信度 (Questionable)"
                    else:
                        reliability_level = "低信度 (Poor)"
            else:
                if np.isnan(alpha):
                    reliability_level = "N/A"
                else:
                    if alpha >= 0.9:
                        reliability_level = "Excellent"
                    elif alpha >= 0.8:
                        reliability_level = "Good"
                    elif alpha >= 0.7:
                        reliability_level = "Acceptable"
                    elif alpha >= 0.6:
                        reliability_level = "Questionable"
                    else:
                        reliability_level = "Poor"

            # 样本描述统计（科研标准格式）
            desc_stats = numerical_df.describe().T[['mean', 'std', 'min', 'max']]
            if self.current_language == "zh":
                # 中文列名
                columns_map = {'mean': '均值', 'std': '标准差', 'min': '最小值', 'max': '最大值'}
            else:
                # 英文列名（保持原统计学术语）
                columns_map = {'mean': 'Mean', 'std': 'Std', 'min': 'Min', 'max': 'Max'}

            # 应用列名映射
            desc_stats.rename(columns=columns_map, inplace=True)
            sample_sizes = numerical_df.count()

            # 根据当前语言定义文本
            if self.current_language == "zh":
                # 中文文本
                mean_col = "均值"
                stats_name1 = "Cronbach's Alpha系数"
                stats_name2 = "样本量"
                stats_name3 = "均值"
                headers = ["统计量", "统计量值", "p值"]
            else:
                # 英文文本
                mean_col = "Mean"  # 对应英文列名
                stats_name1 = "Cronbach's Alpha Coefficient"
                stats_name2 = "Sample Size"
                stats_name3 = "Mean"
                headers = ["Statistic", "Value", "p-value"]

            # 保留均值用于后续绘图（使用语言对应的列名）
            means = desc_stats[mean_col]

            # 整理数据（使用语言对应的统计量名称）
            data = [
                [stats_name1, alpha, ""],
                [stats_name2, sample_sizes.to_dict(), ""],
                [stats_name3, means.to_dict(), ""]
            ]

            # 添加解释说明
            explanations = LANGUAGES[self.current_language]['explanation']
            interpretations = LANGUAGES[self.current_language]['interpretation']

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加标题
                doc.add_heading('信度分析结果' if self.current_language == 'zh' else 'Reliability Analysis Results', 0)

                # 添加表格
                table = doc.add_table(rows=1, cols=len(headers))
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header

                for row in data:
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)
                if self.current_language == "zh":
                    doc.add_paragraph(f"信度等级: {reliability_level}")
                else:
                    doc.add_paragraph(f"Reliability Level: {reliability_level}")

                # 添加描述统计表到 Word 文档
                doc.add_heading('变量描述统计' if self.current_language == 'zh' else 'Descriptive Statistics', 1)

                table2 = doc.add_table(rows=1, cols=len(desc_stats.columns) + 1)
                hdr_cells = table2.rows[0].cells
                hdr_cells[0].text = '变量' if self.current_language == 'zh' else 'Variable'
                for i, col in enumerate(desc_stats.columns):
                    hdr_cells[i + 1].text = col

                for var, row in desc_stats.iterrows():
                    row_cells = table2.add_row().cells
                    row_cells[0].text = str(var)
                    for i, val in enumerate(row):
                        row_cells[i + 1].text = f"{val:.3f}"

                # 添加解释说明
                doc.add_heading('解释说明' if self.current_language == 'zh' else 'Explanation', 1)
                for key, value in explanations.items():
                    doc.add_paragraph(f"{key}: {value}")

                # 添加分析结果解读
                doc.add_heading('结果解读' if self.current_language == 'zh' else 'Interpretation', 1)
                for key, value in interpretations.items():
                    doc.add_paragraph(f"{key}: {value}")

                # 相关性热图
                fig, ax = plt.subplots(figsize=(10, 8))
                corr = numerical_df.corr()
                im = ax.imshow(corr, cmap='coolwarm')

                # 添加颜色条
                cbar = ax.figure.colorbar(im, ax=ax)
                if self.current_language == "zh":
                    cbar.ax.set_ylabel('相关系数', rotation=-90, va="bottom")
                else:
                    cbar.ax.set_ylabel('Correlation Coefficient', rotation=-90, va="bottom")

                # 标记坐标轴
                ax.set_xticks(np.arange(corr.shape[1]))
                ax.set_yticks(np.arange(corr.shape[0]))
                ax.set_xticklabels(corr.columns, rotation=45, ha="right")
                ax.set_yticklabels(corr.columns)

                # 添加相关系数文本
                for i in range(corr.shape[0]):
                    for j in range(corr.shape[1]):
                        text = ax.text(j, i, f"{corr.iloc[i, j]:.2f}",
                                       ha="center", va="center", color="w")

                ax.set_title('变量相关性热图' if self.current_language == 'zh' else 'Variable Correlation Heatmap')
                plt.tight_layout()

                # 保存图片
                img_path = os.path.splitext(save_path)[0] + '.png'
                plt.savefig(img_path)
                plt.close()

                # 将图片插入 Word 文档
                doc.add_heading('变量相关性热图' if self.current_language == 'zh' else 'Variable Correlation Heatmap',
                                1)
                doc.add_picture(img_path, width=Inches(6))

                # 添加数值标签
                for i, v in enumerate(means):
                    ax.text(i, v + 0.05, f"{v:.2f}", ha='center', fontsize=9)

                # 保存 Word 文档
                doc.save(save_path)

                # 导出 Excel 文件，便于复现
                excel_output = os.path.splitext(save_path)[0] + '_analysis.xlsx'
                with pd.ExcelWriter(excel_output) as writer:
                    numerical_df.to_excel(writer, sheet_name='RawData')
                    desc_stats.to_excel(writer, sheet_name='Descriptive')
                    pd.DataFrame({
                        "CronbachAlpha": [alpha],
                        "ReliabilityLevel": [reliability_level]
                    }).to_excel(writer, sheet_name='Reliability')

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(save_path)
                result_msg += f"\n科研复现数据已保存到 {excel_output}"
                self.result_label.config(text=result_msg, wraplength=400)

            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(LANGUAGES[self.current_language]["title"])
        self.select_button.config(text=LANGUAGES[self.current_language]["select_button"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=LANGUAGES[self.current_language]["analyze_button"])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]["switch_language"])
        self.open_excel_label.config(text=LANGUAGES[self.current_language]["open_excel_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 根据屏幕分辨率动态计算窗口尺寸（占屏幕的60%）
        window_width = int(screen_width * 0.4)
        window_height = int(screen_height * 0.4)

        # 限制最小窗口尺寸，避免过小
        min_width = 500
        min_height = 300
        window_width = max(window_width, min_width)
        window_height = max(window_height, min_height)

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["select_button"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["analyze_button"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建打开Excel文件标签
        self.open_excel_label = ttk.Label(
            frame,
            text=LANGUAGES[self.current_language]["open_excel_button_text"],
            foreground="gray",
            cursor="hand2"
        )
        self.open_excel_label.bind("<Button-1>", lambda e: self.open_template_excel())
        self.open_excel_label.pack(padx=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=LANGUAGES[self.current_language]["switch_language"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = ReliabilityAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()